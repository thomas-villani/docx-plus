"""Insert a section break at an arbitrary paragraph.

python-docx's :meth:`Document.add_section` only appends a new section at
the end of the document. Inserting a section break mid-document requires
moving the existing trailing ``<w:sectPr>``'s properties into the chosen
paragraph and writing a fresh ``<w:type>`` value to mark the break style.
:func:`insert_section_break` does that in one call.

This module imports only from ``docx_plus.core`` and python-docx
internals via the section proxy (SPEC §9.1).
"""

from __future__ import annotations

from copy import deepcopy
from typing import TYPE_CHECKING, Literal, cast

from docx.section import Section

from docx_plus.core.ns import qn
from docx_plus.core.oxml import el, insert_before_first_anchor, remove

if TYPE_CHECKING:
    from docx.oxml.section import CT_SectPr
    from docx.parts.document import DocumentPart
    from docx.text.paragraph import Paragraph

SectionStartType = Literal[
    "nextPage", "continuous", "evenPage", "oddPage", "nextColumn"
]


# Schema siblings later than `w:type` per ECMA-376 17.6.17 CT_SectPr. ``type``
# follows the header/footer references and footnote/endnote properties, so it
# must NOT be jammed to position 0 when those precede it.
_TYPE_LATER_SIBLINGS: tuple[str, ...] = (
    "w:pgSz",
    "w:pgMar",
    "w:paperSrc",
    "w:pgBorders",
    "w:lnNumType",
    "w:pgNumType",
    "w:cols",
    "w:formProt",
    "w:vAlign",
    "w:noEndnote",
    "w:titlePg",
    "w:textDirection",
    "w:bidi",
    "w:rtlGutter",
    "w:docGrid",
    "w:printerSettings",
    "w:sectPrChange",
)


def insert_section_break(
    paragraph: Paragraph,
    *,
    start_type: SectionStartType = "nextPage",
) -> Section:
    """Insert a section break at ``paragraph``.

    The break splits the document so that the chosen paragraph becomes
    the last paragraph of a new section. Section properties (page size,
    margins, header/footer references) are copied from the document's
    trailing ``<w:sectPr>``, then ``<w:type>`` is set on the inserted
    copy so Word knows how the *following* section starts.

    If ``paragraph`` already carries a section break, this call replaces
    it (idempotent re-application is supported, e.g. to switch the
    break type from ``nextPage`` to ``continuous``).

    Args:
        paragraph: A python-docx :class:`~docx.text.paragraph.Paragraph`
            in the main document body.
        start_type: How the *following* section begins. ``"nextPage"``
            (default) puts the new section on a fresh page;
            ``"continuous"`` lets it flow on the same page;
            ``"evenPage"`` / ``"oddPage"`` page-break to the next
            even/odd-numbered page; ``"nextColumn"`` starts a new column.

    Returns:
        A :class:`Section` proxy wrapping the new section's
        ``<w:sectPr>``. Use it to mutate the new section's properties
        (margins, columns, headers).

    Raises:
        ValueError: If ``paragraph`` is not parented to ``<w:body>``
            (e.g. a header/footer paragraph), or if the document has no
            trailing ``<w:sectPr>`` to copy section properties from
            (reachable on a hand-built document that never had one).

    Example:
        >>> from docx import Document
        >>> from docx_plus.layout import insert_section_break
        >>> doc = Document()
        >>> doc.add_paragraph("intro")
        >>> p = doc.add_paragraph("split here")
        >>> doc.add_paragraph("after the break")
        >>> new_section = insert_section_break(p, start_type="continuous")
    """
    p_element = paragraph._p
    body = p_element.getparent()
    if body is None or body.tag != qn("w:body"):
        raise ValueError(
            "insert_section_break requires a paragraph in the main document body"
        )

    sentinel = body.find(qn("w:sectPr"))
    if sentinel is None:
        raise ValueError(
            "document has no trailing sectPr to copy properties from"
        )

    new_sect_pr = cast("CT_SectPr", deepcopy(sentinel))
    new_sect_pr.attrib.clear()
    _set_start_type(new_sect_pr, start_type)
    p_element.set_sectPr(new_sect_pr)

    document_part = cast("DocumentPart", paragraph.part)
    return Section(new_sect_pr, document_part)


def _set_start_type(sect_pr: CT_SectPr, start_type: SectionStartType) -> None:
    """Set ``<w:type w:val="...">`` on ``sect_pr``, replacing any existing one.

    Lands ``w:type`` in its ECMA-376 17.6.17 slot: after any
    ``headerReference`` / ``footerReference`` / ``footnotePr`` /
    ``endnotePr`` and before ``pgSz`` onward. The old "insert at position 0"
    shortcut jammed ``w:type`` ahead of header/footer references on sections
    that carry custom headers, producing schema-invalid output.
    """
    existing = sect_pr.find(qn("w:type"))
    if existing is not None:
        remove(existing)
    type_el = el("w:type", **{"w:val": start_type})
    insert_before_first_anchor(sect_pr, type_el, _TYPE_LATER_SIBLINGS)


__all__ = ["SectionStartType", "insert_section_break"]
