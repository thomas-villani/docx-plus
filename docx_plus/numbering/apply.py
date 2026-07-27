"""Attaching list definitions to paragraphs — ``<w:numPr>``.

A paragraph joins a list by carrying ``<w:numPr>`` in its ``<w:pPr>``,
naming a ``w:numId`` (the *instance*, never the abstract definition) and
a ``w:ilvl`` (the outline depth). Paragraphs sharing a ``numId`` continue
one sequence, in document order.

python-docx models ``w:numPr`` — ``CT_NumPr`` exists — but its
convenience accessors for the two children are commented out in the
source, so callers still reach for raw XML.

This module imports only from ``docx_plus.core`` and its siblings in
``docx_plus.numbering`` (SPEC §9.1).
"""

from __future__ import annotations

from typing import TYPE_CHECKING, cast

from docx_plus.core.errors import DocxPlusError
from docx_plus.core.ns import qn
from docx_plus.core.oxml import (
    body_document_for,
    el,
    insert_before_first_anchor,
    ordered_insert,
    remove,
    sub,
)
from docx_plus.core.parts import NUMBERING_SPEC, get_or_create_part
from docx_plus.numbering.define import _abstract_id_for
from docx_plus.numbering.registry import NumIdRegistry

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph
    from lxml import etree

#: ECMA-376 17.9.18 ``CT_NumPr`` child order.
_NUM_PR_CHILD_ORDER: tuple[str, ...] = ("ilvl", "numId", "numberingChange", "ins")

#: Children of ``w:pPr`` that follow ``w:numPr`` (ECMA-376 17.3.1.26
#: ``CT_PPr``). python-docx generates a ``get_or_add_numPr`` at runtime
#: but does not declare it, so it is invisible to a strict type-check —
#: hence placing the element here rather than borrowing that accessor.
_PPR_AFTER_NUM_PR: tuple[str, ...] = (
    "w:suppressLineNumbers",
    "w:pBdr",
    "w:shd",
    "w:tabs",
    "w:suppressAutoHyphens",
    "w:kinsoku",
    "w:wordWrap",
    "w:overflowPunct",
    "w:topLinePunct",
    "w:autoSpaceDE",
    "w:autoSpaceDN",
    "w:bidi",
    "w:adjustRightInd",
    "w:snapToGrid",
    "w:spacing",
    "w:ind",
    "w:contextualSpacing",
    "w:mirrorIndents",
    "w:suppressOverlap",
    "w:jc",
    "w:textDirection",
    "w:textAlignment",
    "w:textboxTightWrap",
    "w:outlineLvl",
    "w:divId",
    "w:cnfStyle",
    "w:rPr",
    "w:sectPr",
    "w:pPrChange",
)

#: Children of ``w:numbering`` that must follow ``w:num``.
_AFTER_NUM: tuple[str, ...] = ("w:numIdMacAtCleanup",)

#: Inside a ``w:numPr``, ``w:numId`` ``0`` is not a definition reference —
#: it is the sentinel meaning "no numbering", the only way a paragraph
#: can opt out of a list applied by its *style*.
_NO_NUMBERING = "0"


class ListDefinitionNotFoundError(DocxPlusError, KeyError):
    """Raised when a ``numId`` has no ``<w:num>`` in ``numbering.xml``.

    Subclasses ``KeyError`` so existing lookup-style handling still
    catches it; also subclasses :class:`DocxPlusError` per SPEC §9.7.
    """


def apply_list(paragraph: Paragraph, num_id: int, *, level: int = 0) -> None:
    """Put ``paragraph`` into the list identified by ``num_id``.

    Idempotent — re-applying replaces the paragraph's existing
    ``w:numPr`` rather than stacking a second one.

    The ``num_id`` is not validated against ``numbering.xml``. A
    paragraph may legitimately reference a definition another tool will
    supply, and Word itself tolerates a dangling reference by rendering
    the paragraph unnumbered. Use :func:`restart_list` or
    :func:`~docx_plus.numbering.read_list_definitions` if you need the
    reference checked.

    Args:
        paragraph: The python-docx :class:`~docx.text.paragraph.Paragraph`
            to mutate.
        num_id: A ``w:numId`` from
            :func:`~docx_plus.numbering.define_list_definition` or one of
            the presets.
        level: Zero-based outline depth. Level 0 is the outermost.

    Raises:
        ValueError: If ``level`` is negative or beyond the ninth level.

    Example:
        >>> from docx import Document
        >>> from docx_plus.numbering import apply_list, define_numbered_list
        >>> doc = Document()
        >>> num = define_numbered_list(doc, levels=2)
        >>> apply_list(doc.add_paragraph("first"), num)
        >>> apply_list(doc.add_paragraph("nested"), num, level=1)
    """
    _check_level(level)
    num_pr = _get_or_add_num_pr(paragraph)
    _set_num_pr(num_pr, num_id=str(num_id), level=level)


def remove_list(paragraph: Paragraph, *, suppress_style_numbering: bool = False) -> None:
    """Take ``paragraph`` out of any list it is directly a member of.

    Idempotent — a paragraph with no ``w:numPr`` is left alone.

    Args:
        paragraph: The python-docx :class:`~docx.text.paragraph.Paragraph`
            to mutate.
        suppress_style_numbering: Whether to also suppress numbering the
            paragraph's *style* applies.

            Removing a direct ``w:numPr`` reverts the paragraph to
            whatever its style says, and for a style like ``ListBullet``
            that means it stays bulleted — usually a surprise. Passing
            ``True`` writes ``<w:numPr><w:numId w:val="0"/></w:numPr>``
            instead of removing the element, which is the sentinel Word
            uses for "definitely not numbered".

    Example:
        >>> from docx import Document
        >>> from docx_plus.numbering import remove_list
        >>> doc = Document()
        >>> p = doc.add_paragraph("plain", style="List Bullet")
        >>> remove_list(p, suppress_style_numbering=True)
    """
    if suppress_style_numbering:
        num_pr = _get_or_add_num_pr(paragraph)
        _set_num_pr(num_pr, num_id=_NO_NUMBERING, level=None)
        return

    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is None:
        return
    existing = ppr.find(qn("w:numPr"))
    if existing is not None:
        remove(existing)


def restart_list(
    paragraph: Paragraph,
    num_id: int,
    *,
    level: int = 0,
    start: int = 1,
    num_registry: NumIdRegistry | None = None,
) -> int:
    """Begin a fresh sequence over the same definition, at ``paragraph``.

    Restarting is not a paragraph property in OOXML — there is nowhere to
    say "count from 1 again here". What Word does, and what this does, is
    add a *second* ``<w:num>`` pointing at the same ``<w:abstractNum>``
    and carrying a ``<w:startOverride>``. Two instances of one definition
    are independent counters that look identical.

    ``paragraph`` and every later paragraph you want in the new sequence
    must use the returned id; this call only moves ``paragraph`` itself.

    Args:
        paragraph: The paragraph that begins the new sequence.
        num_id: The existing list to branch from.
        level: Zero-based outline depth to restart.
        start: Value to restart the counter at.
        num_registry: Pre-existing ``w:numId`` allocator to share across
            an editing session.

    Returns:
        The new ``w:numId``. Apply it to the rest of the run of
        paragraphs that should share the restarted sequence.

    Raises:
        ListDefinitionNotFoundError: If ``num_id`` has no ``<w:num>``, or
            that entry has no ``<w:abstractNumId>`` to branch from.
            Subclasses :class:`KeyError`.
        ValueError: If ``level`` is negative or beyond the ninth level,
            or ``start`` is negative.

    Example:
        >>> from docx import Document
        >>> from docx_plus.numbering import apply_list, define_numbered_list, restart_list
        >>> doc = Document()
        >>> num = define_numbered_list(doc)
        >>> apply_list(doc.add_paragraph("one"), num)
        >>> apply_list(doc.add_paragraph("two"), num)
        >>> second = restart_list(doc.add_paragraph("one again"), num)
        >>> apply_list(doc.add_paragraph("two again"), second)
    """
    _check_level(level)
    if start < 0:
        raise ValueError(f"start must be non-negative; got {start!r}")

    doc = body_document_for(paragraph, operation="restart_list")
    _, root = get_or_create_part(doc, NUMBERING_SPEC)

    abstract_id = _abstract_id_for(root, num_id)
    if abstract_id is None:
        raise ListDefinitionNotFoundError(
            f"no <w:num w:numId='{num_id}'> with an <w:abstractNumId> in numbering.xml; "
            f"restart_list needs an existing definition to branch from"
        )

    if num_registry is None:
        num_registry = NumIdRegistry(doc)
    new_num_id = num_registry.next_sequential()

    num = el("w:num", **{"w:numId": str(new_num_id)})
    sub(num, "w:abstractNumId", **{"w:val": abstract_id})
    override = sub(num, "w:lvlOverride", **{"w:ilvl": str(level)})
    sub(override, "w:startOverride", **{"w:val": str(start)})
    _insert_num(root, num)

    apply_list(paragraph, new_num_id, level=level)
    return new_num_id


# ---------------------------------------------------------------------------
# Internals.
# ---------------------------------------------------------------------------


def _check_level(level: int) -> None:
    from docx_plus.numbering.define import MAX_LEVELS

    if not 0 <= level < MAX_LEVELS:
        raise ValueError(
            f"level is zero-based and a definition holds at most {MAX_LEVELS} levels, "
            f"so it must be 0 to {MAX_LEVELS - 1}; got {level!r}"
        )


def _get_or_add_num_pr(paragraph: Paragraph) -> etree._Element:
    """Return the paragraph's ``w:numPr``, creating it in schema position.

    ``w:pPr`` itself comes from python-docx's ``CT_P.get_or_add_pPr``,
    which places it ahead of the runs. ``w:numPr`` is placed here: the
    generated ``get_or_add_numPr`` exists at runtime but is undeclared,
    so a strict type-check cannot see it.
    """
    ppr = cast("etree._Element", paragraph._p.get_or_add_pPr())
    existing = ppr.find(qn("w:numPr"))
    if existing is not None:
        return existing
    num_pr = el("w:numPr")
    insert_before_first_anchor(ppr, num_pr, _PPR_AFTER_NUM_PR)
    return num_pr


def _set_num_pr(num_pr: etree._Element, *, num_id: str, level: int | None) -> None:
    """Write ``w:ilvl`` / ``w:numId`` into an existing ``w:numPr``.

    ``level=None`` omits ``w:ilvl`` entirely, which is what the
    "no numbering" sentinel wants — a depth is meaningless without a
    list.
    """
    if level is None:
        existing = num_pr.find(qn("w:ilvl"))
        if existing is not None:
            remove(existing)
    else:
        ordered_insert(num_pr, el("w:ilvl", **{"w:val": str(level)}), _NUM_PR_CHILD_ORDER)
    ordered_insert(num_pr, el("w:numId", **{"w:val": num_id}), _NUM_PR_CHILD_ORDER)


def _insert_num(root: etree._Element, num: etree._Element) -> None:
    """Append a ``w:num``, keeping it ahead of ``w:numIdMacAtCleanup``."""
    insert_before_first_anchor(root, num, _AFTER_NUM)


__all__ = ["ListDefinitionNotFoundError", "apply_list", "remove_list", "restart_list"]
