"""Authoring list definitions in ``numbering.xml``.

python-docx cannot write a list definition at all. It ships a
``NumberingPart``, but ``docx/oxml/numbering.py`` defines classes only
for ``w:numbering``, ``w:num``, ``w:lvlOverride``, and ``w:numPr`` —
there is no ``CT_AbstractNum`` and no ``CT_Lvl``, so nothing in it can
express what a list *looks like*: the number format, the level text, the
start value, the indents, the bullet glyph. Callers hand-write XML.

The OOXML model has two halves:

- ``<w:abstractNum>`` — the definition. Up to nine ``<w:lvl>`` children,
  each describing one outline level.
- ``<w:num>`` — an instance pointing at an abstract definition by id.
  Paragraphs reference *this* id, never the abstract one.

The indirection is what makes restarting possible: a second ``w:num``
over the same ``w:abstractNum`` is an independent sequence with
identical formatting. See :func:`~docx_plus.numbering.restart_list`.

This module imports only from ``docx_plus.core`` and its sibling
``docx_plus.numbering.registry`` (SPEC §9.1).
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import TYPE_CHECKING, Literal

from docx_plus.core.errors import DocxPlusError
from docx_plus.core.ns import qn
from docx_plus.core.oxml import el, insert_before_first_anchor, ordered_insert, sub
from docx_plus.core.parts import NUMBERING_SPEC, get_or_create_part
from docx_plus.numbering.registry import AbstractNumIdRegistry, NumIdRegistry

if TYPE_CHECKING:
    from docx.document import Document
    from lxml import etree

#: Word refuses to render past nine levels (ECMA-376 17.9.1 caps
#: ``w:lvl`` at nine children of one ``w:abstractNum``).
MAX_LEVELS = 9

#: ECMA-376 17.18.59 ``ST_NumberFormat``. The full enumeration — a typo
#: here produces a list that silently renders wrong rather than failing,
#: so this is checked exactly rather than shape-matched.
_NUMBER_FORMATS = frozenset(
    {
        "decimal",
        "upperRoman",
        "lowerRoman",
        "upperLetter",
        "lowerLetter",
        "ordinal",
        "cardinalText",
        "ordinalText",
        "hex",
        "chicago",
        "ideographDigital",
        "japaneseCounting",
        "aiueo",
        "iroha",
        "decimalFullWidth",
        "decimalHalfWidth",
        "japaneseLegal",
        "japaneseDigitalTenThousand",
        "decimalEnclosedCircle",
        "decimalFullWidth2",
        "aiueoFullWidth",
        "irohaFullWidth",
        "decimalZero",
        "bullet",
        "ganada",
        "chosung",
        "decimalEnclosedFullstop",
        "decimalEnclosedParen",
        "decimalEnclosedCircleChinese",
        "ideographEnclosedCircle",
        "ideographTraditional",
        "ideographZodiac",
        "ideographZodiacTraditional",
        "taiwaneseCounting",
        "ideographLegalTraditional",
        "taiwaneseCountingThousand",
        "taiwaneseDigital",
        "chineseCounting",
        "chineseLegalSimplified",
        "chineseCountingThousand",
        "koreanDigital",
        "koreanCounting",
        "koreanLegal",
        "koreanDigital2",
        "vietnameseCounting",
        "russianLower",
        "russianUpper",
        "none",
        "numberInDash",
        "hebrew1",
        "hebrew2",
        "arabicAlpha",
        "arabicAbjad",
        "hindiVowels",
        "hindiConsonants",
        "hindiNumbers",
        "hindiCounting",
        "thaiLetters",
        "thaiNumbers",
        "thaiCounting",
        "bahtText",
        "dollarText",
        "custom",
    }
)

_SUFFIXES = frozenset({"tab", "space", "nothing"})
_JUSTIFICATIONS = frozenset({"left", "center", "right", "start", "end"})
_MULTI_LEVEL_TYPES = frozenset({"singleLevel", "multilevel", "hybridMultilevel"})

#: ``%1`` .. ``%9`` placeholders inside ``w:lvlText``.
_PLACEHOLDER_RE = re.compile(r"%([1-9])")

#: ECMA-376 17.9.6 ``CT_Lvl`` child order.
_LVL_CHILD_ORDER: tuple[str, ...] = (
    "start",
    "numFmt",
    "lvlRestart",
    "pStyle",
    "isLgl",
    "suff",
    "lvlText",
    "lvlPicBulletId",
    "legacy",
    "lvlJc",
    "pPr",
    "rPr",
)

#: ECMA-376 17.9.1 ``CT_AbstractNum`` child order.
_ABSTRACT_NUM_CHILD_ORDER: tuple[str, ...] = (
    "nsid",
    "multiLevelType",
    "tmpl",
    "name",
    "styleLink",
    "numStyleLink",
    "lvl",
)

#: Children of ``w:numbering`` that must follow ``w:abstractNum``
#: (ECMA-376 17.9.17: ``numPicBullet*, abstractNum*, num*,
#: numIdMacAtCleanup?``). Nothing in python-docx inserts an
#: ``abstractNum`` at all, so getting this right is on us.
_AFTER_ABSTRACT_NUM: tuple[str, ...] = ("w:num", "w:numIdMacAtCleanup")
_AFTER_NUM: tuple[str, ...] = ("w:numIdMacAtCleanup",)

Suffix = Literal["tab", "space", "nothing"]
Justification = Literal["left", "center", "right", "start", "end"]
MultiLevelType = Literal["singleLevel", "multilevel", "hybridMultilevel"]


class InvalidLevelError(DocxPlusError, ValueError):
    """Raised for a malformed :class:`LevelDefinition` or level list.

    Subclasses ``ValueError`` so existing ``except ValueError:`` clauses
    still catch it; also subclasses :class:`DocxPlusError` per SPEC §9.7.
    """


@dataclass(frozen=True)
class LevelDefinition:
    """One outline level of a list definition — a ``<w:lvl>``.

    Attributes:
        fmt: ECMA-376 17.18.59 ``ST_NumberFormat`` name. ``"decimal"``,
            ``"lowerLetter"``, ``"lowerRoman"``, ``"upperRoman"``, and
            ``"bullet"`` cover almost every list; the full enumeration
            has 60+ entries.
        text: The ``w:lvlText`` pattern. ``%N`` interpolates the counter
            for level ``N`` (**1-based**, so level 0's own counter is
            ``%1``). ``"%1."`` gives ``1.``, ``2.``; ``"%1.%2."`` on
            level 1 gives ``1.1``, ``1.2`` — the legal-outline shape. For
            ``fmt="bullet"`` this is the literal glyph, not a pattern.
        start: First value of the counter. Defaults to ``1``.
        indent: Left indent in **twips** (1/20 pt; 720 = 0.5"). ``None``
            omits the ``w:ind``, inheriting from the style.
        hanging: Hanging indent in twips — the width reserved for the
            number, measured back from ``indent``. Written only when
            ``indent`` is also set.

            **Make this wider than the rendered ``text``.** The gap
            between number and text is a tab stop sitting at ``indent``,
            so when the number is wider than ``hanging`` the tab has
            nowhere to advance to and collapses to nothing — a
            cumulative outline renders ``1.1.1.On-call lead`` rather than
            ``1.1.1. On-call lead``. Deeper levels of a ``%1.%2.%3.``
            outline therefore need progressively larger values, not the
            same 360 that suits a single digit.
        justify: How the number is aligned within the hanging indent.
        suffix: What separates the number from the text —
            ``"tab"`` (default, what Word writes), ``"space"``, or
            ``"nothing"``.
        restart_after: The ``w:lvlRestart`` value: this level's counter
            restarts whenever the level with this (1-based) number
            increments. ``0`` means *never* restart. ``None`` omits the
            element, which is Word's implicit "restart after the
            immediately preceding level".
        font: Font applied to the number or bullet glyph only, not the
            paragraph text. Required in practice for symbol bullets —
            ``"Symbol"`` and ``"Wingdings"`` render as Latin letters
            without it.

    Raises:
        InvalidLevelError: If any field is outside its ECMA-376 type.
    """

    fmt: str = "decimal"
    text: str = "%1."
    start: int = 1
    indent: int | None = None
    hanging: int | None = None
    justify: Justification = "left"
    suffix: Suffix = "tab"
    restart_after: int | None = None
    font: str | None = None

    def __post_init__(self) -> None:
        """Validate the fields against their ECMA-376 simple types."""
        if self.fmt not in _NUMBER_FORMATS:
            raise InvalidLevelError(
                f"LevelDefinition.fmt must be an ECMA-376 ST_NumberFormat name "
                f"(e.g. 'decimal', 'bullet', 'lowerRoman'); got {self.fmt!r}"
            )
        if self.suffix not in _SUFFIXES:
            raise InvalidLevelError(
                f"LevelDefinition.suffix must be one of {sorted(_SUFFIXES)}; got {self.suffix!r}"
            )
        if self.justify not in _JUSTIFICATIONS:
            raise InvalidLevelError(
                f"LevelDefinition.justify must be one of {sorted(_JUSTIFICATIONS)}; "
                f"got {self.justify!r}"
            )
        if self.start < 0:
            raise InvalidLevelError(
                f"LevelDefinition.start must be non-negative; got {self.start!r}"
            )
        if self.restart_after is not None and not 0 <= self.restart_after <= MAX_LEVELS:
            raise InvalidLevelError(
                f"LevelDefinition.restart_after must be 0 (never) or a 1-based level "
                f"number up to {MAX_LEVELS}; got {self.restart_after!r}"
            )
        if self.hanging is not None and self.indent is None:
            raise InvalidLevelError(
                "LevelDefinition.hanging has no effect without indent; set both or "
                "neither (a w:ind with only w:hanging is ignored by Word)"
            )


def define_list_definition(
    doc: Document,
    *,
    levels: list[LevelDefinition] | tuple[LevelDefinition, ...],
    name: str | None = None,
    style_link: str | None = None,
    num_style_link: str | None = None,
    multi_level_type: MultiLevelType | None = None,
    num_registry: NumIdRegistry | None = None,
    abstract_registry: AbstractNumIdRegistry | None = None,
) -> int:
    """Write a list definition and return the ``numId`` to apply.

    Creates one ``<w:abstractNum>`` holding ``levels``, plus one
    ``<w:num>`` instance pointing at it. The returned id is the
    instance's — that is what :func:`~docx_plus.numbering.apply_list`
    takes and what a paragraph's ``w:numPr`` stores.

    ``numbering.xml`` is created if the document has none. That is not
    the same as python-docx's ``doc.part.numbering_part``, which
    fabricates through an unimplemented stub and raises; see
    :data:`~docx_plus.core.parts.NUMBERING_SPEC`.

    Args:
        doc: The python-docx :class:`~docx.document.Document` to mutate.
        levels: One :class:`LevelDefinition` per outline level, outermost
            first. At least one, at most nine.
        name: Optional ``w:name`` for the definition. Cosmetic; Word does
            not surface it.
        style_link: Style id this definition is the numbering *for* —
            the paired half of a "list style". Mutually exclusive with
            ``num_style_link``.
        num_style_link: Style id whose numbering this definition
            *defers to*. Mutually exclusive with ``style_link``.
        multi_level_type: ``"singleLevel"``, ``"multilevel"``, or
            ``"hybridMultilevel"``. Defaults to ``"singleLevel"`` for one
            level and ``"multilevel"`` beyond that, which is what Word
            writes.
        num_registry: Pre-existing ``w:numId`` allocator to share across
            an editing session.
        abstract_registry: Pre-existing ``w:abstractNumId`` allocator to
            share across an editing session.

    Returns:
        The ``w:numId`` of the new instance.

    Raises:
        InvalidLevelError: If ``levels`` is empty, longer than nine, or
            contains a ``w:lvlText`` placeholder referencing a deeper
            level than its own.
        ValueError: If both ``style_link`` and ``num_style_link`` are given.

    Example:
        >>> from docx import Document
        >>> from docx_plus.numbering import LevelDefinition, apply_list, define_list_definition
        >>> doc = Document()
        >>> num = define_list_definition(doc, levels=[
        ...     LevelDefinition(fmt="decimal", text="%1.", indent=720, hanging=360),
        ...     LevelDefinition(fmt="lowerLetter", text="%2)", indent=1440, hanging=360),
        ... ])
        >>> apply_list(doc.add_paragraph("top level"), num)
        >>> apply_list(doc.add_paragraph("nested"), num, level=1)
    """
    levels = tuple(levels)
    _validate_levels(levels)
    if style_link is not None and num_style_link is not None:
        raise ValueError(
            "style_link and num_style_link are the two halves of a style/numbering "
            "pair and cannot both be set on one definition"
        )

    _, root = get_or_create_part(doc, NUMBERING_SPEC)

    if abstract_registry is None:
        abstract_registry = AbstractNumIdRegistry(doc)
    if num_registry is None:
        num_registry = NumIdRegistry(doc)

    abstract_id = abstract_registry.next_sequential()
    num_id = num_registry.next_sequential()

    if multi_level_type is None:
        multi_level_type = "singleLevel" if len(levels) == 1 else "multilevel"
    elif multi_level_type not in _MULTI_LEVEL_TYPES:
        raise ValueError(
            f"multi_level_type must be one of {sorted(_MULTI_LEVEL_TYPES)}; "
            f"got {multi_level_type!r}"
        )

    abstract_num = _build_abstract_num(
        abstract_id,
        levels,
        name=name,
        style_link=style_link,
        num_style_link=num_style_link,
        multi_level_type=multi_level_type,
    )
    # w:abstractNum must precede every w:num. python-docx's own helpers
    # only ever append, so this ordering is entirely on us.
    insert_before_first_anchor(root, abstract_num, _AFTER_ABSTRACT_NUM)

    num = el("w:num", **{"w:numId": str(num_id)})
    sub(num, "w:abstractNumId", **{"w:val": str(abstract_id)})
    insert_before_first_anchor(root, num, _AFTER_NUM)

    return num_id


def define_bullet_list(
    doc: Document,
    *,
    levels: int = 1,
    indent_step: int = 720,
    hanging: int = 360,
    num_registry: NumIdRegistry | None = None,
    abstract_registry: AbstractNumIdRegistry | None = None,
) -> int:
    """Define a bulleted list with Word's default glyph cycle.

    Word cycles three bullets by depth — a filled round bullet, a hollow
    ``o``, then a filled square — each needing its own symbol font to
    render as anything but a Latin letter.

    Args:
        doc: The python-docx :class:`~docx.document.Document` to mutate.
        levels: How many outline levels to define, 1 to 9.
        indent_step: Twips of left indent added per level (720 = 0.5").
        hanging: Hanging indent in twips for every level.
        num_registry: Pre-existing ``w:numId`` allocator.
        abstract_registry: Pre-existing ``w:abstractNumId`` allocator.

    Returns:
        The ``w:numId`` to pass to
        :func:`~docx_plus.numbering.apply_list`.

    Raises:
        InvalidLevelError: If ``levels`` is outside 1 to 9.

    Example:
        >>> from docx import Document
        >>> from docx_plus.numbering import apply_list, define_bullet_list
        >>> doc = Document()
        >>> bullets = define_bullet_list(doc, levels=2)
        >>> apply_list(doc.add_paragraph("first"), bullets)
    """
    return define_list_definition(
        doc,
        levels=[
            _preset_level(_BULLET_CYCLE[index % len(_BULLET_CYCLE)], index, indent_step, hanging)
            for index in range(_checked_level_count(levels))
        ],
        num_registry=num_registry,
        abstract_registry=abstract_registry,
    )


def define_numbered_list(
    doc: Document,
    *,
    levels: int = 1,
    indent_step: int = 720,
    hanging: int = 360,
    num_registry: NumIdRegistry | None = None,
    abstract_registry: AbstractNumIdRegistry | None = None,
) -> int:
    """Define a numbered list with Word's default format cycle.

    Word cycles ``1.`` → ``a.`` → ``i.`` by depth. Each level's counter
    stands alone; for the legal-outline shape (``1.1``, ``1.1.1``) build
    the levels yourself with ``text="%1.%2."`` and pass them to
    :func:`define_list_definition`.

    Args:
        doc: The python-docx :class:`~docx.document.Document` to mutate.
        levels: How many outline levels to define, 1 to 9.
        indent_step: Twips of left indent added per level (720 = 0.5").
        hanging: Hanging indent in twips for every level.
        num_registry: Pre-existing ``w:numId`` allocator.
        abstract_registry: Pre-existing ``w:abstractNumId`` allocator.

    Returns:
        The ``w:numId`` to pass to
        :func:`~docx_plus.numbering.apply_list`.

    Raises:
        InvalidLevelError: If ``levels`` is outside 1 to 9.

    Example:
        >>> from docx import Document
        >>> from docx_plus.numbering import apply_list, define_numbered_list
        >>> doc = Document()
        >>> steps = define_numbered_list(doc, levels=3)
        >>> apply_list(doc.add_paragraph("step one"), steps)
    """
    return define_list_definition(
        doc,
        levels=[
            _preset_level(_NUMBER_CYCLE[index % len(_NUMBER_CYCLE)], index, indent_step, hanging)
            for index in range(_checked_level_count(levels))
        ],
        num_registry=num_registry,
        abstract_registry=abstract_registry,
    )


# ---------------------------------------------------------------------------
# Internals.
# ---------------------------------------------------------------------------

# Word's default bullet cycle: (glyph, font). U+F0B7 and U+F0A7 are
# private-use codepoints that render as a filled bullet and a filled
# square *only* in the paired symbol font.
_BULLET_CYCLE: tuple[tuple[str, str, str], ...] = (
    ("bullet", "", "Symbol"),
    ("bullet", "o", "Courier New"),
    ("bullet", "", "Wingdings"),
)

# Word's default numbered cycle by depth.
_NUMBER_CYCLE: tuple[tuple[str, str, str], ...] = (
    ("decimal", "%{n}.", ""),
    ("lowerLetter", "%{n}.", ""),
    ("lowerRoman", "%{n}.", ""),
)


def _checked_level_count(levels: int) -> int:
    if not 1 <= levels <= MAX_LEVELS:
        raise InvalidLevelError(f"levels must be between 1 and {MAX_LEVELS}; got {levels!r}")
    return levels


def _preset_level(
    spec: tuple[str, str, str],
    index: int,
    indent_step: int,
    hanging: int,
) -> LevelDefinition:
    """Build one preset level. ``spec`` is ``(fmt, text, font)``."""
    fmt, text, font = spec
    return LevelDefinition(
        fmt=fmt,
        # The numbered cycle's text is a template because the placeholder
        # is level-dependent; the bullet cycle's is a literal glyph.
        text=text.replace("{n}", str(index + 1)),
        indent=indent_step * (index + 1),
        hanging=hanging,
        font=font or None,
    )


def _validate_levels(levels: tuple[LevelDefinition, ...]) -> None:
    """Check the level list as a whole — count, and placeholder depth."""
    if not levels:
        raise InvalidLevelError("a list definition needs at least one level")
    if len(levels) > MAX_LEVELS:
        raise InvalidLevelError(
            f"a w:abstractNum holds at most {MAX_LEVELS} levels; got {len(levels)}"
        )
    for index, level in enumerate(levels):
        if level.fmt == "bullet":
            # lvlText is a literal glyph for bullets, so a '%' in it is
            # just a percent sign.
            continue
        for match in _PLACEHOLDER_RE.finditer(level.text):
            referenced = int(match.group(1))
            if referenced > index + 1:
                raise InvalidLevelError(
                    f"level {index} references %{referenced} in its lvlText, but a "
                    f"level can only interpolate counters at or above its own depth "
                    f"(%1 to %{index + 1})"
                )


def _build_abstract_num(
    abstract_id: int,
    levels: tuple[LevelDefinition, ...],
    *,
    name: str | None,
    style_link: str | None,
    num_style_link: str | None,
    multi_level_type: str,
) -> etree._Element:
    """Build a complete ``<w:abstractNum>`` in schema order.

    ``w:nsid`` and ``w:tmpl`` are deliberately omitted. Both are
    optional, Word regenerates them, and minting one would mean inventing
    a fourth id namespace for no behavioural gain.
    """
    abstract_num = el("w:abstractNum", **{"w:abstractNumId": str(abstract_id)})

    ordered_insert(
        abstract_num,
        el("w:multiLevelType", **{"w:val": multi_level_type}),
        _ABSTRACT_NUM_CHILD_ORDER,
    )
    if name is not None:
        ordered_insert(abstract_num, el("w:name", **{"w:val": name}), _ABSTRACT_NUM_CHILD_ORDER)
    if style_link is not None:
        ordered_insert(
            abstract_num,
            el("w:styleLink", **{"w:val": style_link}),
            _ABSTRACT_NUM_CHILD_ORDER,
        )
    if num_style_link is not None:
        ordered_insert(
            abstract_num,
            el("w:numStyleLink", **{"w:val": num_style_link}),
            _ABSTRACT_NUM_CHILD_ORDER,
        )

    # ordered_insert replaces same-tag siblings, which is wrong for the
    # nine repeating w:lvl children — they are appended in order instead,
    # and w:lvl is last in the schema sequence anyway.
    for index, level in enumerate(levels):
        abstract_num.append(_build_level(index, level))

    return abstract_num


def _build_level(ilvl: int, level: LevelDefinition) -> etree._Element:
    """Build one ``<w:lvl>`` in ECMA-376 17.9.6 child order."""
    lvl = el("w:lvl", **{"w:ilvl": str(ilvl)})

    def place(tag: str, **attrs: str) -> None:
        ordered_insert(lvl, el(tag, **attrs), _LVL_CHILD_ORDER)

    place("w:start", **{"w:val": str(level.start)})
    place("w:numFmt", **{"w:val": level.fmt})
    if level.restart_after is not None:
        place("w:lvlRestart", **{"w:val": str(level.restart_after)})
    if level.suffix != "tab":
        # "tab" is the schema default; Word omits the element for it.
        place("w:suff", **{"w:val": level.suffix})
    place("w:lvlText", **{"w:val": level.text})
    place("w:lvlJc", **{"w:val": level.justify})

    if level.indent is not None:
        ppr = el("w:pPr")
        ind_attrs = {"w:left": str(level.indent)}
        if level.hanging is not None:
            ind_attrs["w:hanging"] = str(level.hanging)
        sub(ppr, "w:ind", **ind_attrs)
        ordered_insert(lvl, ppr, _LVL_CHILD_ORDER)

    if level.font is not None:
        rpr = el("w:rPr")
        sub(
            rpr,
            "w:rFonts",
            **{
                "w:ascii": level.font,
                "w:hAnsi": level.font,
                # Without an explicit hint Word may substitute a theme
                # font for the symbol font and render the bullet as a
                # Latin letter.
                "w:hint": "default",
            },
        )
        ordered_insert(lvl, rpr, _LVL_CHILD_ORDER)

    return lvl


def _abstract_id_for(root: etree._Element, num_id: int) -> str | None:
    """Return the ``abstractNumId`` a ``w:num`` points at, or ``None``."""
    from docx_plus.core.oxml import xpath

    matches = xpath(root, "./w:num[@w:numId=$nid]", nid=str(num_id))
    if not matches:
        return None
    ref = matches[0].find(qn("w:abstractNumId"))
    if ref is None:
        return None
    value = ref.get(qn("w:val"))
    return str(value) if value is not None else None


__all__ = [
    "MAX_LEVELS",
    "InvalidLevelError",
    "Justification",
    "LevelDefinition",
    "MultiLevelType",
    "Suffix",
    "define_bullet_list",
    "define_list_definition",
    "define_numbered_list",
]
