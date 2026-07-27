r"""Cross-references to bookmarks via ``REF`` / ``PAGEREF`` fields.

``REF bookmark_name`` inserts the text Word reads from the bookmark's
range; ``PAGEREF bookmark_name`` inserts the page number Word renders
for the bookmark. Both are complex fields built on top of the same
plumbing :mod:`docx_plus.fields` uses for page numbers and dates.

The switches matter more than they look. A bare ``REF`` yields the
bookmark's *text*, which is what "see Figure 3" wants when the bookmark
brackets a caption's label and number — see
:func:`docx_plus.publishing.add_caption`'s ``bookmark_name``. Add ``\p``
and the same reference yields "above" or "below" instead; add ``\r`` and
it yields the target paragraph's *number*, which is what a numbered
heading or a list built with :mod:`docx_plus.numbering` gives you.

This module imports only from ``docx_plus.core`` (SPEC §9.1).
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Literal

from lxml import etree

from docx_plus.core.oxml import build_complex_field, validate_bookmark_name

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

CrossReferenceKind = Literal["text", "page"]

#: How much surrounding context a paragraph-number reference carries.
#:
#: - ``"plain"`` (``\n``) — the number with no context: ``3``
#: - ``"relative"`` (``\r``) — relative context: ``2.3``
#: - ``"full"`` (``\w``) — full context: ``Chapter 4, 2.3``
#:
#: Word's own UI calls these "Paragraph number", "Paragraph number (no
#: context)", and "Paragraph number (full context)"; the mapping between
#: its labels and the switches is not one-to-one, so the switch letters
#: are what these names track.
NumberContext = Literal["plain", "relative", "full"]

_NUMBER_SWITCHES: dict[str, str] = {
    "plain": "n",
    "relative": "r",
    "full": "w",
}


def add_cross_reference(
    paragraph: Paragraph,
    *,
    bookmark: str,
    kind: CrossReferenceKind = "text",
    hyperlink: bool = True,
    number: NumberContext | None = None,
    position: bool = False,
    suppress_non_delimiters: bool = False,
    numeric_format: str | None = None,
    preserve_formatting: bool = False,
) -> etree._Element:
    r"""Append a cross-reference complex field to ``paragraph``.

    Args:
        paragraph: A python-docx :class:`~docx.text.paragraph.Paragraph`
            where the cross-reference field is appended after any
            existing runs.
        bookmark: The target bookmark's ``w:name`` attribute. Must match
            an existing bookmark for Word to resolve the field;
            unresolved cross-references render as ``"Error! Reference
            source not found."``. Validated against Word's name grammar,
            since a name only Word's UI would reject produces exactly
            that silent failure.
        kind: ``"text"`` (default) inserts a ``REF`` field that resolves
            to the bookmark's text content; ``"page"`` inserts a
            ``PAGEREF`` field that resolves to the page number where
            the bookmark sits.
        hyperlink: ``True`` (default) appends ``\h`` so Word makes the
            cross-reference a clickable link to the bookmark.
        number: Resolve to the target paragraph's *number* rather than its
            text, with the given amount of context — see
            :data:`NumberContext`. Only meaningful when the target is a
            numbered paragraph (a numbered heading, or a list from
            :mod:`docx_plus.numbering`). ``REF`` only.
        position: Append ``\p``, so the field resolves to ``"above"`` or
            ``"below"`` depending on where the target sits relative to the
            reference. Combines with ``number`` — Word then renders
            something like ``"2.3 above"``. ``REF`` only.
        suppress_non_delimiters: Append ``\t``, which drops the target's
            non-delimiter text and keeps only its numbering. Rarely
            wanted on its own; pairs with ``number``. ``REF`` only.
        numeric_format: A ``\#`` numeric-picture switch such as
            ``"0.00"``, applied when the target resolves to a number.
            Must not contain a double-quote, which would terminate the
            switch.
        preserve_formatting: Append ``\* MERGEFORMAT``, so the character
            formatting a user applies to the field's result survives the
            next recalculation. Word's UI sets this by default; this
            library does not, to keep the emitted instruction minimal.

    Returns:
        The begin ``w:r`` run that marks the start of the field, same
        contract as :func:`docx_plus.fields.add_page_number_field`.

    Raises:
        ValueError: If ``bookmark`` is not a valid bookmark name; if
            ``numeric_format`` contains a double-quote; or if
            ``kind="page"`` is combined with ``number``, ``position``, or
            ``suppress_non_delimiters``, none of which ``PAGEREF``
            accepts.

    Example:
        >>> from docx import Document
        >>> from docx_plus.bookmarks import add_bookmark, add_cross_reference
        >>> doc = Document()
        >>> p1 = doc.add_paragraph("Section 1")
        >>> _ = add_bookmark(p1, "sec_1")
        >>> p2 = doc.add_paragraph("See ")
        >>> _ = add_cross_reference(p2, bookmark="sec_1", kind="text")
        >>> p2.add_run(" on page ")
        <docx.text.run.Run object at 0x...>
        >>> _ = add_cross_reference(p2, bookmark="sec_1", kind="page")

    Notes:
        Fields are cached: Word displays the previously-computed result
        until ``w:updateFields="true"`` triggers recalculation on open.
        Pair calls to :func:`add_cross_reference` with
        :func:`docx_plus.fields.mark_fields_dirty` so the new
        cross-references resolve on first open.
    """
    validate_bookmark_name(bookmark, arg_name="bookmark")

    if kind == "page" and (number is not None or position or suppress_non_delimiters):
        raise ValueError(
            "PAGEREF resolves to a page number and accepts none of number / "
            'position / suppress_non_delimiters; use kind="text" for those'
        )
    if number is not None and number not in _NUMBER_SWITCHES:
        raise ValueError(f"number must be one of {sorted(_NUMBER_SWITCHES)}; got {number!r}")
    if numeric_format is not None and '"' in numeric_format:
        raise ValueError(
            f"numeric_format must not contain a double-quote, which would "
            f"terminate the switch; got {numeric_format!r}"
        )

    keyword = "REF" if kind == "text" else "PAGEREF"
    switches: list[str] = []
    if number is not None:
        switches.append(f"\\{_NUMBER_SWITCHES[number]}")
    if position:
        switches.append("\\p")
    if suppress_non_delimiters:
        switches.append("\\t")
    if hyperlink:
        switches.append("\\h")
    if numeric_format is not None:
        switches.append(f'\\# "{numeric_format}"')
    if preserve_formatting:
        switches.append("\\* MERGEFORMAT")

    tail = f" {' '.join(switches)}" if switches else ""
    instruction = f" {keyword} {bookmark}{tail} "
    return build_complex_field(paragraph._p, instruction, "")


__all__ = ["CrossReferenceKind", "NumberContext", "add_cross_reference"]
