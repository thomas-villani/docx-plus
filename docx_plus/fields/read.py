r"""Read the fields already in a document.

The read half of :mod:`docx_plus.fields.simple`. A complex field is not a
single element — it is a *run sequence* delimited by ``w:fldChar`` markers,
with the instruction spread across however many ``w:instrText`` elements
Word happened to split it into. Reading one back means walking that
sequence, which is why this is a capability rather than a two-line xpath at
each call site.

Anything auditing a document's cross-references, captions, or table of
contents starts here: those are all fields, and their *instruction* is the
only place their meaning is recorded. The cached result text is whatever
Word last rendered and may be arbitrarily stale.

This module imports only from ``docx_plus.core`` (SPEC §9.1).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import TYPE_CHECKING

from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath

if TYPE_CHECKING:
    from docx.document import Document
    from lxml import etree


# Field instruction arguments are whitespace-separated, except that a
# quoted string is one argument. Switches (\h, \* MERGEFORMAT) are captured
# as ordinary arguments — a caller filtering on the leading backslash is
# clearer than a parser guessing which switches take a value.
_ARGUMENT = re.compile(r'"[^"]*"|\S+')


@dataclass(frozen=True)
class FieldInfo:
    r"""One complex field, as recorded in the document.

    Attributes:
        keyword: The field type, upper-cased — ``"REF"``, ``"PAGEREF"``,
            ``"SEQ"``, ``"TOC"``, ``"PAGE"``. Empty for a field with no
            instruction at all, which is malformed but occurs in the wild.
        arguments: The remaining whitespace-separated tokens, with any
            surrounding quotes stripped. Switches are included in order,
            so ``REF fig1 \\h`` gives ``["fig1", "\\h"]``.
        instruction: The raw instruction text, joined across every
            ``w:instrText`` in the field and stripped.
        result: The cached result text Word last rendered, or ``""`` for a
            field with no result yet. **Not** authoritative — it is only
            what was displayed when the document was saved.
        paragraph_index: 0-based position of the owning paragraph among
            the body's ``w:p`` elements.
        begin_element: The ``w:r`` carrying the ``begin`` ``w:fldChar``,
            for callers that need to locate or rewrite the field.

    Note:
        ``paragraph_index`` counts every ``w:p`` in the body including
        those inside table cells, matching
        :func:`~docx_plus.bookmarks.read_bookmarks`. That is a different
        numbering from ``doc.paragraphs``.
    """

    keyword: str
    arguments: list[str] = field(default_factory=list)
    instruction: str = ""
    result: str = ""
    paragraph_index: int = -1
    begin_element: etree._Element | None = None

    @property
    def switches(self) -> list[str]:
        r"""Just the ``\x`` switch arguments, in order."""
        return [arg for arg in self.arguments if arg.startswith("\\")]

    @property
    def operands(self) -> list[str]:
        r"""The arguments that are not switches — a bookmark name, a SEQ id."""
        return [arg for arg in self.arguments if not arg.startswith("\\")]


def read_fields(doc: Document, *, keyword: str | None = None) -> list[FieldInfo]:
    r"""Return every complex field in ``doc``'s body, in document order.

    Args:
        doc: The python-docx :class:`~docx.document.Document` to scan.
        keyword: If given, return only fields of this type, matched
            case-insensitively (``"REF"``, ``"SEQ"``, ...).

    Returns:
        One :class:`FieldInfo` per field.

    Note:
        Only the main document body is scanned. A ``PAGE`` field in a
        footer — the usual place for one — is not returned. Headers,
        footers, and notes live in separate parts.

    Note:
        Nested fields (a field inside another field's instruction, which
        Word writes for some ``TOC`` and ``IF`` constructions) are read as
        one field: the inner instruction text is folded into the outer.
        The keyword is still the outer field's, which is what a caller
        filtering by type expects.

    Example:
        >>> from docx import Document
        >>> from docx_plus.bookmarks import add_bookmark, add_cross_reference
        >>> from docx_plus.fields import read_fields
        >>> doc = Document()
        >>> target = doc.add_paragraph("Chapter One")
        >>> _ = add_bookmark(target, "chapter1")
        >>> _ = add_cross_reference(doc.add_paragraph("See "), bookmark="chapter1")
        >>> for found in read_fields(doc):
        ...     print(found.keyword, found.operands)
        REF ['chapter1']
    """
    wanted = keyword.upper() if keyword is not None else None
    fields: list[FieldInfo] = []

    for paragraph_index, p_element in enumerate(xpath(doc.element.body, ".//w:p")):
        for info in _fields_in_paragraph(p_element, paragraph_index):
            if wanted is None or info.keyword == wanted:
                fields.append(info)
    return fields


def _fields_in_paragraph(p_element: etree._Element, paragraph_index: int) -> list[FieldInfo]:
    """Walk one paragraph's runs, assembling each begin..end field sequence."""
    fields: list[FieldInfo] = []
    begin_run: etree._Element | None = None
    depth = 0
    instruction: list[str] = []
    result: list[str] = []
    in_result = False

    for run in xpath(p_element, ".//w:r"):
        fld_char = run.find(qn("w:fldChar"))
        if fld_char is not None:
            char_type = fld_char.get(qn("w:fldCharType"))
            if char_type == "begin":
                depth += 1
                if depth == 1:
                    begin_run, instruction, result, in_result = run, [], [], False
            elif char_type == "separate" and depth == 1:
                in_result = True
            elif char_type == "end":
                depth -= 1
                if depth == 0 and begin_run is not None:
                    fields.append(_build(begin_run, instruction, result, paragraph_index))
                    begin_run, in_result = None, False
                depth = max(depth, 0)
            continue

        if begin_run is None:
            continue
        if in_result:
            result.extend(t.text or "" for t in run.findall(qn("w:t")))
        else:
            instruction.extend(t.text or "" for t in run.findall(qn("w:instrText")))

    return fields


def _build(
    begin_run: etree._Element,
    instruction: list[str],
    result: list[str],
    paragraph_index: int,
) -> FieldInfo:
    text = "".join(instruction).strip()
    tokens = [token.strip('"') for token in _ARGUMENT.findall(text)]
    return FieldInfo(
        keyword=tokens[0].upper() if tokens else "",
        arguments=tokens[1:],
        instruction=text,
        result="".join(result),
        paragraph_index=paragraph_index,
        begin_element=begin_run,
    )


__all__ = ["FieldInfo", "read_fields"]
