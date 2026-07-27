"""Insert OOXML complex fields (PAGE, DATE, generic) into paragraphs.

Word fields use the *complex field* syntax: a sequence of runs containing
``w:fldChar`` markers (``begin``/``separate``/``end``) bracketing the field
instruction text (``w:instrText``) and the result text (``w:t``). This
module emits that five-run sequence and appends it to a paragraph.

Word recalculates field results on open only if ``w:updateFields`` is set in
``settings.xml`` — see :func:`docx_plus.fields.update.mark_fields_dirty`.
Initial text supplied here is what Word shows *before* it recalculates, so
the value is meaningful for offline viewers (e.g. ``"1"`` for a PAGE field).

This module imports only from ``docx_plus.core`` (SPEC §9.1).
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Literal

from lxml import etree

from docx_plus.core.oxml import build_complex_field

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph


PageFieldName = Literal["PAGE", "NUMPAGES", "SECTIONPAGES"]

#: How much surrounding context a ``STYLEREF`` number carries: ``"plain"``
#: (``\n``) gives the bare number, ``"relative"`` (``\r``) adds relative
#: context, ``"full"`` (``\w``) adds full context.
StyleRefNumber = Literal["plain", "relative", "full"]

_STYLEREF_NUMBER_SWITCHES: dict[str, str] = {
    "plain": "n",
    "relative": "r",
    "full": "w",
}

#: Word's outline levels, which ``STYLEREF`` accepts in place of a style
#: name.
_MAX_OUTLINE_LEVEL = 9


def add_page_number_field(
    paragraph: Paragraph,
    *,
    field: PageFieldName = "PAGE",
    format: str | None = None,
) -> etree._Element:
    r"""Append a page-number complex field to ``paragraph``.

    Args:
        paragraph: A python-docx :class:`~docx.text.paragraph.Paragraph`. The
            field runs are appended after the paragraph's existing runs.
        field: Which page-number variant. ``"PAGE"`` (current page),
            ``"NUMPAGES"`` (total pages), or ``"SECTIONPAGES"`` (pages in the
            current section).
        format: Optional field switches appended to the instruction. Example:
            ``r"\* ARABIC"`` forces Arabic numerals, ``r"\* ROMAN"`` Roman.
            See ECMA-376 17.16 for the switch syntax.

    Returns:
        The begin ``w:r`` run that marks the start of the field.

    Example:
        >>> from docx import Document
        >>> from docx_plus.fields import add_page_number_field
        >>> doc = Document()
        >>> p = doc.add_paragraph("Page ")
        >>> _ = add_page_number_field(p)
    """
    if format is None or not format.strip():
        instruction = f" {field} "
    else:
        instruction = f" {field} {format.strip()} "
    return build_complex_field(paragraph._p, instruction, "1")


def add_date_field(
    paragraph: Paragraph,
    *,
    format: str = "MMMM d, yyyy",
    auto_update: bool = True,
) -> etree._Element:
    """Append a date complex field to ``paragraph``.

    Args:
        paragraph: A python-docx :class:`~docx.text.paragraph.Paragraph`.
        format: A Word date-format string. Common values:
            ``"MMMM d, yyyy"`` (default, e.g. *May 19, 2026*),
            ``"M/d/yyyy"`` (numeric short), ``"dddd, MMMM d, yyyy"`` (long
            with weekday).
        auto_update: ``True`` (default) emits a ``DATE`` field that Word
            recalculates on every open. ``False`` emits a ``CREATEDATE`` field
            that freezes the document's creation date.

    Returns:
        The begin ``w:r`` run that marks the start of the field.

    Example:
        >>> from docx import Document
        >>> from docx_plus.fields import add_date_field
        >>> doc = Document()
        >>> p = doc.add_paragraph("Today: ")
        >>> _ = add_date_field(p, format="M/d/yyyy")
    """
    keyword = "DATE" if auto_update else "CREATEDATE"
    instruction = f' {keyword} \\@ "{format}" '
    return build_complex_field(paragraph._p, instruction, "")


def add_style_reference(
    paragraph: Paragraph,
    *,
    style: str | int,
    search_from_bottom: bool = False,
    number: StyleRefNumber | None = None,
    position: bool = False,
    suppress_non_delimiters: bool = False,
    preserve_formatting: bool = True,
) -> etree._Element:
    r"""Append a ``STYLEREF`` complex field to ``paragraph``.

    ``STYLEREF`` resolves to the text of the nearest paragraph carrying a
    given style — the field behind a running header that shows the
    current chapter title, and the one cross-reference kind that needs no
    bookmark at all. Word re-resolves it per page, so the same field in a
    header renders differently on every page.

    Args:
        paragraph: A python-docx :class:`~docx.text.paragraph.Paragraph`.
            Usually one belonging to a header — ``section.header.paragraphs[0]``.
        style: The style to search for. A **string** is the style's *name*
            as Word shows it (``"Heading 1"``, with the space) — not the
            ``w:styleId``, unlike most of this library, because that is
            what the field instruction takes. An **int** is an outline
            level, ``1`` to ``9``, which Word accepts as a shorthand for
            the corresponding built-in heading.
        search_from_bottom: Append ``\l``, so Word takes the *last*
            matching paragraph on the page rather than the first. This is
            what you want for a "to" half of a dictionary-style header
            range; the default first-match suits a chapter title.
        number: Resolve to the matched paragraph's *number* rather than
            its text, with the given amount of context — see
            :data:`StyleRefNumber`. Only meaningful when the matched
            paragraphs are numbered.
        position: Append ``\p``, so the field resolves to ``"above"`` or
            ``"below"`` relative to the field itself.
        suppress_non_delimiters: Append ``\t``, dropping the matched
            paragraph's non-delimiter text and keeping only its numbering.
        preserve_formatting: Append ``\* MERGEFORMAT`` (default ``True``,
            matching what Word's UI emits for this field) so character
            formatting applied to the result survives recalculation.

    Returns:
        The begin ``w:r`` run that marks the start of the field.

    Raises:
        ValueError: If ``style`` is an empty string, contains a
            double-quote (which would terminate the quoted argument), is
            an int outside 1–9, or if ``number`` is not a recognised
            context.

    Example:
        >>> from docx import Document
        >>> from docx_plus.fields import add_style_reference, mark_fields_dirty
        >>> doc = Document()
        >>> header = doc.sections[0].header.paragraphs[0]
        >>> _ = add_style_reference(header, style="Heading 1")
        >>> mark_fields_dirty(doc)

    Notes:
        Like every field, this renders blank until Word recalculates.
        Pair it with :func:`docx_plus.fields.mark_fields_dirty`.
    """
    argument = _style_reference_argument(style)

    if number is not None and number not in _STYLEREF_NUMBER_SWITCHES:
        raise ValueError(
            f"number must be one of {sorted(_STYLEREF_NUMBER_SWITCHES)}; got {number!r}"
        )

    switches: list[str] = []
    if number is not None:
        switches.append(f"\\{_STYLEREF_NUMBER_SWITCHES[number]}")
    if search_from_bottom:
        switches.append("\\l")
    if position:
        switches.append("\\p")
    if suppress_non_delimiters:
        switches.append("\\t")
    if preserve_formatting:
        switches.append("\\* MERGEFORMAT")

    tail = f" {' '.join(switches)}" if switches else ""
    return build_complex_field(paragraph._p, f" STYLEREF {argument}{tail} ", "")


def _style_reference_argument(style: str | int) -> str:
    """Render ``style`` as the field's first argument.

    A style *name* is quoted, because names contain spaces. An outline
    level is a bare integer. ``bool`` is rejected explicitly — it is an
    ``int`` subclass, and ``style=True`` would silently become outline
    level 1.
    """
    if isinstance(style, bool):
        raise ValueError(
            f"style must be a style name or an outline level, not a bool; got {style!r}"
        )
    if isinstance(style, int):
        if not 1 <= style <= _MAX_OUTLINE_LEVEL:
            raise ValueError(
                f"an outline-level style reference must be 1 to {_MAX_OUTLINE_LEVEL}; got {style!r}"
            )
        return str(style)
    if not style.strip():
        raise ValueError("style must be a non-empty style name")
    if '"' in style:
        raise ValueError(
            f"style must not contain a double-quote, which would terminate the "
            f"quoted field argument; got {style!r}"
        )
    return f'"{style}"'


def add_field(
    paragraph: Paragraph,
    *,
    instruction: str,
    initial_text: str = "",
) -> etree._Element:
    r"""Append a generic complex field to ``paragraph``.

    Use this for fields without a dedicated helper (``TOC``, ``REF``,
    ``HYPERLINK``, ``MERGEFIELD``, etc.). The ``instruction`` is wrapped in
    leading/trailing spaces if you don't supply them, since Word's field
    parser requires them.

    Args:
        paragraph: A python-docx :class:`~docx.text.paragraph.Paragraph`.
        instruction: The raw field instruction text without the surrounding
            ``{ }`` braces Word shows in its UI. Example: ``'REF Bookmark1'``
            or ``'TOC \o "1-3" \h'``.
        initial_text: Optional placeholder shown before Word recalculates.

    Returns:
        The begin ``w:r`` run that marks the start of the field.

    Raises:
        ValueError: If ``instruction`` is empty or whitespace-only —
            Word renders the field as a silent blank, which is almost
            never intended (see issues.md M1).

    Example:
        >>> from docx import Document
        >>> from docx_plus.fields import add_field
        >>> doc = Document()
        >>> p = doc.add_paragraph()
        >>> _ = add_field(p, instruction='TOC \\o "1-3" \\h', initial_text="(TOC)")
    """
    stripped = instruction.strip()
    if not stripped:
        raise ValueError(f"add_field requires a non-empty instruction; got {instruction!r}")
    wrapped = f" {stripped} "
    return build_complex_field(paragraph._p, wrapped, initial_text)


__all__ = [
    "PageFieldName",
    "StyleRefNumber",
    "add_date_field",
    "add_field",
    "add_page_number_field",
    "add_style_reference",
]
