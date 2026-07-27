"""``people.xml`` — comment author presence.

Word records one entry per comment author in a separate
``/word/people.xml`` part:

.. code-block:: xml

    <w15:people>
      <w15:person w15:author="Thomas Villani">
        <w15:presenceInfo w15:providerId="AD"
                          w15:userId="S::thomas@example.com::541bd2ef-..."/>
      </w15:person>
    </w15:people>

The part drives the presence indicator beside a comment in Word's
reviewing pane — the dot that says whether the author is online, and the
identity a click resolves to. **It is purely cosmetic.** Comments,
threading, and resolution all work without it; Word neither requires it
nor complains when it is missing.

That is why nothing here runs automatically. :func:`~docx_plus.comments.add_comment`
deliberately does *not* write this part: registering an author means
inventing a ``userId`` for someone the library knows nothing about, and
a fabricated directory identity is worse than an absent one. Call
:func:`set_author_presence` explicitly when you want the entry.

Word does **not** prune stale authors when their last comment is
deleted, and neither does this module — see :func:`clear_author_presence`.

This module imports only from ``docx_plus.core`` (SPEC §9.1).
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import TYPE_CHECKING, cast

from docx.opc.part import XmlPart

from docx_plus.core.ns import qn
from docx_plus.core.oxml import el, remove, sub, xpath
from docx_plus.core.parts import PEOPLE_SPEC, RT_PEOPLE, get_or_create_part

if TYPE_CHECKING:
    from docx.document import Document
    from lxml import etree

#: What Word writes for an author with no directory identity behind it.
#: The Active Directory form is ``"AD"``; ``"Windows Live"`` and
#: ``"Office365"`` also appear in the wild.
LOCAL_PROVIDER = "None"


@dataclass(frozen=True)
class AuthorPresence:
    """One author's entry in ``people.xml``.

    Attributes:
        author: The author name, matching the ``w:author`` attribute on
            that person's comments. This is the join key — the part
            carries no comment ids.
        provider_id: The identity provider (``"AD"``, ``"Windows Live"``,
            ``"Office365"``, or :data:`LOCAL_PROVIDER` for an author with
            no directory behind them). ``None`` when the entry has no
            ``<w15:presenceInfo>`` child at all, which is legal.
        user_id: The provider-scoped identity string. ``None`` under the
            same condition as ``provider_id``.
    """

    author: str
    provider_id: str | None = None
    user_id: str | None = None


def set_author_presence(
    doc: Document,
    author: str,
    *,
    provider_id: str = LOCAL_PROVIDER,
    user_id: str | None = None,
) -> AuthorPresence | None:
    """Record ``author`` in ``people.xml``, creating the part on first use.

    Idempotent — re-registering an author replaces that author's entry
    rather than appending a second one.

    Args:
        doc: The python-docx :class:`~docx.document.Document` to write.
        author: Author name. Must match the ``w:author`` on that
            person's comments exactly, since the name is the only join
            between the two parts. An empty name is a no-op: Word writes
            no entry for one, and ``<w15:person w15:author=""/>`` is
            worse than nothing.
        provider_id: Identity provider. Defaults to
            :data:`LOCAL_PROVIDER`, which is what Word writes for an
            author who is not a directory user.
        user_id: Provider-scoped identity. Defaults to ``author``, which
            is what Word writes alongside :data:`LOCAL_PROVIDER`.

    Returns:
        The :class:`AuthorPresence` written, or ``None`` when ``author``
        is empty.

    Example:
        >>> from docx import Document
        >>> from docx_plus.comments import add_comment, set_author_presence
        >>> doc = Document()
        >>> add_comment(doc.add_paragraph("text"), "note", author="Reviewer")
        CommentRef(comment_id=..., body_element=...)
        >>> set_author_presence(doc, "Reviewer")
        AuthorPresence(author='Reviewer', provider_id='None', user_id='Reviewer')
    """
    if not author:
        return None

    _, root = get_or_create_part(doc, PEOPLE_SPEC)
    for existing in xpath(root, "./w15:person[@w15:author=$name]", name=author):
        remove(existing)

    resolved_user_id = author if user_id is None else user_id
    person = el("w15:person", **{"w15:author": author})
    sub(
        person,
        "w15:presenceInfo",
        **{"w15:providerId": provider_id, "w15:userId": resolved_user_id},
    )
    root.append(person)
    return AuthorPresence(author=author, provider_id=provider_id, user_id=resolved_user_id)


def read_author_presence(doc: Document) -> list[AuthorPresence]:
    """Return every author entry in ``doc``'s ``people.xml``.

    Args:
        doc: The python-docx :class:`~docx.document.Document` to read.

    Returns:
        One :class:`AuthorPresence` per ``<w15:person>``, in document
        order. Empty when the part is absent — the state of every
        document this library wrote before v0.5, and of anything not
        produced by Word. Entries with no ``w15:author`` are skipped;
        an entry with no ``<w15:presenceInfo>`` reports ``None`` for
        both provider and user id.

    Example:
        >>> from docx import Document
        >>> from docx_plus.comments import read_author_presence
        >>> read_author_presence(Document())
        []
    """
    root = _people_root(doc)
    if root is None:
        return []

    people: list[AuthorPresence] = []
    for person in xpath(root, "./w15:person"):
        author = person.get(qn("w15:author"))
        if not author:
            continue
        info = person.find(qn("w15:presenceInfo"))
        people.append(
            AuthorPresence(
                author=author,
                provider_id=None if info is None else info.get(qn("w15:providerId")),
                user_id=None if info is None else info.get(qn("w15:userId")),
            )
        )
    return people


def clear_author_presence(doc: Document, *, remove_part: bool = False) -> None:
    """Remove every author entry, optionally tearing down the part.

    Deliberately *not* wired into
    :func:`~docx_plus.comments.delete_comment`. Pruning on delete would
    need the author ref-counted across every surviving comment, and Word
    does not do it either — a document that once had a comment from a
    given author keeps that author's entry. Call this explicitly if a
    stale name is a problem.

    Args:
        doc: The python-docx :class:`~docx.document.Document` to scrub.
        remove_part: When ``False`` (default) the emptied part is left
            in place, so a later :func:`set_author_presence` reuses the
            relationship. When ``True`` the part and its relationship
            are torn down entirely. Mirrors
            :func:`~docx_plus.comments.clear_all_comments`.

    Example:
        >>> from docx import Document
        >>> from docx_plus.comments import clear_author_presence
        >>> clear_author_presence(Document())
    """
    root = _people_root(doc)
    if root is None:
        return

    if remove_part:
        for rid, rel in list(doc.part.rels.items()):
            if rel.reltype == RT_PEOPLE:
                doc.part.drop_rel(rid)
        return

    for person in xpath(root, "./w15:person"):
        remove(person)


def _people_root(doc: Document) -> etree._Element | None:
    """Return the ``<w15:people>`` root, or ``None`` if the part is absent."""
    try:
        part = cast("XmlPart", doc.part.part_related_by(RT_PEOPLE))
    except KeyError:
        return None
    return cast("etree._Element", part.element)


__all__ = [
    "LOCAL_PROVIDER",
    "AuthorPresence",
    "clear_author_presence",
    "read_author_presence",
    "set_author_presence",
]
