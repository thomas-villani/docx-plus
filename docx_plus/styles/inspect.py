"""Cascade resolver: ``resolve_effective_formatting``.

Walks the layers of OOXML formatting precedence (SPEC §4) and returns a
fully-resolved :class:`ResolvedFormatting` describing what a paragraph, run,
or cell would render with right now. Later layers override earlier ones,
except toggle properties (bold, italic, etc.), which follow the rule of
ECMA-376 17.7.3 — see :func:`_resolve_toggle`.

Provenance tracking is plumbed through the same walk gated by the
``include_provenance`` flag; with the flag off, the resolver's value output
is identical (verified by ``test_provenance_does_not_change_values``).
"""

from __future__ import annotations

from dataclasses import dataclass, field, replace
from typing import TYPE_CHECKING, Any, Literal, get_args

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

from docx_plus.core import DocxPlusError
from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.styles.theme import (
    ThemeColors,
    load_theme,
    resolve_theme_color,
    resolve_theme_font,
)

if TYPE_CHECKING:
    from docx.document import Document
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run


_MAX_STYLE_CHAIN_DEPTH = 11


@dataclass(frozen=True)
class TableContext:
    """A cell's position within its table — for conditional table-style formatting.

    ECMA-376 17.7.6.5 lets a ``<w:style w:type="table">`` carry
    conditional formatting branches (``<w:tblStylePr w:type="firstRow"/>``,
    ``"lastRow"``, ``"firstCol"``, ``"lastCol"``, ``"band1Horz"``,
    ``"band1Vert"``, ``"band2Horz"``, ``"band2Vert"``,
    ``"nwCell"`` / ``"neCell"`` / ``"swCell"`` / ``"seCell"``). To pick
    the right branches the cascade resolver needs to know where in the
    table the target lives.

    Construct manually for an out-of-band query, or pass a ``_Cell`` to
    :func:`resolve_effective_formatting` to derive the context
    automatically from the cell's parent row / table.

    **Position is not enough.** A table carries a ``<w:tblLook>`` saying
    which conditional branches it wants — the "Header Row" / "First
    Column" / "Banded Rows" tick-boxes in Word's Table Design tab. A cell
    in the first row of a table whose ``tblLook`` clears ``firstRow`` gets
    no ``firstRow`` formatting at all. The four ``*_enabled`` attributes
    carry those flags; the derived context reads them from the table, and
    they default to True so a hand-built context behaves like a table with
    no ``<w:tblLook>`` (which Word treats as "everything enabled").

    Banding is folded into the ``is_band*`` attributes rather than
    exposed as flags, because deciding them needs the whole table:

    - Bands exist only when a ``<w:tblStyleRowBandSize>`` /
      ``<w:tblStyleColBandSize>`` is declared, on the table instance or
      anywhere in its style chain. **Absent means no banding** — it is
      not a band size of 1. Instance beats style.
    - The stripe sequence starts at row / column 0 *unless* the matching
      ``firstRow`` / ``firstCol`` conditional actually paints that line,
      in which case it starts at 1. The ``tblLook`` flag alone does not
      shift it; the style must also define the branch.

    Scope: this context selects which ``<w:tblStylePr>`` branches apply,
    but only their **run / paragraph** properties are resolved. Cell-,
    row-, and table-level properties (``<w:tcPr>`` / ``<w:trPr>`` /
    ``<w:tblPr>``) from a table style are not surfaced — see the
    :func:`resolve_effective_formatting` note.

    Auto-derivation limitation: when a row wraps its cells in a
    ``<w:sdt>`` (a content control around table cells), the derived
    column index cannot be computed and an empty (all-False)
    :class:`TableContext` is returned. Pass an explicit context in that
    case. Nested tables resolve against the **inner** cell's position.

    Attributes:
        is_first_row: Cell is in the first ``<w:tr>`` of its table.
        is_last_row: Cell is in the last ``<w:tr>``.
        is_first_col: Cell is the first ``<w:tc>`` of its row.
        is_last_col: Cell is the last ``<w:tc>`` of its row.
        is_band_row: Cell is in a "band1" horizontal stripe.
        is_band_col: Cell is in a "band1" vertical stripe.
        is_band2_row: Cell is in a "band2" horizontal stripe.
        is_band2_col: Cell is in a "band2" vertical stripe.
        first_row_enabled: The table's ``tblLook`` asks for ``firstRow``
            formatting. When False the ``firstRow`` (and ``nwCell`` /
            ``neCell``) branches are suppressed even in row 0.
        last_row_enabled: As above for ``lastRow`` / ``swCell`` /
            ``seCell``.
        first_col_enabled: As above for ``firstCol`` / ``nwCell`` /
            ``swCell``.
        last_col_enabled: As above for ``lastCol`` / ``neCell`` /
            ``seCell``.
    """

    is_first_row: bool = False
    is_last_row: bool = False
    is_first_col: bool = False
    is_last_col: bool = False
    is_band_row: bool = False
    is_band_col: bool = False
    is_band2_row: bool = False
    is_band2_col: bool = False
    first_row_enabled: bool = True
    last_row_enabled: bool = True
    first_col_enabled: bool = True
    last_col_enabled: bool = True


# ``<w:tblStylePr w:type=...>`` values in application order: later entries
# override earlier ones. Which of them are candidates for a given cell is
# :func:`_matching_conditional_types`.
#
# This order was measured against Word, and it is NOT the order ECMA-376
# 17.7.6.5 lists — the spec puts the vertical bands before the horizontal
# ones and the row branches before the column ones, and Word does neither.
# Word also *rewrites* a style's branches into the spec's order when it
# saves, so the document order is no guide either. What Word actually does:
# a vertical band beats a horizontal one, a row branch beats a column
# branch, and the corners beat everything.
#
# ``wholeTable`` is deliberately absent. Word discards
# ``<w:tblStylePr w:type="wholeTable">`` outright — it neither renders its
# ``rPr`` / ``pPr`` nor keeps the element on save. Whole-table formatting
# lives on the style's own ``w:rPr`` / ``w:pPr``, which the base pass
# already applies.
_TBL_STYLE_PR_ORDER: tuple[str, ...] = (
    "band1Horz",
    "band2Horz",
    "band1Vert",
    "band2Vert",
    "firstCol",
    "lastCol",
    "firstRow",
    "lastRow",
    "nwCell",
    "neCell",
    "swCell",
    "seCell",
)

# ``<w:tblLook>`` bit values per ECMA-376 17.4.56, for the legacy ``w:val``
# bitmask form that Word still writes alongside the named attributes.
_TBL_LOOK_BITS: dict[str, int] = {
    "firstRow": 0x0020,
    "lastRow": 0x0040,
    "firstColumn": 0x0080,
    "lastColumn": 0x0100,
    "noHBand": 0x0200,
    "noVBand": 0x0400,
}

# Toggle properties combine per ECMA-376 17.7.3 rather than by override.
# Mapped from rPr child name to ResolvedFormatting field name.
_TOGGLE_RPR: dict[str, str] = {
    "b": "bold",
    "i": "italic",
    "bCs": "cs_bold",
    "iCs": "cs_italic",
    "caps": "caps",
    "smallCaps": "small_caps",
    "strike": "strike",
    "vanish": "vanish",
    "emboss": "emboss",
    "imprint": "imprint",
    "outline": "outline",
    "shadow": "shadow",
}
# dstrike is intentionally excluded — per ECMA-376 17.7.3 the toggle list is
# the twelve above. dstrike is handled in :func:`_apply_rpr` as a non-toggle
# property (last writer wins) and surfaced on
# :class:`ResolvedFormatting.double_strike`.


Layer = Literal[
    "docDefaults",
    "tableStyle",
    "paragraphStyle",
    "styleNumbering",
    "numbering",
    "directParagraph",
    "runStyle",
    "directRun",
]
_LAYER_ORDER: tuple[Layer, ...] = get_args(Layer)
"""The layers in ascending precedence, derived from :data:`Layer` itself so
the two cannot drift apart."""

# There is deliberately no ``linkedCharStyle`` layer. A paragraph style's
# ``w:link`` partner (``Heading1`` / ``Heading1Char``) is not a cascade layer
# at all: it exists so a user can apply the paragraph style's character half
# to a selection, and Word never consults it when rendering runs inside the
# paragraph. ``Heading1`` carries its own ``<w:b/>``, which is where heading
# bold actually comes from. Measured against Word: a paragraph style whose
# formatting lives *only* on its Char half renders as though the Char half
# did not exist.


def _includes(stop_below: Layer | None, layer: Layer) -> bool:
    """Whether ``layer`` still applies to a walk stopping below ``stop_below``."""
    if stop_below is None:
        return True
    return _LAYER_ORDER.index(layer) < _LAYER_ORDER.index(stop_below)


# ``styleNumbering`` vs ``numbering``: both describe the numbering layer,
# and the distinction is *where the reference came from*. A paragraph's own
# ``w:numPr`` reports ``numbering``; one inherited from its style chain
# reports ``styleNumbering`` and carries the supplying ``style_id``. That
# matters to callers auditing a document — a hand-numbered paragraph and a
# correctly-styled list paragraph are otherwise indistinguishable. The
# formatting the numbering *level* contributes (its indent — the level's
# ``w:rPr`` belongs to the glyph, not the text) is always ``numbering``,
# since its precedence is the same either way.


class StyleCascadeError(DocxPlusError):
    """Raised when the basedOn chain cycles or exceeds Word's depth limit."""


class MissingPartError(DocxPlusError):
    """Raised when a referenced document part is absent.

    Nothing in the cascade resolver raises this today, and
    :func:`resolve_effective_formatting` used to promise it for a
    ``numPr`` whose ``numbering.xml`` is missing. That promise was wrong:
    an unresolvable numbering reference is normal in real documents and
    the resolver degrades instead — see that function's ``Note``.

    Retained as a public symbol because it has been exported since v0.1
    and callers may name it in an ``except`` clause. New code should not
    expect it from the cascade.
    """


@dataclass(frozen=True)
class FormattingSource:
    """Identifies the cascade layer that contributed a resolved property.

    ``layer`` is the cascade layer the value came from. For style layers,
    ``style_id`` names the specific style (the lowest one in the basedOn
    chain that set the value); ``chain_depth`` records how many basedOn hops
    away that style was from the target. ``is_toggle_resolved`` is True when
    the value was computed by the ECMA-376 17.7.3 toggle rule across more
    than one contributing layer, rather than stated by one of them.
    """

    layer: Layer
    style_id: str | None = None
    is_toggle_resolved: bool = False
    chain_depth: int | None = None


@dataclass(frozen=True)
class ResolvedFormatting:
    """The effective formatting for a paragraph, run, or table cell.

    Every field is ``None`` until some layer of the cascade sets it. Toggle
    properties carry the value the ECMA-376 17.7.3 rule produces. SPEC §4
    specifies the fields.

    All twelve ECMA-376 17.7.3 toggle properties are surfaced: the six
    base toggles (``bold``, ``italic``, ``caps``, ``small_caps``,
    ``strike``, ``vanish``) and the six complex-script / decorative
    variants (``cs_bold``, ``cs_italic``, ``emboss``, ``imprint``,
    ``outline``, ``shadow``). All combine by the same rule — see
    :func:`_resolve_toggle`.
    """

    # Identity
    style_id: str | None = None
    style_name: str | None = None

    # Paragraph-level
    alignment: str | None = None
    indent_left: int | None = None
    indent_right: int | None = None
    indent_first_line: int | None = None
    spacing_before: int | None = None
    spacing_after: int | None = None
    line_spacing: float | None = None
    line_spacing_rule: str | None = None

    # ``<w:contextualSpacing>`` as the cascade resolves it. It does not
    # change ``spacing_before`` / ``spacing_after``, which stay the values
    # the cascade declares: whether either is actually *applied* depends on
    # the paragraph's neighbours, which is a layout question rather than a
    # cascade one. :func:`resolve_paragraph_spacing` answers that.
    contextual_spacing: bool | None = None

    keep_with_next: bool | None = None
    keep_lines: bool | None = None
    page_break_before: bool | None = None
    outline_level: int | None = None

    # Run-level
    font_name: str | None = None
    font_size: float | None = None
    bold: bool | None = None
    italic: bool | None = None
    cs_bold: bool | None = None
    cs_italic: bool | None = None
    underline: str | None = None
    strike: bool | None = None
    double_strike: bool | None = None
    color_rgb: str | None = None
    highlight: str | None = None
    caps: bool | None = None
    small_caps: bool | None = None
    vanish: bool | None = None
    emboss: bool | None = None
    imprint: bool | None = None
    outline: bool | None = None
    shadow: bool | None = None
    vert_align: str | None = None

    # The ``w:lang`` Latin-script language tag (``"en-GB"``, ``"fr-FR"``).
    # Invisible on the page and decisive for proofing: a run tagged with the
    # wrong language is skipped or wrongly flagged by every spell checker.
    lang: str | None = None

    # Numbering
    #
    # ``num_id`` resolves through the style chain like every other property:
    # a paragraph's own ``w:numPr`` wins, otherwise the nearest style in the
    # basedOn chain that supplies one. ``0`` is not an ordinary id — it is
    # the ECMA-376 17.9.18 sentinel for *explicitly not numbered*, the only
    # way to opt a paragraph out of numbering its style applies, and is
    # surfaced faithfully rather than flattened to ``None``. So ``None``
    # means "no numbering information anywhere" and ``0`` means "numbering
    # deliberately suppressed"; check ``provenance["num_id"].layer`` to tell
    # a direct reference from a style-supplied one.
    num_id: int | None = None
    num_level: int | None = None

    # Meta
    partial: bool = False
    provenance: dict[str, FormattingSource] | None = None


def resolve_effective_formatting(
    target: Paragraph | Run | _Cell,
    *,
    include_provenance: bool = False,
    table_context: TableContext | None = None,
    stop_below: Layer | None = None,
) -> ResolvedFormatting:
    """Resolve the effective formatting for ``target``.

    Walks the cascade layers in precedence order, returning a fully
    resolved :class:`ResolvedFormatting`. Toggle properties combine per
    ECMA-376 17.7.3 rather than overriding — see
    :func:`_resolve_toggle`. Theme colors are resolved against the
    document's theme part; if the theme is missing or malformed, the result's
    ``partial`` flag is set and unresolved theme names are returned in place
    of hex values.

    When ``target`` is in a table cell, table-style **conditional
    formatting** (``<w:tblStylePr>`` branches: ``firstRow``, ``lastRow``,
    ``firstCol``, ``lastCol``, the four band branches and the four
    corners) is applied on top of the base table style — but only the
    branches the table's ``<w:tblLook>`` asks for. See
    :class:`TableContext` for how that gating and band membership are
    worked out.

    Note:
        Only **run- and paragraph-level** properties are resolved (the
        ``<w:rPr>`` / ``<w:pPr>`` carried by a style's base and its
        ``<w:tblStylePr>`` branches). Cell-, row-, and table-level
        properties (``<w:tcPr>`` cell shading and margins, ``<w:trPr>``
        row heights, ``<w:tblPr>`` table defaults) declared by a table
        style are **not** surfaced on :class:`ResolvedFormatting` — that
        belongs to a separate cell-formatting resolver deferred to v0.3+.

    Args:
        target: A python-docx :class:`~docx.text.paragraph.Paragraph`,
            :class:`~docx.text.run.Run`, or :class:`~docx.table._Cell`.
        include_provenance: If True, populate ``.provenance`` with the cascade
            layer that set each field. Default False.
        table_context: Optional override for the cell's position within
            its table. When ``None`` (default), the resolver derives it
            from the target's parent ``<w:tr>`` / ``<w:tbl>`` chain;
            pass an explicit :class:`TableContext` to query a hypothetical
            position (e.g. "what would the formatting be if this cell
            were in the first row?").
        stop_below: Stop the walk *below* this :data:`Layer`, so the named
            layer and everything above it contribute nothing. ``None``
            (default) walks the whole cascade. See the note below.

    Returns:
        A :class:`ResolvedFormatting` snapshot.

    Raises:
        StyleCascadeError: If the basedOn chain has a cycle or exceeds Word's
            depth limit of 11.
        ValueError: If ``stop_below`` is not one of the :data:`Layer` names.

    Note:
        A paragraph whose ``w:numPr`` references a numbering id that
        cannot be resolved — because ``numbering.xml`` is absent, or the
        ``numId`` is dangling — is **not** an error. ``num_id`` and
        ``num_level`` are still reported; only the formatting the
        numbering level would have contributed is missing. Word behaves
        the same way, and documents in the wild routinely carry a
        ``numPr`` with no matching definition.

    Note:
        Numbering resolves through the **style chain**, not just the
        paragraph's own ``w:numPr``: a paragraph styled ``List Bullet``
        reports the ``num_id`` that style supplies. ``w:numId`` and
        ``w:ilvl`` resolve independently, so a paragraph overriding only
        the level keeps its style's list. A resolved ``num_id`` of ``0``
        is the ECMA-376 sentinel for *explicitly not numbered*, distinct
        from ``None`` for "no numbering information at all"; with
        ``include_provenance``, the ``num_id`` layer is ``"numbering"``
        for a direct reference and ``"styleNumbering"`` for an inherited
        one.

    Note:
        ``stop_below`` answers **"what would this look like without that
        layer?"**, which provenance alone cannot: provenance names the
        layer that *won*, not the value that would have surfaced in its
        absence. Resolving a run with ``stop_below="directRun"`` gives
        exactly what it would render as if its own ``<w:rPr>`` were
        deleted — character style and all — so a caller can tell direct
        formatting that changes nothing from direct formatting that
        overrides the style. That comparison is the basis of every
        consistency rule in :mod:`docx_plus.lint`.

        ``style_id`` and ``style_name`` are identity rather than
        formatting, so they are reported regardless of where the walk
        stops — a caller resolving beneath the paragraph style still needs
        to know which style it excluded.

    Example:
        >>> from docx import Document
        >>> from docx_plus.styles.inspect import resolve_effective_formatting
        >>> doc = Document()
        >>> p = doc.add_paragraph("Hello")
        >>> resolved = resolve_effective_formatting(p)
        >>> resolved.font_size  # e.g. 11.0 from docDefaults
        11.0

        A run whose direct bold merely restates its style:

        >>> run = p.add_run("bold")
        >>> run.bold = True
        >>> resolve_effective_formatting(run).bold
        True
        >>> resolve_effective_formatting(run, stop_below="directRun").bold is None
        True
    """
    # Classify first so a wrong-typed target raises TypeError from
    # _classify_target, before _document_of reaches for ``.part`` and turns
    # it into an AttributeError. _resolve_with_cache classifies again; three
    # isinstance checks are nothing against a full cascade walk.
    _classify_target(target)
    if stop_below is not None and stop_below not in _LAYER_ORDER:
        raise ValueError(f"stop_below must be one of {', '.join(_LAYER_ORDER)}; got {stop_below!r}")
    return _resolve_with_cache(
        _ResolverCache.for_document(_document_of(target)),
        target,
        include_provenance=include_provenance,
        table_context=table_context,
        stop_below=stop_below,
    )


def _resolve_with_cache(
    cache: _ResolverCache,
    target: Paragraph | Run | _Cell,
    *,
    include_provenance: bool = False,
    table_context: TableContext | None = None,
    stop_below: Layer | None = None,
) -> ResolvedFormatting:
    """Resolve ``target`` against an existing cache.

    The whole of :func:`resolve_effective_formatting` apart from building
    the cache and validating ``stop_below``, so a document-wide sweep shares
    one walk implementation rather than a parallel copy that could drift.
    """
    target_kind, target_el = _classify_target(target)

    # ``partial`` is set lazily — only when a theme reference actually fails
    # to resolve (inside _resolve_color / _resolve_font_theme). A missing
    # theme part is not, on its own, an incomplete resolution: a document
    # with no theme refs resolves fully even without a theme (SPEC §4).
    acc = _Accumulator(theme=cache.theme, want_provenance=include_provenance)

    # _classify_target returns the underlying element alongside the kind, so
    # the union-attr access happens once where isinstance has already narrowed
    # the type — no per-branch type: ignore needed here.
    if target_kind == "paragraph":
        ctx = table_context or _derive_table_context_from_element(target_el, cache)
        _apply_paragraph_cascade(acc, cache, target_el, table_context=ctx, stop_below=stop_below)
    elif target_kind == "run":
        paragraph_element = _enclosing_paragraph(target_el)
        ctx = table_context or _derive_table_context_from_element(paragraph_element, cache)
        _apply_paragraph_cascade(
            acc,
            cache,
            paragraph_element,
            run_element=target_el,
            table_context=ctx,
            stop_below=stop_below,
        )
    else:  # cell
        ctx = table_context or _derive_table_context_from_element(target_el, cache)
        _apply_cell_cascade(acc, cache, target_el, table_context=ctx, stop_below=stop_below)

    return acc.freeze()


# --------------------------------------------------------------------------
# Paragraph spacing: the one place the cascade is not the whole answer.
# --------------------------------------------------------------------------

# Elements that carry no content of their own and so cannot separate two
# paragraphs. Word steps straight over them. Anything *not* listed here and
# not a paragraph or content control -- a table, an ``altChunk`` -- does
# separate them, and stops the search.
_SPACING_TRANSPARENT = frozenset(
    {
        "bookmarkStart",
        "bookmarkEnd",
        "commentRangeStart",
        "commentRangeEnd",
        "proofErr",
        "permStart",
        "permEnd",
    }
)

# A content control can nest, and each level costs a recursion. Word's own
# limit is far lower than anything a real document reaches; this only stops
# a malformed part from recursing without end.
_MAX_SDT_NESTING = 32


def _sibling(node: etree._Element, *, forward: bool) -> etree._Element | None:
    return node.getnext() if forward else node.getprevious()


def _edge_paragraph(
    container: etree._Element, *, forward: bool, depth: int = 0
) -> etree._Element | None:
    """The first (or last) paragraph ``container`` holds.

    ``None`` when the content on that edge is something else — a table, say
    — since that is content, and content between two paragraphs stops them
    being adjacent.
    """
    if depth > _MAX_SDT_NESTING:
        return None
    children = list(container)
    if not forward:
        children.reverse()
    for child in children:
        if not isinstance(child.tag, str):  # comment / processing instruction
            continue
        local = etree.QName(child.tag).localname
        if local == "p":
            return child
        if local == "sdt":
            content = child.find(qn("w:sdtContent"))
            if content is None:
                return None
            return _edge_paragraph(content, forward=forward, depth=depth + 1)
        if local in _SPACING_TRANSPARENT:
            continue
        return None
    return None


def _adjacent_paragraph(node: etree._Element, *, forward: bool) -> etree._Element | None:
    """The paragraph immediately before or after ``node``, if there is one.

    Measured against Word:

    * A table between two same-style contextual paragraphs stops the
      suppression; two paragraphs inside one table cell suppress normally.
    * A **content control is transparent**. A ``<w:sdt>`` wrapping the
      neighbour, or sitting between the pair with a paragraph inside it,
      leaves the paragraphs adjacent — so the search descends into
      ``<w:sdtContent>`` and climbs back out of it.
    """
    current = node
    for _ in range(_MAX_SDT_NESTING):
        sibling = _sibling(current, forward=forward)
        while sibling is not None:
            if not isinstance(sibling.tag, str):  # comment / processing instruction
                sibling = _sibling(sibling, forward=forward)
                continue
            local = etree.QName(sibling.tag).localname
            if local == "p":
                return sibling
            if local == "sdt":
                content = sibling.find(qn("w:sdtContent"))
                if content is None:
                    return None
                return _edge_paragraph(content, forward=forward)
            if local not in _SPACING_TRANSPARENT:
                return None
            sibling = _sibling(sibling, forward=forward)
        # Out of siblings. If this is the inside of a content control, the
        # paragraph next door is outside it.
        parent = current.getparent()
        if parent is None or etree.QName(parent.tag).localname != "sdtContent":
            return None
        grandparent = parent.getparent()
        if grandparent is None:
            return None
        current = grandparent
    return None


@dataclass(frozen=True)
class ParagraphSpacing:
    """How much vertical space Word actually puts above and below a paragraph.

    :class:`ResolvedFormatting` answers what the *cascade* declares.
    That is not the whole story for spacing, for two reasons measured
    against Word rather than inferred:

    * ``<w:contextualSpacing>`` makes a paragraph drop its own space
      before/after when the neighbour on that side carries the **same
      ``styleId``**. Only the paragraph's own flag governs its own edges,
      and numbering plays no part — two paragraphs in different lists still
      suppress, and two related-by-``basedOn`` styles do not.
    * Word does not *add* one paragraph's space-after to the next one's
      space-before. It lays down the space-after, then tops it up to the
      space-before if that is larger — so an unsuppressed pair sits
      ``max(after, before)`` apart, not ``after + before``.

    The two interact: the top-up is measured against the **declared**
    space-after even when that space-after was itself suppressed. A
    contextual paragraph with ``after=20pt`` followed by a non-contextual
    one with ``before=30pt`` leaves 10pt, not 30pt.

    ``space_above`` and ``space_below`` fold all of that together, so
    ``space_below`` of one paragraph always equals ``space_above`` of the
    next. Attributes are twips.
    """

    declared_before: int
    declared_after: int
    contextual_spacing: bool
    before_suppressed: bool
    after_suppressed: bool
    space_above: int
    space_below: int


def _gap(
    first: ResolvedFormatting,
    second: ResolvedFormatting,
) -> int:
    """The twips Word leaves between two adjacent paragraphs."""
    same_style = first.style_id == second.style_id
    after = first.spacing_after or 0
    before = second.spacing_before or 0
    contributed = 0 if (same_style and first.contextual_spacing) else after
    top_up = 0 if (same_style and second.contextual_spacing) else max(0, before - after)
    return contributed + top_up


def resolve_paragraph_spacing(paragraph: Paragraph) -> ParagraphSpacing:
    """Resolve the vertical space actually applied around ``paragraph``.

    Args:
        paragraph: A python-docx :class:`~docx.text.paragraph.Paragraph`.

    Returns:
        A :class:`ParagraphSpacing` snapshot, in twips.

    Note:
        A paragraph with no neighbour on a side keeps its declared space
        there — nothing can suppress it. Word's separate rule about
        space-before at the top of a *page* is layout the resolver does not
        model, since it depends on pagination.
    """
    element = paragraph._p
    cache = _ResolverCache.for_document(_document_of(paragraph))
    own = _resolve_with_cache(cache, paragraph)
    declared_before = own.spacing_before or 0
    declared_after = own.spacing_after or 0

    previous_el = _adjacent_paragraph(element, forward=False)
    next_el = _adjacent_paragraph(element, forward=True)
    previous = _resolve_paragraph_element(cache, previous_el)
    following = _resolve_paragraph_element(cache, next_el)

    before_suppressed = bool(
        previous is not None and own.contextual_spacing and previous.style_id == own.style_id
    )
    after_suppressed = bool(
        following is not None and own.contextual_spacing and following.style_id == own.style_id
    )
    return ParagraphSpacing(
        declared_before=declared_before,
        declared_after=declared_after,
        contextual_spacing=bool(own.contextual_spacing),
        before_suppressed=before_suppressed,
        after_suppressed=after_suppressed,
        space_above=declared_before if previous is None else _gap(previous, own),
        space_below=declared_after if following is None else _gap(own, following),
    )


def _resolve_paragraph_element(
    cache: _ResolverCache, element: etree._Element | None
) -> ResolvedFormatting | None:
    """Resolve a bare ``<w:p>`` element, bypassing the python-docx wrapper."""
    if element is None:
        return None
    acc = _Accumulator(theme=cache.theme, want_provenance=False)
    ctx = _derive_table_context_from_element(element, cache)
    _apply_paragraph_cascade(acc, cache, element, table_context=ctx)
    return acc.freeze()


# --------------------------------------------------------------------------
# Accumulator: in-progress resolved state, with optional provenance tracking.
# --------------------------------------------------------------------------


def _resolve_toggle(
    base: tuple[bool, FormattingSource] | None,
    levels: list[tuple[bool, FormattingSource]],
    direct: tuple[bool, FormattingSource] | None,
) -> tuple[bool, FormattingSource] | None:
    """Combine one toggle property's contributions per ECMA-376 17.7.3.

    ``base`` is what ``docDefaults`` stated, ``levels`` what each *style*
    level stated (each already flattened over its own basedOn chain by plain
    override), and ``direct`` what direct formatting stated. ``None`` means
    "not specified there".

    The rule, measured against Word rather than inferred:

    * Direct formatting is **absolute**. ``<w:b/>`` on a run is bold and
      ``<w:b w:val="0"/>`` is not, whatever the styles underneath said. It
      never participates in the XOR.
    * Otherwise the result is the ``docDefaults`` value flipped once for
      every style level whose value **differs from it**. A level restating
      the default is inert, which is why ``<w:b w:val="0"/>`` on a style
      does nothing when nothing is bold to begin with, and why a bare
      ``<w:b/>`` does nothing when the default is already bold.

    Returns ``None`` when no layer mentioned the property at all, which is
    how an unset toggle stays ``None`` rather than becoming ``False``.
    """
    if direct is not None:
        return direct
    if base is None and not levels:
        return None
    base_value = base[0] if base is not None else False
    flips = sum(1 for value, _ in levels if value != base_value)
    effective = base_value != (flips % 2 == 1)

    # Provenance goes to the most specific layer that had an opinion, and
    # records whether more than one did — a caller asking "why is this bold?"
    # needs to know the answer was computed rather than stated.
    winner = levels[-1][1] if levels else base[1]  # type: ignore[index]
    contributors = len(levels) + (1 if base is not None else 0)
    return effective, replace(winner, is_toggle_resolved=contributors > 1)


@dataclass
class _Accumulator:
    """Mutable in-progress state during the cascade walk.

    Non-toggle properties resolve by override as the walk proceeds. Toggle
    properties cannot: their rule needs one value *per style level*, so they
    are collected into per-level buckets and combined in :meth:`freeze`.
    """

    theme: ThemeColors | None
    want_provenance: bool
    values: dict[str, Any] = field(default_factory=dict)
    provenance: dict[str, FormattingSource] = field(default_factory=dict)
    partial: bool = False

    # Toggle collection. ``toggle_levels`` holds one bucket per style level
    # entered via :meth:`toggle_level`, in cascade order; within a bucket the
    # last writer wins, which is what makes a basedOn chain override rather
    # than alternate.
    toggle_base: dict[str, tuple[bool, FormattingSource]] = field(default_factory=dict)
    toggle_levels: list[dict[str, tuple[bool, FormattingSource]]] = field(default_factory=list)
    toggle_direct: dict[str, tuple[bool, FormattingSource]] = field(default_factory=dict)
    _sink: dict[str, tuple[bool, FormattingSource]] | None = None

    def set(self, name: str, value: Any, source: FormattingSource) -> None:
        """Set a non-toggle property, recording provenance if requested."""
        if value is None:
            return
        self.values[name] = value
        if self.want_provenance:
            self.provenance[name] = source

    def toggle(self, name: str, val_attr: str | None, source: FormattingSource) -> None:
        """Record a toggle specification into the sink currently in scope.

        ``val_attr`` is the ``w:val`` attribute on the toggle element, or
        ``None`` if absent (which means "on").
        """
        sink = self.toggle_direct if self._sink is None else self._sink
        sink[name] = (val_attr not in ("0", "false"), source)

    def toggle_base_scope(self) -> _ToggleScope:
        """Direct toggles into the ``docDefaults`` base bucket."""
        return _ToggleScope(self, self.toggle_base)

    def toggle_level_scope(self) -> _ToggleScope:
        """Open a fresh style level for toggles to accumulate into."""
        bucket: dict[str, tuple[bool, FormattingSource]] = {}
        self.toggle_levels.append(bucket)
        return _ToggleScope(self, bucket)

    def freeze(self) -> ResolvedFormatting:
        """Snapshot into an immutable :class:`ResolvedFormatting`."""
        kwargs: dict[str, Any] = dict(self.values)
        for name in _TOGGLE_RPR.values():
            resolved = _resolve_toggle(
                self.toggle_base.get(name),
                [bucket[name] for bucket in self.toggle_levels if name in bucket],
                self.toggle_direct.get(name),
            )
            if resolved is None:
                continue
            kwargs[name], source = resolved
            if self.want_provenance:
                self.provenance[name] = source
        kwargs["partial"] = self.partial
        kwargs["provenance"] = dict(self.provenance) if self.want_provenance else None
        return ResolvedFormatting(**kwargs)


class _ToggleScope:
    """Context manager routing :meth:`_Accumulator.toggle` at one sink.

    Outside any scope, toggles land in the direct-formatting bucket — which
    is where an unscoped ``_apply_rpr`` call is coming from.
    """

    def __init__(self, acc: _Accumulator, sink: dict[str, tuple[bool, FormattingSource]]) -> None:
        self._acc = acc
        self._sink = sink
        self._previous: dict[str, tuple[bool, FormattingSource]] | None = None

    def __enter__(self) -> None:
        self._previous = self._acc._sink
        self._acc._sink = self._sink

    def __exit__(self, *exc_info: object) -> None:
        self._acc._sink = self._previous


# --------------------------------------------------------------------------
# Per-document memo for the inputs a cascade walk re-reads.
# --------------------------------------------------------------------------


@dataclass
class _ResolverCache:
    """Memoizes the document-level lookups a cascade walk repeats.

    Everything held here is immutable for the lifetime of a resolve: the
    styles part, the theme, the numbering part, and the per-``styleId``
    lookups derived from them. :func:`resolve_effective_formatting` builds
    a throwaway instance per call, so single-target cost is unchanged; a
    document-wide sweep builds one and amortises it across every target.

    This exists because the per-call cost is dominated by work that does
    not vary with the target. Profiling 1500 paragraphs put ``load_theme``
    alone at 39% of resolve time — re-parsing ``theme1.xml`` once per
    target — with most of the rest in repeated ``w:style`` xpath lookups.
    """

    doc: Document
    styles_root: etree._Element
    theme: ThemeColors | None
    _styles: dict[str, etree._Element | None] = field(default_factory=dict)
    _chains: dict[str, list[tuple[str, etree._Element]]] = field(default_factory=dict)
    _names: dict[str, str | None] = field(default_factory=dict)
    _abstract_nums: dict[int, etree._Element | None] = field(default_factory=dict)
    _numbering_root: etree._Element | None = None
    _numbering_read: bool = False
    _doc_defaults: tuple[etree._Element | None, etree._Element | None] | None = None

    @classmethod
    def for_document(cls, doc: Document) -> _ResolverCache:
        """Build a cache over ``doc``, parsing the theme once."""
        return cls(doc=doc, styles_root=doc.styles.element, theme=load_theme(doc))

    def style(self, style_id: str) -> etree._Element | None:
        """The ``w:style`` element for ``style_id``, or None if undefined."""
        if style_id not in self._styles:
            matches = xpath(self.styles_root, "./w:style[@w:styleId=$sid]", sid=style_id)
            self._styles[style_id] = matches[0] if matches else None
        return self._styles[style_id]

    def chain(self, leaf_style_id: str) -> list[tuple[str, etree._Element]]:
        """The basedOn chain from ``leaf_style_id``, leaf-first.

        A cycle or over-deep chain raises out of here every time rather
        than being memoized — the failure is a property of the styles
        part, so re-raising costs nothing and keeps the cache holding
        only well-formed results.
        """
        cached = self._chains.get(leaf_style_id)
        if cached is None:
            cached = _collect_style_chain(self, leaf_style_id)
            self._chains[leaf_style_id] = cached
        return cached

    def style_name(self, style_id: str) -> str | None:
        """The ``w:name`` of ``style_id`` as Word displays it."""
        if style_id not in self._names:
            style_el = self.style(style_id)
            name_el = style_el.find(qn("w:name")) if style_el is not None else None
            self._names[style_id] = name_el.get(qn("w:val")) if name_el is not None else None
        return self._names[style_id]

    # No ``linked_style_id`` lookup: the cascade never consults a ``w:link``
    # partner. ``styles.find_unused_styles`` reads the element directly,
    # because collapsing an unused linked *pair* into one finding is a
    # reporting decision rather than a cascade one.

    def doc_defaults(self) -> tuple[etree._Element | None, etree._Element | None]:
        """The ``(rPr, pPr)`` under ``w:docDefaults``, either possibly None."""
        if self._doc_defaults is None:
            defaults = self.styles_root.find(qn("w:docDefaults"))
            if defaults is None:
                self._doc_defaults = (None, None)
            else:
                rpr_default = defaults.find(qn("w:rPrDefault"))
                ppr_default = defaults.find(qn("w:pPrDefault"))
                self._doc_defaults = (
                    rpr_default.find(qn("w:rPr")) if rpr_default is not None else None,
                    ppr_default.find(qn("w:pPr")) if ppr_default is not None else None,
                )
        return self._doc_defaults

    def numbering_root(self) -> etree._Element | None:
        """The ``w:numbering`` root, or None if the part is absent."""
        if not self._numbering_read:
            self._numbering_root = _read_numbering_root(self.doc)
            self._numbering_read = True
        return self._numbering_root

    def abstract_num(self, num_id: int) -> etree._Element | None:
        """The ``w:abstractNum`` a ``w:num`` id resolves to, or None."""
        if num_id not in self._abstract_nums:
            numbering_root = self.numbering_root()
            self._abstract_nums[num_id] = (
                None if numbering_root is None else _resolve_abstract_num(numbering_root, num_id)
            )
        return self._abstract_nums[num_id]


# --------------------------------------------------------------------------
# Cascade entry points.
# --------------------------------------------------------------------------


def _apply_paragraph_cascade(
    acc: _Accumulator,
    cache: _ResolverCache,
    p_element: etree._Element,
    run_element: etree._Element | None = None,
    table_context: TableContext | None = None,
    stop_below: Layer | None = None,
) -> None:
    """Walk layers 1, 3, 4, 5 (and 6 if run_element) for a paragraph target.

    ``stop_below`` drops the named layer and everything above it, so a
    caller can ask what the target would look like without a given layer.
    """
    # Layer 1: docDefaults
    if _includes(stop_below, "docDefaults"):
        _apply_doc_defaults(acc, cache)

    # Layer 2: table style (if inside a table)
    enclosing_tc = _enclosing_cell(p_element)
    if enclosing_tc is not None and _includes(stop_below, "tableStyle"):
        table_element = _enclosing_table(enclosing_tc)
        if table_element is not None:
            _apply_table_style_chain(acc, cache, table_element, table_context=table_context)

    # Layer 3: paragraph style chain. The identity fields are set whatever
    # ``stop_below`` says — they name the style rather than describing
    # formatting, and a caller resolving beneath the paragraph style still
    # needs to know which style that was.
    p_style_id = _paragraph_style_id(p_element)
    if p_style_id is not None:
        acc.set(
            "style_id",
            p_style_id,
            FormattingSource(layer="paragraphStyle", style_id=p_style_id, chain_depth=0),
        )
        style_name = cache.style_name(p_style_id)
        if style_name is not None:
            acc.set(
                "style_name",
                style_name,
                FormattingSource(layer="paragraphStyle", style_id=p_style_id, chain_depth=0),
            )
        if _includes(stop_below, "paragraphStyle"):
            _apply_style_chain(acc, cache, p_style_id, "paragraphStyle")

    # Layer 4: numbering — the paragraph's own w:numPr, or the nearest one
    # its style chain supplies. The latter is how Word's stock List Bullet /
    # List Number styles number anything at all.
    _apply_numbering(acc, cache, p_element, p_style_id, stop_below=stop_below)

    # Layer 5: direct paragraph formatting
    direct_ppr = p_element.find(qn("w:pPr"))
    if direct_ppr is not None and _includes(stop_below, "directParagraph"):
        _apply_ppr(acc, direct_ppr, FormattingSource(layer="directParagraph"))
        # rPr inside pPr (paragraph mark formatting) — affects whole-paragraph runs
        direct_ppr_rpr = direct_ppr.find(qn("w:rPr"))
        if direct_ppr_rpr is not None and run_element is None:
            _apply_rpr(acc, direct_ppr_rpr, FormattingSource(layer="directParagraph"))

    if run_element is not None:
        # The paragraph style's ``w:link`` partner is deliberately not applied
        # here — see the note beside :data:`Layer`.

        # Run-level rStyle reference (character style applied to one run).
        # Per ECMA-376 17.3.2.29 this is a style layer that sits BELOW direct
        # run formatting — direct rPr on the run must override it.
        run_style_id = _run_style_id(run_element)
        if run_style_id is not None and _includes(stop_below, "runStyle"):
            _apply_style_chain(acc, cache, run_style_id, "runStyle")

        # Layer 6: direct run formatting (highest precedence for the run).
        run_rpr = run_element.find(qn("w:rPr"))
        if run_rpr is not None and _includes(stop_below, "directRun"):
            _apply_rpr(acc, run_rpr, FormattingSource(layer="directRun"))


def _apply_cell_cascade(
    acc: _Accumulator,
    cache: _ResolverCache,
    tc_element: etree._Element,
    table_context: TableContext | None = None,
    stop_below: Layer | None = None,
) -> None:
    """Resolve formatting for a table cell — table style chain only, for now.

    Skips the numbering layer entirely (unlike
    :func:`_apply_paragraph_cascade`): cells carry no paragraph-level
    numbering of their own.
    """
    if _includes(stop_below, "docDefaults"):
        _apply_doc_defaults(acc, cache)
    table_element = _enclosing_table(tc_element)
    if table_element is not None and _includes(stop_below, "tableStyle"):
        _apply_table_style_chain(acc, cache, table_element, table_context=table_context)


# --------------------------------------------------------------------------
# Layer helpers.
# --------------------------------------------------------------------------


def _apply_doc_defaults(acc: _Accumulator, cache: _ResolverCache) -> None:
    rpr, ppr = cache.doc_defaults()
    source = FormattingSource(layer="docDefaults")
    if rpr is not None:
        # docDefaults is the toggle rule's *base*, not one of its levels.
        with acc.toggle_base_scope():
            _apply_rpr(acc, rpr, source)
    if ppr is not None:
        _apply_ppr(acc, ppr, source)


def _apply_style_chain(
    acc: _Accumulator,
    cache: _ResolverCache,
    leaf_style_id: str,
    layer: Layer,
) -> None:
    """Walk the basedOn chain and apply each style's pPr/rPr ancestors-first.

    The whole chain is **one** level of the toggle rule. A child re-asserting
    its parent's ``<w:b/>`` overrides rather than cancels — inheritance is not
    a hierarchy boundary.
    """
    chain = cache.chain(leaf_style_id)
    with acc.toggle_level_scope():
        # Apply in reverse: deepest ancestor first so leaf (most specific) wins.
        for depth, (style_id, style_el) in enumerate(reversed(chain)):
            chain_depth = len(chain) - 1 - depth
            source = FormattingSource(layer=layer, style_id=style_id, chain_depth=chain_depth)
            ppr = style_el.find(qn("w:pPr"))
            if ppr is not None:
                _apply_ppr(acc, ppr, source)
            rpr = style_el.find(qn("w:rPr"))
            if rpr is not None:
                _apply_rpr(acc, rpr, source)


def _collect_style_chain(
    cache: _ResolverCache, leaf_style_id: str
) -> list[tuple[str, etree._Element]]:
    """Return [(id, element), ...] from leaf to root, with cycle/depth checks."""
    chain: list[tuple[str, etree._Element]] = []
    visited: set[str] = set()
    current_id: str | None = leaf_style_id
    while current_id is not None:
        if current_id in visited:
            # basedOn is single-valued, so the chain is linear: this path is
            # the real basedOn sequence up to the repeat, never a diamond.
            # A self-cycle (X basedOn X) prints as "X -> X".
            cycle_path = " -> ".join([sid for sid, _ in chain] + [current_id])
            raise StyleCascadeError(f"cycle in basedOn chain: {cycle_path}")
        if len(chain) > _MAX_STYLE_CHAIN_DEPTH:
            chain_ids = " -> ".join(sid for sid, _ in chain)
            raise StyleCascadeError(
                f"basedOn chain exceeds depth {_MAX_STYLE_CHAIN_DEPTH}: {chain_ids}"
            )
        style_el = cache.style(current_id)
        if style_el is None:
            break
        chain.append((current_id, style_el))
        visited.add(current_id)
        based_on = style_el.find(qn("w:basedOn"))
        current_id = based_on.get(qn("w:val")) if based_on is not None else None
    return chain


def _apply_table_style_chain(
    acc: _Accumulator,
    cache: _ResolverCache,
    tbl_element: etree._Element,
    table_context: TableContext | None = None,
) -> None:
    """Apply the table's style chain in spec-correct interleaved order.

    Walks the basedOn chain ONCE, ancestors-first. For each style level
    apply its base ``pPr`` / ``rPr`` then — when a
    :class:`TableContext` is provided — its matching
    ``<w:tblStylePr w:type="...">`` branches in :data:`_TBL_STYLE_PR_ORDER`
    (bands → first/last col → first/last row → corners). This ensures the
    per-level invariant "conditional branches override that level's base"
    holds while still letting a child level's everything (base +
    conditional) override a parent level's everything.
    """
    tbl_pr = tbl_element.find(qn("w:tblPr"))
    if tbl_pr is None:
        return
    tbl_style = tbl_pr.find(qn("w:tblStyle"))
    if tbl_style is None:
        return
    style_id = tbl_style.get(qn("w:val"))
    if style_id is None:
        return

    chain = cache.chain(style_id)
    matching = _matching_conditional_types(table_context) if table_context is not None else set()

    # The chain and its conditional branches are all one toggle level: a
    # firstRow branch re-stating the base style's bold overrides it rather
    # than cancelling it.
    with acc.toggle_level_scope():
        # Ancestors-first: reverse the leaf-to-root chain.
        for depth, (sid, style_el) in enumerate(reversed(chain)):
            chain_depth = len(chain) - 1 - depth
            source = FormattingSource(layer="tableStyle", style_id=sid, chain_depth=chain_depth)

            # 1. Base pPr / rPr for this style level.
            ppr = style_el.find(qn("w:pPr"))
            if ppr is not None:
                _apply_ppr(acc, ppr, source)
            rpr = style_el.find(qn("w:rPr"))
            if rpr is not None:
                _apply_rpr(acc, rpr, source)

            # 2. Matching conditional branches for this style level, in spec order.
            if not matching:
                continue
            branches: dict[str, etree._Element] = {}
            for branch in style_el.findall(qn("w:tblStylePr")):
                type_attr = branch.get(qn("w:type"))
                if type_attr is not None:
                    branches[type_attr] = branch
            for cond_type in _TBL_STYLE_PR_ORDER:
                if cond_type not in matching or cond_type not in branches:
                    continue
                branch = branches[cond_type]
                branch_ppr = branch.find(qn("w:pPr"))
                if branch_ppr is not None:
                    _apply_ppr(acc, branch_ppr, source)
                branch_rpr = branch.find(qn("w:rPr"))
                if branch_rpr is not None:
                    _apply_rpr(acc, branch_rpr, source)


def _matching_conditional_types(ctx: TableContext) -> set[str]:
    """Return the set of ``<w:tblStylePr w:type=...>`` values that apply.

    A positional flag activates its type only when the table's
    ``<w:tblLook>`` also asks for it, so a cell can sit in row 0 and still
    take no ``firstRow`` formatting. A corner needs *both* of its axes
    enabled: with ``firstColumn`` cleared, the top-left cell takes
    ``firstRow``, not ``nwCell``.

    ``wholeTable`` is never returned — see :data:`_TBL_STYLE_PR_ORDER`.
    """
    first_row = ctx.is_first_row and ctx.first_row_enabled
    last_row = ctx.is_last_row and ctx.last_row_enabled
    first_col = ctx.is_first_col and ctx.first_col_enabled
    last_col = ctx.is_last_col and ctx.last_col_enabled

    types: set[str] = set()
    if ctx.is_band_col:
        types.add("band1Vert")
    if ctx.is_band2_col:
        types.add("band2Vert")
    if ctx.is_band_row:
        types.add("band1Horz")
    if ctx.is_band2_row:
        types.add("band2Horz")
    if first_col:
        types.add("firstCol")
    if last_col:
        types.add("lastCol")
    if first_row:
        types.add("firstRow")
    if last_row:
        types.add("lastRow")
    if first_row and first_col:
        types.add("nwCell")
    if first_row and last_col:
        types.add("neCell")
    if last_row and first_col:
        types.add("swCell")
    if last_row and last_col:
        types.add("seCell")
    return types


def _read_tbl_look(tbl: etree._Element) -> dict[str, bool]:
    """Read a table's ``<w:tblLook>`` into positive flags.

    Returns ``firstRow`` / ``lastRow`` / ``firstColumn`` / ``lastColumn``
    / ``hBand`` / ``vBand``, the last two inverted from the element's
    ``noHBand`` / ``noVBand``.

    Three forms exist in the wild and Word honours all of them:

    - The named attributes (Word 2010+). Any attribute present means the
      named form is in use, and the ones left out default to off.
    - The legacy ``w:val`` hex bitmask alone (Word 2007). Measured: Word
      still obeys it, so a document that only carries ``val`` is gated
      exactly as one carrying attributes.
    - No ``<w:tblLook>`` at all, which Word treats as **everything
      enabled** rather than everything off. An unparseable ``val`` falls
      here too — no usable gating information, so gate nothing.
    """
    all_on = dict.fromkeys(
        ("firstRow", "lastRow", "firstColumn", "lastColumn", "hBand", "vBand"), True
    )
    tbl_pr = tbl.find(qn("w:tblPr"))
    look = tbl_pr.find(qn("w:tblLook")) if tbl_pr is not None else None
    if look is None:
        return all_on

    attrs = {name: look.get(qn(f"w:{name}")) for name in _TBL_LOOK_BITS}
    if any(value is not None for value in attrs.values()):
        flags = {name: value in ("1", "true", "on") for name, value in attrs.items()}
    else:
        raw = look.get(qn("w:val"))
        try:
            bits = int(raw, 16) if raw is not None else None
        except ValueError:
            bits = None
        if bits is None:
            return all_on
        flags = {name: bool(bits & bit) for name, bit in _TBL_LOOK_BITS.items()}

    return {
        "firstRow": flags["firstRow"],
        "lastRow": flags["lastRow"],
        "firstColumn": flags["firstColumn"],
        "lastColumn": flags["lastColumn"],
        "hBand": not flags["noHBand"],
        "vBand": not flags["noVBand"],
    }


def _table_style_chain(
    tbl: etree._Element, cache: _ResolverCache
) -> list[tuple[str, etree._Element]]:
    """The basedOn chain of ``tbl``'s table style, leaf-first, or empty."""
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return []
    tbl_style = tbl_pr.find(qn("w:tblStyle"))
    if tbl_style is None:
        return []
    style_id = tbl_style.get(qn("w:val"))
    if style_id is None:
        return []
    return cache.chain(style_id)


def _read_band_size(
    tbl: etree._Element, chain: list[tuple[str, etree._Element]], child_name: str
) -> int:
    """The effective ``<w:tblStyleRowBandSize>`` / ``ColBandSize``, 0 if none.

    The table instance's own ``<w:tblPr>`` wins, then each style in the
    chain leaf-first. **Zero is the default, and zero means no banding at
    all** — measured against Word, which paints no bands for a style that
    declares band branches but no band size. That is not a quirk of
    hand-built files: it is why Word's own table styles all carry an
    explicit ``<w:tblStyleRowBandSize w:val="1"/>``.
    """
    sources = [tbl.find(qn("w:tblPr"))]
    sources += [style.find(qn("w:tblPr")) for _, style in chain]
    for tbl_pr in sources:
        if tbl_pr is None:
            continue
        size_el = tbl_pr.find(qn(child_name))
        if size_el is None:
            continue
        raw = size_el.get(qn("w:val"))
        if raw is None:
            continue
        try:
            n = int(raw)
        except ValueError:
            continue
        return max(n, 0)
    return 0


def _defines_branch(chain: list[tuple[str, etree._Element]], branch_type: str) -> bool:
    """Whether any style in ``chain`` defines a ``<w:tblStylePr>`` of that type."""
    return any(
        branch.get(qn("w:type")) == branch_type
        for _, style in chain
        for branch in style.findall(qn("w:tblStylePr"))
    )


def _band_membership(index: int, offset: int, size: int, enabled: bool) -> tuple[bool, bool]:
    """Return ``(is_band1, is_band2)`` for a row / column index.

    ``offset`` is where the stripe sequence starts — 1 when the leading
    line is taken by a ``firstRow`` / ``firstCol`` conditional, else 0.
    Even stripes are band1.
    """
    if not enabled or size <= 0 or index < offset:
        return False, False
    stripe = (index - offset) // size
    return stripe % 2 == 0, stripe % 2 == 1


def _derive_table_context_from_element(node: etree._Element, cache: _ResolverCache) -> TableContext:
    """Derive a :class:`TableContext` from a body element's table position.

    Walks up from ``node`` to find the enclosing ``<w:tc>``, then derives
    row / column indices, the table's ``<w:tblLook>`` flags, and band
    membership. Returns an empty (all-False) :class:`TableContext` when
    ``node`` is not inside a table — note that suppresses every
    conditional branch, which is what "not in a table" should mean.

    The band sequence starts at row / column 0 unless the leading line is
    already claimed by a ``firstRow`` / ``firstCol`` conditional, which
    needs both the ``tblLook`` flag *and* a branch defined in the style
    chain. Measured: the flag on its own does not shift the stripes.
    """
    if isinstance(node.tag, str) and etree.QName(node.tag).localname == "tc":
        tc: etree._Element | None = node
    else:
        tc = _enclosing_cell(node)
    if tc is None:
        return TableContext()
    tr = tc.getparent()
    if tr is None or tr.tag != qn("w:tr"):
        return TableContext()
    tbl = tr.getparent()
    if tbl is None or tbl.tag != qn("w:tbl"):
        return TableContext()

    rows = [child for child in tbl if child.tag == qn("w:tr")]
    cells = [child for child in tr if child.tag == qn("w:tc")]
    try:
        row_idx = rows.index(tr)
        col_idx = cells.index(tc)
    except ValueError:
        # tr/tc not a direct child of its parent — happens when a <w:sdt>
        # wraps the row's cells. Position is indeterminate; fall back to an
        # empty context (caller may pass an explicit one). See TableContext.
        return TableContext()

    look = _read_tbl_look(tbl)
    chain = _table_style_chain(tbl, cache)

    row_offset = 1 if look["firstRow"] and _defines_branch(chain, "firstRow") else 0
    col_offset = 1 if look["firstColumn"] and _defines_branch(chain, "firstCol") else 0
    is_band_row, is_band2_row = _band_membership(
        row_idx, row_offset, _read_band_size(tbl, chain, "w:tblStyleRowBandSize"), look["hBand"]
    )
    is_band_col, is_band2_col = _band_membership(
        col_idx, col_offset, _read_band_size(tbl, chain, "w:tblStyleColBandSize"), look["vBand"]
    )

    return TableContext(
        is_first_row=row_idx == 0,
        is_last_row=row_idx == len(rows) - 1,
        is_first_col=col_idx == 0,
        is_last_col=col_idx == len(cells) - 1,
        is_band_row=is_band_row,
        is_band_col=is_band_col,
        is_band2_row=is_band2_row,
        is_band2_col=is_band2_col,
        first_row_enabled=look["firstRow"],
        last_row_enabled=look["lastRow"],
        first_col_enabled=look["firstColumn"],
        last_col_enabled=look["lastColumn"],
    )


def _num_pr_values(num_pr: etree._Element) -> tuple[int | None, int | None]:
    """Return ``(numId, ilvl)`` from a ``w:numPr``, each None if absent or malformed.

    Both children are independently optional per ECMA-376 17.3.1.19, so a
    partial ``w:numPr`` is legal and each half resolves separately.
    """
    values: list[int | None] = []
    for tag in ("w:numId", "w:ilvl"):
        value: int | None = None
        child = num_pr.find(qn(tag))
        if child is not None:
            raw = child.get(qn("w:val"))
            if raw is not None:
                try:
                    value = int(raw)
                except ValueError:
                    value = None
        values.append(value)
    return values[0], values[1]


def _style_chain_num_pr(
    cache: _ResolverCache, leaf_style_id: str
) -> tuple[int | None, FormattingSource | None, int | None, FormattingSource | None]:
    """Resolve numId / ilvl from the nearest style in the basedOn chain setting each.

    Returns ``(num_id, num_id_source, ilvl, ilvl_source)``. The chain from
    :func:`_collect_style_chain` runs leaf-first, so the first style to
    supply a value is the most specific one — matching how every other
    property resolves through the chain.
    """
    num_id: int | None = None
    ilvl: int | None = None
    num_id_source: FormattingSource | None = None
    ilvl_source: FormattingSource | None = None

    for depth, (style_id, style_el) in enumerate(cache.chain(leaf_style_id)):
        num_pr = style_el.find(f"./{qn('w:pPr')}/{qn('w:numPr')}")
        if num_pr is None:
            continue
        candidate_num_id, candidate_ilvl = _num_pr_values(num_pr)
        source = FormattingSource(layer="styleNumbering", style_id=style_id, chain_depth=depth)
        if num_id is None and candidate_num_id is not None:
            num_id, num_id_source = candidate_num_id, source
        if ilvl is None and candidate_ilvl is not None:
            ilvl, ilvl_source = candidate_ilvl, source
        if num_id is not None and ilvl is not None:
            break

    return num_id, num_id_source, ilvl, ilvl_source


def _apply_numbering(
    acc: _Accumulator,
    cache: _ResolverCache,
    p_element: etree._Element,
    p_style_id: str | None,
    stop_below: Layer | None = None,
) -> None:
    """Apply the paragraph's effective numbering, direct winning over style-supplied.

    ``w:numId`` and ``w:ilvl`` resolve **independently**, so a paragraph
    carrying only an ``ilvl`` keeps the ``numId`` its style supplies. Both
    children are optional per ECMA-376 17.3.1.19, and the alternative —
    treating ``w:numPr`` as atomic — would mean a partial one silently
    strips the style's list rather than demoting within it.

    The spec does not state merge semantics for a compound property across
    the style / direct boundary, so this was settled against Word 2016: a
    ``List Bullet`` paragraph given a bare ``<w:ilvl w:val="2"/>`` renders
    as a third-level bullet of the style's own list, not as unnumbered
    body text.

    The two halves gate separately on ``stop_below``, which is the point of
    splitting the numbering layer in two: resolving with
    ``stop_below="numbering"`` drops the paragraph's own ``w:numPr`` and
    reports the list its style would have given it — the only way to see a
    direct numbering reference overriding a style-supplied one.
    """
    num_id: int | None = None
    ilvl: int | None = None
    num_id_source: FormattingSource | None = None
    ilvl_source: FormattingSource | None = None

    direct_num_pr = p_element.find(f"./{qn('w:pPr')}/{qn('w:numPr')}")
    if direct_num_pr is not None and _includes(stop_below, "numbering"):
        num_id, ilvl = _num_pr_values(direct_num_pr)
        direct_source = FormattingSource(layer="numbering")
        num_id_source = direct_source if num_id is not None else None
        ilvl_source = direct_source if ilvl is not None else None

    if (
        (num_id is None or ilvl is None)
        and p_style_id is not None
        and _includes(stop_below, "styleNumbering")
    ):
        style_num_id, style_num_id_source, style_ilvl, style_ilvl_source = _style_chain_num_pr(
            cache, p_style_id
        )
        if num_id is None:
            num_id, num_id_source = style_num_id, style_num_id_source
        if ilvl is None:
            ilvl, ilvl_source = style_ilvl, style_ilvl_source

    if num_id is None or num_id_source is None:
        # An ilvl with no numId behind it references nothing.
        return

    effective_ilvl = ilvl if ilvl is not None else 0
    acc.set("num_id", num_id, num_id_source)
    acc.set("num_level", effective_ilvl, ilvl_source or num_id_source)

    if num_id == 0 or not _includes(stop_below, "numbering"):
        # numId 0 is the ECMA-376 17.9.18 "no numbering" sentinel: there is
        # no w:num to resolve, and the point of the reference is to suppress
        # the level formatting a style would otherwise contribute. The level's
        # own formatting sits at the ``numbering`` layer however the reference
        # was reached, so it drops out with that layer too.
        return

    # An unmaterialised numbering part, or a dangling numId, is common when
    # Word hasn't authored the definition yet. Not fatal — the reference is
    # already recorded; only the level's own formatting is missing.
    abstract_num = cache.abstract_num(num_id)
    if abstract_num is None:
        return
    lvl_el = _find_level(abstract_num, effective_ilvl)
    if lvl_el is None:
        return
    # The level's own formatting sits at the numbering layer regardless of
    # how the reference was reached — only the reference is attributed to a
    # style. See the Layer note above.
    level_source = FormattingSource(layer="numbering")
    lvl_ppr = lvl_el.find(qn("w:pPr"))
    if lvl_ppr is not None:
        _apply_ppr(acc, lvl_ppr, level_source)
    # The level's ``w:rPr`` is deliberately *not* applied. It formats the
    # number or bullet glyph, not the paragraph's text: a level carrying
    # ``<w:b/>`` renders a bold bullet in front of unbolded prose. Applying
    # it here reported the glyph's formatting as the run's.


# --------------------------------------------------------------------------
# pPr / rPr property extraction.
# --------------------------------------------------------------------------


def _apply_ppr(acc: _Accumulator, ppr: etree._Element, source: FormattingSource) -> None:
    jc = ppr.find(qn("w:jc"))
    if jc is not None:
        acc.set("alignment", jc.get(qn("w:val")), source)

    ind = ppr.find(qn("w:ind"))
    if ind is not None:
        _apply_indent(acc, ind, source)

    spacing = ppr.find(qn("w:spacing"))
    if spacing is not None:
        _apply_spacing(acc, spacing, source)

    for tag, field_name in (
        ("keepNext", "keep_with_next"),
        ("keepLines", "keep_lines"),
        ("pageBreakBefore", "page_break_before"),
        ("contextualSpacing", "contextual_spacing"),
    ):
        flag_el = ppr.find(qn(f"w:{tag}"))
        if flag_el is not None:
            raw = flag_el.get(qn("w:val"))
            acc.set(field_name, raw not in ("0", "false"), source)

    outline = ppr.find(qn("w:outlineLvl"))
    if outline is not None:
        raw = outline.get(qn("w:val"))
        if raw is not None:
            try:
                acc.set("outline_level", int(raw), source)
            except ValueError:
                pass


def _apply_indent(acc: _Accumulator, ind: etree._Element, source: FormattingSource) -> None:
    left = ind.get(qn("w:left")) or ind.get(qn("w:start"))
    right = ind.get(qn("w:right")) or ind.get(qn("w:end"))
    first_line = ind.get(qn("w:firstLine"))
    hanging = ind.get(qn("w:hanging"))
    if left is not None:
        try:
            acc.set("indent_left", int(left), source)
        except ValueError:
            pass
    if right is not None:
        try:
            acc.set("indent_right", int(right), source)
        except ValueError:
            pass
    if hanging is not None:
        try:
            acc.set("indent_first_line", -int(hanging), source)
        except ValueError:
            pass
    elif first_line is not None:
        try:
            acc.set("indent_first_line", int(first_line), source)
        except ValueError:
            pass


def _apply_spacing(acc: _Accumulator, spacing: etree._Element, source: FormattingSource) -> None:
    before = spacing.get(qn("w:before"))
    after = spacing.get(qn("w:after"))
    line = spacing.get(qn("w:line"))
    line_rule = spacing.get(qn("w:lineRule"))
    if before is not None:
        try:
            acc.set("spacing_before", int(before), source)
        except ValueError:
            pass
    if after is not None:
        try:
            acc.set("spacing_after", int(after), source)
        except ValueError:
            pass
    if line is not None:
        try:
            line_val = int(line)
        except ValueError:
            return
        rule = line_rule or "auto"
        if rule == "auto":
            acc.set("line_spacing", line_val / 240.0, source)
        else:
            acc.set("line_spacing", float(line_val), source)
        acc.set("line_spacing_rule", rule, source)


def _apply_rpr(acc: _Accumulator, rpr: etree._Element, source: FormattingSource) -> None:
    for child in rpr:
        if not isinstance(child.tag, str):
            continue
        local = etree.QName(child.tag).localname
        if local in _TOGGLE_RPR:
            field_name = _TOGGLE_RPR[local]
            acc.toggle(field_name, child.get(qn("w:val")), source)
            continue
        if local == "rFonts":
            # A theme token (w:asciiTheme) resolves against the theme's
            # font scheme; a literal face (w:ascii / w:hAnsi / w:cs) is used
            # verbatim. Theme attributes take precedence — that is what Word
            # writes when a font is theme-bound.
            ascii_theme = child.get(qn("w:asciiTheme"))
            if ascii_theme is not None:
                acc.set("font_name", _resolve_font_theme(ascii_theme, acc), source)
            else:
                literal = (
                    child.get(qn("w:ascii")) or child.get(qn("w:hAnsi")) or child.get(qn("w:cs"))
                )
                if literal is not None:
                    acc.set("font_name", literal, source)
        elif local == "sz":
            raw = child.get(qn("w:val"))
            if raw is not None:
                try:
                    acc.set("font_size", int(raw) / 2.0, source)
                except ValueError:
                    pass
        elif local == "color":
            color_val = _resolve_color(child, acc)
            if color_val is not None:
                acc.set("color_rgb", color_val, source)
        elif local == "u":
            val = child.get(qn("w:val"))
            if val is not None:
                acc.set("underline", val, source)
        elif local == "dstrike":
            # ECMA-376 17.3.2.10: not a toggle (last writer wins).
            val = child.get(qn("w:val"))
            acc.set("double_strike", val != "false" and val != "0", source)
        elif local == "highlight":
            val = child.get(qn("w:val"))
            if val is not None:
                acc.set("highlight", val, source)
        elif local == "vertAlign":
            val = child.get(qn("w:val"))
            if val is not None:
                acc.set("vert_align", val, source)
        elif local == "lang":
            # Only w:val (the Latin-script language) is surfaced. w:eastAsia
            # and w:bidi are separate properties for separate scripts, and
            # collapsing three languages into one field would be a lie about
            # which one a proofing tool will use.
            val = child.get(qn("w:val"))
            if val is not None:
                acc.set("lang", val, source)


def _resolve_color(color_el: etree._Element, acc: _Accumulator) -> str | None:
    """Resolve a ``<w:color>`` element to an uppercase ``RRGGBB`` hex string.

    Handles the two theme transforms ``<w:color>`` can carry —
    ``themeTint`` and ``themeShade`` (ECMA-376 CT_Color). The DrawingML
    ``lumMod`` / ``lumOff`` transforms are not applicable here: the
    ``w:color`` schema cannot carry them (see :mod:`docx_plus.styles.theme`).

    On an unresolvable theme reference the result is flagged
    ``partial`` (SPEC §4). The unresolved name is surfaced as the value
    **only** when the theme part is entirely absent — so callers without a
    theme can still log which color was wanted. When the theme loaded but
    the name is not in its scheme (a typo such as ``"accent7"``, or the
    explicit ``"none"`` sentinel), no value is returned: a bare name would
    land a non-hex string in ``color_rgb`` that the style writers reject.
    """
    theme_name = color_el.get(qn("w:themeColor"))
    if theme_name is not None:
        if theme_name == "none":
            # Explicit "no theme color" — not a resolution failure.
            return None
        tint = color_el.get(qn("w:themeTint"))
        shade = color_el.get(qn("w:themeShade"))
        resolved = resolve_theme_color(acc.theme, theme_name, tint=tint, shade=shade)
        if resolved is not None:
            return resolved
        acc.partial = True
        if acc.theme is None:
            return theme_name
        return None
    val = color_el.get(qn("w:val"))
    if val and val.lower() != "auto":
        return val.upper()
    return None


def _resolve_font_theme(token: str, acc: _Accumulator) -> str:
    """Resolve a ``w:asciiTheme`` font token to its concrete typeface.

    Reads the theme's ``a:fontScheme`` (e.g. ``"minorHAnsi"`` -> ``"Calibri"``).
    When the theme is absent or the token has no scheme entry the token is
    surfaced unchanged and the result is flagged ``partial`` — the same
    contract :func:`_resolve_color` uses, so a ``partial=True`` result
    reliably means "a theme reference did not resolve to a concrete value"
    (SPEC §4).
    """
    resolved = resolve_theme_font(acc.theme, token)
    if resolved is not None:
        return resolved
    acc.partial = True
    return token


# --------------------------------------------------------------------------
# Document traversal / metadata helpers.
# --------------------------------------------------------------------------


def _classify_target(
    target: object,
) -> tuple[Literal["paragraph", "run", "cell"], etree._Element]:
    """Classify ``target`` and return ``(kind, underlying_element)``.

    Returning the element here — where ``isinstance`` has narrowed the type
    — lets the caller avoid a ``type: ignore[union-attr]`` on each of
    ``._p`` / ``._r`` / ``._tc``.
    """
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

    if isinstance(target, Paragraph):
        return "paragraph", target._p
    if isinstance(target, Run):
        return "run", target._r
    if isinstance(target, _Cell):
        return "cell", target._tc
    kind = type(target).__name__
    raise TypeError(f"resolve_effective_formatting expects Paragraph, Run, or _Cell; got {kind}")


def _document_of(target: Paragraph | Run | _Cell) -> Document:
    """Return the owning ``Document`` for a paragraph/run/cell.

    python-docx exposes ``.document`` on the main document part. We cast
    through ``Any`` because the base ``Part`` class in python-docx is not
    typed with that attribute even though concrete subclasses provide it.
    """
    part: Any = target.part
    doc: Document = part.document
    return doc


def _paragraph_style_id(p_element: etree._Element) -> str | None:
    pstyle = p_element.find(f"./{qn('w:pPr')}/{qn('w:pStyle')}")
    if pstyle is None:
        return None
    return pstyle.get(qn("w:val"))


def _run_style_id(r_element: etree._Element) -> str | None:
    rstyle = r_element.find(f"./{qn('w:rPr')}/{qn('w:rStyle')}")
    if rstyle is None:
        return None
    return rstyle.get(qn("w:val"))


def _enclosing_paragraph(r_element: etree._Element) -> etree._Element:
    node: etree._Element | None = r_element
    while node is not None:
        if isinstance(node.tag, str) and etree.QName(node.tag).localname == "p":
            return node
        node = node.getparent()
    raise StyleCascadeError("run is not inside a paragraph")


def _enclosing_cell(p_element: etree._Element) -> etree._Element | None:
    node: etree._Element | None = p_element.getparent()
    while node is not None:
        if isinstance(node.tag, str) and etree.QName(node.tag).localname == "tc":
            return node
        node = node.getparent()
    return None


def _enclosing_table(node: etree._Element) -> etree._Element | None:
    cursor: etree._Element | None = node
    while cursor is not None:
        if isinstance(cursor.tag, str) and etree.QName(cursor.tag).localname == "tbl":
            return cursor
        cursor = cursor.getparent()
    return None


# --------------------------------------------------------------------------
# Numbering helpers.
# --------------------------------------------------------------------------


def _read_numbering_root(doc: Document) -> etree._Element | None:
    """Return the ``w:numbering`` root, or ``None`` if the part is absent.

    Deliberately does **not** go through ``doc.part.numbering_part``.
    That property fabricates a missing part via ``NumberingPart.new()``,
    which is an unimplemented stub in python-docx (1.2.0) and raises a
    bare ``NotImplementedError`` — a crash for any document carrying a
    ``w:numPr`` without a ``numbering.xml``, which LibreOffice, Pandoc,
    and stripped templates all produce. Reading the relationship directly
    keeps the resolver read-only and lets an absent part read as "no
    numbering information", which is what the cascade wants.
    """
    try:
        numbering_part = doc.part.part_related_by(RT.NUMBERING)
    except KeyError:
        return None
    element = getattr(numbering_part, "element", None)
    if isinstance(element, etree._Element):
        return element
    return None


def _resolve_abstract_num(numbering_root: etree._Element, num_id: int) -> etree._Element | None:
    num_matches = xpath(numbering_root, "./w:num[@w:numId=$nid]", nid=str(num_id))
    if not num_matches:
        return None
    num_el = num_matches[0]
    abstract_ref = num_el.find(qn("w:abstractNumId"))
    if abstract_ref is None:
        return None
    abstract_id = abstract_ref.get(qn("w:val"))
    if abstract_id is None:
        return None
    abstract_matches = xpath(
        numbering_root,
        "./w:abstractNum[@w:abstractNumId=$aid]",
        aid=abstract_id,
    )
    return abstract_matches[0] if abstract_matches else None


def _find_level(abstract_num: etree._Element, ilvl: int) -> etree._Element | None:
    matches = xpath(abstract_num, "./w:lvl[@w:ilvl=$lvl]", lvl=str(ilvl))
    return matches[0] if matches else None


__all__ = [
    "FormattingSource",
    "MissingPartError",
    "ResolvedFormatting",
    "StyleCascadeError",
    "TableContext",
    "resolve_effective_formatting",
]
