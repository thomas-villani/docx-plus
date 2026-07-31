"""Document-wide cascade sweep: ``iter_resolved_paragraphs``.

:func:`~docx_plus.styles.resolve_effective_formatting` resolves one target
at a time and rebuilds every document-level lookup on each call — the theme,
the styles part, each ``basedOn`` chain. That is the right shape for asking
about a single paragraph and the wrong one for asking about all of them:
profiling put ``load_theme`` alone at 39% of per-call cost.

This module walks a whole document against one shared cache, in document
order, yielding each paragraph's resolved formatting alongside its runs'.
It is the read half any whole-document analysis needs — the effective
formatting of everything, cheaply enough to then compare targets against
each other.
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import TYPE_CHECKING, cast

from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

from docx_plus.core.ns import qn
from docx_plus.styles.inspect import (
    ResolvedFormatting,
    _resolve_with_cache,
    _ResolverCache,
)

if TYPE_CHECKING:
    from collections.abc import Iterator

    from docx.document import Document
    from docx.oxml.text.run import CT_R
    from docx.table import _Cell


@dataclass(frozen=True)
class ResolvedRun:
    """A run's resolved formatting, with its position in the paragraph.

    Attributes:
        run: The python-docx :class:`~docx.text.run.Run`.
        index: 0-based position within the owning paragraph.
        formatting: The run's fully-resolved formatting.
        baseline: The same run resolved with ``stop_below="directRun"`` —
            what it would render as if its own ``<w:rPr>`` were deleted,
            character style and all. ``None`` unless the sweep was run with
            ``include_baseline=True``.
    """

    run: Run
    index: int
    formatting: ResolvedFormatting
    baseline: ResolvedFormatting | None = None


@dataclass(frozen=True)
class ResolvedParagraph:
    """A paragraph's resolved formatting, its runs', and where it sits.

    Attributes:
        paragraph: The python-docx :class:`~docx.text.paragraph.Paragraph`.
        index: 0-based position in the sweep's own ordering — see the note
            on :func:`iter_resolved_paragraphs` about how this relates to
            ``doc.paragraphs``.
        formatting: The paragraph's fully-resolved formatting. Run-level
            properties here reflect the paragraph mark, not any one run.
        runs: One :class:`ResolvedRun` per run, in order, **including runs
            inside a ``<w:hyperlink>``** — which ``Paragraph.runs`` omits.
            Together they cover exactly the text :attr:`text` reports. Empty
            when the sweep was run with ``include_runs=False``, and for an
            empty paragraph.
        table_depth: 0 for a body-level paragraph, 1 inside a table, 2
            inside a table nested in a table, and so on.
        baseline: The same paragraph resolved with
            ``stop_below="directParagraph"``, which is what makes "this
            direct override deviates from the style" answerable. ``None``
            unless the sweep was run with ``include_baseline=True``.

            Note this is *not* the same as deleting the paragraph's
            ``<w:pPr>``: the ``numbering`` layer sits **below**
            ``directParagraph``, so a direct ``<w:numPr>`` still supplies
            ``num_id`` and the indents it implies. The baseline excludes
            the direct paragraph layer, not the whole element.
    """

    paragraph: Paragraph
    index: int
    formatting: ResolvedFormatting
    runs: tuple[ResolvedRun, ...]
    table_depth: int
    baseline: ResolvedFormatting | None = None

    @property
    def in_table(self) -> bool:
        """True if this paragraph sits inside a table cell."""
        return self.table_depth > 0

    @property
    def text(self) -> str:
        """The paragraph's text, for convenience when scanning content.

        Covers exactly the runs in :attr:`runs` — both include the inside of
        a ``<w:hyperlink>`` and both exclude ``<w:ins>`` / ``<w:del>`` /
        ``<w:sdt>``, matching python-docx's own ``Paragraph.text``. A rule
        may therefore index into this string and expect a run to answer for
        every offset.
        """
        return self.paragraph.text


def iter_resolved_paragraphs(
    doc: Document,
    *,
    include_provenance: bool = False,
    include_runs: bool = True,
    include_tables: bool = True,
    include_baseline: bool = False,
) -> Iterator[ResolvedParagraph]:
    """Resolve every paragraph in ``doc``, sharing one cascade cache.

    Yields lazily in document order, so a caller can stop early or stream a
    large document without materialising every result. Wrap in ``list()``
    for the whole set.

    Equivalent to calling
    :func:`~docx_plus.styles.resolve_effective_formatting` on each paragraph
    and run — the same walk over the same code path — but with the theme,
    the styles part, and every ``basedOn`` chain resolved once for the whole
    document rather than once per target.

    Args:
        doc: The python-docx :class:`~docx.document.Document` to sweep.
        include_provenance: Populate each result's ``.provenance``. Costs
            extra work per target, so it is off by default.
        include_runs: Resolve each paragraph's runs as well. Set False when
            only paragraph-level properties matter — runs are usually the
            bulk of the work in a text-heavy document.
        include_tables: Descend into table cells (including nested tables).
            Set False to sweep body-level paragraphs only.
        include_baseline: Also resolve each target with its own direct
            formatting layer excluded, populating ``.baseline``. Roughly
            doubles the resolve work, so it is off by default; turn it on
            for the "is this direct override doing anything?" question.

    Yields:
        One :class:`ResolvedParagraph` per paragraph, in document order.

    Raises:
        StyleCascadeError: If a ``basedOn`` chain has a cycle or exceeds
            Word's depth limit of 11. Raised from the first paragraph that
            reaches the bad chain, so a partial sweep may already have been
            yielded.

    Note:
        ``index`` counts the paragraphs **this sweep yields**, in document
        order. With ``include_tables=True`` (the default) that includes
        table-cell paragraphs, which ``doc.paragraphs`` omits — so the two
        indexings diverge at the first table. Pass ``include_tables=False``
        to get indices that line up with ``doc.paragraphs``, and note the
        CLI's ``inspect`` command numbers from 1 rather than 0.

    Note:
        Headers, footers, footnotes, endnotes, and comments are **not**
        swept — only the main document body. Those live in separate parts
        with their own style references.

    Example:
        >>> from docx import Document
        >>> from docx_plus.styles import iter_resolved_paragraphs
        >>> doc = Document()
        >>> _ = doc.add_paragraph("Hello")
        >>> for resolved in iter_resolved_paragraphs(doc):
        ...     print(resolved.index, resolved.formatting.font_size)
        0 11.0
    """
    cache = _ResolverCache.for_document(doc)
    counter = 0

    for paragraph, depth in _walk(doc, include_tables=include_tables):
        runs: tuple[ResolvedRun, ...] = ()
        if include_runs:
            runs = tuple(
                ResolvedRun(
                    run=run,
                    index=run_index,
                    formatting=_resolve_with_cache(
                        cache, run, include_provenance=include_provenance
                    ),
                    baseline=(
                        _resolve_with_cache(cache, run, stop_below="directRun")
                        if include_baseline
                        else None
                    ),
                )
                for run_index, run in enumerate(_iter_runs(paragraph))
            )
        yield ResolvedParagraph(
            paragraph=paragraph,
            index=counter,
            formatting=_resolve_with_cache(cache, paragraph, include_provenance=include_provenance),
            runs=runs,
            table_depth=depth,
            baseline=(
                _resolve_with_cache(cache, paragraph, stop_below="directParagraph")
                if include_baseline
                else None
            ),
        )
        counter += 1


def _iter_runs(paragraph: Paragraph) -> Iterator[Run]:
    """Yield the paragraph's runs in document order, hyperlinks included.

    ``Paragraph.runs`` is direct ``<w:r>`` children only, so it misses every
    run inside a ``<w:hyperlink>`` — while ``Paragraph.text`` *includes* that
    text, because ``CT_P.text`` walks ``w:r | w:hyperlink``. Taking the
    narrower list left :attr:`ResolvedParagraph.text` reporting text whose
    formatting had never been resolved, and made every run-level lint rule
    blind to the inside of a link. This matches the two back up.

    Deliberately built on the XML rather than ``iter_inner_content()`` /
    ``Hyperlink``, which python-docx only grew in 1.1.0 — the supported
    floor is 1.0.0.

    Runs inside ``<w:ins>``, ``<w:del>``, and ``<w:sdt>`` stay out, because
    ``Paragraph.text`` excludes them too; ``text`` and ``runs`` covering the
    same content is the invariant worth holding.
    """
    for child in paragraph._p:
        if child.tag == qn("w:r"):
            yield Run(cast("CT_R", child), paragraph)
        elif child.tag == qn("w:hyperlink"):
            for sub_child in child:
                if sub_child.tag == qn("w:r"):
                    yield Run(cast("CT_R", sub_child), paragraph)


def _walk(
    container: Document | _Cell,
    *,
    include_tables: bool,
    depth: int = 0,
) -> Iterator[tuple[Paragraph, int]]:
    """Yield ``(paragraph, table_depth)`` in document order, tables inline.

    Uses python-docx's ``iter_inner_content``, which is what preserves the
    interleaving of paragraphs and tables; ``doc.paragraphs`` drops the
    tables and ``doc.tables`` drops the ordering.
    """
    for item in container.iter_inner_content():
        if isinstance(item, Paragraph):
            yield item, depth
        elif include_tables and isinstance(item, Table):
            # ``row.cells`` returns a merged cell once per grid position it
            # spans — the same ``w:tc`` object each time — so a spanned cell's
            # paragraphs would otherwise be yielded once per covered column.
            seen: set[etree._Element] = set()
            for row in item.rows:
                for cell in row.cells:
                    if cell._tc in seen:
                        continue
                    seen.add(cell._tc)
                    yield from _walk(cell, include_tables=True, depth=depth + 1)


__all__ = [
    "ResolvedParagraph",
    "ResolvedRun",
    "iter_resolved_paragraphs",
]
