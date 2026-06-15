"""Author tracked insertions and deletions — wrap existing runs.

python-docx writes runs but cannot mark them as tracked changes. This
module wraps run(s) already present in the document in the revision
container OOXML python-docx skips:

- :func:`mark_insertion` wraps the target run(s) in ``<w:ins>`` — the text
  is shown as an inserted revision; accepting it makes the text permanent.
- :func:`mark_deletion` wraps the target run(s) in ``<w:del>`` and retags
  each ``<w:t>`` to ``<w:delText>`` — the text is shown struck-through;
  accepting the deletion removes it.

Both target a single run, a whole paragraph, or a ``(start_run, end_run)``
range — the same target shapes as ``comments.add_comment``. A range must
lie within one paragraph: ``w:ins`` / ``w:del`` are run-level containers
and cannot span a paragraph boundary.

Runs wrapped in ``w:ins`` / ``w:del`` are *not* visible through
python-docx's ``paragraph.runs`` (its ``CT_P`` descriptor only sees direct
``w:r`` children); read them back with
:func:`~docx_plus.revisions.read_revisions`.

This module imports only from ``docx_plus.core`` and the sibling
``docx_plus.revisions.registry`` (SPEC §9.1).
"""

from __future__ import annotations

import datetime as dt
from dataclasses import dataclass
from typing import TYPE_CHECKING

from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

from docx_plus.core import DocxPlusError
from docx_plus.core.ns import qn
from docx_plus.core.oxml import body_document_for, el, remove, xpath
from docx_plus.revisions.registry import RevisionIdRegistry

if TYPE_CHECKING:
    from docx.document import Document


RevisionTarget = Run | Paragraph | tuple[Run, Run]


class RevisionNotFoundError(DocxPlusError, KeyError):
    """Raised when no revision element with the requested ``w:id`` exists.

    Subclasses :class:`KeyError` so existing ``except KeyError:`` clauses
    still catch it; also :class:`DocxPlusError` per SPEC §9.7.
    """


@dataclass(frozen=True)
class RevisionRef:
    """Handle for an authored revision.

    Attributes:
        revision_id: The ``w:id`` assigned to the wrapping element.
        body_element: The ``<w:ins>`` or ``<w:del>`` lxml element placed in
            the document body. Mutate it directly for advanced edits the
            v0.3 surface does not yet expose.
    """

    revision_id: int
    body_element: etree._Element


def mark_insertion(
    target: RevisionTarget,
    *,
    author: str = "",
    date: dt.datetime | None = None,
    id_registry: RevisionIdRegistry | None = None,
) -> RevisionRef:
    """Mark existing run(s) as a tracked insertion.

    Wraps the target run(s) in ``<w:ins w:id w:author w:date>``. The run
    content (including each ``<w:t>``) is preserved unchanged; only the
    wrapping container is added.

    Args:
        target: Where the insertion applies.

            - A :class:`~docx.text.run.Run` wraps that one run.
            - A :class:`~docx.text.paragraph.Paragraph` wraps the contiguous
              span from its first run to its last run; the paragraph must
              contain at least one run.
            - A ``(start_run, end_run)`` tuple wraps the contiguous span
              from ``start_run`` to ``end_run`` inclusive. Both runs must
              share one paragraph and ``start_run`` must not appear after
              ``end_run``.
        author: Author recorded in ``w:author``. The empty string is legal.
        date: Timestamp recorded in ``w:date``. ``None`` uses the current
            UTC time (millisecond precision). A supplied :class:`datetime`
            is formatted the same way.
        id_registry: Pre-existing registry to share across an editing
            session. A fresh :class:`RevisionIdRegistry` is built from the
            target's document if not supplied.

    Returns:
        A :class:`RevisionRef` with the assigned id and the ``<w:ins>``
        element.

    Raises:
        ValueError: If a paragraph target has no runs, or a range is
            reversed or spans more than one paragraph.
        TypeError: If ``target`` is not a Run, Paragraph, or ``(Run, Run)``
            tuple.

    Example:
        >>> from docx import Document
        >>> from docx_plus.revisions import mark_insertion
        >>> doc = Document()
        >>> p = doc.add_paragraph()
        >>> run = p.add_run("freshly added")
        >>> ref = mark_insertion(run, author="Reviewer")
    """
    return _wrap_revision(target, "w:ins", author=author, date=date, id_registry=id_registry)


def mark_deletion(
    target: RevisionTarget,
    *,
    author: str = "",
    date: dt.datetime | None = None,
    id_registry: RevisionIdRegistry | None = None,
) -> RevisionRef:
    """Mark existing run(s) as a tracked deletion.

    Wraps the target run(s) in ``<w:del w:id w:author w:date>`` and retags
    every ``<w:t>`` in the span to ``<w:delText>`` (the OOXML element for
    deleted run text). Word renders the span struck-through; accepting the
    deletion removes it, rejecting it restores live ``<w:t>``.

    Args:
        target: Where the deletion applies — same shapes as
            :func:`mark_insertion`.
        author: Author recorded in ``w:author``. The empty string is legal.
        date: Timestamp recorded in ``w:date``. ``None`` uses the current
            UTC time (millisecond precision).
        id_registry: Pre-existing registry to share across an editing
            session. A fresh :class:`RevisionIdRegistry` is built from the
            target's document if not supplied.

    Returns:
        A :class:`RevisionRef` with the assigned id and the ``<w:del>``
        element.

    Raises:
        ValueError: If a paragraph target has no runs, or a range is
            reversed or spans more than one paragraph.
        TypeError: If ``target`` is not a Run, Paragraph, or ``(Run, Run)``
            tuple.
    """
    return _wrap_revision(target, "w:del", author=author, date=date, id_registry=id_registry)


# ---------------------------------------------------------------------------
# Internals.
# ---------------------------------------------------------------------------


def _wrap_revision(
    target: RevisionTarget,
    wrapper_tag: str,
    *,
    author: str,
    date: dt.datetime | None,
    id_registry: RevisionIdRegistry | None,
) -> RevisionRef:
    """Wrap the target's run span in ``wrapper_tag`` (``w:ins`` or ``w:del``)."""
    span, doc = _normalize_target(target)

    if id_registry is None:
        id_registry = RevisionIdRegistry(doc)
    revision_id = id_registry.next()

    wrapper = el(
        wrapper_tag,
        **{"w:id": str(revision_id), "w:author": author, "w:date": _format_date(date)},
    )
    # Insert the wrapper where the span starts, then move each element of the
    # span into it. ``append`` on an already-parented node relocates it, so
    # document order is preserved.
    span[0].addprevious(wrapper)
    for elem in span:
        wrapper.append(elem)

    if wrapper_tag == "w:del":
        _retag_text_to_deltext(wrapper)

    return RevisionRef(revision_id=revision_id, body_element=wrapper)


def _normalize_target(target: RevisionTarget) -> tuple[list[etree._Element], Document]:
    """Resolve ``target`` to ``(span, doc)``.

    ``span`` is the contiguous list of sibling elements (runs and any
    interleaved inline content) to move into the revision wrapper, in
    document order.
    """
    if isinstance(target, Run):
        doc = body_document_for(target, operation="mark revision")
        return [target._r], doc

    if isinstance(target, Paragraph):
        runs = [child for child in target._p if child.tag == qn("w:r")]
        if not runs:
            raise ValueError("mark revision requires a paragraph with at least one run")
        doc = body_document_for(target, operation="mark revision")
        return _sibling_span(runs[0], runs[-1]), doc

    if isinstance(target, tuple) and len(target) == 2:
        first, second = target
        if not (isinstance(first, Run) and isinstance(second, Run)):
            raise TypeError("range target must be a tuple of (Run, Run)")
        doc = body_document_for(first, operation="mark revision")
        return _sibling_span(first._r, second._r), doc

    raise TypeError(
        f"mark revision target must be Run, Paragraph, or (Run, Run); "
        f"got {type(target).__name__}"
    )


def _sibling_span(first: etree._Element, last: etree._Element) -> list[etree._Element]:
    """Return the contiguous children from ``first`` to ``last`` inclusive.

    Both elements must share a parent (a single paragraph — ``w:ins`` /
    ``w:del`` cannot cross a paragraph boundary) and ``first`` must not
    appear after ``last`` in document order.
    """
    if first is last:
        return [first]

    parent = first.getparent()
    if parent is None or last.getparent() is not parent:
        raise ValueError(
            "revision range must lie within a single paragraph "
            "(w:ins/w:del cannot span paragraphs)"
        )

    children = list(parent)
    start = children.index(first)
    stop = children.index(last)
    if stop < start:
        raise ValueError("revision range is reversed: start run appears after end run")
    return children[start : stop + 1]


def _retag_text_to_deltext(scope: etree._Element) -> None:
    """Retag every ``<w:t>`` under ``scope`` to ``<w:delText>`` in place.

    Deleted run text uses ``<w:delText>`` rather than ``<w:t>`` (ECMA-376
    17.3.3.7). lxml cannot rename a tag, so each ``<w:t>`` is replaced by a
    fresh ``<w:delText>`` carrying the same text and ``xml:space``.
    """
    for t in xpath(scope, ".//w:t"):
        new = el("w:delText")
        space = t.get(qn("xml:space"))
        if space is not None:
            new.set(qn("xml:space"), space)
        new.text = t.text
        t.addprevious(new)
        remove(t)


def _format_date(date: dt.datetime | None) -> str:
    """Format ``date`` as ``xsd:dateTime`` (UTC, ms precision, trailing ``Z``).

    ``None`` uses the current time. A naive datetime is assumed UTC; an
    aware one is converted to UTC. Millisecond precision keeps two marks
    authored in the same wall-clock second from colliding on ``w:date``.
    """
    if date is None:
        date = dt.datetime.now(dt.timezone.utc)
    elif date.tzinfo is None:
        date = date.replace(tzinfo=dt.timezone.utc)
    else:
        date = date.astimezone(dt.timezone.utc)
    return date.isoformat(timespec="milliseconds").replace("+00:00", "Z")


__all__ = [
    "RevisionNotFoundError",
    "RevisionRef",
    "RevisionTarget",
    "mark_deletion",
    "mark_insertion",
]
