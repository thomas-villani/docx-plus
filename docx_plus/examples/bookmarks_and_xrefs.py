"""Bookmark + cross-reference example.

Anchors a bookmark to a heading, then inserts a ``REF`` field that
resolves to the heading's text and a ``PAGEREF`` field for its page
number. Pairs the writes with :func:`mark_fields_dirty` so Word
recalculates on open.

Usage::

    python -m docx_plus.examples.bookmarks_and_xrefs              # writes ./bookmarks.docx
    python -m docx_plus.examples.bookmarks_and_xrefs path/out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from docx_plus.bookmarks import (
    add_bookmark,
    add_cross_reference,
    read_bookmarks,
)
from docx_plus.fields import mark_fields_dirty


def build_bookmark_document(out_path: Path) -> Path:
    """Build a small doc with one bookmark and two cross-references."""
    doc = Document()

    heading = doc.add_heading("Introduction", level=1)
    # Bookmark wraps the heading's entire content.
    add_bookmark(heading, "intro_section")

    doc.add_paragraph(
        "This is the introductory section. The following paragraph contains "
        "cross-references back to this heading by both text and page number."
    )

    ref_para = doc.add_paragraph("See ")
    add_cross_reference(ref_para, bookmark="intro_section", kind="text")
    ref_para.add_run(" on page ")
    add_cross_reference(ref_para, bookmark="intro_section", kind="page")
    ref_para.add_run(".")

    # Mark fields dirty so Word resolves the REF and PAGEREF on open.
    mark_fields_dirty(doc)

    doc.save(str(out_path))
    return out_path


def main(argv: list[str] | None = None) -> int:
    """Entry point. One optional positional arg: the output docx path."""
    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) > 1:
        print(
            "usage: python -m docx_plus.examples.bookmarks_and_xrefs [output.docx]",
            file=sys.stderr,
        )
        return 2

    out_path = (
        Path(args[0]).expanduser().resolve() if args else Path.cwd() / "bookmarks.docx"
    )
    written = build_bookmark_document(out_path)
    print(f"# wrote: {written}")
    reopened = Document(str(written))
    bookmarks = read_bookmarks(reopened)
    print(f"# bookmarks: {[b.name for b in bookmarks]}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
