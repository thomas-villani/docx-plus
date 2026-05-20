"""Publishing-layout example: TOC + captioned figures + Table of Figures.

Builds a small document with headings, a Table of Contents at the top,
three figure captions (which would normally sit beneath images), and a
Table of Figures at the end. Marks fields dirty so Word populates the
TOC and ToF when the file is opened.

Usage::

    python -m docx_plus.examples.publishing_layout              # writes ./publishing.docx
    python -m docx_plus.examples.publishing_layout path/out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from docx_plus.fields import mark_fields_dirty
from docx_plus.publishing import add_caption, add_table_of_figures, add_toc


def build_publishing_document(out_path: Path) -> Path:
    """Build a doc with a TOC, three captioned figures, and a ToF."""
    doc = Document()

    # ----- Table of Contents at the top -----
    doc.add_heading("Contents", level=1)
    add_toc(doc.add_paragraph(), levels=(1, 2))

    # ----- Body with headings + captioned figures -----
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Opening remarks.")

    doc.add_heading("Architecture", level=1)
    doc.add_paragraph("High-level diagram below.")
    cap1 = doc.add_paragraph()
    add_caption(cap1, caption_type="Figure")
    cap1.add_run(": System architecture overview.")

    doc.add_heading("Data flow", level=2)
    doc.add_paragraph("Step-by-step trace.")
    cap2 = doc.add_paragraph()
    add_caption(cap2, caption_type="Figure")
    cap2.add_run(": Request lifecycle.")

    doc.add_heading("Deployment topology", level=2)
    doc.add_paragraph("Production layout.")
    cap3 = doc.add_paragraph()
    add_caption(cap3, caption_type="Figure")
    cap3.add_run(": Regional rollout map.")

    # ----- Table of Figures at the end -----
    doc.add_heading("List of Figures", level=1)
    add_table_of_figures(doc.add_paragraph(), caption_type="Figure")

    # Mark fields dirty so Word recalculates TOC, SEQ, and ToF on open.
    mark_fields_dirty(doc)

    doc.save(str(out_path))
    return out_path


def main(argv: list[str] | None = None) -> int:
    """Entry point. One optional positional arg: the output docx path."""
    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) > 1:
        print(
            "usage: python -m docx_plus.examples.publishing_layout [output.docx]",
            file=sys.stderr,
        )
        return 2

    out_path = Path(args[0]).expanduser().resolve() if args else Path.cwd() / "publishing.docx"
    written = build_publishing_document(out_path)
    print(f"# wrote: {written}")
    reopened = Document(str(written))
    print(f"# paragraphs: {len(reopened.paragraphs)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
