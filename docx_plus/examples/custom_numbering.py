"""Build a document with custom numbered and bulleted lists.

Demonstrates the v0.5 numbering surface: :func:`define_numbered_list` and
:func:`define_bullet_list` for Word's own defaults,
:func:`define_list_definition` for a legal-outline definition Word's UI
does not offer in one click, :func:`apply_list` to place paragraphs, and
:func:`restart_list` to begin a second sequence over the same definition.

Open the result in Word: the numbered procedure runs 1, 2, 3; the
appendix restarts at 1 with identical formatting; the outline reads
1., 1.1., 1.1.1.; and the bullets show Word's round / hollow / square
cycle by depth.

Usage::

    python -m docx_plus.examples.custom_numbering              # ./numbering.docx
    python -m docx_plus.examples.custom_numbering path/out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from docx_plus.numbering import (
    AbstractNumIdRegistry,
    LevelDefinition,
    NumIdRegistry,
    apply_list,
    define_bullet_list,
    define_list_definition,
    define_numbered_list,
    read_list_definitions,
    restart_list,
)


def build_numbered_document(out_path: Path) -> Path:
    """Build a document exercising each list shape."""
    doc = Document()
    doc.add_heading("Deployment Runbook", level=1)

    # One allocator pair shared across every definition, so the ids stay
    # dense and no two definitions collide.
    nums = NumIdRegistry(doc)
    abstracts = AbstractNumIdRegistry(doc)

    # 1. A plain numbered procedure.
    doc.add_heading("Procedure", level=2)
    steps = define_numbered_list(doc, num_registry=nums, abstract_registry=abstracts)
    for text in ("Drain the connection pool.", "Deploy the new image.", "Re-enable traffic."):
        apply_list(doc.add_paragraph(text), steps)

    # 2. A second sequence over the same definition. Restarting is not a
    #    paragraph property in OOXML -- Word adds another w:num pointing at
    #    the same w:abstractNum, which is what restart_list does.
    doc.add_heading("Rollback", level=2)
    first = doc.add_paragraph("Stop the rollout.")
    rollback = restart_list(first, steps, num_registry=nums)
    apply_list(doc.add_paragraph("Redeploy the previous image."), rollback)

    # 3. A legal outline: each level accumulates the ones above it.
    #    The hanging indent has to grow with the number: it is the width
    #    reserved for "1.1.1.", and if the number overflows it the tab
    #    that separates number from text collapses to nothing.
    doc.add_heading("Sign-off", level=2)
    outline = define_list_definition(
        doc,
        levels=[
            LevelDefinition(text="%1.", indent=720, hanging=360),
            LevelDefinition(text="%1.%2.", indent=1584, hanging=504),
            LevelDefinition(text="%1.%2.%3.", indent=2448, hanging=792),
        ],
        name="Sign-off outline",
        num_registry=nums,
        abstract_registry=abstracts,
    )
    apply_list(doc.add_paragraph("Engineering"), outline)
    apply_list(doc.add_paragraph("Service owner"), outline, level=1)
    apply_list(doc.add_paragraph("On-call lead"), outline, level=2)
    apply_list(doc.add_paragraph("Security"), outline)

    # 4. Bullets, using Word's glyph cycle by depth.
    doc.add_heading("Notes", level=2)
    bullets = define_bullet_list(doc, levels=3, num_registry=nums, abstract_registry=abstracts)
    apply_list(doc.add_paragraph("Window is 30 minutes."), bullets)
    apply_list(doc.add_paragraph("Extendable once."), bullets, level=1)
    apply_list(doc.add_paragraph("Ask the release manager."), bullets, level=2)

    doc.save(str(out_path))
    return out_path


def main() -> None:
    """Write the document, then read its definitions back."""
    out_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("numbering.docx")
    build_numbered_document(out_path)
    print(f"wrote {out_path}")

    reopened = Document(str(out_path))
    definitions = read_list_definitions(reopened)
    print(f"\n{len(definitions)} list definitions in the saved file")
    print("(the first nine ship in python-docx's default template)\n")

    for definition in definitions:
        if definition.num_id < 10:
            continue
        label = definition.name or "(unnamed)"
        restarted = " restarted" if definition.start_overrides else ""
        formats = ", ".join(level.fmt or "?" for level in definition.levels)
        print(f"  numId {definition.num_id}: {label}{restarted} -- [{formats}]")


if __name__ == "__main__":
    main()
