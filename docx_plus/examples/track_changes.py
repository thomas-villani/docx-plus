"""Author, read, and resolve tracked changes.

Demonstrates :func:`docx_plus.revisions.enable_track_changes`,
:func:`mark_insertion` / :func:`mark_deletion` against the three target
shapes (a single run, a whole paragraph, a run-range), then
:func:`read_revisions` round-tripping and :func:`accept_all_revisions`.

Usage::

    python -m docx_plus.examples.track_changes              # writes ./tracked.docx
    python -m docx_plus.examples.track_changes path/out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from docx_plus.revisions import (
    RevisionIdRegistry,
    accept_all_revisions,
    enable_track_changes,
    mark_deletion,
    mark_insertion,
    read_revisions,
)


def build_tracked_document(out_path: Path) -> Path:
    """Build a document with a few tracked insertions and deletions."""
    doc = Document()
    doc.add_heading("Draft with revisions", level=1)

    # Turn on track-changes so Word keeps tracking edits the reader makes.
    enable_track_changes(doc)

    p1 = doc.add_paragraph()
    p1.add_run("This sentence keeps ")
    inserted = p1.add_run("a freshly inserted phrase ")
    p1.add_run("in place.")

    p2 = doc.add_paragraph()
    p2.add_run("Here a ")
    start = p2.add_run("multi-run ")
    p2.add_run("range is ")
    end = p2.add_run("struck out.")

    p3 = doc.add_paragraph("This whole paragraph is marked as inserted.")

    # Share one registry across the calls so ids stay unique (single namespace).
    registry = RevisionIdRegistry(doc)
    mark_insertion(inserted, author="Alice", id_registry=registry)
    mark_deletion((start, end), author="Bob", id_registry=registry)
    mark_insertion(p3, author="Carol", id_registry=registry)

    doc.save(str(out_path))
    return out_path


def main(argv: list[str] | None = None) -> int:
    """Entry point. One optional positional arg: the output docx path."""
    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) > 1:
        print(
            "usage: python -m docx_plus.examples.track_changes [output.docx]",
            file=sys.stderr,
        )
        return 2

    out_path = Path(args[0]).expanduser().resolve() if args else Path.cwd() / "tracked.docx"
    written = build_tracked_document(out_path)
    print(f"# wrote: {written}")

    reopened = Document(str(written))
    revisions = read_revisions(reopened)
    print(f"# round-tripped {len(revisions)} revisions:")
    for rv in revisions:
        text = rv.text or "(no text)"
        print(f"#   [{rv.author}] {rv.revision_type}: {text!r}")

    # Demonstrate resolving every change into final text.
    accept_all_revisions(reopened)
    remaining = read_revisions(reopened)
    print(f"# after accept_all: {len(remaining)} revisions remain")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
