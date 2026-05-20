"""Add anchored comments to a document.

Demonstrates :func:`docx_plus.comments.add_comment` against three target
shapes (a single run, a whole paragraph, a run-range) plus
:func:`read_comments` round-tripping.

Usage::

    python -m docx_plus.examples.add_comments              # writes ./commented.docx
    python -m docx_plus.examples.add_comments path/out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from docx_plus.comments import (
    CommentIdRegistry,
    add_comment,
    read_comments,
)


def build_commented_document(out_path: Path) -> Path:
    """Build a small document with three anchored comments. Returns the path."""
    doc = Document()
    doc.add_heading("Review notes", level=1)

    p1 = doc.add_paragraph()
    p1.add_run("This first sentence ")
    inline = p1.add_run("contains a phrase")
    p1.add_run(" we'd like to flag.")

    p2 = doc.add_paragraph("Whole paragraph wrap — every run gets included.")
    p3 = doc.add_paragraph()
    p3.add_run("Range example: ")
    start = p3.add_run("from here ")
    p3.add_run("through ")
    end = p3.add_run("to here.")

    # Share one registry across the three calls so ids stay unique.
    registry = CommentIdRegistry(doc)
    add_comment(
        inline,
        "Comment on a single run.",
        author="Alice",
        initials="A",
        id_registry=registry,
    )
    add_comment(
        p2,
        "Whole-paragraph comment.",
        author="Bob",
        initials="B",
        id_registry=registry,
    )
    add_comment(
        (start, end),
        "Run-range comment spanning multiple runs.",
        author="Carol",
        initials="C",
        id_registry=registry,
    )

    doc.save(str(out_path))
    return out_path


def main(argv: list[str] | None = None) -> int:
    """Entry point. One optional positional arg: the output docx path."""
    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) > 1:
        print(
            "usage: python -m docx_plus.examples.add_comments [output.docx]",
            file=sys.stderr,
        )
        return 2

    out_path = Path(args[0]).expanduser().resolve() if args else Path.cwd() / "commented.docx"
    written = build_commented_document(out_path)
    print(f"# wrote: {written}")

    reopened = Document(str(written))
    comments = read_comments(reopened)
    print(f"# round-tripped {len(comments)} comments:")
    for c in comments:
        text = c.anchored_text or "(unanchored)"
        print(f"#   [{c.author}] '{text}' → {c.text!r}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
