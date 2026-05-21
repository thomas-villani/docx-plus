"""Add footnotes and endnotes to a document.

Demonstrates :func:`docx_plus.notes.add_footnote` and
:func:`add_endnote`, plus :func:`read_footnotes` / :func:`read_endnotes`
for round-tripping.

Usage::

    python -m docx_plus.examples.footnotes_and_endnotes              # writes ./notes.docx
    python -m docx_plus.examples.footnotes_and_endnotes path/out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from docx_plus.notes import (
    add_endnote,
    add_footnote,
    read_endnotes,
    read_footnotes,
)


def build_notes_document(out_path: Path) -> Path:
    """Build a small doc with footnotes and endnotes."""
    doc = Document()
    doc.add_heading("On the structure of OOXML notes", level=1)

    p1 = doc.add_paragraph("Footnotes appear at the bottom of the page where they are referenced")
    add_footnote(p1, "Foot of the same page is the conventional placement.")

    p2 = doc.add_paragraph("Endnotes are collected at the end of the document or section")
    add_endnote(p2, "End-of-document is the typical placement for endnotes.")

    p3 = doc.add_paragraph("Both kinds share the same part-plus-body-marker shape")
    add_footnote(p3, "The reference marker lives inline; the text lives in a separate part.")
    add_endnote(p3, "Same pattern, different parts.")

    doc.save(str(out_path))
    return out_path


def main(argv: list[str] | None = None) -> int:
    """Entry point. One optional positional arg: the output docx path."""
    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) > 1:
        print(
            "usage: python -m docx_plus.examples.footnotes_and_endnotes [output.docx]",
            file=sys.stderr,
        )
        return 2

    out_path = Path(args[0]).expanduser().resolve() if args else Path.cwd() / "notes.docx"
    written = build_notes_document(out_path)
    print(f"# wrote: {written}")

    reopened = Document(str(written))
    fns = read_footnotes(reopened)
    ens = read_endnotes(reopened)
    print(f"# footnotes: {len(fns)} -- {[f.text for f in fns]}")
    print(f"# endnotes:  {len(ens)} -- {[e.text for e in ens]}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
