"""Build a document with ruled, shaded, and merged tables.

Demonstrates the v0.5 tables surface: :func:`set_table_borders` and
:func:`set_cell_borders` for the ruling python-docx cannot express,
:func:`set_row_shading` for a header band, :func:`merge_cells` for a
spanning title, :func:`unmerge_cell` for the inverse python-docx has no
notion of, and :func:`read_table_formatting` to read it all back.

Open the result in Word: the budget table has a navy header row, hairline
inside rules, a heavier outer box, and a merged banner across the top;
the summary table below shows the same region after being split apart
again.

Usage::

    python -m docx_plus.examples.table_formatting              # ./tables.docx
    python -m docx_plus.examples.table_formatting path/out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from docx_plus.core import Border
from docx_plus.tables import (
    Shading,
    merge_cells,
    read_table_formatting,
    set_cell_borders,
    set_row_shading,
    set_table_borders,
    unmerge_cell,
)

HEADER_FILL = "2F5496"
ROWS = [
    ("Engineering", "1,240,000", "1,198,400"),
    ("Operations", "610,000", "655,200"),
    ("Research", "480,000", "471,900"),
]


def build_table_document(out_path: Path) -> Path:
    """Build a document exercising each table-formatting helper."""
    doc = Document()
    doc.add_heading("Annual Budget", level=1)

    # Row 0 becomes a merged banner, row 1 the column headers.
    table = doc.add_table(rows=len(ROWS) + 2, cols=3)

    banner = merge_cells(table.cell(0, 0), table.cell(0, 2))
    banner.text = "Fiscal Year 2026"
    set_row_shading(table.rows[0], Shading(fill=HEADER_FILL))
    # Shading is a cell property; the contrasting text on top of it is
    # ordinary python-docx run formatting.
    banner_run = banner.paragraphs[0].runs[0]
    banner_run.bold = True
    banner_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for column, heading in enumerate(("Department", "Budgeted", "Actual")):
        table.cell(1, column).text = heading
    set_row_shading(table.rows[1], Shading(fill="D9E2F3"))

    for offset, row in enumerate(ROWS, start=2):
        for column, value in enumerate(row):
            table.cell(offset, column).text = value

    # A heavier box around the whole table, hairlines inside it.
    hairline = Border(style="single", size=4, color="8EAADB")
    set_table_borders(
        table,
        all_edges=Border(style="single", size=12, color=HEADER_FILL),
        inside_h=hairline,
        inside_v=hairline,
    )

    # One cell called out with a double rule above the figure.
    set_cell_borders(table.cell(3, 2), top=Border(style="double", size=6, color="C00000"))

    doc.add_paragraph()
    doc.add_paragraph("The same region, unmerged:").runs[0].font.size = Pt(11)

    # Merge then unmerge, to show the inverse python-docx does not offer.
    split = doc.add_table(rows=2, cols=3)
    merged = merge_cells(split.cell(0, 0), split.cell(1, 1))
    merged.text = "merged 2x2"
    unmerge_cell(merged)
    split.cell(0, 1).text = "restored"
    set_table_borders(split, all_edges=hairline)

    doc.save(str(out_path))
    return out_path


def main() -> None:
    """Write the document, then read its table formatting back."""
    out_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("tables.docx")
    build_table_document(out_path)
    print(f"wrote {out_path}")

    reopened = Document(str(out_path))
    for index, table in enumerate(reopened.tables):
        read = read_table_formatting(table)
        edges = ", ".join(sorted(read.borders)) or "(none)"
        print(f"\ntable {index}: {len(read.cells)} cells, borders [{edges}]")

        shaded = [c for c in read.cells if c.shading is not None]
        if shaded:
            fills = sorted({c.shading.fill for c in shaded if c.shading is not None})
            print(f"  shaded cells: {len(shaded)}, fills [{', '.join(fills)}]")

        spans = [c for c in read.cells if c.grid_span > 1 or c.vertical_merge is not None]
        print(f"  merged cells: {len(spans)}")


if __name__ == "__main__":
    main()
