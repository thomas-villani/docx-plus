"""Demonstrate ``docx_plus.layout`` — columns, mid-document section
breaks, and the doc-level even/odd-headers flag.

Usage::

    python -m docx_plus.examples.multi_column_layout                # writes ./multicol.docx
    python -m docx_plus.examples.multi_column_layout path/out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from docx_plus.layout import (
    enable_distinct_even_odd_headers,
    insert_section_break,
    set_columns,
)


def build_multi_column_document(out_path: Path) -> Path:
    """Build a small document with two sections — single-column then two-column."""
    doc = Document()
    doc.add_heading("Single-column intro", level=1)
    doc.add_paragraph(
        "This first section is a regular single-column flow. After the next "
        "paragraph there's a section break, and the rest of the document "
        "renders in two columns with a separator line between them."
    )
    split_anchor = doc.add_paragraph("Break point ↓")

    new_section = insert_section_break(split_anchor, start_type="continuous")
    set_columns(new_section, 2, space=720, separator=True)

    doc.add_heading("Two-column body", level=1)
    for i in range(1, 6):
        doc.add_paragraph(
            f"Paragraph {i}: lorem ipsum dolor sit amet, consectetur adipiscing "
            "elit. Sed do eiusmod tempor incididunt ut labore et dolore magna "
            "aliqua."
        )

    # Doc-level flag — distinct even-page vs odd-page headers and footers.
    enable_distinct_even_odd_headers(doc)

    doc.save(str(out_path))
    return out_path


def main(argv: list[str] | None = None) -> int:
    """Entry point. One optional positional arg: the output docx path."""
    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) > 1:
        print(
            "usage: python -m docx_plus.examples.multi_column_layout [output.docx]",
            file=sys.stderr,
        )
        return 2

    out_path = (
        Path(args[0]).expanduser().resolve() if args else Path.cwd() / "multicol.docx"
    )
    written = build_multi_column_document(out_path)
    print(f"# wrote: {written}")
    reopened = Document(str(written))
    print(f"# sections: {len(reopened.sections)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
