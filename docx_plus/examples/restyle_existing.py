"""Restyle an existing document by modifying its Heading 1 definition.

Demonstrates the Word-native workflow: change the *style*, not each
paragraph. Every paragraph already using ``Heading1`` picks up the new
color, size, and spacing at next open — and so will any heading added
later. This is the entire point of paragraph styles, and it's the workflow
python-docx makes awkward enough that most code reaches for direct
formatting instead.

Usage::

    python -m docx_plus.examples.restyle_existing                    # demo
    python -m docx_plus.examples.restyle_existing in.docx out.docx   # restyle a real doc

The no-arg form writes ``restyled.docx`` next to the script and re-resolves
one of the restyled headings to prove the cascade reflects the change.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObj

from docx_plus.styles import (
    apply_style,
    ensure_style,
    modify_style,
    resolve_effective_formatting,
)


def _build_demo_doc() -> DocumentObj:
    """Build a small in-memory document with several Heading 1 paragraphs."""
    doc = Document()
    ensure_style(doc, "Heading1")

    doc.add_paragraph("Quarterly Update")  # body intro

    for section_title, body in [
        ("Reliability", "p99 latency down 38%."),
        ("Cost", "Storage spend down 12% after the dedup migration."),
        ("Headcount", "Two SRE hires closed; one open requisition."),
    ]:
        heading = doc.add_paragraph(section_title)
        apply_style(heading, "Heading1")
        doc.add_paragraph(body)

    return doc


def restyle_heading_one(doc: DocumentObj) -> None:
    """Mutate the ``Heading1`` style in-place.

    After this call, every paragraph already using ``Heading1`` re-renders
    with the new properties at next document open — no per-paragraph
    rewrites required.
    """
    ensure_style(doc, "Heading1")
    modify_style(
        doc,
        "Heading1",
        font_size=20.0,
        color_rgb="C00000",  # Word's "Dark Red" accent
        bold=True,
        spacing_before=480,
        spacing_after=120,
    )


def _print_heading_after_restyle(doc: DocumentObj) -> None:
    """Print the resolved formatting of the first Heading1 paragraph."""
    for paragraph in doc.paragraphs:
        resolved = resolve_effective_formatting(paragraph)
        if resolved.style_id == "Heading1":
            print(f'  first Heading1: "{paragraph.text}"')
            print(f"    font_size : {resolved.font_size}")
            print(f"    color_rgb : {resolved.color_rgb}")
            print(f"    bold      : {resolved.bold}")
            return
    print("  (no Heading1 paragraphs found)")


def main(argv: list[str] | None = None) -> int:
    """Entry point. Accepts ``in.docx out.docx`` or runs the no-arg demo."""
    args = list(sys.argv[1:] if argv is None else argv)

    if len(args) == 0:
        in_doc = _build_demo_doc()
        out_path = Path.cwd() / "restyled.docx"
        print("# restyling: (built-in demo document)")
    elif len(args) == 2:
        in_path = Path(args[0]).expanduser().resolve()
        if not in_path.exists():
            print(f"error: {in_path} not found", file=sys.stderr)
            return 1
        in_doc = Document(str(in_path))
        out_path = Path(args[1]).expanduser().resolve()
        print(f"# restyling: {in_path}")
    else:
        print(
            "usage: python -m docx_plus.examples.restyle_existing [input.docx output.docx]",
            file=sys.stderr,
        )
        return 2

    restyle_heading_one(in_doc)
    in_doc.save(str(out_path))
    print(f"# wrote: {out_path}")
    print()
    print("# verifying via resolve_effective_formatting:")
    _print_heading_after_restyle(Document(str(out_path)))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
