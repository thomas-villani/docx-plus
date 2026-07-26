"""Build a document with threaded, resolvable comments.

Demonstrates the v0.4 comment surface: :func:`reply_to_comment` to build a
thread, :func:`resolve_comment` / :func:`reopen_comment` to toggle a
thread's resolved state, and :func:`read_threads` to read the nesting back
after a round-trip.

Open the result in Word and check the review pane: the first thread shows
a root with two nested replies and is greyed out as resolved; the second
is still open.

Usage::

    python -m docx_plus.examples.threaded_comments              # ./threaded.docx
    python -m docx_plus.examples.threaded_comments path/out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from docx_plus.comments import (
    CommentIdRegistry,
    add_comment,
    read_threads,
    reopen_comment,
    reply_to_comment,
    resolve_comment,
)
from docx_plus.core.ids import ParaIdRegistry


def build_threaded_document(out_path: Path) -> Path:
    """Build a document with one resolved thread and one open one."""
    doc = Document()
    doc.add_heading("Draft under review", level=1)

    p1 = doc.add_paragraph()
    p1.add_run("The estimate assumes ")
    claim = p1.add_run("a six-week schedule")
    p1.add_run(", which drives the rest of the plan.")

    p2 = doc.add_paragraph("Budget figures are carried over from last quarter.")

    # Share both registries so ids and paraIds stay unique across the batch.
    ids = CommentIdRegistry(doc)
    para_ids = ParaIdRegistry(doc)

    schedule = add_comment(
        claim,
        "Where does the six-week figure come from?",
        author="Alice",
        initials="A",
        id_registry=ids,
        para_id_registry=para_ids,
    )
    reply_to_comment(
        doc,
        schedule.comment_id,
        "From the Q2 capacity model.",
        author="Bob",
        initials="B",
        id_registry=ids,
        para_id_registry=para_ids,
    )
    reply_to_comment(
        doc,
        schedule.comment_id,
        "Thanks - that settles it.",
        author="Alice",
        initials="A",
        id_registry=ids,
        para_id_registry=para_ids,
    )
    # The question was answered, so close the thread. Both replies close with it.
    resolve_comment(doc, schedule.comment_id)

    budget = add_comment(
        p2,
        "These need refreshing before we circulate.",
        author="Carol",
        initials="C",
        id_registry=ids,
        para_id_registry=para_ids,
    )
    # Demonstrate that reopening is the exact inverse: resolve, then undo it.
    resolve_comment(doc, budget.comment_id)
    reopen_comment(doc, budget.comment_id)

    doc.save(str(out_path))
    return out_path


def main(argv: list[str] | None = None) -> int:
    """Entry point. One optional positional arg: the output docx path."""
    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) > 1:
        print(
            "usage: python -m docx_plus.examples.threaded_comments [output.docx]",
            file=sys.stderr,
        )
        return 2

    out_path = Path(args[0]).expanduser().resolve() if args else Path.cwd() / "threaded.docx"
    written = build_threaded_document(out_path)
    print(f"# wrote: {written}")

    reopened = Document(str(written))
    threads = read_threads(reopened)
    print(f"# round-tripped {len(threads)} threads:")
    for thread in threads:
        state = "resolved" if thread.resolved else "open"
        anchor = thread.root.anchored_text or "(unanchored)"
        print(f"#   [{state}] on '{anchor}'")
        print(f"#     {thread.root.author}: {thread.root.text}")
        for reply in thread.replies:
            print(f"#       -> {reply.author}: {reply.text}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
