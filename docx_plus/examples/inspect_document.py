"""Inspect a document's effective formatting paragraph-by-paragraph.

Demonstrates :func:`docx_plus.styles.resolve_effective_formatting` with
``include_provenance=True``: for each paragraph in the document, prints the
fields the cascade actually resolved and, for each one, the cascade layer
that produced it.

Usage::

    python -m docx_plus.examples.inspect_document               # built-in demo
    python -m docx_plus.examples.inspect_document path/to.docx  # any docx

Output shape::

    [1] "Document Title"
        style: Title (paragraph)
        font_name: Calibri Light       <- paragraphStyle: Title
        font_size: 28.0                <- paragraphStyle: Title
        bold: True                     <- paragraphStyle: Title (toggle)
        ...

When run with no path the script generates a small demo document inline so
the cascade has interesting things to print: a Title, a Heading1, and a
plain body paragraph. Pointing it at a real document is how you answer the
"why does this paragraph look the way it does?" question on real content.
"""

from __future__ import annotations

import sys
from dataclasses import fields as dataclass_fields
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObj

from docx_plus.styles import (
    FormattingSource,
    ResolvedFormatting,
    apply_style,
    ensure_style,
    resolve_effective_formatting,
)

# Fields excluded from the printed output. style_id/style_name are shown on
# their own header line; partial/provenance are meta-fields.
_META_FIELDS = frozenset({"style_id", "style_name", "partial", "provenance"})


def _format_source(src: FormattingSource) -> str:
    """Render a FormattingSource as a one-line suffix."""
    parts: list[str] = [src.layer]
    if src.style_id is not None:
        parts.append(f": {src.style_id}")
    extras: list[str] = []
    if src.chain_depth is not None and src.chain_depth > 0:
        extras.append(f"chain_depth={src.chain_depth}")
    if src.is_toggle_resolved:
        extras.append("toggle XOR")
    if extras:
        parts.append(f" ({', '.join(extras)})")
    return "".join(parts)


def _print_resolved(index: int, text: str, resolved: ResolvedFormatting) -> None:
    """Print one paragraph's resolved formatting in SPEC §11 shape."""
    snippet = text if len(text) <= 60 else text[:57] + "..."
    print(f'[{index}] "{snippet}"')

    style_label = resolved.style_id or "(no style)"
    if resolved.style_name and resolved.style_name != resolved.style_id:
        style_label = f"{resolved.style_id} ({resolved.style_name})"
    print(f"    style: {style_label}")

    provenance = resolved.provenance or {}
    rows: list[tuple[str, str, str]] = []
    for f in dataclass_fields(resolved):
        if f.name in _META_FIELDS:
            continue
        value = getattr(resolved, f.name)
        if value is None:
            continue
        src = provenance.get(f.name)
        suffix = f"  <- {_format_source(src)}" if src is not None else ""
        rows.append((f.name, repr(value), suffix))

    if not rows:
        print("    (no fields set)")
        return

    name_w = max(len(name) for name, _, _ in rows)
    value_w = max(len(v) for _, v, _ in rows)
    for name, value, suffix in rows:
        print(f"    {name:<{name_w}} : {value:<{value_w}}{suffix}")

    if resolved.partial:
        print("    (partial: theme part missing or malformed)")


def _build_demo_doc() -> DocumentObj:
    """Build a small in-memory document for the no-arg demo run."""
    doc = Document()
    ensure_style(doc, "Title")
    ensure_style(doc, "Heading1")

    title = doc.add_paragraph("Q3 Engineering Review")
    apply_style(title, "Title")

    heading = doc.add_paragraph("Reliability work")
    apply_style(heading, "Heading1")

    doc.add_paragraph(
        "We reduced p99 request latency by 38% across the ingest tier this quarter, "
        "driven mostly by the buffered-batch rewrite shipped in week 7."
    )
    return doc


def inspect_document(doc: DocumentObj) -> None:
    """Print effective formatting + provenance for every paragraph in ``doc``."""
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text or "(empty paragraph)"
        resolved = resolve_effective_formatting(paragraph, include_provenance=True)
        _print_resolved(index, text, resolved)
        print()


def main(argv: list[str] | None = None) -> int:
    """Entry point. Pass a docx path or run with no args for a built-in demo."""
    args = list(sys.argv[1:] if argv is None else argv)
    if args:
        path = Path(args[0]).expanduser().resolve()
        if not path.exists():
            print(f"error: {path} not found", file=sys.stderr)
            return 1
        print(f"# inspecting: {path}")
        doc = Document(str(path))
    else:
        print("# inspecting: (built-in demo document)")
        doc = _build_demo_doc()
    print()
    inspect_document(doc)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
