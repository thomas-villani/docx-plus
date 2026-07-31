"""Audit a document's formatting, then describe the repair.

Demonstrates :func:`docx_plus.lint.lint` and
:func:`docx_plus.lint.plan_fixes` — the two halves of the linter, both of
them pure reads. ``lint`` reports what is wrong; ``plan_fixes`` turns those
findings into an ordered, serializable description of what a repair pass
*would* change, and stops there. Nothing in this example writes a document.

Usage::

    python -m docx_plus.examples.lint_document               # built-in demo
    python -m docx_plus.examples.lint_document path/to.docx  # any docx

Output shape::

    == findings ==
    W paragraph 1  heading-level-skip
        Outline jumps from level 1 to level 3, skipping level 2.
        > "Deep Dive"
    ...

    == plan ==
    1. [review] paragraph 2  double-space
       Collapse 1 run of spaces to a single space.
       - replace-paragraph-text
    ...

    2 unfixable finding(s): heading-level-skip, manual-list

The demo document is built with defects on purpose: an outline that skips a
level, a typed list marker, doubled spaces, a space before a period, a run
whose direct size is the value it would inherit anyway, and a paragraph
overriding its style's spacing. Pointing the script at a real document is
how you see what it would say about real content.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObj
from docx.shared import Pt

from docx_plus.lint import Finding, FixPlan, lint, plan_fixes
from docx_plus.styles import apply_style, ensure_style

# One-character severity marks keep a long report scannable, and keep the
# output cp1252-safe for a default Windows console.
_MARK = {"error": "E", "warning": "W", "info": "i"}


def _build_demo_doc() -> DocumentObj:
    """Build a small document carrying one instance of several defects."""
    doc = Document()
    ensure_style(doc, "Heading1")
    ensure_style(doc, "Heading3")

    apply_style(doc.add_paragraph("Quarterly Report"), "Heading1")
    # Heading 1 straight to Heading 3: the outline skips a level.
    apply_style(doc.add_paragraph("Deep Dive"), "Heading3")

    doc.add_paragraph("Body with  two spaces and a space .")
    doc.add_paragraph("1. First typed item")

    # A run whose direct size is exactly what it would inherit anyway.
    redundant = doc.add_paragraph()
    redundant.add_run("redundant").font.size = Pt(11)

    # A paragraph overriding what its style says about spacing.
    drifting = doc.add_paragraph("Drifting paragraph")
    drifting.paragraph_format.space_after = Pt(24)

    return doc


def _print_findings(findings: list[Finding]) -> None:
    """Print each finding with its severity mark, location, and excerpt."""
    print("== findings ==")
    if not findings:
        print("(clean)")
        return

    for finding in findings:
        print(f"{_MARK[finding.severity]} {finding.location.describe()}  {finding.rule}")
        print(f"    {finding.message}")
        if finding.location.excerpt:
            print(f'    > "{finding.location.excerpt}"')

    counts = {level: sum(f.severity == level for f in findings) for level in _MARK}
    summary = ", ".join(f"{n} {level}" for level, n in counts.items() if n)
    print(f"\n{len(findings)} finding(s): {summary}.")


def _print_plan(plan: FixPlan) -> None:
    """Print the planned repair: the edits, then what was left out and why."""
    print("== plan ==")
    if not plan.fixes:
        print("(no edits)")

    for position, planned in enumerate(plan.fixes, start=1):
        where = planned.finding.location.describe()
        print(f"{position}. [{planned.safety}] {where}  {planned.rule}")
        print(f"   {planned.fix.summary}")
        for operation in planned.operations:
            print(f"   - {operation.op}")

    # An edit that deletes a paragraph or a style changes what the document
    # *contains*, not how it looks, so plan_fixes withholds it by default.
    if plan.deferred:
        print(f"\n{len(plan.deferred)} withheld (pass allow_content=True to include):")
        for planned in plan.deferred:
            print(f"   {planned.rule}: {planned.fix.summary}")

    if plan.conflicts:
        print(f"\n{len(plan.conflicts)} dropped for colliding with an earlier edit:")
        for conflict in plan.conflicts:
            print(f"   {conflict.dropped.rule} lost to {conflict.kept.rule}: {conflict.reason}")

    if plan.unfixable:
        rules = sorted({finding.rule for finding in plan.unfixable})
        print(f"\n{len(plan.unfixable)} unfixable finding(s): {', '.join(rules)}")


def lint_document(doc: DocumentObj) -> None:
    """Audit ``doc``, print the findings, then print the planned repair."""
    findings = lint(doc)
    _print_findings(findings)
    print()
    _print_plan(plan_fixes(findings))


def main(argv: list[str] | None = None) -> int:
    """Entry point. Pass a docx path or run with no args for a built-in demo."""
    args = list(sys.argv[1:] if argv is None else argv)
    if args:
        path = Path(args[0]).expanduser().resolve()
        if not path.exists():
            print(f"error: {path} not found", file=sys.stderr)
            return 1
        print(f"# linting: {path}")
        doc = Document(str(path))
    else:
        print("# linting: (built-in demo document)")
        doc = _build_demo_doc()
    print()
    lint_document(doc)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
