"""Merging, unmerging, and legacy ``<w:hMerge>`` spans.

Merging itself is one of the few table features python-docx implements
properly: ``_Cell.merge`` grows a rectangular region, refusing L- and
T-shaped selections. This module does **not** re-implement it. What it
adds is the other three-quarters of the story:

- :func:`merge_cells` — the same operation behind a typed error, so a
  caller can ``except DocxPlusError`` uniformly (SPEC §9.7) instead of
  reaching into ``docx.exceptions``.
- :func:`unmerge_cell` — the inverse, which python-docx has no notion
  of. Nothing in the package removes a ``w:gridSpan`` or a
  ``w:vMerge``, so a merge is one-way.
- :func:`normalize_horizontal_merges` — rewrites the *other* horizontal
  merge encoding. OOXML has two (17.4.22 ``w:hMerge`` and 17.4.17
  ``w:gridSpan``); python-docx models only ``gridSpan``, so a table
  written with ``hMerge`` — older Word versions and several converters
  do — reads back as separate cells that Word draws as one.

This module imports only from ``docx_plus.core`` (SPEC §9.1).
"""

from __future__ import annotations

from typing import TYPE_CHECKING, cast

from docx.exceptions import InvalidSpanError
from docx.oxml.table import CT_Tc
from docx.shared import Length

from docx_plus.core.errors import DocxPlusError
from docx_plus.core.ns import qn
from docx_plus.core.oxml import remove, xpath

if TYPE_CHECKING:
    from docx.table import Table, _Cell
    from lxml import etree

#: ECMA-376 17.18.57 ``ST_Merge``.
_RESTART = "restart"
_CONTINUE = "continue"


class InvalidMergeError(DocxPlusError, ValueError):
    """Raised when a merge cannot be performed as requested.

    Covers a non-rectangular selection — the case python-docx signals
    with ``InvalidSpanError``, which this wraps — and a normalization
    that would otherwise discard cell content.

    Subclasses ``ValueError`` so existing handling still catches it;
    also subclasses :class:`DocxPlusError` per SPEC §9.7.
    """


def merge_cells(start: _Cell, end: _Cell) -> _Cell:
    """Merge the rectangular region with ``start`` and ``end`` as corners.

    A thin wrapper over python-docx's ``_Cell.merge`` that translates
    its ``InvalidSpanError`` into a :class:`DocxPlusError` subclass.
    Content from every cell in the region is moved into the top-left
    one, which is what Word does.

    Args:
        start: One corner of the region.
        end: The diagonally opposite corner. May be ``start`` itself,
            in which case nothing changes.

    Returns:
        The merged cell — the top-left of the region, which is not
        necessarily ``start``.

    Raises:
        InvalidMergeError: If the two cells do not define a rectangular
            region, because one of them is already part of a merge that
            makes the selection L- or T-shaped. Subclasses
            :class:`ValueError`.

    Example:
        >>> from docx import Document
        >>> from docx_plus.tables import merge_cells
        >>> doc = Document()
        >>> table = doc.add_table(rows=2, cols=3)
        >>> banner = merge_cells(table.cell(0, 0), table.cell(0, 2))
        >>> banner.text = "Quarterly results"
    """
    try:
        # `_Cell.merge` carries no return annotation in python-docx, so
        # its result arrives as Any.
        return cast("_Cell", start.merge(end))
    except InvalidSpanError as exc:
        raise InvalidMergeError(
            f"the region between these two cells is not rectangular: {exc}"
        ) from exc


def unmerge_cell(cell: _Cell) -> None:
    """Split a merged region back into individual cells.

    The inverse of :func:`merge_cells`, which python-docx does not
    provide in any form. Works from *any* cell in the region — pass a
    vertical continuation and the whole span is still resolved and
    undone.

    Content stays in the original top-left cell; the cells restored
    around it are empty, matching Word's "Split Cells". A merged cell's
    width is divided evenly among them, since the individual widths
    were summed away when the merge happened and cannot be recovered.

    Idempotent: a cell that is not merged is left untouched.

    Args:
        cell: Any cell in the merged region.

    Example:
        >>> from docx import Document
        >>> from docx_plus.tables import merge_cells, unmerge_cell
        >>> doc = Document()
        >>> table = doc.add_table(rows=2, cols=3)
        >>> merged = merge_cells(table.cell(0, 0), table.cell(1, 1))
        >>> unmerge_cell(merged)
        >>> len(table.rows[0].cells)
        3
    """
    tc = cell._tc
    offset = tc.grid_offset
    rows = tc._tbl.tr_lst
    anchor = rows[tc.top].tc_at_grid_offset(offset)

    # Collect the vertical run before mutating anything: splitting a
    # cell horizontally shifts the grid offsets of its right-hand
    # neighbours, so the lookups have to happen first.
    spanned = [anchor]
    for row in rows[tc.top + 1 :]:
        try:
            below = row.tc_at_grid_offset(offset)
        except ValueError:
            # A row with no cell starting at this offset cannot be part
            # of the span — ragged rows (w:gridBefore / w:gridAfter) end
            # it just as a non-continuation cell does.
            break
        if below.vMerge != _CONTINUE:
            break
        spanned.append(below)

    for member in spanned:
        member.vMerge = None
        _split_horizontally(member)


def normalize_horizontal_merges(table: Table, *, discard_content: bool = False) -> int:
    """Rewrite ``<w:hMerge>`` spans in ``table`` as ``<w:gridSpan>``.

    OOXML can express a horizontal merge two ways. ``w:gridSpan`` widens
    one cell over several grid columns; ``w:hMerge`` keeps one
    ``<w:tc>`` per column and marks the followers as continuations.
    Word renders both identically, but python-docx's grid model only
    understands the first — so on an ``hMerge`` table, ``Table.cell``
    hands back cells that look separate and are not, and
    ``Row.cells`` reports a column count Word never shows.

    This converts the second form into the first, in place, leaving the
    rendered table unchanged.

    Args:
        table: A python-docx :class:`~docx.table.Table`.
        discard_content: Whether to drop text held in a continuation
            cell. Such text is invisible in Word — the cell it lives in
            is merged away — so keeping it would make hidden content
            appear, and dropping it silently would lose data. The
            default refuses rather than choosing for you.

    Returns:
        The number of merged regions converted. ``0`` means the table
        used ``gridSpan`` already, which is the common case.

    Raises:
        InvalidMergeError: If a continuation cell holds content and
            ``discard_content`` is ``False``.

    Example:
        >>> from docx import Document
        >>> from docx_plus.tables import normalize_horizontal_merges
        >>> doc = Document()
        >>> normalize_horizontal_merges(doc.add_table(rows=1, cols=2))
        0
    """
    converted = 0
    for row in table._tbl.tr_lst:
        cells = list(row.tc_lst)
        index = 0
        while index < len(cells):
            anchor = cells[index]
            anchor_mark = _h_merge(anchor)
            if anchor_mark is None or _merge_val(anchor_mark) != _RESTART:
                index += 1
                continue

            followers: list[CT_Tc] = []
            probe = index + 1
            while probe < len(cells) and _h_merge_val(cells[probe]) == _CONTINUE:
                followers.append(cells[probe])
                probe += 1

            _absorb(anchor, anchor_mark, followers, discard_content=discard_content)
            if followers:
                converted += 1
            index = probe
    return converted


# ---------------------------------------------------------------------------
# Internals.
# ---------------------------------------------------------------------------


def _h_merge(tc: CT_Tc) -> etree._Element | None:
    """Return the cell's ``<w:hMerge>`` element, or ``None`` if absent."""
    tc_pr = tc.tcPr
    if tc_pr is None:
        return None
    return tc_pr.find(qn("w:hMerge"))


def _merge_val(mark: etree._Element) -> str:
    """Read a ``ST_Merge`` element's value.

    ``w:hMerge`` with no ``w:val`` means ``"continue"`` — ECMA-376
    17.18.57 makes that the default for the attribute, which is why the
    encoding is so easy to miss.
    """
    return mark.get(qn("w:val")) or _CONTINUE


def _h_merge_val(tc: CT_Tc) -> str | None:
    """Return the cell's ``w:hMerge`` value, or ``None`` if absent."""
    mark = _h_merge(tc)
    return None if mark is None else _merge_val(mark)


def _has_content(tc: CT_Tc) -> bool:
    """Whether ``tc`` holds anything a reader would see.

    An empty run does not count. Every cell has at least one ``<w:p>``,
    and assigning ``cell.text = ""`` leaves a ``<w:r><w:t/></w:r>``
    behind, so testing for the presence of a run would call every
    ordinary cell "occupied".
    """
    if xpath(tc, "./w:tbl | .//w:drawing | .//w:pict | .//w:object"):
        return True
    return any(str(text).strip() for text in xpath(tc, ".//w:t/text()"))


def _absorb(
    anchor: CT_Tc,
    anchor_mark: etree._Element,
    followers: list[CT_Tc],
    *,
    discard_content: bool,
) -> None:
    """Widen ``anchor`` over ``followers`` and drop them from the row."""
    if not discard_content:
        holding = [tc for tc in followers if _has_content(tc)]
        if holding:
            raise InvalidMergeError(
                f"{len(holding)} continuation cell(s) of an <w:hMerge> span hold "
                f"content, which Word does not render because the cells are merged "
                f"away; pass discard_content=True to drop it"
            )

    span = anchor.grid_span
    for follower in followers:
        span += follower.grid_span
        if anchor.width is not None and follower.width is not None:
            anchor.width = Length(anchor.width + follower.width)
        remove(follower)

    remove(anchor_mark)
    anchor.grid_span = span


def _split_horizontally(tc: CT_Tc) -> None:
    """Restore a ``w:gridSpan`` cell to that many single-column cells."""
    span = tc.grid_span
    if span < 2:
        return

    total = tc.width
    share = Length(total // span) if total is not None else None

    tc.grid_span = 1
    if share is not None:
        tc.width = share

    cursor = tc
    for _ in range(span - 1):
        fresh = CT_Tc.new()
        if share is not None:
            fresh.width = share
        cursor.addnext(fresh)
        cursor = fresh


__all__ = ["InvalidMergeError", "merge_cells", "normalize_horizontal_merges", "unmerge_cell"]
