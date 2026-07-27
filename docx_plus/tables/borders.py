"""Table and cell borders — ``<w:tblBorders>`` and ``<w:tcBorders>``.

python-docx has no element class for any of this: there is no
``CT_Border``, no ``CT_TblBorders``, and no ``CT_TcBorders`` anywhere in
the package, and none of those tags is registered, so a border set by
hand round-trips as an anonymous ``lxml`` element. Drawing a ruled table
has meant writing the XML yourself.

ECMA-376 uses the same ``CT_Border`` shape here as it does for page
borders, so both go through :class:`~docx_plus.core.borders.Border` and
:func:`~docx_plus.core.borders.border_attrs`. Tables add the two
*inside* edges (17.4.39) and cells add the two diagonals (17.4.67).

.. note::
   ``Border.space`` is ignored for tables and cells — both writers emit
   ``w:space="0"``, which is what Word does and the only value its UI
   can produce. The dataclass default of ``24`` is a *page*-border
   default and would otherwise leak a third of an inch onto every table
   edge.

This module imports only from ``docx_plus.core`` (SPEC §9.1).
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx_plus.core.borders import Border, border_attrs
from docx_plus.core.ns import qn
from docx_plus.core.oxml import el, insert_before_first_anchor, remove, sub

if TYPE_CHECKING:
    from docx.table import Table, _Cell
    from lxml import etree

#: ECMA-376 17.4.39 ``CT_TblBorders`` child order.
_TBL_BORDER_EDGES: tuple[str, ...] = (
    "w:top",
    "w:left",
    "w:bottom",
    "w:right",
    "w:insideH",
    "w:insideV",
)

#: ECMA-376 17.4.67 ``CT_TcBorders`` child order. The four sides and two
#: inside edges match the table container; ``tl2br`` / ``tr2bl`` are the
#: diagonals, which only cells have.
_TC_BORDER_EDGES: tuple[str, ...] = (
    "w:top",
    "w:left",
    "w:bottom",
    "w:right",
    "w:insideH",
    "w:insideV",
    "w:tl2br",
    "w:tr2bl",
)

#: Children of ``w:tblPr`` that follow ``w:tblBorders`` (ECMA-376
#: 17.4.60 ``CT_TblPrBase``). python-docx spells this sequence out in
#: ``CT_TblPr._tag_seq`` but deletes the attribute after building its
#: accessors, so it cannot be borrowed at runtime.
_TBL_PR_AFTER_BORDERS: tuple[str, ...] = (
    "w:shd",
    "w:tblLayout",
    "w:tblCellMar",
    "w:tblLook",
    "w:tblCaption",
    "w:tblDescription",
    "w:tblPrChange",
)

#: Children of ``w:tcPr`` that follow ``w:tcBorders`` (ECMA-376 17.4.70
#: ``CT_TcPr``), for the same reason.
_TC_PR_AFTER_BORDERS: tuple[str, ...] = (
    "w:shd",
    "w:noWrap",
    "w:tcMar",
    "w:textDirection",
    "w:tcFitText",
    "w:vAlign",
    "w:hideMark",
    "w:headers",
    "w:cellIns",
    "w:cellDel",
    "w:cellMerge",
    "w:tcPrChange",
)

#: The value Word writes for ``w:space`` on every table and cell border.
_TABLE_BORDER_SPACE = "0"


def set_table_borders(
    table: Table,
    *,
    all_edges: Border | None = None,
    top: Border | None = None,
    bottom: Border | None = None,
    left: Border | None = None,
    right: Border | None = None,
    inside_h: Border | None = None,
    inside_v: Border | None = None,
) -> None:
    """Set the table-level borders on ``table``.

    A full replacement, not a merge: any existing ``<w:tblBorders>`` is
    discarded first, so an edge left unset ends up absent. Calling with
    every edge ``None`` removes the element rather than writing an empty
    container.

    These are the *table's* borders. A cell's own ``<w:tcBorders>``
    overrides them for that cell, and a table *style* supplies them when
    neither is present — see :func:`set_cell_borders`.

    Args:
        table: A python-docx :class:`~docx.table.Table`.
        all_edges: Applied to the four outer edges and both inside
            edges, as a shorthand for the common "rule everything the
            same way" case. Any explicit edge below overrides it.
        top: The table's top edge.
        bottom: The table's bottom edge.
        left: The table's left edge.
        right: The table's right edge.
        inside_h: The horizontal rules *between* rows.
        inside_v: The vertical rules *between* columns.

    Example:
        >>> from docx import Document
        >>> from docx_plus.core import Border
        >>> from docx_plus.tables import set_table_borders
        >>> doc = Document()
        >>> table = doc.add_table(rows=2, cols=2)
        >>> hairline = Border(style="single", size=4, color="808080")
        >>> set_table_borders(table, all_edges=hairline,
        ...                   top=Border(style="single", size=12))
    """
    tbl_pr = table._tbl.tblPr
    edges = _resolve_edges(
        all_edges,
        {
            "w:top": top,
            "w:left": left,
            "w:bottom": bottom,
            "w:right": right,
            "w:insideH": inside_h,
            "w:insideV": inside_v,
        },
    )
    _write_borders(tbl_pr, "w:tblBorders", edges, _TBL_BORDER_EDGES, _TBL_PR_AFTER_BORDERS)


def set_cell_borders(
    cell: _Cell,
    *,
    all_edges: Border | None = None,
    top: Border | None = None,
    bottom: Border | None = None,
    left: Border | None = None,
    right: Border | None = None,
    tl2br: Border | None = None,
    tr2bl: Border | None = None,
) -> None:
    """Set the borders on a single ``cell``.

    A full replacement, exactly as :func:`set_table_borders` is. Cell
    borders take precedence over the table's, which is how a single
    emphasized row or a boxed total is expressed.

    Args:
        cell: A python-docx :class:`~docx.table._Cell`.
        all_edges: Applied to the four sides. The diagonals are
            deliberately *not* included — they are a decorative
            "crossed-out cell" mark, never something a caller means by
            "all borders".
        top: The cell's top edge.
        bottom: The cell's bottom edge.
        left: The cell's left edge.
        right: The cell's right edge.
        tl2br: Diagonal from the top-left to the bottom-right corner.
        tr2bl: Diagonal from the top-right to the bottom-left corner.

    Example:
        >>> from docx import Document
        >>> from docx_plus.core import Border
        >>> from docx_plus.tables import set_cell_borders
        >>> doc = Document()
        >>> table = doc.add_table(rows=2, cols=2)
        >>> set_cell_borders(table.cell(1, 1),
        ...                  top=Border(style="double", size=6))
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    edges = _resolve_edges(
        all_edges,
        {
            "w:top": top,
            "w:left": left,
            "w:bottom": bottom,
            "w:right": right,
        },
    )
    edges.update({tag: b for tag, b in (("w:tl2br", tl2br), ("w:tr2bl", tr2bl)) if b is not None})
    _write_borders(tc_pr, "w:tcBorders", edges, _TC_BORDER_EDGES, _TC_PR_AFTER_BORDERS)


# ---------------------------------------------------------------------------
# Internals.
# ---------------------------------------------------------------------------


def _resolve_edges(
    all_edges: Border | None,
    explicit: dict[str, Border | None],
) -> dict[str, Border]:
    """Fold ``all_edges`` under the explicitly named edges.

    An edge named explicitly always wins, including over an
    ``all_edges`` that would otherwise cover it.
    """
    resolved: dict[str, Border] = {}
    for tag, border in explicit.items():
        chosen = border if border is not None else all_edges
        if chosen is not None:
            resolved[tag] = chosen
    return resolved


def _table_border_attrs(border: Border) -> dict[str, str]:
    """Serialize ``border`` with ``w:space`` pinned to Word's table value."""
    attrs = border_attrs(border)
    attrs["w:space"] = _TABLE_BORDER_SPACE
    return attrs


def _write_borders(
    parent: etree._Element,
    container_tag: str,
    edges: dict[str, Border],
    order: tuple[str, ...],
    later_siblings: tuple[str, ...],
) -> None:
    """Replace ``parent``'s border container with one holding ``edges``.

    Writes nothing and removes any existing container when ``edges`` is
    empty, so "no borders" never leaves an empty element behind.
    """
    existing = parent.find(qn(container_tag))
    if existing is not None:
        remove(existing)
    if not edges:
        return

    container = el(container_tag)
    for tag in order:
        border = edges.get(tag)
        if border is not None:
            sub(container, tag, **_table_border_attrs(border))
    insert_before_first_anchor(parent, container, later_siblings)


__all__ = ["set_cell_borders", "set_table_borders"]
