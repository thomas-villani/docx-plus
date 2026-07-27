"""Reading table and cell formatting back out.

The inverse of :mod:`docx_plus.tables.borders`,
:mod:`docx_plus.tables.shading`, and :mod:`docx_plus.tables.merge`:
report the borders, shading, and merge state actually written on a
table, its rows, and its cells.

.. warning::
   This reports **direct formatting only** — what is present on the
   ``<w:tblPr>`` and ``<w:tcPr>`` elements themselves. It does not
   resolve the cell-formatting cascade (table style →
   ``<w:tblStylePr>`` conditional branch → direct properties), so a
   table whose ruling comes entirely from a style such as
   ``Table Grid`` reads back with no borders at all — which is the
   truth about its XML, not about its appearance.

   That resolver is a considerably larger piece of work than every
   writer in this package put together, and is scoped out in the same
   terms by :mod:`docx_plus.styles.inspect`, which resolves the
   paragraph and run cascade but explicitly not this one.

This module imports only from ``docx_plus.core`` and its siblings in
``docx_plus.tables`` (SPEC §9.1).
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import TYPE_CHECKING

from docx_plus.core.borders import Border
from docx_plus.core.ns import qn
from docx_plus.tables.borders import _TBL_BORDER_EDGES, _TC_BORDER_EDGES
from docx_plus.tables.shading import Shading

if TYPE_CHECKING:
    from docx.table import Table
    from lxml import etree


@dataclass(frozen=True)
class CellFormatting:
    """Direct formatting and merge state of one ``<w:tc>``.

    Attributes:
        row: Zero-based index of the row the cell appears in. For a
            vertical continuation this is the continuation's own row,
            not the row the span started in.
        column: Zero-based *grid* offset of the cell's left edge. Cells
            to the right of a merged one are offset by the span, so
            this is not the cell's position in ``Row.cells``.
        grid_span: Number of grid columns the cell covers. ``1`` for an
            unmerged cell.
        vertical_merge: ``"restart"`` on the top cell of a vertical
            span, ``"continue"`` on the cells beneath it, ``None`` when
            the cell is not vertically merged.
        horizontal_merge: The same for the legacy ``<w:hMerge>``
            encoding, which is ``None`` on anything Word wrote
            recently. See
            :func:`~docx_plus.tables.normalize_horizontal_merges`.
        borders: Edge name (``"top"``, ``"insideV"``, ``"tl2br"``, …)
            to :class:`~docx_plus.core.borders.Border`. Empty when the
            cell carries no ``<w:tcBorders>``.
        shading: The cell's :class:`~docx_plus.tables.Shading`, or
            ``None`` when it carries no ``<w:shd>``.
    """

    row: int
    column: int
    grid_span: int = 1
    vertical_merge: str | None = None
    horizontal_merge: str | None = None
    borders: dict[str, Border] = field(default_factory=dict)
    shading: Shading | None = None


@dataclass(frozen=True)
class TableFormatting:
    """Direct formatting of a table and each of its cells.

    Attributes:
        style: The ``w:tblStyle`` id, or ``None`` if the table names no
            style. Note that a style is exactly what ``borders`` and
            ``shading`` below do *not* account for.
        borders: Edge name to :class:`~docx_plus.core.borders.Border`
            from the table's ``<w:tblBorders>``. Empty when absent.
        shading: The table's :class:`~docx_plus.tables.Shading`, or
            ``None``.
        cells: Every cell in the table, in row-major document order.
            One entry per ``<w:tc>`` element, so a merged cell appears
            once rather than once per grid column it covers.
    """

    style: str | None = None
    borders: dict[str, Border] = field(default_factory=dict)
    shading: Shading | None = None
    cells: tuple[CellFormatting, ...] = ()


def read_table_formatting(table: Table) -> TableFormatting:
    """Read the direct border, shading, and merge state of ``table``.

    Args:
        table: A python-docx :class:`~docx.table.Table`.

    Returns:
        A :class:`TableFormatting` describing the table and every cell
        in it. Formatting inherited from a table style is **not**
        resolved — see the module warning.

    Example:
        >>> from docx import Document
        >>> from docx_plus.core import Border
        >>> from docx_plus.tables import read_table_formatting, set_table_borders
        >>> doc = Document()
        >>> table = doc.add_table(rows=1, cols=2)
        >>> set_table_borders(table, all_edges=Border(style="single"))
        >>> read_table_formatting(table).borders["insideV"].style
        'single'
    """
    tbl_pr = table._tbl.tblPr
    cells: list[CellFormatting] = []
    for row_index, row in enumerate(table._tbl.tr_lst):
        for tc in row.tc_lst:
            tc_pr = tc.tcPr
            cells.append(
                CellFormatting(
                    row=row_index,
                    column=tc.grid_offset,
                    grid_span=tc.grid_span,
                    vertical_merge=tc.vMerge,
                    horizontal_merge=_merge_val(tc_pr, "w:hMerge"),
                    borders=_read_borders(tc_pr, "w:tcBorders", _TC_BORDER_EDGES),
                    shading=_read_shading(tc_pr),
                )
            )

    return TableFormatting(
        style=table._tbl.tblStyle_val,
        borders=_read_borders(tbl_pr, "w:tblBorders", _TBL_BORDER_EDGES),
        shading=_read_shading(tbl_pr),
        cells=tuple(cells),
    )


# ---------------------------------------------------------------------------
# Internals.
# ---------------------------------------------------------------------------


def _merge_val(parent: etree._Element | None, tag: str) -> str | None:
    """Read a ``ST_Merge`` child's value, defaulting an absent one to continue."""
    if parent is None:
        return None
    node = parent.find(qn(tag))
    if node is None:
        return None
    return node.get(qn("w:val")) or "continue"


def _read_borders(
    parent: etree._Element | None,
    container_tag: str,
    order: tuple[str, ...],
) -> dict[str, Border]:
    """Read a ``<w:tblBorders>`` / ``<w:tcBorders>`` into edge-keyed Borders."""
    if parent is None:
        return {}
    container = parent.find(qn(container_tag))
    if container is None:
        return {}

    borders: dict[str, Border] = {}
    for tag in order:
        edge = container.find(qn(tag))
        if edge is None:
            continue
        borders[tag.partition(":")[2]] = Border(
            style=edge.get(qn("w:val")) or "single",
            size=_int_attr(edge, "w:sz", default=4),
            color=edge.get(qn("w:color")) or "auto",
            space=_int_attr(edge, "w:space", default=0),
        )
    return borders


def _read_shading(parent: etree._Element | None) -> Shading | None:
    """Read a ``<w:shd>`` into a :class:`Shading`, or ``None`` if absent."""
    if parent is None:
        return None
    shd = parent.find(qn("w:shd"))
    if shd is None:
        return None
    return Shading(
        fill=shd.get(qn("w:fill")) or "auto",
        pattern=shd.get(qn("w:val")) or "clear",
        color=shd.get(qn("w:color")) or "auto",
    )


def _int_attr(node: etree._Element, name: str, *, default: int) -> int:
    """Read an integer attribute, falling back when absent or malformed.

    Border attributes are optional and a foreign producer may write
    something outside ``ST_DecimalNumber``; a formatting *reader* should
    report what it can rather than raise on a file it did not write.
    """
    raw = node.get(qn(name))
    if raw is None:
        return default
    try:
        return int(raw)
    except ValueError:
        return default


__all__ = ["CellFormatting", "TableFormatting", "read_table_formatting"]
