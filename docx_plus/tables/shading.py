"""Table, row, and cell shading — ``<w:shd>``.

python-docx has no ``CT_Shd`` class and does not register the ``w:shd``
tag, so cell background fills — the single most common thing anyone
wants from a table beyond its text — have to be written as raw XML.

ECMA-376 17.4.32 models shading as three attributes rather than one
colour: a ``w:fill`` (the background), a ``w:val`` pattern drawn over
it, and a ``w:color`` for that pattern's foreground. Solid fills, which
is what nearly everyone means, are ``pattern="clear"`` with only
``fill`` set — so :class:`Shading` defaults to exactly that and a plain
``Shading(fill="D9E2F3")`` does the obvious thing.

.. note::
   **Rows have no shading element.** ``CT_TrPr`` (17.4.82) has no
   ``w:shd`` child — Word implements "shade this row" by writing the
   same ``<w:shd>`` into every cell in it, and so does
   :func:`set_row_shading`.

This module imports only from ``docx_plus.core`` (SPEC §9.1).
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import TYPE_CHECKING

from docx_plus.core.ns import qn
from docx_plus.core.oxml import el, insert_before_first_anchor, remove

if TYPE_CHECKING:
    from docx.table import Table, _Cell, _Row
    from lxml import etree

# ECMA-376 17.18.79 ST_HexColor, the same simple type `Border.color`
# takes: "auto" or six hex digits.
_HEX_COLOR_RE = re.compile(r"^(auto|[0-9A-Fa-f]{6})$")

# ECMA-376 17.18.78 ST_Shd values are bare identifiers, several of which
# carry digits ("pct25", "thinHorzStripe").
_PATTERN_RE = re.compile(r"^[A-Za-z][A-Za-z0-9]*$")

#: Children of ``w:tblPr`` that follow ``w:shd`` (ECMA-376 17.4.60).
_TBL_PR_AFTER_SHD: tuple[str, ...] = (
    "w:tblLayout",
    "w:tblCellMar",
    "w:tblLook",
    "w:tblCaption",
    "w:tblDescription",
    "w:tblPrChange",
)

#: Children of ``w:tcPr`` that follow ``w:shd`` (ECMA-376 17.4.70).
_TC_PR_AFTER_SHD: tuple[str, ...] = (
    "w:noWrap",
    "w:tcMar",
    "w:textDirection",
    "w:tcFitText",
    "w:vAlign",
    "w:hideMark",
    "w:headers",
    "w:cellIns",
    "w:cellDel",
    "w:cellMerge",
    "w:tcPrChange",
)


@dataclass(frozen=True)
class Shading:
    """A background fill, optionally with a pattern drawn over it.

    Attributes:
        fill: Background colour as ``"RRGGBB"`` hex, or ``"auto"``
            (default) to let the consumer choose. This is the attribute
            that produces a solid block of colour.
        pattern: ECMA-376 17.18.78 ``ST_Shd`` value. ``"clear"``
            (default) means no pattern — just ``fill``. ``"nil"``
            removes shading, and the ``"pct5"``…``"pct95"`` family
            blends ``color`` into ``fill`` by that percentage.
        color: Foreground colour of ``pattern``, as ``"RRGGBB"`` hex or
            ``"auto"`` (default). Has no visible effect while
            ``pattern`` is ``"clear"``.

    Raises:
        ValueError: If ``fill`` or ``color`` is not ``"auto"`` or a
            six-hex-digit ``"RRGGBB"`` string, or if ``pattern`` is not
            a bare identifier.

    Example:
        >>> from docx_plus.tables import Shading
        >>> header = Shading(fill="2F5496")
        >>> hatched = Shading(fill="FFFFFF", pattern="pct25", color="808080")
    """

    fill: str = "auto"
    pattern: str = "clear"
    color: str = "auto"

    def __post_init__(self) -> None:
        """Validate the fields against their ECMA-376 simple types."""
        for name in ("fill", "color"):
            value = getattr(self, name)
            if not _HEX_COLOR_RE.match(value):
                raise ValueError(
                    f"Shading.{name} must be 'auto' or a six-hex-digit 'RRGGBB' "
                    f"string; got {value!r}"
                )
        if not _PATTERN_RE.match(self.pattern):
            raise ValueError(
                "Shading.pattern must be an ECMA-376 17.18.78 ST_Shd value such as "
                f"clear/nil/solid/pct25; got {self.pattern!r}"
            )


def shading_attrs(shading: Shading) -> dict[str, str]:
    """Serialize ``shading`` to the ``CT_Shd`` attribute mapping.

    Args:
        shading: The shading to serialize.

    Returns:
        A ``{"w:val": ..., "w:color": ..., "w:fill": ...}`` mapping in
        the attribute order Word writes.
    """
    return {
        "w:val": shading.pattern,
        "w:color": shading.color,
        "w:fill": shading.fill,
    }


def set_table_shading(table: Table, shading: Shading | None) -> None:
    """Set the table-level shading on ``table``.

    Idempotent — replaces any existing ``<w:shd>``. Passing ``None``
    removes it.

    Table-level shading sits *below* row-banding from a table style and
    below any cell's own shading, so it reads as a default rather than
    an override.

    Args:
        table: A python-docx :class:`~docx.table.Table`.
        shading: The shading to apply, or ``None`` to remove it.

    Example:
        >>> from docx import Document
        >>> from docx_plus.tables import Shading, set_table_shading
        >>> doc = Document()
        >>> set_table_shading(doc.add_table(rows=1, cols=1), Shading(fill="F2F2F2"))
    """
    _write_shading(table._tbl.tblPr, shading, _TBL_PR_AFTER_SHD)


def set_cell_shading(cell: _Cell, shading: Shading | None) -> None:
    """Set the shading on a single ``cell``.

    Idempotent — replaces any existing ``<w:shd>``. Passing ``None``
    removes it.

    Args:
        cell: A python-docx :class:`~docx.table._Cell`.
        shading: The shading to apply, or ``None`` to remove it.

    Example:
        >>> from docx import Document
        >>> from docx_plus.tables import Shading, set_cell_shading
        >>> doc = Document()
        >>> table = doc.add_table(rows=2, cols=2)
        >>> set_cell_shading(table.cell(0, 0), Shading(fill="2F5496"))
    """
    _write_shading(cell._tc.get_or_add_tcPr(), shading, _TC_PR_AFTER_SHD)


def set_row_shading(row: _Row, shading: Shading | None) -> None:
    """Shade every cell in ``row``.

    There is no row-level shading element in the format — ``CT_TrPr``
    has no ``w:shd`` child — so this writes the same ``<w:shd>`` into
    each of the row's cells, which is what Word's UI does for a selected
    row.

    Iterates the row's ``<w:tc>`` elements directly rather than
    ``Row.cells``, so a cell spanning several grid columns is visited
    once rather than once per column it covers.

    Args:
        row: A python-docx :class:`~docx.table._Row`.
        shading: The shading to apply, or ``None`` to remove it from
            every cell in the row.

    Example:
        >>> from docx import Document
        >>> from docx_plus.tables import Shading, set_row_shading
        >>> doc = Document()
        >>> table = doc.add_table(rows=2, cols=3)
        >>> set_row_shading(table.rows[0], Shading(fill="2F5496"))
    """
    for tc in row._tr.tc_lst:
        _write_shading(tc.get_or_add_tcPr(), shading, _TC_PR_AFTER_SHD)


# ---------------------------------------------------------------------------
# Internals.
# ---------------------------------------------------------------------------


def _write_shading(
    parent: etree._Element,
    shading: Shading | None,
    later_siblings: tuple[str, ...],
) -> None:
    """Replace ``parent``'s ``<w:shd>``, or remove it when ``shading`` is None."""
    existing = parent.find(qn("w:shd"))
    if existing is not None:
        remove(existing)
    if shading is None:
        return
    insert_before_first_anchor(parent, el("w:shd", **shading_attrs(shading)), later_siblings)


__all__ = [
    "Shading",
    "set_cell_shading",
    "set_row_shading",
    "set_table_shading",
    "shading_attrs",
]
