r"""Table of Figures — ``TOC \c "Figure"`` complex field.

A Table of Figures is structurally the same as a Table of Contents,
except the field instruction uses the ``\\c "<caption_type>"`` switch
to collect ``SEQ`` captions of the named type instead of paragraphs
with outline levels. The caption helpers in
``docx_plus.publishing.captions`` produce the entries this table
picks up.

This module imports only from ``docx_plus.core`` (SPEC §9.1).
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx_plus.core.oxml import build_complex_field
from docx_plus.publishing._validate import validate_seq_identifier

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph
    from lxml import etree


def add_table_of_figures(
    paragraph: Paragraph,
    *,
    caption_type: str = "Figure",
    hyperlink: bool = True,
) -> etree._Element:
    r"""Append a Table of Figures complex field to ``paragraph``.

    The field lists every caption whose ``SEQ`` name equals
    ``caption_type``. Word populates the result on next open; call
    :func:`docx_plus.fields.mark_fields_dirty` before saving so the
    table fills in automatically.

    Args:
        paragraph: A python-docx :class:`~docx.text.paragraph.Paragraph`
            where the field is appended.
        caption_type: ``SEQ`` name to match — must equal the
            ``caption_type`` passed to
            :func:`docx_plus.publishing.add_caption`. Must conform to
            the SEQ identifier rule (ASCII letter/underscore start,
            then letters/digits/underscores).
        hyperlink: When ``True`` (default), entries become clickable
            hyperlinks via the ``\h`` switch.

    Returns:
        The ``<w:r>`` element wrapping the field's ``begin`` ``fldChar``.

    Raises:
        ValueError: If ``caption_type`` is empty or violates the SEQ
            identifier rule (issues.md H11).

    Example:
        >>> from docx import Document
        >>> from docx_plus.publishing import add_caption, add_table_of_figures
        >>> from docx_plus.fields import mark_fields_dirty
        >>> doc = Document()
        >>> doc.add_paragraph("List of Figures")
        >>> add_table_of_figures(doc.add_paragraph())
        >>> # ... captions added elsewhere via add_caption(..., caption_type="Figure")
        >>> mark_fields_dirty(doc)
    """
    validate_seq_identifier(caption_type, arg_name="caption_type")

    instruction = f' TOC \\c "{caption_type}"'
    if hyperlink:
        instruction += " \\h"
    instruction += " \\z"
    instruction += " "
    return build_complex_field(paragraph._p, instruction, "")


__all__ = ["add_table_of_figures"]
