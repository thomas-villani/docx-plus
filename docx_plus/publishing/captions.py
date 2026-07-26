r"""Figure / table captions — leading text + ``SEQ`` complex field.

A Word caption is a paragraph that opens with a label run (e.g.
``"Figure "``) followed by a ``SEQ`` complex field that auto-numbers
items of the same caption type. The Table of Figures (see
``docx_plus.publishing.figures``) picks up captions whose ``SEQ`` name
matches its ``\c`` switch.

Making a caption *referenceable* ("see Figure 3") needs one more thing,
and it is not obvious: a ``REF`` field cannot point at a ``SEQ`` field.
It can only point at a **bookmark**. So the caption has to be bracketed
by one, which is what ``bookmark_name`` does — see
:func:`add_caption`.

This module imports only from ``docx_plus.core`` (SPEC §9.1).
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from lxml import etree

from docx_plus.core.ids import BookmarkIdRegistry
from docx_plus.core.oxml import (
    body_document_for,
    build_bookmark,
    build_complex_field,
    el,
    sub,
    validate_bookmark_name,
)
from docx_plus.publishing._validate import (
    validate_numbering_picture,
    validate_seq_identifier,
)

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph


def add_caption(
    paragraph: Paragraph,
    label: str | None = None,
    *,
    caption_type: str = "Figure",
    numbering: str = "ARABIC",
    bookmark_name: str | None = None,
    bookmark_id_registry: BookmarkIdRegistry | None = None,
) -> etree._Element:
    r"""Append a caption (label run + auto-numbered ``SEQ`` field).

    The label is emitted as a literal text run; the number is a
    ``SEQ`` complex field that Word re-numbers on open. After all
    captions are inserted, call
    :func:`docx_plus.fields.mark_fields_dirty` so Word recalculates
    the SEQ values.

    Args:
        paragraph: A python-docx :class:`~docx.text.paragraph.Paragraph`
            where the caption is appended. Typically a fresh paragraph
            beneath the figure / table being captioned.
        label: Leading text shown before the number, including any
            trailing whitespace. When omitted (``None``, the default),
            uses ``f"{caption_type} "`` — the common case. Pass an
            empty string ``""`` to suppress the label run entirely
            (e.g. when the surrounding paragraph already supplies it).
        caption_type: The ``SEQ`` field's name. Items sharing this name
            are numbered together (so all ``"Figure"`` captions number
            ``1, 2, 3, …``, independent of all ``"Table"`` captions).
            Must match the ``\c`` switch on any downstream Table of
            Figures, and must conform to the SEQ identifier rule
            (ASCII letter/underscore start, then letters/digits/
            underscores).
        numbering: Word numbering format token for the ``\* <picture>``
            switch. Common values: ``"ARABIC"`` (default — ``1, 2, 3,
            …``), ``"ROMAN"`` (``I, II, III, …``), ``"roman"``
            (``i, ii, iii, …``), ``"ALPHABETIC"`` (``A, B, C, …``).
            See ECMA-376 17.16.4.1 for the full token list.
        bookmark_name: Bracket the label and number in a bookmark of this
            name, making the caption referenceable.

            Needed because a ``REF`` field **cannot point at a ``SEQ``
            field** — only at a bookmark. Without this, "see Figure 3" is
            not expressible: there is nothing for the reference to target.

            The bookmark spans exactly the label run plus the ``SEQ``
            field, so a bare
            :func:`~docx_plus.bookmarks.add_cross_reference` to it
            resolves to ``"Figure 3"`` — the same extent Word's own
            "Only label and number" option uses. Add any descriptive text
            with ``paragraph.add_run(...)`` *after* this call and it stays
            outside the bookmark.

            For an anchor the reader never sees, mint a hidden Word-style
            name with
            :meth:`~docx_plus.bookmarks.BookmarkNameRegistry.next_ref_name`.
        bookmark_id_registry: Pre-existing bookmark-id allocator to share
            across an editing session. Only consulted when
            ``bookmark_name`` is given.

    Returns:
        The ``<w:r>`` element wrapping the field's ``begin`` ``fldChar``.

    Raises:
        ValueError: If ``caption_type`` is empty or violates the SEQ
            identifier rule, if ``numbering`` is not a recognised
            format token (issues.md H11, M16), or if ``bookmark_name``
            violates Word's bookmark-name grammar.

    Note:
        The caption's paragraph is *not* automatically restyled to
        Word's built-in ``Caption`` paragraph style. Apply it yourself
        if you want the conventional italic-grey rendering:
        ``paragraph.style = doc.styles["Caption"]``.

    Example:
        >>> from docx import Document
        >>> from docx_plus.publishing import add_caption
        >>> doc = Document()
        >>> p = doc.add_paragraph()
        >>> add_caption(p, caption_type="Figure")  # label defaults to "Figure "
        >>> p.add_run(": Architecture overview")
        <docx.text.run.Run object at 0x...>

    Example:
        A referenceable caption, and the reference to it::

            >>> from docx_plus.bookmarks import add_cross_reference
            >>> from docx_plus.publishing import add_caption
            >>> doc = Document()
            >>> cap = doc.add_paragraph()
            >>> _ = add_caption(cap, bookmark_name="fig_arch")
            >>> cap.add_run(": Architecture overview")
            <docx.text.run.Run object at 0x...>
            >>> body = doc.add_paragraph("As shown in ")
            >>> _ = add_cross_reference(body, bookmark="fig_arch")
    """
    validate_seq_identifier(caption_type, arg_name="caption_type")
    validate_numbering_picture(numbering)
    if bookmark_name is not None:
        validate_bookmark_name(bookmark_name, arg_name="bookmark_name")

    if label is None:
        label = f"{caption_type} "

    first_element: etree._Element | None = None
    if label:
        label_run = el("w:r")
        label_t = sub(label_run, "w:t", **{"xml:space": "preserve"})
        label_t.text = label
        paragraph._p.append(label_run)
        first_element = label_run

    instruction = f" SEQ {caption_type} \\* {numbering} "
    begin_run = build_complex_field(paragraph._p, instruction, "1")
    if first_element is None:
        first_element = begin_run

    if bookmark_name is not None:
        # build_complex_field appends its five runs, so the field's `end`
        # run is now the paragraph's last child — the far edge of what the
        # bookmark should span.
        document = body_document_for(paragraph, operation="add_caption(bookmark_name=...)")
        if bookmark_id_registry is None:
            bookmark_id_registry = BookmarkIdRegistry(document)
        build_bookmark(
            first_element,
            paragraph._p[-1],
            bookmark_id=bookmark_id_registry.next(),
            name=bookmark_name,
        )

    return begin_run


__all__ = ["add_caption"]
