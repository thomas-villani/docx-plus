"""Theme resolution edge cases: missing, malformed, or partial theme parts.

SPEC §4 "Theme references": theme resolution failures must *not* raise. They
set ``partial=True`` on the result. When the theme part is entirely absent
the unresolved name passes through as the value (so debugging output stays
useful); when the theme loaded but the name is unknown, no value is stored
(a bare name is not valid hex — see M9).
"""

from __future__ import annotations

from docx import Document

from docx_plus.core.oxml import sub
from docx_plus.styles.inspect import resolve_effective_formatting
from docx_plus.styles.theme import (
    ThemeColors,
    load_theme,
    resolve_theme_color,
)

_THEME_RELTYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"


def _strip_theme(doc: Document) -> None:
    """Remove every theme relationship from ``doc``'s main part."""
    rids = [rid for rid, rel in doc.part.rels.items() if rel.reltype == _THEME_RELTYPE]
    for rid in rids:
        del doc.part.rels[rid]


def _corrupt_theme(doc: Document) -> None:
    """Overwrite the theme part's blob with garbage that's not valid XML."""
    for rel in doc.part.rels.values():
        if rel.reltype == _THEME_RELTYPE:
            # python-docx's Part keeps the blob in a private attribute.
            rel.target_part._blob = b"<<<not xml<<<"
            return


# --------------------------------------------------------------------------
# load_theme: defensive against missing/malformed parts.
# --------------------------------------------------------------------------


def test_load_theme_returns_none_when_no_theme_part() -> None:
    doc = Document()
    _strip_theme(doc)
    assert load_theme(doc) is None


def test_load_theme_returns_none_when_theme_part_malformed() -> None:
    doc = Document()
    _corrupt_theme(doc)
    assert load_theme(doc) is None


# --------------------------------------------------------------------------
# resolve_theme_color: unknown names return None silently.
# --------------------------------------------------------------------------


def test_resolve_theme_color_unknown_name_returns_none() -> None:
    theme = ThemeColors(scheme={"accent1": "4F81BD"})
    assert resolve_theme_color(theme, "accent99") is None


def test_resolve_theme_color_known_name_missing_from_scheme_returns_none() -> None:
    """If accent2 isn't in the scheme dict at all, resolution returns None."""
    theme = ThemeColors(scheme={"accent1": "4F81BD"})
    assert resolve_theme_color(theme, "accent2") is None


# --------------------------------------------------------------------------
# Resolver: missing theme part propagates partial=True end-to-end.
# --------------------------------------------------------------------------


def test_resolver_marks_partial_when_theme_part_missing() -> None:
    """A theme-color reference on a doc without a theme part yields partial=True."""
    doc = Document()
    _strip_theme(doc)

    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "Themed"})
    sub(s, "w:name", **{"w:val": "Themed"})
    rpr = sub(s, "w:rPr")
    sub(
        rpr,
        "w:color",
        **{
            "w:val": "auto",
            "w:themeColor": "accent1",
            "w:themeShade": "80",
        },
    )

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "Themed"})

    resolved = resolve_effective_formatting(p)
    assert resolved.partial is True
    # Unresolved theme name passes through so callers can still log it.
    assert resolved.color_rgb == "accent1"


def test_resolver_marks_partial_when_theme_part_malformed() -> None:
    doc = Document()
    _corrupt_theme(doc)

    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "T"})
    sub(s, "w:name", **{"w:val": "T"})
    rpr = sub(s, "w:rPr")
    sub(rpr, "w:color", **{"w:val": "auto", "w:themeColor": "accent1"})

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "T"})

    resolved = resolve_effective_formatting(p)
    assert resolved.partial is True
    assert resolved.color_rgb == "accent1"


def test_resolver_marks_partial_for_unknown_theme_name() -> None:
    """A themeColor name that isn't in ST_ThemeColor still parses but is partial."""
    doc = Document()
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "T2"})
    sub(s, "w:name", **{"w:val": "T2"})
    rpr = sub(s, "w:rPr")
    sub(rpr, "w:color", **{"w:val": "auto", "w:themeColor": "madeup"})

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "T2"})

    resolved = resolve_effective_formatting(p)
    assert resolved.partial is True
    # M9: theme loaded but the name is unknown -> no value stored (a bare
    # theme name is not valid hex). The result is still flagged partial.
    assert resolved.color_rgb is None


def test_resolver_not_partial_for_normal_color() -> None:
    """A doc with an explicit hex color and a valid theme is not partial."""
    doc = Document()
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "H"})
    sub(s, "w:name", **{"w:val": "H"})
    rpr = sub(s, "w:rPr")
    sub(rpr, "w:color", **{"w:val": "FF0000"})

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "H"})

    resolved = resolve_effective_formatting(p)
    assert resolved.partial is False
    assert resolved.color_rgb == "FF0000"


def test_color_val_auto_does_not_set_color() -> None:
    """``w:color w:val="auto"`` without themeColor leaves color unset."""
    doc = Document()
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "A"})
    sub(s, "w:name", **{"w:val": "A"})
    rpr = sub(s, "w:rPr")
    sub(rpr, "w:color", **{"w:val": "auto"})

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "A"})

    resolved = resolve_effective_formatting(p)
    assert resolved.color_rgb is None


# --------------------------------------------------------------------------
# M9 / M10 / M13: precise partial semantics.
# --------------------------------------------------------------------------


def test_resolver_theme_color_none_is_not_partial() -> None:
    """M9: themeColor="none" is an explicit no-color, not a failed resolution."""
    doc = Document()
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "N"})
    sub(s, "w:name", **{"w:val": "N"})
    rpr = sub(s, "w:rPr")
    sub(rpr, "w:color", **{"w:themeColor": "none"})

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "N"})

    resolved = resolve_effective_formatting(p)
    assert resolved.color_rgb is None
    assert resolved.partial is False


def test_resolver_font_token_partial_when_theme_missing() -> None:
    """M10: an unresolvable font token surfaces the token and flags partial."""
    doc = Document()
    _strip_theme(doc)
    p = doc.add_paragraph("text")  # docDefaults uses w:asciiTheme="minorHAnsi"

    resolved = resolve_effective_formatting(p)
    assert resolved.font_name == "minorHAnsi"  # token surfaced for diagnostics
    assert resolved.partial is True


def test_resolver_not_partial_when_no_theme_refs_and_no_theme_part() -> None:
    """M13: a theme-less doc with zero theme references resolves fully."""
    from docx_plus.core.ns import qn

    doc = Document()
    _strip_theme(doc)

    # Neutralise the default docDefaults font *token* so the cascade carries
    # no theme reference at all — without this the (legitimately) missing
    # theme would make font resolution partial.
    rfonts_path = "/".join(qn(t) for t in ("w:docDefaults", "w:rPrDefault", "w:rPr", "w:rFonts"))
    rfonts = doc.styles.element.find(rfonts_path)
    assert rfonts is not None
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]
    rfonts.set(qn("w:ascii"), "Arial")

    p = doc.add_paragraph("text")
    resolved = resolve_effective_formatting(p)
    assert resolved.font_name == "Arial"
    assert resolved.partial is False
