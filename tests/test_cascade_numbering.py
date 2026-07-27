"""Tests for cascade layer 4 (numbering) in :func:`resolve_effective_formatting`.

The ``numbered`` fixture has one paragraph whose ``w:numPr`` references a
custom ``abstractNum`` carrying both pPr (indent 720 left, -360 first-line)
and rPr (bold) at ``lvl[0]``. These tests verify the resolver picks all of
that up, attributes it to the ``numbering`` layer, and degrades gracefully
when the references can't be followed.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.paragraph import Paragraph
from lxml import etree

from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub
from docx_plus.styles import resolve_effective_formatting


def test_numbering_layer_resolves_indent_and_bold(numbered_docx_path: Path) -> None:
    """Both pPr (indent) and rPr (bold) on lvl[0] flow into ResolvedFormatting."""
    doc = Document(str(numbered_docx_path))
    resolved = resolve_effective_formatting(doc.paragraphs[0])
    assert resolved.num_id == 100
    assert resolved.num_level == 0
    assert resolved.indent_left == 720
    assert resolved.indent_first_line == -360
    assert resolved.bold is True


def test_numbering_provenance_marks_layer(numbered_docx_path: Path) -> None:
    """Provenance attributes the indent + bold to the ``numbering`` layer."""
    doc = Document(str(numbered_docx_path))
    resolved = resolve_effective_formatting(doc.paragraphs[0], include_provenance=True)
    prov = resolved.provenance or {}
    assert prov["indent_left"].layer == "numbering"
    assert prov["bold"].layer == "numbering"
    assert prov["num_id"].layer == "numbering"


def test_numbering_absent_part_does_not_crash() -> None:
    """Regression: the resolver used to raise a bare ``NotImplementedError``.

    ``doc.part.numbering_part`` fabricates a missing part through
    ``NumberingPart.new()``, an unimplemented stub in python-docx 1.2.0.
    ``_numbering_root`` reached it via ``getattr(..., None)``, which
    swallows only ``AttributeError``, so any document carrying a
    ``w:numPr`` with no ``numbering.xml`` — LibreOffice, Pandoc, and
    stripped templates all produce these — crashed here.
    """
    doc = Document()
    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype == RT.NUMBERING:
            doc.part.drop_rel(rid)

    para = doc.add_paragraph("item")
    ppr = sub(para._p, "w:pPr")
    num_pr = sub(ppr, "w:numPr")
    sub(num_pr, "w:ilvl", **{"w:val": "0"})
    sub(num_pr, "w:numId", **{"w:val": "3"})

    resolved = resolve_effective_formatting(para)

    # The reference is still reported; only the formatting the level would
    # have contributed is missing.
    assert resolved.num_id == 3
    assert resolved.num_level == 0


def test_numbering_unknown_num_id_skips_silently(numbered_docx_path: Path) -> None:
    """A numPr that references an absent num is non-fatal — the layer just no-ops."""
    doc = Document(str(numbered_docx_path))
    para = doc.paragraphs[0]
    num_pr = para._p.find(f"./{qn('w:pPr')}/{qn('w:numPr')}")
    assert num_pr is not None
    num_id_el = num_pr.find(qn("w:numId"))
    assert num_id_el is not None
    num_id_el.set(qn("w:val"), "9999")  # not in numbering.xml
    resolved = resolve_effective_formatting(para)
    # num_id is still recorded from the paragraph's numPr, but no pPr/rPr
    # flows in because the resolver can't find the abstractNum.
    assert resolved.num_id == 9999
    assert resolved.bold is None


def test_numbering_unparsable_numid_is_ignored(numbered_docx_path: Path) -> None:
    """A non-numeric w:numId/@w:val short-circuits cleanly without raising."""
    doc = Document(str(numbered_docx_path))
    para = doc.paragraphs[0]
    num_pr = para._p.find(f"./{qn('w:pPr')}/{qn('w:numPr')}")
    assert num_pr is not None
    num_id_el = num_pr.find(qn("w:numId"))
    assert num_id_el is not None
    num_id_el.set(qn("w:val"), "not-a-number")
    resolved = resolve_effective_formatting(para)
    assert resolved.num_id is None
    assert resolved.bold is None


def test_numbering_unparsable_ilvl_falls_back_to_zero(
    numbered_docx_path: Path,
) -> None:
    """A bogus ilvl falls back to level 0 rather than raising."""
    doc = Document(str(numbered_docx_path))
    para = doc.paragraphs[0]
    num_pr = para._p.find(f"./{qn('w:pPr')}/{qn('w:numPr')}")
    assert num_pr is not None
    ilvl_el = num_pr.find(qn("w:ilvl"))
    assert ilvl_el is not None
    ilvl_el.set(qn("w:val"), "garbage")
    resolved = resolve_effective_formatting(para)
    # Level fallback kicked in; the indent + bold at lvl[0] still apply.
    assert resolved.num_level == 0
    assert resolved.bold is True


def test_numbering_missing_numid_attribute_short_circuits(
    numbered_docx_path: Path,
) -> None:
    """w:numId without a w:val attribute exits early without setting num_id."""
    doc = Document(str(numbered_docx_path))
    para = doc.paragraphs[0]
    num_pr = para._p.find(f"./{qn('w:pPr')}/{qn('w:numPr')}")
    assert num_pr is not None
    num_id_el = num_pr.find(qn("w:numId"))
    assert num_id_el is not None
    # Remove the val attribute entirely.
    del num_id_el.attrib[qn("w:val")]
    resolved = resolve_effective_formatting(para)
    assert resolved.num_id is None


def test_numbering_paragraph_without_numpr_skips_layer(
    multistyle_docx_path: Path,
) -> None:
    """A plain paragraph (no w:numPr) doesn't trigger the numbering layer."""
    doc = Document(str(multistyle_docx_path))
    resolved = resolve_effective_formatting(doc.paragraphs[0], include_provenance=True)
    assert resolved.num_id is None
    prov = resolved.provenance or {}
    # No numbering-layer entries appear.
    assert not any(src.layer == "numbering" for src in prov.values())


# --------------------------------------------------------------------------
# Style-supplied numbering.
#
# Layer 4 used to read only the paragraph's *direct* w:numPr, so a correctly
# styled list paragraph reported num_id=None — indistinguishable from one
# where a bullet glyph was typed by hand, and a break in the contract every
# other ResolvedFormatting field keeps (they all walk the style chain).
# --------------------------------------------------------------------------


def _ppr(parent: etree._Element) -> etree._Element:
    """Return ``parent``'s ``w:pPr``, creating it if absent."""
    existing = parent.find(qn("w:pPr"))
    return existing if existing is not None else sub(parent, "w:pPr")


def _set_num_pr(parent: etree._Element, *, num_id: str | None, ilvl: str | None) -> None:
    """Write a ``w:pPr/w:numPr`` onto a style or paragraph element."""
    num_pr = sub(_ppr(parent), "w:numPr")
    if ilvl is not None:
        sub(num_pr, "w:ilvl", **{"w:val": ilvl})
    if num_id is not None:
        sub(num_pr, "w:numId", **{"w:val": num_id})


def _add_style(
    doc: Document,
    style_id: str,
    *,
    based_on: str | None = None,
    num_id: str | None = None,
) -> None:
    """Append a paragraph style to styles.xml, optionally carrying numbering.

    Built as raw lxml rather than through ``create_style`` because linking
    a numbering definition into a style is a separate, still-unshipped
    writer; the resolver reads the XML either way.
    """
    style_el = sub(doc.styles.element, "w:style", **{"w:type": "paragraph", "w:styleId": style_id})
    sub(style_el, "w:name", **{"w:val": style_id})
    if based_on is not None:
        sub(style_el, "w:basedOn", **{"w:val": based_on})
    if num_id is not None:
        _set_num_pr(style_el, num_id=num_id, ilvl=None)


def _apply_style(para: Paragraph, style_id: str) -> None:
    """Point a paragraph at a style by ``w:pStyle``.

    Deliberately not ``para.style = doc.styles[...]``: python-docx's style
    factory needs its own ``CT_Style`` class, which hand-built elements
    aren't. The resolver reads ``w:pStyle`` regardless.
    """
    sub(_ppr(para._p), "w:pStyle", **{"w:val": style_id})


def _set_paragraph_numbering(
    para: Paragraph, *, num_id: str | None, ilvl: str | None = None
) -> None:
    """Give a paragraph a direct w:numPr."""
    _set_num_pr(para._p, num_id=num_id, ilvl=ilvl)


def test_style_supplied_numbering_resolves() -> None:
    """A stock ``List Bullet`` paragraph reports the numId its style supplies.

    Regression for the documented gap: this returned ``None`` before, even
    though the bundled template links ``numId`` 1 on that style.
    """
    doc = Document()
    para = doc.add_paragraph("bulleted", style="List Bullet")
    resolved = resolve_effective_formatting(para, include_provenance=True)

    assert resolved.num_id == 1
    assert resolved.num_level == 0
    prov = resolved.provenance or {}
    assert prov["num_id"].layer == "styleNumbering"
    assert prov["num_id"].style_id == "ListBullet"
    assert prov["num_id"].chain_depth == 0


def test_style_supplied_numbering_applies_level_formatting() -> None:
    """The abstractNum level's own pPr flows in for a style-supplied reference too.

    Reaching the numbering definition at all is what was previously
    impossible, so the indent it contributes never applied either.
    """
    doc = Document()
    para = doc.add_paragraph("bulleted", style="List Bullet")
    resolved = resolve_effective_formatting(para, include_provenance=True)

    assert resolved.indent_left == 360
    prov = resolved.provenance or {}
    # The *level's* formatting is the numbering layer regardless of how the
    # reference was reached; only the reference itself is style-attributed.
    assert prov["indent_left"].layer == "numbering"


def test_direct_numpr_overrides_style_supplied() -> None:
    """A paragraph's own numId wins over the one its style supplies."""
    doc = Document()
    para = doc.add_paragraph("item", style="List Bullet")
    _set_paragraph_numbering(para, num_id="77")

    resolved = resolve_effective_formatting(para, include_provenance=True)

    assert resolved.num_id == 77
    prov = resolved.provenance or {}
    assert prov["num_id"].layer == "numbering"
    assert prov["num_id"].style_id is None


def test_direct_ilvl_merges_with_style_supplied_numid() -> None:
    """numId and ilvl resolve independently — a demoted item keeps its style's list.

    Both w:numPr children are optional per ECMA-376 17.3.1.19, and the spec
    does not state merge semantics across the style / direct boundary.
    Verified against Word 2016: a ``List Bullet`` paragraph carrying a bare
    ``<w:ilvl w:val="2"/>`` renders as a third-level bullet of the style's
    own list (marker U+F0B7), not as unnumbered body text.
    """
    doc = Document()
    para = doc.add_paragraph("sub-item", style="List Bullet")
    _set_paragraph_numbering(para, num_id=None, ilvl="2")

    resolved = resolve_effective_formatting(para, include_provenance=True)

    assert resolved.num_id == 1  # inherited from ListBullet
    assert resolved.num_level == 2  # overridden directly
    prov = resolved.provenance or {}
    assert prov["num_id"].layer == "styleNumbering"
    assert prov["num_level"].layer == "numbering"


def test_style_numbering_walks_basedon_chain() -> None:
    """Numbering is inherited through basedOn like every other pPr property."""
    doc = Document()
    _add_style(doc, "BaseList", num_id="42")
    _add_style(doc, "DerivedList", based_on="BaseList")

    para = doc.add_paragraph("item")
    _apply_style(para, "DerivedList")

    resolved = resolve_effective_formatting(para, include_provenance=True)

    assert resolved.num_id == 42
    prov = resolved.provenance or {}
    assert prov["num_id"].style_id == "BaseList"
    assert prov["num_id"].chain_depth == 1


def test_nearest_style_in_chain_wins() -> None:
    """The most specific style supplying a numId wins, matching every other property."""
    doc = Document()
    _add_style(doc, "OuterList", num_id="10")
    _add_style(doc, "InnerList", based_on="OuterList", num_id="20")

    para = doc.add_paragraph("item")
    _apply_style(para, "InnerList")

    resolved = resolve_effective_formatting(para, include_provenance=True)

    assert resolved.num_id == 20
    prov = resolved.provenance or {}
    assert prov["num_id"].style_id == "InnerList"
    assert prov["num_id"].chain_depth == 0


def test_numid_zero_sentinel_suppresses_style_numbering() -> None:
    """A direct numId=0 is surfaced as 0, not flattened to None.

    ECMA-376 17.9.18: zero is the "explicitly not numbered" sentinel and the
    only way to opt out of numbering a style applies. Reporting it faithfully
    is what lets a caller tell deliberate suppression from a paragraph that
    was never numbered — and it must not drag the level's indent in with it.
    """
    doc = Document()
    para = doc.add_paragraph("opted out", style="List Bullet")
    _set_paragraph_numbering(para, num_id="0")

    resolved = resolve_effective_formatting(para, include_provenance=True)

    assert resolved.num_id == 0
    prov = resolved.provenance or {}
    assert prov["num_id"].layer == "numbering"
    # The suppressed style's level formatting must not leak through.
    assert resolved.indent_left != 360


def test_style_supplied_numid_zero_also_surfaces() -> None:
    """The sentinel resolves through the style chain too."""
    doc = Document()
    _add_style(doc, "NoNumberList", num_id="0")

    para = doc.add_paragraph("item")
    _apply_style(para, "NoNumberList")

    resolved = resolve_effective_formatting(para, include_provenance=True)

    assert resolved.num_id == 0
    prov = resolved.provenance or {}
    assert prov["num_id"].layer == "styleNumbering"


def test_ilvl_with_no_numid_anywhere_resolves_nothing() -> None:
    """An ilvl with no numId behind it references nothing and is dropped."""
    doc = Document()
    para = doc.add_paragraph("orphan")
    _set_paragraph_numbering(para, num_id=None, ilvl="3")

    resolved = resolve_effective_formatting(para)

    assert resolved.num_id is None
    assert resolved.num_level is None


def test_unstyled_paragraph_reports_no_numbering() -> None:
    """The common case stays unaffected: no style numbering, no numPr, no result."""
    doc = Document()
    resolved = resolve_effective_formatting(doc.add_paragraph("plain"), include_provenance=True)

    assert resolved.num_id is None
    prov = resolved.provenance or {}
    assert not any(src.layer in ("numbering", "styleNumbering") for src in prov.values())
