"""Tests for ``docx_plus.layout.enable_distinct_even_odd_headers`` /
``disable_distinct_even_odd_headers``."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.layout import (
    disable_distinct_even_odd_headers,
    enable_distinct_even_odd_headers,
)


def test_enable_writes_even_and_odd_headers() -> None:
    doc = Document()
    enable_distinct_even_odd_headers(doc)
    element = doc.settings.element.find(qn("w:evenAndOddHeaders"))
    assert element is not None


def test_enable_is_idempotent() -> None:
    doc = Document()
    enable_distinct_even_odd_headers(doc)
    enable_distinct_even_odd_headers(doc)
    matches = xpath(doc.settings.element, "./w:evenAndOddHeaders")
    assert len(matches) == 1


def test_disable_removes_element() -> None:
    doc = Document()
    enable_distinct_even_odd_headers(doc)
    disable_distinct_even_odd_headers(doc)
    assert doc.settings.element.find(qn("w:evenAndOddHeaders")) is None


def test_disable_is_idempotent_when_absent() -> None:
    doc = Document()
    disable_distinct_even_odd_headers(doc)  # no-op
    disable_distinct_even_odd_headers(doc)


def test_enable_lands_before_w_compat() -> None:
    """python-docx's default settings.xml already has ``w:compat``;
    a blank-doc enable should place the new element before it."""
    doc = Document()
    enable_distinct_even_odd_headers(doc)
    compat = doc.settings.element.find(qn("w:compat"))
    assert compat is not None
    prev = compat.getprevious()
    while prev is not None and prev.tag != qn("w:evenAndOddHeaders"):
        prev = prev.getprevious()
    assert prev is not None
    assert prev.tag == qn("w:evenAndOddHeaders")


def test_enable_appends_when_no_anchor() -> None:
    """With no later-sibling anchor, fall back to append."""
    doc = Document()
    settings = doc.settings.element
    # Strip every recognised later-sibling.
    from docx_plus.layout.settings import _EVEN_AND_ODD_HEADERS_LATER_SIBLINGS

    for child in list(settings):
        if child.tag.startswith(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        ):
            local = child.tag.rpartition("}")[2]
            if f"w:{local}" in _EVEN_AND_ODD_HEADERS_LATER_SIBLINGS:
                settings.remove(child)

    enable_distinct_even_odd_headers(doc)
    new_element = settings.find(qn("w:evenAndOddHeaders"))
    assert new_element is not None
    assert list(settings)[-1] is new_element


def test_enable_round_trip(tmp_path: Path) -> None:
    doc = Document()
    enable_distinct_even_odd_headers(doc)
    out = tmp_path / "eo.docx"
    doc.save(str(out))
    reopened = Document(str(out))
    assert reopened.settings.element.find(qn("w:evenAndOddHeaders")) is not None
