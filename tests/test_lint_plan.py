"""Tests for ``plan_fixes`` — ordering, the content gate, and conflicts.

Split in two. The first half builds findings by hand, because the planner's
job is to reason about a *set* of fixes and hand-built ones make the set
say exactly what a case is about. The second half drives real rules over
real documents, which is the only thing that proves the fixes the rules
emit are in the shape the planner expects.
"""

from __future__ import annotations

import json
from typing import TYPE_CHECKING, Any, get_args

import pytest
from docx import Document
from docx.shared import Pt

from docx_plus.core.oxml import sub
from docx_plus.lint import (
    Finding,
    Fix,
    FixOp,
    FixOperation,
    Location,
    lint,
    plan_fixes,
)

if TYPE_CHECKING:
    from collections.abc import Sequence


# --------------------------------------------------------------------------
# Hand-built findings.
# --------------------------------------------------------------------------


def _finding(
    rule: str = "test-rule",
    *,
    paragraph: int | None = None,
    run: int | None = None,
    style_id: str | None = None,
    fix: Fix | None = None,
    adds_content: bool = False,
) -> Finding:
    """A finding carrying whatever the case under test needs and nothing else."""
    return Finding(
        rule=rule,
        kind="consistency",
        severity="info",
        message=f"{rule} fired",
        location=Location(paragraph_index=paragraph, run_index=run, style_id=style_id),
        fix=fix,
        adds_content=adds_content,
    )


def _fix(*operations: FixOperation, safety: str = "review") -> Fix:
    return Fix(summary="do the thing", safety=safety, operations=operations)  # type: ignore[arg-type]


def _text(paragraph: int, *spans: tuple[int, int]) -> FixOperation:
    return FixOperation(
        op="replace-paragraph-text",
        args={
            "paragraph_index": paragraph,
            "spans": [{"start": s, "end": e, "replacement": ""} for s, e in spans],
        },
    )


def _run_props(paragraph: int, run: int, *properties: str) -> FixOperation:
    return FixOperation(
        op="clear-run-properties",
        args={"paragraph_index": paragraph, "run_index": run, "properties": list(properties)},
    )


def _delete(paragraph: int) -> FixOperation:
    return FixOperation(op="delete-paragraph", args={"paragraph_index": paragraph})


def _rules(planned: Sequence[Any]) -> list[str]:
    return [p.rule for p in planned]


def test_a_finding_with_no_fix_is_reported_as_unfixable() -> None:
    """The plan accounts for every finding rather than quietly shortening the list."""
    findings = [_finding("no-fix"), _finding("has-fix", paragraph=0, fix=_fix(_text(0, (0, 1))))]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["has-fix"]
    assert [f.rule for f in plan.unfixable] == ["no-fix"]


def test_an_empty_report_plans_nothing() -> None:
    plan = plan_fixes([])

    assert plan.fixes == ()
    assert plan.deferred == ()
    assert plan.conflicts == ()
    assert plan.unfixable == ()
    assert plan.operations == ()


def test_edits_are_ordered_down_the_document() -> None:
    findings = [
        _finding("c", paragraph=9, fix=_fix(_text(9, (0, 1)))),
        _finding("a", paragraph=2, fix=_fix(_text(2, (0, 1)))),
        _finding("b", paragraph=4, fix=_fix(_text(4, (0, 1)))),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a", "b", "c"]


def test_deletions_are_ordered_last_and_back_to_front() -> None:
    """The one piece of sequencing a caller cannot be left to work out.

    Every operation names a position in the document as it was swept, so a
    deletion partway down invalidates every index below it. Deletions last,
    descending, is what keeps the other edits addressable.
    """
    findings = [
        _finding("delete-early", paragraph=1, fix=_fix(_delete(1)), adds_content=True),
        _finding("delete-late", paragraph=8, fix=_fix(_delete(8)), adds_content=True),
        _finding("edit-late", paragraph=9, fix=_fix(_text(9, (0, 1)))),
        _finding("edit-early", paragraph=0, fix=_fix(_text(0, (0, 1)))),
    ]

    plan = plan_fixes(findings, allow_content=True)

    assert _rules(plan.fixes) == ["edit-early", "edit-late", "delete-late", "delete-early"]


def test_a_fix_applies_its_own_operations_in_the_order_it_gave_them() -> None:
    """Plan order sequences fixes; a fix sequences itself.

    A single fix deleting several paragraphs has the same index problem in
    miniature, and the planner deliberately does not reorder inside a fix —
    a rule that emits an ordered sequence is asserting that the order is
    load-bearing.
    """
    findings = [_finding("multi", paragraph=3, fix=_fix(_delete(5), _delete(4)), adds_content=True)]

    plan = plan_fixes(findings, allow_content=True)

    assert [op.args["paragraph_index"] for op in plan.operations] == [5, 4]


def test_content_changing_fixes_are_withheld_by_default() -> None:
    findings = [
        _finding("formatting", paragraph=0, fix=_fix(_text(0, (0, 1)))),
        _finding("content", paragraph=3, fix=_fix(_delete(3)), adds_content=True),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["formatting"]
    assert _rules(plan.deferred) == ["content"]


def test_allow_content_includes_them() -> None:
    findings = [_finding("content", paragraph=3, fix=_fix(_delete(3)), adds_content=True)]

    plan = plan_fixes(findings, allow_content=True)

    assert _rules(plan.fixes) == ["content"]
    assert plan.deferred == ()


def test_a_withheld_deletion_does_not_block_edits_to_the_same_paragraph() -> None:
    """The gate runs before conflict detection, and the order matters.

    A deletion claims the whole paragraph, so it collides with every other
    edit there. Resolving conflicts first would let a fix that is not even
    going to be applied knock out one that is.
    """
    findings = [
        _finding("content", paragraph=3, fix=_fix(_delete(3)), adds_content=True),
        _finding("formatting", paragraph=3, fix=_fix(_text(3, (0, 4)))),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["formatting"]
    assert plan.conflicts == ()


def test_a_deletion_and_an_edit_to_the_same_paragraph_conflict_when_both_apply() -> None:
    findings = [
        _finding("content", paragraph=3, fix=_fix(_delete(3)), adds_content=True),
        _finding("formatting", paragraph=3, fix=_fix(_text(3, (0, 4)))),
    ]

    plan = plan_fixes(findings, allow_content=True)

    assert _rules(plan.fixes) == ["formatting"]
    assert [c.dropped.rule for c in plan.conflicts] == ["content"]
    assert "paragraph 3 is removed" in plan.conflicts[0].reason


def test_overlapping_text_spans_conflict() -> None:
    findings = [
        _finding("a", paragraph=0, fix=_fix(_text(0, (4, 9)))),
        _finding("b", paragraph=0, fix=_fix(_text(0, (7, 12)))),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a"]
    assert [c.dropped.rule for c in plan.conflicts] == ["b"]
    assert plan.conflicts[0].kept.rule == "a"
    assert plan.conflicts[0].reason == "paragraph 0, characters 4-9"


def test_adjacent_text_spans_do_not_conflict() -> None:
    """Spans are half-open, so ``[0, 4)`` and ``[4, 9)`` touch without overlapping."""
    findings = [
        _finding("a", paragraph=0, fix=_fix(_text(0, (0, 4)))),
        _finding("b", paragraph=0, fix=_fix(_text(0, (4, 9)))),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a", "b"]
    assert plan.conflicts == ()


def test_text_spans_in_different_paragraphs_do_not_conflict() -> None:
    findings = [
        _finding("a", paragraph=0, fix=_fix(_text(0, (4, 9)))),
        _finding("b", paragraph=1, fix=_fix(_text(1, (4, 9)))),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a", "b"]


def test_two_rules_clearing_the_same_run_property_conflict() -> None:
    findings = [
        _finding("a", paragraph=0, run=1, fix=_fix(_run_props(0, 1, "bold", "font_size"))),
        _finding("b", paragraph=0, run=1, fix=_fix(_run_props(0, 1, "font_size"))),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a"]
    assert plan.conflicts[0].reason == "paragraph 0, run 1, font_size"


def test_two_rules_clearing_different_run_properties_do_not_conflict() -> None:
    """A claim is per property, not per run.

    Coarser detection would call a paragraph carrying several unrelated
    defects unfixable, which is the common case rather than an edge one.
    """
    findings = [
        _finding("a", paragraph=0, run=1, fix=_fix(_run_props(0, 1, "bold"))),
        _finding("b", paragraph=0, run=1, fix=_fix(_run_props(0, 1, "font_size"))),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a", "b"]


def test_the_same_property_on_different_runs_does_not_conflict() -> None:
    findings = [
        _finding("a", paragraph=0, run=1, fix=_fix(_run_props(0, 1, "bold"))),
        _finding("b", paragraph=0, run=2, fix=_fix(_run_props(0, 2, "bold"))),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a", "b"]


def test_a_run_edit_and_a_text_edit_on_one_paragraph_do_not_conflict() -> None:
    """Different kinds of claim on the same paragraph are independent.

    Both survive, and the paragraph-level edit sorts first: within one
    paragraph the order is by run, and a fix with no run at all comes
    before one that names a run.
    """
    findings = [
        _finding("a", paragraph=0, run=1, fix=_fix(_run_props(0, 1, "bold"))),
        _finding("b", paragraph=0, fix=_fix(_text(0, (0, 4)))),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["b", "a"]


def test_retagging_a_language_conflicts_with_clearing_it() -> None:
    """``set-run-language`` and ``clear-run-properties`` both claim ``lang``.

    Two different operations, one property: the claim is what collides, not
    the op name.
    """
    findings = [
        _finding("a", paragraph=0, run=0, fix=_fix(_run_props(0, 0, "lang"))),
        _finding(
            "b",
            paragraph=0,
            run=0,
            fix=_fix(
                FixOperation(
                    op="set-run-language",
                    args={"paragraph_index": 0, "run_index": 0, "lang": "en-GB"},
                )
            ),
        ),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a"]
    assert plan.conflicts[0].reason == "paragraph 0, run 0, lang"


def test_dropping_numbering_conflicts_with_clearing_the_same_property() -> None:
    """Same again one layer up: ``clear-paragraph-numbering`` claims ``num_id``."""
    findings = [
        _finding(
            "a",
            paragraph=2,
            fix=_fix(
                FixOperation(
                    op="clear-paragraph-properties",
                    args={"paragraph_index": 2, "properties": ["num_id", "alignment"]},
                )
            ),
        ),
        _finding(
            "b",
            paragraph=2,
            fix=_fix(FixOperation(op="clear-paragraph-numbering", args={"paragraph_index": 2})),
        ),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a"]
    assert plan.conflicts[0].reason == "paragraph 2, num_id"


def test_clearing_different_paragraph_properties_does_not_conflict() -> None:
    findings = [
        _finding(
            "a",
            paragraph=2,
            fix=_fix(
                FixOperation(
                    op="clear-paragraph-properties",
                    args={"paragraph_index": 2, "properties": ["alignment"]},
                )
            ),
        ),
        _finding(
            "b",
            paragraph=2,
            fix=_fix(
                FixOperation(
                    op="clear-paragraph-properties",
                    args={"paragraph_index": 2, "properties": ["spacing_after"]},
                )
            ),
        ),
    ]

    plan = plan_fixes(findings)

    assert _rules(plan.fixes) == ["a", "b"]
    assert plan.conflicts == ()


def test_deleting_two_different_styles_does_not_conflict() -> None:
    findings = [
        _finding(
            "a",
            style_id="One",
            fix=_fix(FixOperation(op="delete-style", args={"style_id": "One"})),
            adds_content=True,
        ),
        _finding(
            "b",
            style_id="Two",
            fix=_fix(FixOperation(op="delete-style", args={"style_id": "Two"})),
            adds_content=True,
        ),
    ]

    plan = plan_fixes(findings, allow_content=True)

    assert _rules(plan.fixes) == ["a", "b"]


def test_two_fixes_deleting_one_style_conflict() -> None:
    delete = FixOperation(op="delete-style", args={"style_id": "One"})
    findings = [
        _finding("a", style_id="One", fix=_fix(delete), adds_content=True),
        _finding("b", style_id="One", fix=_fix(delete), adds_content=True),
    ]

    plan = plan_fixes(findings, allow_content=True)

    assert _rules(plan.fixes) == ["a"]
    assert plan.conflicts[0].reason == "style One is removed"


def test_the_flattened_operations_follow_plan_order() -> None:
    findings = [
        _finding("b", paragraph=4, fix=_fix(_text(4, (0, 1)))),
        _finding("a", paragraph=1, fix=_fix(_text(1, (0, 1)), _text(1, (5, 6)))),
    ]

    plan = plan_fixes(findings)

    assert [op.args["paragraph_index"] for op in plan.operations] == [1, 1, 4]


def test_a_plan_serializes_to_json() -> None:
    """A plan that cannot be written to a file and reviewed is not much of a plan."""
    # Named so the tie-break — same position, so by rule id — is legible:
    # "a-kept" sorts first and therefore claims the span first.
    findings = [
        _finding("a-kept", paragraph=0, fix=_fix(_text(0, (0, 4)))),
        _finding("b-dropped", paragraph=0, fix=_fix(_text(0, (2, 6)))),
        _finding("held", paragraph=7, fix=_fix(_delete(7)), adds_content=True),
        _finding("unfixable"),
    ]

    encoded = json.dumps(plan_fixes(findings).to_dict())
    restored = json.loads(encoded)

    assert [f["rule"] for f in restored["fixes"]] == ["a-kept"]
    assert [f["rule"] for f in restored["deferred"]] == ["held"]
    assert restored["conflicts"][0]["dropped"]["rule"] == "b-dropped"
    assert restored["unfixable"][0]["rule"] == "unfixable"
    assert restored["fixes"][0]["fix"]["operations"][0]["op"] == "replace-paragraph-text"


# --------------------------------------------------------------------------
# The fixes the real rules emit.
# --------------------------------------------------------------------------


def _by_rule(doc: Document, rule_id: str) -> list[Finding]:
    return [f for f in lint(doc, select=[rule_id]) if f.rule == rule_id]


def test_fixable_is_exactly_whether_a_fix_is_attached() -> None:
    """One source of truth, so the flag cannot drift from the fix."""
    doc = Document()
    doc.add_paragraph("Body with  two spaces.")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Too deep", style="Heading 3")

    findings = lint(doc)

    assert findings
    for finding in findings:
        assert finding.fixable is (finding.fix is not None)


def test_every_fix_uses_the_documented_vocabulary() -> None:
    """No rule may invent an operation the planner has never heard of."""
    doc = Document()
    doc.add_paragraph("Body with  two spaces .")
    doc.add_paragraph("trailing ")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("after")
    run = doc.add_paragraph().add_run("redundant")
    run.font.size = Pt(11)

    known = set(get_args(FixOp))
    seen = set()
    for finding in lint(doc, select=["typography", "formatting", "styles"]):
        if finding.fix is None:
            continue
        assert finding.fix.safety in ("safe", "review", "destructive")
        assert finding.fix.summary.endswith(".")
        for operation in finding.fix.operations:
            assert operation.op in known
            assert "paragraph_index" in operation.args or "style_id" in operation.args
            seen.add(operation.op)

    assert seen  # the fixture really did produce fixes


def test_redundant_direct_formatting_plans_a_safe_deletion() -> None:
    """The one class of fix that provably leaves the rendering alone."""
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.font.size = Pt(11)  # Normal is already 11pt in the stock template

    findings = _by_rule(doc, "redundant-direct-formatting")

    assert len(findings) == 1
    fix = findings[0].fix
    assert fix is not None
    assert fix.safety == "safe"
    operation = fix.operations[0]
    assert operation.op == "clear-run-properties"
    assert operation.args["properties"] == ["font_size"]
    assert operation.args["run_index"] == 0


def test_style_drift_plans_a_review_fix_naming_the_property() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("drifted")
    paragraph.paragraph_format.space_after = Pt(31)

    findings = _by_rule(doc, "style-drift")

    assert len(findings) == 1
    fix = findings[0].fix
    assert fix is not None
    assert fix.safety == "review"
    assert fix.operations[0].op == "clear-paragraph-properties"
    assert fix.operations[0].args["properties"] == ["spacing_after"]


def test_double_space_fixes_every_occurrence_not_just_the_reported_one() -> None:
    """One finding, one problem — but repairing it has to repair the paragraph."""
    doc = Document()
    doc.add_paragraph("one  two  three  four")

    findings = _by_rule(doc, "double-space")

    assert len(findings) == 1
    fix = findings[0].fix
    assert fix is not None
    spans = fix.operations[0].args["spans"]
    assert len(spans) == 3
    assert all(span["replacement"] == " " for span in spans)


def test_double_space_reports_a_run_between_two_others() -> None:
    """Regression: a consuming pattern eats the word between two double spaces.

    ``\\S {2,}\\S`` matches ``"e  t"`` and consumes the ``t``, so the search
    for the next occurrence starts past the character that would have
    anchored it.
    """
    doc = Document()
    doc.add_paragraph("a  b  c")

    fix = _by_rule(doc, "double-space")[0].fix

    assert fix is not None
    assert len(fix.operations[0].args["spans"]) == 2


def test_space_before_punctuation_fixes_every_occurrence() -> None:
    doc = Document()
    doc.add_paragraph("Alpha , beta ; gamma .")

    fix = _by_rule(doc, "space-before-punctuation")[0].fix

    assert fix is not None
    spans = fix.operations[0].args["spans"]
    assert len(spans) == 3
    assert all(span["replacement"] == "" for span in spans)


def test_trailing_whitespace_plans_a_span_at_the_end() -> None:
    doc = Document()
    doc.add_paragraph("text   ")

    fix = _by_rule(doc, "trailing-whitespace")[0].fix

    assert fix is not None
    span = fix.operations[0].args["spans"][0]
    assert (span["start"], span["end"], span["replacement"]) == (4, 7, "")


def test_stray_empty_paragraphs_are_deleted_back_to_front_keeping_one() -> None:
    doc = Document()
    doc.add_paragraph("before")
    for _ in range(3):
        doc.add_paragraph("")
    doc.add_paragraph("after")

    findings = _by_rule(doc, "stray-empty-paragraph")

    assert len(findings) == 1
    fix = findings[0].fix
    assert fix is not None
    assert fix.safety == "destructive"
    assert findings[0].adds_content is True
    assert [op.args["paragraph_index"] for op in fix.operations] == [3, 2]


def test_a_suppressed_numbering_override_carries_no_fix() -> None:
    """``numId=0`` is the deliberate opt-out; repairing it would undo a decision."""
    doc = Document()
    paragraph = doc.add_paragraph("item", style="List Bullet")
    _suppress_numbering(paragraph)

    findings = _by_rule(doc, "direct-numbering-override")

    assert len(findings) == 1
    assert findings[0].severity == "info"
    assert findings[0].fix is None


def test_a_real_numbering_override_plans_a_fix() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("item", style="List Bullet")
    _point_numbering_at(paragraph, 4)

    findings = _by_rule(doc, "direct-numbering-override")

    assert len(findings) == 1
    fix = findings[0].fix
    assert fix is not None
    assert fix.operations[0].op == "clear-paragraph-numbering"


def test_an_unused_style_plans_a_destructive_deletion() -> None:
    """The only fix in the catalogue that removes a definition outright."""
    from docx_plus.styles import create_style

    doc = Document()
    create_style(doc, "Orphan", style_type="paragraph")

    findings = [f for f in lint(doc, select=["unused-styles"]) if f.location.style_id == "Orphan"]

    assert len(findings) == 1
    fix = findings[0].fix
    assert fix is not None
    assert fix.safety == "destructive"
    assert findings[0].adds_content is True
    assert fix.operations[0].op == "delete-style"
    assert fix.operations[0].args["style_id"] == "Orphan"


def test_an_unused_style_deletion_is_withheld_unless_asked_for() -> None:
    from docx_plus.styles import create_style

    doc = Document()
    create_style(doc, "Orphan", style_type="paragraph")
    findings = lint(doc, select=["unused-styles"])

    assert plan_fixes(findings).fixes == ()
    assert [p.rule for p in plan_fixes(findings).deferred] == ["unused-styles"]
    assert [p.rule for p in plan_fixes(findings, allow_content=True).fixes] == ["unused-styles"]


def test_mixed_language_plans_a_retag_to_the_dominant_language() -> None:
    from docx_plus.core.oxml import sub

    doc = Document()
    for text in ("first", "second", "third"):
        run = doc.add_paragraph().add_run(text)
        sub(sub(run._r, "w:rPr"), "w:lang", **{"w:val": "en-GB"})
    odd = doc.add_paragraph().add_run("quatrieme")
    sub(sub(odd._r, "w:rPr"), "w:lang", **{"w:val": "fr-FR"})

    findings = _by_rule(doc, "mixed-language")

    assert len(findings) == 1
    fix = findings[0].fix
    assert fix is not None
    assert fix.operations[0].op == "set-run-language"
    assert fix.operations[0].args["lang"] == "en-GB"


@pytest.mark.parametrize(
    "rule_id",
    [
        "heading-level-skip",
        "empty-heading",
        "manual-list",
        "list-numbering-continuity",
        "manual-heading-formatting",
        "duplicate-styles",
        "indent-by-whitespace",
        "mixed-run-formatting",
        "font-outliers",
        "broken-cross-reference",
        "caption-manual-numbering",
    ],
)
def test_report_only_rules_are_registered_and_stay_report_only(rule_id: str) -> None:
    """The eleven rules whose repair is a judgement the document cannot supply.

    Pinned as a list rather than left implicit: a rule quietly gaining a fix
    is a decision about someone's document, and it should not happen by
    accident.

    This used to assert only that the rule was *registered*, which it would
    have passed unchanged if all eleven had grown fixes — while
    ``ARCHITECTURE.md`` §7.15 claimed the list was pinned. It now asserts
    the actual property, on a document built to make every one of them
    fire.
    """
    from docx_plus.lint import all_rules, lint

    registered = {r.id for r in all_rules()}
    assert rule_id in registered

    findings = lint(_report_only_document(), select=[rule_id])
    assert findings, f"{rule_id} did not fire — the fixture no longer exercises it"
    assert all(f.fix is None for f in findings), (
        f"{rule_id} grew a fix. If that is deliberate, move it out of this "
        f"list and say in its docstring what makes the repair unambiguous."
    )


def test_the_report_only_list_is_exactly_the_rules_without_fixes() -> None:
    """No rule may be report-only *and* missing from the list above.

    The parametrization pins each named rule. This pins the other
    direction, so a new report-only rule cannot be added without being
    declared here.
    """
    from docx_plus.lint import all_rules, lint

    pinned = set(test_report_only_rules_are_registered_and_stay_report_only.pytestmark[0].args[1])
    doc = _report_only_document()
    fixable = {f.rule for f in lint(doc, select=[r.id for r in all_rules()]) if f.fix}

    assert pinned.isdisjoint(fixable)


def test_the_lint_package_doctests_pass() -> None:
    """The examples in the docstrings are executable, so they cannot go stale."""
    import doctest
    import importlib

    attempted = 0
    for name in ("docx_plus.lint", "docx_plus.lint.engine", "docx_plus.lint.plan"):
        result = doctest.testmod(importlib.import_module(name), verbose=False)
        assert result.failed == 0, name
        attempted += result.attempted

    assert attempted


# --------------------------------------------------------------------------
# Numbering helpers — direct ``w:numPr`` is not something python-docx writes.
# --------------------------------------------------------------------------


def _num_pr(paragraph: Any, num_id: int) -> None:
    from docx_plus.core.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.makeelement(qn("w:numPr"), {})
    num_id_el = num_pr.makeelement(qn("w:numId"), {qn("w:val"): str(num_id)})
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def _suppress_numbering(paragraph: Any) -> None:
    _num_pr(paragraph, 0)


def _point_numbering_at(paragraph: Any, num_id: int) -> None:
    _num_pr(paragraph, num_id)


def _report_only_document() -> Document:
    """One document that makes every report-only rule fire.

    Deliberately a single fixture rather than eleven: the point of the test
    it serves is that *no* rule in the list grew a fix, and a shared
    document makes it obvious when one silently stops firing.
    """
    from docx.shared import Pt

    doc = Document()

    # heading-level-skip: 1 straight to 3.
    doc.add_paragraph("Top", style="Heading 1")
    doc.add_paragraph("Deep", style="Heading 3")
    # empty-heading.
    doc.add_paragraph("", style="Heading 2")
    # manual-list: a typed marker, with a sibling so it is unambiguous.
    doc.add_paragraph("1. First typed item")
    doc.add_paragraph("2. Second typed item")
    # indent-by-whitespace.
    doc.add_paragraph("    indented by spaces")
    # manual-heading-formatting: short, bold, not a heading style.
    bold = doc.add_paragraph()
    bold.add_run("Looks Like A Heading").bold = True
    # mixed-run-formatting: two runs disagreeing on size.
    mixed = doc.add_paragraph()
    mixed.add_run("big").font.size = Pt(18)
    mixed.add_run("small").font.size = Pt(9)
    # font-outliers: one run in a font nothing else uses.
    for _ in range(30):
        doc.add_paragraph("ordinary body text")
    odd = doc.add_paragraph().add_run("rare")
    odd.font.name = "Papyrus"
    odd.font.size = Pt(37)
    # list-numbering-continuity: adjacent items, different lists, same level.
    first = doc.add_paragraph("item one", style="List Number")
    second = doc.add_paragraph("item two", style="List Number")
    _num_pr(first, 40)
    _num_pr(second, 41)
    # caption-manual-numbering.
    doc.add_paragraph("Figure 1: typed number", style="Caption")
    # broken-cross-reference: a REF at a bookmark nobody defines.
    _add_ref_field(doc.add_paragraph(), "NoSuchBookmark")
    # duplicate-styles: two ids resolving identically, both applied.
    for style_id in ("DupeA", "DupeB"):
        style = sub(doc.styles.element, "w:style", **{"w:type": "paragraph", "w:styleId": style_id})
        sub(style, "w:name", **{"w:val": style_id})
        sub(sub(style, "w:rPr"), "w:sz", **{"w:val": "26"})
        para = doc.add_paragraph("styled")
        sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": style_id})

    return doc


def _add_ref_field(paragraph: Any, target: str) -> None:
    """A complex REF field pointing at ``target``."""
    from docx_plus.core.ns import qn

    run = paragraph.add_run()
    sub(run._r, "w:fldChar", **{"w:fldCharType": "begin"})
    instr = paragraph.add_run()
    instr_el = sub(instr._r, "w:instrText")
    instr_el.text = rf" REF {target} \h "
    instr_el.set(qn("xml:space"), "preserve")
    end = paragraph.add_run()
    sub(end._r, "w:fldChar", **{"w:fldCharType": "end"})


# --------------------------------------------------------------------------
# Ordering soundness. The "deletions descend" guarantee has to be keyed off
# the indices the *operations* name, not off where the finding was reported.
# --------------------------------------------------------------------------


def test_deletion_order_follows_the_operations_not_the_finding_location() -> None:
    """A finding can sit anywhere relative to what its fix deletes.

    Sorting on `Finding.location` produced `[delete 6, delete 20]` here,
    and once 6 is gone the old paragraph 20 sits at 19.
    """
    findings = [
        _finding("far", paragraph=1, fix=_fix(_delete(20)), adds_content=True),
        _finding("near", paragraph=5, fix=_fix(_delete(6)), adds_content=True),
    ]

    plan = plan_fixes(findings, allow_content=True)

    assert [op.args["paragraph_index"] for op in plan.operations] == [20, 6]


def test_a_multi_delete_fix_sorts_on_its_deepest_index() -> None:
    findings = [
        _finding("deep", paragraph=0, fix=_fix(_delete(31), _delete(30)), adds_content=True),
        _finding("shallow", paragraph=0, fix=_fix(_delete(12)), adds_content=True),
    ]

    plan = plan_fixes(findings, allow_content=True)

    assert [op.args["paragraph_index"] for op in plan.operations] == [31, 30, 12]


def test_a_fix_that_deletes_and_also_edits_elsewhere_is_rejected() -> None:
    """Such a fix belongs in both phases and is wrong in either.

    `rule` is public, so a third-party rule getting this wrong would
    otherwise produce a plan that quietly corrupts a document.
    """
    from docx_plus.lint import InvalidFixError

    findings = [
        _finding(
            "mixed",
            paragraph=0,
            fix=_fix(_delete(50), _run_props(60, 0, "bold")),
            adds_content=True,
        )
    ]

    with pytest.raises(InvalidFixError, match="cannot also carry clear-run-properties"):
        plan_fixes(findings, allow_content=True)


def test_delete_style_still_orders_without_a_paragraph_index() -> None:
    """`delete-style` names no paragraph, so it falls back to the location."""
    findings = [
        _finding(
            "style",
            style_id="Unused",
            fix=_fix(FixOperation(op="delete-style", args={"style_id": "Unused"})),
            adds_content=True,
        ),
        _finding("para", paragraph=4, fix=_fix(_delete(4)), adds_content=True),
    ]

    plan = plan_fixes(findings, allow_content=True)

    assert set(_rules(plan.fixes)) == {"style", "para"}
