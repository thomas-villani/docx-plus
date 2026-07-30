"""Paragraph spacing as measured against live Microsoft Word.

Every number in this file came out of Word's own layout, not the spec.
``<w:contextualSpacing>`` suppresses space the cascade declares, so COM's
``ParagraphFormat.SpaceBefore`` cannot answer the question — it reports what
the cascade says, which is exactly what is in dispute. The probes were
exported to PDF (``ExportAsFixedFormat(path, 17)``) and the paragraph
baselines read back with PyMuPDF instead. Each probe paragraph is a single
line at ``w:lineRule="exact"`` 12pt, so the baseline-to-baseline distance is
12pt plus whatever space Word put between the two.

The resolver was blind to ``<w:contextualSpacing>`` entirely before this,
and the measurements turned up a second divergence nobody was looking for:
Word does **not** add one paragraph's space-after to the next one's
space-before.

Re-measuring: ``scratchpad/build_ctxspacing_probe{,2,3,4}.py`` build the
documents and ``read_ctxspacing{,3}_pdf.py`` score them.
"""

from __future__ import annotations

import itertools

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from docx_plus.core.oxml import sub
from docx_plus.styles import resolve_paragraph_spacing

LINE = 240  # twips; w:lineRule="exact" so the line box is exactly this tall

# The probe grid swept space-after on the first paragraph and space-before
# on the second over these, crossed with contextualSpacing on each.
VALUES_PT = (0, 10, 20, 30)


def _pt(twips: int) -> float:
    return twips / 20


def _add_style(doc: Document, style_id: str, *, contextual: bool = False) -> None:
    style = sub(doc.styles.element, "w:style", **{"w:type": "paragraph", "w:styleId": style_id})
    sub(style, "w:name", **{"w:val": style_id})
    ppr = sub(style, "w:pPr")
    sub(
        ppr,
        "w:spacing",
        **{"w:before": "0", "w:after": "0", "w:line": str(LINE), "w:lineRule": "exact"},
    )
    if contextual:
        sub(ppr, "w:contextualSpacing")


def _add_paragraph(
    doc: Document,
    style_id: str,
    *,
    before: int = 0,
    after: int = 0,
    contextual: bool | None = None,
) -> Paragraph:
    """A single-line paragraph with the given direct spacing.

    ``contextual`` of ``None`` declares nothing, so the style decides;
    ``False`` writes an explicit ``w:val="0"``.
    """
    para = doc.add_paragraph()
    para.add_run("x")
    ppr = para._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": style_id})
    sub(
        ppr,
        "w:spacing",
        **{
            "w:before": str(before),
            "w:after": str(after),
            "w:line": str(LINE),
            "w:lineRule": "exact",
        },
    )
    if contextual is not None:
        sub(ppr, "w:contextualSpacing", **({} if contextual else {"w:val": "0"}))
    return para


def _pair(
    *,
    after: int,
    before: int,
    ctx_first: bool,
    ctx_second: bool,
    second_style: str = "Base",
) -> tuple[Paragraph, Paragraph]:
    doc = Document()
    for style_id in {"Base", second_style}:
        _add_style(doc, style_id)
    first = _add_paragraph(doc, "Base", after=after, contextual=ctx_first or None)
    second = _add_paragraph(doc, second_style, before=before, contextual=ctx_second or None)
    return first, second


# --------------------------------------------------------------------------
# The arithmetic: Word tops space-after up to space-before, it does not add.
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("after_pt", "before_pt", "measured_pt"),
    [
        # Measured with no contextualSpacing anywhere. The sum model predicts
        # 40, 40 and 40 for the first three; Word gives the max.
        (10, 30, 30),
        (30, 10, 30),
        (20, 20, 20),
        (0, 30, 30),
        (30, 0, 30),
    ],
)
def test_word_takes_the_larger_of_space_after_and_space_before(
    after_pt: int, before_pt: int, measured_pt: int
) -> None:
    first, second = _pair(
        after=after_pt * 20, before=before_pt * 20, ctx_first=False, ctx_second=False
    )
    assert _pt(resolve_paragraph_spacing(first).space_below) == measured_pt
    assert _pt(resolve_paragraph_spacing(second).space_above) == measured_pt


# --------------------------------------------------------------------------
# The full 4 x 4 x 2 x 2 grid, exactly as Word laid it out.
# --------------------------------------------------------------------------

# Rows are space-after on the first paragraph, columns space-before on the
# second, both in points. One grid per (contextual on first, on second).
MEASURED_GRIDS: dict[tuple[bool, bool], tuple[tuple[int, ...], ...]] = {
    # Neither paragraph is contextual: the plain max.
    (False, False): (
        (0, 10, 20, 30),
        (10, 10, 20, 30),
        (20, 20, 20, 30),
        (30, 30, 30, 30),
    ),
    # Only the second is contextual: its space-before is gone, so all that
    # is left is the first paragraph's space-after.
    (False, True): (
        (0, 0, 0, 0),
        (10, 10, 10, 10),
        (20, 20, 20, 20),
        (30, 30, 30, 30),
    ),
    # Only the first is contextual. This is the quadrant that gives the
    # mechanism away: not max(0, before) but max(0, before - after). The
    # suppressed space-after still sets the level the top-up measures from.
    (True, False): (
        (0, 10, 20, 30),
        (0, 0, 10, 20),
        (0, 0, 0, 10),
        (0, 0, 0, 0),
    ),
    # Both contextual: nothing survives.
    (True, True): (
        (0, 0, 0, 0),
        (0, 0, 0, 0),
        (0, 0, 0, 0),
        (0, 0, 0, 0),
    ),
}


@pytest.mark.parametrize(
    ("ctx_first", "ctx_second", "after_pt", "before_pt"),
    [
        (ctx_first, ctx_second, after_pt, before_pt)
        for ctx_first, ctx_second in MEASURED_GRIDS
        for after_pt, before_pt in itertools.product(VALUES_PT, VALUES_PT)
    ],
)
def test_measured_spacing_grid(
    ctx_first: bool, ctx_second: bool, after_pt: int, before_pt: int
) -> None:
    expected = MEASURED_GRIDS[(ctx_first, ctx_second)][VALUES_PT.index(after_pt)][
        VALUES_PT.index(before_pt)
    ]
    first, second = _pair(
        after=after_pt * 20,
        before=before_pt * 20,
        ctx_first=ctx_first,
        ctx_second=ctx_second,
    )
    assert _pt(resolve_paragraph_spacing(first).space_below) == expected
    assert _pt(resolve_paragraph_spacing(second).space_above) == expected


def test_space_below_and_space_above_agree_across_the_whole_grid() -> None:
    """The gap belongs to the boundary, not to either paragraph."""
    for (ctx_first, ctx_second), grid in MEASURED_GRIDS.items():
        for row, after_pt in enumerate(VALUES_PT):
            for col, before_pt in enumerate(VALUES_PT):
                first, second = _pair(
                    after=after_pt * 20,
                    before=before_pt * 20,
                    ctx_first=ctx_first,
                    ctx_second=ctx_second,
                )
                below = resolve_paragraph_spacing(first).space_below
                above = resolve_paragraph_spacing(second).space_above
                assert below == above == grid[row][col] * 20


def test_the_suppressed_edge_still_anchors_the_top_up() -> None:
    """The single measurement that rules out "zero the edge, then take the max".

    A contextual paragraph with 20pt after, followed by a non-contextual one
    with 30pt before, leaves 10pt. Zeroing the suppressed edge first would
    leave the full 30.
    """
    first, second = _pair(after=400, before=600, ctx_first=True, ctx_second=False)
    assert resolve_paragraph_spacing(first).space_below == 200
    assert resolve_paragraph_spacing(first).after_suppressed is True
    assert resolve_paragraph_spacing(second).before_suppressed is False


# --------------------------------------------------------------------------
# What counts as "the same style", and what counts as adjacent.
# --------------------------------------------------------------------------


def test_a_different_style_id_never_suppresses() -> None:
    """Measured: CtxA followed by CtxB keeps its space, though both declare it."""
    doc = Document()
    _add_style(doc, "CtxA", contextual=True)
    _add_style(doc, "CtxB", contextual=True)
    first = _add_paragraph(doc, "CtxA", after=600)
    second = _add_paragraph(doc, "CtxB", before=600)
    assert resolve_paragraph_spacing(first).space_below == 600
    assert resolve_paragraph_spacing(second).space_above == 600


def test_a_based_on_child_is_a_different_style() -> None:
    """Sharing a basedOn parent is not sharing a style: measured at 30pt kept."""
    doc = Document()
    _add_style(doc, "CtxA", contextual=True)
    child = sub(doc.styles.element, "w:style", **{"w:type": "paragraph", "w:styleId": "CtxChild"})
    sub(child, "w:name", **{"w:val": "CtxChild"})
    sub(child, "w:basedOn", **{"w:val": "CtxA"})
    first = _add_paragraph(doc, "CtxA", after=600)
    second = _add_paragraph(doc, "CtxChild", before=600)
    assert resolve_paragraph_spacing(first).space_below == 600
    # ...but two of the child suppress, so it does inherit the flag itself.
    third = _add_paragraph(doc, "CtxChild", after=600)
    fourth = _add_paragraph(doc, "CtxChild", before=600)
    assert resolve_paragraph_spacing(second).contextual_spacing is True
    assert resolve_paragraph_spacing(third).space_below == 0
    assert resolve_paragraph_spacing(fourth).space_above == 0


@pytest.mark.parametrize(
    ("first_num", "second_num"),
    [
        ((1, 0), (1, 0)),
        ((1, 0), (1, 1)),  # same list, different level
        ((1, 0), (2, 0)),  # different lists entirely
        ((1, 0), None),
        (None, (1, 0)),
    ],
)
def test_numbering_plays_no_part_in_the_same_style_test(
    first_num: tuple[int, int] | None, second_num: tuple[int, int] | None
) -> None:
    """Measured: every one of these still suppresses.

    The rule is styleId identity. Two ``ListParagraph`` paragraphs in
    unrelated lists collapse into each other exactly as two in one list do.
    """
    doc = Document()
    _add_style(doc, "Base", contextual=True)
    first = _add_paragraph(doc, "Base", after=600)
    second = _add_paragraph(doc, "Base", before=600)
    for para, num in ((first, first_num), (second, second_num)):
        if num is None:
            continue
        numpr = sub(para._p.get_or_add_pPr(), "w:numPr")
        sub(numpr, "w:ilvl", **{"w:val": str(num[1])})
        sub(numpr, "w:numId", **{"w:val": str(num[0])})
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0


def test_a_table_between_two_paragraphs_stops_the_suppression() -> None:
    """Measured: the gap was identical with and without contextualSpacing."""
    doc = Document()
    _add_style(doc, "Base", contextual=True)
    first = _add_paragraph(doc, "Base", after=600)
    doc.add_table(rows=1, cols=1)
    second = _add_paragraph(doc, "Base", before=600)
    assert resolve_paragraph_spacing(first).after_suppressed is False
    assert resolve_paragraph_spacing(first).space_below == 600
    assert resolve_paragraph_spacing(second).before_suppressed is False
    assert resolve_paragraph_spacing(second).space_above == 600


def test_two_paragraphs_in_one_cell_suppress_normally() -> None:
    """Measured: 42pt without the flag, 12pt with it."""
    doc = Document()
    _add_style(doc, "Base", contextual=True)
    cell = doc.add_table(rows=1, cols=1).rows[0].cells[0]
    first, second = cell.paragraphs[0], cell.add_paragraph()
    for para, before, after in ((first, 0, 600), (second, 600, 0)):
        ppr = para._p.get_or_add_pPr()
        sub(ppr, "w:pStyle", **{"w:val": "Base"})
        sub(
            ppr,
            "w:spacing",
            **{
                "w:before": str(before),
                "w:after": str(after),
                "w:line": str(LINE),
                "w:lineRule": "exact",
            },
        )
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0


def test_a_content_control_is_transparent() -> None:
    """Measured: an sdt wrapping the neighbour still suppresses (42pt -> 12pt)."""
    doc = Document()
    _add_style(doc, "Base", contextual=True)
    first = _add_paragraph(doc, "Base", after=600)
    second = _add_paragraph(doc, "Base", before=600)
    sdt = sub(doc.element.body, "w:sdt")
    sub(sdt, "w:sdtPr")
    content = sub(sdt, "w:sdtContent")
    second._p.addprevious(sdt)
    content.append(second._p)
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0


def test_a_content_control_between_the_pair_is_transparent() -> None:
    """Measured: 84pt -> 24pt, i.e. the wrapped paragraph is an ordinary neighbour."""
    doc = Document()
    _add_style(doc, "Base", contextual=True)
    first = _add_paragraph(doc, "Base", after=600)
    middle = _add_paragraph(doc, "Base", contextual=False)
    second = _add_paragraph(doc, "Base", before=600)
    sdt = sub(doc.element.body, "w:sdt")
    sub(sdt, "w:sdtPr")
    content = sub(sdt, "w:sdtContent")
    middle._p.addprevious(sdt)
    content.append(middle._p)
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(middle).space_above == 0
    assert resolve_paragraph_spacing(middle).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0


# --------------------------------------------------------------------------
# The flag itself resolves through the ordinary cascade.
# --------------------------------------------------------------------------


def test_the_flag_reaches_a_paragraph_through_doc_defaults() -> None:
    """Measured on a document whose only declaration is in docDefaults."""
    doc = Document()
    _add_style(doc, "Base")
    defaults = doc.styles.element.find(qn("w:docDefaults"))
    sub(defaults.find(qn("w:pPrDefault")).find(qn("w:pPr")), "w:contextualSpacing")
    first = _add_paragraph(doc, "Base", after=600)
    second = _add_paragraph(doc, "Base", before=600)
    assert resolve_paragraph_spacing(first).contextual_spacing is True
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0


@pytest.mark.parametrize(
    ("opted_out", "measured_pt"),
    [
        # 10pt after, 30pt before, both paragraphs styled contextual, one of
        # them turning it back off. Which one decides which edge survives.
        ("first", 10),  # after survives, before is suppressed
        ("second", 20),  # after is suppressed, before tops up from it
    ],
)
def test_an_explicit_off_on_one_paragraph_frees_only_its_own_edge(
    opted_out: str, measured_pt: int
) -> None:
    """Measured: each edge answers to its own paragraph and nothing else."""
    doc = Document()
    _add_style(doc, "Base", contextual=True)
    first = _add_paragraph(doc, "Base", after=200, contextual=opted_out != "first")
    second = _add_paragraph(doc, "Base", before=600, contextual=opted_out != "second")
    spacing = resolve_paragraph_spacing(first)
    assert spacing.after_suppressed is (opted_out != "first")
    assert resolve_paragraph_spacing(second).before_suppressed is (opted_out != "second")
    assert _pt(spacing.space_below) == measured_pt
    assert resolve_paragraph_spacing(second).space_above == spacing.space_below
