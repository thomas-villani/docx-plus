"""Theme colour resolution as measured against live Microsoft Word.

Every expected hex here came out of Word, not out of the spec: a probe
document was built with one run per theme reference, opened in Word, and
exported both as filtered HTML and as PDF. The two exports agree on every
value, which is what makes them trustworthy — Word's COM reports a theme
colour as a theme-*encoded* integer (0xD0 + slot index in the high byte)
and never as the RGB it paints, so `Font.Color` cannot answer this
question at all.

Two things are pinned here:

  * ``<w:clrSchemeMapping>`` — which scheme slot each ``w:themeColor``
    name resolves through. Word honours it; ignoring it made ``text1``
    resolve to black in a document that renders it white.
  * ``themeTint`` / ``themeShade`` arithmetic, to within a known and
    enumerated residual of one unit per channel.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from docx_plus.core.oxml import sub
from docx_plus.styles import resolve_effective_formatting
from docx_plus.styles.theme import apply_theme_shade, apply_theme_tint, load_theme

# The Office 2007 scheme that ships in python-docx's default template.
BASES = {
    "dark1": "000000",
    "light1": "FFFFFF",
    "dark2": "1F497D",
    "light2": "EEECE1",
    "text1": "000000",
    "background1": "FFFFFF",
    "text2": "1F497D",
    "background2": "EEECE1",
    "accent1": "4F81BD",
    "accent2": "C0504D",
    "accent3": "9BBB59",
    "accent4": "8064A2",
    "accent5": "4BACC6",
    "accent6": "F79646",
    "hyperlink": "0000FF",
    "followedHyperlink": "800080",
}

# (name, tint, shade, what Word rendered)
TRANSFORMS: list[tuple[str, str | None, str | None, str]] = [
    ("accent1", "33", None, "DBE5F1"),
    ("accent1", "66", None, "B8CCE4"),
    ("accent1", "99", None, "95B3D7"),
    ("accent1", "BF", None, "7BA0CD"),
    ("accent1", "E6", None, "608CC3"),
    ("accent1", None, "40", "122030"),
    ("accent1", None, "80", "244061"),
    ("accent1", None, "BF", "365F91"),
    ("accent2", "33", None, "F2DBDB"),
    ("accent2", "66", None, "E5B8B7"),
    ("accent2", "99", None, "D99594"),
    ("accent2", "BF", None, "CF7B79"),
    ("accent2", "E6", None, "C6605E"),
    ("accent2", None, "40", "311211"),
    ("accent2", None, "80", "632423"),
    ("accent2", None, "BF", "943634"),
    ("text1", "33", None, "CCCCCC"),
    ("text1", "66", None, "999999"),
    ("text1", "99", None, "666666"),
    ("text1", "BF", None, "404040"),
    ("text1", "E6", None, "191919"),
    ("text1", None, "40", "000000"),
    ("text1", None, "80", "000000"),
    ("text1", None, "BF", "000000"),
    ("background1", "33", None, "FFFFFF"),
    ("background1", "E6", None, "FFFFFF"),
    ("background1", None, "40", "404040"),
    ("background1", None, "80", "808080"),
    ("background1", None, "BF", "BFBFBF"),
    ("dark2", "33", None, "C6D9F1"),
    ("dark2", "66", None, "8DB3E2"),
    ("dark2", "99", None, "548DD4"),
    ("dark2", "BF", None, "3071C3"),
    ("dark2", "E6", None, "265898"),
    ("dark2", None, "40", "07121F"),
    ("dark2", None, "80", "0F243E"),
    ("dark2", None, "BF", "17365D"),
    ("light2", "33", None, "FBFBF8"),
    ("light2", "66", None, "F8F7F2"),
    ("light2", "99", None, "F4F3EC"),
    ("light2", "BF", None, "F2F0E8"),
    ("light2", "E6", None, "EFEDE3"),
    ("light2", None, "40", "4A442A"),
    ("light2", None, "80", "948A54"),
    ("light2", None, "BF", "C4BC96"),
]

# Cases where the resolver lands within one unit on a channel but not on
# Word's exact value. Word's rounding at these boundaries was not
# reverse-engineered; the transform is HSL luminance scaling and the
# residual is imperceptible, but it is real and enumerated rather than
# waved at. Shrinking this list is the definition of progress here —
# nothing should ever be added to it.
KNOWN_OFF_BY_ONE = {
    ("accent1", "E6", None),
    ("accent1", None, "BF"),
    ("accent2", "33", None),
    ("accent2", "66", None),
    ("accent2", "99", None),
    ("accent2", "E6", None),
    ("accent2", None, "BF"),
    ("dark2", "33", None),
    ("dark2", "BF", None),
    ("dark2", "E6", None),
    ("light2", "33", None),
    ("light2", "66", None),
    ("light2", "99", None),
    ("light2", None, "40"),
    ("light2", None, "BF"),
}


def _themed_run(name: str, tint: str | None = None, shade: str | None = None):
    """A document whose single run references a theme colour."""
    doc = Document()
    run = doc.add_paragraph().add_run("x")
    attrs = {"w:val": "000000", "w:themeColor": name}
    if tint:
        attrs["w:themeTint"] = tint
    if shade:
        attrs["w:themeShade"] = shade
    sub(run._r.get_or_add_rPr(), "w:color", **attrs)
    return doc, run


def _channels(value: str) -> tuple[int, ...]:
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


# --------------------------------------------------------------------------
# Base colours, and the clrSchemeMapping that decides which slot they use.
# --------------------------------------------------------------------------


@pytest.mark.parametrize(("name", "expected"), sorted(BASES.items()))
def test_base_theme_colour_matches_word(name: str, expected: str) -> None:
    _, run = _themed_run(name)
    assert resolve_effective_formatting(run).color_rgb == expected


# What Word rendered with t1/bg1, t2/bg2, accent1/accent3 and the two
# hyperlink slots swapped in <w:clrSchemeMapping>.
SWAPPED = {
    "dark1": "000000",  # a direct slot name — never remapped
    "light1": "FFFFFF",
    "dark2": "1F497D",
    "light2": "EEECE1",
    "text1": "FFFFFF",  # ...whereas the semantic names all follow the map
    "background1": "000000",
    "text2": "EEECE1",
    "background2": "1F497D",
    "accent1": "9BBB59",
    "accent2": "C0504D",
    "accent3": "4F81BD",
    "accent4": "8064A2",
    "accent5": "4BACC6",
    "accent6": "F79646",
    "hyperlink": "800080",
    "followedHyperlink": "0000FF",
}

SWAP = {
    "t1": "light1",
    "bg1": "dark1",
    "t2": "light2",
    "bg2": "dark2",
    "accent1": "accent3",
    "accent3": "accent1",
    "hyperlink": "followedHyperlink",
    "followedHyperlink": "hyperlink",
}


@pytest.mark.parametrize(("name", "expected"), sorted(SWAPPED.items()))
def test_clr_scheme_mapping_redirects_theme_colours(name: str, expected: str) -> None:
    """Word resolves w:themeColor through <w:clrSchemeMapping>.

    Only the semantic names follow it. ``dark1`` and friends name a scheme
    slot outright and stay put even in the same document — which is why
    this cannot be modelled as a simple rename of the scheme.
    """
    doc, run = _themed_run(name)
    mapping = doc.settings.element.find(qn("w:clrSchemeMapping"))
    for attr, slot in SWAP.items():
        mapping.set(qn(f"w:{attr}"), slot)

    assert resolve_effective_formatting(run).color_rgb == expected


def test_absent_clr_scheme_mapping_uses_word_defaults() -> None:
    doc, run = _themed_run("text1")
    settings = doc.settings.element
    settings.remove(settings.find(qn("w:clrSchemeMapping")))

    assert resolve_effective_formatting(run).color_rgb == "000000"


def test_partial_clr_scheme_mapping_defaults_the_rest() -> None:
    """An element that names only some slots leaves the others alone."""
    doc, run = _themed_run("text2")
    mapping = doc.settings.element.find(qn("w:clrSchemeMapping"))
    for attr in list(mapping.attrib):
        del mapping.attrib[attr]
    mapping.set(qn("w:t1"), "light1")

    assert resolve_effective_formatting(run).color_rgb == "1F497D"


def test_theme_mapping_is_read_onto_the_loaded_theme() -> None:
    doc, _ = _themed_run("text1")
    doc.settings.element.find(qn("w:clrSchemeMapping")).set(qn("w:t1"), "accent2")
    theme = load_theme(doc)

    assert theme is not None
    assert theme.mapping["t1"] == "accent2"
    assert theme.base("text1") == "C0504D"


# --------------------------------------------------------------------------
# themeTint / themeShade arithmetic.
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("name", "tint", "shade", "expected"),
    TRANSFORMS,
    ids=[f"{n}-{t or ''}{s or ''}" for n, t, s, _ in TRANSFORMS],
)
def test_tint_and_shade_match_word(
    name: str, tint: str | None, shade: str | None, expected: str
) -> None:
    base = BASES[name]
    got = apply_theme_tint(base, tint) if tint else apply_theme_shade(base, shade or "FF")

    if (name, tint, shade) in KNOWN_OFF_BY_ONE:
        deltas = [abs(a - b) for a, b in zip(_channels(got), _channels(expected), strict=True)]
        assert got != expected, "this case now matches exactly — drop it from KNOWN_OFF_BY_ONE"
        assert max(deltas) <= 1, f"{name} {tint or shade}: {got} is more than 1 off {expected}"
    else:
        assert got == expected


def test_tint_of_ff_is_a_no_op() -> None:
    """The identity that the arithmetic has to preserve.

    Truncating the final channel is only safe because the rational
    round-trip is lossless; a float implementation loses this.
    """
    for base in BASES.values():
        assert apply_theme_tint(base, "FF") == base
        assert apply_theme_shade(base, "FF") == base


def test_exact_arithmetic_hits_integer_boundaries() -> None:
    """Black lightened by E6 lands on exactly 25 per channel.

    In binary floating point ``1 - 0xE6/255`` is 0.09803921568627449, and
    255 times that is 24.999999999999996 — one below what Word paints.
    """
    assert apply_theme_tint("000000", "E6") == "191919"
