"""Tests for ``docx_plus.styles.theme``.

Covers the read-only theme color resolution surface:
:func:`load_theme`, :func:`resolve_theme_color`, the :class:`ThemeColors`
container, and the four transform helpers
(``themeTint``/``themeShade``/``lumMod``/``lumOff``).
"""

from __future__ import annotations

import pytest
from docx import Document

from docx_plus.styles.theme import (
    ThemeColors,
    ThemeError,
    apply_lum_mod,
    apply_lum_off,
    apply_theme_shade,
    apply_theme_tint,
    load_theme,
    resolve_theme_color,
)


def test_load_theme_returns_scheme_for_default_doc() -> None:
    theme = load_theme(Document())
    assert theme is not None
    # python-docx ships the Office 2010 default theme; accent1 = 4F81BD.
    assert theme.base("accent1") == "4F81BD"


def test_load_theme_populates_all_standard_scheme_keys() -> None:
    theme = load_theme(Document())
    assert theme is not None
    expected = {
        "dk1", "lt1", "dk2", "lt2",
        "accent1", "accent2", "accent3", "accent4", "accent5", "accent6",
        "hlink", "folHlink",
    }
    assert expected.issubset(theme.scheme.keys())


def test_theme_color_aliases() -> None:
    theme = load_theme(Document())
    assert theme is not None
    # ECMA-376 17.18.97: text1 -> dk1, background1 -> lt1, etc.
    assert theme.base("text1") == theme.base("dark1")
    assert theme.base("background1") == theme.base("light1")
    assert theme.base("text2") == theme.base("dark2")
    assert theme.base("background2") == theme.base("light2")


def test_theme_base_unknown_name_returns_none() -> None:
    theme = load_theme(Document())
    assert theme is not None
    assert theme.base("nosuchcolor") is None


def test_resolve_theme_color_with_none_theme() -> None:
    assert resolve_theme_color(None, "accent1") is None


def test_resolve_theme_color_with_literal_none_name() -> None:
    theme = load_theme(Document())
    assert resolve_theme_color(theme, "none") is None


def test_resolve_theme_color_applies_shade() -> None:
    theme = ThemeColors(scheme={"accent1": "FFFFFF"})
    # shade="80" darkens L by ~half; white (L=1) -> grey (L≈0.502)
    assert resolve_theme_color(theme, "accent1", shade="80") == "808080"


def test_resolve_theme_color_applies_tint() -> None:
    theme = ThemeColors(scheme={"accent1": "000000"})
    # tint="80" lightens L from 0 to ~0.498
    assert resolve_theme_color(theme, "accent1", tint="80") == "7F7F7F"


def test_apply_theme_shade_identity_at_ff() -> None:
    assert apply_theme_shade("4F81BD", "FF") == "4F81BD"


def test_apply_theme_shade_floors_to_black_at_00() -> None:
    assert apply_theme_shade("FF0000", "00") == "000000"


def test_apply_theme_tint_identity_at_ff() -> None:
    assert apply_theme_tint("4F81BD", "FF") == "4F81BD"


def test_apply_theme_tint_ceils_to_white_at_00() -> None:
    assert apply_theme_tint("000000", "00") == "FFFFFF"


def test_apply_theme_shade_known_value() -> None:
    # accent1 = 4F81BD (Office default); themeShade="80" → 254062
    # Verified against Word's rendering of the same input.
    assert apply_theme_shade("4F81BD", "80") == "254062"


def test_apply_lum_mod_halves_lightness() -> None:
    # Red has L=0.5; lumMod=50000 (factor 0.5) → L=0.25 → 800000.
    assert apply_lum_mod("FF0000", 50000) == "800000"


def test_apply_lum_mod_identity_at_100000() -> None:
    assert apply_lum_mod("4F81BD", 100000) == "4F81BD"


def test_apply_lum_off_clamps_at_one() -> None:
    assert apply_lum_off("FFFFFF", 80000) == "FFFFFF"


def test_apply_lum_off_floors_at_zero() -> None:
    # Negative offset clamps to L=0 → black.
    assert apply_lum_off("FF0000", -100000) == "000000"


def test_apply_theme_shade_rejects_non_hex_byte() -> None:
    with pytest.raises(ThemeError):
        apply_theme_shade("FF0000", "not-hex")


def test_apply_theme_shade_rejects_out_of_range_byte() -> None:
    with pytest.raises(ThemeError):
        apply_theme_shade("FF0000", "100")  # 0x100 = 256, out of byte range


def test_apply_theme_shade_rejects_malformed_hex_color() -> None:
    with pytest.raises(ThemeError):
        apply_theme_shade("notacolor", "80")


def test_apply_theme_tint_strips_leading_hash() -> None:
    assert apply_theme_tint("#4F81BD", "FF") == "4F81BD"


def test_themed_fixture_round_trips(themed_docx_path: "object") -> None:
    """The themed.docx fixture's accent1 color is the same as Document() default."""
    from pathlib import Path

    assert isinstance(themed_docx_path, Path)
    theme = load_theme(Document(str(themed_docx_path)))
    assert theme is not None
    # The fixture is built off a fresh Document() so the theme is identical.
    assert theme.base("accent1") == "4F81BD"
