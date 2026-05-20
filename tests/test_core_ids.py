"""Tests for ``docx_plus.core.ids.IdRegistry``."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus.core.ids import DuplicateIdError, IdRegistry
from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub


def test_docx_plus_error_reexport_identity() -> None:
    """L14: DocxPlusError lives in core.errors and is re-exported by core.

    Guards the import-cycle refactor: the package-level name and the
    errors-module name must be the same object, and the id errors must
    still subclass it.
    """
    from docx_plus.core import DocxPlusError as PackageError
    from docx_plus.core.errors import DocxPlusError as ErrorsError
    from docx_plus.core.ids import IdRangeError

    assert PackageError is ErrorsError
    assert issubclass(DuplicateIdError, PackageError)
    assert issubclass(IdRangeError, PackageError)


def test_registry_empty_on_fresh_doc(empty_docx_path: Path) -> None:
    doc = Document(str(empty_docx_path))
    reg = IdRegistry(doc)
    assert reg.issued() == frozenset()


def test_next_returns_unique_positive_31bit() -> None:
    doc = Document()
    reg = IdRegistry(doc)
    issued = {reg.next() for _ in range(100)}
    assert len(issued) == 100
    for value in issued:
        assert 1 <= value <= 2**31 - 1


def test_reserve_records_value() -> None:
    doc = Document()
    reg = IdRegistry(doc)
    assert reg.reserve(42) == 42
    assert 42 in reg.issued()


def test_reserve_rejects_duplicate() -> None:
    doc = Document()
    reg = IdRegistry(doc)
    reg.reserve(7)
    with pytest.raises(DuplicateIdError):
        reg.reserve(7)


def test_duplicate_id_error_is_value_error() -> None:
    doc = Document()
    reg = IdRegistry(doc)
    reg.reserve(1)
    with pytest.raises(ValueError):
        reg.reserve(1)


def test_reserve_rejects_out_of_range() -> None:
    doc = Document()
    reg = IdRegistry(doc)
    with pytest.raises(ValueError):
        reg.reserve(0)
    with pytest.raises(ValueError):
        reg.reserve(-1)
    with pytest.raises(ValueError):
        reg.reserve(2**31)


def test_next_avoids_reserved_values() -> None:
    doc = Document()
    reg = IdRegistry(doc)
    reserved = {reg.reserve(i) for i in range(1, 200)}
    for _ in range(50):
        value = reg.next()
        assert value not in reserved


def test_registry_seeds_from_existing_sdt_ids() -> None:
    doc = Document()
    para = doc.add_paragraph()
    p = para._p
    sdt = sub(p, "w:sdt")
    sdt_pr = sub(sdt, "w:sdtPr")
    sub(sdt_pr, "w:id", **{"w:val": "12345"})
    sub(sdt, "w:sdtContent")

    reg = IdRegistry(doc)
    assert 12345 in reg.issued()

    with pytest.raises(DuplicateIdError):
        reg.reserve(12345)


def test_registry_seeds_ignore_unparseable_ids() -> None:
    doc = Document()
    para = doc.add_paragraph()
    sdt = sub(para._p, "w:sdt")
    sdt_pr = sub(sdt, "w:sdtPr")
    sub(sdt_pr, "w:id", **{"w:val": "not-an-int"})
    sub(sdt, "w:sdtContent")

    reg = IdRegistry(doc)
    assert reg.issued() == frozenset()


def test_registry_ignores_id_on_non_sdt() -> None:
    doc = Document()
    para = doc.add_paragraph()
    bookmark = sub(para._p, "w:bookmarkStart", **{"w:id": "999", "w:name": "x"})
    assert bookmark.get(qn("w:id")) == "999"

    reg = IdRegistry(doc)
    assert 999 not in reg.issued()
