"""Tests for the ``docx-plus`` command-line interface.

Every test drives the public entry point :func:`docx_plus.cli.main` with an
explicit ``argv`` list and inspects the captured stdout/stderr, the returned
exit code, and any output file. Fixtures are built into ``tmp_path`` so each
test controls the exact styles and control tags it exercises.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from docx_plus.cli import main
from docx_plus.comments import add_comment, reply_to_comment, resolve_comment
from docx_plus.controls import FormBuilder
from docx_plus.styles import apply_style, ensure_style


@pytest.fixture
def styled_doc(tmp_path: Path) -> Path:
    """A document with a Title paragraph and a plain body paragraph."""
    doc = Document()
    ensure_style(doc, "Title")
    apply_style(doc.add_paragraph("The Title"), "Title")
    doc.add_paragraph("Body text.")
    path = tmp_path / "styled.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def form_doc(tmp_path: Path) -> Path:
    """A form with text / dropdown / date / checkbox controls."""
    fb = FormBuilder()
    p = fb.doc.add_paragraph("Name: ")
    fb.add_text_control(p, tag="name", alias="Full name", placeholder="Type name")
    p = fb.doc.add_paragraph("Dept: ")
    fb.add_dropdown(p, tag="dept", items=["Eng", "Design"])
    p = fb.doc.add_paragraph("Start: ")
    fb.add_date_picker(p, tag="start")
    p = fb.doc.add_paragraph("Subscribed: ")
    fb.add_checkbox(p, tag="subscribed")
    path = tmp_path / "form.docx"
    fb.save(str(path))
    return path


@pytest.fixture
def commented_doc(tmp_path: Path) -> Path:
    """Two threads: a resolved one with a reply, and an open one."""
    doc = Document()
    first = doc.add_paragraph("First point.")
    second = doc.add_paragraph("Second point.")

    root = add_comment(first, "Needs a citation.", author="Alice", initials="A")
    reply_to_comment(doc, root.comment_id, "Added one.", author="Bob", initials="B")
    resolve_comment(doc, root.comment_id)

    add_comment(second, "Still unclear.", author="Carol", initials="C")

    path = tmp_path / "commented.docx"
    doc.save(str(path))
    return path


# --------------------------------------------------------------------------
# inspect
# --------------------------------------------------------------------------


def test_inspect_text(styled_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["inspect", str(styled_doc)]) == 0
    out = capsys.readouterr().out
    assert '[1] "The Title"' in out
    assert "style: Title" in out
    assert "font_size" in out


def test_inspect_provenance_annotates_layer(
    styled_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    assert main(["inspect", str(styled_doc), "--provenance"]) == 0
    out = capsys.readouterr().out
    assert "<- paragraphStyle: Title" in out


def test_inspect_json(styled_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["inspect", str(styled_doc), "--json", "--provenance"]) == 0
    records = json.loads(capsys.readouterr().out)
    assert records[0]["index"] == 1
    assert records[0]["text"] == "The Title"
    assert records[0]["style_id"] == "Title"
    assert records[0]["fields"]["font_size"] == 26.0
    assert records[0]["provenance"]["font_size"] == "paragraphStyle: Title"


def test_inspect_json_without_provenance_omits_key(
    styled_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    assert main(["inspect", str(styled_doc), "--json"]) == 0
    records = json.loads(capsys.readouterr().out)
    assert "provenance" not in records[0]


def test_inspect_missing_file(capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["inspect", "nope.docx"]) == 1
    assert "not found" in capsys.readouterr().err


# --------------------------------------------------------------------------
# restyle
# --------------------------------------------------------------------------


def test_restyle_writes_output_and_reports_mapping(
    styled_doc: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "restyled.docx"
    assert main(["restyle", str(styled_doc), "--target", "Title", "-o", str(out)]) == 0
    assert out.is_file()
    output = capsys.readouterr().out
    assert f"wrote {out}" in output
    assert "-> Title" in output


def test_restyle_json(styled_doc: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]) -> None:
    out = tmp_path / "restyled.docx"
    code = main(["restyle", str(styled_doc), "--target", "Title", "-o", str(out), "--json"])
    assert code == 0
    mapping = json.loads(capsys.readouterr().out)
    assert mapping == {"Title": "Title"}


def test_restyle_requires_output(styled_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["restyle", str(styled_doc), "--target", "Title"]) == 1
    assert "specify -o/--output or --in-place" in capsys.readouterr().err


def test_restyle_in_place(styled_doc: Path) -> None:
    assert main(["restyle", str(styled_doc), "--target", "Title", "--in-place"]) == 0
    assert styled_doc.is_file()


def test_restyle_create_missing(styled_doc: Path, tmp_path: Path) -> None:
    out = tmp_path / "restyled.docx"
    code = main(
        [
            "restyle",
            str(styled_doc),
            "--target",
            "Heading1",
            "--create-missing",
            "-o",
            str(out),
        ]
    )
    assert code == 0


def test_restyle_bad_map(
    styled_doc: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "restyled.docx"
    code = main(["restyle", str(styled_doc), "--target", "Title", "--map", "bad", "-o", str(out)])
    assert code == 1
    assert "invalid --map" in capsys.readouterr().err


def test_restyle_map_hint(styled_doc: Path, tmp_path: Path) -> None:
    out = tmp_path / "restyled.docx"
    code = main(
        ["restyle", str(styled_doc), "--target", "Title", "--map", "Title=Title", "-o", str(out)]
    )
    assert code == 0


# --------------------------------------------------------------------------
# controls list
# --------------------------------------------------------------------------


def test_controls_list_text(form_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["controls", "list", str(form_doc)]) == 0
    out = capsys.readouterr().out
    assert "name: text" in out
    assert "subscribed: checkbox = False" in out


def test_controls_list_json(form_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["controls", "list", str(form_doc), "--json"]) == 0
    records = json.loads(capsys.readouterr().out)
    by_tag = {r["tag"]: r for r in records}
    assert by_tag["name"]["control_type"] == "text"
    assert by_tag["name"]["is_placeholder"] is True
    assert by_tag["subscribed"]["value"] is False


def test_controls_list_by_alias(form_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["controls", "list", str(form_doc), "--by", "alias"]) == 0
    out = capsys.readouterr().out
    # Only the text control has an alias; the others are skipped.
    assert "Full name:" in out
    assert "subscribed:" not in out


# --------------------------------------------------------------------------
# controls set / clear
# --------------------------------------------------------------------------


def test_controls_set_text(form_doc: Path, tmp_path: Path) -> None:
    out = tmp_path / "set.docx"
    assert (
        main(["controls", "set", str(form_doc), "--tag", "name", "--value", "Ada", "-o", str(out)])
        == 0
    )
    reread = Document(str(out))
    from docx_plus.controls import read_controls

    assert read_controls(reread)["name"].value == "Ada"


def test_controls_set_checkbox(form_doc: Path, tmp_path: Path) -> None:
    out = tmp_path / "set.docx"
    code = main(
        ["controls", "set", str(form_doc), "--tag", "subscribed", "--value", "yes", "-o", str(out)]
    )
    assert code == 0
    from docx_plus.controls import read_controls

    assert read_controls(Document(str(out)))["subscribed"].value is True


def test_controls_set_date(form_doc: Path, tmp_path: Path) -> None:
    out = tmp_path / "set.docx"
    code = main(
        [
            "controls",
            "set",
            str(form_doc),
            "--tag",
            "start",
            "--value",
            "2026-06-15",
            "-o",
            str(out),
        ]
    )
    assert code == 0
    from docx_plus.controls import read_controls

    value = read_controls(Document(str(out)))["start"].value
    assert value is not None
    assert value.year == 2026 and value.month == 6 and value.day == 15


def test_controls_set_bad_checkbox(
    form_doc: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "set.docx"
    code = main(
        [
            "controls",
            "set",
            str(form_doc),
            "--tag",
            "subscribed",
            "--value",
            "maybe",
            "-o",
            str(out),
        ]
    )
    assert code == 1
    assert "checkbox value must be true/false" in capsys.readouterr().err


def test_controls_set_bad_date(
    form_doc: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "set.docx"
    code = main(
        [
            "controls",
            "set",
            str(form_doc),
            "--tag",
            "start",
            "--value",
            "not-a-date",
            "-o",
            str(out),
        ]
    )
    assert code == 1
    assert "date value must be ISO 8601" in capsys.readouterr().err


def test_controls_set_unknown_tag(
    form_doc: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "set.docx"
    code = main(
        ["controls", "set", str(form_doc), "--tag", "ghost", "--value", "x", "-o", str(out)]
    )
    assert code == 1
    assert "no control with tag 'ghost'" in capsys.readouterr().err


def test_controls_clear(form_doc: Path, tmp_path: Path) -> None:
    out = tmp_path / "cleared.docx"
    # First fill it, then clear it.
    main(["controls", "set", str(form_doc), "--tag", "name", "--value", "Ada", "-o", str(out)])
    assert main(["controls", "clear", str(out), "--tag", "name", "--in-place"]) == 0
    from docx_plus.controls import read_controls

    assert read_controls(Document(str(out)))["name"].is_placeholder is True


def test_controls_set_requires_output(form_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    code = main(["controls", "set", str(form_doc), "--tag", "name", "--value", "Ada"])
    assert code == 1
    assert "specify -o/--output or --in-place" in capsys.readouterr().err


# --------------------------------------------------------------------------
# comments
# --------------------------------------------------------------------------


def test_comments_list_text(commented_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["comments", "list", str(commented_doc)]) == 0
    out = capsys.readouterr().out
    assert "Alice [resolved]: Needs a citation." in out
    assert "Bob: Added one." in out
    assert "Carol: Still unclear." in out
    assert "on paragraph 0: 'First point.'" in out


def test_comments_list_nests_replies_under_their_root(
    commented_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    main(["comments", "list", str(commented_doc)])
    lines = capsys.readouterr().out.splitlines()
    root_line = next(i for i, line in enumerate(lines) if "Needs a citation." in line)
    reply_line = next(i for i, line in enumerate(lines) if "Bob: Added one." in line)
    assert reply_line > root_line
    assert lines[reply_line].startswith("  [")
    assert not lines[root_line].startswith(" ")


def test_comments_list_json(commented_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["comments", "list", str(commented_doc), "--json"]) == 0
    payload = json.loads(capsys.readouterr().out)
    assert len(payload) == 2
    resolved = next(thread for thread in payload if thread["resolved"])
    assert resolved["author"] == "Alice"
    assert [reply["author"] for reply in resolved["replies"]] == ["Bob"]


def test_comments_list_unresolved_filter(
    commented_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    assert main(["comments", "list", str(commented_doc), "--unresolved", "--json"]) == 0
    payload = json.loads(capsys.readouterr().out)
    assert [thread["author"] for thread in payload] == ["Carol"]


def test_comments_list_flags_an_orphaned_comment(
    commented_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    from docx_plus.core.oxml import xpath

    doc = Document(str(commented_doc))
    for expr in (".//w:commentRangeStart", ".//w:commentRangeEnd"):
        for marker in xpath(doc.element.body, expr):
            marker.getparent().remove(marker)
    doc.save(str(commented_doc))

    assert main(["comments", "list", str(commented_doc)]) == 0
    assert "(orphaned - no anchor in the document body)" in capsys.readouterr().out


def test_comments_list_on_document_without_comments(
    styled_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    assert main(["comments", "list", str(styled_doc)]) == 0
    assert "(no comments)" in capsys.readouterr().out


def test_comments_resolve(commented_doc: Path, tmp_path: Path) -> None:
    from docx_plus.comments import read_threads

    out = tmp_path / "resolved.docx"
    open_id = next(
        thread.root.comment_id
        for thread in read_threads(Document(str(commented_doc)))
        if not thread.resolved
    )
    assert main(["comments", "resolve", str(commented_doc), str(open_id), "-o", str(out)]) == 0
    assert all(thread.resolved for thread in read_threads(Document(str(out))))


def test_comments_reopen_in_place(commented_doc: Path) -> None:
    from docx_plus.comments import read_threads

    resolved_id = next(
        thread.root.comment_id
        for thread in read_threads(Document(str(commented_doc)))
        if thread.resolved
    )
    code = main(["comments", "reopen", str(commented_doc), str(resolved_id), "--in-place"])
    assert code == 0
    assert not any(thread.resolved for thread in read_threads(Document(str(commented_doc))))


def test_comments_resolve_unknown_id(
    commented_doc: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "nope.docx"
    assert main(["comments", "resolve", str(commented_doc), "999999", "-o", str(out)]) == 1
    assert "CommentNotFoundError" in capsys.readouterr().err
    assert not out.exists()


def test_comments_resolve_requires_output(
    commented_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    assert main(["comments", "resolve", str(commented_doc), "1"]) == 1
    assert "specify -o/--output or --in-place" in capsys.readouterr().err


# --------------------------------------------------------------------------
# top-level dispatch
# --------------------------------------------------------------------------


def test_no_command_prints_help(capsys: pytest.CaptureFixture[str]) -> None:
    assert main([]) == 2
    assert "usage: docx-plus" in capsys.readouterr().out


def test_version(capsys: pytest.CaptureFixture[str]) -> None:
    with pytest.raises(SystemExit) as exc:
        main(["--version"])
    assert exc.value.code == 0
    assert "docx-plus" in capsys.readouterr().out


def test_unknown_command_exits_2() -> None:
    with pytest.raises(SystemExit) as exc:
        main(["frobnicate"])
    assert exc.value.code == 2


# ---------------------------------------------------------------------------
# skill — locate / read / install the packaged agent skill.
# ---------------------------------------------------------------------------


EXPECTED_TOPICS = {
    "cli",
    "comments",
    "forms",
    "layout",
    "numbering",
    "publishing",
    "revisions",
    "styles",
    "tables",
}


class TestSkillPath:
    def test_prints_an_existing_directory(self, capsys: pytest.CaptureFixture[str]) -> None:
        assert main(["skill", "path"]) == 0
        printed = Path(capsys.readouterr().out.strip())
        assert printed.is_dir()
        assert (printed / "SKILL.md").is_file()

    def test_points_inside_the_package(self, capsys: pytest.CaptureFixture[str]) -> None:
        """The skill must resolve from the package, not the repo root."""
        main(["skill", "path"])
        printed = Path(capsys.readouterr().out.strip())
        assert printed.parent.name == "docx_plus"


class TestSkillList:
    def test_lists_every_topic(self, capsys: pytest.CaptureFixture[str]) -> None:
        assert main(["skill", "list"]) == 0
        assert set(capsys.readouterr().out.split()) == EXPECTED_TOPICS

    def test_is_sorted(self, capsys: pytest.CaptureFixture[str]) -> None:
        main(["skill", "list"])
        topics = capsys.readouterr().out.split()
        assert topics == sorted(topics)


class TestSkillShow:
    def test_no_topic_prints_the_entry_point(self, capsys: pytest.CaptureFixture[str]) -> None:
        assert main(["skill", "show"]) == 0
        out = capsys.readouterr().out
        assert out.startswith("---")
        assert "name: docx-plus" in out

    def test_named_topic(self, capsys: pytest.CaptureFixture[str]) -> None:
        assert main(["skill", "show", "tables"]) == 0
        assert "# Tables" in capsys.readouterr().out

    def test_md_suffix_is_tolerated(self, capsys: pytest.CaptureFixture[str]) -> None:
        assert main(["skill", "show", "tables.md"]) == 0
        assert "# Tables" in capsys.readouterr().out

    def test_unknown_topic_lists_the_valid_ones(self, capsys: pytest.CaptureFixture[str]) -> None:
        assert main(["skill", "show", "bogus"]) == 1
        err = capsys.readouterr().err
        assert "unknown topic 'bogus'" in err
        assert "tables" in err


class TestSkillInstall:
    def test_writes_the_whole_tree(
        self, tmp_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        dest = tmp_path / "skills"
        assert main(["skill", "install", "--dest", str(dest)]) == 0
        installed = dest / "docx-plus"
        assert (installed / "SKILL.md").is_file()
        assert {p.stem for p in (installed / "reference").glob("*.md")} == EXPECTED_TOPICS
        assert "installed 10 files" in capsys.readouterr().out

    def test_content_matches_the_package(self, tmp_path: Path) -> None:
        dest = tmp_path / "skills"
        main(["skill", "install", "--dest", str(dest)])
        main(["skill", "show", "numbering"])
        packaged = (
            Path(__file__).resolve().parent.parent
            / "docx_plus"
            / "skill"
            / "reference"
            / "numbering.md"
        )
        assert (dest / "docx-plus" / "reference" / "numbering.md").read_bytes() == (
            packaged.read_bytes()
        )

    def test_refuses_to_overwrite(self, tmp_path: Path, capsys: pytest.CaptureFixture[str]) -> None:
        dest = tmp_path / "skills"
        main(["skill", "install", "--dest", str(dest)])
        capsys.readouterr()
        assert main(["skill", "install", "--dest", str(dest)]) == 1
        assert "pass --force to overwrite" in capsys.readouterr().err

    def test_force_overwrites(self, tmp_path: Path) -> None:
        dest = tmp_path / "skills"
        main(["skill", "install", "--dest", str(dest)])
        stray = dest / "docx-plus" / "stale.md"
        stray.write_text("removed by a forced reinstall", encoding="utf-8")
        assert main(["skill", "install", "--dest", str(dest), "--force"]) == 0
        assert not stray.exists()

    def test_user_and_dest_are_mutually_exclusive(
        self, tmp_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        assert main(["skill", "install", "--user", "--dest", str(tmp_path)]) == 1
        assert "mutually exclusive" in capsys.readouterr().err

    def test_defaults_to_the_project_skills_dir(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.chdir(tmp_path)
        assert main(["skill", "install"]) == 0
        assert (tmp_path / ".claude" / "skills" / "docx-plus" / "SKILL.md").is_file()

    def test_user_flag_targets_the_home_skills_dir(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        from docx_plus.cli import skill as skill_cmd

        monkeypatch.setattr(skill_cmd, "USER_SKILLS_DIR", tmp_path / "home" / "skills")
        assert main(["skill", "install", "--user"]) == 0
        assert (tmp_path / "home" / "skills" / "docx-plus" / "SKILL.md").is_file()


class TestSkillFrontmatter:
    """The frontmatter is what makes an agent discover the skill at all."""

    def test_skill_md_has_name_and_description(self) -> None:
        from docx_plus.cli.skill import _read, _skill_root

        text = _read(_skill_root() / "SKILL.md")
        head = text.split("---")[1]
        assert "name: docx-plus" in head
        assert "description:" in head

    def test_every_topic_is_referenced_from_skill_md(self) -> None:
        from docx_plus.cli.skill import _read, _skill_root, _topics

        text = _read(_skill_root() / "SKILL.md")
        missing = [t for t in _topics() if f"reference/{t}.md" not in text]
        assert not missing, f"topics not linked from SKILL.md: {missing}"


class TestSkillZipimportFallback:
    """A zipped distribution has no on-disk path for `skill path` to print."""

    def test_path_reports_a_clean_error(
        self, capsys: pytest.CaptureFixture[str], monkeypatch: pytest.MonkeyPatch
    ) -> None:
        from importlib.resources import files as resource_files

        from docx_plus.cli import skill as skill_cmd

        zipped = resource_files("docx_plus") / "skill" / "does-not-exist-on-disk"
        monkeypatch.setattr(skill_cmd, "_skill_root", lambda: zipped)
        assert main(["skill", "path"]) == 1
        assert "zipped distribution" in capsys.readouterr().err
