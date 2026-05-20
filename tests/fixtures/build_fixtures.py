"""Generate the .docx fixtures used by the test suite.

Fixtures are generated, not committed (SPEC §10). The **canonical** path is
``conftest.py``, which builds each fixture lazily into a per-session tmp dir.
This module's :func:`main` is a separate, manual inspection helper: it builds
into a fresh temp dir (or a directory you name on the command line) and prints
the paths — it never writes into the committed ``tests/fixtures/`` tree, so a
stray ``python -m`` run cannot leave generated ``.docx`` files in source.

Each builder function is small and explicit so the fixture's contents are
obvious from reading this file.
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

from docx import Document

from docx_plus.core.oxml import sub


def build_empty(path: Path) -> Path:
    """Write a minimal valid ``.docx`` with no body content."""
    doc = Document()
    doc.save(path)
    return path


def build_multistyle(path: Path) -> Path:
    """Write a doc with a three-level paragraph style chain Base -> Mid -> Top.

    The chain is intentionally shallow so toggle and basedOn behavior can be
    verified end-to-end by the Phase 2 cascade resolver. One paragraph is
    styled ``Top`` so the cascade has a real target.
    """
    doc = Document()
    styles_element = doc.styles.element

    base = sub(
        styles_element,
        "w:style",
        **{"w:type": "paragraph", "w:styleId": "Base"},
    )
    sub(base, "w:name", **{"w:val": "Base"})
    base_rpr = sub(base, "w:rPr")
    sub(base_rpr, "w:b")  # toggle: bold ON at Base

    mid = sub(
        styles_element,
        "w:style",
        **{"w:type": "paragraph", "w:styleId": "Mid"},
    )
    sub(mid, "w:name", **{"w:val": "Mid"})
    sub(mid, "w:basedOn", **{"w:val": "Base"})
    mid_rpr = sub(mid, "w:rPr")
    sub(mid_rpr, "w:i")  # italic ON at Mid

    top = sub(
        styles_element,
        "w:style",
        **{"w:type": "paragraph", "w:styleId": "Top"},
    )
    sub(top, "w:name", **{"w:val": "Top"})
    sub(top, "w:basedOn", **{"w:val": "Mid"})
    top_rpr = sub(top, "w:rPr")
    sub(top_rpr, "w:b")  # toggle: bold XOR through chain -> effective FALSE

    para = doc.add_paragraph("Styled by Top, which is basedOn Mid, basedOn Base.")
    para_pr = para._p.get_or_add_pPr()
    sub(para_pr, "w:pStyle", **{"w:val": "Top"})

    doc.save(path)
    return path


def build_themed(path: Path) -> Path:
    """Write a doc whose paragraph style uses a themeColor + themeShade.

    Exercises the Phase 2 theme-resolution path end-to-end: the style
    ``Themed`` sets ``w:color`` with ``themeColor="accent1"`` and a
    ``themeShade`` byte, so the resolver has to fetch ``accent1`` from the
    document's theme part and darken it.
    """
    doc = Document()
    styles_element = doc.styles.element

    themed_style = sub(
        styles_element,
        "w:style",
        **{"w:type": "paragraph", "w:styleId": "Themed"},
    )
    sub(themed_style, "w:name", **{"w:val": "Themed"})
    themed_rpr = sub(themed_style, "w:rPr")
    sub(
        themed_rpr,
        "w:color",
        **{
            "w:val": "auto",
            "w:themeColor": "accent1",
            "w:themeShade": "80",
        },
    )

    para = doc.add_paragraph("Themed paragraph.")
    para_pr = para._p.get_or_add_pPr()
    sub(para_pr, "w:pStyle", **{"w:val": "Themed"})

    doc.save(path)
    return path


def build_existing_form(path: Path) -> Path:
    """Write a doc with three SDTs built by hand, *not* via FormBuilder.

    Used by the read-side tests to verify ``read_controls`` works on
    documents that did not originate from ``docx_plus.controls.builder`` —
    third-party tools, Word itself, or earlier hand-rolled scripts.
    """
    doc = Document()
    para = doc.add_paragraph("Name: ")
    para_two = doc.add_paragraph("Region: ")
    para_three = doc.add_paragraph("Subscribe: ")

    # Filled text SDT (no showingPlcHdr).
    text_sdt = sub(para._p, "w:sdt")
    text_pr = sub(text_sdt, "w:sdtPr")
    sub(text_pr, "w:tag", **{"w:val": "name"})
    sub(text_pr, "w:id", **{"w:val": "100"})
    sub(text_pr, "w:text")
    text_content = sub(text_sdt, "w:sdtContent")
    text_run = sub(text_content, "w:r")
    text_t = sub(text_run, "w:t")
    text_t.text = "Ada Lovelace"

    # Dropdown SDT in placeholder state.
    dd_sdt = sub(para_two._p, "w:sdt")
    dd_pr = sub(dd_sdt, "w:sdtPr")
    sub(dd_pr, "w:alias", **{"w:val": "Region selector"})
    sub(dd_pr, "w:tag", **{"w:val": "region"})
    sub(dd_pr, "w:id", **{"w:val": "200"})
    sub(dd_pr, "w:showingPlcHdr")
    dd_list = sub(dd_pr, "w:dropDownList")
    sub(dd_list, "w:listItem", **{"w:displayText": "Choose a region", "w:value": ""})
    sub(dd_list, "w:listItem", **{"w:displayText": "North", "w:value": "N"})
    sub(dd_list, "w:listItem", **{"w:displayText": "South", "w:value": "S"})
    dd_content = sub(dd_sdt, "w:sdtContent")
    dd_run = sub(dd_content, "w:r")
    dd_run_pr = sub(dd_run, "w:rPr")
    sub(dd_run_pr, "w:rStyle", **{"w:val": "PlaceholderText"})
    dd_t = sub(dd_run, "w:t")
    dd_t.text = "Choose a region"

    # Checkbox SDT, checked.
    cb_sdt = sub(para_three._p, "w:sdt")
    cb_pr = sub(cb_sdt, "w:sdtPr")
    sub(cb_pr, "w:tag", **{"w:val": "subscribe"})
    sub(cb_pr, "w:id", **{"w:val": "300"})
    cb_box = sub(cb_pr, "w14:checkbox")
    sub(cb_box, "w14:checked", **{"w14:val": "1"})
    sub(cb_box, "w14:checkedState", **{"w14:val": "2612", "w14:font": "MS Gothic"})
    sub(cb_box, "w14:uncheckedState", **{"w14:val": "2610", "w14:font": "MS Gothic"})
    cb_content = sub(cb_sdt, "w:sdtContent")
    cb_run = sub(cb_content, "w:r")
    cb_t = sub(cb_run, "w:t")
    cb_t.text = "☒"

    doc.save(path)
    return path


def build_numbered(path: Path) -> Path:
    """Write a doc with one ordered-list paragraph that uses ``w:numPr``.

    Exercises layer 4 of the cascade (numbering). python-docx's default
    template already includes a populated ``numbering.xml`` and registered
    relationship, so we just inject our own ``abstractNum`` (with indent on
    ``pPr`` and bold on ``rPr`` at ``lvl[0]``) and ``num`` into the existing
    part, then point the paragraph at the new ``numId``.
    """
    doc = Document()

    numbering_root = doc.part.numbering_part.element
    # Use IDs above python-docx's default range (the bundled template uses
    # 0-8) to avoid colliding with built-in styles like ListBullet.
    abstract = sub(numbering_root, "w:abstractNum", **{"w:abstractNumId": "100"})
    lvl = sub(abstract, "w:lvl", **{"w:ilvl": "0"})
    sub(lvl, "w:start", **{"w:val": "1"})
    sub(lvl, "w:numFmt", **{"w:val": "decimal"})
    sub(lvl, "w:lvlText", **{"w:val": "%1."})
    lvl_ppr = sub(lvl, "w:pPr")
    sub(lvl_ppr, "w:ind", **{"w:left": "720", "w:hanging": "360"})
    lvl_rpr = sub(lvl, "w:rPr")
    sub(lvl_rpr, "w:b")

    num = sub(numbering_root, "w:num", **{"w:numId": "100"})
    sub(num, "w:abstractNumId", **{"w:val": "100"})

    para = doc.add_paragraph("Numbered item one.")
    para_pr = para._p.get_or_add_pPr()
    num_pr = sub(para_pr, "w:numPr")
    sub(num_pr, "w:ilvl", **{"w:val": "0"})
    sub(num_pr, "w:numId", **{"w:val": "100"})

    doc.save(path)
    return path


def build_all(out_dir: Path) -> dict[str, Path]:
    """Build every fixture into ``out_dir`` (required) and return their paths."""
    out_dir.mkdir(parents=True, exist_ok=True)
    return {
        "empty": build_empty(out_dir / "empty.docx"),
        "multistyle": build_multistyle(out_dir / "multistyle.docx"),
        "themed": build_themed(out_dir / "themed.docx"),
        "existing_form": build_existing_form(out_dir / "existing_form.docx"),
        "numbered": build_numbered(out_dir / "numbered.docx"),
    }


def main(argv: list[str] | None = None) -> int:
    """Manual inspection helper for ``python -m tests.fixtures.build_fixtures``.

    Tests do not use this — they go through ``conftest.py``. Pass an output
    directory to inspect the generated files there; omit it to build into a
    fresh temp dir whose path is printed. Never writes into ``tests/fixtures/``.
    """
    args = sys.argv[1:] if argv is None else argv
    out_dir = Path(args[0]) if args else Path(tempfile.mkdtemp(prefix="docx_plus_fixtures_"))
    built = build_all(out_dir)
    for name, path in built.items():
        print(f"built {name}: {path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
