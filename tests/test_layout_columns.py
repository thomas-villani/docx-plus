"""Tests for ``docx_plus.layout.set_columns``."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.layout import set_columns


def _cols(section):
    return section._sectPr.find(qn("w:cols"))


def test_set_columns_writes_w_cols() -> None:
    doc = Document()
    set_columns(doc.sections[0], 2)
    cols = _cols(doc.sections[0])
    assert cols is not None
    assert cols.get(qn("w:num")) == "2"


def test_set_columns_default_space_is_720() -> None:
    doc = Document()
    set_columns(doc.sections[0], 2)
    assert _cols(doc.sections[0]).get(qn("w:space")) == "720"


def test_set_columns_custom_space() -> None:
    doc = Document()
    set_columns(doc.sections[0], 3, space=1440)
    assert _cols(doc.sections[0]).get(qn("w:space")) == "1440"


def test_set_columns_no_separator_by_default() -> None:
    doc = Document()
    set_columns(doc.sections[0], 2)
    assert _cols(doc.sections[0]).get(qn("w:sep")) is None


def test_set_columns_separator_true() -> None:
    doc = Document()
    set_columns(doc.sections[0], 2, separator=True)
    assert _cols(doc.sections[0]).get(qn("w:sep")) == "1"


def test_set_columns_unequal_widths() -> None:
    doc = Document()
    set_columns(doc.sections[0], 3, space=360, widths=[2000, 3000, 4000])
    cols = _cols(doc.sections[0])
    assert cols.get(qn("w:equalWidth")) == "0"
    children = xpath(cols, "./w:col")
    assert len(children) == 3
    assert [c.get(qn("w:w")) for c in children] == ["2000", "3000", "4000"]
    # First two columns carry trailing space; the last does not.
    assert children[0].get(qn("w:space")) == "360"
    assert children[1].get(qn("w:space")) == "360"
    assert children[2].get(qn("w:space")) is None


def test_set_columns_is_idempotent() -> None:
    doc = Document()
    set_columns(doc.sections[0], 2)
    set_columns(doc.sections[0], 4)  # second call overrides
    cols_list = xpath(doc.sections[0]._sectPr, "./w:cols")
    assert len(cols_list) == 1
    assert cols_list[0].get(qn("w:num")) == "4"


def test_set_columns_rejects_zero() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="num >= 1"):
        set_columns(doc.sections[0], 0)


def test_set_columns_rejects_widths_length_mismatch() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="widths"):
        set_columns(doc.sections[0], 3, widths=[100, 200])


def test_set_columns_round_trip(tmp_path: Path) -> None:
    doc = Document()
    set_columns(doc.sections[0], 2, space=720, separator=True)
    out = tmp_path / "cols.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    cols = _cols(reopened.sections[0])
    assert cols is not None
    assert cols.get(qn("w:num")) == "2"
    assert cols.get(qn("w:sep")) == "1"
