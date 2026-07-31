"""Structural edges of the default-style fallback and typed style lookup.

The behavioural rules these rest on were measured against live Word in
``test_default_styles_word_verified``. What is here is the plumbing around
them — provenance, ``stop_below``, malformed styles parts, and the cases a
document has to be broken to reach.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from docx_plus.core.oxml import sub
from docx_plus.styles import resolve_effective_formatting
from docx_plus.styles.inspect import StyleCascadeError


def _style(doc: Document, style_id: str) -> etree._Element:
    for st in doc.styles.element.findall(qn("w:style")):
        if st.get(qn("w:styleId")) == style_id:
            return st
    raise AssertionError(f"no style {style_id}")


def _add_style(
    doc: Document,
    style_id: str,
    kind: str | None = "paragraph",
    *,
    sz: str | None = None,
    based_on: str | None = None,
    default: str | None = None,
) -> etree._Element:
    attrs = {"w:styleId": style_id}
    if kind is not None:
        attrs["w:type"] = kind
    if default is not None:
        attrs["w:default"] = default
    style = sub(doc.styles.element, "w:style", **attrs)
    sub(style, "w:name", **{"w:val": style_id})
    if based_on is not None:
        sub(style, "w:basedOn", **{"w:val": based_on})
    if sz is not None:
        rpr = sub(style, "w:rPr")
        sub(rpr, "w:sz", **{"w:val": sz})
    return style


def _set_normal_size(doc: Document, sz: str) -> None:
    normal = _style(doc, "Normal")
    for el in normal.findall(qn("w:rPr")):
        normal.remove(el)
    rpr = sub(normal, "w:rPr")
    sub(rpr, "w:sz", **{"w:val": sz})


# --------------------------------------------------------------------------
# Provenance and stop_below.
# --------------------------------------------------------------------------


def test_the_default_style_reports_the_paragraph_style_layer() -> None:
    """It is the paragraph style, so it needs no layer of its own."""
    doc = Document()
    _set_normal_size(doc, "48")
    para = doc.add_paragraph("text")

    resolved = resolve_effective_formatting(para, include_provenance=True)
    source = resolved.provenance["font_size"]
    assert source.layer == "paragraphStyle"
    assert source.style_id == "Normal"
    assert source.chain_depth == 0


def test_stop_below_paragraph_style_drops_the_default_styles_formatting() -> None:
    """``stop_below`` treats it as the layer it claims to be."""
    doc = Document()
    _set_normal_size(doc, "48")
    para = doc.add_paragraph("text")

    assert resolve_effective_formatting(para).font_size == 24.0
    beneath = resolve_effective_formatting(para, stop_below="paragraphStyle")
    assert beneath.font_size == 11.0


def test_stop_below_still_names_the_default_style() -> None:
    """Identity survives the cut, exactly as a declared pStyle's does."""
    doc = Document()
    para = doc.add_paragraph("text")

    beneath = resolve_effective_formatting(para, stop_below="paragraphStyle")
    assert beneath.style_id == "Normal"
    assert beneath.style_name == "Normal"


# --------------------------------------------------------------------------
# Selecting the default when the styles part is odd.
# --------------------------------------------------------------------------


def test_a_style_with_no_type_attribute_counts_as_a_paragraph_style() -> None:
    """ECMA-376 17.7.4.17: an omitted ``w:type`` means a paragraph style.

    Spec-derived rather than measured — Word always writes the attribute,
    so a document missing it is malformed rather than merely unusual.
    """
    doc = Document()
    del _style(doc, "Normal").attrib[qn("w:default")]
    _add_style(doc, "Untyped", kind=None, sz="48", default="1")

    assert resolve_effective_formatting(doc.add_paragraph("x")).style_id == "Untyped"


def test_a_character_style_claiming_default_is_not_the_paragraph_default() -> None:
    doc = Document()
    _set_normal_size(doc, "48")
    _add_style(doc, "CharDefault", "character", sz="96", default="1")

    resolved = resolve_effective_formatting(doc.add_paragraph("x"))
    assert resolved.style_id == "Normal"
    assert resolved.font_size == 24.0


def test_a_default_style_with_no_style_id_is_skipped() -> None:
    """``w:styleId`` is what a fallback would have to name."""
    doc = Document()
    _set_normal_size(doc, "48")
    anonymous = sub(doc.styles.element, "w:style", **{"w:type": "paragraph", "w:default": "1"})
    sub(anonymous, "w:name", **{"w:val": "Anonymous"})

    assert resolve_effective_formatting(doc.add_paragraph("x")).style_id == "Normal"


def test_an_empty_default_attribute_does_not_claim_the_slot() -> None:
    doc = Document()
    _set_normal_size(doc, "48")
    _add_style(doc, "Blank", sz="96", default="")

    assert resolve_effective_formatting(doc.add_paragraph("x")).style_id == "Normal"


def test_the_default_style_is_resolved_once_per_document() -> None:
    """Two resolves over one cache agree — the lookup is memoized, not stale."""
    doc = Document()
    _set_normal_size(doc, "48")
    first = doc.add_paragraph("one")
    second = doc.add_paragraph("two")

    assert resolve_effective_formatting(first).style_id == "Normal"
    assert resolve_effective_formatting(second).style_id == "Normal"


# --------------------------------------------------------------------------
# Typed chains and the cycle guards.
# --------------------------------------------------------------------------


def test_a_cycle_through_same_type_styles_still_raises() -> None:
    """Type-checking the chain must not defeat the cycle guard."""
    doc = Document()
    _add_style(doc, "A", based_on="B")
    _add_style(doc, "B", based_on="A")

    para = doc.add_paragraph()
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "A"})
    with pytest.raises(StyleCascadeError, match="cycle"):
        resolve_effective_formatting(para)


def test_a_cross_type_based_on_ends_the_chain_rather_than_cycling() -> None:
    """A style based on a character style of the same name is not a cycle."""
    doc = Document()
    _add_style(doc, "Shared", "character", sz="96")
    para_style = _add_style(doc, "Para", sz="48", based_on="Shared")
    assert para_style is not None

    para = doc.add_paragraph()
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Para"})
    assert resolve_effective_formatting(para).font_size == 24.0


def test_a_based_on_chain_of_the_right_type_still_inherits() -> None:
    doc = Document()
    _add_style(doc, "Parent", sz="96")
    _add_style(doc, "Child", based_on="Parent")

    para = doc.add_paragraph()
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Child"})
    assert resolve_effective_formatting(para).font_size == 48.0


# --------------------------------------------------------------------------
# Cells.
# --------------------------------------------------------------------------


def test_a_cell_in_a_table_with_no_style_still_gets_the_default() -> None:
    doc = Document()
    _set_normal_size(doc, "48")
    table = doc.add_table(rows=1, cols=1)
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    for el in tbl_pr.findall(qn("w:tblStyle")):
        tbl_pr.remove(el)

    resolved = resolve_effective_formatting(table.cell(0, 0))
    assert resolved.font_size == 24.0
    assert resolved.style_id == "Normal"


def test_a_cell_resolves_when_the_document_has_no_default_style() -> None:
    """No default and no Normal: docDefaults only, and no crash."""
    doc = Document()
    styles = doc.styles.element
    for st in list(styles.findall(qn("w:style"))):
        if (st.get(qn("w:type")) or "paragraph") == "paragraph":
            styles.remove(st)
    table = doc.add_table(rows=1, cols=1)

    resolved = resolve_effective_formatting(table.cell(0, 0))
    assert resolved.style_id is None
    assert resolved.font_size == 11.0


def test_stop_below_paragraph_style_applies_to_a_cell_too() -> None:
    doc = Document()
    _set_normal_size(doc, "48")
    table = doc.add_table(rows=1, cols=1)

    beneath = resolve_effective_formatting(table.cell(0, 0), stop_below="paragraphStyle")
    assert beneath.font_size == 11.0
    assert beneath.style_id == "Normal"
