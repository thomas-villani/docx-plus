"""Tests for ``docx_plus.publishing.add_caption``."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus._testing.ooxml_asserts import assert_field_dirty, assert_field_not_dirty
from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.fields import mark_fields_dirty
from docx_plus.publishing import add_caption


def _instruction_texts(p_element):
    return [t.text for t in xpath(p_element, ".//w:instrText") if t.text is not None]


def _label_texts(p_element):
    """Plain text runs (the ``<w:r><w:t>`` ones, not the field runs)."""
    out: list[str] = []
    for r in p_element.findall(qn("w:r")):
        if r.find(qn("w:fldChar")) is not None or r.find(qn("w:instrText")) is not None:
            continue
        t = r.find(qn("w:t"))
        if t is not None and t.text is not None:
            out.append(t.text)
    return out


# --------------------------------------------------------------------------
# Emission — label text run + SEQ field.
# --------------------------------------------------------------------------


def test_add_caption_writes_label_run() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, "Figure ")
    labels = _label_texts(p._p)
    assert "Figure " in labels


def test_add_caption_writes_seq_field() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, "Figure ", caption_type="Figure")
    instructions = _instruction_texts(p._p)
    assert len(instructions) == 1
    assert "SEQ Figure" in instructions[0]
    assert "\\* ARABIC" in instructions[0]


def test_add_caption_custom_caption_type() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, "Table ", caption_type="Table")
    instructions = _instruction_texts(p._p)
    assert "SEQ Table" in instructions[0]


def test_add_caption_custom_numbering() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, "Item ", caption_type="Item", numbering="ROMAN")
    instructions = _instruction_texts(p._p)
    assert "\\* ROMAN" in instructions[0]


def test_add_caption_empty_label_suppresses_run() -> None:
    """An empty label string emits no label run, just the SEQ field.

    The only visible ``<w:t>`` content is the SEQ field's initial
    result text (``"1"``) between the ``separate`` and ``end``
    ``fldChar`` markers.
    """
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, "", caption_type="X")
    visible_text = "".join((t.text or "") for t in xpath(p._p, ".//w:r/w:t"))
    assert visible_text == "1"
    assert len(_instruction_texts(p._p)) == 1


def test_add_caption_label_preserves_whitespace() -> None:
    """The trailing space in ``"Figure "`` survives Word's XML normaliser."""
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, "Figure ")
    t = xpath(p._p, ".//w:r/w:t")[0]
    assert t.get(qn("xml:space")) == "preserve"
    assert t.text == "Figure "


# --------------------------------------------------------------------------
# Field structure — five-run complex-field sequence.
# --------------------------------------------------------------------------


def test_add_caption_seq_has_full_complex_field() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, "Figure ")

    fld_chars = xpath(p._p, ".//w:fldChar")
    char_types = [c.get(qn("w:fldCharType")) for c in fld_chars]
    assert char_types == ["begin", "separate", "end"]


def test_add_caption_returns_begin_run() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = add_caption(p, "Figure ")
    assert ref.tag == qn("w:r")
    fld_char = ref.find(qn("w:fldChar"))
    assert fld_char is not None
    assert fld_char.get(qn("w:fldCharType")) == "begin"


# --------------------------------------------------------------------------
# Round-trip — save / reopen preserves the field.
# --------------------------------------------------------------------------


def test_add_caption_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, "Figure ", caption_type="Figure")
    out = tmp_path / "cap.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    body = reopened.element.body
    instructions = _instruction_texts(body)
    assert any("SEQ Figure" in i for i in instructions)


# --------------------------------------------------------------------------
# Validation — H11 (identifier + numbering picture) and M15 (label default).
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "bad_caption_type",
    [
        "",                       # empty
        "1Figure",                # starts with digit
        "Figure 1",               # contains space
        'Figure" \\f "evil',      # injection attempt — terminates identifier
        "Figure,evil",            # contains comma
    ],
)
def test_add_caption_rejects_bad_caption_type(bad_caption_type: str) -> None:
    doc = Document()
    p = doc.add_paragraph()
    with pytest.raises(ValueError, match="caption_type"):
        add_caption(p, caption_type=bad_caption_type)


@pytest.mark.parametrize(
    "bad_numbering",
    ["", "lower roman", "junk", "ARABIC "],
)
def test_add_caption_rejects_bad_numbering(bad_numbering: str) -> None:
    doc = Document()
    p = doc.add_paragraph()
    with pytest.raises(ValueError, match="numbering"):
        add_caption(p, caption_type="Figure", numbering=bad_numbering)


def test_add_caption_label_defaults_to_caption_type_plus_space() -> None:
    """M15: omitting ``label`` falls back to ``f'{caption_type} '``."""
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, caption_type="Table")
    assert "Table " in _label_texts(p._p)


def test_add_caption_explicit_empty_label_suppresses_run() -> None:
    """Passing ``""`` (vs ``None``) still suppresses the label run.

    With label suppressed, the only visible <w:t> is the SEQ field's
    result text "1" (mirrors test_add_caption_empty_label_suppresses_run).
    """
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, "", caption_type="Figure")
    visible_text = "".join((t.text or "") for t in xpath(p._p, ".//w:r/w:t"))
    assert visible_text == "1"


# --------------------------------------------------------------------------
# L20: add_caption must NOT auto-mark fields dirty.
# --------------------------------------------------------------------------


def test_add_caption_does_not_mark_fields_dirty() -> None:
    doc = Document()
    add_caption(doc.add_paragraph(), "Figure ")
    assert_field_not_dirty(doc)


def test_add_caption_then_mark_fields_dirty_sets_flag() -> None:
    doc = Document()
    add_caption(doc.add_paragraph(), "Figure ")
    mark_fields_dirty(doc)
    assert_field_dirty(doc)
