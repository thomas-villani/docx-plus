"""Tests for cross-references that don't target a plain bookmark.

Two halves, both v0.5:

- ``STYLEREF`` (``fields.add_style_reference``) — the one cross-reference
  kind needing no bookmark at all.
- Referenceable captions (``publishing.add_caption(bookmark_name=...)``)
  plus the switch surface on ``bookmarks.add_cross_reference`` that makes
  "see Figure 3 above" expressible.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from docx.document import Document as DocumentObject

from docx_plus._testing.ooxml_asserts import field_instruction_text
from docx_plus.bookmarks import (
    BookmarkNameRegistry,
    DuplicateBookmarkNameError,
    add_bookmark,
    add_cross_reference,
    read_bookmarks,
)
from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.fields import add_style_reference, mark_fields_dirty
from docx_plus.publishing import add_caption

# --------------------------------------------------------------------------
# Helpers.
# --------------------------------------------------------------------------


def _instruction(paragraph) -> str:
    text = field_instruction_text(paragraph._p)
    assert text is not None, "no w:instrText in the paragraph"
    return text


def _instructions(paragraph) -> list[str]:
    return [t.text or "" for t in xpath(paragraph._p, ".//w:instrText")]


def _round_trip(doc: DocumentObject) -> DocumentObject:
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Document(buffer)


# --------------------------------------------------------------------------
# STYLEREF.
# --------------------------------------------------------------------------


def test_style_reference_quotes_the_style_name() -> None:
    doc = Document()
    p = doc.sections[0].header.paragraphs[0]
    add_style_reference(p, style="Heading 1")
    # The name is quoted because style names contain spaces, and
    # MERGEFORMAT is on by default to match Word's own UI.
    assert _instruction(p) == r' STYLEREF "Heading 1" \* MERGEFORMAT '


def test_style_reference_accepts_an_outline_level_unquoted() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_style_reference(p, style=2, preserve_formatting=False)
    assert _instruction(p) == " STYLEREF 2 "


def test_style_reference_can_omit_merge_format() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_style_reference(p, style="Heading 1", preserve_formatting=False)
    assert _instruction(p) == r' STYLEREF "Heading 1" '


@pytest.mark.parametrize(
    ("kwargs", "expected"),
    [
        ({"search_from_bottom": True}, r"\l"),
        ({"position": True}, r"\p"),
        ({"suppress_non_delimiters": True}, r"\t"),
        ({"number": "plain"}, r"\n"),
        ({"number": "relative"}, r"\r"),
        ({"number": "full"}, r"\w"),
    ],
)
def test_style_reference_switches(kwargs: dict[str, object], expected: str) -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_style_reference(p, style="Heading 1", preserve_formatting=False, **kwargs)  # type: ignore[arg-type]
    assert _instruction(p) == f' STYLEREF "Heading 1" {expected} '


def test_style_reference_orders_switches_deterministically() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_style_reference(
        p,
        style="Heading 1",
        number="relative",
        search_from_bottom=True,
        position=True,
    )
    assert _instruction(p) == r' STYLEREF "Heading 1" \r \l \p \* MERGEFORMAT '


def test_style_reference_rejects_a_quote_in_the_style_name() -> None:
    """A double-quote would terminate the quoted argument."""
    doc = Document()
    with pytest.raises(ValueError, match="double-quote"):
        add_style_reference(doc.add_paragraph(), style='Heading "1"')


@pytest.mark.parametrize("style", ["", "   "])
def test_style_reference_rejects_an_empty_style_name(style: str) -> None:
    doc = Document()
    with pytest.raises(ValueError, match="non-empty style name"):
        add_style_reference(doc.add_paragraph(), style=style)


@pytest.mark.parametrize("level", [0, 10, -1])
def test_style_reference_rejects_an_out_of_range_outline_level(level: int) -> None:
    doc = Document()
    with pytest.raises(ValueError, match="outline-level"):
        add_style_reference(doc.add_paragraph(), style=level)


def test_style_reference_rejects_a_bool_style() -> None:
    """``bool`` is an ``int`` subclass; ``style=True`` would mean level 1."""
    doc = Document()
    with pytest.raises(ValueError, match="not a bool"):
        add_style_reference(doc.add_paragraph(), style=True)


def test_style_reference_rejects_an_unknown_number_context() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="number must be one of"):
        add_style_reference(
            doc.add_paragraph(),
            style="Heading 1",
            number="absolute",  # type: ignore[arg-type]
        )


def test_style_reference_survives_a_round_trip() -> None:
    doc = Document()
    header = doc.sections[0].header.paragraphs[0]
    add_style_reference(header, style="Heading 1")
    mark_fields_dirty(doc)

    reopened = _round_trip(doc)
    instructions = _instructions(reopened.sections[0].header.paragraphs[0])
    assert instructions == [r' STYLEREF "Heading 1" \* MERGEFORMAT ']


# --------------------------------------------------------------------------
# add_cross_reference switches.
# --------------------------------------------------------------------------


def test_cross_reference_default_is_ref_with_hyperlink() -> None:
    doc = Document()
    add_bookmark(doc.add_paragraph("Section 1"), "sec_1")
    p = doc.add_paragraph("See ")
    add_cross_reference(p, bookmark="sec_1")
    assert _instruction(p) == r" REF sec_1 \h "


@pytest.mark.parametrize(
    ("kwargs", "expected"),
    [
        ({"number": "plain"}, r"\n \h"),
        ({"number": "relative"}, r"\r \h"),
        ({"number": "full"}, r"\w \h"),
        ({"position": True}, r"\p \h"),
        ({"suppress_non_delimiters": True}, r"\t \h"),
        ({"preserve_formatting": True}, r"\h \* MERGEFORMAT"),
        ({"hyperlink": False}, ""),
    ],
)
def test_cross_reference_switches(kwargs: dict[str, object], expected: str) -> None:
    doc = Document()
    add_bookmark(doc.add_paragraph("Section 1"), "sec_1")
    p = doc.add_paragraph()
    add_cross_reference(p, bookmark="sec_1", **kwargs)  # type: ignore[arg-type]
    tail = f" {expected} " if expected else " "
    assert _instruction(p) == f" REF sec_1{tail}"


def test_cross_reference_numeric_picture_is_quoted() -> None:
    doc = Document()
    add_bookmark(doc.add_paragraph("2"), "qty")
    p = doc.add_paragraph()
    add_cross_reference(p, bookmark="qty", numeric_format="0.00")
    assert _instruction(p) == r' REF qty \h \# "0.00" '


def test_cross_reference_rejects_a_quote_in_the_numeric_picture() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="numeric_format"):
        add_cross_reference(doc.add_paragraph(), bookmark="x", numeric_format='0"0')


def test_pageref_still_takes_only_the_hyperlink_switch() -> None:
    doc = Document()
    add_bookmark(doc.add_paragraph("Section 1"), "sec_1")
    p = doc.add_paragraph()
    add_cross_reference(p, bookmark="sec_1", kind="page")
    assert _instruction(p) == r" PAGEREF sec_1 \h "


@pytest.mark.parametrize(
    "kwargs",
    [
        {"number": "plain"},
        {"position": True},
        {"suppress_non_delimiters": True},
    ],
)
def test_pageref_rejects_ref_only_switches(kwargs: dict[str, object]) -> None:
    """PAGEREF resolves to a page number; these switches are meaningless."""
    doc = Document()
    with pytest.raises(ValueError, match="PAGEREF"):
        add_cross_reference(
            doc.add_paragraph(),
            bookmark="sec_1",
            kind="page",
            **kwargs,  # type: ignore[arg-type]
        )


def test_cross_reference_rejects_an_unknown_number_context() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="number must be one of"):
        add_cross_reference(
            doc.add_paragraph(),
            bookmark="sec_1",
            number="absolute",  # type: ignore[arg-type]
        )


def test_cross_reference_validates_the_bookmark_name() -> None:
    """An invalid name produces a silently unresolved field in Word."""
    doc = Document()
    with pytest.raises(ValueError, match="bookmark"):
        add_cross_reference(doc.add_paragraph(), bookmark="has spaces")


# --------------------------------------------------------------------------
# Referenceable captions. The point of the whole exercise: a REF field
# cannot target a SEQ field, only a bookmark.
# --------------------------------------------------------------------------


def test_caption_without_a_bookmark_name_is_unchanged() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, caption_type="Figure")
    assert xpath(p._p, ".//w:bookmarkStart") == []


def test_caption_bookmark_brackets_the_label_and_the_field() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, caption_type="Figure", bookmark_name="fig_arch")

    children = [child.tag for child in p._p]
    assert children[0] == qn("w:bookmarkStart")
    assert children[-1] == qn("w:bookmarkEnd")

    start = p._p[0]
    assert start.get(qn("w:name")) == "fig_arch"
    assert start.get(qn("w:id")) == p._p[-1].get(qn("w:id"))


def test_caption_bookmark_excludes_description_added_afterwards() -> None:
    """Word's "Only label and number" extent — text added later is outside."""
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, caption_type="Figure", bookmark_name="fig_arch")
    p.add_run(": Architecture overview")

    end_index = list(p._p).index(p._p.find(qn("w:bookmarkEnd")))
    assert end_index < len(list(p._p)) - 1, "description should sit after the end marker"

    # read_bookmarks joins w:t only, so the bookmark reads as the label
    # plus the field's cached result -- not the description.
    info = {b.name: b.anchored_text for b in read_bookmarks(doc)}
    assert info["fig_arch"] == "Figure 1"


def test_caption_bookmark_works_with_a_suppressed_label() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_caption(p, label="", caption_type="Figure", bookmark_name="fig_bare")

    children = [child.tag for child in p._p]
    assert children[0] == qn("w:bookmarkStart")
    assert children[-1] == qn("w:bookmarkEnd")
    # No label run: the bookmark opens directly on the field.
    assert xpath(p._p, ".//w:t") == [] or all(t.text != "Figure " for t in xpath(p._p, ".//w:t"))


def test_caption_bookmark_validates_the_name() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="bookmark_name"):
        add_caption(doc.add_paragraph(), bookmark_name="not a name")


def test_caption_bookmark_ids_are_unique_across_captions() -> None:
    doc = Document()
    for index in range(4):
        add_caption(doc.add_paragraph(), bookmark_name=f"fig_{index}")

    ids = [start.get(qn("w:id")) for start in xpath(doc.element.body, ".//w:bookmarkStart")]
    assert len(ids) == 4
    assert len(set(ids)) == 4


def test_caption_bookmark_shares_an_id_registry() -> None:
    from docx_plus.bookmarks import BookmarkIdRegistry

    doc = Document()
    registry = BookmarkIdRegistry(doc)
    add_caption(doc.add_paragraph(), bookmark_name="fig_a", bookmark_id_registry=registry)
    add_caption(doc.add_paragraph(), bookmark_name="fig_b", bookmark_id_registry=registry)

    ids = [
        int(start.get(qn("w:id")) or 0) for start in xpath(doc.element.body, ".//w:bookmarkStart")
    ]
    assert set(ids) <= registry.issued()


def test_see_figure_three_end_to_end() -> None:
    """The workflow the whole feature exists for."""
    doc = Document()
    names = BookmarkNameRegistry(doc)

    captions = []
    for title in ("First", "Second", "Third"):
        para = doc.add_paragraph()
        name = names.next_ref_name()
        add_caption(para, caption_type="Figure", bookmark_name=name)
        para.add_run(f": {title}")
        captions.append(name)

    body = doc.add_paragraph("As shown in ")
    add_cross_reference(body, bookmark=captions[2])
    body.add_run(", and on page ")
    add_cross_reference(body, bookmark=captions[2], kind="page")
    mark_fields_dirty(doc)

    reopened = _round_trip(doc)
    found = {b.name for b in read_bookmarks(reopened)}
    assert set(captions) <= found

    instructions = _instructions(reopened.paragraphs[3])
    assert instructions == [
        f" REF {captions[2]} \\h ",
        f" PAGEREF {captions[2]} \\h ",
    ]


# --------------------------------------------------------------------------
# BookmarkNameRegistry.
# --------------------------------------------------------------------------


def test_name_registry_seeds_from_existing_bookmarks() -> None:
    doc = Document()
    add_bookmark(doc.add_paragraph("x"), "intro")
    registry = BookmarkNameRegistry(doc)
    assert "intro" in registry
    assert registry.issued() == frozenset({"intro"})


def test_name_registry_rejects_a_duplicate_reserve() -> None:
    doc = Document()
    add_bookmark(doc.add_paragraph("x"), "intro")
    registry = BookmarkNameRegistry(doc)
    with pytest.raises(DuplicateBookmarkNameError, match="already in use"):
        registry.reserve("intro")


def test_name_registry_reserve_echoes_and_claims() -> None:
    registry = BookmarkNameRegistry(Document())
    assert registry.reserve("fig_1") == "fig_1"
    with pytest.raises(DuplicateBookmarkNameError):
        registry.reserve("fig_1")


def test_duplicate_bookmark_name_error_subclasses_value_error() -> None:
    registry = BookmarkNameRegistry(Document())
    registry.reserve("x")
    with pytest.raises(ValueError):
        registry.reserve("x")


def test_next_ref_name_mints_hidden_word_style_names() -> None:
    registry = BookmarkNameRegistry(Document())
    names = {registry.next_ref_name() for _ in range(50)}

    assert len(names) == 50, "minted names must be unique"
    for name in names:
        # The leading underscore is what keeps these out of Word's
        # Bookmark dialog.
        assert name.startswith("_Ref")
        assert name[4:].isdigit()
        assert len(name[4:]) == 9


def test_next_ref_name_avoids_names_already_in_the_document() -> None:
    doc = Document()
    add_bookmark(doc.add_paragraph("x"), "_Ref000000001")
    registry = BookmarkNameRegistry(doc)
    assert registry.next_ref_name() != "_Ref000000001"


def test_minted_names_pass_words_bookmark_grammar() -> None:
    doc = Document()
    registry = BookmarkNameRegistry(doc)
    name = registry.next_ref_name()
    # add_bookmark validates; a minted name must be usable as-is.
    ref = add_bookmark(doc.add_paragraph("x"), name)
    assert ref.name == name


def test_bookmark_registries_are_re_exported_from_bookmarks() -> None:
    """They moved to ``core.ids`` in v0.5; the old import must still work."""
    from docx_plus.bookmarks.registry import BookmarkIdRegistry as FromBookmarks
    from docx_plus.core.ids import BookmarkIdRegistry as FromCore

    assert FromBookmarks is FromCore


def test_caption_bookmark_round_trips(tmp_path: Path) -> None:
    doc = Document()
    para = doc.add_paragraph()
    add_caption(para, caption_type="Table", bookmark_name="tbl_costs")
    para.add_run(": Cost breakdown")
    body = doc.add_paragraph("See ")
    add_cross_reference(body, bookmark="tbl_costs")
    mark_fields_dirty(doc)

    out = tmp_path / "captions.docx"
    doc.save(str(out))
    reopened = Document(str(out))

    assert [b.name for b in read_bookmarks(reopened)] == ["tbl_costs"]
    assert _instructions(reopened.paragraphs[1]) == [r" REF tbl_costs \h "]
