"""Tests for the rules whose subject is a style definition, not a paragraph.

These are the first rules to read a part the sweep does not cover, so they
are also where the composing-layer split is under test: the OOXML knowledge
lives in ``styles/`` and the rule only decides what is worth reporting.

The load-bearing tests here are the *negative* ones. A stock python-docx
template materialises 165 style definitions and uses almost none of them,
so a naive `unused-styles` reports 165 findings on an empty document.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.shared import Pt
from lxml import etree

from docx_plus.core.oxml import sub
from docx_plus.lint import Finding, lint
from docx_plus.styles import create_style, find_unused_styles, modify_style


def _rules_fired(doc: Document, rule_id: str) -> list[Finding]:
    return lint(doc, select=[rule_id])


def _paragraph_styled(doc: Document, text: str, style_id: str) -> None:
    para = doc.add_paragraph(text)
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": style_id})


# ---------------------------------------------------------------------------
# find_unused_styles — the capability the rule composes.
# ---------------------------------------------------------------------------


def test_find_unused_styles_reports_an_orphan() -> None:
    doc = Document()
    create_style(doc, "Orphan", style_type="paragraph")

    assert "Orphan" in {info.style_id for info in find_unused_styles(doc)}


def test_find_unused_styles_ignores_an_applied_style() -> None:
    doc = Document()
    create_style(doc, "Applied", style_type="paragraph")
    _paragraph_styled(doc, "text", "Applied")

    assert "Applied" not in {info.style_id for info in find_unused_styles(doc)}


def test_find_unused_styles_follows_based_on() -> None:
    """A style reached through a used style's basedOn chain is used."""
    doc = Document()
    create_style(doc, "Parent", style_type="paragraph")
    create_style(doc, "Child", style_type="paragraph", based_on="Parent")
    _paragraph_styled(doc, "text", "Child")

    unused = {info.style_id for info in find_unused_styles(doc)}

    assert "Parent" not in unused
    assert "Child" not in unused


def test_find_unused_styles_is_transitive() -> None:
    """A style referenced only by an unused style is itself unused.

    The reachable set is grown from the body outwards for exactly this: a
    single pass would call ``Grandparent`` used, because ``Parent`` points
    at it, even though nothing points at ``Parent``.
    """
    doc = Document()
    create_style(doc, "Grandparent", style_type="paragraph")
    create_style(doc, "Parent", style_type="paragraph", based_on="Grandparent")

    unused = {info.style_id for info in find_unused_styles(doc)}

    assert {"Grandparent", "Parent"} <= unused


def test_find_unused_styles_keeps_default_styles() -> None:
    """Word applies a default style with no reference at all."""
    doc = Document()

    assert "Normal" not in {info.style_id for info in find_unused_styles(doc)}


def test_find_unused_styles_collapses_an_unused_linked_pair() -> None:
    """A linked pair is one style, so an unused pair is one finding."""
    doc = Document()
    unused = {info.style_id for info in find_unused_styles(doc)}

    assert "Heading1" in unused
    assert "Heading1Char" not in unused


def test_find_unused_styles_keeps_a_linked_partner_of_a_used_style() -> None:
    doc = Document()
    doc.add_paragraph("Heading", style="Heading 1")

    unused = {info.style_id for info in find_unused_styles(doc)}

    assert "Heading1" not in unused
    assert "Heading1Char" not in unused


def test_find_unused_styles_counts_header_references() -> None:
    """A style used only in a header is used — deleting it would break it."""
    doc = Document()
    create_style(doc, "HeaderOnly", style_type="paragraph")
    header_para = doc.sections[0].header.paragraphs[0]
    sub(sub(header_para._p, "w:pPr"), "w:pStyle", **{"w:val": "HeaderOnly"})

    assert "HeaderOnly" not in {info.style_id for info in find_unused_styles(doc)}


def test_is_builtin_distinguishes_template_styles_from_authored_ones() -> None:
    """``w:customStyle="1"`` is the ECMA-376 17.7.4.9 marker."""
    from docx_plus.styles import list_styles

    doc = Document()
    create_style(doc, "Authored", style_type="paragraph")
    by_id = {info.style_id: info for info in list_styles(doc)}

    assert by_id["Authored"].is_builtin is False
    assert by_id["Normal"].is_builtin is True
    # The table-style gallery is most of what a template ships, and none of
    # it is in the known-built-ins table — so the marker, not the table, is
    # what the classification has to rest on.
    assert by_id["LightShading"].is_builtin is True


# ---------------------------------------------------------------------------
# unused-styles.
# ---------------------------------------------------------------------------


def test_unused_styles_fires_on_an_authored_orphan() -> None:
    doc = Document()
    create_style(doc, "Orphan", style_type="paragraph")

    findings = _rules_fired(doc, "unused-styles")

    assert len(findings) == 1
    assert findings[0].location.style_id == "Orphan"
    assert findings[0].location.paragraph_index is None


def test_unused_styles_ignores_the_entire_stock_template() -> None:
    """The negative case that decides whether the rule is usable at all.

    An empty document is not a document with 165 defects in it.
    """
    doc = Document()
    doc.add_paragraph("Just some text.")

    assert _rules_fired(doc, "unused-styles") == []


def test_unused_styles_is_off_by_default() -> None:
    doc = Document()
    create_style(doc, "Orphan", style_type="paragraph")

    assert "unused-styles" not in {f.rule for f in lint(doc)}


def test_unused_styles_marks_the_fix_as_content_changing() -> None:
    """Deleting a definition removes something, so it is gated."""
    doc = Document()
    create_style(doc, "Orphan", style_type="paragraph")

    assert _rules_fired(doc, "unused-styles")[0].adds_content is True


# ---------------------------------------------------------------------------
# duplicate-styles.
# ---------------------------------------------------------------------------


def test_duplicate_styles_fires_on_two_identical_styles() -> None:
    doc = Document()
    create_style(doc, "BodyA", style_type="paragraph", based_on="Normal")
    create_style(doc, "BodyB", style_type="paragraph", based_on="Normal")
    _paragraph_styled(doc, "one", "BodyA")
    _paragraph_styled(doc, "two", "BodyB")

    findings = _rules_fired(doc, "duplicate-styles")

    assert {f.location.style_id for f in findings} == {"BodyA", "BodyB"}


def test_duplicate_styles_ignores_styles_that_differ() -> None:
    doc = Document()
    create_style(doc, "BodyA", style_type="paragraph", based_on="Normal")
    create_style(doc, "BodyB", style_type="paragraph", based_on="Normal")
    modify_style(doc, "BodyB", font_size=Pt(18))
    _paragraph_styled(doc, "one", "BodyA")
    _paragraph_styled(doc, "two", "BodyB")

    assert _rules_fired(doc, "duplicate-styles") == []


def test_duplicate_styles_matches_across_different_based_on_routes() -> None:
    """Resolved formatting is the comparison, not the style elements."""
    doc = Document()
    create_style(doc, "Direct", style_type="paragraph", based_on="Normal")
    modify_style(doc, "Direct", font_size=Pt(18))
    create_style(doc, "Middle", style_type="paragraph", based_on="Normal")
    modify_style(doc, "Middle", font_size=Pt(18))
    create_style(doc, "Inherited", style_type="paragraph", based_on="Middle")
    _paragraph_styled(doc, "one", "Direct")
    _paragraph_styled(doc, "two", "Inherited")

    assert {f.location.style_id for f in _rules_fired(doc, "duplicate-styles")} == {
        "Direct",
        "Inherited",
    }


def test_duplicate_styles_ignores_a_document_using_one_style() -> None:
    doc = Document()
    doc.add_paragraph("one")
    doc.add_paragraph("two")

    assert _rules_fired(doc, "duplicate-styles") == []


def test_duplicate_styles_ignores_direct_formatting() -> None:
    """The baseline excludes the author's own overrides.

    Two paragraphs of one style nudged apart by hand are not two styles,
    and two paragraphs of different styles nudged into looking alike are
    still two styles.
    """
    doc = Document()
    create_style(doc, "BodyA", style_type="paragraph", based_on="Normal")
    create_style(doc, "BodyB", style_type="paragraph", based_on="Normal")
    modify_style(doc, "BodyB", font_size=Pt(18))
    _paragraph_styled(doc, "one", "BodyA")
    para = doc.add_paragraph("two")
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "BodyB"})
    para.runs[0].font.size = Pt(11)  # made to *look* like BodyA

    assert _rules_fired(doc, "duplicate-styles") == []


def test_duplicate_styles_skips_table_and_list_paragraphs() -> None:
    """Their baselines carry the table style and numbering level too.

    Neither belongs to the paragraph style being compared, so including
    them would both miss duplicates and invent them.
    """
    doc = Document()
    create_style(doc, "BodyA", style_type="paragraph", based_on="Normal")
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    sub(sub(cell.paragraphs[0]._p, "w:pPr"), "w:pStyle", **{"w:val": "BodyA"})
    doc.add_paragraph("item", style="List Bullet")

    assert _rules_fired(doc, "duplicate-styles") == []


@pytest.mark.parametrize("rule_id", ["duplicate-styles", "unused-styles"])
def test_style_rules_do_not_mutate_the_document(rule_id: str) -> None:
    """These read a mutable part directly, so the read-only promise is worth asserting."""
    doc = Document()
    create_style(doc, "Orphan", style_type="paragraph")
    before = etree.tostring(doc.styles.element)

    _rules_fired(doc, rule_id)

    assert etree.tostring(doc.styles.element) == before
