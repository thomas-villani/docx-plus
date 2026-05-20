"""Tests for ``docx_plus.comments`` — anchored-comment insertion, read,
delete, and the comment-id registry."""

from __future__ import annotations

import datetime as dt
from pathlib import Path

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from docx_plus.comments import (
    AnchoredComment,
    CommentIdRegistry,
    CommentNotFoundError,
    CommentRef,
    add_comment,
    clear_all_comments,
    delete_comment,
    edit_comment,
    read_comments,
)
from docx_plus.core.ids import DuplicateIdError
from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath

# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------


def _body(doc):
    return doc.element.body


def _range_starts(doc, cid):
    return xpath(_body(doc), ".//w:commentRangeStart[@w:id=$cid]", cid=str(cid))


def _range_ends(doc, cid):
    return xpath(_body(doc), ".//w:commentRangeEnd[@w:id=$cid]", cid=str(cid))


def _reference_runs(doc, cid):
    return xpath(_body(doc), ".//w:commentReference[@w:id=$cid]", cid=str(cid))


def _comment_part_entries(doc, cid):
    part = doc.part.part_related_by(RT.COMMENTS)
    return xpath(part.element, "./w:comment[@w:id=$cid]", cid=str(cid))


# --------------------------------------------------------------------------
# add_comment — body-side anchoring on a single run.
# --------------------------------------------------------------------------


def test_add_comment_to_run_writes_all_three_body_markers() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("anchored text")
    ref = add_comment(run, "comment body")
    assert isinstance(ref, CommentRef)
    cid = ref.comment_id
    assert len(_range_starts(doc, cid)) == 1
    assert len(_range_ends(doc, cid)) == 1
    assert len(_reference_runs(doc, cid)) == 1


def test_add_comment_to_run_brackets_the_run_element() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("hello")
    ref = add_comment(run, "x")
    cid = str(ref.comment_id)
    siblings = list(p._p)
    types = [child.tag.rpartition("}")[2] for child in siblings]
    # Order: commentRangeStart, r, commentRangeEnd, r (reference run)
    assert "commentRangeStart" in types
    assert "commentRangeEnd" in types
    rs_idx = types.index("commentRangeStart")
    r_idx = types.index("r", rs_idx)
    re_idx = types.index("commentRangeEnd", r_idx)
    assert rs_idx < r_idx < re_idx
    # commentRangeStart and commentRangeEnd both carry the right id
    assert siblings[rs_idx].get(qn("w:id")) == cid
    assert siblings[re_idx].get(qn("w:id")) == cid


def test_add_comment_writes_comment_body_to_comments_part() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("hello")
    ref = add_comment(run, "this is the comment", author="Alice", initials="A")

    entries = _comment_part_entries(doc, ref.comment_id)
    assert len(entries) == 1
    body = entries[0]
    assert body.get(qn("w:author")) == "Alice"
    assert body.get(qn("w:initials")) == "A"
    # Extract body text.
    texts = [t.text for t in xpath(body, ".//w:t") if t.text]
    assert "".join(texts) == "this is the comment"


def test_add_comment_default_initials_from_author() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("x")
    ref = add_comment(run, "y", author="Wendy")
    body = _comment_part_entries(doc, ref.comment_id)[0]
    assert body.get(qn("w:initials")) == "W"


def test_add_comment_no_initials_with_empty_author() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("x")
    ref = add_comment(run, "y")
    body = _comment_part_entries(doc, ref.comment_id)[0]
    assert body.get(qn("w:initials")) is None


def test_add_comment_emits_iso_utc_timestamp() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("x")
    ref = add_comment(run, "y")
    body = _comment_part_entries(doc, ref.comment_id)[0]
    date = body.get(qn("w:date"))
    assert date is not None and date.endswith("Z")
    # Round-trips through fromisoformat after stripping Z.
    parsed = dt.datetime.fromisoformat(date.replace("Z", "+00:00"))
    assert parsed.tzinfo is not None


def test_add_comment_preserves_whitespace_in_text() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("x")
    ref = add_comment(run, "  leading and trailing  ")
    body = _comment_part_entries(doc, ref.comment_id)[0]
    text_runs = xpath(body, ".//w:t")
    # The user text is in the second w:t (first is the annotation-ref run's empty content)
    user_t = next(t for t in text_runs if t.text and "leading" in t.text)
    from docx_plus.core.ns import XML

    assert user_t.get(f"{{{XML}}}space") == "preserve"


def test_add_comment_returns_unique_ids() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("a")
    r2 = p.add_run("b")
    ref1 = add_comment(r1, "one")
    ref2 = add_comment(r2, "two")
    assert ref1.comment_id != ref2.comment_id


# --------------------------------------------------------------------------
# add_comment — Paragraph and run-range targets.
# --------------------------------------------------------------------------


def test_add_comment_to_paragraph_wraps_all_runs() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("one ")
    p.add_run("two ")
    p.add_run("three")
    add_comment(p, "wrap me")

    siblings = list(p._p)
    types = [child.tag.rpartition("}")[2] for child in siblings]
    rs_idx = types.index("commentRangeStart")
    re_idx = types.index("commentRangeEnd")
    runs_between = [t for t in types[rs_idx + 1 : re_idx] if t == "r"]
    assert len(runs_between) == 3


def test_add_comment_to_paragraph_with_no_runs_raises() -> None:
    doc = Document()
    p = doc.add_paragraph()  # no runs
    with pytest.raises(ValueError, match="at least one run"):
        add_comment(p, "x")


def test_add_comment_to_run_range_wraps_inclusive() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("first ")
    p.add_run("middle ")
    r3 = p.add_run("last")
    add_comment((r1, r3), "wrap range")

    types = [c.tag.rpartition("}")[2] for c in p._p]
    rs_idx = types.index("commentRangeStart")
    re_idx = types.index("commentRangeEnd")
    runs_between = [t for t in types[rs_idx + 1 : re_idx] if t == "r"]
    assert len(runs_between) == 3


def test_add_comment_with_bad_target_type_raises() -> None:
    with pytest.raises(TypeError, match="Run, Paragraph"):
        add_comment("not a run", "x")  # type: ignore[arg-type]


def test_add_comment_with_bad_tuple_raises() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("hi")
    with pytest.raises(TypeError, match="tuple of"):
        add_comment((r1, "not a run"), "x")  # type: ignore[arg-type]


# --------------------------------------------------------------------------
# add_comment — comments part is created on first use, reused after.
# --------------------------------------------------------------------------


def test_first_add_comment_creates_comments_part() -> None:
    doc = Document()
    # No comments part exists yet.
    with pytest.raises(KeyError):
        doc.part.part_related_by(RT.COMMENTS)
    p = doc.add_paragraph()
    add_comment(p.add_run("x"), "first")
    # Now it exists.
    assert doc.part.part_related_by(RT.COMMENTS) is not None


def test_second_add_comment_reuses_comments_part() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_comment(p.add_run("a"), "one")
    part_after_first = doc.part.part_related_by(RT.COMMENTS)
    add_comment(p.add_run("b"), "two")
    part_after_second = doc.part.part_related_by(RT.COMMENTS)
    assert part_after_first is part_after_second
    assert len(xpath(part_after_second.element, "./w:comment")) == 2


# --------------------------------------------------------------------------
# CommentIdRegistry — seeding, sharing.
# --------------------------------------------------------------------------


def test_registry_seeds_from_existing_comments_part() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("x"), "y")
    reg = CommentIdRegistry(doc)
    with pytest.raises(DuplicateIdError):
        reg.reserve(ref.comment_id)


def test_registry_seeds_from_orphaned_body_anchors() -> None:
    """A `commentRangeStart` without a matching comment body still blocks reuse."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("hi")
    add_comment(run, "real one")
    # Manually inject an orphaned commentRangeStart with a known id.
    from docx_plus.core.oxml import el

    orphan_id = 999_111
    orphan = el("w:commentRangeStart", **{"w:id": str(orphan_id)})
    run._r.addprevious(orphan)

    reg = CommentIdRegistry(doc)
    with pytest.raises(DuplicateIdError):
        reg.reserve(orphan_id)


def test_shared_registry_produces_disjoint_ids() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("a")
    r2 = p.add_run("b")
    reg = CommentIdRegistry(doc)
    a = add_comment(r1, "1", id_registry=reg)
    b = add_comment(r2, "2", id_registry=reg)
    assert a.comment_id != b.comment_id


# --------------------------------------------------------------------------
# read_comments — happy path, orphans, ordering.
# --------------------------------------------------------------------------


def test_read_comments_empty_when_no_comments_part() -> None:
    assert read_comments(Document()) == []


def test_read_comments_returns_anchored_text() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("before ")
    target = p.add_run("ANCHORED")
    p.add_run(" after")
    add_comment(target, "this part", author="Alice")

    comments = read_comments(doc)
    assert len(comments) == 1
    only = comments[0]
    assert isinstance(only, AnchoredComment)
    assert only.author == "Alice"
    assert only.initials == "A"
    assert only.text == "this part"
    assert only.anchored_text == "ANCHORED"
    assert only.paragraph_index == 0
    assert only.timestamp is not None


def test_read_comments_handles_paragraph_range() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("alpha ")
    p.add_run("beta")
    add_comment(p, "wrap whole paragraph")

    comments = read_comments(doc)
    assert comments[0].anchored_text == "alpha beta"


def test_read_comments_handles_orphans() -> None:
    """A comment with no matching range markers reads as anchored_text=''."""
    doc = Document()
    p = doc.add_paragraph()
    add_comment(p.add_run("x"), "y")
    # Strip the body-side markers, leaving an orphan in comments.xml.
    body = doc.element.body
    from docx_plus.core.oxml import remove as remove_el

    for tag in ("w:commentRangeStart", "w:commentRangeEnd"):
        for el in xpath(body, f".//{tag}"):
            remove_el(el)

    comments = read_comments(doc)
    assert comments[0].anchored_text == ""
    assert comments[0].paragraph_index == -1


def test_read_comments_paragraph_index_tracks_position() -> None:
    doc = Document()
    doc.add_paragraph("first paragraph")
    doc.add_paragraph("second paragraph")
    p3 = doc.add_paragraph()
    p3.add_run("third")
    add_comment(p3, "anchor here")

    comments = read_comments(doc)
    assert comments[0].paragraph_index == 2


def test_read_comments_preserves_comments_xml_order() -> None:
    doc = Document()
    p = doc.add_paragraph()
    refs = [add_comment(p.add_run(f"r{i}"), f"comment {i}") for i in range(3)]
    comments = read_comments(doc)
    assert [c.comment_id for c in comments] == [r.comment_id for r in refs]


# --------------------------------------------------------------------------
# delete_comment — idempotency, full cleanup.
# --------------------------------------------------------------------------


def test_delete_comment_removes_all_body_anchors() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("x"), "y")
    delete_comment(doc, ref.comment_id)
    assert _range_starts(doc, ref.comment_id) == []
    assert _range_ends(doc, ref.comment_id) == []
    assert _reference_runs(doc, ref.comment_id) == []


def test_delete_comment_removes_comments_part_entry() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("x"), "y")
    delete_comment(doc, ref.comment_id)
    assert _comment_part_entries(doc, ref.comment_id) == []


def test_delete_comment_is_idempotent_for_missing_id() -> None:
    doc = Document()
    delete_comment(doc, 999)  # no comments part — no-op
    p = doc.add_paragraph()
    add_comment(p.add_run("x"), "y")
    delete_comment(doc, 12345)  # comments part exists but id absent — no-op


def test_delete_comment_leaves_other_comments_intact() -> None:
    doc = Document()
    p = doc.add_paragraph()
    keep = add_comment(p.add_run("a"), "keep me")
    drop = add_comment(p.add_run("b"), "drop me")
    delete_comment(doc, drop.comment_id)
    assert _comment_part_entries(doc, keep.comment_id)
    assert _range_starts(doc, keep.comment_id)


# --------------------------------------------------------------------------
# Round-trip — save / reopen / read.
# --------------------------------------------------------------------------


def test_comment_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("prefix ")
    p.add_run("ANCHORED")
    p.add_run(" suffix")
    target = [r for r in p.runs if r.text == "ANCHORED"][0]
    add_comment(target, "review this", author="Bob", initials="B")
    out = tmp_path / "round.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    comments = read_comments(reopened)
    assert len(comments) == 1
    assert comments[0].author == "Bob"
    assert comments[0].text == "review this"
    assert comments[0].anchored_text == "ANCHORED"


def test_comment_delete_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref_keep = add_comment(p.add_run("keep"), "keep")
    ref_drop = add_comment(p.add_run("drop"), "drop")
    delete_comment(doc, ref_drop.comment_id)
    out = tmp_path / "del.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    comments = read_comments(reopened)
    assert len(comments) == 1
    assert comments[0].comment_id == ref_keep.comment_id


# --------------------------------------------------------------------------
# clear_all_comments — bulk removal of every comment in the document.
# --------------------------------------------------------------------------


def test_clear_all_comments_on_empty_doc_is_noop() -> None:
    doc = Document()
    clear_all_comments(doc)  # no comments part exists yet — must not raise
    assert read_comments(doc) == []


def test_clear_all_comments_removes_every_anchor_and_body() -> None:
    doc = Document()
    p = doc.add_paragraph()
    refs = [add_comment(p.add_run(f"chunk-{i}"), f"note {i}", author=f"A{i}") for i in range(3)]
    clear_all_comments(doc)

    for ref in refs:
        assert _range_starts(doc, ref.comment_id) == []
        assert _range_ends(doc, ref.comment_id) == []
        assert _reference_runs(doc, ref.comment_id) == []
        assert _comment_part_entries(doc, ref.comment_id) == []
    assert read_comments(doc) == []


def test_clear_all_comments_preserves_unrelated_content() -> None:
    doc = Document()
    doc.add_paragraph("untouched paragraph 1")
    p = doc.add_paragraph()
    p.add_run("prefix ")
    p.add_run("anchored")
    p.add_run(" suffix")
    add_comment(p.runs[1], "review")
    doc.add_paragraph("untouched paragraph 2")

    clear_all_comments(doc)

    texts = [para.text for para in doc.paragraphs]
    assert "untouched paragraph 1" in texts
    assert "untouched paragraph 2" in texts
    # The anchored paragraph keeps its three runs' text intact.
    anchored_paragraph_text = "".join(r.text for r in p.runs)
    assert anchored_paragraph_text == "prefix anchored suffix"


def test_clear_all_comments_is_idempotent_after_clearing() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_comment(p.add_run("x"), "y")
    clear_all_comments(doc)
    clear_all_comments(doc)  # second call must be a clean no-op
    assert read_comments(doc) == []


def test_clear_all_comments_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_comment(p.add_run("a"), "first")
    add_comment(p.add_run("b"), "second")
    clear_all_comments(doc)
    out = tmp_path / "cleared.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    assert read_comments(reopened) == []


# --------------------------------------------------------------------------
# edit_comment — in-place body replacement preserving metadata.
# --------------------------------------------------------------------------


def test_edit_comment_replaces_text() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("x"), "draft", author="Bob")

    edit_comment(doc, ref.comment_id, "final")

    comments = read_comments(doc)
    assert [c.text for c in comments] == ["final"]


def test_edit_comment_preserves_author_date_initials() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("x"), "draft", author="Alice", initials="A")
    comment_el = _comment_part_entries(doc, ref.comment_id)[0]
    original_author = comment_el.get(qn("w:author"))
    original_date = comment_el.get(qn("w:date"))
    original_initials = comment_el.get(qn("w:initials"))

    edit_comment(doc, ref.comment_id, "final")

    assert comment_el.get(qn("w:author")) == original_author
    assert comment_el.get(qn("w:date")) == original_date
    assert comment_el.get(qn("w:initials")) == original_initials


def test_edit_comment_preserves_body_side_anchors() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("anchored"), "draft")

    edit_comment(doc, ref.comment_id, "final")

    # The three body-side anchors must remain.
    assert _range_starts(doc, ref.comment_id)
    assert _range_ends(doc, ref.comment_id)
    assert _reference_runs(doc, ref.comment_id)


def test_edit_comment_leaves_other_comments_intact() -> None:
    doc = Document()
    p = doc.add_paragraph()
    keep = add_comment(p.add_run("a"), "keep me")
    target = add_comment(p.add_run("b"), "before")

    edit_comment(doc, target.comment_id, "after")

    comments = {c.comment_id: c.text for c in read_comments(doc)}
    assert comments[keep.comment_id] == "keep me"
    assert comments[target.comment_id] == "after"


def test_edit_comment_missing_id_raises() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_comment(p.add_run("x"), "exists")  # forces the comments part to exist

    with pytest.raises(CommentNotFoundError):
        edit_comment(doc, 9999, "no such comment")


def test_edit_comment_with_no_part_raises() -> None:
    doc = Document()  # no comments part has been created
    with pytest.raises(CommentNotFoundError):
        edit_comment(doc, 1, "nope")


def test_comment_not_found_is_a_key_error() -> None:
    """``CommentNotFoundError`` subclasses ``KeyError`` per SPEC §16."""
    doc = Document()
    p = doc.add_paragraph()
    add_comment(p.add_run("x"), "exists")

    with pytest.raises(KeyError):
        edit_comment(doc, 9999, "no")


def test_edit_comment_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("x"), "draft", author="Bob", initials="B")
    edit_comment(doc, ref.comment_id, "final review")

    out = tmp_path / "edited.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    comments = read_comments(reopened)
    assert len(comments) == 1
    assert comments[0].text == "final review"
    assert comments[0].author == "Bob"


def test_registry_seeds_from_orphaned_range_end() -> None:
    """M3: a lone `commentRangeEnd` (rangeStart stripped) still blocks reuse."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("hi")
    add_comment(run, "real one")
    from docx_plus.core.oxml import el

    orphan_id = 888_222
    orphan = el("w:commentRangeEnd", **{"w:id": str(orphan_id)})
    run._r.addnext(orphan)

    reg = CommentIdRegistry(doc)
    with pytest.raises(DuplicateIdError):
        reg.reserve(orphan_id)


def test_delete_comment_preserves_text_in_shared_reference_run() -> None:
    """M6: cleanup removes only the marker, not sibling text in the run."""
    from docx_plus.core.oxml import sub

    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("anchor"), "note")
    # Simulate a hand-edited doc: the reference run also carries text.
    ref_run = _reference_runs(doc, ref.comment_id)[0].getparent()
    t = sub(ref_run, "w:t")
    t.text = "KEEP"

    delete_comment(doc, ref.comment_id)

    assert _reference_runs(doc, ref.comment_id) == []
    assert ref_run.getparent() is not None  # run itself survives
    surviving = ref_run.find(qn("w:t"))
    assert surviving is not None and surviving.text == "KEEP"


def test_clear_all_comments_preserves_text_in_shared_reference_run() -> None:
    """M6: bulk cleanup also keeps sibling text in a shared reference run."""
    from docx_plus.core.oxml import sub

    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("anchor"), "note")
    ref_run = _reference_runs(doc, ref.comment_id)[0].getparent()
    sub(ref_run, "w:t").text = "KEEP"

    clear_all_comments(doc)

    assert _reference_runs(doc, ref.comment_id) == []
    assert ref_run.getparent() is not None
    assert ref_run.find(qn("w:t")).text == "KEEP"


def test_delete_comment_collapses_marker_only_reference_run() -> None:
    """M6 regression: the internally-built reference run still fully collapses."""
    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("x"), "y")
    ref_run = _reference_runs(doc, ref.comment_id)[0].getparent()

    delete_comment(doc, ref.comment_id)

    assert ref_run.getparent() is None  # rPr-only run removed entirely


def test_now_iso_has_millisecond_precision() -> None:
    """L1: timestamps carry sub-second precision and still round-trip."""
    from docx_plus.comments.anchor import _now_iso

    stamp = _now_iso()
    assert stamp.endswith("Z")
    assert "." in stamp  # fractional seconds present
    parsed = dt.datetime.fromisoformat(stamp.replace("Z", "+00:00"))
    assert parsed.tzinfo is not None


def test_clear_all_comments_remove_part_tears_down_relationship() -> None:
    """L17: remove_part=True drops the comments part and its relationship."""
    doc = Document()
    p = doc.add_paragraph()
    add_comment(p.add_run("x"), "y")
    assert doc.part.part_related_by(RT.COMMENTS) is not None

    clear_all_comments(doc, remove_part=True)

    with pytest.raises(KeyError):
        doc.part.part_related_by(RT.COMMENTS)


def test_clear_all_comments_default_keeps_empty_part() -> None:
    """L17: the default leaves the (now empty) part connected for reuse."""
    doc = Document()
    p = doc.add_paragraph()
    add_comment(p.add_run("x"), "y")
    clear_all_comments(doc)
    part = doc.part.part_related_by(RT.COMMENTS)
    assert xpath(part.element, "./w:comment") == []


def test_clear_all_comments_remove_part_round_trip(tmp_path: Path) -> None:
    """L17: a torn-down comments part stays absent through save / reopen."""
    doc = Document()
    p = doc.add_paragraph()
    add_comment(p.add_run("a"), "first")
    add_comment(p.add_run("b"), "second")
    clear_all_comments(doc, remove_part=True)
    out = tmp_path / "noparts.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    with pytest.raises(KeyError):
        reopened.part.part_related_by(RT.COMMENTS)
    assert read_comments(reopened) == []


def test_edit_comment_strips_non_paragraph_children() -> None:
    """H6 regression: comments can contain tables / SDTs, not just <w:p>.

    ECMA-376 17.13.4.2 (`CT_Comment`) extends `EG_BlockLevelElts`, which
    includes `<w:tbl>`, `<w:sdt>`, etc. The edit helper must strip all
    children, not just paragraphs, or the old block-level content
    survives next to the new paragraph.
    """
    from docx_plus.core.oxml import sub

    doc = Document()
    p = doc.add_paragraph()
    ref = add_comment(p.add_run("hi"), "first draft")

    # Inject a `<w:tbl>` directly into the comment body to simulate a
    # comment authored elsewhere that contains a block table.
    comments_part = doc.part.part_related_by(RT.COMMENTS)
    comment_el = comments_part.element.find(qn("w:comment"))
    assert comment_el is not None
    sub(comment_el, "w:tbl")
    assert comment_el.find(qn("w:tbl")) is not None  # sanity

    edit_comment(doc, ref.comment_id, "rewritten")

    # After edit: exactly one child, the new <w:p>, no leftover <w:tbl>.
    children = list(comment_el)
    assert len(children) == 1
    assert children[0].tag == qn("w:p")
    assert comment_el.find(qn("w:tbl")) is None
