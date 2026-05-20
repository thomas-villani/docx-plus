"""Tests for ``docx_plus.layout.set_page_borders`` and ``Border``."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus.core.ns import qn
from docx_plus.core.oxml import el, sub, xpath
from docx_plus.layout import Border, set_page_borders


def _pg_borders(section):
    return section._sectPr.find(qn("w:pgBorders"))


# --------------------------------------------------------------------------
# Emission — element exists with one child per requested side.
# --------------------------------------------------------------------------


def test_set_page_borders_writes_element() -> None:
    doc = Document()
    set_page_borders(doc.sections[0], top=Border())
    assert _pg_borders(doc.sections[0]) is not None


def test_set_page_borders_all_four_sides() -> None:
    doc = Document()
    rule = Border(style="single", size=8, color="2F5496", space=24)
    set_page_borders(doc.sections[0], top=rule, bottom=rule, left=rule, right=rule)

    pg = _pg_borders(doc.sections[0])
    assert pg is not None
    side_tags = {child.tag for child in pg}
    assert side_tags == {qn(t) for t in ("w:top", "w:bottom", "w:left", "w:right")}

    for child in pg:
        assert child.get(qn("w:val")) == "single"
        assert child.get(qn("w:sz")) == "8"
        assert child.get(qn("w:color")) == "2F5496"
        assert child.get(qn("w:space")) == "24"


def test_set_page_borders_child_order_is_top_left_bottom_right() -> None:
    """ECMA-376 17.6.10 CT_PageBorders requires this child sequence (C3)."""
    doc = Document()
    rule = Border()
    set_page_borders(doc.sections[0], top=rule, bottom=rule, left=rule, right=rule)

    pg = _pg_borders(doc.sections[0])
    assert pg is not None
    actual_order = [child.tag for child in pg]
    expected_order = [qn(t) for t in ("w:top", "w:left", "w:bottom", "w:right")]
    assert actual_order == expected_order


def test_set_page_borders_emits_offset_from_page_by_default() -> None:
    """Default ``offset_from="page"`` matches Word's UI emission (H7)."""
    doc = Document()
    set_page_borders(doc.sections[0], top=Border())
    pg = _pg_borders(doc.sections[0])
    assert pg is not None
    assert pg.get(qn("w:offsetFrom")) == "page"


def test_set_page_borders_accepts_offset_from_text() -> None:
    doc = Document()
    set_page_borders(doc.sections[0], top=Border(), offset_from="text")
    pg = _pg_borders(doc.sections[0])
    assert pg is not None
    assert pg.get(qn("w:offsetFrom")) == "text"


def test_set_page_borders_only_top() -> None:
    """Sides set to ``None`` are omitted from the emitted XML."""
    doc = Document()
    set_page_borders(doc.sections[0], top=Border(style="thick"))

    pg = _pg_borders(doc.sections[0])
    assert pg is not None
    side_tags = [child.tag for child in pg]
    assert side_tags == [qn("w:top")]
    assert pg.find(qn("w:top")).get(qn("w:val")) == "thick"


def test_border_defaults() -> None:
    """Bare ``Border()`` has the documented defaults."""
    b = Border()
    assert b.style == "single"
    assert b.size == 4
    assert b.color == "auto"
    assert b.space == 24


def test_border_rejects_invalid_color() -> None:
    """M5: color must be 'auto' or a six-hex-digit RRGGBB string."""
    for bad in ("red", "#FF0000", "FF00", "12345", "GGGGGG", ""):
        with pytest.raises(ValueError, match="Border.color"):
            Border(color=bad)


def test_border_accepts_auto_and_hex_colors() -> None:
    assert Border(color="auto").color == "auto"
    assert Border(color="2F5496").color == "2F5496"
    assert Border(color="ffffff").color == "ffffff"  # lowercase hex is valid


# --------------------------------------------------------------------------
# Replacement — second call overrides; all-None removes.
# --------------------------------------------------------------------------


def test_set_page_borders_replaces_existing() -> None:
    doc = Document()
    set_page_borders(doc.sections[0], top=Border(style="single"))
    set_page_borders(doc.sections[0], top=Border(style="double"))

    found = xpath(doc.sections[0]._sectPr, "./w:pgBorders")
    assert len(found) == 1
    assert found[0].find(qn("w:top")).get(qn("w:val")) == "double"


def test_set_page_borders_all_none_removes_element() -> None:
    doc = Document()
    set_page_borders(doc.sections[0], top=Border())
    assert _pg_borders(doc.sections[0]) is not None

    set_page_borders(doc.sections[0])  # all sides default to None
    assert _pg_borders(doc.sections[0]) is None


def test_set_page_borders_all_none_on_empty_is_noop() -> None:
    """All-``None`` on a section with no existing borders is a no-op."""
    doc = Document()
    set_page_borders(doc.sections[0])
    assert _pg_borders(doc.sections[0]) is None


# --------------------------------------------------------------------------
# Schema-strict insertion — pgBorders lands before its later siblings.
# --------------------------------------------------------------------------


def test_set_page_borders_lands_before_cols() -> None:
    doc = Document()
    sect_pr = doc.sections[0]._sectPr
    cols = el("w:cols", **{"w:num": "1"})
    sect_pr.append(cols)

    set_page_borders(doc.sections[0], top=Border())

    children = list(sect_pr)
    pg_idx = next(i for i, c in enumerate(children) if c.tag == qn("w:pgBorders"))
    cols_idx = next(i for i, c in enumerate(children) if c.tag == qn("w:cols"))
    assert pg_idx < cols_idx


def test_set_page_borders_lands_before_lnNumType() -> None:
    doc = Document()
    sect_pr = doc.sections[0]._sectPr
    ln = el("w:lnNumType", **{"w:countBy": "1"})
    sect_pr.append(ln)

    set_page_borders(doc.sections[0], top=Border())

    children = list(sect_pr)
    pg_idx = next(i for i, c in enumerate(children) if c.tag == qn("w:pgBorders"))
    ln_idx = next(i for i, c in enumerate(children) if c.tag == qn("w:lnNumType"))
    assert pg_idx < ln_idx


def test_set_page_borders_lands_before_cols_and_docGrid() -> None:
    """L19: with both cols and docGrid present, pgBorders precedes both."""
    doc = Document()
    sect_pr = doc.sections[0]._sectPr
    assert sect_pr.find(qn("w:cols")) is not None
    assert sect_pr.find(qn("w:docGrid")) is not None

    set_page_borders(doc.sections[0], top=Border())

    tags = [c.tag for c in sect_pr]
    pg_idx = tags.index(qn("w:pgBorders"))
    assert pg_idx < tags.index(qn("w:cols"))
    assert pg_idx < tags.index(qn("w:docGrid"))


def test_set_page_borders_replaces_preexisting_element() -> None:
    """L18: a pre-seeded pgBorders (as if loaded from Word) is replaced in place."""
    doc = Document()
    sect_pr = doc.sections[0]._sectPr
    seeded = el("w:pgBorders", **{"w:offsetFrom": "text"})
    sub(seeded, "w:top", **{"w:val": "dotted", "w:sz": "2", "w:color": "auto", "w:space": "1"})
    sect_pr.append(seeded)

    set_page_borders(doc.sections[0], top=Border(style="double", size=12, color="2F5496"))

    found = xpath(sect_pr, "./w:pgBorders")
    assert len(found) == 1
    assert found[0].get(qn("w:offsetFrom")) == "page"
    top = found[0].find(qn("w:top"))
    assert top.get(qn("w:val")) == "double"
    assert top.get(qn("w:sz")) == "12"
    assert top.get(qn("w:color")) == "2F5496"


# --------------------------------------------------------------------------
# Round-trip — save / reopen preserves the element and attributes.
# --------------------------------------------------------------------------


def test_set_page_borders_round_trip(tmp_path: Path) -> None:
    doc = Document()
    set_page_borders(
        doc.sections[0],
        top=Border(style="double", size=12, color="FF0000", space=10),
        bottom=Border(style="single", size=4, color="auto", space=24),
    )
    out = tmp_path / "borders.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    pg = _pg_borders(reopened.sections[0])
    assert pg is not None

    top = pg.find(qn("w:top"))
    assert top.get(qn("w:val")) == "double"
    assert top.get(qn("w:sz")) == "12"
    assert top.get(qn("w:color")) == "FF0000"
    assert top.get(qn("w:space")) == "10"

    bottom = pg.find(qn("w:bottom"))
    assert bottom.get(qn("w:val")) == "single"
    assert pg.find(qn("w:left")) is None
    assert pg.find(qn("w:right")) is None
