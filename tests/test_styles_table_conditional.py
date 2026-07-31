"""Conditional table-style formatting — ``<w:tblStylePr>`` branches.

ECMA-376 17.7.6.5 lets a table style carry zero or more ``<w:tblStylePr>``
children, each with a ``w:type`` that names the conditional region
(``firstRow``, ``band1Horz``, ``nwCell``, etc.). Which of them actually
reach a cell depends on three things: the cell's position, the table's
``<w:tblLook>``, and — for the band branches — a declared band size.

These tests verify the wiring end-to-end: build a document with a custom
table style, attach a table, then resolve formatting at different cell
positions and confirm the right branch was applied. The behaviour they
assert was measured against live Word; ``test_tables_word_verified.py``
holds the raw grids those measurements produced.

Note that python-docx's ``add_table`` emits Word's own default
``<w:tblLook>`` — first row and first column on, last row and last column
off, horizontal banding on, vertical banding off — so a test that wants
``lastRow`` or a vertical band has to say so.
"""

from __future__ import annotations

import dataclasses
from typing import Any

import pytest
from docx import Document
from docx.oxml.ns import qn

from docx_plus.core.oxml import sub
from docx_plus.styles import TableContext, resolve_effective_formatting
from docx_plus.styles.inspect import _grid_column, _grid_span

# What python-docx (and Word) put on a new table.
WORD_DEFAULT_LOOK = {
    "firstRow": 1,
    "lastRow": 0,
    "firstColumn": 1,
    "lastColumn": 0,
    "noHBand": 0,
    "noVBand": 1,
}

# --------------------------------------------------------------------------
# Helpers — build a synthetic table style with conditional branches.
# --------------------------------------------------------------------------


def _add_table_style(
    doc: Document,
    style_id: str,
    *,
    base_rpr: dict[str, Any] | None = None,
    branches: dict[str, dict[str, Any]] | None = None,
    band_sizes: tuple[int, int] | None = None,
    based_on: str | None = None,
) -> None:
    """Append a ``w:style w:type="table"`` to the doc's styles part.

    Args:
        doc: Document to mutate.
        style_id: ``w:styleId`` for the new style.
        base_rpr: Optional dict of (rpr-child-tag -> attrs) for the
            base ``w:rPr`` of the style.
        branches: Mapping of conditional ``w:type`` → dict of
            ``w:rPr`` children, e.g. ``{"firstRow": {"w:b": None,
            "w:color": {"w:val": "FF0000"}}}``.
        band_sizes: Optional ``(row, col)`` band sizes for the style's
            own ``w:tblPr``. Without one of these somewhere, no band
            branch applies at all.
        based_on: Optional parent style id.
    """
    styles_el = doc.styles.element
    style_el = sub(styles_el, "w:style", **{"w:type": "table", "w:styleId": style_id})
    sub(style_el, "w:name", **{"w:val": style_id})
    if based_on:
        sub(style_el, "w:basedOn", **{"w:val": based_on})

    if base_rpr:
        base = sub(style_el, "w:rPr")
        for tag, attrs in base_rpr.items():
            sub(base, tag, **(attrs or {}))

    if band_sizes:
        tbl_pr = sub(style_el, "w:tblPr")
        sub(tbl_pr, "w:tblStyleRowBandSize", **{"w:val": str(band_sizes[0])})
        sub(tbl_pr, "w:tblStyleColBandSize", **{"w:val": str(band_sizes[1])})

    if branches:
        for cond_type, rpr_children in branches.items():
            branch = sub(style_el, "w:tblStylePr", **{"w:type": cond_type})
            rpr = sub(branch, "w:rPr")
            for tag, attrs in rpr_children.items():
                sub(rpr, tag, **(attrs or {}))


def _add_table_with_style(
    doc: Document,
    style_id: str,
    rows: int,
    cols: int,
    *,
    look: dict[str, int] | None = None,
    band_sizes: tuple[int, int] | None = None,
):
    """Add a fresh table to ``doc`` with ``style_id`` applied to ``tblPr``.

    ``look`` names only the ``tblLook`` flags to change; the rest keep
    Word's defaults. Pass ``look={}`` to restate the defaults explicitly.
    """
    table = doc.add_table(rows=rows, cols=cols)
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = sub(table._tbl, "w:tblPr")
    sub(tbl_pr, "w:tblStyle", **{"w:val": style_id})

    if band_sizes:
        sub(tbl_pr, "w:tblStyleRowBandSize", **{"w:val": str(band_sizes[0])})
        sub(tbl_pr, "w:tblStyleColBandSize", **{"w:val": str(band_sizes[1])})

    if look is not None:
        for existing in tbl_pr.findall(qn("w:tblLook")):
            tbl_pr.remove(existing)
        flags = {**WORD_DEFAULT_LOOK, **look}
        sub(tbl_pr, "w:tblLook", **{f"w:{k}": str(v) for k, v in flags.items()})
    return table


def _drop_look(table) -> None:
    """Remove the table's ``<w:tblLook>`` entirely."""
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    for existing in tbl_pr.findall(qn("w:tblLook")):
        tbl_pr.remove(existing)


def _colors(table, cells) -> list[str | None]:
    """Resolve ``color_rgb`` for a list of ``(row, col)`` positions."""
    return [resolve_effective_formatting(table.rows[r].cells[c]).color_rgb for r, c in cells]


# --------------------------------------------------------------------------
# TableContext — defaults and field values.
# --------------------------------------------------------------------------


def test_table_context_position_defaults_are_all_false() -> None:
    ctx = TableContext()
    assert ctx.is_first_row is False
    assert ctx.is_last_row is False
    assert ctx.is_first_col is False
    assert ctx.is_last_col is False
    assert ctx.is_band_row is False
    assert ctx.is_band_col is False
    assert ctx.is_band2_row is False
    assert ctx.is_band2_col is False


def test_table_context_look_flags_default_to_enabled() -> None:
    """A hand-built context gates nothing.

    Word treats a table with no ``<w:tblLook>`` as wanting every branch,
    so the defaults match that rather than the all-False position fields.
    """
    ctx = TableContext()
    assert ctx.first_row_enabled is True
    assert ctx.last_row_enabled is True
    assert ctx.first_col_enabled is True
    assert ctx.last_col_enabled is True


def test_table_context_is_frozen() -> None:
    """TableContext is immutable (frozen dataclass)."""
    ctx = TableContext(is_first_row=True)
    with pytest.raises(dataclasses.FrozenInstanceError):
        ctx.is_first_row = False  # type: ignore[misc]


# --------------------------------------------------------------------------
# Auto-derive — TableContext is built from the cell's table position.
# --------------------------------------------------------------------------


def test_first_row_branch_applies_to_top_cell() -> None:
    """A 3x3 table with a ``firstRow`` branch sets bold on row 0 cells."""
    doc = Document()
    _add_table_style(doc, "ConditionalA", branches={"firstRow": {"w:b": None}})
    table = _add_table_with_style(doc, "ConditionalA", rows=3, cols=3)

    assert resolve_effective_formatting(table.rows[0].cells[0]).bold is True
    assert resolve_effective_formatting(table.rows[1].cells[1]).bold is None


def test_last_row_branch_applies_to_bottom_cell() -> None:
    """``lastRow`` needs its tblLook flag — Word's default leaves it off."""
    doc = Document()
    _add_table_style(doc, "ConditionalB", branches={"lastRow": {"w:i": None}})
    table = _add_table_with_style(doc, "ConditionalB", rows=3, cols=3, look={"lastRow": 1})

    assert resolve_effective_formatting(table.rows[2].cells[1]).italic is True
    assert resolve_effective_formatting(table.rows[1].cells[1]).italic is None


# --------------------------------------------------------------------------
# tblLook gates every conditional branch.
# --------------------------------------------------------------------------


def test_tbl_look_suppresses_the_first_row_branch() -> None:
    """Row 0 of a table whose tblLook clears firstRow takes no firstRow rPr."""
    doc = Document()
    _add_table_style(doc, "Gated", branches={"firstRow": {"w:b": None}})
    table = _add_table_with_style(doc, "Gated", rows=3, cols=3, look={"firstRow": 0})

    assert resolve_effective_formatting(table.rows[0].cells[0]).bold is None


def test_tbl_look_suppresses_the_last_row_branch_by_default() -> None:
    doc = Document()
    _add_table_style(doc, "GatedLast", branches={"lastRow": {"w:b": None}})
    table = _add_table_with_style(doc, "GatedLast", rows=3, cols=3)

    assert resolve_effective_formatting(table.rows[2].cells[0]).bold is None


def test_absent_tbl_look_enables_every_branch() -> None:
    """No ``<w:tblLook>`` means "everything on", not "everything off"."""
    doc = Document()
    _add_table_style(
        doc,
        "NoLook",
        branches={"firstRow": {"w:b": None}, "lastRow": {"w:i": None}},
    )
    table = _add_table_with_style(doc, "NoLook", rows=3, cols=3)
    _drop_look(table)

    assert resolve_effective_formatting(table.rows[0].cells[0]).bold is True
    assert resolve_effective_formatting(table.rows[2].cells[0]).italic is True


def test_legacy_val_bitmask_gates_branches_when_attributes_are_absent() -> None:
    """Word 2007 wrote only ``w:val``; Word still honours it.

    0x0040 is lastRow alone, so firstRow must be suppressed even though
    the cell sits in row 0.
    """
    doc = Document()
    _add_table_style(
        doc,
        "Legacy",
        branches={"firstRow": {"w:b": None}, "lastRow": {"w:i": None}},
    )
    table = _add_table_with_style(doc, "Legacy", rows=3, cols=3)
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    for existing in tbl_pr.findall(qn("w:tblLook")):
        tbl_pr.remove(existing)
    sub(tbl_pr, "w:tblLook", **{"w:val": "0040"})

    assert resolve_effective_formatting(table.rows[0].cells[0]).bold is None
    assert resolve_effective_formatting(table.rows[2].cells[0]).italic is True


def test_unparseable_val_falls_back_to_enabling_everything() -> None:
    doc = Document()
    _add_table_style(doc, "BadVal", branches={"lastRow": {"w:i": None}})
    table = _add_table_with_style(doc, "BadVal", rows=3, cols=3)
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    for existing in tbl_pr.findall(qn("w:tblLook")):
        tbl_pr.remove(existing)
    sub(tbl_pr, "w:tblLook", **{"w:val": "not-hex"})

    assert resolve_effective_formatting(table.rows[2].cells[0]).italic is True


def test_corner_branch_needs_both_of_its_axes_enabled() -> None:
    """With firstColumn cleared, the top-left cell is firstRow, not nwCell."""
    doc = Document()
    _add_table_style(
        doc,
        "CornerGated",
        branches={
            "firstRow": {"w:color": {"w:val": "AAAAAA"}},
            "nwCell": {"w:color": {"w:val": "BBBBBB"}},
        },
    )
    table = _add_table_with_style(doc, "CornerGated", rows=3, cols=3, look={"firstColumn": 0})

    assert resolve_effective_formatting(table.rows[0].cells[0]).color_rgb == "AAAAAA"


# --------------------------------------------------------------------------
# Banding — requires a declared band size, and starts at row/column 0.
# --------------------------------------------------------------------------


def test_band_branches_need_a_declared_band_size() -> None:
    """No ``tblStyleRowBandSize`` anywhere means no banding at all.

    Absent is zero, not one. Word's own table styles all declare an
    explicit ``<w:tblStyleRowBandSize w:val="1"/>`` for exactly this
    reason.
    """
    doc = Document()
    _add_table_style(doc, "NoSize", branches={"band1Horz": {"w:color": {"w:val": "FF0000"}}})
    table = _add_table_with_style(doc, "NoSize", rows=4, cols=2)

    assert _colors(table, [(0, 0), (1, 0), (2, 0), (3, 0)]) == [None] * 4


def test_unparseable_band_size_falls_through_to_the_style_chain() -> None:
    """A junk instance value is skipped, not treated as zero."""
    doc = Document()
    _add_table_style(
        doc,
        "JunkSize",
        branches={"band1Horz": {"w:color": {"w:val": "FF0000"}}},
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "JunkSize", rows=4, cols=2)
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    sub(tbl_pr, "w:tblStyleRowBandSize", **{"w:val": "not-a-number"})
    sub(tbl_pr, "w:tblStyleColBandSize")  # no w:val at all

    # The style's own size of 1 still applies.
    assert _colors(table, [(0, 0), (1, 0), (2, 0), (3, 0)]) == ["FF0000", None, "FF0000", None]


def test_negative_band_size_disables_banding() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "NegSize",
        branches={"band1Horz": {"w:color": {"w:val": "FF0000"}}},
        band_sizes=(-2, -2),
    )
    table = _add_table_with_style(doc, "NegSize", rows=4, cols=2)

    assert _colors(table, [(0, 0), (1, 0), (2, 0), (3, 0)]) == [None] * 4


def test_table_style_reference_without_a_value_is_ignored() -> None:
    """A ``<w:tblStyle/>`` with no ``w:val`` resolves to no table style."""
    doc = Document()
    _add_table_style(doc, "Unreferenced", branches={"firstRow": {"w:b": None}})
    table = doc.add_table(rows=2, cols=2)
    sub(table._tbl.find(qn("w:tblPr")), "w:tblStyle")

    assert resolve_effective_formatting(table.rows[0].cells[0]).bold is None


def test_explicit_zero_band_size_disables_banding() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "ZeroSize",
        branches={"band1Horz": {"w:color": {"w:val": "FF0000"}}},
        band_sizes=(0, 0),
    )
    table = _add_table_with_style(doc, "ZeroSize", rows=4, cols=2)

    assert _colors(table, [(0, 0), (1, 0), (2, 0), (3, 0)]) == [None] * 4


def test_band_size_is_read_from_the_style_chain() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "StyleSized",
        branches={"band1Horz": {"w:color": {"w:val": "FF0000"}}},
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "StyleSized", rows=4, cols=2)

    # No firstRow branch, so the stripes start at row 0: band1 on 0 and 2.
    assert _colors(table, [(0, 0), (1, 0), (2, 0), (3, 0)]) == ["FF0000", None, "FF0000", None]


def test_instance_band_size_overrides_the_style_band_size() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "Conflict",
        branches={
            "band1Horz": {"w:color": {"w:val": "111111"}},
            "band2Horz": {"w:color": {"w:val": "222222"}},
        },
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "Conflict", rows=6, cols=1, band_sizes=(3, 3))

    expected = ["111111"] * 3 + ["222222"] * 3
    assert _colors(table, [(r, 0) for r in range(6)]) == expected


def test_horizontal_bands_start_at_row_zero_without_a_first_row_branch() -> None:
    """The tblLook flag alone does not shift the stripe sequence.

    ``firstRow`` is set in the default look here, but the style defines no
    ``firstRow`` branch, so nothing claims row 0 and it bands normally.
    """
    doc = Document()
    _add_table_style(
        doc,
        "BandFromZero",
        branches={
            "band1Horz": {"w:color": {"w:val": "111111"}},
            "band2Horz": {"w:color": {"w:val": "222222"}},
        },
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "BandFromZero", rows=4, cols=1)

    assert _colors(table, [(r, 0) for r in range(4)]) == [
        "111111",
        "222222",
        "111111",
        "222222",
    ]


def test_horizontal_bands_start_at_row_one_when_first_row_claims_row_zero() -> None:
    """A live ``firstRow`` branch takes row 0 out of the stripe sequence."""
    doc = Document()
    _add_table_style(
        doc,
        "BandAfterHeader",
        branches={
            "firstRow": {"w:color": {"w:val": "FFFFFF"}},
            "band1Horz": {"w:color": {"w:val": "111111"}},
            "band2Horz": {"w:color": {"w:val": "222222"}},
        },
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "BandAfterHeader", rows=4, cols=1)

    assert _colors(table, [(r, 0) for r in range(4)]) == [
        "FFFFFF",
        "111111",
        "222222",
        "111111",
    ]


def test_vertical_bands_need_their_tbl_look_flag() -> None:
    """Word's default look clears vertical banding."""
    doc = Document()
    _add_table_style(
        doc,
        "VertDefault",
        branches={"band1Vert": {"w:color": {"w:val": "00FF00"}}},
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "VertDefault", rows=1, cols=4)

    assert _colors(table, [(0, c) for c in range(4)]) == [None] * 4


def test_vertical_bands_start_at_column_zero_without_a_first_col_branch() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "VertBands",
        branches={
            "band1Vert": {"w:color": {"w:val": "111111"}},
            "band2Vert": {"w:color": {"w:val": "222222"}},
        },
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "VertBands", rows=1, cols=4, look={"noVBand": 0})

    assert _colors(table, [(0, c) for c in range(4)]) == [
        "111111",
        "222222",
        "111111",
        "222222",
    ]


def test_vertical_bands_start_at_column_one_when_first_col_claims_column_zero() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "VertAfterHeader",
        branches={
            "firstCol": {"w:color": {"w:val": "FFFFFF"}},
            "band1Vert": {"w:color": {"w:val": "111111"}},
            "band2Vert": {"w:color": {"w:val": "222222"}},
        },
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "VertAfterHeader", rows=1, cols=4, look={"noVBand": 0})

    assert _colors(table, [(0, c) for c in range(4)]) == [
        "FFFFFF",
        "111111",
        "222222",
        "111111",
    ]


def test_row_band_size_two_groups_rows_in_pairs() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "Pairs",
        branches={
            "band1Horz": {"w:color": {"w:val": "111111"}},
            "band2Horz": {"w:color": {"w:val": "222222"}},
        },
    )
    table = _add_table_with_style(doc, "Pairs", rows=7, cols=1, band_sizes=(2, 2))

    expected = [
        "111111",  # row 0 — band1 stripe 0
        "111111",  # row 1 — band1 stripe 0
        "222222",  # row 2 — band2 stripe 1
        "222222",  # row 3 — band2 stripe 1
        "111111",  # row 4 — band1 stripe 2
        "111111",  # row 5 — band1 stripe 2
        "222222",  # row 6 — band2 stripe 3
    ]
    assert _colors(table, [(r, 0) for r in range(7)]) == expected


def test_col_band_size_three_groups_columns_in_triples() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "Triples",
        branches={"band1Vert": {"w:color": {"w:val": "333333"}}},
    )
    table = _add_table_with_style(
        doc, "Triples", rows=1, cols=7, look={"noVBand": 0}, band_sizes=(3, 3)
    )

    # Cols 0-2 → band1 stripe 0; 3-5 → band2 stripe 1; 6 → band1 stripe 2.
    expected = ["333333", "333333", "333333", None, None, None, "333333"]
    assert _colors(table, [(0, c) for c in range(7)]) == expected


# --------------------------------------------------------------------------
# Precedence — measured against Word, and NOT the order ECMA-376 lists.
# --------------------------------------------------------------------------


def test_vertical_band_overrides_horizontal_band() -> None:
    """A cell in both stripes takes the vertical branch.

    The spec lists the vertical bands *before* the horizontal ones, which
    would make horizontal win. Word does the opposite.
    """
    doc = Document()
    _add_table_style(
        doc,
        "BothBands",
        branches={
            "band1Horz": {"w:color": {"w:val": "111111"}},
            "band1Vert": {"w:color": {"w:val": "222222"}},
        },
        band_sizes=(2, 2),
    )
    table = _add_table_with_style(doc, "BothBands", rows=2, cols=2, look={"noVBand": 0})

    assert resolve_effective_formatting(table.rows[0].cells[0]).color_rgb == "222222"


def test_first_col_overrides_a_horizontal_band() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "ColOverBand",
        branches={
            "band1Horz": {"w:color": {"w:val": "111111"}},
            "firstCol": {"w:color": {"w:val": "222222"}},
        },
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "ColOverBand", rows=2, cols=2)

    assert resolve_effective_formatting(table.rows[0].cells[0]).color_rgb == "222222"
    assert resolve_effective_formatting(table.rows[0].cells[1]).color_rgb == "111111"


def test_first_row_overrides_a_vertical_band() -> None:
    doc = Document()
    _add_table_style(
        doc,
        "RowOverBand",
        branches={
            "band1Vert": {"w:color": {"w:val": "111111"}},
            "firstRow": {"w:color": {"w:val": "222222"}},
        },
        band_sizes=(1, 1),
    )
    table = _add_table_with_style(doc, "RowOverBand", rows=2, cols=2, look={"noVBand": 0})

    assert resolve_effective_formatting(table.rows[0].cells[0]).color_rgb == "222222"
    assert resolve_effective_formatting(table.rows[1].cells[0]).color_rgb == "111111"


def test_first_row_overrides_first_col_at_corner_without_corner_branch() -> None:
    """Row branches win at row/col intersections.

    ECMA-376 17.7.6.5 lists the column branches after the row ones, which
    would make ``firstCol`` win here. Measured against Word, the row
    branch wins — the spec's listing is not the application order.
    """
    doc = Document()
    _add_table_style(
        doc,
        "RowVsCol",
        branches={
            "firstRow": {"w:color": {"w:val": "AAAAAA"}},
            "firstCol": {"w:color": {"w:val": "BBBBBB"}},
        },
    )
    table = _add_table_with_style(doc, "RowVsCol", rows=3, cols=3)

    nw_cell = table.rows[0].cells[0]  # matches both firstRow and firstCol
    ne_cell = table.rows[0].cells[2]  # matches firstRow only
    sw_cell = table.rows[2].cells[0]  # matches firstCol only

    assert resolve_effective_formatting(nw_cell).color_rgb == "AAAAAA"  # firstRow wins
    assert resolve_effective_formatting(ne_cell).color_rgb == "AAAAAA"
    assert resolve_effective_formatting(sw_cell).color_rgb == "BBBBBB"


def test_single_row_table_last_row_overrides_first_row() -> None:
    """A 1-row table matches both firstRow and lastRow — lastRow wins."""
    doc = Document()
    _add_table_style(
        doc,
        "OneRow",
        branches={
            "firstRow": {"w:color": {"w:val": "111111"}},
            "lastRow": {"w:color": {"w:val": "222222"}},
        },
    )
    table = _add_table_with_style(doc, "OneRow", rows=1, cols=3, look={"lastRow": 1})

    assert resolve_effective_formatting(table.rows[0].cells[1]).color_rgb == "222222"


def test_single_column_table_last_col_overrides_first_col() -> None:
    """A 1-column table matches both firstCol and lastCol — lastCol wins."""
    doc = Document()
    _add_table_style(
        doc,
        "OneCol",
        branches={
            "firstCol": {"w:color": {"w:val": "333333"}},
            "lastCol": {"w:color": {"w:val": "444444"}},
        },
    )
    table = _add_table_with_style(doc, "OneCol", rows=3, cols=1, look={"lastColumn": 1})

    assert resolve_effective_formatting(table.rows[1].cells[0]).color_rgb == "444444"


def test_corner_overrides_first_row() -> None:
    """``nwCell`` beats ``firstRow`` at the top-left cell."""
    doc = Document()
    _add_table_style(
        doc,
        "Corners",
        branches={
            "firstRow": {"w:color": {"w:val": "AAAAAA"}},
            "nwCell": {"w:color": {"w:val": "BBBBBB"}},
        },
    )
    table = _add_table_with_style(doc, "Corners", rows=3, cols=3)

    assert resolve_effective_formatting(table.rows[0].cells[0]).color_rgb == "BBBBBB"
    # NE matches firstRow but not nwCell — falls back to firstRow color.
    assert resolve_effective_formatting(table.rows[0].cells[2]).color_rgb == "AAAAAA"


def test_child_base_overrides_parent_conditional_branch() -> None:
    """H9 regression: per-level base + conditional must interleave.

    Per ECMA-376 17.7.6.5 each style level computes (base then matching
    conditionals); the resulting per-level state then cascades child-
    over-parent. So a child style's BASE rPr must override a parent's
    matching conditional branch. The buggy implementation walked the
    whole chain for base first, then the whole chain for conditionals,
    inverting this at the parent/child boundary.
    """
    doc = Document()
    _add_table_style(
        doc,
        "ParentWithFirstRow",
        branches={"firstRow": {"w:color": {"w:val": "FFA500"}}},
    )
    _add_table_style(
        doc,
        "ChildBaseGreen",
        base_rpr={"w:color": {"w:val": "00FF00"}},
        based_on="ParentWithFirstRow",
    )

    table = _add_table_with_style(doc, "ChildBaseGreen", rows=3, cols=2)
    # Order at row 0: parent base (none) → parent firstRow (ORANGE) →
    # child base (GREEN) → child firstRow (none). GREEN wins.
    assert resolve_effective_formatting(table.rows[0].cells[0]).color_rgb == "00FF00"


def test_conditional_branches_inherit_down_a_based_on_chain() -> None:
    """A child style picks up its parent's branches and band size."""
    doc = Document()
    _add_table_style(
        doc,
        "ChainParent",
        branches={"firstRow": {"w:color": {"w:val": "AAAAAA"}}},
        band_sizes=(1, 1),
    )
    _add_table_style(
        doc,
        "ChainChild",
        based_on="ChainParent",
        branches={"band1Horz": {"w:color": {"w:val": "BBBBBB"}}},
    )
    table = _add_table_with_style(doc, "ChainChild", rows=3, cols=1)

    # The inherited firstRow branch claims row 0, so the bands start at 1.
    assert _colors(table, [(0, 0), (1, 0), (2, 0)]) == ["AAAAAA", "BBBBBB", None]


def test_whole_table_branch_is_ignored() -> None:
    """Word discards ``<w:tblStylePr w:type="wholeTable">`` outright.

    It neither renders the branch's ``rPr`` nor keeps the element when it
    saves the file. Whole-table formatting belongs on the style's own
    ``w:rPr``, which the base pass already applies — so the style's 12pt
    base wins here and the branch's 10pt is never seen.
    """
    doc = Document()
    _add_table_style(
        doc,
        "BaseAndWhole",
        base_rpr={"w:sz": {"w:val": "24"}},  # 12pt
        branches={
            "wholeTable": {"w:sz": {"w:val": "20"}},  # 10pt — inert
            "firstRow": {"w:b": None},
        },
    )
    table = _add_table_with_style(doc, "BaseAndWhole", rows=2, cols=2)

    top = resolve_effective_formatting(table.rows[0].cells[0])
    assert top.bold is True
    assert top.font_size == 12.0

    middle = resolve_effective_formatting(table.rows[1].cells[1])
    assert middle.bold is None
    assert middle.font_size == 12.0


def test_whole_table_branch_alone_contributes_nothing() -> None:
    """With no base rPr to fall back on, the value simply never appears."""
    doc = Document()
    _add_table_style(doc, "WholeOnly", branches={"wholeTable": {"w:color": {"w:val": "FF0000"}}})
    table = _add_table_with_style(doc, "WholeOnly", rows=2, cols=2)

    assert resolve_effective_formatting(table.rows[0].cells[0]).color_rgb is None


# --------------------------------------------------------------------------
# Manual override — passing an explicit TableContext.
# --------------------------------------------------------------------------


def test_manual_table_context_overrides_auto_derived() -> None:
    """An explicit ``table_context`` arg supersedes auto-derivation."""
    doc = Document()
    _add_table_style(doc, "FirstRowBold", branches={"firstRow": {"w:b": None}})
    table = _add_table_with_style(doc, "FirstRowBold", rows=3, cols=2)

    # Middle cell — auto-derived TableContext has is_first_row=False.
    middle_cell = table.rows[1].cells[0]
    assert resolve_effective_formatting(middle_cell).bold is None

    # Same cell, but the caller forces the first-row position.
    overridden = resolve_effective_formatting(
        middle_cell, table_context=TableContext(is_first_row=True)
    )
    assert overridden.bold is True


def test_manual_table_context_can_gate_a_branch_off() -> None:
    """The ``*_enabled`` flags work on a hand-built context too."""
    doc = Document()
    _add_table_style(doc, "GatedManual", branches={"firstRow": {"w:b": None}})
    table = _add_table_with_style(doc, "GatedManual", rows=3, cols=2)

    cell = table.rows[0].cells[0]
    gated = resolve_effective_formatting(
        cell, table_context=TableContext(is_first_row=True, first_row_enabled=False)
    )
    assert gated.bold is None


def test_non_table_target_has_no_conditional_formatting() -> None:
    """A regular paragraph never picks up table-style branches."""
    doc = Document()
    _add_table_style(doc, "ShouldNotApply", branches={"firstRow": {"w:b": None}})
    p = doc.add_paragraph("standalone text")

    assert resolve_effective_formatting(p).bold is None


# --------------------------------------------------------------------------
# Paragraphs and runs inside table cells also receive conditional formatting.
# --------------------------------------------------------------------------


def test_conditional_formatting_reaches_paragraph_in_cell() -> None:
    doc = Document()
    _add_table_style(doc, "FirstRowItalic", branches={"firstRow": {"w:i": None}})
    table = _add_table_with_style(doc, "FirstRowItalic", rows=2, cols=2)

    top_cell = table.rows[0].cells[0]
    top_cell.text = "header"

    assert resolve_effective_formatting(top_cell.paragraphs[0]).italic is True


def test_conditional_formatting_reaches_run_in_cell() -> None:
    doc = Document()
    _add_table_style(doc, "FirstColBold", branches={"firstCol": {"w:b": None}})
    table = _add_table_with_style(doc, "FirstColBold", rows=2, cols=3)

    run = table.rows[1].cells[0].paragraphs[0].add_run("data")

    assert resolve_effective_formatting(run).bold is True


# --------------------------------------------------------------------------
# gridSpan. A horizontally merged cell occupies several grid columns, so a
# row's `w:tc` positions and its grid columns diverge after the first merge —
# and it is the grid column that decides vertical banding and is_last_col.
# --------------------------------------------------------------------------


def _merge_first_two_cells(table) -> None:
    """Turn cells 0 and 1 of row 1 into one `w:tc` spanning two grid columns."""
    row = table.rows[1]
    first, second = row._tr.findall(qn("w:tc"))[:2]
    tc_pr = first.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = sub(first, "w:tcPr")
    sub(tc_pr, "w:gridSpan", **{"w:val": "2"})
    row._tr.remove(second)


def test_vertical_banding_counts_grid_columns_not_tc_elements() -> None:
    """The cell after a two-column merge is in the band its grid position says.

    With `colBandSize=1` the stripe alternates every grid column. Counting
    `w:tc` elements put the cell right of a `gridSpan="2"` merge one stripe
    early — reported band2Vert where the grid says band1Vert.
    """
    doc = Document()
    _add_table_style(
        doc,
        "Banded",
        branches={
            "band1Vert": {"w:color": {"w:val": "111111"}},
            "band2Vert": {"w:color": {"w:val": "222222"}},
        },
    )
    table = _add_table_with_style(
        doc,
        "Banded",
        rows=2,
        cols=4,
        look={"noVBand": 0, "firstRow": 0, "firstColumn": 0},
        band_sizes=(1, 1),
    )
    _merge_first_two_cells(table)

    # Row 1 now holds three `w:tc`: grid columns 0-1 (merged), 2, and 3.
    cells = table.rows[1]._tr.findall(qn("w:tc"))
    resolved = [resolve_effective_formatting(_cell_paragraph(c, table)).color_rgb for c in cells]

    # Grid column 0 -> band1, grid column 2 -> band1, grid column 3 -> band2.
    assert resolved == ["111111", "111111", "222222"]


def test_is_last_col_accounts_for_a_merge_widening_the_row() -> None:
    """`lastCol` still lands on the rightmost cell once a merge widens the row.

    Not a discriminating test on its own — counting `w:tc` and counting grid
    columns agree about which cell is *last*. It is here because the merge
    path is what changed, and this is the property that must not regress
    while band membership is being recomputed from the grid.
    """
    doc = Document()
    _add_table_style(doc, "LastCol", branches={"lastCol": {"w:color": {"w:val": "333333"}}})
    table = _add_table_with_style(doc, "LastCol", rows=2, cols=4, look={"lastColumn": 1})
    _merge_first_two_cells(table)

    cells = table.rows[1]._tr.findall(qn("w:tc"))
    resolved = [resolve_effective_formatting(_cell_paragraph(c, table)).color_rgb for c in cells]

    assert resolved[-1] == "333333"
    assert resolved[:-1] == [None, None]


def test_grid_column_counts_spans_not_cells() -> None:
    """The arithmetic itself: `w:tc` position and grid column diverge at a merge."""
    doc = Document()
    table = _add_table_with_style(doc, "Plain", rows=2, cols=4, look={})
    _merge_first_two_cells(table)
    cells = table.rows[1]._tr.findall(qn("w:tc"))

    assert [_grid_column(cells, tc) for tc in cells] == [0, 2, 3]
    assert [_grid_span(tc) for tc in cells] == [2, 1, 1]


def test_grid_column_rejects_a_cell_from_another_row() -> None:
    doc = Document()
    table = _add_table_with_style(doc, "Plain", rows=2, cols=2, look={})
    row0 = table.rows[0]._tr.findall(qn("w:tc"))
    stranger = table.rows[1]._tr.findall(qn("w:tc"))[0]

    with pytest.raises(ValueError, match="not in this row"):
        _grid_column(row0, stranger)


@pytest.mark.parametrize("raw", ["0", "-3", "notanumber"])
def test_a_malformed_grid_span_falls_back_to_one(raw: str) -> None:
    """A span below 1 would make two cells claim the same grid column."""
    doc = Document()
    table = _add_table_with_style(doc, "Plain", rows=1, cols=2, look={})
    tc = table.rows[0]._tr.findall(qn("w:tc"))[0]
    sub(sub(tc, "w:tcPr"), "w:gridSpan", **{"w:val": raw})

    assert _grid_span(tc) == 1


def _cell_paragraph(tc, table):
    """Wrap a raw `w:tc` as a _Cell so the resolver accepts it."""
    from docx.table import _Cell

    return _Cell(tc, table)
