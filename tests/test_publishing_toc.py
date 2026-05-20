"""Tests for ``docx_plus.publishing.add_toc``."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.publishing import add_toc


def _instruction(field_run):
    """Extract the instruction string from the field that ``field_run`` opens."""
    p = field_run.getparent()
    instr = p.find(qn("w:r") + "/" + qn("w:instrText"))
    # Fallback: search any run for instrText (the begin-run has no instrText).
    if instr is None:
        for r in p.findall(qn("w:r")):
            t = r.find(qn("w:instrText"))
            if t is not None:
                return t.text
        return None
    return instr.text


# --------------------------------------------------------------------------
# Emission — instruction string assembly.
# --------------------------------------------------------------------------


def test_add_toc_defaults_match_word_output() -> None:
    doc = Document()
    field = add_toc(doc.add_paragraph())
    instr = _instruction(field)
    assert instr is not None
    assert 'TOC \\o "1-3"' in instr
    assert "\\h" in instr
    assert "\\z" in instr
    assert "\\u" in instr
    assert "\\n" not in instr  # page_numbers default True


def test_add_toc_custom_levels() -> None:
    doc = Document()
    field = add_toc(doc.add_paragraph(), levels=(2, 5))
    instr = _instruction(field)
    assert '\\o "2-5"' in instr


def test_add_toc_no_hyperlink() -> None:
    doc = Document()
    field = add_toc(doc.add_paragraph(), hyperlink=False)
    instr = _instruction(field)
    assert "\\h" not in instr


def test_add_toc_no_page_numbers_emits_n_switch() -> None:
    doc = Document()
    field = add_toc(doc.add_paragraph(), page_numbers=False)
    instr = _instruction(field)
    assert "\\n" in instr


# --------------------------------------------------------------------------
# Field structure — five-run complex-field sequence.
# --------------------------------------------------------------------------


def test_add_toc_writes_full_complex_field() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_toc(p)

    fld_chars = xpath(p._p, ".//w:fldChar")
    char_types = [c.get(qn("w:fldCharType")) for c in fld_chars]
    assert char_types == ["begin", "separate", "end"]


def test_add_toc_returns_begin_run() -> None:
    doc = Document()
    field = add_toc(doc.add_paragraph())
    assert field.tag == qn("w:r")
    fld_char = field.find(qn("w:fldChar"))
    assert fld_char is not None
    assert fld_char.get(qn("w:fldCharType")) == "begin"


# --------------------------------------------------------------------------
# Round-trip — save / reopen preserves the field.
# --------------------------------------------------------------------------


def test_add_toc_round_trip(tmp_path: Path) -> None:
    doc = Document()
    add_toc(doc.add_paragraph(), levels=(1, 4), page_numbers=False)
    out = tmp_path / "toc.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    body = reopened.element.body
    instr_runs = xpath(body, ".//w:instrText")
    assert len(instr_runs) == 1
    text = instr_runs[0].text
    assert 'TOC \\o "1-4"' in text
    assert "\\n" in text


# --------------------------------------------------------------------------
# Validation — H12 (levels range) and H13 (additional_styles).
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "bad_levels",
    [
        (3, 1),       # reversed
        (0, 3),       # below 1
        (1, 10),      # above 9
        (-1, 3),      # negative
        (1,),         # wrong arity (1-tuple)
        (1, 2, 3),    # wrong arity (3-tuple)
        "1-3",        # not a tuple
        5,            # bare int
        (1.0, 3.0),   # non-int
    ],
)
def test_add_toc_rejects_bad_levels(bad_levels) -> None:
    doc = Document()
    with pytest.raises(ValueError, match="levels"):
        add_toc(doc.add_paragraph(), levels=bad_levels)


def test_add_toc_additional_styles_emits_t_switch() -> None:
    """H13: ``additional_styles`` plumbs to ``\\t "Style1,1,Style2,2"``."""
    doc = Document()
    p = doc.add_paragraph()
    add_toc(p, levels=(1, 2), additional_styles=[("Caption", 4), ("Quote", 5)])
    instructions = xpath(p._p, ".//w:instrText")
    assert len(instructions) == 1
    text = instructions[0].text
    assert 'TOC \\o "1-2"' in text
    assert '\\t "Caption,4,Quote,5"' in text


@pytest.mark.parametrize(
    "bad_styles",
    [
        [("Caption",)],                   # arity 1
        [("Caption", 4, "extra")],        # arity 3
        [("Caption", 0)],                 # level below 1
        [("Caption", 10)],                # level above 9
        [("", 4)],                        # empty style name
        [('Caption" \\o "1-9', 4)],       # double-quote injection
        [("Cap,tion", 4)],                # comma in name
        [(123, 4)],                       # non-str name
        [("Caption", "4")],               # non-int level
        "not iterable as pairs",          # bare str
    ],
)
def test_add_toc_rejects_bad_additional_styles(bad_styles) -> None:
    doc = Document()
    with pytest.raises(ValueError, match="additional_styles"):
        add_toc(doc.add_paragraph(), additional_styles=bad_styles)


def test_add_toc_additional_styles_none_omits_t_switch() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_toc(p)  # additional_styles default is None
    instructions = xpath(p._p, ".//w:instrText")
    assert "\\t" not in instructions[0].text
