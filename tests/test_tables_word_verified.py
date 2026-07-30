"""Conditional table-style formatting as measured against live Microsoft Word.

Every expected grid in this file was read out of Word over COM, not derived
from the spec: a probe document was built for each shape, opened in Word,
and each cell's effective font size recorded. The resolver previously
followed ECMA-376 17.7.6.5's prose and diverged from Word on six separate
counts — see :data:`docx_plus.styles.inspect._TBL_STYLE_PR_ORDER`.

The measurements are expressed as font sizes because a size is an absolute
value: one distinct size per conditional branch means a cell's size names
the branch that won, with no ambiguity about how values combined.

Re-measuring: build the same shapes, open the document in Word, and read
each cell with ``wordlive read format --anchor-id table:N:R:C``.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from docx_plus.core.oxml import sub
from docx_plus.styles import resolve_effective_formatting

# One distinct point size per branch, so a resolved size names the winner.
# 10.0 is the style's own base rPr — i.e. "no conditional branch applied".
BRANCH_SIZE = {
    "base": 10.0,
    "band1Vert": 11.0,
    "band2Vert": 12.0,
    "band1Horz": 13.0,
    "band2Horz": 14.0,
    "firstRow": 15.0,
    "lastRow": 16.0,
    "firstCol": 17.0,
    "lastCol": 18.0,
    "nwCell": 19.0,
    "neCell": 20.0,
    "swCell": 21.0,
    "seCell": 22.0,
}
ALL_BRANCHES = tuple(name for name in BRANCH_SIZE if name != "base")

# tblLook flags. "off" is every conditional disabled; the probes name only
# what they turn on.
LOOK_OFF = {
    "firstRow": 0,
    "lastRow": 0,
    "firstColumn": 0,
    "lastColumn": 0,
    "noHBand": 1,
    "noVBand": 1,
}


def _look(**overrides: int) -> dict[str, int]:
    return {**LOOK_OFF, **overrides}


def _build(
    branches: tuple[str, ...],
    look: dict[str, int] | None,
    rows: int,
    cols: int,
    *,
    band_sizes: tuple[int, int] | None = (1, 1),
    instance_band_sizes: tuple[int, int] | None = None,
):
    """A document with one probe table style and one table using it."""
    doc = Document()
    style = sub(doc.styles.element, "w:style", **{"w:type": "table", "w:styleId": "T"})
    sub(style, "w:name", **{"w:val": "T"})
    rpr = sub(style, "w:rPr")
    sub(rpr, "w:sz", **{"w:val": str(int(BRANCH_SIZE["base"] * 2))})
    if band_sizes is not None:
        tbl_pr = sub(style, "w:tblPr")
        sub(tbl_pr, "w:tblStyleRowBandSize", **{"w:val": str(band_sizes[0])})
        sub(tbl_pr, "w:tblStyleColBandSize", **{"w:val": str(band_sizes[1])})
    for branch in branches:
        branch_el = sub(style, "w:tblStylePr", **{"w:type": branch})
        branch_rpr = sub(branch_el, "w:rPr")
        sub(branch_rpr, "w:sz", **{"w:val": str(int(BRANCH_SIZE[branch] * 2))})

    table = doc.add_table(rows=rows, cols=cols)
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    # python-docx's StyleFactory cannot wrap a hand-built style element, so
    # the reference goes straight into tblPr rather than through .style.
    sub(tbl_pr, "w:tblStyle", **{"w:val": "T"})
    if instance_band_sizes is not None:
        sub(tbl_pr, "w:tblStyleRowBandSize", **{"w:val": str(instance_band_sizes[0])})
        sub(tbl_pr, "w:tblStyleColBandSize", **{"w:val": str(instance_band_sizes[1])})
    for existing in tbl_pr.findall(qn("w:tblLook")):
        tbl_pr.remove(existing)
    if look is not None:
        sub(tbl_pr, "w:tblLook", **{f"w:{k}": str(v) for k, v in look.items()})
    return table


def _grid(table) -> list[list[float | None]]:
    return [
        [resolve_effective_formatting(cell).font_size for cell in row.cells] for row in table.rows
    ]


# --------------------------------------------------------------------------
# tblLook gating, over a style defining every branch. 4x4, so every cell is
# a corner, an edge, or interior.
# --------------------------------------------------------------------------

GATING_CASES: list[tuple[str, dict[str, int] | None, list[list[float]]]] = [
    (
        "all-on",
        _look(firstRow=1, lastRow=1, firstColumn=1, lastColumn=1, noHBand=0, noVBand=0),
        [
            [19.0, 15.0, 15.0, 20.0],
            [17.0, 10.0, 10.0, 18.0],
            [17.0, 10.0, 10.0, 18.0],
            [21.0, 16.0, 16.0, 22.0],
        ],
    ),
    (
        # Every conditional off: the base rPr is all that is left.
        "all-off",
        _look(),
        [[10.0] * 4 for _ in range(4)],
    ),
    (
        # Word's own default for a new table. lastRow / lastColumn are off,
        # so the NE / SW / SE corners never appear.
        "word-default",
        _look(firstRow=1, firstColumn=1, noHBand=0),
        [
            [19.0, 15.0, 15.0, 15.0],
            [17.0, 10.0, 10.0, 10.0],
            [17.0, 10.0, 10.0, 10.0],
            [17.0, 10.0, 10.0, 10.0],
        ],
    ),
    (
        # firstRow alone: the NW corner needs firstColumn too, so row 0
        # takes firstRow rather than nwCell.
        "row-only",
        _look(firstRow=1),
        [
            [15.0, 15.0, 15.0, 15.0],
            [10.0, 10.0, 10.0, 10.0],
            [10.0, 10.0, 10.0, 10.0],
            [10.0, 10.0, 10.0, 10.0],
        ],
    ),
    (
        # Banding enabled by the look, but this style declares no band size,
        # so the band branches stay inert and nothing paints.
        "bands-enabled-but-unsized",
        _look(noHBand=0, noVBand=0),
        [[10.0] * 4 for _ in range(4)],
    ),
]


@pytest.mark.parametrize(
    ("case_id", "look", "expected"), GATING_CASES, ids=[c[0] for c in GATING_CASES]
)
def test_tbl_look_gates_conditional_branches(
    case_id: str, look: dict[str, int] | None, expected: list[list[float]]
) -> None:
    """What Word rendered for a 4x4 table under each tblLook.

    The style defines every branch but no band size, which is what makes
    the interior cells a clean read of the positional branches alone.
    """
    table = _build(ALL_BRANCHES, look, rows=4, cols=4, band_sizes=None)
    assert _grid(table) == expected


def test_absent_tbl_look_enables_everything() -> None:
    """No ``<w:tblLook>`` renders as though every flag were set.

    firstRow paints row 0 and the horizontal bands run underneath it —
    both of which a look of all-zeroes would have suppressed.
    """
    table = _build(("firstRow", "band1Horz", "band2Horz"), None, rows=5, cols=3)
    assert [row[0] for row in _grid(table)] == [15.0, 13.0, 14.0, 13.0, 14.0]


# --------------------------------------------------------------------------
# Banding: needs a declared size, and the stripe sequence's starting line
# depends on whether a firstRow / firstCol branch claims it.
# --------------------------------------------------------------------------


def test_band_branches_are_inert_without_a_declared_band_size() -> None:
    """Measured: Word paints no bands when no band size exists anywhere."""
    table = _build(
        ("band1Horz", "band2Horz"),
        _look(noHBand=0),
        rows=6,
        cols=2,
        band_sizes=None,
    )
    assert _grid(table) == [[10.0, 10.0] for _ in range(6)]


def test_horizontal_bands_run_from_row_zero() -> None:
    table = _build(("band1Horz", "band2Horz"), _look(noHBand=0), rows=7, cols=2)
    assert [row[0] for row in _grid(table)] == [13.0, 14.0, 13.0, 14.0, 13.0, 14.0, 13.0]


def test_vertical_bands_run_from_column_zero() -> None:
    table = _build(("band1Vert", "band2Vert"), _look(noVBand=0), rows=2, cols=5)
    assert _grid(table)[0] == [11.0, 12.0, 11.0, 12.0, 11.0]


def test_a_live_first_row_branch_shifts_the_horizontal_bands() -> None:
    """Row 0 leaves the stripe sequence when firstRow actually paints it."""
    table = _build(
        ("band1Horz", "band2Horz", "firstRow", "lastRow"),
        _look(firstRow=1, lastRow=1, noHBand=0),
        rows=7,
        cols=2,
    )
    assert [row[0] for row in _grid(table)] == [15.0, 13.0, 14.0, 13.0, 14.0, 13.0, 16.0]


def test_the_flag_alone_does_not_shift_the_bands() -> None:
    """firstRow set but no firstRow branch: nothing claims row 0."""
    table = _build(
        ("band1Horz", "band2Horz"),
        _look(firstRow=1, noHBand=0),
        rows=6,
        cols=2,
    )
    assert [row[0] for row in _grid(table)] == [13.0, 14.0, 13.0, 14.0, 13.0, 14.0]


def test_a_live_first_col_branch_shifts_the_vertical_bands() -> None:
    table = _build(
        ("band1Vert", "band2Vert", "firstCol", "lastCol"),
        _look(firstColumn=1, lastColumn=1, noVBand=0),
        rows=2,
        cols=4,
    )
    assert _grid(table)[0] == [17.0, 11.0, 12.0, 18.0]


def test_band_size_two_widens_the_stripes() -> None:
    table = _build(("band1Horz", "band2Horz"), _look(noHBand=0), rows=7, cols=2, band_sizes=(2, 2))
    assert [row[0] for row in _grid(table)] == [13.0, 13.0, 14.0, 14.0, 13.0, 13.0, 14.0]


def test_instance_band_size_beats_the_style_band_size() -> None:
    table = _build(
        ("band1Horz", "band2Horz"),
        _look(noHBand=0),
        rows=8,
        cols=2,
        band_sizes=(1, 1),
        instance_band_sizes=(3, 3),
    )
    assert [row[0] for row in _grid(table)] == [
        13.0,
        13.0,
        13.0,
        14.0,
        14.0,
        14.0,
        13.0,
        13.0,
    ]


# --------------------------------------------------------------------------
# Precedence contests. Each pairs two branches that can reach one cell.
# --------------------------------------------------------------------------


def test_vertical_band_beats_horizontal_band() -> None:
    """The spec's listing implies the opposite; Word prefers vertical."""
    table = _build(
        ("band1Horz", "band2Horz", "band1Vert", "band2Vert"),
        _look(noHBand=0, noVBand=0),
        rows=7,
        cols=4,
    )
    assert _grid(table) == [[11.0, 12.0, 11.0, 12.0] for _ in range(7)]


def test_first_row_beats_a_vertical_band() -> None:
    table = _build(
        ("firstRow", "band1Vert", "band2Vert"),
        _look(firstRow=1, noVBand=0),
        rows=5,
        cols=5,
    )
    grid = _grid(table)
    assert grid[0] == [15.0] * 5
    assert grid[1] == [11.0, 12.0, 11.0, 12.0, 11.0]


def test_first_col_beats_a_horizontal_band() -> None:
    table = _build(
        ("firstCol", "band1Horz", "band2Horz"),
        _look(firstColumn=1, noHBand=0),
        rows=5,
        cols=5,
    )
    grid = _grid(table)
    assert [row[0] for row in grid] == [17.0] * 5
    assert [row[1] for row in grid] == [13.0, 14.0, 13.0, 14.0, 13.0]


def test_last_row_beats_a_horizontal_band() -> None:
    table = _build(
        ("lastRow", "band1Horz", "band2Horz"),
        _look(lastRow=1, noHBand=0),
        rows=6,
        cols=3,
    )
    assert [row[0] for row in _grid(table)] == [13.0, 14.0, 13.0, 14.0, 13.0, 16.0]


def test_row_branches_beat_column_branches() -> None:
    """No corner branch defined, so the row/col intersections are a contest."""
    table = _build(
        ("firstRow", "lastRow", "firstCol", "lastCol"),
        _look(firstRow=1, lastRow=1, firstColumn=1, lastColumn=1),
        rows=4,
        cols=4,
    )
    assert _grid(table) == [
        [15.0, 15.0, 15.0, 15.0],
        [17.0, 10.0, 10.0, 18.0],
        [17.0, 10.0, 10.0, 18.0],
        [16.0, 16.0, 16.0, 16.0],
    ]


# --------------------------------------------------------------------------
# wholeTable. Word discards the branch on load and does not keep it on save.
# --------------------------------------------------------------------------


def _whole_table_style(doc: Document, *, base_size: float | None) -> None:
    """A table style whose wholeTable branch disagrees with its base rPr."""
    style = sub(doc.styles.element, "w:style", **{"w:type": "table", "w:styleId": "W"})
    sub(style, "w:name", **{"w:val": "W"})
    if base_size is not None:
        rpr = sub(style, "w:rPr")
        sub(rpr, "w:sz", **{"w:val": str(int(base_size * 2))})
    branch = sub(style, "w:tblStylePr", **{"w:type": "wholeTable"})
    branch_rpr = sub(branch, "w:rPr")
    sub(branch_rpr, "w:sz", **{"w:val": "18"})  # 9pt, never rendered


def _whole_table_cell(doc: Document):
    table = doc.add_table(rows=2, cols=2)
    sub(table._tbl.find(qn("w:tblPr")), "w:tblStyle", **{"w:val": "W"})
    return table.rows[0].cells[0]


def test_whole_table_branch_loses_to_the_styles_own_base() -> None:
    doc = Document()
    _whole_table_style(doc, base_size=12.0)
    assert resolve_effective_formatting(_whole_table_cell(doc)).font_size == 12.0


def test_whole_table_branch_contributes_nothing_on_its_own() -> None:
    """Inert even as the only branch: the size falls through to docDefaults."""
    doc = Document()
    _whole_table_style(doc, base_size=None)
    assert resolve_effective_formatting(_whole_table_cell(doc)).font_size != 9.0
