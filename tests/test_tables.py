"""Tests for the `tables/` capability — borders, shading, merging, reading."""

from __future__ import annotations

import dataclasses

import pytest
from docx import Document
from docx.exceptions import InvalidSpanError
from docx.shared import Inches
from lxml import etree

from docx_plus.core import Border, DocxPlusError
from docx_plus.core.ns import qn
from docx_plus.tables import (
    CellFormatting,
    InvalidMergeError,
    Shading,
    TableFormatting,
    merge_cells,
    normalize_horizontal_merges,
    read_table_formatting,
    set_cell_borders,
    set_cell_shading,
    set_row_shading,
    set_table_borders,
    set_table_shading,
    shading_attrs,
    unmerge_cell,
)


@pytest.fixture
def table():
    """A 3x3 table in a fresh document."""
    return Document().add_table(rows=3, cols=3)


def child_names(element):
    """Local names of an element's children, in document order."""
    return [etree.QName(child).localname for child in element]


def borders_of(parent, container_tag):
    """The border container under `parent`, or None."""
    return parent.find(qn(container_tag))


# ---------------------------------------------------------------------------
# Borders.
# ---------------------------------------------------------------------------


class TestTableBorders:
    def test_all_edges_writes_six_edges(self, table):
        set_table_borders(table, all_edges=Border(style="single", size=8, color="2F5496"))
        container = borders_of(table._tbl.tblPr, "w:tblBorders")
        assert child_names(container) == [
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV",
        ]

    def test_explicit_edge_overrides_all_edges(self, table):
        set_table_borders(
            table,
            all_edges=Border(style="single", size=4),
            top=Border(style="double", size=12),
        )
        read = read_table_formatting(table).borders
        assert read["top"] == Border(style="double", size=12, space=0)
        assert read["bottom"] == Border(style="single", size=4, space=0)

    def test_individual_edges_only(self, table):
        set_table_borders(table, top=Border(), inside_v=Border(style="dotted"))
        container = borders_of(table._tbl.tblPr, "w:tblBorders")
        assert child_names(container) == ["top", "insideV"]

    def test_no_edges_removes_container(self, table):
        set_table_borders(table, all_edges=Border())
        set_table_borders(table)
        assert borders_of(table._tbl.tblPr, "w:tblBorders") is None

    def test_no_edges_on_bare_table_writes_nothing(self, table):
        set_table_borders(table)
        assert borders_of(table._tbl.tblPr, "w:tblBorders") is None

    def test_idempotent_replace(self, table):
        set_table_borders(table, all_edges=Border(style="single"))
        set_table_borders(table, all_edges=Border(style="double"))
        assert len(table._tbl.tblPr.findall(qn("w:tblBorders"))) == 1
        assert read_table_formatting(table).borders["top"].style == "double"

    def test_space_is_pinned_to_zero(self, table):
        """`Border.space` defaults to a *page* value; tables must write 0."""
        assert Border().space == 24
        set_table_borders(table, all_edges=Border())
        edge = borders_of(table._tbl.tblPr, "w:tblBorders").find(qn("w:top"))
        assert edge.get(qn("w:space")) == "0"

    def test_explicit_space_is_still_pinned(self, table):
        set_table_borders(table, top=Border(space=12))
        edge = borders_of(table._tbl.tblPr, "w:tblBorders").find(qn("w:top"))
        assert edge.get(qn("w:space")) == "0"

    def test_lands_in_schema_position(self, table):
        table.style = "Table Grid"
        set_table_shading(table, Shading(fill="EEEEEE"))
        set_table_borders(table, all_edges=Border())
        assert child_names(table._tbl.tblPr) == [
            "tblStyle",
            "tblW",
            "tblBorders",
            "shd",
            "tblLook",
        ]

    def test_attributes_written(self, table):
        set_table_borders(table, top=Border(style="dashed", size=16, color="FF0000"))
        edge = borders_of(table._tbl.tblPr, "w:tblBorders").find(qn("w:top"))
        assert edge.get(qn("w:val")) == "dashed"
        assert edge.get(qn("w:sz")) == "16"
        assert edge.get(qn("w:color")) == "FF0000"


class TestCellBorders:
    def test_sides_and_diagonals(self, table):
        set_cell_borders(
            table.cell(0, 0),
            all_edges=Border(),
            tl2br=Border(style="dotted"),
            tr2bl=Border(style="dashed"),
        )
        container = borders_of(table.cell(0, 0)._tc.tcPr, "w:tcBorders")
        assert child_names(container) == ["top", "left", "bottom", "right", "tl2br", "tr2bl"]

    def test_all_edges_excludes_diagonals(self, table):
        set_cell_borders(table.cell(0, 0), all_edges=Border())
        container = borders_of(table.cell(0, 0)._tc.tcPr, "w:tcBorders")
        assert child_names(container) == ["top", "left", "bottom", "right"]

    def test_diagonal_only(self, table):
        set_cell_borders(table.cell(1, 1), tl2br=Border(style="single"))
        assert read_table_formatting(table).cells[4].borders == {
            "tl2br": Border(style="single", size=4, color="auto", space=0)
        }

    def test_removal(self, table):
        cell = table.cell(0, 0)
        set_cell_borders(cell, all_edges=Border())
        set_cell_borders(cell)
        assert borders_of(cell._tc.tcPr, "w:tcBorders") is None

    def test_lands_before_shading(self, table):
        cell = table.cell(0, 0)
        set_cell_shading(cell, Shading(fill="EEEEEE"))
        set_cell_borders(cell, all_edges=Border())
        assert child_names(cell._tc.tcPr) == ["tcW", "tcBorders", "shd"]

    def test_space_pinned_to_zero(self, table):
        set_cell_borders(table.cell(0, 0), top=Border(space=31))
        edge = borders_of(table.cell(0, 0)._tc.tcPr, "w:tcBorders").find(qn("w:top"))
        assert edge.get(qn("w:space")) == "0"


# ---------------------------------------------------------------------------
# Shading.
# ---------------------------------------------------------------------------


class TestShadingDataclass:
    def test_defaults(self):
        assert Shading() == Shading(fill="auto", pattern="clear", color="auto")

    def test_attrs_order_and_values(self):
        assert shading_attrs(Shading(fill="D9E2F3", pattern="pct25", color="808080")) == {
            "w:val": "pct25",
            "w:color": "808080",
            "w:fill": "D9E2F3",
        }

    @pytest.mark.parametrize("bad", ["red", "#FF0000", "FFF", "0x00FF00", ""])
    def test_bad_fill_rejected(self, bad):
        with pytest.raises(ValueError, match="Shading.fill"):
            Shading(fill=bad)

    @pytest.mark.parametrize("bad", ["red", "#FF0000", "FFF"])
    def test_bad_color_rejected(self, bad):
        with pytest.raises(ValueError, match="Shading.color"):
            Shading(color=bad)

    @pytest.mark.parametrize("bad", ["25%", "", "pct 25", "-clear"])
    def test_bad_pattern_rejected(self, bad):
        with pytest.raises(ValueError, match="Shading.pattern"):
            Shading(pattern=bad)

    @pytest.mark.parametrize("good", ["clear", "nil", "solid", "pct25", "thinHorzStripe"])
    def test_valid_patterns_accepted(self, good):
        assert Shading(pattern=good).pattern == good

    def test_frozen(self):
        with pytest.raises(dataclasses.FrozenInstanceError):
            Shading().fill = "FFFFFF"


class TestShadingWriters:
    def test_table_shading(self, table):
        set_table_shading(table, Shading(fill="F2F2F2"))
        shd = table._tbl.tblPr.find(qn("w:shd"))
        assert shd.get(qn("w:fill")) == "F2F2F2"
        assert shd.get(qn("w:val")) == "clear"

    def test_table_shading_removal(self, table):
        set_table_shading(table, Shading(fill="F2F2F2"))
        set_table_shading(table, None)
        assert table._tbl.tblPr.find(qn("w:shd")) is None

    def test_table_shading_removal_when_absent(self, table):
        set_table_shading(table, None)
        assert table._tbl.tblPr.find(qn("w:shd")) is None

    def test_table_shading_idempotent(self, table):
        set_table_shading(table, Shading(fill="F2F2F2"))
        set_table_shading(table, Shading(fill="2F5496"))
        assert len(table._tbl.tblPr.findall(qn("w:shd"))) == 1
        assert read_table_formatting(table).shading == Shading(fill="2F5496")

    def test_cell_shading(self, table):
        set_cell_shading(table.cell(1, 2), Shading(fill="2F5496"))
        assert read_table_formatting(table).cells[5].shading == Shading(fill="2F5496")

    def test_cell_shading_removal(self, table):
        cell = table.cell(0, 0)
        set_cell_shading(cell, Shading(fill="2F5496"))
        set_cell_shading(cell, None)
        assert cell._tc.tcPr.find(qn("w:shd")) is None

    def test_row_shading_covers_every_cell(self, table):
        set_row_shading(table.rows[0], Shading(fill="2F5496"))
        row_zero = [c for c in read_table_formatting(table).cells if c.row == 0]
        assert all(c.shading == Shading(fill="2F5496") for c in row_zero)
        assert all(c.shading is None for c in read_table_formatting(table).cells if c.row == 1)

    def test_row_shading_removal(self, table):
        set_row_shading(table.rows[0], Shading(fill="2F5496"))
        set_row_shading(table.rows[0], None)
        assert all(c.shading is None for c in read_table_formatting(table).cells)

    def test_row_shading_visits_a_merged_cell_once(self, table):
        merge_cells(table.cell(0, 0), table.cell(0, 2))
        set_row_shading(table.rows[0], Shading(fill="2F5496"))
        assert len(table.rows[0]._tr.tc_lst) == 1
        row_zero = [c for c in read_table_formatting(table).cells if c.row == 0]
        assert len(row_zero) == 1
        assert row_zero[0].shading == Shading(fill="2F5496")


# ---------------------------------------------------------------------------
# Merging.
# ---------------------------------------------------------------------------


class TestMergeCells:
    def test_horizontal(self, table):
        merged = merge_cells(table.cell(0, 0), table.cell(0, 2))
        assert merged.grid_span == 3
        assert len(table.rows[0]._tr.tc_lst) == 1

    def test_vertical(self, table):
        merge_cells(table.cell(0, 0), table.cell(2, 0))
        vals = [c.vertical_merge for c in read_table_formatting(table).cells if c.column == 0]
        assert vals == ["restart", "continue", "continue"]

    def test_block(self, table):
        merged = merge_cells(table.cell(0, 0), table.cell(1, 1))
        assert merged.grid_span == 2
        assert merged._tc.vMerge == "restart"

    def test_content_moves_to_anchor(self, table):
        table.cell(0, 1).text = "second"
        merged = merge_cells(table.cell(0, 0), table.cell(0, 2))
        assert "second" in merged.text

    def test_merging_a_cell_with_itself_is_a_noop(self, table):
        merged = merge_cells(table.cell(1, 1), table.cell(1, 1))
        assert merged.grid_span == 1

    def test_non_rectangular_raises_typed_error(self, table):
        # Widening (0, 0) across two columns makes a selection down to
        # (1, 0) an inverted L: same left edge, different right edge.
        merge_cells(table.cell(0, 0), table.cell(0, 1))
        with pytest.raises(InvalidMergeError, match="not rectangular"):
            merge_cells(table.cell(0, 0), table.cell(1, 0))

    def test_error_is_a_docx_plus_error_and_value_error(self):
        assert issubclass(InvalidMergeError, DocxPlusError)
        assert issubclass(InvalidMergeError, ValueError)

    def test_error_chains_the_original(self, table):
        merge_cells(table.cell(0, 0), table.cell(0, 1))
        with pytest.raises(InvalidMergeError) as caught:
            merge_cells(table.cell(0, 0), table.cell(1, 0))
        assert isinstance(caught.value.__cause__, InvalidSpanError)


class TestUnmergeCell:
    def test_horizontal_restores_columns(self, table):
        merge_cells(table.cell(0, 0), table.cell(0, 2))
        unmerge_cell(table.cell(0, 0))
        assert len(table.rows[0]._tr.tc_lst) == 3
        assert [c.grid_span for c in read_table_formatting(table).cells if c.row == 0] == [1, 1, 1]

    def test_vertical_restores_rows(self, table):
        merge_cells(table.cell(0, 0), table.cell(2, 0))
        unmerge_cell(table.cell(0, 0))
        vals = [c.vertical_merge for c in read_table_formatting(table).cells if c.column == 0]
        assert vals == [None, None, None]

    def test_block_restores_grid(self, table):
        merge_cells(table.cell(0, 0), table.cell(1, 1))
        unmerge_cell(table.cell(0, 0))
        cells = read_table_formatting(table).cells
        assert len(cells) == 9
        assert all(c.grid_span == 1 and c.vertical_merge is None for c in cells)

    def test_works_from_a_continuation_cell(self, table):
        merge_cells(table.cell(0, 0), table.cell(2, 0))
        # `Table.cell(2, 0)` resolves to the continuation `w:tc`.
        unmerge_cell(table.cell(2, 0))
        vals = [c.vertical_merge for c in read_table_formatting(table).cells if c.column == 0]
        assert vals == [None, None, None]

    def test_content_stays_in_the_anchor(self, table):
        merged = merge_cells(table.cell(0, 0), table.cell(0, 2))
        merged.text = "banner"
        unmerge_cell(table.cell(0, 0))
        assert [c.text for c in table.rows[0].cells] == ["banner", "", ""]

    def test_idempotent_on_an_unmerged_cell(self, table):
        unmerge_cell(table.cell(1, 1))
        assert len(read_table_formatting(table).cells) == 9

    def test_repeated_unmerge_is_stable(self, table):
        merge_cells(table.cell(0, 0), table.cell(1, 1))
        unmerge_cell(table.cell(0, 0))
        unmerge_cell(table.cell(0, 0))
        assert len(read_table_formatting(table).cells) == 9

    def test_width_is_divided_evenly(self):
        table = Document().add_table(rows=1, cols=3)
        for cell in table.rows[0].cells:
            cell.width = Inches(1)
        merge_cells(table.cell(0, 0), table.cell(0, 2))
        assert table.cell(0, 0).width == Inches(3)
        unmerge_cell(table.cell(0, 0))
        assert [c.width for c in table.rows[0].cells] == [Inches(1)] * 3

    def test_survives_when_widths_are_absent(self, table):
        merge_cells(table.cell(0, 0), table.cell(0, 2))
        for tc in table.rows[0]._tr.tc_lst:
            tc_pr = tc.get_or_add_tcPr()
            for width in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(width)
        unmerge_cell(table.cell(0, 0))
        assert len(table.rows[0]._tr.tc_lst) == 3

    def test_stops_at_a_row_with_no_cell_at_that_offset(self, table):
        """A ragged row ends the vertical run instead of raising."""
        merge_cells(table.cell(1, 0), table.cell(1, 1))
        # Row 1 now starts cells at offsets 0 and 2; nothing begins at 1.
        table.cell(0, 1)._tc.vMerge = "restart"
        unmerge_cell(table.cell(0, 1))
        assert table.cell(0, 1)._tc.vMerge is None


class TestNormalizeHorizontalMerges:
    @staticmethod
    def mark_h_merge(cell, value):
        """Write a `w:hMerge` onto `cell`, omitting `w:val` when None."""
        tc_pr = cell._tc.get_or_add_tcPr()
        h_merge = etree.SubElement(tc_pr, qn("w:hMerge"))
        if value is not None:
            h_merge.set(qn("w:val"), value)
        tc_pr.insert(0, h_merge)

    def test_grid_span_table_is_untouched(self, table):
        merge_cells(table.cell(0, 0), table.cell(0, 2))
        assert normalize_horizontal_merges(table) == 0
        assert len(table.rows[0]._tr.tc_lst) == 1

    def test_plain_table_reports_zero(self, table):
        assert normalize_horizontal_merges(table) == 0

    def test_converts_to_grid_span(self, table):
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), "continue")
        assert normalize_horizontal_merges(table) == 1
        row_zero = [c for c in read_table_formatting(table).cells if c.row == 0]
        assert [(c.column, c.grid_span) for c in row_zero] == [(0, 2), (2, 1)]
        assert all(c.horizontal_merge is None for c in row_zero)

    def test_omitted_val_means_continue(self, table):
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), None)
        assert normalize_horizontal_merges(table) == 1
        assert len(table.rows[0]._tr.tc_lst) == 2

    def test_refuses_to_discard_content(self, table):
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), "continue")
        table.cell(0, 1).text = "hidden"
        with pytest.raises(InvalidMergeError, match="discard_content=True"):
            normalize_horizontal_merges(table)
        assert len(table.rows[0]._tr.tc_lst) == 3

    def test_discards_content_when_asked(self, table):
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), "continue")
        table.cell(0, 1).text = "hidden"
        assert normalize_horizontal_merges(table, discard_content=True) == 1
        assert "hidden" not in table.rows[0].cells[0].text

    def test_anchor_content_is_kept(self, table):
        table.cell(0, 0).text = "kept"
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), "continue")
        normalize_horizontal_merges(table)
        assert table.rows[0].cells[0].text == "kept"

    def test_lone_restart_is_cleared(self, table):
        self.mark_h_merge(table.cell(0, 0), "restart")
        assert normalize_horizontal_merges(table) == 0
        assert read_table_formatting(table).cells[0].horizontal_merge is None

    def test_widths_are_summed(self):
        table = Document().add_table(rows=1, cols=3)
        for cell in table.rows[0].cells:
            cell.width = Inches(1)
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), "continue")
        normalize_horizontal_merges(table)
        assert table.rows[0]._tr.tc_lst[0].width == Inches(2)

    def test_two_regions_in_one_row(self):
        table = Document().add_table(rows=1, cols=4)
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), "continue")
        self.mark_h_merge(table.cell(0, 2), "restart")
        self.mark_h_merge(table.cell(0, 3), "continue")
        assert normalize_horizontal_merges(table) == 2
        assert [c.grid_span for c in read_table_formatting(table).cells] == [2, 2]

    @pytest.mark.parametrize("blank", ["", "   "])
    def test_an_empty_run_is_not_content(self, table, blank):
        """`cell.text = ""` leaves a run behind; it must not block the merge."""
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), "continue")
        table.cell(0, 1).text = blank
        assert normalize_horizontal_merges(table) == 1

    def test_a_drawing_counts_as_content(self, table):
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), "continue")
        run = table.cell(0, 1).paragraphs[0].add_run()
        etree.SubElement(run._r, qn("w:drawing"))
        with pytest.raises(InvalidMergeError):
            normalize_horizontal_merges(table)

    def test_cells_without_tc_pr_are_skipped(self, table):
        for tc in table.rows[0]._tr.tc_lst:
            tc.remove(tc.tcPr)
        assert normalize_horizontal_merges(table) == 0

    def test_a_nested_table_counts_as_content(self, table):
        self.mark_h_merge(table.cell(0, 0), "restart")
        self.mark_h_merge(table.cell(0, 1), "continue")
        table.cell(0, 1).add_table(rows=1, cols=1)
        with pytest.raises(InvalidMergeError):
            normalize_horizontal_merges(table)


# ---------------------------------------------------------------------------
# Reading.
# ---------------------------------------------------------------------------


class TestReadTableFormatting:
    def test_bare_table(self, table):
        read = read_table_formatting(table)
        assert isinstance(read, TableFormatting)
        assert read.style is None
        assert read.borders == {}
        assert read.shading is None
        assert len(read.cells) == 9
        assert all(isinstance(c, CellFormatting) for c in read.cells)

    def test_style_id(self, table):
        table.style = "Table Grid"
        assert read_table_formatting(table).style == "TableGrid"

    def test_round_trips_borders_with_space_zeroed(self, table):
        set_table_borders(table, all_edges=Border(style="dashed", size=12, color="112233"))
        read = read_table_formatting(table).borders
        assert read["insideH"] == Border(style="dashed", size=12, color="112233", space=0)

    def test_round_trips_shading(self, table):
        set_table_shading(table, Shading(fill="D9E2F3", pattern="pct25", color="808080"))
        assert read_table_formatting(table).shading == Shading(
            fill="D9E2F3", pattern="pct25", color="808080"
        )

    def test_cells_are_row_major(self, table):
        positions = [(c.row, c.column) for c in read_table_formatting(table).cells]
        assert positions == [(r, c) for r in range(3) for c in range(3)]

    def test_merged_cell_appears_once(self, table):
        merge_cells(table.cell(0, 0), table.cell(0, 2))
        read = read_table_formatting(table)
        assert len([c for c in read.cells if c.row == 0]) == 1
        assert read.cells[0].grid_span == 3

    def test_column_is_a_grid_offset(self, table):
        merge_cells(table.cell(1, 0), table.cell(1, 1))
        row_one = [c for c in read_table_formatting(table).cells if c.row == 1]
        assert [c.column for c in row_one] == [0, 2]

    def test_reports_h_merge_before_normalization(self, table):
        TestNormalizeHorizontalMerges.mark_h_merge(table.cell(0, 0), "restart")
        TestNormalizeHorizontalMerges.mark_h_merge(table.cell(0, 1), None)
        cells = read_table_formatting(table).cells
        assert cells[0].horizontal_merge == "restart"
        assert cells[1].horizontal_merge == "continue"
        assert cells[2].horizontal_merge is None

    def test_missing_attributes_fall_back(self, table):
        tc_pr = table.cell(0, 0)._tc.get_or_add_tcPr()
        container = etree.SubElement(tc_pr, qn("w:tcBorders"))
        etree.SubElement(container, qn("w:top"))
        assert read_table_formatting(table).cells[0].borders["top"] == Border(
            style="single", size=4, color="auto", space=0
        )

    def test_malformed_integer_attributes_fall_back(self, table):
        tc_pr = table.cell(0, 0)._tc.get_or_add_tcPr()
        container = etree.SubElement(tc_pr, qn("w:tcBorders"))
        edge = etree.SubElement(container, qn("w:top"))
        edge.set(qn("w:sz"), "thick")
        edge.set(qn("w:space"), "")
        assert read_table_formatting(table).cells[0].borders["top"].size == 4

    def test_empty_border_container_reads_as_no_borders(self, table):
        tc_pr = table.cell(0, 0)._tc.get_or_add_tcPr()
        etree.SubElement(tc_pr, qn("w:tcBorders"))
        assert read_table_formatting(table).cells[0].borders == {}

    def test_cell_without_tc_pr(self, table):
        assert table.cell(0, 0)._tc.tcPr is not None  # python-docx writes w:tcW
        for tc in table.rows[0]._tr.tc_lst:
            tc.remove(tc.tcPr)
        row_zero = [c for c in read_table_formatting(table).cells if c.row == 0]
        assert all(c.borders == {} and c.shading is None for c in row_zero)
        assert all(c.vertical_merge is None and c.horizontal_merge is None for c in row_zero)


# ---------------------------------------------------------------------------
# Round-trip through a saved file.
# ---------------------------------------------------------------------------


def test_survives_a_save_and_reload(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    set_table_borders(table, all_edges=Border(style="single", size=8, color="2F5496"))
    set_row_shading(table.rows[0], Shading(fill="2F5496"))
    set_cell_borders(table.cell(0, 1), tl2br=Border(style="dotted"))
    merge_cells(table.cell(1, 0), table.cell(1, 2))

    path = tmp_path / "tables.docx"
    doc.save(path)

    reloaded = read_table_formatting(Document(path).tables[0])
    assert reloaded.borders["insideV"] == Border(style="single", size=8, color="2F5496", space=0)
    assert reloaded.cells[0].shading == Shading(fill="2F5496")
    assert reloaded.cells[1].borders == {"tl2br": Border(style="dotted", space=0)}
    # Row 1 collapsed to a single `w:tc` under the merge.
    assert len(reloaded.cells) == 4
    assert reloaded.cells[3].grid_span == 3
