"""Toggle property semantics — the highest-risk part of the cascade.

ECMA-376 17.7.3 gives certain run properties (bold, italic, caps, smallCaps,
strike, vanish, ...) a combination rule that is not plain override. The rule
was settled by measuring live Word, because the spec prose is ambiguous
enough that three plausible readings all fit a casual sample:

* Each *style level* — table style, paragraph style, character style — is
  flattened over its own ``w:basedOn`` chain by **override**. Inheritance is
  not a hierarchy boundary, so a child re-asserting its parent's ``<w:b/>``
  stays bold.
* The result is the ``docDefaults`` value flipped once per style level whose
  value **differs** from it. A level restating the default is inert.
* Direct formatting is **absolute** and never participates.

The consequences look pedantic and are exactly the bugs that survive
otherwise: a run inside a stock bold ``Heading 1`` must resolve bold, and a
paragraph style saying ``<w:b w:val="0"/>`` when nothing is bold must not
somehow become meaningful.
"""

from __future__ import annotations

from typing import Any

import pytest
from docx import Document

from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub
from docx_plus.styles.inspect import resolve_effective_formatting


def _add_paragraph_style(
    doc: Document,
    style_id: str,
    *,
    based_on: str | None = None,
    rpr_children: list[tuple[str, dict[str, str] | None]] | None = None,
) -> None:
    """Append a paragraph style to the doc's styles part.

    Args:
        doc: Document to mutate.
        style_id: ``w:styleId`` for the new style.
        based_on: Optional basedOn target.
        rpr_children: List of ``(tag, attrs|None)`` to add inside ``w:rPr``.
    """
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": style_id})
    sub(s, "w:name", **{"w:val": style_id})
    if based_on is not None:
        sub(s, "w:basedOn", **{"w:val": based_on})
    if rpr_children:
        rpr = sub(s, "w:rPr")
        for tag, attrs in rpr_children:
            kwargs: dict[str, Any] = attrs or {}
            sub(rpr, tag, **kwargs)


def _add_character_style(
    doc: Document,
    style_id: str,
    *,
    based_on: str | None = None,
    rpr_children: list[tuple[str, dict[str, str] | None]] | None = None,
) -> None:
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "character", "w:styleId": style_id})
    sub(s, "w:name", **{"w:val": style_id})
    if based_on is not None:
        sub(s, "w:basedOn", **{"w:val": based_on})
    if rpr_children:
        rpr = sub(s, "w:rPr")
        for tag, attrs in rpr_children:
            sub(rpr, tag, **(attrs or {}))


def _styled_paragraph(doc: Document, style_id: str) -> Any:
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": style_id})
    return p


def _styled_run(paragraph: Any, style_id: str, text: str = "text") -> Any:
    run = paragraph.add_run(text)
    sub(run._r.get_or_add_rPr(), "w:rStyle", **{"w:val": style_id})
    return run


def _set_doc_default(doc: Document, tag: str, attrs: dict[str, str] | None = None) -> None:
    """Put a toggle into ``w:docDefaults`` — the base the rule counts from."""
    defaults = doc.styles.element.find(qn("w:docDefaults"))
    rpr = defaults.find(qn("w:rPrDefault")).find(qn("w:rPr"))
    sub(rpr, tag, **(attrs or {}))


# --------------------------------------------------------------------------
# One style level: a basedOn chain flattens by override, not by parity.
# --------------------------------------------------------------------------


def test_single_style_bold_is_bold() -> None:
    doc = Document()
    _add_paragraph_style(doc, "BoldOnly", rpr_children=[("w:b", None)])
    p = _styled_paragraph(doc, "BoldOnly")

    resolved = resolve_effective_formatting(p)
    assert resolved.bold is True


def test_bold_survives_a_basedon_chain_that_reasserts_it() -> None:
    """The case that used to XOR to False. Word renders it bold.

    A child style restating its parent's ``<w:b/>`` is the single most
    common shape in a real styles.xml, and cancelling it made every
    derived heading style resolve unbold.
    """
    doc = Document()
    _add_paragraph_style(doc, "ABold", rpr_children=[("w:b", None)])
    _add_paragraph_style(doc, "BBoldChild", based_on="ABold", rpr_children=[("w:b", None)])
    p = _styled_paragraph(doc, "BBoldChild")

    assert resolve_effective_formatting(p).bold is True


def test_three_level_chain_stays_bold() -> None:
    doc = Document()
    _add_paragraph_style(doc, "L1", rpr_children=[("w:b", None)])
    _add_paragraph_style(doc, "L2", based_on="L1", rpr_children=[("w:b", None)])
    _add_paragraph_style(doc, "L3", based_on="L2", rpr_children=[("w:b", None)])
    p = _styled_paragraph(doc, "L3")

    assert resolve_effective_formatting(p).bold is True


def test_explicit_false_in_a_chain_overrides_the_parent() -> None:
    """Within one level the last writer wins, so the child's off sticks."""
    doc = Document()
    _add_paragraph_style(doc, "ABold2", rpr_children=[("w:b", None)])
    _add_paragraph_style(
        doc,
        "BUnboldChild",
        based_on="ABold2",
        rpr_children=[("w:b", {"w:val": "false"})],
    )
    p = _styled_paragraph(doc, "BUnboldChild")

    assert resolve_effective_formatting(p).bold is False


def test_value_zero_treated_as_explicit_false() -> None:
    """`w:val="0"` is the legacy form of `w:val="false"` per the schema."""
    doc = Document()
    _add_paragraph_style(doc, "BoldZ", rpr_children=[("w:b", None)])
    _add_paragraph_style(
        doc,
        "UnboldZ",
        based_on="BoldZ",
        rpr_children=[("w:b", {"w:val": "0"})],
    )
    p = _styled_paragraph(doc, "UnboldZ")

    assert resolve_effective_formatting(p).bold is False


def test_value_true_is_equivalent_to_no_val() -> None:
    doc = Document()
    _add_paragraph_style(doc, "BoldTrue", rpr_children=[("w:b", {"w:val": "true"})])
    p = _styled_paragraph(doc, "BoldTrue")

    assert resolve_effective_formatting(p).bold is True


def test_toggles_are_independent_of_each_other() -> None:
    doc = Document()
    _add_paragraph_style(doc, "BoldOnce", rpr_children=[("w:b", None), ("w:i", None)])
    _add_paragraph_style(doc, "BoldTwice", based_on="BoldOnce", rpr_children=[("w:b", None)])
    p = _styled_paragraph(doc, "BoldTwice")

    resolved = resolve_effective_formatting(p)
    assert resolved.bold is True
    assert resolved.italic is True


# --------------------------------------------------------------------------
# Between style levels: this is where the parity rule actually lives.
# --------------------------------------------------------------------------


def test_paragraph_style_and_character_style_cancel() -> None:
    """The spec's own example, and the only boundary that is widely known."""
    doc = Document()
    _add_paragraph_style(doc, "PBold", rpr_children=[("w:b", None)])
    _add_character_style(doc, "CBold", rpr_children=[("w:b", None)])
    run = _styled_run(_styled_paragraph(doc, "PBold"), "CBold")

    assert resolve_effective_formatting(run).bold is False


def test_a_character_styles_own_chain_still_flattens_first() -> None:
    """Flatten each level, then compare levels — not one XOR over every style."""
    doc = Document()
    _add_paragraph_style(doc, "PBold2", rpr_children=[("w:b", None)])
    _add_character_style(doc, "CBase", rpr_children=[("w:b", None)])
    _add_character_style(doc, "CDerived", based_on="CBase", rpr_children=[("w:b", None)])
    run = _styled_run(_styled_paragraph(doc, "PBold2"), "CDerived")

    # The character level flattens to a single "on", which cancels the
    # paragraph level's. Were the chain XOR'd first it would flatten to
    # "off" and the run would come out bold.
    assert resolve_effective_formatting(run).bold is False


def test_an_off_level_is_inert_when_the_base_is_off() -> None:
    """``<w:b w:val="0"/>`` matching the default contributes nothing.

    So a character style turning bold off does *not* defeat a paragraph
    style turning it on — the paragraph style is the only level that
    differs from the default, and one difference means on.
    """
    doc = Document()
    _add_paragraph_style(doc, "PBold3", rpr_children=[("w:b", None)])
    _add_character_style(doc, "CBoldOff", rpr_children=[("w:b", {"w:val": "0"})])
    run = _styled_run(_styled_paragraph(doc, "PBold3"), "CBoldOff")

    assert resolve_effective_formatting(run).bold is True


def test_an_off_level_alone_resolves_off() -> None:
    doc = Document()
    _add_paragraph_style(doc, "PPlain")
    _add_character_style(doc, "CBoldOff2", rpr_children=[("w:b", {"w:val": "0"})])
    run = _styled_run(_styled_paragraph(doc, "PPlain"), "CBoldOff2")

    assert resolve_effective_formatting(run).bold is False


# --------------------------------------------------------------------------
# docDefaults is the base the levels are counted against, not a level.
# --------------------------------------------------------------------------


def test_doc_defaults_alone_supplies_the_value() -> None:
    doc = Document()
    _set_doc_default(doc, "w:i")

    assert resolve_effective_formatting(doc.add_paragraph("text")).italic is True


def test_a_style_restating_the_document_default_is_inert() -> None:
    """Were docDefaults a level, this would cancel to not-italic."""
    doc = Document()
    _set_doc_default(doc, "w:i")
    _add_paragraph_style(doc, "PItalic", rpr_children=[("w:i", None)])

    assert resolve_effective_formatting(_styled_paragraph(doc, "PItalic")).italic is True


def test_two_levels_restating_the_document_default_are_both_inert() -> None:
    doc = Document()
    _set_doc_default(doc, "w:i")
    _add_paragraph_style(doc, "PItalic2", rpr_children=[("w:i", None)])
    _add_character_style(doc, "CItalic", rpr_children=[("w:i", None)])
    run = _styled_run(_styled_paragraph(doc, "PItalic2"), "CItalic")

    assert resolve_effective_formatting(run).italic is True


def test_a_style_contradicting_the_document_default_turns_it_off() -> None:
    doc = Document()
    _set_doc_default(doc, "w:i")
    _add_paragraph_style(doc, "PItalicOff", rpr_children=[("w:i", {"w:val": "0"})])

    assert resolve_effective_formatting(_styled_paragraph(doc, "PItalicOff")).italic is False


# --------------------------------------------------------------------------
# Direct formatting states a value outright.
# --------------------------------------------------------------------------


def test_direct_bold_on_unbold_style() -> None:
    doc = Document()
    _add_paragraph_style(doc, "PlainStyle")  # no w:b
    p = _styled_paragraph(doc, "PlainStyle")
    r = p.add_run("text")
    sub(r._r.get_or_add_rPr(), "w:b")

    assert resolve_effective_formatting(r).bold is True


def test_direct_unbold_on_bold_style() -> None:
    doc = Document()
    _add_paragraph_style(doc, "BoldStyleA", rpr_children=[("w:b", None)])
    p = _styled_paragraph(doc, "BoldStyleA")
    r = p.add_run("text")
    sub(r._r.get_or_add_rPr(), "w:b", **{"w:val": "false"})

    assert resolve_effective_formatting(r).bold is False


def test_direct_bold_over_a_bold_style_stays_bold() -> None:
    """Direct formatting is absolute: it states, it does not flip.

    Word's UI writes ``w:val="0"`` when a user un-bolds bold text, so a
    bare ``<w:b/>`` sitting over a bold style means bold — and treating it
    as a flip made every such run resolve unbold.
    """
    doc = Document()
    _add_paragraph_style(doc, "BoldStyleB", rpr_children=[("w:b", None)])
    p = _styled_paragraph(doc, "BoldStyleB")
    r = p.add_run("text")
    sub(r._r.get_or_add_rPr(), "w:b")

    assert resolve_effective_formatting(r).bold is True


def test_direct_bold_beats_a_cancelled_pair_of_levels() -> None:
    doc = Document()
    _add_paragraph_style(doc, "PBold4", rpr_children=[("w:b", None)])
    _add_character_style(doc, "CBold4", rpr_children=[("w:b", None)])
    run = _styled_run(_styled_paragraph(doc, "PBold4"), "CBold4")
    sub(run._r.get_or_add_rPr(), "w:b")

    assert resolve_effective_formatting(run).bold is True


def test_a_run_in_a_stock_heading_resolves_bold() -> None:
    """The end-to-end shape all of this exists to get right."""
    doc = Document()
    para = doc.add_paragraph("Heading text", style="Heading 1")

    assert resolve_effective_formatting(para).bold is True
    assert resolve_effective_formatting(para.runs[0]).bold is True


# --------------------------------------------------------------------------
# Every toggle behaves alike: the six complex-script / decorative variants
# (bCs, iCs, emboss, imprint, outline, shadow) plus the six base ones.
# --------------------------------------------------------------------------


_NEW_TOGGLES = [
    ("w:bCs", "cs_bold"),
    ("w:iCs", "cs_italic"),
    ("w:emboss", "emboss"),
    ("w:imprint", "imprint"),
    ("w:outline", "outline"),
    ("w:shadow", "shadow"),
]


@pytest.mark.parametrize(("tag", "field_name"), _NEW_TOGGLES)
def test_new_toggle_direct_application(tag: str, field_name: str) -> None:
    """Single style declares the toggle -> field resolves True."""
    doc = Document()
    style_id = f"OnceOnly_{field_name}"
    _add_paragraph_style(doc, style_id, rpr_children=[(tag, None)])
    p = _styled_paragraph(doc, style_id)

    resolved = resolve_effective_formatting(p)
    assert getattr(resolved, field_name) is True


@pytest.mark.parametrize(("tag", "field_name"), _NEW_TOGGLES)
def test_new_toggle_overrides_through_a_chain(tag: str, field_name: str) -> None:
    """Two styles in one chain both declare it -> still True."""
    doc = Document()
    parent_id = f"Parent_{field_name}"
    child_id = f"Child_{field_name}"
    _add_paragraph_style(doc, parent_id, rpr_children=[(tag, None)])
    _add_paragraph_style(doc, child_id, based_on=parent_id, rpr_children=[(tag, None)])
    p = _styled_paragraph(doc, child_id)

    assert getattr(resolve_effective_formatting(p), field_name) is True


@pytest.mark.parametrize(("tag", "field_name"), _NEW_TOGGLES)
def test_new_toggle_cancels_across_levels(tag: str, field_name: str) -> None:
    """Paragraph level and character level both declare it -> False."""
    doc = Document()
    para_id = f"XorPara_{field_name}"
    char_id = f"XorChar_{field_name}"
    _add_paragraph_style(doc, para_id, rpr_children=[(tag, None)])
    _add_character_style(doc, char_id, rpr_children=[(tag, None)])
    run = _styled_run(_styled_paragraph(doc, para_id), char_id)

    assert getattr(resolve_effective_formatting(run), field_name) is False


@pytest.mark.parametrize(("tag", "field_name"), _NEW_TOGGLES)
def test_new_toggle_explicit_false_overrides_within_a_chain(tag: str, field_name: str) -> None:
    doc = Document()
    parent_id = f"ResetParent_{field_name}"
    child_id = f"ResetChild_{field_name}"
    _add_paragraph_style(doc, parent_id, rpr_children=[(tag, None)])
    _add_paragraph_style(
        doc,
        child_id,
        based_on=parent_id,
        rpr_children=[(tag, {"w:val": "false"})],
    )
    p = _styled_paragraph(doc, child_id)

    resolved = resolve_effective_formatting(p)
    assert getattr(resolved, field_name) is False


def test_an_untouched_toggle_stays_none() -> None:
    """Unset must not collapse to False — the linter distinguishes them."""
    doc = Document()

    assert resolve_effective_formatting(doc.add_paragraph("text")).caps is None


# --------------------------------------------------------------------------
# dstrike — non-toggle (last-writer-wins) per ECMA-376 17.7.3 + 17.3.2.10.
# Regression coverage for H2.
# --------------------------------------------------------------------------


def test_dstrike_direct_resolves_double_strike_true() -> None:
    doc = Document()
    _add_paragraph_style(doc, "Dstrike", rpr_children=[("w:dstrike", None)])
    p = _styled_paragraph(doc, "Dstrike")

    resolved = resolve_effective_formatting(p)
    assert resolved.double_strike is True


def test_dstrike_explicit_false_resolves_false() -> None:
    doc = Document()
    _add_paragraph_style(doc, "DstrikeOff", rpr_children=[("w:dstrike", {"w:val": "false"})])
    p = _styled_paragraph(doc, "DstrikeOff")

    resolved = resolve_effective_formatting(p)
    assert resolved.double_strike is False


def test_dstrike_child_overrides_parent_last_writer_wins() -> None:
    """Two layers both set dstrike — child wins (non-toggle, no parity)."""
    doc = Document()
    _add_paragraph_style(doc, "DstrikeParent", rpr_children=[("w:dstrike", None)])
    _add_paragraph_style(
        doc,
        "DstrikeChild",
        based_on="DstrikeParent",
        rpr_children=[("w:dstrike", {"w:val": "false"})],
    )
    p = _styled_paragraph(doc, "DstrikeChild")

    resolved = resolve_effective_formatting(p)
    assert resolved.double_strike is False


def test_dstrike_does_not_cancel_across_levels() -> None:
    """The parity rule is for toggles only; dstrike is not one."""
    doc = Document()
    _add_paragraph_style(doc, "DsPara", rpr_children=[("w:dstrike", None)])
    _add_character_style(doc, "DsChar", rpr_children=[("w:dstrike", None)])
    run = _styled_run(_styled_paragraph(doc, "DsPara"), "DsChar")

    assert resolve_effective_formatting(run).double_strike is True


def test_dstrike_and_strike_are_independent() -> None:
    """``strike`` and ``dstrike`` are separate properties — both can be set.

    Word's UI enforces mutual exclusivity but ECMA-376 does not; the
    resolver preserves whatever the cascade actually emits.
    """
    doc = Document()
    _add_paragraph_style(
        doc,
        "Both",
        rpr_children=[("w:strike", None), ("w:dstrike", None)],
    )
    p = _styled_paragraph(doc, "Both")

    resolved = resolve_effective_formatting(p)
    assert resolved.strike is True
    assert resolved.double_strike is True
