"""Tests for ``docx_plus.notes`` — footnotes + endnotes insert / read."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.notes import (
    EndnoteIdRegistry,
    EndnoteRef,
    FootnoteIdRegistry,
    FootnoteRef,
    NoteContent,
    NoteNotFoundError,
    add_endnote,
    add_footnote,
    edit_endnote,
    edit_footnote,
    read_endnotes,
    read_footnotes,
)


def _footnote_part(doc):
    return doc.part.part_related_by(RT.FOOTNOTES)


def _endnote_part(doc):
    return doc.part.part_related_by(RT.ENDNOTES)


# --------------------------------------------------------------------------
# add_footnote — body reference + part entry.
# --------------------------------------------------------------------------


def test_add_footnote_writes_reference_run() -> None:
    doc = Document()
    p = doc.add_paragraph("anchored")
    ref = add_footnote(p, "footnote text")
    assert isinstance(ref, FootnoteRef)
    refs = xpath(p._p, ".//w:footnoteReference")
    assert len(refs) == 1
    assert refs[0].get(qn("w:id")) == str(ref.note_id)


def test_add_footnote_writes_part_entry() -> None:
    doc = Document()
    p = doc.add_paragraph("anchored")
    ref = add_footnote(p, "the note")
    part = _footnote_part(doc)
    notes = xpath(part.element, "./w:footnote[@w:id=$i]", i=str(ref.note_id))
    assert len(notes) == 1
    texts = [t.text for t in xpath(notes[0], ".//w:t") if t.text]
    assert "".join(texts) == "the note"


def test_add_footnote_part_entry_has_para_style() -> None:
    doc = Document()
    p = doc.add_paragraph("a")
    ref = add_footnote(p, "x")
    part = _footnote_part(doc)
    style = xpath(
        part.element,
        "./w:footnote[@w:id=$i]/w:p/w:pPr/w:pStyle",
        i=str(ref.note_id),
    )
    assert len(style) == 1
    assert style[0].get(qn("w:val")) == "FootnoteText"


def test_add_footnote_reference_marker_style() -> None:
    """Body-side reference run uses rStyle val='FootnoteReference'."""
    doc = Document()
    p = doc.add_paragraph("a")
    add_footnote(p, "x")
    rstyle = xpath(
        p._p,
        ".//w:rPr/w:rStyle[@w:val='FootnoteReference']",
    )
    assert len(rstyle) >= 1


def test_add_footnote_returns_unique_ids() -> None:
    doc = Document()
    p = doc.add_paragraph("a")
    a = add_footnote(p, "one")
    b = add_footnote(p, "two")
    assert a.note_id != b.note_id


def test_add_footnote_first_call_creates_part() -> None:
    doc = Document()
    with pytest.raises(KeyError):
        _footnote_part(doc)
    p = doc.add_paragraph("a")
    add_footnote(p, "x")
    assert _footnote_part(doc) is not None


def test_add_footnote_reuses_part() -> None:
    doc = Document()
    p = doc.add_paragraph("a")
    add_footnote(p, "x")
    first_part = _footnote_part(doc)
    add_footnote(p, "y")
    assert _footnote_part(doc) is first_part
    # 2 seeded separator entries (-1, 0) + 2 user-added footnotes = 4.
    assert len(xpath(first_part.element, "./w:footnote")) == 4


# --------------------------------------------------------------------------
# add_endnote — same shape, endnote variants of tags.
# --------------------------------------------------------------------------


def test_add_endnote_writes_reference_run() -> None:
    doc = Document()
    p = doc.add_paragraph("anchored")
    ref = add_endnote(p, "endnote text")
    assert isinstance(ref, EndnoteRef)
    refs = xpath(p._p, ".//w:endnoteReference")
    assert len(refs) == 1
    assert refs[0].get(qn("w:id")) == str(ref.note_id)


def test_add_endnote_writes_part_entry() -> None:
    doc = Document()
    p = doc.add_paragraph("a")
    ref = add_endnote(p, "the note")
    part = _endnote_part(doc)
    notes = xpath(part.element, "./w:endnote[@w:id=$i]", i=str(ref.note_id))
    assert len(notes) == 1


def test_add_endnote_uses_endnote_styles() -> None:
    doc = Document()
    p = doc.add_paragraph("a")
    add_endnote(p, "x")
    p_style = xpath(
        _endnote_part(doc).element,
        ".//w:pPr/w:pStyle[@w:val='EndnoteText']",
    )
    r_style = xpath(p._p, ".//w:rPr/w:rStyle[@w:val='EndnoteReference']")
    assert len(p_style) == 1
    assert len(r_style) >= 1


def test_add_footnote_and_endnote_use_separate_id_spaces() -> None:
    """Footnote id 1 and endnote id 1 can coexist."""
    doc = Document()
    doc.add_paragraph("a")
    foot_reg = FootnoteIdRegistry(doc)
    end_reg = EndnoteIdRegistry(doc)
    # Reserve id 1 in each registry — they're disjoint, so this works.
    foot_reg.reserve(1)
    end_reg.reserve(1)
    # ...and a separate reserve in one doesn't pollute the other
    foot_reg.reserve(2)
    end_reg.reserve(2)


# --------------------------------------------------------------------------
# Reserved id 0 / -1 cannot be issued (range check fires first).
# --------------------------------------------------------------------------


def test_footnote_registry_rejects_reserved_id_zero() -> None:
    doc = Document()
    reg = FootnoteIdRegistry(doc)
    with pytest.raises(ValueError):
        reg.reserve(0)
    with pytest.raises(ValueError):
        reg.reserve(-1)


def test_endnote_registry_rejects_reserved_id_zero() -> None:
    doc = Document()
    reg = EndnoteIdRegistry(doc)
    with pytest.raises(ValueError):
        reg.reserve(0)
    with pytest.raises(ValueError):
        reg.reserve(-1)


# --------------------------------------------------------------------------
# read_footnotes / read_endnotes
# --------------------------------------------------------------------------


def test_read_footnotes_empty() -> None:
    assert read_footnotes(Document()) == []
    assert read_endnotes(Document()) == []


def test_read_footnotes_returns_user_notes() -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_footnote(p, "note one")
    notes = read_footnotes(doc)
    assert len(notes) == 1
    only = notes[0]
    assert isinstance(only, NoteContent)
    assert only.text == "note one"
    assert only.paragraph_index == 0


def test_read_footnotes_skips_reserved_separators() -> None:
    """If a document has separator entries (id <= 0), filter them out."""
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_footnote(p, "real")
    # Manually inject a separator entry to verify filtering.
    from docx_plus.core.oxml import el as make_el

    part = _footnote_part(doc)
    sep = make_el("w:footnote", **{"w:id": "-1", "w:type": "separator"})
    part.element.insert(0, sep)
    cont = make_el("w:footnote", **{"w:id": "0", "w:type": "continuationSeparator"})
    part.element.insert(1, cont)

    notes = read_footnotes(doc)
    assert len(notes) == 1
    assert notes[0].text == "real"


def test_read_endnotes_returns_user_notes() -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_endnote(p, "tail")
    notes = read_endnotes(doc)
    assert len(notes) == 1
    assert notes[0].text == "tail"


def test_read_footnotes_paragraph_index_tracks_position() -> None:
    doc = Document()
    doc.add_paragraph("first")
    doc.add_paragraph("second")
    p3 = doc.add_paragraph("third")
    add_footnote(p3, "tail")
    notes = read_footnotes(doc)
    assert notes[0].paragraph_index == 2


# --------------------------------------------------------------------------
# Round-trip
# --------------------------------------------------------------------------


def test_footnote_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_footnote(p, "saved note")
    out = tmp_path / "fn.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    notes = read_footnotes(reopened)
    assert len(notes) == 1
    assert notes[0].text == "saved note"


def test_endnote_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_endnote(p, "tail note")
    out = tmp_path / "en.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    notes = read_endnotes(reopened)
    assert len(notes) == 1
    assert notes[0].text == "tail note"


def test_mixed_footnotes_and_endnotes_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_footnote(p, "fn1")
    add_endnote(p, "en1")
    add_footnote(p, "fn2")
    out = tmp_path / "mix.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    fns = read_footnotes(reopened)
    ens = read_endnotes(reopened)
    assert sorted(n.text for n in fns) == ["fn1", "fn2"]
    assert [n.text for n in ens] == ["en1"]


# --------------------------------------------------------------------------
# edit_footnote / edit_endnote — in-place body replacement.
# --------------------------------------------------------------------------


def test_edit_footnote_replaces_text() -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    ref = add_footnote(p, "original")

    edit_footnote(doc, ref.note_id, "rewritten")

    notes = read_footnotes(doc)
    assert [n.text for n in notes] == ["rewritten"]


def test_edit_endnote_replaces_text() -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    ref = add_endnote(p, "v1")

    edit_endnote(doc, ref.note_id, "v2")

    notes = read_endnotes(doc)
    assert [n.text for n in notes] == ["v2"]


def test_edit_footnote_leaves_other_notes_intact() -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    keep_ref = add_footnote(p, "keep")
    edit_ref = add_footnote(p, "before")
    add_footnote(p, "trailing")

    edit_footnote(doc, edit_ref.note_id, "after")

    notes = {n.note_id: n.text for n in read_footnotes(doc)}
    assert notes[keep_ref.note_id] == "keep"
    assert notes[edit_ref.note_id] == "after"


def test_edit_footnote_preserves_reference_marker_in_body() -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    ref = add_footnote(p, "original")

    edit_footnote(doc, ref.note_id, "rewritten")

    # The body-side reference run is untouched (still a child of p._p).
    refs_in_body = xpath(
        doc.element.body, ".//w:footnoteReference[@w:id=$nid]", nid=str(ref.note_id)
    )
    assert len(refs_in_body) == 1


def test_edit_footnote_missing_id_raises() -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_footnote(p, "real")  # ensures the part exists

    with pytest.raises(NoteNotFoundError):
        edit_footnote(doc, 9999, "won't land")


def test_edit_footnote_with_no_part_raises() -> None:
    doc = Document()  # no footnotes part has been created
    with pytest.raises(NoteNotFoundError):
        edit_footnote(doc, 1, "nope")


def test_edit_endnote_missing_id_raises() -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_endnote(p, "real")

    with pytest.raises(NoteNotFoundError):
        edit_endnote(doc, 9999, "won't land")


@pytest.mark.parametrize("reserved_id", [-1, 0])
def test_edit_footnote_rejects_reserved_ids(reserved_id: int) -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_footnote(p, "real")

    with pytest.raises(ValueError):
        edit_footnote(doc, reserved_id, "no separator edits")


@pytest.mark.parametrize("reserved_id", [-1, 0])
def test_edit_endnote_rejects_reserved_ids(reserved_id: int) -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_endnote(p, "real")

    with pytest.raises(ValueError):
        edit_endnote(doc, reserved_id, "no separator edits")


def test_note_not_found_is_a_key_error() -> None:
    """``NoteNotFoundError`` subclasses ``KeyError`` per SPEC §16."""
    doc = Document()
    p = doc.add_paragraph("anchor")
    add_footnote(p, "real")

    with pytest.raises(KeyError):
        edit_footnote(doc, 9999, "no")


def test_edit_footnote_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("anchor")
    ref = add_footnote(p, "draft")
    edit_footnote(doc, ref.note_id, "final")

    out = tmp_path / "edited.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    assert [n.text for n in read_footnotes(reopened)] == ["final"]


def test_edit_footnote_strips_non_paragraph_children() -> None:
    """H6 regression: footnote bodies can contain tables, not just <w:p>."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    from docx_plus.core.oxml import sub

    doc = Document()
    p = doc.add_paragraph("anchor")
    ref = add_footnote(p, "draft")

    # Inject a <w:tbl> into the footnote body to simulate richer authoring.
    footnotes_part = doc.part.part_related_by(RT.FOOTNOTES)
    note_el = None
    for el_ in footnotes_part.element.findall(qn("w:footnote")):
        if el_.get(qn("w:id")) == str(ref.note_id):
            note_el = el_
            break
    assert note_el is not None
    sub(note_el, "w:tbl")
    assert note_el.find(qn("w:tbl")) is not None  # sanity

    edit_footnote(doc, ref.note_id, "rewritten")

    children = list(note_el)
    assert len(children) == 1
    assert children[0].tag == qn("w:p")
    assert note_el.find(qn("w:tbl")) is None
