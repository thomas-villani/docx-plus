"""Tests for ``fields.read_fields``.

A complex field is a *run sequence*, not an element, so the load-bearing
cases are the ones where that sequence is awkward: an instruction Word split
across several ``w:instrText`` elements, two fields in one paragraph, and a
nested field.
"""

from __future__ import annotations

from docx import Document

from docx_plus.core.oxml import build_complex_field, sub
from docx_plus.fields import read_fields


def _instr_run(p_element: object, text: str) -> None:
    """Append a bare instruction run, for building split instructions by hand."""
    run = sub(p_element, "w:r")  # type: ignore[arg-type]
    instr = sub(run, "w:instrText", **{"xml:space": "preserve"})
    instr.text = text


def _fld_char(p_element: object, char_type: str) -> None:
    run = sub(p_element, "w:r")  # type: ignore[arg-type]
    sub(run, "w:fldChar", **{"w:fldCharType": char_type})


def test_reads_a_simple_field() -> None:
    doc = Document()
    para = doc.add_paragraph()
    build_complex_field(para._p, " PAGE ", "1")

    fields = read_fields(doc)

    assert len(fields) == 1
    assert fields[0].keyword == "PAGE"
    assert fields[0].instruction == "PAGE"
    assert fields[0].result == "1"


def test_reads_arguments_and_switches() -> None:
    doc = Document()
    build_complex_field(doc.add_paragraph()._p, r" REF chapter1 \h ", "")

    found = read_fields(doc)[0]

    assert found.keyword == "REF"
    assert found.operands == ["chapter1"]
    assert found.switches == ["\\h"]


def test_strips_quotes_from_a_quoted_argument() -> None:
    doc = Document()
    build_complex_field(doc.add_paragraph()._p, ' SEQ "Figure Number" ', "")

    assert read_fields(doc)[0].operands == ["Figure Number"]


def test_joins_an_instruction_split_across_runs() -> None:
    """Word splits an instruction at arbitrary points, so joining is required."""
    doc = Document()
    para = doc.add_paragraph()
    _fld_char(para._p, "begin")
    _instr_run(para._p, " REF chap")
    _instr_run(para._p, "ter1 ")
    _fld_char(para._p, "separate")
    para.add_run("Chapter One")
    _fld_char(para._p, "end")

    found = read_fields(doc)[0]

    assert found.keyword == "REF"
    assert found.operands == ["chapter1"]
    assert found.result == "Chapter One"


def test_reads_two_fields_in_one_paragraph() -> None:
    doc = Document()
    para = doc.add_paragraph()
    build_complex_field(para._p, " PAGE ", "1")
    build_complex_field(para._p, " NUMPAGES ", "9")

    assert [f.keyword for f in read_fields(doc)] == ["PAGE", "NUMPAGES"]


def test_a_nested_field_reads_as_one() -> None:
    """Word nests fields inside TOC and IF instructions.

    The outer keyword is what a caller filtering by type is asking about,
    and the inner begin/end must not close the outer field early.
    """
    doc = Document()
    para = doc.add_paragraph()
    _fld_char(para._p, "begin")
    _instr_run(para._p, " IF ")
    _fld_char(para._p, "begin")
    _instr_run(para._p, " PAGE ")
    _fld_char(para._p, "end")
    _instr_run(para._p, ' = 1 "first" "later" ')
    _fld_char(para._p, "separate")
    para.add_run("first")
    _fld_char(para._p, "end")

    fields = read_fields(doc)

    assert len(fields) == 1
    assert fields[0].keyword == "IF"


def test_filters_by_keyword_case_insensitively() -> None:
    doc = Document()
    build_complex_field(doc.add_paragraph()._p, " PAGE ", "")
    build_complex_field(doc.add_paragraph()._p, " SEQ Figure ", "")

    assert [f.keyword for f in read_fields(doc, keyword="seq")] == ["SEQ"]


def test_reports_the_paragraph_index() -> None:
    doc = Document()
    doc.add_paragraph("first")
    build_complex_field(doc.add_paragraph()._p, " PAGE ", "")

    assert read_fields(doc)[0].paragraph_index == 1


def test_finds_a_field_inside_a_table_cell() -> None:
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    build_complex_field(cell.paragraphs[0]._p, " PAGE ", "")

    assert [f.keyword for f in read_fields(doc)] == ["PAGE"]


def test_a_field_with_no_instruction_reads_as_empty() -> None:
    """Malformed, but it occurs — it must not raise."""
    doc = Document()
    para = doc.add_paragraph()
    _fld_char(para._p, "begin")
    _fld_char(para._p, "end")

    found = read_fields(doc)[0]

    assert found.keyword == ""
    assert found.operands == []


def test_an_unclosed_field_is_not_reported() -> None:
    """A begin with no end is not a field, and must not consume the rest."""
    doc = Document()
    para = doc.add_paragraph()
    _fld_char(para._p, "begin")
    _instr_run(para._p, " PAGE ")

    assert read_fields(doc) == []


def test_ignores_ordinary_text() -> None:
    doc = Document()
    doc.add_paragraph("No fields here at all.")

    assert read_fields(doc) == []


def test_does_not_mutate_the_document() -> None:
    doc = Document()
    build_complex_field(doc.add_paragraph()._p, " PAGE ", "1")
    from lxml import etree

    before = etree.tostring(doc.element.body)
    read_fields(doc)

    assert etree.tostring(doc.element.body) == before
