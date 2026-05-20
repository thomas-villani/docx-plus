"""Tests for ``docx_plus.publishing.add_table_of_figures``."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.publishing import add_caption, add_table_of_figures


def _instructions(body):
    return [t.text for t in xpath(body, ".//w:instrText") if t.text is not None]


# --------------------------------------------------------------------------
# Emission — instruction string assembly.
# --------------------------------------------------------------------------


def test_add_tof_default_caption_type() -> None:
    doc = Document()
    add_table_of_figures(doc.add_paragraph())
    body = doc.element.body
    instructions = _instructions(body)
    assert any('TOC \\c "Figure"' in i for i in instructions)


def test_add_tof_custom_caption_type() -> None:
    doc = Document()
    add_table_of_figures(doc.add_paragraph(), caption_type="Table")
    body = doc.element.body
    instructions = _instructions(body)
    assert any('TOC \\c "Table"' in i for i in instructions)


def test_add_tof_hyperlink_switch_default() -> None:
    doc = Document()
    add_table_of_figures(doc.add_paragraph())
    body = doc.element.body
    instructions = _instructions(body)
    assert any("\\h" in i for i in instructions)


def test_add_tof_no_hyperlink() -> None:
    doc = Document()
    add_table_of_figures(doc.add_paragraph(), hyperlink=False)
    body = doc.element.body
    instr = _instructions(body)[0]
    assert "\\h" not in instr


# --------------------------------------------------------------------------
# Field structure.
# --------------------------------------------------------------------------


def test_add_tof_writes_full_complex_field() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_table_of_figures(p)

    fld_chars = xpath(p._p, ".//w:fldChar")
    char_types = [c.get(qn("w:fldCharType")) for c in fld_chars]
    assert char_types == ["begin", "separate", "end"]


def test_add_tof_returns_begin_run() -> None:
    doc = Document()
    field = add_table_of_figures(doc.add_paragraph())
    assert field.tag == qn("w:r")
    fld_char = field.find(qn("w:fldChar"))
    assert fld_char is not None
    assert fld_char.get(qn("w:fldCharType")) == "begin"


# --------------------------------------------------------------------------
# End-to-end with captions — the caption_type contract.
# --------------------------------------------------------------------------


def test_tof_caption_type_matches_caption_seq_name(tmp_path: Path) -> None:
    """``add_caption`` and ``add_table_of_figures`` share a name vocabulary."""
    doc = Document()
    add_table_of_figures(doc.add_paragraph(), caption_type="Figure")
    for n in range(3):
        cap = doc.add_paragraph()
        add_caption(cap, "Figure ", caption_type="Figure")
        cap.add_run(f": item {n}")

    out = tmp_path / "tof.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    instructions = _instructions(reopened.element.body)
    # One TOC field instruction plus three SEQ caption instructions.
    tof_count = sum(1 for i in instructions if i and "TOC" in i)
    seq_count = sum(1 for i in instructions if i and "SEQ Figure" in i)
    assert tof_count == 1
    assert seq_count == 3


# --------------------------------------------------------------------------
# Validation — H11 (identifier injection at caption_type).
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "bad_caption_type",
    ["", "1Bad", "has space", 'Figure" \\o "1-9'],
)
def test_add_tof_rejects_bad_caption_type(bad_caption_type: str) -> None:
    doc = Document()
    with pytest.raises(ValueError, match="caption_type"):
        add_table_of_figures(doc.add_paragraph(), caption_type=bad_caption_type)
