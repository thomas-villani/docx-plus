"""Tests for ``docx_plus.layout.set_line_numbering``."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus.core.ns import qn
from docx_plus.core.oxml import el, xpath
from docx_plus.layout import set_line_numbering


def _ln(section):
    return section._sectPr.find(qn("w:lnNumType"))


# --------------------------------------------------------------------------
# Emission — element shape and attribute values.
# --------------------------------------------------------------------------


def test_set_line_numbering_writes_element() -> None:
    doc = Document()
    set_line_numbering(doc.sections[0])
    assert _ln(doc.sections[0]) is not None


def test_set_line_numbering_default_attrs() -> None:
    doc = Document()
    set_line_numbering(doc.sections[0])
    node = _ln(doc.sections[0])
    assert node.get(qn("w:countBy")) == "1"
    assert node.get(qn("w:start")) == "1"
    assert node.get(qn("w:restart")) == "newPage"
    # distance is intentionally omitted so Word picks its built-in default
    assert node.get(qn("w:distance")) is None


def test_set_line_numbering_custom_values() -> None:
    doc = Document()
    set_line_numbering(
        doc.sections[0],
        count_by=5,
        restart="continuous",
        start=11,
        distance=360,
    )
    node = _ln(doc.sections[0])
    assert node.get(qn("w:countBy")) == "5"
    assert node.get(qn("w:restart")) == "continuous"
    assert node.get(qn("w:start")) == "11"
    assert node.get(qn("w:distance")) == "360"


# --------------------------------------------------------------------------
# Replacement — second call overrides rather than stacks.
# --------------------------------------------------------------------------


def test_set_line_numbering_is_idempotent() -> None:
    doc = Document()
    set_line_numbering(doc.sections[0], count_by=1)
    set_line_numbering(doc.sections[0], count_by=10)
    found = xpath(doc.sections[0]._sectPr, "./w:lnNumType")
    assert len(found) == 1
    assert found[0].get(qn("w:countBy")) == "10"


# --------------------------------------------------------------------------
# Validation — argument guards.
# --------------------------------------------------------------------------


def test_set_line_numbering_rejects_zero_count_by() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="count_by >= 1"):
        set_line_numbering(doc.sections[0], count_by=0)


def test_set_line_numbering_rejects_zero_start() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="start >= 1"):
        set_line_numbering(doc.sections[0], start=0)


def test_set_line_numbering_rejects_unknown_restart() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="restart must be"):
        set_line_numbering(doc.sections[0], restart="bogus")  # type: ignore[arg-type]


def test_set_line_numbering_rejects_negative_distance() -> None:
    """L15: a negative distance is rejected at the API boundary."""
    doc = Document()
    with pytest.raises(ValueError, match="distance >= 0"):
        set_line_numbering(doc.sections[0], distance=-1)


def test_set_line_numbering_accepts_zero_distance() -> None:
    doc = Document()
    set_line_numbering(doc.sections[0], distance=0)
    ln = doc.sections[0]._sectPr.find(qn("w:lnNumType"))
    assert ln is not None and ln.get(qn("w:distance")) == "0"


# --------------------------------------------------------------------------
# Schema-strict insertion — lnNumType lands before its later siblings.
# --------------------------------------------------------------------------


def test_set_line_numbering_lands_before_cols() -> None:
    """``lnNumType`` is schema-positioned before ``cols``."""
    doc = Document()
    sect_pr = doc.sections[0]._sectPr
    # Pre-seed an empty `<w:cols>` so insert_before_first_anchor has a target.
    cols = el("w:cols", **{"w:num": "1"})
    sect_pr.append(cols)

    set_line_numbering(doc.sections[0], count_by=2)

    children = list(sect_pr)
    ln_idx = next(i for i, c in enumerate(children) if c.tag == qn("w:lnNumType"))
    cols_idx = next(i for i, c in enumerate(children) if c.tag == qn("w:cols"))
    assert ln_idx < cols_idx


def test_set_line_numbering_lands_before_cols_and_docGrid() -> None:
    """L19: with both cols and docGrid present, lnNumType precedes both.

    A fresh Document() already carries cols + docGrid in schema order, so
    this exercises the multi-anchor case that would catch a misordered
    ``_LATER_SIBLINGS`` (e.g. dropping ``w:cols``).
    """
    doc = Document()
    sect_pr = doc.sections[0]._sectPr
    assert sect_pr.find(qn("w:cols")) is not None
    assert sect_pr.find(qn("w:docGrid")) is not None

    set_line_numbering(doc.sections[0], count_by=1)

    tags = [c.tag for c in sect_pr]
    ln_idx = tags.index(qn("w:lnNumType"))
    assert ln_idx < tags.index(qn("w:cols"))
    assert ln_idx < tags.index(qn("w:docGrid"))


def test_set_line_numbering_replaces_preexisting_element() -> None:
    """L18: a pre-seeded lnNumType (as if loaded from Word) is replaced in place."""
    doc = Document()
    sect_pr = doc.sections[0]._sectPr
    sect_pr.append(
        el("w:lnNumType", **{"w:countBy": "1", "w:restart": "newPage", "w:start": "1"})
    )

    set_line_numbering(doc.sections[0], count_by=7, restart="continuous", start=3)

    found = xpath(sect_pr, "./w:lnNumType")
    assert len(found) == 1
    assert found[0].get(qn("w:countBy")) == "7"
    assert found[0].get(qn("w:restart")) == "continuous"
    assert found[0].get(qn("w:start")) == "3"


# --------------------------------------------------------------------------
# Round-trip — save / reopen preserves the element.
# --------------------------------------------------------------------------


def test_set_line_numbering_round_trip(tmp_path: Path) -> None:
    doc = Document()
    set_line_numbering(doc.sections[0], count_by=5, restart="newSection", start=21)
    out = tmp_path / "ln.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    node = _ln(reopened.sections[0])
    assert node is not None
    assert node.get(qn("w:countBy")) == "5"
    assert node.get(qn("w:restart")) == "newSection"
    assert node.get(qn("w:start")) == "21"
