"""Tests for ``docx_plus.controls.builder.FormBuilder``."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus._testing.ooxml_asserts import assert_ids_unique, count_controls
from docx_plus.controls import FormBuilder, MissingNamespaceError
from docx_plus.core.ids import IdRegistry
from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath

# --------------------------------------------------------------------------
# Construction.
# --------------------------------------------------------------------------


def test_formbuilder_blank_document() -> None:
    fb = FormBuilder()
    assert fb.doc is not None
    # PlaceholderText style is materialised on construction.
    styles = fb.doc.styles.element
    matches = xpath(styles, "./w:style[@w:styleId='PlaceholderText']")
    assert len(matches) == 1


def test_formbuilder_from_path(tmp_path: Path, empty_docx_path: Path) -> None:
    fb = FormBuilder(str(empty_docx_path))
    assert fb.doc is not None


def test_formbuilder_from_pathlike(empty_docx_path: Path) -> None:
    fb = FormBuilder(empty_docx_path)
    assert fb.doc is not None


def test_formbuilder_from_existing_document() -> None:
    doc = Document()
    fb = FormBuilder(doc)
    assert fb.doc is doc


def test_formbuilder_uses_provided_id_registry() -> None:
    doc = Document()
    registry = IdRegistry(doc)
    registry.reserve(42)
    fb = FormBuilder(doc, id_registry=registry)
    p = fb.doc.add_paragraph()
    fb.add_text_control(p, tag="t")
    assert 42 in registry.issued()
    # Builder issued one new id; registry holds at least 2 (42 + new).
    assert len(registry.issued()) >= 2


# --------------------------------------------------------------------------
# Per-control-type SDT structure (round-trip via save+reopen).
# --------------------------------------------------------------------------


def _build_and_reload(fb: FormBuilder, tmp_path: Path) -> Document:
    out = tmp_path / "out.docx"
    fb.save(out)
    return Document(out)


def test_text_control_round_trip(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph("Name: ")
    fb.add_text_control(p, tag="full_name", placeholder="Type your name")
    doc = _build_and_reload(fb, tmp_path)
    assert count_controls(doc, "text") == 1
    assert_ids_unique(doc)
    sdt = xpath(doc.element.body, ".//w:sdt")[0]
    sdt_pr = sdt.find(qn("w:sdtPr"))
    assert sdt_pr.find(qn("w:text")) is not None
    assert sdt_pr.find(qn("w:showingPlcHdr")) is not None


def test_text_control_multiline(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_text_control(p, tag="comments", multiline=True)
    doc = _build_and_reload(fb, tmp_path)
    text_el = xpath(doc.element.body, ".//w:sdtPr/w:text")[0]
    assert text_el.get(qn("w:multiLine")) == "1"


def test_dropdown_round_trip(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_dropdown(p, tag="dept", items=["Eng", "Sales", "Ops"])
    doc = _build_and_reload(fb, tmp_path)
    assert count_controls(doc, "dropdown") == 1
    list_el = xpath(doc.element.body, ".//w:dropDownList")[0]
    items = list_el.findall(qn("w:listItem"))
    # Includes the auto-prepended placeholder list-item.
    assert len(items) == 4
    # First is placeholder with empty value.
    assert items[0].get(qn("w:value")) == ""
    assert items[1].get(qn("w:displayText")) == "Eng"
    assert items[1].get(qn("w:value")) == "Eng"


def test_dropdown_with_display_value_pairs(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_dropdown(
        p,
        tag="priority",
        items=[("High", "P1"), ("Medium", "P2"), ("Low", "P3")],
    )
    doc = _build_and_reload(fb, tmp_path)
    items = xpath(doc.element.body, ".//w:dropDownList/w:listItem")
    real_items = [it for it in items if it.get(qn("w:value")) != ""]
    assert real_items[0].get(qn("w:displayText")) == "High"
    assert real_items[0].get(qn("w:value")) == "P1"


def test_dropdown_editable_emits_combobox(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_dropdown(p, tag="referral", items=["Web", "Friend"], editable=True)
    doc = _build_and_reload(fb, tmp_path)
    assert count_controls(doc, "combobox") == 1
    assert count_controls(doc, "dropdown") == 0
    assert xpath(doc.element.body, ".//w:dropDownList") == []
    assert len(xpath(doc.element.body, ".//w:comboBox")) == 1


def test_dropdown_rejects_invalid_item() -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    with pytest.raises(TypeError):
        fb.add_dropdown(p, tag="bad", items=[123])  # type: ignore[list-item]


def test_date_picker_round_trip(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_date_picker(p, tag="visit", date_format="yyyy-MM-dd", lcid="en-GB")
    doc = _build_and_reload(fb, tmp_path)
    assert count_controls(doc, "date") == 1
    date_el = xpath(doc.element.body, ".//w:date")[0]
    assert date_el.find(qn("w:dateFormat")).get(qn("w:val")) == "yyyy-MM-dd"
    assert date_el.find(qn("w:lid")).get(qn("w:val")) == "en-GB"
    assert date_el.find(qn("w:calendar")).get(qn("w:val")) == "gregorian"


def test_checkbox_unchecked_round_trip(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_checkbox(p, tag="agree")
    doc = _build_and_reload(fb, tmp_path)
    assert count_controls(doc, "checkbox") == 1
    checked = xpath(doc.element.body, ".//w14:checkbox/w14:checked")[0]
    assert checked.get(qn("w14:val")) == "0"
    glyph_t = xpath(doc.element.body, ".//w:sdt//w:t")[0]
    assert glyph_t.text == "☐"


def test_checkbox_checked_round_trip(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_checkbox(p, tag="urgent", checked=True)
    doc = _build_and_reload(fb, tmp_path)
    checked = xpath(doc.element.body, ".//w14:checkbox/w14:checked")[0]
    assert checked.get(qn("w14:val")) == "1"
    glyph_t = xpath(doc.element.body, ".//w:sdt//w:t")[0]
    assert glyph_t.text == "☒"


# --------------------------------------------------------------------------
# Multi-control documents and IdRegistry continuity.
# --------------------------------------------------------------------------


def test_multiple_controls_unique_ids(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_text_control(p, tag="t1")
    fb.add_text_control(p, tag="t2")
    fb.add_dropdown(p, tag="d1", items=["a"])
    fb.add_date_picker(p, tag="dt1")
    fb.add_checkbox(p, tag="cb1")
    doc = _build_and_reload(fb, tmp_path)
    assert count_controls(doc) == 5
    assert_ids_unique(doc)


def test_controls_in_separate_paragraphs(tmp_path: Path) -> None:
    fb = FormBuilder()
    p1 = fb.doc.add_paragraph("Name: ")
    fb.add_text_control(p1, tag="name")
    p2 = fb.doc.add_paragraph("Visit date: ")
    fb.add_date_picker(p2, tag="visit_date")
    doc = _build_and_reload(fb, tmp_path)
    assert count_controls(doc) == 2


def test_existing_id_seeded_from_doc(tmp_path: Path) -> None:
    """If a doc already has an SDT with id=999, the registry won't reuse it."""
    fb = FormBuilder()
    # Inject a hand-rolled SDT with a known id.
    p = fb.doc.add_paragraph()
    from docx_plus.core.oxml import sub

    sdt = sub(p._p, "w:sdt")
    pr = sub(sdt, "w:sdtPr")
    sub(pr, "w:tag", **{"w:val": "preexisting"})
    sub(pr, "w:id", **{"w:val": "999"})
    sub(pr, "w:text")
    sub(sdt, "w:sdtContent")

    # New builder seeded from this doc should not reissue 999.
    fb2 = FormBuilder(fb.doc)
    p2 = fb2.doc.add_paragraph()
    fb2.add_text_control(p2, tag="new")
    doc = _build_and_reload(fb2, tmp_path)
    ids = [int(el.get(qn("w:val"))) for el in xpath(doc.element.body, ".//w:sdt/w:sdtPr/w:id")]
    assert 999 in ids
    assert len(ids) == 2
    assert len(set(ids)) == 2


# --------------------------------------------------------------------------
# Schema-strict sdtPr child order.
# --------------------------------------------------------------------------


def test_sdtpr_child_order_text(tmp_path: Path) -> None:
    """sdtPr children must follow the order [alias?], tag, id, [showingPlcHdr?], <type-marker>."""
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_text_control(p, tag="t", alias="My text")
    doc = _build_and_reload(fb, tmp_path)
    sdt_pr = xpath(doc.element.body, ".//w:sdtPr")[0]
    local_names = [child.tag.rsplit("}", 1)[-1] for child in sdt_pr]
    assert local_names == ["alias", "tag", "id", "showingPlcHdr", "text"]


def test_sdtpr_child_order_dropdown_no_alias(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_dropdown(p, tag="d", items=["a"])
    doc = _build_and_reload(fb, tmp_path)
    sdt_pr = xpath(doc.element.body, ".//w:sdtPr")[0]
    local_names = [child.tag.rsplit("}", 1)[-1] for child in sdt_pr]
    assert local_names == ["tag", "id", "showingPlcHdr", "dropDownList"]


def test_sdtpr_child_order_checkbox(tmp_path: Path) -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_checkbox(p, tag="c")
    doc = _build_and_reload(fb, tmp_path)
    sdt_pr = xpath(doc.element.body, ".//w:sdtPr")[0]
    local_names = [child.tag.rsplit("}", 1)[-1] for child in sdt_pr]
    # Checkboxes have no showingPlcHdr (no placeholder concept).
    assert local_names == ["tag", "id", "checkbox"]


# --------------------------------------------------------------------------
# Placeholder style materialisation.
# --------------------------------------------------------------------------


def test_placeholder_style_materialised_once() -> None:
    fb = FormBuilder()
    for i in range(5):
        p = fb.doc.add_paragraph()
        fb.add_text_control(p, tag=f"t{i}")
    matches = xpath(fb.doc.styles.element, "./w:style[@w:styleId='PlaceholderText']")
    assert len(matches) == 1


def test_placeholder_style_preserved_if_already_defined() -> None:
    """If the doc already defines PlaceholderText, FormBuilder leaves it alone."""
    doc = Document()
    from docx_plus.core.oxml import sub

    custom = sub(
        doc.styles.element,
        "w:style",
        **{"w:type": "character", "w:styleId": "PlaceholderText"},
    )
    sub(custom, "w:name", **{"w:val": "Custom Placeholder"})

    FormBuilder(doc)
    matches = xpath(doc.styles.element, "./w:style[@w:styleId='PlaceholderText']")
    assert len(matches) == 1
    assert matches[0].find(qn("w:name")).get(qn("w:val")) == "Custom Placeholder"


# --------------------------------------------------------------------------
# Save returns the path as a string.
# --------------------------------------------------------------------------


def test_save_returns_path_string(tmp_path: Path) -> None:
    fb = FormBuilder()
    fb.doc.add_paragraph("Hello")
    out = tmp_path / "x.docx"
    result = fb.save(out)
    assert result == str(out)
    assert out.exists()


# --------------------------------------------------------------------------
# Namespace-declaration guard.
# --------------------------------------------------------------------------


def test_missing_w14_namespace_raises() -> None:
    """If a doc somehow lacks w14, FormBuilder construction must raise."""

    class _FakeDoc:
        class _Element:
            nsmap: dict[str | None, str] = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }

        element = _Element()
        # Stub everything FormBuilder.__init__ touches before _verify.
        styles = _Element()  # not really used because we'll fail earlier

    fake = _FakeDoc()
    fake.element.body = None  # type: ignore[attr-defined]

    with pytest.raises(MissingNamespaceError):
        # Bypass the IdRegistry seeding by passing a registry directly; the
        # IdRegistry constructor would fail on the fake first otherwise.
        # We get the same effect by using a real Document but mutating its
        # nsmap in place — which is impossible in lxml. So this is the
        # cleanest test: just call _verify_w14_declared directly via FormBuilder.
        from docx_plus.controls.builder import _verify_w14_declared

        _verify_w14_declared(fake)  # type: ignore[arg-type]
