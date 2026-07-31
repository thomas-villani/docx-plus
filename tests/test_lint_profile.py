"""Tests for the lint profile — the one place a house opinion may live.

Two things are under test and they pull in opposite directions. A profile
has to be able to change what the linter reports, or it is useless; and it
must never be able to stop someone asking a direct question of a single
document, or a checked-in file becomes a way to hide findings.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from docx_plus.lint import (
    DEFAULT_PROFILE_NAME,
    InvalidProfileError,
    Profile,
    UnknownRuleError,
    lint,
)


@pytest.fixture
def spaced_doc() -> Document:
    """A document with exactly one defect: a double space."""
    doc = Document()
    doc.add_paragraph("Body with  two spaces.")
    return doc


def _write(directory: Path, payload: object, name: str = DEFAULT_PROFILE_NAME) -> Path:
    path = directory / name
    path.write_text(json.dumps(payload), encoding="utf-8")
    return path


# --------------------------------------------------------------------------
# Loading.
# --------------------------------------------------------------------------


def test_no_profile_is_the_empty_profile() -> None:
    """So every caller can pass its argument straight through without branching."""
    profile = Profile.load(None)

    assert profile.rules == {}
    assert profile.enabled("double-space", default=True) is True
    assert profile.severity("double-space", default="info") == "info"


def test_a_mapping_loads_without_touching_the_filesystem() -> None:
    profile = Profile.load({"rules": {"double-space": {"enabled": False}}})

    assert profile.enabled("double-space", default=True) is False


def test_a_path_loads(tmp_path: Path) -> None:
    path = _write(tmp_path, {"rules": {"double-space": {"severity": "error"}}})

    profile = Profile.load(path)

    assert profile.severity("double-space", default="info") == "error"


def test_a_rule_the_profile_does_not_mention_keeps_its_registered_behaviour() -> None:
    """A profile is a set of deltas, not a replacement catalogue."""
    profile = Profile.load({"rules": {"double-space": {"enabled": False}}})

    assert profile.enabled("trailing-whitespace", default=True) is True
    assert profile.enabled("font-outliers", default=False) is False


def test_options_are_carried_but_unread() -> None:
    """The hook policy rules will read their target through; nothing uses it yet."""
    profile = Profile.load({"rules": {"font-outliers": {"options": {"max_share": 0.02}}}})

    assert profile.option("font-outliers", "max_share") == 0.02
    assert profile.option("font-outliers", "missing", "fallback") == "fallback"
    assert profile.option("double-space", "max_share") is None


# --------------------------------------------------------------------------
# Rejecting a broken profile — loudly, and on load.
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("payload", "fragment"),
    [
        (["double-space"], "must be a JSON object"),
        ({"ruels": {}}, "unknown profile key"),
        ({"rules": []}, "keyed by rule id"),
        ({"rules": {"double-space": []}}, "must be an object"),
        ({"rules": {"double-space": {"enabld": True}}}, "unknown key"),
        ({"rules": {"double-space": {"enabled": "yes"}}}, "must be true or false"),
        ({"rules": {"double-space": {"severity": "fatal"}}}, "must be one of"),
        ({"rules": {"double-space": {"options": 3}}}, "must be an object"),
    ],
)
def test_a_malformed_profile_is_rejected(payload: object, fragment: str) -> None:
    with pytest.raises(InvalidProfileError, match=fragment):
        Profile.load(payload)  # type: ignore[arg-type]


def test_a_profile_that_is_not_json_is_rejected(tmp_path: Path) -> None:
    path = tmp_path / DEFAULT_PROFILE_NAME
    path.write_text("{not json", encoding="utf-8")

    with pytest.raises(InvalidProfileError, match="not valid JSON"):
        Profile.load(path)


def test_a_profile_that_is_not_an_object_is_rejected(tmp_path: Path) -> None:
    path = _write(tmp_path, ["double-space"])

    with pytest.raises(InvalidProfileError, match="must be a JSON object"):
        Profile.load(path)


def test_an_unreadable_profile_is_rejected(tmp_path: Path) -> None:
    with pytest.raises(InvalidProfileError, match="cannot read profile"):
        Profile.load(tmp_path / "nothing-here.json")


def test_a_profile_naming_an_unknown_rule_is_rejected(spaced_doc: Document) -> None:
    """A typo would otherwise configure nothing and read exactly like success."""
    with pytest.raises(UnknownRuleError, match="unknown rule"):
        lint(spaced_doc, profile={"rules": {"duoble-space": {"enabled": False}}})


def test_a_profile_may_not_configure_a_tag(spaced_doc: Document) -> None:
    """Selectors accept tags; a profile's per-rule settings do not.

    "Apply this severity to whatever carries the tag today" is not a stable
    thing to check into a repository.
    """
    with pytest.raises(UnknownRuleError, match="unknown rule"):
        lint(spaced_doc, profile={"rules": {"typography": {"severity": "error"}}})


# --------------------------------------------------------------------------
# Discovery.
# --------------------------------------------------------------------------


def test_discover_finds_a_profile_beside_the_document(tmp_path: Path) -> None:
    _write(tmp_path, {"rules": {"double-space": {"enabled": False}}})

    profile = Profile.discover(tmp_path / "report.docx")

    assert profile.enabled("double-space", default=True) is False


def test_discover_walks_up_to_the_project_root(tmp_path: Path) -> None:
    _write(tmp_path, {"rules": {"double-space": {"enabled": False}}})
    nested = tmp_path / "drafts" / "2026"
    nested.mkdir(parents=True)

    profile = Profile.discover(nested)

    assert profile.enabled("double-space", default=True) is False


def test_discover_returns_the_empty_profile_when_there_is_none(tmp_path: Path) -> None:
    assert Profile.discover(tmp_path).rules == {}


def test_discover_still_rejects_a_broken_profile(tmp_path: Path) -> None:
    """Silently ignoring a broken checked-in profile is worse than having none."""
    (tmp_path / DEFAULT_PROFILE_NAME).write_text("{", encoding="utf-8")

    with pytest.raises(InvalidProfileError):
        Profile.discover(tmp_path)


# --------------------------------------------------------------------------
# What a profile does to a lint run.
# --------------------------------------------------------------------------


def test_a_profile_can_disable_a_rule(spaced_doc: Document) -> None:
    assert {f.rule for f in lint(spaced_doc)} == {"double-space"}

    findings = lint(spaced_doc, profile={"rules": {"double-space": {"enabled": False}}})

    assert findings == []


def test_a_profile_can_enable_an_off_by_default_rule() -> None:
    doc = Document()
    doc.add_paragraph("before")
    doc.add_paragraph("")
    doc.add_paragraph("")

    assert "stray-empty-paragraph" not in {f.rule for f in lint(doc)}

    findings = lint(doc, profile={"rules": {"stray-empty-paragraph": {"enabled": True}}})

    assert "stray-empty-paragraph" in {f.rule for f in findings}


def test_a_profile_can_change_a_rule_s_severity(spaced_doc: Document) -> None:
    findings = lint(spaced_doc, profile={"rules": {"double-space": {"severity": "error"}}})

    assert [f.severity for f in findings] == ["error"]


def test_a_findings_own_severity_survives_a_profile() -> None:
    """A rule downgrading one finding is saying something about *that* finding.

    ``direct-numbering-override`` reports the ``numId=0`` opt-out at
    ``info`` because it is the one legitimate reason to override. A blanket
    "treat this rule as an error" is not an answer to that.
    """
    from docx_plus.core.ns import qn

    doc = Document()
    paragraph = doc.add_paragraph("item", style="List Bullet")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.makeelement(qn("w:numPr"), {})
    num_pr.append(num_pr.makeelement(qn("w:numId"), {qn("w:val"): "0"}))
    p_pr.append(num_pr)

    findings = lint(
        doc,
        select=["direct-numbering-override"],
        profile={"rules": {"direct-numbering-override": {"severity": "error"}}},
    )

    assert [f.severity for f in findings] == ["info"]


def test_naming_a_rule_explicitly_beats_a_profile_that_disabled_it(
    spaced_doc: Document,
) -> None:
    """Configuration never gets to veto a direct question about one document."""
    findings = lint(
        spaced_doc,
        select=["double-space"],
        profile={"rules": {"double-space": {"enabled": False}}},
    )

    assert [f.rule for f in findings] == ["double-space"]


def test_exclude_still_wins_over_a_profile_that_enabled_a_rule(spaced_doc: Document) -> None:
    """``exclude`` is applied last and always wins — profile included."""
    findings = lint(
        spaced_doc,
        exclude=["double-space"],
        profile={"rules": {"double-space": {"enabled": True}}},
    )

    assert findings == []


def test_a_loaded_profile_object_is_accepted_as_is(spaced_doc: Document) -> None:
    profile = Profile.load({"rules": {"double-space": {"enabled": False}}})

    assert lint(spaced_doc, profile=profile) == []
