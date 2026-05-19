"""Tests for ``docx_plus.protection`` — document-level protection enforcement."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus._testing.ooxml_asserts import assert_protected
from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.protection import (
    is_protected,
    protect_document,
    unprotect_document,
)

# --------------------------------------------------------------------------
# protect_document — basic emission.
# --------------------------------------------------------------------------


def test_protect_document_writes_documentProtection_element() -> None:
    doc = Document()
    protect_document(doc)
    matches = xpath(doc.settings.element, "./w:documentProtection")
    assert len(matches) == 1
    element = matches[0]
    assert element.get(qn("w:edit")) == "forms"
    assert element.get(qn("w:enforcement")) == "1"


@pytest.mark.parametrize("mode", ["forms", "readOnly", "comments", "trackedChanges"])
def test_protect_document_writes_correct_mode(mode: str) -> None:
    doc = Document()
    protect_document(doc, mode=mode)  # type: ignore[arg-type]
    assert_protected(doc, mode=mode)


# --------------------------------------------------------------------------
# Schema position — w:documentProtection precedes w:defaultTabStop.
# --------------------------------------------------------------------------


def test_protect_document_precedes_defaultTabStop() -> None:
    """SPEC §8: ``w:documentProtection`` must come before ``w:defaultTabStop``."""
    doc = Document()
    protect_document(doc)
    settings = doc.settings.element
    children = list(settings)
    protection_idx = None
    tabstop_idx = None
    for i, child in enumerate(children):
        if child.tag == qn("w:documentProtection"):
            protection_idx = i
        elif child.tag == qn("w:defaultTabStop"):
            tabstop_idx = i
    assert protection_idx is not None, "w:documentProtection missing"
    assert tabstop_idx is not None, "w:defaultTabStop missing (python-docx default)"
    assert protection_idx < tabstop_idx, (
        f"w:documentProtection (idx={protection_idx}) is not before "
        f"w:defaultTabStop (idx={tabstop_idx})"
    )


# --------------------------------------------------------------------------
# Idempotency — second call replaces, doesn't stack.
# --------------------------------------------------------------------------


def test_protect_document_is_idempotent() -> None:
    doc = Document()
    protect_document(doc, mode="forms")
    protect_document(doc, mode="forms")
    matches = xpath(doc.settings.element, "./w:documentProtection")
    assert len(matches) == 1


def test_protect_document_second_call_replaces_mode() -> None:
    doc = Document()
    protect_document(doc, mode="forms")
    protect_document(doc, mode="readOnly")
    matches = xpath(doc.settings.element, "./w:documentProtection")
    assert len(matches) == 1
    assert matches[0].get(qn("w:edit")) == "readOnly"


# --------------------------------------------------------------------------
# unprotect_document — removal + idempotency.
# --------------------------------------------------------------------------


def test_unprotect_document_removes_element() -> None:
    doc = Document()
    protect_document(doc)
    unprotect_document(doc)
    matches = xpath(doc.settings.element, "./w:documentProtection")
    assert matches == []


def test_unprotect_document_is_idempotent_on_fresh_doc() -> None:
    doc = Document()
    # Should not raise.
    unprotect_document(doc)
    assert is_protected(doc) is False


def test_unprotect_then_reprotect_works() -> None:
    doc = Document()
    protect_document(doc, mode="forms")
    unprotect_document(doc)
    protect_document(doc, mode="comments")
    assert_protected(doc, mode="comments")


# --------------------------------------------------------------------------
# is_protected predicate.
# --------------------------------------------------------------------------


def test_is_protected_false_on_fresh_doc() -> None:
    doc = Document()
    assert is_protected(doc) is False


def test_is_protected_true_after_protect() -> None:
    doc = Document()
    protect_document(doc)
    assert is_protected(doc) is True


def test_is_protected_false_after_unprotect() -> None:
    doc = Document()
    protect_document(doc)
    unprotect_document(doc)
    assert is_protected(doc) is False


# --------------------------------------------------------------------------
# Round-trip — save and reopen preserves protection.
# --------------------------------------------------------------------------


def test_protect_document_round_trip(tmp_path: Path) -> None:
    doc = Document()
    protect_document(doc, mode="forms")
    out = tmp_path / "protected.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    assert is_protected(reopened) is True
    assert_protected(reopened, mode="forms")


def test_unprotect_round_trip(tmp_path: Path) -> None:
    doc = Document()
    protect_document(doc, mode="forms")
    out1 = tmp_path / "protected.docx"
    doc.save(str(out1))

    reopened = Document(str(out1))
    unprotect_document(reopened)
    out2 = tmp_path / "unprotected.docx"
    reopened.save(str(out2))

    final = Document(str(out2))
    assert is_protected(final) is False


# --------------------------------------------------------------------------
# assert_protected helper catches the failure modes it claims to.
# --------------------------------------------------------------------------


def test_assert_protected_fails_on_unprotected_doc() -> None:
    doc = Document()
    with pytest.raises(AssertionError, match="w:documentProtection is not present"):
        assert_protected(doc)


def test_assert_protected_fails_on_mode_mismatch() -> None:
    doc = Document()
    protect_document(doc, mode="forms")
    with pytest.raises(AssertionError, match="w:edit is 'forms'"):
        assert_protected(doc, mode="readOnly")
