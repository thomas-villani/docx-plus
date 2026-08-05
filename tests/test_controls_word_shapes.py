"""Read-side regression tests against the SDT shapes Word itself writes.

``FormBuilder`` writes tidy controls: every one tagged, every tag unique, every
type marker present, all in the body. Real Word output is looser on all four
counts, and testing only against our own builder let those gaps through. Each
fixture below is a shape observed in Word-authored documents:

- ``<w:tag w:val=""/>`` — what the Developer ribbon writes unless the author
  types a tag, so most controls in a real form share the empty tag.
- no ``w:tag`` element at all — legal; ``w:tag`` is optional in ECMA-376.
- no type marker — a rich-text control, the default for the ``w:sdtPr``
  choice group and the most common control Word inserts.
- controls in headers, footers, footnotes, and endnotes.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from docx_plus.controls import (
    ControlNotFoundError,
    DuplicateTagError,
    clear_control,
    list_controls,
    read_controls,
    set_control_value,
)
from docx_plus.core.ns import NSMAP
from docx_plus.core.parts import ENDNOTES_SPEC, FOOTNOTES_SPEC, get_or_create_part

W = NSMAP["w"]
W14 = NSMAP["w14"]
W15 = NSMAP["w15"]


# --------------------------------------------------------------------------
# Helpers.
# --------------------------------------------------------------------------


def _sdt(inner_pr: str, *, content: str = "<w:r><w:t>x</w:t></w:r>") -> etree._Element:
    """Build a ``w:sdt`` element from a ``w:sdtPr`` fragment."""
    xml = (
        f'<w:sdt xmlns:w="{W}" xmlns:w14="{W14}" xmlns:w15="{W15}">'
        f"<w:sdtPr>{inner_pr}</w:sdtPr>"
        f"<w:sdtContent>{content}</w:sdtContent>"
        f"</w:sdt>"
    )
    return etree.fromstring(xml)


def _doc_with(*sdt_prs: str) -> Document:
    """A document with one paragraph-hosted control per ``w:sdtPr`` fragment."""
    doc = Document()
    for pr in sdt_prs:
        para = doc.add_paragraph()
        para._p.append(_sdt(pr))
    return doc


#: The empty-tag plain-text control Word writes for an untagged form field.
_EMPTY_TAG_TEXT = '<w:alias w:val="{alias}"/><w:tag w:val=""/><w:id w:val="{id}"/><w:text/>'


# --------------------------------------------------------------------------
# Empty and absent tags.
# --------------------------------------------------------------------------


def test_repeated_empty_tags_do_not_crash_read_controls() -> None:
    """Ten empty-tag controls is normal Word output, not a duplicate-key error."""
    doc = _doc_with(*(_EMPTY_TAG_TEXT.format(alias=f"Field {n}", id=n) for n in range(1, 11)))
    assert read_controls(doc) == {}
    assert len(list_controls(doc)) == 10


def test_empty_tag_is_empty_string_not_none() -> None:
    """``<w:tag w:val=""/>`` and a missing ``w:tag`` are different states."""
    doc = _doc_with(
        '<w:tag w:val=""/><w:id w:val="1"/><w:text/>',
        '<w:id w:val="2"/><w:text/>',
    )
    present, absent = list_controls(doc)
    assert present.tag == ""
    assert absent.tag is None


def test_untagged_control_is_reported_not_dropped() -> None:
    """A control with no ``w:tag`` used to vanish from the read entirely."""
    doc = _doc_with('<w:id w:val="7"/><w:text/>')
    (only,) = list_controls(doc)
    assert only.tag is None
    assert only.control_id == 7
    assert only.value == "x"


def test_read_controls_still_raises_on_real_duplicates() -> None:
    """An empty tag is unkeyable; a repeated *real* tag is genuine ambiguity."""
    doc = _doc_with(
        '<w:tag w:val="dup"/><w:id w:val="1"/><w:text/>',
        '<w:tag w:val="dup"/><w:id w:val="2"/><w:text/>',
    )
    with pytest.raises(DuplicateTagError, match="dup"):
        read_controls(doc)


def test_read_controls_by_alias_raises_on_duplicate_alias() -> None:
    """Aliases are UI labels and repeat freely, so this path collides too."""
    doc = _doc_with(
        '<w:alias w:val="Date"/><w:tag w:val="a"/><w:id w:val="1"/><w:text/>',
        '<w:alias w:val="Date"/><w:tag w:val="b"/><w:id w:val="2"/><w:text/>',
    )
    with pytest.raises(DuplicateTagError, match="Date"):
        read_controls(doc, by="alias")


# --------------------------------------------------------------------------
# Type coverage.
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("marker", "expected"),
    [
        ("<w:text/>", "text"),
        ("<w:dropDownList/>", "dropdown"),
        ("<w:comboBox/>", "combobox"),
        ("<w:date/>", "date"),
        ("<w14:checkbox/>", "checkbox"),
        ("", "richtext"),
        ("<w:richText/>", "richtext"),
        ("<w:picture/>", "picture"),
        ("<w:group/>", "group"),
        ("<w15:repeatingSection/>", "repeating"),
        ("<w15:repeatingSectionItem/>", "repeatingitem"),
        ("<w:docPartObj/>", "docpart"),
        ("<w:docPartList/>", "docpart"),
        ("<w:citation/>", "citation"),
        ("<w:bibliography/>", "bibliography"),
        ("<w:equation/>", "equation"),
    ],
)
def test_every_sdt_type_marker_is_classified(marker: str, expected: str) -> None:
    doc = _doc_with(f'<w:tag w:val="t"/><w:id w:val="1"/>{marker}')
    (only,) = list_controls(doc)
    assert only.control_type == expected


def test_container_controls_are_read_only() -> None:
    """A repeating section wraps block content; there is no scalar to set."""
    doc = _doc_with('<w:tag w:val="rows"/><w:id w:val="1"/><w15:repeatingSection/>')
    with pytest.raises(TypeError, match="repeating"):
        set_control_value(doc, "rows", "nope")
    with pytest.raises(TypeError, match="repeating"):
        clear_control(doc, "rows")


# --------------------------------------------------------------------------
# Stories beyond the body.
# --------------------------------------------------------------------------


def test_controls_in_header_and_footer_are_found() -> None:
    doc = Document()
    section = doc.sections[0]
    for hdrftr, tag in ((section.header, "hdr"), (section.footer, "ftr")):
        hdrftr.is_linked_to_previous = False
        para = hdrftr.paragraphs[0]
        para._p.append(_sdt(f'<w:tag w:val="{tag}"/><w:id w:val="1"/><w:text/>'))

    found = {c.tag: c.location for c in list_controls(doc)}
    assert found == {"hdr": "header:1:primary", "ftr": "footer:1:primary"}


def test_controls_in_footnotes_are_found() -> None:
    doc = Document()
    _part, root = get_or_create_part(doc, FOOTNOTES_SPEC)
    note = etree.SubElement(root, f"{{{W}}}footnote")
    note.set(f"{{{W}}}id", "2")
    para = etree.SubElement(note, f"{{{W}}}p")
    para.append(_sdt('<w:tag w:val="note"/><w:id w:val="1"/><w:text/>'))

    (only,) = list_controls(doc)
    assert only.tag == "note"
    assert only.location == "footnotes"


def test_controls_in_endnotes_are_found() -> None:
    doc = Document()
    _part, root = get_or_create_part(doc, ENDNOTES_SPEC)
    note = etree.SubElement(root, f"{{{W}}}endnote")
    note.set(f"{{{W}}}id", "2")
    para = etree.SubElement(note, f"{{{W}}}p")
    para.append(_sdt('<w:tag w:val="end"/><w:id w:val="1"/><w:text/>'))

    (only,) = list_controls(doc)
    assert only.location == "endnotes"


def test_reading_does_not_create_header_parts() -> None:
    """python-docx materialises an absent header on access; the read must not."""
    doc = Document()
    before = {p.partname for p in doc.part.package.iter_parts()}
    list_controls(doc)
    after = {p.partname for p in doc.part.package.iter_parts()}
    assert before == after


def test_body_is_walked_before_headers() -> None:
    """``index`` follows the documented story order, body first."""
    doc = Document()
    doc.add_paragraph()._p.append(_sdt('<w:tag w:val="b"/><w:id w:val="1"/><w:text/>'))
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    header.paragraphs[0]._p.append(_sdt('<w:tag w:val="h"/><w:id w:val="2"/><w:text/>'))

    assert [(c.index, c.tag) for c in list_controls(doc)] == [(0, "b"), (1, "h")]


# --------------------------------------------------------------------------
# Ambiguous writes.
# --------------------------------------------------------------------------


def test_write_to_ambiguous_tag_is_refused() -> None:
    """Silently writing to the first of N matches corrupted the other N-1."""
    doc = _doc_with(*(_EMPTY_TAG_TEXT.format(alias=f"Field {n}", id=n) for n in (1, 2, 3)))
    with pytest.raises(DuplicateTagError, match="matches 3 controls"):
        set_control_value(doc, "", "value")
    with pytest.raises(DuplicateTagError, match="matches 3 controls"):
        clear_control(doc, "")


def test_control_id_disambiguates_a_repeated_tag(tmp_path: Path) -> None:
    doc = _doc_with(*(_EMPTY_TAG_TEXT.format(alias=f"Field {n}", id=n) for n in (1, 2, 3)))
    set_control_value(doc, None, "picked", control_id=2)

    out = tmp_path / "out.docx"
    doc.save(out)
    values = [c.value for c in list_controls(Document(out))]
    assert values == ["x", "picked", "x"]


def test_unknown_control_id_raises() -> None:
    doc = _doc_with('<w:tag w:val="a"/><w:id w:val="1"/><w:text/>')
    with pytest.raises(ControlNotFoundError, match="id 99"):
        set_control_value(doc, None, "v", control_id=99)


def test_neither_selector_raises() -> None:
    doc = _doc_with('<w:tag w:val="a"/><w:id w:val="1"/><w:text/>')
    with pytest.raises(ControlNotFoundError, match="tag or control_id"):
        set_control_value(doc, None, "v")


def test_non_numeric_control_id_reads_as_none() -> None:
    """Word writes numeric ids, but the schema does not guarantee it."""
    doc = _doc_with('<w:tag w:val="a"/><w:id w:val="not-a-number"/><w:text/>')
    (only,) = list_controls(doc)
    assert only.control_id is None


def test_sdt_without_sdtpr_is_skipped_not_raised() -> None:
    """A ``w:sdt`` with no ``w:sdtPr`` is malformed; it must not break the walk."""
    doc = Document()
    para = doc.add_paragraph()
    broken = etree.fromstring(
        f'<w:sdt xmlns:w="{W}"><w:sdtContent><w:r><w:t>x</w:t></w:r></w:sdtContent></w:sdt>'
    )
    para._p.append(broken)
    doc.add_paragraph()._p.append(_sdt('<w:tag w:val="ok"/><w:id w:val="1"/><w:text/>'))

    assert [c.tag for c in list_controls(doc)] == ["ok"]
    with pytest.raises(ControlNotFoundError):
        set_control_value(doc, "missing", "x")


def test_control_with_no_id_is_still_writable_by_tag() -> None:
    """``w:id`` is optional too, so the tag has to remain a usable selector."""
    doc = _doc_with('<w:tag w:val="only"/><w:text/>')
    set_control_value(doc, "only", "written")
    (only,) = list_controls(doc)
    assert only.control_id is None
    assert only.value == "written"


def test_control_id_selector_reports_the_type_error_by_id() -> None:
    """The rejection message names whichever selector the caller used."""
    doc = _doc_with('<w:tag w:val="rich"/><w:id w:val="5"/>')
    with pytest.raises(TypeError, match="id 5"):
        set_control_value(doc, None, "x", control_id=5)
