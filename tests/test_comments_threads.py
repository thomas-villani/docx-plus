"""Tests for ``docx_plus.comments.threads`` — replies, resolve / reopen,
thread reads, and the ``commentsExtended.xml`` plumbing behind them."""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from docx_plus._testing.ooxml_asserts import assert_para_ids_unique
from docx_plus.comments import (
    CommentNotFoundError,
    CommentThread,
    _extended,
    add_comment,
    clear_all_comments,
    delete_comment,
    edit_comment,
    read_comments,
    read_threads,
    reopen_comment,
    reply_to_comment,
    resolve_comment,
)
from docx_plus.core.ids import ParaIdRegistry
from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.core.parts import RT_COMMENTS_EXTENDED

# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------


def _threaded_doc(text: str = "Hello world"):
    """Return ``(doc, root_ref)`` — one paragraph carrying one comment."""
    doc = Document()
    paragraph = doc.add_paragraph(text)
    return doc, add_comment(paragraph, "Is this right?", author="Reviewer")


def _round_trip(doc):
    """Save and reload ``doc`` so assertions run against parsed-from-disk XML."""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Document(buffer)


def _ex_entries(doc):
    root = _extended.extended_root(doc)
    return [] if root is None else xpath(root, "./w15:commentEx")


def _entry_for(doc, comment_id):
    keys, _ = _extended.key_maps(doc)
    root = _extended.extended_root(doc)
    assert root is not None
    return _extended.find_comment_ex(root, keys[comment_id])


def _markers(doc, comment_id):
    """Return the body-side marker triple counts for ``comment_id``."""
    body = doc.element.body
    cid = str(comment_id)
    return (
        len(xpath(body, ".//w:commentRangeStart[@w:id=$cid]", cid=cid)),
        len(xpath(body, ".//w:commentRangeEnd[@w:id=$cid]", cid=cid)),
        len(xpath(body, ".//w:commentReference[@w:id=$cid]", cid=cid)),
    )


# --------------------------------------------------------------------------
# add_comment now writes thread metadata eagerly.
# --------------------------------------------------------------------------


def test_add_comment_stamps_para_id_on_the_body_paragraph() -> None:
    doc, ref = _threaded_doc()
    para_id = ref.body_element.find(qn("w:p")).get(qn("w14:paraId"))
    assert para_id is not None
    assert len(para_id) == 8
    assert 0 < int(para_id, 16) < 2**31


def test_add_comment_creates_the_extended_part_with_an_unresolved_entry() -> None:
    doc, ref = _threaded_doc()
    entries = _ex_entries(doc)
    assert len(entries) == 1
    assert entries[0].get(qn("w15:done")) == "0"
    assert entries[0].get(qn("w15:paraIdParent")) is None


def test_add_comment_entry_keys_off_the_comment_body_para_id() -> None:
    doc, ref = _threaded_doc()
    body_para_id = ref.body_element.find(qn("w:p")).get(qn("w14:paraId"))
    assert _ex_entries(doc)[0].get(qn("w15:paraId")) == body_para_id


def test_add_comment_reuses_the_extended_part_across_inserts() -> None:
    doc = Document()
    for index in range(3):
        add_comment(doc.add_paragraph(f"line {index}"), f"note {index}")
    assert len(_ex_entries(doc)) == 3
    assert doc.part.part_related_by(RT_COMMENTS_EXTENDED) is not None


def test_add_comment_accepts_a_shared_para_id_registry() -> None:
    doc = Document()
    registry = ParaIdRegistry(doc)
    for index in range(4):
        add_comment(doc.add_paragraph(f"line {index}"), "note", para_id_registry=registry)
    assert_para_ids_unique(doc)


def test_comments_part_root_declares_the_w14_prefix() -> None:
    doc, _ = _threaded_doc()
    root = doc.part.part_related_by(RT.COMMENTS).element
    assert root.nsmap.get("w14") == "http://schemas.microsoft.com/office/word/2010/wordml"
    assert "w14" in (root.get(qn("mc:Ignorable")) or "")


# --------------------------------------------------------------------------
# reply_to_comment
# --------------------------------------------------------------------------


def test_reply_links_to_its_parent_in_the_extended_part() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.", author="Author")

    keys, _ = _extended.key_maps(doc)
    entry = _entry_for(doc, reply.comment_id)
    assert entry is not None
    assert entry.get(qn("w15:paraIdParent")) == keys[root.comment_id]


def test_reply_gets_its_own_comment_id_and_body() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.", author="Author")
    assert reply.comment_id != root.comment_id

    texts = {comment.comment_id: comment.text for comment in read_comments(doc)}
    assert texts[reply.comment_id] == "Yes."


def test_reply_mirrors_the_parent_anchor_range() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")

    assert _markers(doc, reply.comment_id) == (1, 1, 1)
    comments = {comment.comment_id: comment for comment in read_comments(doc)}
    assert comments[reply.comment_id].anchored_text == "Hello world"
    assert comments[reply.comment_id].paragraph_index == 0


def test_reply_range_start_nests_inside_the_parent_range_start() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")

    paragraph = doc.paragraphs[0]._p
    starts = [
        child.get(qn("w:id")) for child in paragraph if child.tag == qn("w:commentRangeStart")
    ]
    assert starts == [str(root.comment_id), str(reply.comment_id)]


def test_reply_range_end_follows_the_parent_reference_run() -> None:
    doc, root = _threaded_doc()
    reply_to_comment(doc, root.comment_id, "Yes.")

    paragraph = doc.paragraphs[0]._p
    order = [
        child.tag.rpartition("}")[2]
        for child in paragraph
        if child.tag != qn("w:commentRangeStart")
    ]
    # text run, parent end, parent reference run, reply end, reply reference run
    assert order == ["r", "commentRangeEnd", "r", "commentRangeEnd", "r"]


def test_multiple_replies_all_parent_to_the_named_comment() -> None:
    doc, root = _threaded_doc()
    first = reply_to_comment(doc, root.comment_id, "Yes.")
    second = reply_to_comment(doc, root.comment_id, "Agreed.")

    state = _extended.thread_state(doc)
    assert state[first.comment_id][0] == root.comment_id
    assert state[second.comment_id][0] == root.comment_id


def test_reply_to_a_reply_parents_to_that_reply_not_the_root() -> None:
    doc, root = _threaded_doc()
    first = reply_to_comment(doc, root.comment_id, "Yes.")
    nested = reply_to_comment(doc, first.comment_id, "Still yes.")

    state = _extended.thread_state(doc)
    assert state[nested.comment_id][0] == first.comment_id
    assert _extended.root_id_of(state, nested.comment_id) == root.comment_id


def test_reply_to_unknown_comment_raises() -> None:
    doc, _ = _threaded_doc()
    with pytest.raises(CommentNotFoundError):
        reply_to_comment(doc, 999999, "orphan")


def test_reply_to_unknown_comment_is_catchable_as_key_error() -> None:
    doc, _ = _threaded_doc()
    with pytest.raises(KeyError):
        reply_to_comment(doc, 999999, "orphan")


def test_reply_with_no_comments_part_at_all_raises() -> None:
    doc = Document()
    with pytest.raises(CommentNotFoundError):
        reply_to_comment(doc, 1, "orphan")


def test_reply_to_orphaned_parent_writes_no_body_anchors() -> None:
    # An orphaned comment — body present, body-side markers stripped — is the
    # state python-docx's own add_comment produces. The reply inherits it
    # rather than inventing an anchor.
    doc, root = _threaded_doc()
    for expr in (".//w:commentRangeStart", ".//w:commentRangeEnd"):
        for marker in xpath(doc.element.body, expr):
            marker.getparent().remove(marker)

    reply = reply_to_comment(doc, root.comment_id, "Yes.")
    assert _markers(doc, reply.comment_id) == (0, 0, 0)
    state = _extended.thread_state(doc)
    assert state[reply.comment_id][0] == root.comment_id


def test_reply_shares_a_comment_id_registry_without_collision() -> None:
    doc, root = _threaded_doc()
    replies = [reply_to_comment(doc, root.comment_id, f"r{i}") for i in range(5)]
    ids = [reply.comment_id for reply in replies] + [root.comment_id]
    assert len(set(ids)) == len(ids)


def test_reply_para_ids_stay_unique() -> None:
    doc, root = _threaded_doc()
    for index in range(5):
        reply_to_comment(doc, root.comment_id, f"reply {index}")
    assert_para_ids_unique(doc)


def _strip_threading(doc):
    """Reduce ``doc`` to what a pre-2013 producer would have written.

    Drops the extended part and every ``w14:paraId`` stamp, leaving only
    the flat ``comments.xml`` model python-docx and old Word emit.
    """
    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype == RT_COMMENTS_EXTENDED:
            doc.part.drop_rel(rid)
    part = doc.part.part_related_by(RT.COMMENTS)
    for paragraph in xpath(part.element, ".//w:p"):
        if paragraph.get(qn("w14:paraId")) is not None:
            del paragraph.attrib[qn("w14:paraId")]


def test_reply_to_a_foreign_comment_materializes_the_thread() -> None:
    doc, root = _threaded_doc()
    _strip_threading(doc)

    reply = reply_to_comment(doc, root.comment_id, "Yes.")

    # The parent had neither a paraId nor an entry; replying creates both,
    # and the parent lands as an unresolved root rather than inheriting a
    # stale flag.
    assert len(_ex_entries(doc)) == 2
    state = _extended.thread_state(doc)
    assert state[root.comment_id] == (None, False)
    assert state[reply.comment_id][0] == root.comment_id
    assert_para_ids_unique(doc)


def test_resolving_a_foreign_comment_materializes_its_entry() -> None:
    doc, root = _threaded_doc()
    _strip_threading(doc)

    resolve_comment(doc, root.comment_id)
    assert _extended.thread_state(doc)[root.comment_id][1] is True


def test_stamped_comment_without_an_entry_reads_as_an_unresolved_root() -> None:
    # paraId present, commentsExtended.xml entry missing — the halfway state
    # a tool that rewrites comments.xml alone would leave behind.
    doc, root = _threaded_doc()
    entry = _ex_entries(doc)[0]
    entry.getparent().remove(entry)

    assert _extended.thread_state(doc)[root.comment_id] == (None, False)
    assert len(read_threads(doc)) == 1


def test_reply_falls_back_to_the_range_end_when_the_reference_run_is_gone() -> None:
    doc, root = _threaded_doc()
    for ref in xpath(doc.element.body, ".//w:commentReference"):
        ref.getparent().getparent().remove(ref.getparent())

    reply = reply_to_comment(doc, root.comment_id, "Yes.")
    assert _markers(doc, reply.comment_id) == (1, 1, 1)
    assert read_comments(doc)[-1].anchored_text == "Hello world"


# --------------------------------------------------------------------------
# resolve / reopen
# --------------------------------------------------------------------------


def test_resolve_sets_done_on_the_whole_thread() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")

    resolve_comment(doc, root.comment_id)
    assert all(entry.get(qn("w15:done")) == "1" for entry in _ex_entries(doc))
    state = _extended.thread_state(doc)
    assert state[root.comment_id][1] is True
    assert state[reply.comment_id][1] is True


def test_resolving_a_reply_resolves_the_whole_thread() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")

    resolve_comment(doc, reply.comment_id)
    assert all(entry.get(qn("w15:done")) == "1" for entry in _ex_entries(doc))


def test_reopen_clears_done_across_the_thread() -> None:
    doc, root = _threaded_doc()
    reply_to_comment(doc, root.comment_id, "Yes.")

    resolve_comment(doc, root.comment_id)
    reopen_comment(doc, root.comment_id)
    assert all(entry.get(qn("w15:done")) == "0" for entry in _ex_entries(doc))


def test_resolve_leaves_a_sibling_thread_alone() -> None:
    doc = Document()
    first = add_comment(doc.add_paragraph("one"), "a")
    second = add_comment(doc.add_paragraph("two"), "b")

    resolve_comment(doc, first.comment_id)
    state = _extended.thread_state(doc)
    assert state[first.comment_id][1] is True
    assert state[second.comment_id][1] is False


def test_resolve_is_idempotent() -> None:
    doc, root = _threaded_doc()
    resolve_comment(doc, root.comment_id)
    resolve_comment(doc, root.comment_id)
    assert len(_ex_entries(doc)) == 1
    assert _ex_entries(doc)[0].get(qn("w15:done")) == "1"


def test_resolve_unknown_comment_raises() -> None:
    doc, _ = _threaded_doc()
    with pytest.raises(CommentNotFoundError):
        resolve_comment(doc, 999999)


def test_reopen_unknown_comment_raises() -> None:
    doc, _ = _threaded_doc()
    with pytest.raises(CommentNotFoundError):
        reopen_comment(doc, 999999)


# --------------------------------------------------------------------------
# read_threads
# --------------------------------------------------------------------------


def test_read_threads_groups_replies_under_their_root() -> None:
    doc, root = _threaded_doc()
    first = reply_to_comment(doc, root.comment_id, "Yes.")
    second = reply_to_comment(doc, root.comment_id, "Agreed.")

    threads = read_threads(doc)
    assert len(threads) == 1
    assert isinstance(threads[0], CommentThread)
    assert threads[0].root.comment_id == root.comment_id
    assert {reply.comment_id for reply in threads[0].replies} == {
        first.comment_id,
        second.comment_id,
    }


def test_read_threads_reports_resolved_state() -> None:
    doc, root = _threaded_doc()
    reply_to_comment(doc, root.comment_id, "Yes.")
    resolve_comment(doc, root.comment_id)

    thread = read_threads(doc)[0]
    assert thread.resolved is True
    assert all(reply.resolved for reply in thread.replies)


def test_read_threads_flattens_a_deeper_chain_under_the_root() -> None:
    doc, root = _threaded_doc()
    first = reply_to_comment(doc, root.comment_id, "Yes.")
    nested = reply_to_comment(doc, first.comment_id, "Still yes.")

    threads = read_threads(doc)
    assert len(threads) == 1
    assert [reply.comment_id for reply in threads[0].replies] == [
        first.comment_id,
        nested.comment_id,
    ]


def test_read_threads_returns_one_thread_per_independent_comment() -> None:
    doc = Document()
    add_comment(doc.add_paragraph("one"), "a")
    add_comment(doc.add_paragraph("two"), "b")

    threads = read_threads(doc)
    assert len(threads) == 2
    assert all(thread.replies == () for thread in threads)


def test_read_threads_on_document_without_comments_is_empty() -> None:
    assert read_threads(Document()) == []


def test_read_threads_without_extended_part_treats_every_comment_as_a_root() -> None:
    doc, root = _threaded_doc()
    reply_to_comment(doc, root.comment_id, "Yes.")
    # Simulate a producer that never wrote commentsExtended.xml.
    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype == RT_COMMENTS_EXTENDED:
            doc.part.drop_rel(rid)

    threads = read_threads(doc)
    assert len(threads) == 2
    assert all(thread.replies == () for thread in threads)
    assert all(thread.resolved is False for thread in threads)


def test_read_comments_exposes_parent_and_resolved_fields() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")
    resolve_comment(doc, root.comment_id)

    comments = {comment.comment_id: comment for comment in read_comments(doc)}
    assert comments[root.comment_id].parent_id is None
    assert comments[reply.comment_id].parent_id == root.comment_id
    assert comments[root.comment_id].resolved is True


def test_read_comments_defaults_are_root_and_unresolved() -> None:
    doc = Document()
    add_comment(doc.add_paragraph("text"), "note")
    comment = read_comments(doc)[0]
    assert comment.parent_id is None
    assert comment.resolved is False


# --------------------------------------------------------------------------
# Round-trip through the package.
# --------------------------------------------------------------------------


def test_thread_survives_a_save_and_reload() -> None:
    doc, root = _threaded_doc()
    reply_to_comment(doc, root.comment_id, "Yes.", author="Author")
    reply_to_comment(doc, root.comment_id, "Agreed.", author="Third")
    resolve_comment(doc, root.comment_id)

    reloaded = _round_trip(doc)
    threads = read_threads(reloaded)
    assert len(threads) == 1
    assert threads[0].resolved is True
    assert len(threads[0].replies) == 2
    assert threads[0].root.anchored_text == "Hello world"
    assert all(reply.anchored_text == "Hello world" for reply in threads[0].replies)


def test_extended_part_deserializes_as_xml_not_a_blob() -> None:
    doc, root = _threaded_doc()
    reply_to_comment(doc, root.comment_id, "Yes.")
    reloaded = _round_trip(doc)

    part = reloaded.part.part_related_by(RT_COMMENTS_EXTENDED)
    assert part.element is not None
    assert part.element.tag == qn("w15:commentsEx")


def test_replying_after_a_reload_extends_the_existing_thread() -> None:
    doc, root = _threaded_doc()
    reply_to_comment(doc, root.comment_id, "Yes.")
    reloaded = _round_trip(doc)

    ids = {comment.comment_id for comment in read_comments(reloaded)}
    root_id = next(
        comment.comment_id for comment in read_comments(reloaded) if comment.parent_id is None
    )
    third = reply_to_comment(reloaded, root_id, "One more.")
    assert third.comment_id not in ids
    assert len(read_threads(reloaded)[0].replies) == 2
    assert_para_ids_unique(reloaded)


# --------------------------------------------------------------------------
# Interaction with the shipped edit / delete / clear surface.
# --------------------------------------------------------------------------


def test_edit_comment_preserves_the_thread_link() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")
    resolve_comment(doc, root.comment_id)

    edit_comment(doc, root.comment_id, "Rewritten question.")

    threads = read_threads(doc)
    assert len(threads) == 1
    assert threads[0].root.text == "Rewritten question."
    assert threads[0].resolved is True
    assert [r.comment_id for r in threads[0].replies] == [reply.comment_id]


def test_edit_comment_without_a_para_id_leaves_it_unstamped() -> None:
    doc, root = _threaded_doc()
    paragraph = root.body_element.find(qn("w:p"))
    del paragraph.attrib[qn("w14:paraId")]

    edit_comment(doc, root.comment_id, "Rewritten.")
    assert root.body_element.find(qn("w:p")).get(qn("w14:paraId")) is None


def test_delete_comment_removes_its_thread_entry() -> None:
    doc, root = _threaded_doc()
    delete_comment(doc, root.comment_id)
    assert _ex_entries(doc) == []


def test_delete_root_cascades_to_replies_by_default() -> None:
    doc, root = _threaded_doc()
    first = reply_to_comment(doc, root.comment_id, "Yes.")
    second = reply_to_comment(doc, first.comment_id, "Still yes.")

    delete_comment(doc, root.comment_id)
    assert read_comments(doc) == []
    assert _ex_entries(doc) == []
    for comment_id in (root.comment_id, first.comment_id, second.comment_id):
        assert _markers(doc, comment_id) == (0, 0, 0)


def test_delete_root_without_cascade_promotes_replies_to_roots() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")

    delete_comment(doc, root.comment_id, include_replies=False)

    remaining = read_comments(doc)
    assert [comment.comment_id for comment in remaining] == [reply.comment_id]
    # The reply's paraIdParent now dangles; a dangling parent reads as a root.
    assert remaining[0].parent_id is None
    assert len(read_threads(doc)) == 1


def test_delete_reply_leaves_the_root_thread_intact() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")

    delete_comment(doc, reply.comment_id)
    threads = read_threads(doc)
    assert len(threads) == 1
    assert threads[0].root.comment_id == root.comment_id
    assert threads[0].replies == ()


def test_clear_all_comments_empties_the_extended_part() -> None:
    doc, root = _threaded_doc()
    reply_to_comment(doc, root.comment_id, "Yes.")

    clear_all_comments(doc)
    assert _ex_entries(doc) == []
    assert read_threads(doc) == []
    assert doc.part.part_related_by(RT_COMMENTS_EXTENDED) is not None


def test_clear_all_comments_with_remove_part_drops_the_extended_part() -> None:
    doc, root = _threaded_doc()
    reply_to_comment(doc, root.comment_id, "Yes.")

    clear_all_comments(doc, remove_part=True)
    with pytest.raises(KeyError):
        doc.part.part_related_by(RT_COMMENTS_EXTENDED)


def test_clear_all_comments_on_a_clean_document_is_a_no_op() -> None:
    doc = Document()
    clear_all_comments(doc)
    clear_all_comments(doc, remove_part=True)
    assert read_threads(doc) == []


# --------------------------------------------------------------------------
# _extended internals — tolerance for foreign / malformed input.
# --------------------------------------------------------------------------


def test_stamp_para_ids_does_not_overwrite_existing_values() -> None:
    doc, root = _threaded_doc()
    paragraph = root.body_element.find(qn("w:p"))
    original = paragraph.get(qn("w14:paraId"))

    _extended.stamp_para_ids(root.body_element, ParaIdRegistry(doc))
    assert paragraph.get(qn("w14:paraId")) == original


def test_stamp_para_ids_on_a_bodyless_comment_returns_empty() -> None:
    doc, root = _threaded_doc()
    for child in list(root.body_element):
        root.body_element.remove(child)
    assert _extended.stamp_para_ids(root.body_element, ParaIdRegistry(doc)) == ""
    assert _extended.thread_key(root.body_element) is None


def test_stamp_para_ids_keys_off_the_last_paragraph() -> None:
    doc, root = _threaded_doc()
    from docx_plus.core.oxml import sub

    sub(root.body_element, "w:p")
    key = _extended.stamp_para_ids(root.body_element, ParaIdRegistry(doc))
    paragraphs = xpath(root.body_element, ".//w:p")
    assert key == paragraphs[-1].get(qn("w14:paraId"))
    assert key != paragraphs[0].get(qn("w14:paraId"))


def test_upsert_with_empty_para_id_is_a_no_op() -> None:
    doc = Document()
    assert _extended.upsert_comment_ex(doc, "", done=True) is None
    with pytest.raises(KeyError):
        doc.part.part_related_by(RT_COMMENTS_EXTENDED)


def test_upsert_can_clear_parentage() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")
    keys, _ = _extended.key_maps(doc)

    _extended.upsert_comment_ex(doc, keys[reply.comment_id], parent_para_id="")
    assert _extended.thread_state(doc)[reply.comment_id][0] is None
    assert len(read_threads(doc)) == 2


def test_upsert_leaves_done_alone_when_not_specified() -> None:
    doc, root = _threaded_doc()
    resolve_comment(doc, root.comment_id)
    keys, _ = _extended.key_maps(doc)

    _extended.upsert_comment_ex(doc, keys[root.comment_id])
    assert _extended.thread_state(doc)[root.comment_id][1] is True


def test_drop_comment_ex_is_idempotent_and_part_safe() -> None:
    doc = Document()
    _extended.drop_comment_ex(doc, "DEADBEEF")  # no part at all
    doc, root = _threaded_doc()
    keys, _ = _extended.key_maps(doc)
    _extended.drop_comment_ex(doc, keys[root.comment_id])
    _extended.drop_comment_ex(doc, keys[root.comment_id])
    assert _ex_entries(doc) == []


def test_done_accepts_every_st_on_off_spelling() -> None:
    doc, root = _threaded_doc()
    entry = _ex_entries(doc)[0]
    for spelling, expected in (
        ("1", True),
        ("true", True),
        ("True", True),
        ("on", True),
        ("0", False),
        ("false", False),
        ("off", False),
    ):
        entry.set(qn("w15:done"), spelling)
        assert _extended.thread_state(doc)[root.comment_id][1] is expected


def test_missing_done_attribute_reads_as_unresolved() -> None:
    doc, root = _threaded_doc()
    entry = _ex_entries(doc)[0]
    del entry.attrib[qn("w15:done")]
    assert _extended.thread_state(doc)[root.comment_id][1] is False


def test_dangling_parent_key_reads_as_a_root() -> None:
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")
    _entry_for(doc, reply.comment_id).set(qn("w15:paraIdParent"), "FFFFFFFE")

    assert _extended.thread_state(doc)[reply.comment_id][0] is None
    assert len(read_threads(doc)) == 2


def test_self_parented_entry_reads_as_a_root() -> None:
    doc, root = _threaded_doc()
    keys, _ = _extended.key_maps(doc)
    _entry_for(doc, root.comment_id).set(qn("w15:paraIdParent"), keys[root.comment_id])

    assert _extended.thread_state(doc)[root.comment_id][0] is None
    assert len(read_threads(doc)) == 1


def test_parentage_cycle_terminates() -> None:
    # Two comments pointing at each other: malformed, but it must not spin.
    doc, root = _threaded_doc()
    reply = reply_to_comment(doc, root.comment_id, "Yes.")
    keys, _ = _extended.key_maps(doc)
    _entry_for(doc, root.comment_id).set(qn("w15:paraIdParent"), keys[reply.comment_id])

    state = _extended.thread_state(doc)
    assert _extended.root_id_of(state, reply.comment_id) in (root.comment_id, reply.comment_id)
    assert _extended.descendant_ids(state, root.comment_id) == [reply.comment_id]


def test_comment_elements_skips_unparseable_ids() -> None:
    doc, root = _threaded_doc()
    root_element = doc.part.part_related_by(RT.COMMENTS).element
    comment = xpath(root_element, "./w:comment")[0]
    comment.set(qn("w:id"), "not-a-number")
    assert _extended.comment_elements(doc) == {}


def test_comment_elements_skips_comments_with_no_id_at_all() -> None:
    doc, root = _threaded_doc()
    root_element = doc.part.part_related_by(RT.COMMENTS).element
    comment = xpath(root_element, "./w:comment")[0]
    del comment.attrib[qn("w:id")]
    assert _extended.comment_elements(doc) == {}


def test_comment_elements_without_a_comments_part_is_empty() -> None:
    assert _extended.comment_elements(Document()) == {}


def test_key_maps_skip_unstamped_comments() -> None:
    doc, root = _threaded_doc()
    paragraph = root.body_element.find(qn("w:p"))
    del paragraph.attrib[qn("w14:paraId")]

    by_id, by_key = _extended.key_maps(doc)
    assert by_id == {}
    assert by_key == {}
    # Still reported, just as an unresolved root.
    assert _extended.thread_state(doc)[root.comment_id] == (None, False)


# --------------------------------------------------------------------------
# ParaIdRegistry seeding.
# --------------------------------------------------------------------------


def test_para_id_registry_seeds_from_body_paragraphs() -> None:
    doc = Document()
    doc.add_paragraph("text")._p.set(qn("w14:paraId"), "0000002A")
    assert 0x2A in ParaIdRegistry(doc).issued()


def test_para_id_registry_seeds_from_the_comments_part() -> None:
    doc, root = _threaded_doc()
    existing = int(root.body_element.find(qn("w:p")).get(qn("w14:paraId")), 16)
    assert existing in ParaIdRegistry(doc).issued()


def test_para_id_registry_ignores_unparseable_values() -> None:
    doc = Document()
    doc.add_paragraph("text")._p.set(qn("w14:paraId"), "ZZZZZZZZ")
    assert ParaIdRegistry(doc).issued() == frozenset()


def test_next_hex_is_eight_uppercase_hex_digits() -> None:
    value = ParaIdRegistry(Document()).next_hex()
    assert len(value) == 8
    assert value == value.upper()
    assert 0 < int(value, 16) < 2**31


def test_next_hex_never_repeats() -> None:
    registry = ParaIdRegistry(Document())
    values = [registry.next_hex() for _ in range(200)]
    assert len(set(values)) == 200
