"""Tests for the built-in lint rules.

Each rule gets a positive case and the negative case most likely to make it
misfire, since a linter that cries wolf is worse than no linter. Rules are
exercised through :func:`~docx_plus.lint.lint` with an explicit ``select``,
so what is under test is the rule as a user actually invokes it.
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Twips
from docx.text.paragraph import Paragraph

from docx_plus.core.ns import qn
from docx_plus.core.oxml import build_complex_field, sub
from docx_plus.lint import Finding, LintContext, lint
from docx_plus.styles import iter_resolved_paragraphs, resolve_effective_formatting


def _rules_fired(doc: Document, rule_id: str) -> list[Finding]:
    return lint(doc, select=[rule_id])


# --------------------------------------------------------------------------
# Typography.
# --------------------------------------------------------------------------


def test_double_space_fires_between_words() -> None:
    doc = Document()
    doc.add_paragraph("Two  spaces between these words.")

    assert len(_rules_fired(doc, "double-space")) == 1


def test_double_space_ignores_single_spacing() -> None:
    doc = Document()
    doc.add_paragraph("Ordinary single spacing throughout.")

    assert _rules_fired(doc, "double-space") == []


def test_double_space_skips_verbatim_styles() -> None:
    """Runs of spaces in preformatted text are content, not sloppiness."""
    doc = Document()
    para = doc.add_paragraph("indent    matters here")
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "HTMLPreformatted"})

    assert _rules_fired(doc, "double-space") == []


def test_trailing_whitespace_fires() -> None:
    doc = Document()
    doc.add_paragraph("Ends with a space ")

    assert len(_rules_fired(doc, "trailing-whitespace")) == 1


def test_trailing_whitespace_ignores_empty_paragraphs() -> None:
    """An empty paragraph is not 'trailing whitespace' — that is a different rule."""
    doc = Document()
    doc.add_paragraph("   ")

    assert _rules_fired(doc, "trailing-whitespace") == []


def test_space_before_punctuation_fires() -> None:
    doc = Document()
    doc.add_paragraph("A sentence ends oddly .")

    findings = _rules_fired(doc, "space-before-punctuation")

    assert len(findings) == 1
    assert findings[0].expected == "."


def test_space_before_punctuation_ignores_normal_text() -> None:
    doc = Document()
    doc.add_paragraph("Commas, colons: and semicolons; all fine.")

    assert _rules_fired(doc, "space-before-punctuation") == []


def test_indent_by_whitespace_fires_on_tab() -> None:
    doc = Document()
    doc.add_paragraph("\tIndented with a tab")

    assert len(_rules_fired(doc, "indent-by-whitespace")) == 1


def test_indent_by_whitespace_fires_on_multiple_spaces() -> None:
    doc = Document()
    doc.add_paragraph("    Indented with spaces")

    assert len(_rules_fired(doc, "indent-by-whitespace")) == 1


def test_indent_by_whitespace_tolerates_a_single_leading_space() -> None:
    """One stray leading space is a typo, not a layout decision."""
    doc = Document()
    doc.add_paragraph(" Just one space")

    assert _rules_fired(doc, "indent-by-whitespace") == []


def test_stray_empty_paragraph_fires_on_a_run() -> None:
    doc = Document()
    doc.add_paragraph("Before")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("After")

    findings = _rules_fired(doc, "stray-empty-paragraph")

    assert len(findings) == 1
    assert findings[0].adds_content is True


def test_stray_empty_paragraph_ignores_a_single_blank() -> None:
    """One blank paragraph is common and usually deliberate."""
    doc = Document()
    doc.add_paragraph("Before")
    doc.add_paragraph("")
    doc.add_paragraph("After")

    assert _rules_fired(doc, "stray-empty-paragraph") == []


def test_stray_empty_paragraph_is_off_by_default() -> None:
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("")

    assert [f.rule for f in lint(doc)] == []


def test_stray_empty_paragraph_fires_at_document_end() -> None:
    """A trailing run of blanks is reported — the walk must flush at the end."""
    doc = Document()
    doc.add_paragraph("Content")
    doc.add_paragraph("")
    doc.add_paragraph("")

    assert len(_rules_fired(doc, "stray-empty-paragraph")) == 1


# --------------------------------------------------------------------------
# Structure.
# --------------------------------------------------------------------------


def test_heading_level_skip_fires() -> None:
    doc = Document()
    doc.add_paragraph("Top", style="Heading 1")
    doc.add_paragraph("Too deep", style="Heading 3")

    findings = _rules_fired(doc, "heading-level-skip")

    assert len(findings) == 1
    assert findings[0].observed == "level 3"
    assert findings[0].expected == "level 2"


def test_heading_level_skip_allows_descending_freely() -> None:
    """Coming back *up* the outline by any amount is fine."""
    doc = Document()
    doc.add_paragraph("One", style="Heading 1")
    doc.add_paragraph("Two", style="Heading 2")
    doc.add_paragraph("Three", style="Heading 3")
    doc.add_paragraph("Back to top", style="Heading 1")

    assert _rules_fired(doc, "heading-level-skip") == []


def test_heading_level_skip_ignores_body_paragraphs_between() -> None:
    """Body text between headings does not reset the outline expectation."""
    doc = Document()
    doc.add_paragraph("One", style="Heading 1")
    doc.add_paragraph("Body text.")
    doc.add_paragraph("Two", style="Heading 2")

    assert _rules_fired(doc, "heading-level-skip") == []


def test_empty_heading_fires() -> None:
    doc = Document()
    doc.add_paragraph("", style="Heading 2")

    assert len(_rules_fired(doc, "empty-heading")) == 1


def test_empty_heading_ignores_empty_body_paragraphs() -> None:
    doc = Document()
    doc.add_paragraph("")

    assert _rules_fired(doc, "empty-heading") == []


def test_manual_list_fires_on_typed_numbers() -> None:
    doc = Document()
    doc.add_paragraph("1. First item")
    doc.add_paragraph("2. Second item")

    assert len(_rules_fired(doc, "manual-list")) == 2


def test_manual_list_fires_on_typed_bullets() -> None:
    doc = Document()
    doc.add_paragraph("- A bullet")

    assert len(_rules_fired(doc, "manual-list")) == 1


def test_manual_list_ignores_a_real_list() -> None:
    """The whole point of resolving numbering through the style chain.

    Before that fix a correctly-styled ``List Bullet`` paragraph reported
    ``num_id=None`` and would have been flagged as hand-typed.
    """
    doc = Document()
    para = doc.add_paragraph("Genuinely bulleted", style="List Bullet")
    assert para.text

    assert _rules_fired(doc, "manual-list") == []


def test_manual_list_ignores_decimals_in_prose() -> None:
    """'1.5x faster' is not a list item — the marker needs trailing space."""
    doc = Document()
    doc.add_paragraph("1.5x faster than before")

    assert _rules_fired(doc, "manual-list") == []


def test_manual_list_respects_the_suppression_sentinel() -> None:
    """numId=0 means 'deliberately not numbered', so a typed marker still fires."""
    doc = Document()
    para = doc.add_paragraph("1. Opted out of numbering")
    num_pr = sub(sub(para._p, "w:pPr"), "w:numPr")
    sub(num_pr, "w:numId", **{"w:val": "0"})

    assert len(_rules_fired(doc, "manual-list")) == 1


# --------------------------------------------------------------------------
# Direct formatting.
# --------------------------------------------------------------------------


def test_redundant_direct_formatting_fires_on_inherited_value() -> None:
    """A run setting the size it would have inherited anyway."""
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.font.size = Pt(11)  # the bundled template's docDefaults size

    findings = _rules_fired(doc, "redundant-direct-formatting")

    assert len(findings) == 1
    assert findings[0].location.run_index == 0
    assert "size" in findings[0].message


def test_redundant_direct_formatting_ignores_a_real_override() -> None:
    """A run genuinely changing the size is not redundant."""
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.font.size = Pt(24)

    assert _rules_fired(doc, "redundant-direct-formatting") == []


def test_redundant_direct_formatting_ignores_undecorated_runs() -> None:
    doc = Document()
    doc.add_paragraph().add_run("plain text")

    assert _rules_fired(doc, "redundant-direct-formatting") == []


def test_redundant_direct_formatting_checks_runs_with_a_character_style() -> None:
    """A run carrying w:rStyle is checked against a baseline that includes it.

    ``stop_below="directRun"`` drops only the run's own ``rPr``, so the
    character style is still in force in the baseline. A direct size that
    merely restates what the cascade gives is therefore still redundant even
    on a styled run — the case the earlier paragraph-level baseline had to
    skip entirely.
    """
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.font.size = Pt(11)  # docDefaults already says 11pt
    sub(run._r.get_or_add_rPr(), "w:rStyle", **{"w:val": "Emphasis"})

    findings = _rules_fired(doc, "redundant-direct-formatting")

    assert len(findings) == 1
    assert "size" in findings[0].message


def test_redundant_direct_formatting_flags_a_restated_character_style_value() -> None:
    """Restating what the character style already supplies changes nothing.

    ``Emphasis`` resolves ``italic=True``, and direct formatting *states* a
    toggle rather than flipping it, so a direct ``<w:i/>`` on top renders
    identically to no direct formatting at all — genuinely removable.
    """
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.italic = True
    sub(run._r.get_or_add_rPr(), "w:rStyle", **{"w:val": "Emphasis"})

    assert any("italic" in f.message for f in _rules_fired(doc, "redundant-direct-formatting"))


def test_redundant_direct_formatting_names_every_redundant_property() -> None:
    """One finding per run, listing every property that changed nothing."""
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("text")
    run.font.size = Pt(11)
    run.bold = False

    findings = _rules_fired(doc, "redundant-direct-formatting")

    assert len(findings) == 1
    assert "size" in findings[0].message
    assert "bold" in findings[0].message


def test_redundant_direct_formatting_treats_toggle_off_as_redundant() -> None:
    """An explicit ``<w:b w:val="0"/>`` over an unset toggle renders identically.

    This is what select-all-then-clear-formatting leaves behind: the
    resolved values differ (``False`` vs ``None``) but the picture does
    not, so it is redundant.
    """
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.bold = False

    findings = _rules_fired(doc, "redundant-direct-formatting")

    assert len(findings) == 1
    assert "bold" in findings[0].message


def test_redundant_direct_formatting_keeps_a_toggle_that_turns_something_off() -> None:
    """Turning a toggle off *against* a style that turns it on is a real override.

    The bundled template's ``Heading 1`` resolves ``bold=True``, so an
    explicit "not bold" on it genuinely changes the picture and must not be
    normalised away as redundant.
    """
    doc = Document()
    para = doc.add_paragraph("Heading text", style="Heading 1")
    para.runs[0].bold = False

    findings = _rules_fired(doc, "redundant-direct-formatting")

    assert all("bold" not in f.message for f in findings)


def test_redundant_direct_formatting_flags_bold_restated_on_a_bold_style() -> None:
    """Direct bold over an already-bold ``Heading 1`` renders identically.

    Direct formatting is absolute rather than a flip, so the run is bold
    either way and the ``<w:b/>`` is safe to drop. The complementary case —
    an explicit *off* on the same style — is a real override, covered by
    :func:`test_redundant_direct_formatting_keeps_a_toggle_that_turns_something_off`.
    """
    doc = Document()
    para = doc.add_paragraph("Heading text", style="Heading 1")
    para.runs[0].bold = True

    assert any("bold" in f.message for f in _rules_fired(doc, "redundant-direct-formatting"))


def test_mixed_run_formatting_fires_on_disagreeing_sizes() -> None:
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("normal ").font.size = Pt(11)
    para.add_run("larger").font.size = Pt(14)

    assert len(_rules_fired(doc, "mixed-run-formatting")) == 1


def test_mixed_run_formatting_ignores_uniform_paragraphs() -> None:
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("one ")
    para.add_run("two")

    assert _rules_fired(doc, "mixed-run-formatting") == []


def test_mixed_run_formatting_ignores_whitespace_only_runs() -> None:
    """A whitespace-only run carrying a stray size is not a visible inconsistency."""
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("text ").font.size = Pt(11)
    para.add_run("   ").font.size = Pt(30)

    assert _rules_fired(doc, "mixed-run-formatting") == []


def test_mixed_run_formatting_is_off_by_default() -> None:
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("a").font.size = Pt(11)
    para.add_run("b").font.size = Pt(14)

    assert "mixed-run-formatting" not in {f.rule for f in lint(doc)}


# --------------------------------------------------------------------------
# Cross-cutting.
# --------------------------------------------------------------------------


def test_rules_do_not_fire_on_a_stock_document() -> None:
    """A default python-docx document with plain prose is clean under every rule.

    The broadest false-positive guard there is: if any rule fires here,
    every real document will be noisy.
    """
    doc = Document()
    doc.add_paragraph("Title", style="Heading 1")
    doc.add_paragraph("A paragraph of ordinary prose, with a comma and a full stop.")
    doc.add_paragraph("Subheading", style="Heading 2")
    doc.add_paragraph("More prose; nothing unusual here.")
    doc.add_paragraph("Bulleted item", style="List Bullet")

    assert lint(doc, select=[r for r in ("typography", "structure", "formatting", "lists")]) == []


def test_findings_reference_only_registered_rule_ids() -> None:
    """Nothing can emit a finding under an id that is not registered."""
    from docx_plus.lint import all_rules

    doc = Document()
    doc.add_paragraph("Messy  text .")
    doc.add_paragraph("Deep", style="Heading 3")

    known = {r.id for r in all_rules()}

    assert {f.rule for f in lint(doc)} <= known


def test_paragraph_mark_formatting_is_not_the_run_baseline() -> None:
    """A run's baseline excludes ``pPr/rPr``, which formats the mark, not the runs.

    Word applies paragraph-mark formatting to the pilcrow alone, so a run
    matching it is a genuine override rather than a redundant restatement.
    The run baseline resolves with ``stop_below="directRun"``, which for a
    run target never applies the mark's ``rPr`` — this guards that.
    """
    doc = Document()
    para = doc.add_paragraph()
    ppr_rpr = sub(sub(para._p, "w:pPr"), "w:rPr")
    sub(ppr_rpr, "w:sz", **{"w:val": "36"})  # 18pt in half-points
    run = para.add_run("text")
    run.font.size = Pt(18)

    assert _rules_fired(doc, "redundant-direct-formatting") == []


def test_verbatim_style_detection_uses_resolved_style_id() -> None:
    """The skip is keyed off the resolved style id, not the raw pStyle text."""
    doc = Document()
    para = doc.add_paragraph("code    block")
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "PlainText"})

    assert _rules_fired(doc, "double-space") == []
    assert para._p.find(f"./{qn('w:pPr')}/{qn('w:pStyle')}") is not None


# --------------------------------------------------------------------------
# style-drift.
# --------------------------------------------------------------------------


def test_style_drift_fires_on_a_direct_paragraph_override() -> None:
    doc = Document()
    para = doc.add_paragraph("text", style="Heading 1")
    para.paragraph_format.space_after = Pt(48)

    findings = _rules_fired(doc, "style-drift")

    assert len(findings) == 1
    assert "space after" in findings[0].message
    assert findings[0].location.style_id == "Heading1"


def test_style_drift_reports_both_sides() -> None:
    """The finding must say what it found *and* what the style says."""
    doc = Document()
    para = doc.add_paragraph("text")
    para.paragraph_format.left_indent = Pt(36)

    finding = _rules_fired(doc, "style-drift")[0]

    assert finding.observed is not None
    assert finding.expected is not None
    assert finding.observed != finding.expected


def test_style_drift_ignores_an_undecorated_paragraph() -> None:
    doc = Document()
    doc.add_paragraph("text", style="Heading 1")

    assert _rules_fired(doc, "style-drift") == []


def test_style_drift_ignores_a_direct_value_matching_the_style() -> None:
    """That is `redundant-direct-formatting`'s finding, not this one.

    The two rules split the same comparison, so a property must never
    produce both findings.
    """
    doc = Document()
    para = doc.add_paragraph("text")
    inherited = resolve_effective_formatting(para).spacing_after
    assert inherited is not None
    para.paragraph_format.space_after = Twips(inherited)

    # It really is direct now, and it really does match what it replaced.
    resolved = resolve_effective_formatting(para, include_provenance=True)
    assert resolved.provenance is not None
    assert resolved.provenance["spacing_after"].layer == "directParagraph"
    assert resolved.spacing_after == inherited

    assert _rules_fired(doc, "style-drift") == []
    assert len(_rules_fired(doc, "redundant-direct-formatting")) == 0  # run-scoped rule


def test_style_drift_ignores_run_level_overrides() -> None:
    """Run-level drift is nearly always deliberate, so it is out of scope."""
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("emphasis")
    run.bold = True
    run.font.size = Pt(24)

    assert _rules_fired(doc, "style-drift") == []


def test_style_drift_does_not_blame_the_numbering_level() -> None:
    """A numbered paragraph's indent comes from the level, not a direct override.

    This is the case a two-layer COM compare gets wrong: the effective
    indent differs from the style's, so it reads as drift, when in fact no
    direct formatting is involved at all.
    """
    doc = Document()
    doc.add_paragraph("item", style="List Bullet")

    assert _rules_fired(doc, "style-drift") == []


def test_style_drift_names_every_drifted_property_in_one_finding() -> None:
    doc = Document()
    para = doc.add_paragraph("text")
    para.paragraph_format.left_indent = Pt(36)
    para.paragraph_format.space_before = Pt(24)

    findings = _rules_fired(doc, "style-drift")

    assert len(findings) == 1
    assert "left indent" in findings[0].message
    assert "space before" in findings[0].message


def test_style_drift_names_the_style_it_deviates_from() -> None:
    doc = Document()
    para = doc.add_paragraph("text", style="Heading 2")
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    assert "heading 2" in _rules_fired(doc, "style-drift")[0].message


def test_lint_context_resolve_reaches_other_layers() -> None:
    """The escape hatch for a rule wanting a slice the sweep did not precompute."""
    doc = Document()
    para = doc.add_paragraph("item")
    num_pr = sub(sub(para._p, "w:pPr"), "w:numPr")
    sub(num_pr, "w:numId", **{"w:val": "7"})

    ctx = LintContext(doc=doc, paragraphs=list(iter_resolved_paragraphs(doc)))

    assert ctx.resolve(para).num_id == 7
    assert ctx.resolve(para, stop_below="numbering").num_id is None


# --------------------------------------------------------------------------
# direct-numbering-override.
# --------------------------------------------------------------------------


def _override_num_id(para: Paragraph, num_id: str) -> None:
    """Point a paragraph's own w:numPr at ``num_id``."""
    num_pr = sub(sub(para._p, "w:pPr"), "w:numPr")
    sub(num_pr, "w:numId", **{"w:val": num_id})


def test_direct_numbering_override_fires() -> None:
    doc = Document()
    para = doc.add_paragraph("Item", style="List Number")
    _override_num_id(para, "42")

    findings = _rules_fired(doc, "direct-numbering-override")

    assert len(findings) == 1
    assert findings[0].observed == "numId=42"
    assert "from the style" in (findings[0].expected or "")


def test_direct_numbering_override_ignores_a_style_supplied_list() -> None:
    """The rule must not flag every correctly-styled list paragraph."""
    doc = Document()
    doc.add_paragraph("Item", style="List Number")

    assert _rules_fired(doc, "direct-numbering-override") == []


def test_direct_numbering_override_ignores_a_list_with_no_style_numbering() -> None:
    """A direct numPr on an unnumbered style overrides nothing.

    This is the shape ``apply_list`` produces, so firing here would flag
    ordinary library output as a defect.
    """
    doc = Document()
    _override_num_id(doc.add_paragraph("Item"), "3")

    assert _rules_fired(doc, "direct-numbering-override") == []


def test_direct_numbering_override_downgrades_the_opt_out_sentinel() -> None:
    """numId=0 is the one legitimate override, so it reports as info."""
    doc = Document()
    para = doc.add_paragraph("Not a list item after all", style="List Number")
    _override_num_id(para, "0")

    findings = _rules_fired(doc, "direct-numbering-override")

    assert len(findings) == 1
    assert findings[0].severity == "info"
    assert "suppresses" in findings[0].message


# --------------------------------------------------------------------------
# list-numbering-continuity.
# --------------------------------------------------------------------------


def test_list_numbering_continuity_fires_on_a_split_run() -> None:
    doc = Document()
    doc.add_paragraph("One", style="List Number")
    _override_num_id(doc.add_paragraph("Two", style="List Number"), "77")

    assert len(_rules_fired(doc, "list-numbering-continuity")) == 1


def test_list_numbering_continuity_ignores_one_unbroken_list() -> None:
    doc = Document()
    for text in ("One", "Two", "Three"):
        doc.add_paragraph(text, style="List Number")

    assert _rules_fired(doc, "list-numbering-continuity") == []


def test_list_numbering_continuity_ignores_lists_separated_by_body_text() -> None:
    """A restart after prose is deliberate often enough not to call it."""
    doc = Document()
    doc.add_paragraph("One", style="List Number")
    doc.add_paragraph("Some intervening prose.")
    _override_num_id(doc.add_paragraph("One again", style="List Number"), "77")

    assert _rules_fired(doc, "list-numbering-continuity") == []


def test_list_numbering_continuity_ignores_a_sublist() -> None:
    """A different level is legitimately a different list."""
    doc = Document()
    doc.add_paragraph("One", style="List Number")
    para = doc.add_paragraph("Sub-item", style="List Number")
    num_pr = sub(sub(para._p, "w:pPr"), "w:numPr")
    sub(num_pr, "w:ilvl", **{"w:val": "1"})
    sub(num_pr, "w:numId", **{"w:val": "77"})

    assert _rules_fired(doc, "list-numbering-continuity") == []


def test_list_numbering_continuity_ignores_the_opt_out_sentinel() -> None:
    """A suppressed paragraph is not a list item, so it breaks the run."""
    doc = Document()
    doc.add_paragraph("One", style="List Number")
    _override_num_id(doc.add_paragraph("Opted out", style="List Number"), "0")

    assert _rules_fired(doc, "list-numbering-continuity") == []


# --------------------------------------------------------------------------
# manual-heading-formatting.
# --------------------------------------------------------------------------


def _bold_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph().add_run(text).bold = True


def test_manual_heading_formatting_fires_on_a_bold_short_line() -> None:
    doc = Document()
    _bold_paragraph(doc, "Background and Scope")
    doc.add_paragraph("Ordinary body text follows here.")

    findings = _rules_fired(doc, "manual-heading-formatting")

    assert len(findings) == 1
    assert findings[0].expected == "a heading style"


def test_manual_heading_formatting_fires_on_an_enlarged_line() -> None:
    """The size comparison is against the document's own body text."""
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Section Two").font.size = Pt(16)
    for _ in range(3):
        doc.add_paragraph("Ordinary body text at the document's usual size.")

    assert len(_rules_fired(doc, "manual-heading-formatting")) == 1


def test_manual_heading_formatting_ignores_a_real_heading() -> None:
    doc = Document()
    doc.add_paragraph("A Real Heading", style="Heading 1")

    assert _rules_fired(doc, "manual-heading-formatting") == []


def test_manual_heading_formatting_ignores_a_bold_sentence() -> None:
    """Ending like a sentence is the strongest signal it is not a heading."""
    doc = Document()
    _bold_paragraph(doc, "This is emphasised prose, not a heading.")

    assert _rules_fired(doc, "manual-heading-formatting") == []


def test_manual_heading_formatting_ignores_long_bold_text() -> None:
    doc = Document()
    _bold_paragraph(doc, "A bold run of text far too long to be mistaken for a heading " * 2)

    assert _rules_fired(doc, "manual-heading-formatting") == []


def test_manual_heading_formatting_ignores_table_cells() -> None:
    """Table cells are full of short bold labels, and none are headings."""
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run("Total").bold = True

    assert _rules_fired(doc, "manual-heading-formatting") == []


def test_manual_heading_formatting_ignores_a_bold_list_item() -> None:
    doc = Document()
    para = doc.add_paragraph(style="List Bullet")
    para.add_run("A bold bullet").bold = True

    assert _rules_fired(doc, "manual-heading-formatting") == []


def test_manual_heading_formatting_ignores_partly_bold_text() -> None:
    """A bolded term inside a line is emphasis, not a heading."""
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("A line with a ")
    para.add_run("bold").bold = True
    para.add_run(" word")

    assert _rules_fired(doc, "manual-heading-formatting") == []


# --------------------------------------------------------------------------
# font-outliers.
# --------------------------------------------------------------------------


def _run_with_font(doc: Document, text: str, name: str, size: float) -> None:
    run = doc.add_paragraph().add_run(text)
    run.font.name = name
    run.font.size = Pt(size)


def test_font_outliers_fires_on_a_thin_combination() -> None:
    doc = Document()
    for index in range(40):
        _run_with_font(doc, f"body {index}", "Calibri", 11)
    _run_with_font(doc, "pasted in", "Times New Roman", 12)

    findings = _rules_fired(doc, "font-outliers")

    assert len(findings) == 1
    assert "Times New Roman" in (findings[0].observed or "")
    assert "Calibri" in (findings[0].expected or "")


def test_font_outliers_ignores_a_uniform_document() -> None:
    doc = Document()
    for index in range(20):
        _run_with_font(doc, f"body {index}", "Calibri", 11)

    assert _rules_fired(doc, "font-outliers") == []


def test_font_outliers_ignores_a_substantial_second_font() -> None:
    """A font used throughout is a design choice, not an outlier."""
    doc = Document()
    for index in range(20):
        _run_with_font(doc, f"body {index}", "Calibri", 11)
        _run_with_font(doc, f"code {index}", "Consolas", 10)

    assert _rules_fired(doc, "font-outliers") == []


def test_font_outliers_ignores_a_document_with_no_dominant_font() -> None:
    """Everything an outlier means there is nothing to be an outlier from."""
    doc = Document()
    for index in range(6):
        _run_with_font(doc, f"line {index}", f"Font{index}", 10 + index)

    assert _rules_fired(doc, "font-outliers") == []


def test_font_outliers_is_off_by_default() -> None:
    doc = Document()
    for index in range(40):
        _run_with_font(doc, f"body {index}", "Calibri", 11)
    _run_with_font(doc, "pasted in", "Times New Roman", 12)

    assert "font-outliers" not in {f.rule for f in lint(doc)}


# --------------------------------------------------------------------------
# False positives found by running the CLI against a realistic document.
# --------------------------------------------------------------------------


def test_trailing_whitespace_ignores_a_paragraph_containing_a_field() -> None:
    """ "See " before an unrendered REF is not trailing whitespace.

    A field contributes its *cached result*, which is empty in a freshly
    written document, so the space doing its job looks like sloppiness.
    Every cross-reference and page number would otherwise report.
    """
    doc = Document()
    para = doc.add_paragraph("See ")
    build_complex_field(para._p, " REF chapter1 ", "")

    assert _rules_fired(doc, "trailing-whitespace") == []


def test_trailing_whitespace_still_fires_without_a_field() -> None:
    """The field guard must not swallow the ordinary case."""
    doc = Document()
    doc.add_paragraph("See ")

    assert len(_rules_fired(doc, "trailing-whitespace")) == 1


def test_manual_heading_formatting_ignores_a_styled_paragraph() -> None:
    """A paragraph carrying any non-default style was styled deliberately.

    The template's ``Caption`` style is bold, so without this the rule
    reports every caption in the document.
    """
    doc = Document()
    para = doc.add_paragraph("Figure 1: A diagram")
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Caption"})

    assert _rules_fired(doc, "manual-heading-formatting") == []


def test_manual_heading_formatting_still_fires_on_the_default_style() -> None:
    """An explicit Normal is still unstyled body text."""
    doc = Document()
    para = doc.add_paragraph()
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Normal"})
    para.add_run("Background and Scope").bold = True

    assert len(_rules_fired(doc, "manual-heading-formatting")) == 1
