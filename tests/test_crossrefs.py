"""Tests for ``docx_plus.bookmarks.add_cross_reference``."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from docx_plus.bookmarks import add_bookmark, add_cross_reference
from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.fields import mark_fields_dirty


def _instr(paragraph) -> str:
    instrs = xpath(paragraph._p, "./w:r/w:instrText")
    assert len(instrs) == 1
    return instrs[0].text or ""


def _field_char_types(paragraph) -> list[str | None]:
    return [fc.get(qn("w:fldCharType")) for fc in xpath(paragraph._p, "./w:r/w:fldChar")]


def test_text_cross_ref_emits_ref_instruction() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_cross_reference(p, bookmark="sec_1", kind="text")
    assert _instr(p) == r" REF sec_1 \h "


def test_page_cross_ref_emits_pageref_instruction() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_cross_reference(p, bookmark="sec_1", kind="page")
    assert _instr(p) == r" PAGEREF sec_1 \h "


def test_no_hyperlink_drops_h_flag() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_cross_reference(p, bookmark="sec_1", hyperlink=False)
    assert _instr(p) == " REF sec_1 "


def test_cross_ref_emits_full_complex_field_sequence() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_cross_reference(p, bookmark="sec_1")
    assert _field_char_types(p) == ["begin", "separate", "end"]


def test_cross_ref_initial_text_is_empty() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_cross_reference(p, bookmark="sec_1")
    t = xpath(p._p, "./w:r/w:t")
    assert len(t) == 1
    assert (t[0].text or "") == ""


def test_cross_ref_resolves_after_round_trip(tmp_path: Path) -> None:
    """A bookmark + a cross-reference should both round-trip cleanly.

    We don't ask Word to actually resolve the cached result (that
    happens on open), but the field instruction and bookmark should
    survive the write/read cycle so Word's resolver has what it needs.
    """
    doc = Document()
    p_target = doc.add_paragraph()
    p_target.add_run("Section 1 intro")
    add_bookmark(p_target, "sec_1")

    p_ref = doc.add_paragraph("See ")
    add_cross_reference(p_ref, bookmark="sec_1")
    mark_fields_dirty(doc)

    out = tmp_path / "xref.docx"
    doc.save(str(out))
    reopened = Document(str(out))

    p2 = reopened.paragraphs[1]
    instr = xpath(p2._p, "./w:r/w:instrText")[0].text
    assert instr == r" REF sec_1 \h "

    starts = xpath(reopened.element.body, ".//w:bookmarkStart[@w:name=$n]", n="sec_1")
    assert len(starts) == 1
