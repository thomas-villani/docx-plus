"""Tests for the lint engine, registry, and selection semantics.

The rules themselves are tested in ``test_lint_rules.py``; this file covers
the machinery around them — how rules are registered, how selectors resolve,
and the contract that a rule cannot misreport its own metadata.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

import pytest
from docx import Document

from docx_plus.lint import (
    Finding,
    Issue,
    Location,
    UnknownRuleError,
    all_rules,
    lint,
    select_rules,
)
from docx_plus.lint.models import Rule

if TYPE_CHECKING:
    from collections.abc import Iterator

    from docx_plus.lint import LintContext


# --------------------------------------------------------------------------
# Registry & selection.
# --------------------------------------------------------------------------


def test_rules_are_registered() -> None:
    """The built-in rules self-register on import."""
    rules = all_rules()
    ids = {r.id for r in rules}

    assert "double-space" in ids
    assert "heading-level-skip" in ids
    assert rules == sorted(rules, key=lambda r: r.id)


def test_every_rule_has_complete_metadata() -> None:
    """Every rule carries the fields the report and CLI depend on."""
    for registered in all_rules():
        assert registered.id == registered.id.lower()
        assert " " not in registered.id, f"{registered.id} should be kebab-case"
        assert registered.kind in ("consistency", "structural", "policy")
        assert registered.severity in ("error", "warning", "info")
        assert registered.description.endswith("."), registered.id
        assert registered.tags, f"{registered.id} has no tags to select it by"


def test_policy_rules_are_never_default_on() -> None:
    """A policy rule has no target without a profile, so it must not run by default.

    This is the invariant that keeps the library from shipping an opinion.
    """
    for registered in all_rules():
        if registered.kind == "policy":
            assert not registered.default_on, f"{registered.id} imposes an opinion by default"


def test_select_none_runs_default_on_rules() -> None:
    """No selector runs exactly the default-on set."""
    chosen = select_rules()

    assert chosen
    assert all(r.default_on for r in chosen)
    assert {r.id for r in chosen} == {r.id for r in all_rules() if r.default_on}


def test_select_by_id() -> None:
    """A selector naming an id runs just that rule."""
    assert [r.id for r in select_rules(["double-space"])] == ["double-space"]


def test_select_by_tag_enables_off_by_default_rules() -> None:
    """Naming a tag opts into that cluster's off-by-default rules.

    This is the mechanism by which heuristic rules stay out of default
    output but remain one flag away.
    """
    off_by_default = {r.id for r in all_rules() if not r.default_on and "whitespace" in r.tags}
    assert off_by_default, "fixture assumption: some whitespace rule ships off"

    chosen = {r.id for r in select_rules(["whitespace"])}

    assert off_by_default <= chosen


def test_exclude_wins_over_select() -> None:
    """``exclude`` is applied last."""
    chosen = select_rules(["typography"], exclude=["double-space"])

    assert chosen
    assert "double-space" not in {r.id for r in chosen}


def test_exclude_by_tag() -> None:
    """Excluding a tag drops the whole cluster."""
    chosen = {r.id for r in select_rules(exclude=["typography"])}

    assert "double-space" not in chosen
    assert "heading-level-skip" in chosen


def test_unknown_selector_raises() -> None:
    """A typo selects nothing, which reads exactly like a clean document — so it raises."""
    with pytest.raises(UnknownRuleError, match="no-such-rule"):
        select_rules(["no-such-rule"])


def test_unknown_exclude_selector_raises() -> None:
    """The same guard applies to exclusions."""
    with pytest.raises(UnknownRuleError):
        select_rules(exclude=["no-such-rule"])


def test_unknown_rule_error_is_a_key_error() -> None:
    """It dual-inherits KeyError, per the project's error convention."""
    assert issubclass(UnknownRuleError, KeyError)


def test_duplicate_rule_id_is_rejected() -> None:
    """Registering the same id twice is a programming error, caught at import."""
    from docx_plus.lint.registry import rule

    with pytest.raises(ValueError, match="duplicate lint rule id"):

        @rule(
            id="double-space",
            kind="consistency",
            severity="info",
            description="Duplicate.",
            tags={"typography"},
        )
        def _clash(ctx: LintContext) -> Iterator[Issue]:  # pragma: no cover
            yield Issue(message="never runs")


def test_rule_matches_id_and_tags() -> None:
    """``matches`` accepts either an id or a tag."""
    registered = Rule(
        id="demo",
        kind="structural",
        severity="info",
        description="Demo.",
        check=lambda ctx: iter(()),
        tags=frozenset({"cluster"}),
    )

    assert registered.matches("demo")
    assert registered.matches("cluster")
    assert not registered.matches("other")


# --------------------------------------------------------------------------
# Engine.
# --------------------------------------------------------------------------


def test_lint_stamps_rule_metadata_onto_findings() -> None:
    """The rule's id / kind / severity come from registration, not the rule body.

    A rule yields an Issue carrying only what it knows; the engine promotes
    it. That is what stops a rule from advertising one severity in
    ``list_rules`` and emitting another.
    """
    doc = Document()
    doc.add_paragraph("Two  spaces.")

    findings = lint(doc, select=["double-space"])

    assert len(findings) == 1
    finding = findings[0]
    assert finding.rule == "double-space"
    assert finding.kind == "consistency"
    assert finding.severity == "info"


def test_lint_sorts_by_severity_then_document_order() -> None:
    """Findings come back worst-first, then in reading order."""
    doc = Document()
    doc.add_paragraph("Body  text.")  # info, paragraph 0
    doc.add_paragraph("Heading 1", style="Heading 1")
    doc.add_paragraph("Skipped", style="Heading 3")  # warning, paragraph 2

    findings = lint(doc)

    assert [f.severity for f in findings] == sorted(
        [f.severity for f in findings], key=lambda s: {"error": 0, "warning": 1, "info": 2}[s]
    )
    assert findings[0].severity == "warning"


def test_lint_is_read_only() -> None:
    """Linting mutates nothing — the document's XML is byte-identical after."""
    from lxml import etree

    doc = Document()
    doc.add_paragraph("Two  spaces and a trailing one ")
    doc.add_paragraph("Heading", style="Heading 1")

    before = etree.tostring(doc.element.body)
    lint(doc)
    after = etree.tostring(doc.element.body)

    assert before == after


def test_lint_clean_document_returns_nothing() -> None:
    """A tidy document produces no findings — no false positives on the trivial case."""
    doc = Document()
    doc.add_paragraph("A clean sentence.")
    doc.add_paragraph("Another one, with punctuation; fine.")

    assert lint(doc) == []


def test_lint_empty_document_returns_nothing() -> None:
    """A document with no content lints cleanly rather than erroring."""
    doc = Document()
    for para in list(doc.paragraphs):
        para._p.getparent().remove(para._p)

    assert lint(doc) == []


def test_lint_respects_include_tables() -> None:
    """Table-cell paragraphs are audited by default and skippable."""
    doc = Document()
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Inside  a cell."

    assert [f.rule for f in lint(doc)] == ["double-space"]
    assert lint(doc, include_tables=False) == []


def test_lint_findings_carry_a_usable_location() -> None:
    """Every positional finding points at a real paragraph and carries an excerpt."""
    doc = Document()
    doc.add_paragraph("Clean first paragraph.")
    doc.add_paragraph("Second  one has the defect.")

    findings = lint(doc, select=["double-space"])

    assert len(findings) == 1
    location = findings[0].location
    assert location.paragraph_index == 1
    assert "Second" in location.excerpt
    assert location.describe() == "paragraph 1"


def test_location_describe_handles_every_shape() -> None:
    """The human-readable position covers run, paragraph, style, and document scope."""
    assert Location(paragraph_index=3).describe() == "paragraph 3"
    assert Location(paragraph_index=3, run_index=1).describe() == "paragraph 3, run 1"
    assert Location(style_id="Heading1").describe() == "style Heading1"
    assert Location().describe() == "document"


def test_finding_sort_key_orders_unpositioned_first() -> None:
    """A style-level finding sorts ahead of positional ones at the same severity."""
    style_level = Finding(
        rule="r", kind="structural", severity="warning", message="m", location=Location()
    )
    positional = Finding(
        rule="r",
        kind="structural",
        severity="warning",
        message="m",
        location=Location(paragraph_index=0),
    )

    assert style_level.sort_key < positional.sort_key


def test_context_excerpt_collapses_whitespace_and_truncates() -> None:
    """Excerpts are one tidy line, short enough for a report column."""
    doc = Document()
    doc.add_paragraph("word " * 40)

    findings = lint(doc, select=["trailing-whitespace"])

    assert len(findings) == 1
    excerpt = findings[0].location.excerpt
    assert len(excerpt) <= 60
    assert "  " not in excerpt
