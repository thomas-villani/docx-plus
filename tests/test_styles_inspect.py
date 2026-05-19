"""Tests for ``docx_plus.styles.inspect.resolve_effective_formatting``.

Per-layer coverage: each cascade layer is exercised in isolation, then in
combination, to verify the precedence and override rules of SPEC §4.
"""

from __future__ import annotations

import dataclasses
from pathlib import Path

import pytest
from docx import Document

from docx_plus.core.oxml import sub
from docx_plus.styles.inspect import (
    ResolvedFormatting,
    StyleCascadeError,
    resolve_effective_formatting,
)

# --------------------------------------------------------------------------
# Layer 1: docDefaults
# --------------------------------------------------------------------------


def test_doc_defaults_provide_font_size() -> None:
    """A fresh Document() has rPrDefault sz=22 (= 11pt)."""
    doc = Document()
    p = doc.add_paragraph("text")
    resolved = resolve_effective_formatting(p)
    assert resolved.font_size == 11.0


def test_doc_defaults_provide_font_name_token() -> None:
    """Theme-font tokens pass through unresolved in Phase 2."""
    doc = Document()
    p = doc.add_paragraph("text")
    resolved = resolve_effective_formatting(p)
    # python-docx's default rPrDefault uses w:asciiTheme="minorHAnsi"
    assert resolved.font_name == "minorHAnsi"


# --------------------------------------------------------------------------
# Layer 3: paragraph style chain
# --------------------------------------------------------------------------


def test_paragraph_style_sets_style_id_and_name() -> None:
    doc = Document()
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "S1"})
    sub(s, "w:name", **{"w:val": "Style One"})
    rpr = sub(s, "w:rPr")
    sub(rpr, "w:sz", **{"w:val": "28"})

    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "S1"})

    resolved = resolve_effective_formatting(p)
    assert resolved.style_id == "S1"
    assert resolved.style_name == "Style One"
    assert resolved.font_size == 14.0  # 28 / 2 = 14pt; overrides docDefaults


def test_paragraph_style_overrides_doc_defaults() -> None:
    """A style's font_size wins over docDefaults."""
    doc = Document()
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "Big"})
    sub(s, "w:name", **{"w:val": "Big"})
    rpr = sub(s, "w:rPr")
    sub(rpr, "w:sz", **{"w:val": "48"})

    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "Big"})

    resolved = resolve_effective_formatting(p)
    assert resolved.font_size == 24.0  # Big style, not docDefaults' 11.0


def test_paragraph_style_chain_walks_basedon(multistyle_docx_path: Path) -> None:
    """multistyle.docx has Base -> Mid -> Top. Top's italic comes from Mid."""
    doc = Document(str(multistyle_docx_path))
    p = doc.paragraphs[0]
    resolved = resolve_effective_formatting(p)
    assert resolved.style_id == "Top"
    # Italic set at Mid (only place); should propagate through Top.
    assert resolved.italic is True


# --------------------------------------------------------------------------
# Layer 5: direct paragraph formatting
# --------------------------------------------------------------------------


def test_direct_paragraph_alignment() -> None:
    doc = Document()
    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:jc", **{"w:val": "center"})

    resolved = resolve_effective_formatting(p)
    assert resolved.alignment == "center"


def test_direct_paragraph_indent_left_and_first_line() -> None:
    doc = Document()
    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:ind", **{"w:left": "720", "w:firstLine": "360"})

    resolved = resolve_effective_formatting(p)
    assert resolved.indent_left == 720
    assert resolved.indent_first_line == 360


def test_direct_paragraph_hanging_indent_is_negative() -> None:
    doc = Document()
    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:ind", **{"w:left": "720", "w:hanging": "360"})

    resolved = resolve_effective_formatting(p)
    assert resolved.indent_first_line == -360


def test_direct_spacing_before_after() -> None:
    doc = Document()
    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:spacing", **{"w:before": "120", "w:after": "240"})

    resolved = resolve_effective_formatting(p)
    assert resolved.spacing_before == 120
    assert resolved.spacing_after == 240


def test_line_spacing_auto_rule() -> None:
    doc = Document()
    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:spacing", **{"w:line": "480", "w:lineRule": "auto"})

    resolved = resolve_effective_formatting(p)
    assert resolved.line_spacing == 2.0  # 480/240
    assert resolved.line_spacing_rule == "auto"


def test_line_spacing_exact_rule_keeps_twips() -> None:
    doc = Document()
    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:spacing", **{"w:line": "480", "w:lineRule": "exact"})

    resolved = resolve_effective_formatting(p)
    assert resolved.line_spacing == 480.0
    assert resolved.line_spacing_rule == "exact"


def test_paragraph_boolean_flags() -> None:
    doc = Document()
    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:keepNext")
    sub(ppr, "w:keepLines")
    sub(ppr, "w:pageBreakBefore")
    sub(ppr, "w:outlineLvl", **{"w:val": "2"})

    resolved = resolve_effective_formatting(p)
    assert resolved.keep_with_next is True
    assert resolved.keep_lines is True
    assert resolved.page_break_before is True
    assert resolved.outline_level == 2


# --------------------------------------------------------------------------
# Layer 6: direct run formatting (and run-level overrides)
# --------------------------------------------------------------------------


def test_direct_run_bold() -> None:
    doc = Document()
    p = doc.add_paragraph("Hello ")
    r = p.add_run("world")
    r.bold = True

    resolved_para = resolve_effective_formatting(p)
    resolved_run = resolve_effective_formatting(r)
    # Paragraph resolution ignores per-run rPr.
    assert resolved_para.bold is None
    # Run resolution picks up the direct run formatting.
    assert resolved_run.bold is True


def test_direct_run_overrides_style_font_size() -> None:
    doc = Document()
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "S"})
    sub(s, "w:name", **{"w:val": "S"})
    rpr = sub(s, "w:rPr")
    sub(rpr, "w:sz", **{"w:val": "20"})

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "S"})

    r = p.add_run("text")
    r.font.size = None  # ensure no python-docx-applied size
    # Apply a direct run sz manually
    r_rpr = r._r.get_or_add_rPr()
    sub(r_rpr, "w:sz", **{"w:val": "36"})

    resolved_run = resolve_effective_formatting(r)
    assert resolved_run.font_size == 18.0  # 36/2; overrides style's 10pt


# --------------------------------------------------------------------------
# Layer 6: theme-colored direct run formatting via fixture
# --------------------------------------------------------------------------


def test_theme_color_resolves_via_theme(themed_docx_path: Path) -> None:
    doc = Document(str(themed_docx_path))
    p = doc.paragraphs[0]
    resolved = resolve_effective_formatting(p)
    # accent1 (4F81BD in the default Office theme) with themeShade=80 -> 254062
    assert resolved.color_rgb == "254062"
    assert resolved.partial is False


# --------------------------------------------------------------------------
# Cycle detection / depth limits
# --------------------------------------------------------------------------


def test_based_on_cycle_raises() -> None:
    doc = Document()
    styles_el = doc.styles.element
    a = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "A"})
    sub(a, "w:name", **{"w:val": "A"})
    sub(a, "w:basedOn", **{"w:val": "B"})

    b = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": "B"})
    sub(b, "w:name", **{"w:val": "B"})
    sub(b, "w:basedOn", **{"w:val": "A"})

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "A"})

    with pytest.raises(StyleCascadeError, match="cycle"):
        resolve_effective_formatting(p)


def test_based_on_depth_limit_raises() -> None:
    doc = Document()
    styles_el = doc.styles.element
    # Build a non-cyclic chain S0 -> S1 -> ... -> S12 (length 13 > 11).
    for i in range(13):
        s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": f"S{i}"})
        sub(s, "w:name", **{"w:val": f"S{i}"})
        if i < 12:
            sub(s, "w:basedOn", **{"w:val": f"S{i + 1}"})

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "S0"})

    with pytest.raises(StyleCascadeError, match="depth"):
        resolve_effective_formatting(p)


def test_missing_style_does_not_raise() -> None:
    """A pStyle pointing at an undefined style is harmless — chain just ends."""
    doc = Document()
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "DoesNotExist"})

    resolved = resolve_effective_formatting(p)
    assert resolved.style_id == "DoesNotExist"
    # docDefaults still apply
    assert resolved.font_size == 11.0


# --------------------------------------------------------------------------
# Target type dispatch
# --------------------------------------------------------------------------


def test_rejects_non_target_type() -> None:
    with pytest.raises(TypeError):
        resolve_effective_formatting(42)  # type: ignore[arg-type]


# --------------------------------------------------------------------------
# Returned dataclass shape
# --------------------------------------------------------------------------


def test_resolved_formatting_is_frozen() -> None:
    doc = Document()
    p = doc.add_paragraph()
    resolved = resolve_effective_formatting(p)
    assert isinstance(resolved, ResolvedFormatting)
    with pytest.raises(dataclasses.FrozenInstanceError):
        resolved.font_size = 99.0  # type: ignore[misc]


def test_provenance_absent_when_flag_off() -> None:
    doc = Document()
    p = doc.add_paragraph()
    resolved = resolve_effective_formatting(p, include_provenance=False)
    assert resolved.provenance is None
