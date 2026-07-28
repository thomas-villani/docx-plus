"""Tests for the field-backed rules and ``mixed-language``.

The interesting cases are where the sweep's paragraph numbering and
``read_fields``' body numbering disagree, which happens as soon as a table
is involved. A finding attributed to the wrong paragraph is worse than no
finding, so that mapping is tested directly.
"""

from __future__ import annotations

from docx import Document

from docx_plus.bookmarks import add_bookmark, add_cross_reference
from docx_plus.core.oxml import build_complex_field, sub
from docx_plus.lint import Finding, lint
from docx_plus.styles import iter_resolved_paragraphs, resolve_effective_formatting


def _rules_fired(doc: Document, rule_id: str) -> list[Finding]:
    return lint(doc, select=[rule_id])


def _caption(doc: Document, text: str) -> object:
    para = doc.add_paragraph(text)
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Caption"})
    return para


def _set_lang(run: object, tag: str) -> None:
    sub(run._r.get_or_add_rPr(), "w:lang", **{"w:val": tag})  # type: ignore[attr-defined]


# ---------------------------------------------------------------------------
# broken-cross-reference.
# ---------------------------------------------------------------------------


def test_broken_cross_reference_fires() -> None:
    doc = Document()
    add_cross_reference(doc.add_paragraph("See "), bookmark="ghost")

    findings = _rules_fired(doc, "broken-cross-reference")

    assert len(findings) == 1
    assert findings[0].severity == "error"
    assert findings[0].observed == "ghost"


def test_broken_cross_reference_ignores_a_resolvable_reference() -> None:
    doc = Document()
    add_bookmark(doc.add_paragraph("Chapter One"), "chapter1")
    add_cross_reference(doc.add_paragraph("See "), bookmark="chapter1")

    assert _rules_fired(doc, "broken-cross-reference") == []


def test_broken_cross_reference_covers_pageref() -> None:
    doc = Document()
    add_cross_reference(doc.add_paragraph("On page "), bookmark="ghost", kind="page")

    assert len(_rules_fired(doc, "broken-cross-reference")) == 1


def test_broken_cross_reference_ignores_other_field_types() -> None:
    doc = Document()
    build_complex_field(doc.add_paragraph()._p, " PAGE ", "1")
    build_complex_field(doc.add_paragraph()._p, " SEQ Figure ", "1")

    assert _rules_fired(doc, "broken-cross-reference") == []


def test_broken_cross_reference_reads_the_instruction_not_the_result() -> None:
    """A stale cached result is exactly how a broken reference hides."""
    doc = Document()
    build_complex_field(doc.add_paragraph()._p, " REF ghost ", "Chapter One")

    findings = _rules_fired(doc, "broken-cross-reference")

    assert len(findings) == 1


def test_broken_cross_reference_locates_a_finding_past_a_table() -> None:
    """The sweep and read_fields number paragraphs differently after a table.

    Mapping by element identity is what keeps the finding on the right
    paragraph; matching raw indices would point at the wrong one.
    """
    doc = Document()
    doc.add_paragraph("before")
    doc.add_table(rows=2, cols=2)
    add_cross_reference(doc.add_paragraph("See "), bookmark="ghost")

    finding = _rules_fired(doc, "broken-cross-reference")[0]
    swept = list(iter_resolved_paragraphs(doc))
    index = finding.location.paragraph_index

    # read_fields sees 6 body paragraphs and calls this one 5; the sweep
    # yields the same paragraph at its own index. Both must name the "See "
    # paragraph, not whichever table cell shares a number with it.
    assert index is not None
    assert swept[index].text.startswith("See")


# ---------------------------------------------------------------------------
# caption-manual-numbering.
# ---------------------------------------------------------------------------


def test_caption_manual_numbering_fires() -> None:
    doc = Document()
    _caption(doc, "Figure 1: A diagram")

    findings = _rules_fired(doc, "caption-manual-numbering")

    assert len(findings) == 1
    assert findings[0].observed == "Figure 1"
    assert "SEQ" in (findings[0].expected or "")


def test_caption_manual_numbering_ignores_a_seq_numbered_caption() -> None:
    doc = Document()
    para = _caption(doc, "Figure ")
    build_complex_field(para._p, " SEQ Figure ", "1")

    assert _rules_fired(doc, "caption-manual-numbering") == []


def test_caption_manual_numbering_ignores_unstyled_prose() -> None:
    """'Table 2 shows...' in body text is prose about a table, not a caption."""
    doc = Document()
    doc.add_paragraph("Table 2 shows the results of the experiment.")

    assert _rules_fired(doc, "caption-manual-numbering") == []


def test_caption_manual_numbering_ignores_an_unnumbered_caption() -> None:
    doc = Document()
    _caption(doc, "A diagram of the process")

    assert _rules_fired(doc, "caption-manual-numbering") == []


def test_caption_manual_numbering_handles_chapter_qualified_numbers() -> None:
    doc = Document()
    _caption(doc, "Table 3.4 - Results")

    assert _rules_fired(doc, "caption-manual-numbering")[0].observed == "Table 3.4"


# ---------------------------------------------------------------------------
# mixed-language.
# ---------------------------------------------------------------------------


def test_mixed_language_fires_on_a_minority_tag() -> None:
    doc = Document()
    for index in range(5):
        _set_lang(doc.add_paragraph().add_run(f"English text {index}"), "en-GB")
    _set_lang(doc.add_paragraph().add_run("Texte francais"), "fr-FR")

    findings = _rules_fired(doc, "mixed-language")

    assert len(findings) == 1
    assert findings[0].observed == "fr-FR"
    assert findings[0].expected == "en-GB"


def test_mixed_language_ignores_a_uniform_document() -> None:
    doc = Document()
    for index in range(3):
        _set_lang(doc.add_paragraph().add_run(f"English text {index}"), "en-GB")

    assert _rules_fired(doc, "mixed-language") == []


def test_mixed_language_ignores_a_document_with_no_tags() -> None:
    doc = Document()
    doc.add_paragraph("Untagged text.")

    assert _rules_fired(doc, "mixed-language") == []


def test_lang_resolves_directly() -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    _set_lang(run, "fr-FR")

    assert resolve_effective_formatting(run).lang == "fr-FR"


def test_lang_resolves_through_the_style_chain() -> None:
    """The rule reads a resolved value, so a style-supplied lang counts.

    That is the case worth having: a whole style tagged with the wrong
    language is invisible in the document and affects every run using it.
    """
    doc = Document()
    style_el = sub(
        doc.styles.element,
        "w:style",
        **{"w:type": "paragraph", "w:styleId": "French"},
    )
    sub(style_el, "w:name", **{"w:val": "French"})
    sub(sub(style_el, "w:rPr"), "w:lang", **{"w:val": "fr-FR"})

    para = doc.add_paragraph()
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "French"})
    run = para.add_run("text")

    resolved = resolve_effective_formatting(run, include_provenance=True)

    assert resolved.lang == "fr-FR"
    assert resolved.provenance is not None
    assert resolved.provenance["lang"].layer == "paragraphStyle"
