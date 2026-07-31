"""Default styles and typed style references, as measured against live Word.

Every answer here came from live Word, read back through COM. Unlike the
spacing work, COM is a perfectly good oracle for this: the question is which
style *won*, and ``Range.Font`` plus ``Range.Style`` report exactly that.

Two divergences turned up, and they run through nearly every document:

- **The default paragraph style was never applied.** A paragraph carrying no
  ``w:pStyle`` — the majority of paragraphs in the wild — resolved to
  ``docDefaults`` alone, so everything ``Normal`` declares was dropped. It
  also sits high enough in the cascade to *beat the table style*, which the
  resolver had no way to express.
- **Style references were followed without checking ``w:type``.** Word
  ignores a ``w:rStyle`` naming a paragraph style, a ``w:tblStyle`` naming a
  character style, and a ``w:basedOn`` crossing between the two. The resolver
  followed all of them.

The distinct-value-per-layer trick is what makes each read self-describing:
``docDefaults`` 10pt, ``Normal`` 20pt, a table style 36pt, and so on, so a
single size names its own winner.

Re-measuring: ``scratchpad/build_default_probe{,2,3,4}.py`` build the
documents, ``read_defaults{,2}.py`` read them back, and ``verify_defaults.py``
scores the lot (92/93 at the time of writing — see
``test_a_document_with_no_normal_style_at_all_reports_no_style`` for the one).
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from lxml import etree

from docx_plus.core.oxml import sub
from docx_plus.styles import resolve_effective_formatting

# Half-points, chosen so no two layers can be confused for one another.
DOC_DEFAULTS_SZ = "20"  # 10pt
NORMAL_SZ = "40"  # 20pt
TABLE_STYLE_SZ = "72"  # 36pt
FIRST_ROW_SZ = "88"  # 44pt


def _set_doc_defaults(doc: Document, sz: str = DOC_DEFAULTS_SZ) -> None:
    """Pin docDefaults to a known size, replacing whatever the template had."""
    defaults = doc.styles.element.find(qn("w:docDefaults"))
    rpr = defaults.find(qn("w:rPrDefault")).find(qn("w:rPr"))
    for tag in ("w:sz", "w:szCs", "w:b"):
        for el in rpr.findall(qn(tag)):
            rpr.remove(el)
    sub(rpr, "w:sz", **{"w:val": sz})
    sub(rpr, "w:szCs", **{"w:val": sz})


def _style(doc: Document, style_id: str) -> etree._Element:
    for st in doc.styles.element.findall(qn("w:style")):
        if st.get(qn("w:styleId")) == style_id:
            return st
    raise AssertionError(f"no style {style_id}")


def _reset_rpr(style: etree._Element) -> etree._Element:
    for el in style.findall(qn("w:rPr")):
        style.remove(el)
    return sub(style, "w:rPr")


def _size_style(doc: Document, style_id: str, sz: str) -> None:
    rpr = _reset_rpr(_style(doc, style_id))
    sub(rpr, "w:sz", **{"w:val": sz})
    sub(rpr, "w:szCs", **{"w:val": sz})


def _add_style(
    doc: Document,
    style_id: str,
    kind: str = "paragraph",
    *,
    sz: str | None = None,
    based_on: str | None = None,
    default: str | None = None,
    name: str | None = None,
) -> etree._Element:
    attrs = {"w:type": kind, "w:styleId": style_id}
    if default is not None:
        attrs["w:default"] = default
    style = sub(doc.styles.element, "w:style", **attrs)
    sub(style, "w:name", **{"w:val": name or style_id})
    if based_on is not None:
        sub(style, "w:basedOn", **{"w:val": based_on})
    if sz is not None:
        rpr = sub(style, "w:rPr")
        sub(rpr, "w:sz", **{"w:val": sz})
        sub(rpr, "w:szCs", **{"w:val": sz})
    return style


def _para(doc: Document, *, pstyle: str | None = None, rstyle: str | None = None):
    para = doc.add_paragraph()
    if pstyle is not None:
        sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": pstyle})
    run = para.add_run("probe")
    if rstyle is not None:
        sub(sub(run._r, "w:rPr"), "w:rStyle", **{"w:val": rstyle})
    return para


def _run(doc: Document, *, pstyle: str | None = None, rstyle: str | None = None):
    """The paragraph's run, which is the only target a ``w:rStyle`` reaches.

    Resolving the *paragraph* never consults ``w:rStyle`` at all, so an
    rStyle assertion made against one passes whatever the run style says.
    """
    return _para(doc, pstyle=pstyle, rstyle=rstyle).runs[0]


def _table(doc: Document, *, tbl_style: str | None, rows: int = 1) -> Table:
    table = doc.add_table(rows=rows, cols=1)
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    for el in tbl_pr.findall(qn("w:tblStyle")):
        tbl_pr.remove(el)
    if tbl_style is not None:
        ref = etree.Element(qn("w:tblStyle"))
        ref.set(qn("w:val"), tbl_style)
        tbl_pr.insert(0, ref)
    return table


def _baseline_doc() -> Document:
    """docDefaults at 10pt, Normal at 20pt, nothing else."""
    doc = Document()
    _set_doc_defaults(doc)
    _size_style(doc, "Normal", NORMAL_SZ)
    return doc


# --------------------------------------------------------------------------
# The default paragraph style applies, and it is the paragraph's own style.
# --------------------------------------------------------------------------


def test_a_paragraph_with_no_pstyle_takes_the_default_style() -> None:
    """Word: 20pt, style Normal. The resolver used to say 10pt and no style."""
    doc = _baseline_doc()
    resolved = resolve_effective_formatting(_para(doc))

    assert resolved.font_size == 20.0
    assert resolved.style_id == "Normal"
    assert resolved.style_name == "Normal"


def test_a_style_that_declares_nothing_does_not_inherit_the_default() -> None:
    """The default applies *instead of* a pStyle, never underneath one.

    ``Loose`` has no ``w:basedOn`` and declares nothing, so Word floors it at
    docDefaults — 10pt — rather than letting Normal's 20pt through.
    """
    doc = _baseline_doc()
    _add_style(doc, "Loose")

    resolved = resolve_effective_formatting(_para(doc, pstyle="Loose"))
    assert resolved.font_size == 10.0
    assert resolved.style_id == "Loose"


@pytest.mark.parametrize(
    ("pstyle", "case"),
    [
        (None, "absent"),
        ("Ghost", "dangling"),
        ("SomeChar", "names a character style"),
    ],
)
def test_the_default_fills_in_for_a_pstyle_that_does_not_resolve(
    pstyle: str | None, case: str
) -> None:
    """Word reported style Normal at 20pt for all three."""
    doc = _baseline_doc()
    _add_style(doc, "SomeChar", "character", sz="88")

    resolved = resolve_effective_formatting(_para(doc, pstyle=pstyle))
    assert resolved.font_size == 20.0, case
    assert resolved.style_id == "Normal", case


def test_the_default_style_is_a_toggle_level_like_any_other() -> None:
    """A bold Normal and a bold character style cancel, per ECMA-376 17.7.3.

    Measured: Word renders this *not bold*. The default style is a genuine
    style level, not a second base alongside docDefaults — had it been a
    base, the character style's bold would have had nothing to cancel
    against and the run would have stayed bold.
    """
    doc = _baseline_doc()
    sub(_style(doc, "Normal").find(qn("w:rPr")), "w:b")
    char_style = _add_style(doc, "CharB", "character")
    sub(sub(char_style, "w:rPr"), "w:b")

    assert resolve_effective_formatting(_para(doc)).bold is True
    assert resolve_effective_formatting(_run(doc, rstyle="CharB")).bold is False


def test_numbering_reaches_a_bare_paragraph_through_the_default_style() -> None:
    """Word numbers it. The default style supplies numbering like any style.

    The first attempt at this measurement was confounded: it reused ``List
    Number``'s ``numId``, whose ``abstractNum`` carries a ``w:styleLink``
    binding it to that style, so Word refused to lend it out and even the
    explicit control went unnumbered.
    """
    doc = Document()
    doc.add_paragraph("seed", style="List Number")
    numbering = doc.part.numbering_part.element

    abstract_ids = [
        int(a.get(qn("w:abstractNumId")))
        for a in numbering.findall(qn("w:abstractNum"))
        if a.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(n.get(qn("w:numId")))
        for n in numbering.findall(qn("w:num"))
        if n.get(qn("w:numId")) is not None
    ]
    abstract_id = str(max(abstract_ids, default=0) + 1)
    num_id = str(max(num_ids, default=0) + 1)

    abstract = etree.Element(qn("w:abstractNum"))
    abstract.set(qn("w:abstractNumId"), abstract_id)
    sub(abstract, "w:multiLevelType", **{"w:val": "singleLevel"})
    lvl = sub(abstract, "w:lvl", **{"w:ilvl": "0"})
    sub(lvl, "w:numFmt", **{"w:val": "decimal"})
    sub(lvl, "w:lvlText", **{"w:val": "%1."})
    numbering.findall(qn("w:abstractNum"))[-1].addnext(abstract)
    num = sub(numbering, "w:num", **{"w:numId": num_id})
    sub(num, "w:abstractNumId", **{"w:val": abstract_id})

    ppr = sub(_style(doc, "Normal"), "w:pPr")
    num_pr = sub(ppr, "w:numPr")
    sub(num_pr, "w:ilvl", **{"w:val": "0"})
    sub(num_pr, "w:numId", **{"w:val": num_id})

    resolved = resolve_effective_formatting(_para(doc))
    assert resolved.num_id == int(num_id)
    assert resolved.num_level == 0


# --------------------------------------------------------------------------
# Which style is "the default".
# --------------------------------------------------------------------------


def test_the_last_style_claiming_default_wins() -> None:
    """Measured both ways round: declaration order is the tie-break.

    With Normal first and AltDefault last Word chose AltDefault; with the two
    swapped it chose Normal. Nothing about *being* Normal mattered.
    """
    doc = _baseline_doc()
    _add_style(doc, "AltDefault", sz="64", default="1")  # appended last

    assert resolve_effective_formatting(_para(doc)).style_id == "AltDefault"


def test_a_default_declared_before_normal_loses_to_it() -> None:
    """The other half of the order measurement."""
    doc = _baseline_doc()
    styles = doc.styles.element
    normal = _style(doc, "Normal")
    alt = _add_style(doc, "AltDefault", sz="64", default="1")
    styles.remove(alt)
    styles.insert(list(styles).index(normal), alt)

    assert resolve_effective_formatting(_para(doc)).style_id == "Normal"


@pytest.mark.parametrize("spelling", ["1", "true", "on"])
def test_default_accepts_every_on_off_spelling(spelling: str) -> None:
    """``w:default="true"`` counts; Word picked that style over Normal."""
    doc = _baseline_doc()
    _style(doc, "Normal").set(qn("w:default"), "0")
    _add_style(doc, "TrueDefault", sz="64", default=spelling)

    assert resolve_effective_formatting(_para(doc)).style_id == "TrueDefault"


def test_default_zero_does_not_claim_the_slot() -> None:
    """An explicit ``w:default="0"`` is not a claim, and nothing else claims."""
    doc = _baseline_doc()
    _style(doc, "Normal").set(qn("w:default"), "0")
    _add_style(doc, "Other", sz="64")

    # Falls through to the style *called* Normal, which is still 20pt.
    assert resolve_effective_formatting(_para(doc)).font_size == 20.0


def test_a_style_called_normal_is_the_fallback_when_none_claims_default() -> None:
    """Word still applied Normal's 20pt with no ``w:default`` anywhere."""
    doc = _baseline_doc()
    del _style(doc, "Normal").attrib[qn("w:default")]

    resolved = resolve_effective_formatting(_para(doc))
    assert resolved.font_size == 20.0
    assert resolved.style_id == "Normal"


def test_a_document_with_no_normal_style_at_all_reports_no_style() -> None:
    """No default and no Normal: docDefaults alone, and no style id.

    This is the one read where the resolver and Word do not literally agree.
    Word's COM reported the style as ``Normal`` — but no such style exists in
    the document, and Word applied none of the two that do (10pt is
    docDefaults). The name is COM naming a phantom, so ``None`` is the
    honest answer and the *formatting* matches.
    """
    doc = _baseline_doc()
    styles = doc.styles.element
    for st in list(styles.findall(qn("w:style"))):
        if (st.get(qn("w:type")) or "paragraph") == "paragraph":
            styles.remove(st)
    _add_style(doc, "Alpha", sz="72")
    _add_style(doc, "Beta", sz="88")

    resolved = resolve_effective_formatting(_para(doc))
    assert resolved.font_size == 10.0
    assert resolved.style_id is None


# --------------------------------------------------------------------------
# The default style beats the table style.
# --------------------------------------------------------------------------


def test_the_default_style_beats_the_table_style() -> None:
    """Word: 20pt (Normal), not 36pt (the table style).

    This is the case that fixes the default style's *position*. Sitting it
    under docDefaults would have been simpler and would have matched every
    other measurement in round 1 — but it would give 36pt here.
    """
    doc = _baseline_doc()
    _add_style(doc, "MyTable", "table", sz=TABLE_STYLE_SZ)
    table = _table(doc, tbl_style="MyTable")
    para = table.cell(0, 0).paragraphs[0]
    para.add_run("probe")

    assert resolve_effective_formatting(para).font_size == 20.0


def test_the_default_style_beats_a_conditional_table_branch_too() -> None:
    """Word: 20pt, not the firstRow branch's 44pt.

    Conditional branches invert precedence against each other in two ways
    (see ``test_tables_word_verified``), so this needed measuring rather
    than assuming — but against a paragraph style they simply lose.
    """
    doc = _baseline_doc()
    style = _add_style(doc, "MyTable", "table", sz=TABLE_STYLE_SZ)
    branch = sub(style, "w:tblStylePr", **{"w:type": "firstRow"})
    branch_rpr = sub(branch, "w:rPr")
    sub(branch_rpr, "w:sz", **{"w:val": FIRST_ROW_SZ})
    sub(branch_rpr, "w:szCs", **{"w:val": FIRST_ROW_SZ})
    _add_style(doc, "Loose")

    table = _table(doc, tbl_style="MyTable", rows=2)
    sub(
        table._tbl.find(qn("w:tblPr")),
        "w:tblLook",
        **{"w:firstRow": "1", "w:noHBand": "1", "w:noVBand": "1"},
    )

    plain = table.cell(0, 0).paragraphs[0]
    plain.add_run("probe")
    assert resolve_effective_formatting(plain).font_size == 20.0

    # ...but the branch does reach a paragraph whose style declares nothing.
    styled = table.cell(1, 0).paragraphs[0]
    sub(sub(styled._p, "w:pPr"), "w:pStyle", **{"w:val": "Loose"})
    styled.add_run("probe")
    table_context = resolve_effective_formatting(styled)
    assert table_context.font_size == 36.0


def test_a_table_style_reaches_a_paragraph_whose_style_is_silent() -> None:
    """36pt: the table style is not dead, it just loses to a declared value."""
    doc = _baseline_doc()
    _add_style(doc, "MyTable", "table", sz=TABLE_STYLE_SZ)
    _add_style(doc, "Loose")

    table = _table(doc, tbl_style="MyTable")
    para = table.cell(0, 0).paragraphs[0]
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Loose"})
    para.add_run("probe")

    assert resolve_effective_formatting(para).font_size == 36.0


# --------------------------------------------------------------------------
# The other two default styles are non-events.
# --------------------------------------------------------------------------


def test_the_default_character_style_never_applies() -> None:
    """DefaultParagraphFont declaring 44pt reached nothing in Word.

    Neither a bare run nor one wearing a character style that declares
    nothing picked it up — so unlike ``w:pStyle`` there is no fallback for
    ``w:rStyle`` at all.
    """
    doc = _baseline_doc()
    _size_style(doc, "DefaultParagraphFont", "88")
    _add_style(doc, "Loose")
    _add_style(doc, "LooseChar", "character")

    bare = resolve_effective_formatting(_para(doc, pstyle="Loose"))
    assert bare.font_size == 10.0

    dressed = resolve_effective_formatting(_run(doc, pstyle="Loose", rstyle="LooseChar"))
    assert dressed.font_size == 10.0


def test_the_default_table_style_never_applies() -> None:
    """TableNormal declaring 28pt reached nothing in a table naming no style."""
    doc = _baseline_doc()
    _size_style(doc, "TableNormal", "56")
    _add_style(doc, "Loose")

    table = _table(doc, tbl_style=None)
    para = table.cell(0, 0).paragraphs[0]
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Loose"})
    para.add_run("probe")

    assert resolve_effective_formatting(para).font_size == 10.0


# --------------------------------------------------------------------------
# A style reference only resolves within its own w:type.
# --------------------------------------------------------------------------


def test_an_rstyle_naming_a_paragraph_style_is_ignored() -> None:
    """Word gave Normal's 20pt, not the paragraph style's 36pt."""
    doc = _baseline_doc()
    _add_style(doc, "PBig", sz=TABLE_STYLE_SZ)

    assert resolve_effective_formatting(_run(doc, rstyle="PBig")).font_size == 20.0


def test_a_dangling_rstyle_contributes_nothing() -> None:
    doc = _baseline_doc()

    assert resolve_effective_formatting(_run(doc, rstyle="GhostChar")).font_size == 20.0


def test_a_tblstyle_naming_a_paragraph_style_is_ignored() -> None:
    """Word gave 10pt — docDefaults — not the paragraph style's 36pt."""
    doc = _baseline_doc()
    _add_style(doc, "PBig", sz=TABLE_STYLE_SZ)
    _add_style(doc, "Loose")

    table = _table(doc, tbl_style="PBig")
    para = table.cell(0, 0).paragraphs[0]
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Loose"})
    para.add_run("probe")

    assert resolve_effective_formatting(para).font_size == 10.0


def test_a_paragraph_style_based_on_a_character_style_inherits_nothing() -> None:
    """Word: 10pt. The style applies; its cross-type ``w:basedOn`` does not.

    The style itself is still the paragraph's — Word named it — so this is
    a severed *link*, not a rejected reference.
    """
    doc = _baseline_doc()
    _add_style(doc, "CBig", "character", sz="88")
    _add_style(doc, "PfromC", based_on="CBig")

    resolved = resolve_effective_formatting(_para(doc, pstyle="PfromC"))
    assert resolved.font_size == 10.0
    assert resolved.style_id == "PfromC"


def test_a_character_style_based_on_a_paragraph_style_inherits_nothing() -> None:
    """Word: 20pt from Normal, nothing from the paragraph style behind it."""
    doc = _baseline_doc()
    _add_style(doc, "PBig", sz=TABLE_STYLE_SZ)
    _add_style(doc, "CfromP", "character", based_on="PBig")

    assert resolve_effective_formatting(_run(doc, rstyle="CfromP")).font_size == 20.0


# --------------------------------------------------------------------------
# Cells.
# --------------------------------------------------------------------------


def test_a_cell_resolves_through_the_default_paragraph_style() -> None:
    """Word reads an untouched cell as Normal, so the cell cascade must too.

    Before this the cell cascade was docDefaults plus the table style and
    nothing else, so it reported the table style's 36pt where Word shows
    20pt.
    """
    doc = _baseline_doc()
    _add_style(doc, "MyTable", "table", sz=TABLE_STYLE_SZ)
    table = _table(doc, tbl_style="MyTable")

    resolved = resolve_effective_formatting(table.cell(0, 0))
    assert resolved.font_size == 20.0
    assert resolved.style_id == "Normal"


def test_a_cell_still_shows_the_table_style_the_default_is_silent_about() -> None:
    """Word: 20pt from Normal, bold from the table style. The two layers split.

    ``w:rPr`` is singular in the schema, so the bold has to go in the one
    :func:`_add_style` already made — a second ``w:rPr`` is simply not found.
    """
    doc = _baseline_doc()
    style = _add_style(doc, "MyTable", "table", sz=TABLE_STYLE_SZ)
    sub(style.find(qn("w:rPr")), "w:b")
    table = _table(doc, tbl_style="MyTable")

    resolved = resolve_effective_formatting(table.cell(0, 0))
    assert resolved.font_size == 20.0
    assert resolved.bold is True
