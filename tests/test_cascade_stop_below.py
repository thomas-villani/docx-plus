"""Tests for ``resolve_effective_formatting(..., stop_below=...)``.

The affordance answers "what would this look like *without* that layer?",
which provenance alone cannot: provenance names the layer that won, not the
value that would have surfaced in its absence.

Also covers the ``w:link`` partner, which is deliberately *not* a layer —
the one place where two style definitions describe the same style, and the
one most likely to be mistaken for a further level of the hierarchy.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.shared import Pt

from docx_plus.core.oxml import sub
from docx_plus.styles import resolve_effective_formatting
from docx_plus.styles.inspect import _LAYER_ORDER, Layer


def test_layer_order_covers_every_layer() -> None:
    """``_LAYER_ORDER`` is derived from ``Layer``, so the two cannot drift."""
    assert set(_LAYER_ORDER) == set(Layer.__args__)  # type: ignore[attr-defined]


def test_none_walks_the_whole_cascade() -> None:
    """The default is unchanged behaviour."""
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.font.size = Pt(24)

    assert resolve_effective_formatting(run, stop_below=None).font_size == 24.0


def test_stop_below_direct_run_drops_the_runs_own_rpr() -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.font.size = Pt(24)

    baseline = resolve_effective_formatting(run, stop_below="directRun")

    assert baseline.font_size == 11.0  # docDefaults


def test_stop_below_direct_run_keeps_the_character_style() -> None:
    """A ``w:rStyle`` sits below direct run formatting, so it survives.

    This is what makes the baseline usable for "is this direct property
    doing anything?": a value the character style supplies is still there.
    """
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.italic = True
    sub(run._r.get_or_add_rPr(), "w:rStyle", **{"w:val": "Emphasis"})

    assert resolve_effective_formatting(run, stop_below="directRun").italic is True


def test_stop_below_direct_paragraph_drops_the_paragraphs_own_ppr() -> None:
    doc = Document()
    para = doc.add_paragraph("text")
    para.paragraph_format.space_after = Pt(18)

    full = resolve_effective_formatting(para)
    baseline = resolve_effective_formatting(para, stop_below="directParagraph")

    assert full.spacing_after == 18 * 20  # twips
    assert baseline.spacing_after != full.spacing_after


def test_stop_below_numbering_drops_a_direct_num_pr() -> None:
    """The two numbering layers gate separately, which is why they are two.

    A paragraph's own ``w:numPr`` is the ``numbering`` layer; one its style
    supplies is ``styleNumbering``. Stopping below ``numbering`` reports the
    list the style would have given it.
    """
    doc = Document()
    para = doc.add_paragraph("item")
    num_pr = sub(sub(para._p, "w:pPr"), "w:numPr")
    sub(num_pr, "w:numId", **{"w:val": "7"})

    assert resolve_effective_formatting(para).num_id == 7
    assert resolve_effective_formatting(para, stop_below="numbering").num_id is None


def test_stop_below_paragraph_style_still_reports_identity() -> None:
    """``style_id`` / ``style_name`` are identity, not formatting.

    A caller resolving beneath the paragraph style still needs to know
    which style it excluded.
    """
    doc = Document()
    para = doc.add_paragraph("text", style="Heading 1")

    resolved = resolve_effective_formatting(para, stop_below="paragraphStyle")

    assert resolved.style_id == "Heading1"
    assert resolved.style_name == "heading 1"
    assert resolved.font_size == 11.0  # the style's 14pt is gone


def test_stop_below_doc_defaults_resolves_nothing() -> None:
    """The lowest layer: excluding it leaves an empty resolve."""
    doc = Document()
    para = doc.add_paragraph("text")

    assert resolve_effective_formatting(para, stop_below="docDefaults").font_size is None


def test_stop_below_records_provenance_for_what_remains() -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.font.size = Pt(24)

    baseline = resolve_effective_formatting(run, stop_below="directRun", include_provenance=True)

    assert baseline.provenance is not None
    assert baseline.provenance["font_size"].layer == "docDefaults"


def test_stop_below_applies_to_a_cell_target() -> None:
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)

    assert resolve_effective_formatting(cell).font_size == 11.0
    assert resolve_effective_formatting(cell, stop_below="docDefaults").font_size is None


def test_rejects_an_unknown_layer() -> None:
    doc = Document()
    para = doc.add_paragraph("text")

    with pytest.raises(ValueError, match="stop_below must be one of"):
        resolve_effective_formatting(para, stop_below="nonsense")  # type: ignore[arg-type]


# ---------------------------------------------------------------------------
# The w:link partner is the same style's other half, not another layer.
# ---------------------------------------------------------------------------


def test_a_run_in_a_bold_heading_resolves_bold() -> None:
    """Regression: the linked Char style must not cancel the toggle.

    Word writes ``Heading1`` and its ``w:link`` partner ``Heading1Char``
    with identical ``w:rPr``. Applying both as independent layers made every
    toggle cancel, so a run inside a stock ``Heading 1`` paragraph resolved
    ``bold=False`` while the paragraph itself resolved ``bold=True``. The
    bold comes from ``Heading1``'s own ``w:rPr``; the partner is inert.
    """
    doc = Document()
    para = doc.add_paragraph("Heading text", style="Heading 1")

    paragraph_bold = resolve_effective_formatting(para).bold
    run_bold = resolve_effective_formatting(para.runs[0]).bold

    assert paragraph_bold is True
    assert run_bold is True


def test_paragraph_and_run_agree_on_a_styles_formatting() -> None:
    """The two targets must not disagree about what the style says."""
    doc = Document()
    para = doc.add_paragraph("Heading text", style="Heading 1")

    para_resolved = resolve_effective_formatting(para)
    run_resolved = resolve_effective_formatting(para.runs[0])

    for prop in ("bold", "italic", "font_size", "font_name", "caps", "small_caps"):
        assert getattr(para_resolved, prop) == getattr(run_resolved, prop), prop


def test_a_link_partners_formatting_is_not_applied() -> None:
    """The ``w:link`` partner is not a cascade layer, so it contributes nothing.

    Measured against Word: a paragraph style whose character formatting
    lives *only* on its Char half renders as though the Char half did not
    exist. The partner is a UI affordance for applying that formatting to a
    selection, not something Word consults to render the paragraph.
    """
    doc = Document()
    styles = doc.styles.element

    para_style = sub(styles, "w:style", **{"w:type": "paragraph", "w:styleId": "OnlyLinked"})
    sub(para_style, "w:name", **{"w:val": "Only Linked"})
    sub(para_style, "w:link", **{"w:val": "OnlyLinkedChar"})

    char_style = sub(styles, "w:style", **{"w:type": "character", "w:styleId": "OnlyLinkedChar"})
    sub(char_style, "w:name", **{"w:val": "Only Linked Char"})
    sub(sub(char_style, "w:rPr"), "w:sz", **{"w:val": "48"})  # 24pt

    para = doc.add_paragraph()
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "OnlyLinked"})
    run = para.add_run("text")

    assert resolve_effective_formatting(run).font_size == 11.0  # docDefaults


def test_a_link_partner_does_not_override_the_style_itself() -> None:
    """When the two halves disagree, the paragraph style's own rPr wins."""
    doc = Document()
    styles = doc.styles.element

    para_style = sub(styles, "w:style", **{"w:type": "paragraph", "w:styleId": "Disagree"})
    sub(para_style, "w:name", **{"w:val": "Disagree"})
    sub(para_style, "w:link", **{"w:val": "DisagreeChar"})
    sub(sub(para_style, "w:rPr"), "w:sz", **{"w:val": "24"})  # 12pt

    char_style = sub(styles, "w:style", **{"w:type": "character", "w:styleId": "DisagreeChar"})
    sub(char_style, "w:name", **{"w:val": "Disagree Char"})
    sub(char_style, "w:link", **{"w:val": "Disagree"})
    sub(sub(char_style, "w:rPr"), "w:sz", **{"w:val": "48"})  # 24pt

    para = doc.add_paragraph()
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Disagree"})
    run = para.add_run("text")

    assert resolve_effective_formatting(run).font_size == 12.0


def test_direct_run_formatting_still_overrides_a_heading_style() -> None:
    doc = Document()
    para = doc.add_paragraph("Heading text", style="Heading 1")
    run = para.runs[0]
    run.font.size = Pt(30)

    assert resolve_effective_formatting(run).font_size == 30.0
