"""Tests for ``docx_plus.core.parts`` — get-or-create plumbing for
separate OOXML parts (comments, footnotes, endnotes)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import PartFactory, XmlPart
from docx.parts.comments import CommentsPart
from docx.parts.numbering import NumberingPart

from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub
from docx_plus.core.parts import (
    COMMENTS_EXTENDED_SPEC,
    COMMENTS_IDS_SPEC,
    COMMENTS_SPEC,
    CT_COMMENTS_EXTENDED,
    CT_COMMENTS_IDS,
    CT_PEOPLE,
    ENDNOTES_SPEC,
    FOOTNOTES_SPEC,
    NUMBERING_SPEC,
    PEOPLE_SPEC,
    RT_COMMENTS_EXTENDED,
    RT_COMMENTS_IDS,
    RT_PEOPLE,
    _CommentsExtendedPart,
    _CommentsIdsPart,
    _EndnotesPart,
    _FootnotesPart,
    _PeoplePart,
    get_or_create_part,
)

# --------------------------------------------------------------------------
# Part-class registration: footnote and endnote parts must round-trip with
# parsed XML rather than as raw blobs.
# --------------------------------------------------------------------------


def test_footnotes_part_class_registered() -> None:
    assert PartFactory.part_type_for[CT.WML_FOOTNOTES] is _FootnotesPart


def test_endnotes_part_class_registered() -> None:
    assert PartFactory.part_type_for[CT.WML_ENDNOTES] is _EndnotesPart


def test_comments_part_class_still_python_docx_default() -> None:
    """python-docx registers ``CommentsPart`` itself; we must not clobber it."""
    assert PartFactory.part_type_for[CT.WML_COMMENTS] is CommentsPart


def test_comments_extended_part_class_registered() -> None:
    assert PartFactory.part_type_for[CT_COMMENTS_EXTENDED] is _CommentsExtendedPart


def test_comments_ids_part_class_registered() -> None:
    assert PartFactory.part_type_for[CT_COMMENTS_IDS] is _CommentsIdsPart


def test_people_part_class_registered() -> None:
    assert PartFactory.part_type_for[CT_PEOPLE] is _PeoplePart


def test_numbering_part_class_still_python_docx_default() -> None:
    """python-docx registers ``NumberingPart``; we must not clobber it."""
    assert PartFactory.part_type_for[CT.WML_NUMBERING] is NumberingPart


# --------------------------------------------------------------------------
# Create paths: fresh document has none of the optional parts.
# --------------------------------------------------------------------------


def test_get_or_create_comments_creates_when_absent() -> None:
    doc = Document()
    part, root = get_or_create_part(doc, COMMENTS_SPEC)
    assert isinstance(part, CommentsPart)
    assert part.content_type == CT.WML_COMMENTS
    assert part.partname == "/word/comments.xml"
    assert root.tag == qn("w:comments")
    assert len(list(root)) == 0  # empty comments root


def test_comments_root_declares_w14_for_para_id_stamping() -> None:
    # Threaded comments write ``w14:paraId`` onto comment body paragraphs,
    # so the fabricated root has to declare the prefix and mark it ignorable.
    _, root = get_or_create_part(Document(), COMMENTS_SPEC)
    assert root.nsmap["w14"] == "http://schemas.microsoft.com/office/word/2010/wordml"
    assert root.get(qn("mc:Ignorable")) == "w14"


def test_get_or_create_comments_extended_creates_when_absent() -> None:
    doc = Document()
    part, root = get_or_create_part(doc, COMMENTS_EXTENDED_SPEC)
    assert isinstance(part, _CommentsExtendedPart)
    assert part.content_type == CT_COMMENTS_EXTENDED
    assert part.partname == "/word/commentsExtended.xml"
    assert root.tag == qn("w15:commentsEx")
    assert len(list(root)) == 0


def test_comments_extended_relationship_is_wired_from_the_document_part() -> None:
    doc = Document()
    part, _ = get_or_create_part(doc, COMMENTS_EXTENDED_SPEC)
    assert doc.part.part_related_by(RT_COMMENTS_EXTENDED) is part


def test_get_or_create_comments_ids_creates_when_absent() -> None:
    doc = Document()
    part, root = get_or_create_part(doc, COMMENTS_IDS_SPEC)
    assert isinstance(part, _CommentsIdsPart)
    assert part.content_type == CT_COMMENTS_IDS
    assert part.partname == "/word/commentsIds.xml"
    assert root.tag == qn("w16cid:commentsIds")
    assert len(list(root)) == 0
    assert doc.part.part_related_by(RT_COMMENTS_IDS) is part


def test_get_or_create_people_creates_when_absent() -> None:
    doc = Document()
    part, root = get_or_create_part(doc, PEOPLE_SPEC)
    assert isinstance(part, _PeoplePart)
    assert part.content_type == CT_PEOPLE
    assert part.partname == "/word/people.xml"
    assert root.tag == qn("w15:people")
    assert len(list(root)) == 0
    assert doc.part.part_related_by(RT_PEOPLE) is part


# --------------------------------------------------------------------------
# Numbering: the one spec whose part python-docx can load but cannot create.
# --------------------------------------------------------------------------


def test_get_or_create_numbering_returns_the_template_part() -> None:
    # python-docx's bundled template already ships numbering.xml, so this is
    # the lookup path rather than the create path.
    doc = Document()
    part, root = get_or_create_part(doc, NUMBERING_SPEC)
    assert isinstance(part, NumberingPart)
    assert root.tag == qn("w:numbering")
    assert doc.part.part_related_by(RT.NUMBERING) is part


def test_get_or_create_numbering_creates_when_absent() -> None:
    """The gap this spec exists to close.

    ``doc.part.numbering_part`` fabricates through ``NumberingPart.new()``,
    an unimplemented stub that raises ``NotImplementedError``. Documents
    from LibreOffice / Pandoc / stripped templates reach that path.
    """
    doc = Document()
    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype == RT.NUMBERING:
            doc.part.drop_rel(rid)

    part, root = get_or_create_part(doc, NUMBERING_SPEC)
    assert isinstance(part, NumberingPart)
    assert part.partname == "/word/numbering.xml"
    assert root.tag == qn("w:numbering")
    assert len(list(root)) == 0


def test_numbering_part_round_trip(tmp_path: Path) -> None:
    doc = Document()
    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype == RT.NUMBERING:
            doc.part.drop_rel(rid)
    _, root = get_or_create_part(doc, NUMBERING_SPEC)
    sub(root, "w:abstractNum", **{"w:abstractNumId": "42"})

    path = tmp_path / "numbering.docx"
    doc.save(path)
    reloaded = Document(path)

    part = reloaded.part.part_related_by(RT.NUMBERING)
    entries = part.element.findall(qn("w:abstractNum"))
    assert [entry.get(qn("w:abstractNumId")) for entry in entries] == ["42"]


def test_comments_extended_part_round_trip(tmp_path: Path) -> None:
    doc = Document()
    _, root = get_or_create_part(doc, COMMENTS_EXTENDED_SPEC)
    sub(root, "w15:commentEx", **{"w15:paraId": "0000002A", "w15:done": "1"})

    path = tmp_path / "extended.docx"
    doc.save(path)
    reloaded = Document(path)

    part = reloaded.part.part_related_by(RT_COMMENTS_EXTENDED)
    assert isinstance(part, _CommentsExtendedPart)
    entries = part.element.findall(qn("w15:commentEx"))
    assert [entry.get(qn("w15:paraId")) for entry in entries] == ["0000002A"]


def test_get_or_create_footnotes_creates_when_absent() -> None:
    doc = Document()
    part, root = get_or_create_part(doc, FOOTNOTES_SPEC)
    assert isinstance(part, _FootnotesPart)
    assert part.content_type == CT.WML_FOOTNOTES
    assert part.partname == "/word/footnotes.xml"
    assert root.tag == qn("w:footnotes")
    # ECMA-376 / Word convention: separator + continuationSeparator (C1).
    note_ids = [n.get(qn("w:id")) for n in root.findall(qn("w:footnote"))]
    assert note_ids == ["-1", "0"]


def test_get_or_create_endnotes_creates_when_absent() -> None:
    doc = Document()
    part, root = get_or_create_part(doc, ENDNOTES_SPEC)
    assert isinstance(part, _EndnotesPart)
    assert part.content_type == CT.WML_ENDNOTES
    assert part.partname == "/word/endnotes.xml"
    assert root.tag == qn("w:endnotes")
    note_ids = [n.get(qn("w:id")) for n in root.findall(qn("w:endnote"))]
    assert note_ids == ["-1", "0"]


def test_create_wires_relationship_from_document_part() -> None:
    doc = Document()
    part, _ = get_or_create_part(doc, FOOTNOTES_SPEC)
    # Round-trip the relationship lookup that the helper itself performs.
    assert doc.part.part_related_by(RT.FOOTNOTES) is part


# --------------------------------------------------------------------------
# Idempotency: second call returns the same part, not a freshly built one.
# --------------------------------------------------------------------------


def test_second_call_returns_same_part() -> None:
    doc = Document()
    first, first_root = get_or_create_part(doc, COMMENTS_SPEC)
    second, second_root = get_or_create_part(doc, COMMENTS_SPEC)
    assert first is second
    assert first_root is second_root


def test_idempotency_preserves_mutations() -> None:
    """Edits between the two calls must survive the second lookup."""
    doc = Document()
    _, root = get_or_create_part(doc, FOOTNOTES_SPEC)
    sub(root, "w:footnote", **{"w:id": "1"})
    _, root_again = get_or_create_part(doc, FOOTNOTES_SPEC)
    assert root is root_again
    # 2 seeded separators (ids -1, 0) + 1 user-added footnote (id 1) = 3.
    assert len(list(root_again)) == 3
    note_ids = [n.get(qn("w:id")) for n in root_again]
    assert note_ids == ["-1", "0", "1"]


# --------------------------------------------------------------------------
# Round-trip: parts survive save/reopen with the registered class.
# --------------------------------------------------------------------------


def test_footnotes_part_round_trip(tmp_path: Path) -> None:
    doc = Document()
    _, root = get_or_create_part(doc, FOOTNOTES_SPEC)
    sub(root, "w:footnote", **{"w:id": "1"})
    out = tmp_path / "with_footnotes.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    part_again, root_again = get_or_create_part(reopened, FOOTNOTES_SPEC)
    assert isinstance(part_again, _FootnotesPart)
    # The footnote we wrote AND the seeded separators (C1) survive round-trip.
    note_ids = [n.get(qn("w:id")) for n in root_again]
    assert note_ids == ["-1", "0", "1"]
    # Verify the separator types are correctly typed for Word to render the
    # horizontal divider line above the footnote area.
    types = {n.get(qn("w:id")): n.get(qn("w:type")) for n in root_again}
    assert types["-1"] == "separator"
    assert types["0"] == "continuationSeparator"
    assert types["1"] is None  # user note has no type attribute


def test_endnotes_part_round_trip(tmp_path: Path) -> None:
    doc = Document()
    _, root = get_or_create_part(doc, ENDNOTES_SPEC)
    sub(root, "w:endnote", **{"w:id": "1"})
    out = tmp_path / "with_endnotes.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    part_again, root_again = get_or_create_part(reopened, ENDNOTES_SPEC)
    assert isinstance(part_again, _EndnotesPart)
    note_ids = [n.get(qn("w:id")) for n in root_again]
    assert note_ids == ["-1", "0", "1"]
    types = {n.get(qn("w:id")): n.get(qn("w:type")) for n in root_again}
    assert types["-1"] == "separator"
    assert types["0"] == "continuationSeparator"


def test_no_part_created_when_not_requested(tmp_path: Path) -> None:
    """``get_or_create_part`` must be the only thing that adds the
    relationship — a save/reopen without calling it must not have one."""
    doc = Document()
    out = tmp_path / "no_optional_parts.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    for rel in reopened.part.rels.values():
        assert rel.reltype not in {RT.COMMENTS, RT.FOOTNOTES, RT.ENDNOTES}


# --------------------------------------------------------------------------
# Falls back to plain XmlPart for unknown content types.
# --------------------------------------------------------------------------


def test_unknown_content_type_falls_back_to_xmlpart() -> None:
    from docx_plus.core.parts import PartSpec

    custom = PartSpec(
        partname="/word/custom.xml",
        content_type="application/vnd.example.custom+xml",
        relationship_type="http://example.com/custom",
        root_xml=(
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            b'<w:custom xmlns:w="http://schemas.openxmlformats.org/wordprocessingml'
            b'/2006/main"/>'
        ),
    )
    doc = Document()
    part, _ = get_or_create_part(doc, custom)
    assert type(part) is XmlPart
