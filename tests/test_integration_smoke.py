"""Phase 1 exit-criterion smoke test.

Build a Document, attach a ``w:sdt`` constructed through ``core/oxml``, reserve
its id through ``core/ids``, save, reopen with python-docx, verify the SDT
survives the round trip. This catches the broadest class of "the XML I'm
emitting is malformed enough python-docx rejects it" bug before any feature
code lands.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from docx_plus._testing.ooxml_asserts import assert_ids_unique
from docx_plus.core.ids import IdRegistry
from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub, xpath


def test_round_trip_sdt_through_core(tmp_path: Path) -> None:
    out = tmp_path / "round_trip.docx"
    doc = Document()
    reg = IdRegistry(doc)
    sdt_id = reg.next()

    para = doc.add_paragraph("Before SDT.")
    sdt = sub(para._p, "w:sdt")
    sdt_pr = sub(sdt, "w:sdtPr")
    sub(sdt_pr, "w:id", **{"w:val": str(sdt_id)})
    sub(sdt_pr, "w:tag", **{"w:val": "smoke"})
    sub(sdt, "w:sdtContent")

    doc.save(out)

    reopened = Document(str(out))
    [reopened_id_el] = xpath(reopened.element.body, ".//w:sdt/w:sdtPr/w:id")
    assert reopened_id_el.get(qn("w:val")) == str(sdt_id)

    [tag_el] = xpath(reopened.element.body, ".//w:sdt/w:sdtPr/w:tag")
    assert tag_el.get(qn("w:val")) == "smoke"

    assert_ids_unique(reopened)
