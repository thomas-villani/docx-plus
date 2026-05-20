"""Tests for ``docx_plus.styles.modify``.

Phase 3 of the build. Coverage strategy: every operation that writes to
styles.xml is verified by reading back through ``resolve_effective_formatting``
on a paragraph that uses the style. The cascade resolver from Phase 2 acts as
the round-trip oracle — if the resolver returns the value we wrote, the
written XML is structurally correct enough for the cascade to pick it up.

Round-trip tests (build → save → re-open via python-docx → assert) sit at the
end and verify the styles.xml survives Word-style serialisation.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub
from docx_plus.styles.inspect import resolve_effective_formatting
from docx_plus.styles.modify import (
    StyleExistsError,
    StyleInfo,
    StyleInUseError,
    StyleNotFoundError,
    StyleProxy,
    UnknownStylePropertyError,
    apply_style,
    create_style,
    delete_style,
    ensure_style,
    find_matching_style,
    list_styles,
    modify_style,
    remap_styles,
)

# --------------------------------------------------------------------------
# create_style: happy path + property writes via cascade round-trip.
# --------------------------------------------------------------------------


def test_create_style_returns_proxy_with_correct_metadata() -> None:
    doc = Document()
    proxy = create_style(doc, "Foo", name="My Foo", ui_priority=42, q_format=True)
    assert isinstance(proxy, StyleProxy)
    assert proxy.style_id == "Foo"
    assert proxy.style_type == "paragraph"
    assert proxy.name == "My Foo"
    assert proxy.ui_priority == 42
    assert proxy.q_format is True


def test_create_style_default_name_falls_back_to_style_id() -> None:
    doc = Document()
    proxy = create_style(doc, "Bar")
    assert proxy.name == "Bar"


def test_create_style_with_basedon_link_next() -> None:
    doc = Document()
    create_style(doc, "Parent")
    proxy = create_style(
        doc,
        "Child",
        based_on="Parent",
        next_style="Parent",
        linked_style="ChildChar",
    )
    assert proxy.based_on == "Parent"
    assert proxy.next_style == "Parent"
    assert proxy.linked_style == "ChildChar"


def test_create_style_marks_custom_by_default() -> None:
    doc = Document()
    proxy = create_style(doc, "Custom")
    assert proxy.element.get(qn("w:customStyle")) == "1"


def test_create_style_custom_false_omits_attr() -> None:
    doc = Document()
    proxy = create_style(doc, "BuiltinLike", custom=False)
    assert proxy.element.get(qn("w:customStyle")) is None


def test_create_style_raises_on_duplicate_id() -> None:
    doc = Document()
    create_style(doc, "Dup")
    with pytest.raises(StyleExistsError, match="Dup"):
        create_style(doc, "Dup")


def test_create_style_unknown_property_raises() -> None:
    doc = Document()
    with pytest.raises(UnknownStylePropertyError, match="not_a_property"):
        create_style(doc, "X", not_a_property=42)  # type: ignore[arg-type]


def test_unknown_style_property_is_a_typeerror() -> None:
    """Per SPEC §5: unrecognized property kwargs raise TypeError."""
    doc = Document()
    with pytest.raises(TypeError):
        create_style(doc, "Y", bogus=1)  # type: ignore[arg-type]


# --------------------------------------------------------------------------
# create_style + cascade resolver — every supported property type.
# --------------------------------------------------------------------------


def _paragraph_with_style(doc: Document, style_id: str) -> object:
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()  # type: ignore[attr-defined]
    sub(ppr, "w:pStyle", **{"w:val": style_id})
    return p


def test_font_size_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", font_size=18.0)
    p = _paragraph_with_style(doc, "S")
    assert resolve_effective_formatting(p).font_size == 18.0


def test_font_size_writes_sz_and_szcs() -> None:
    doc = Document()
    proxy = create_style(doc, "S", font_size=14.0)
    rpr = proxy.element.find(qn("w:rPr"))
    assert rpr is not None
    assert rpr.find(qn("w:sz")).get(qn("w:val")) == "28"
    assert rpr.find(qn("w:szCs")).get(qn("w:val")) == "28"


def test_color_rgb_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", color_rgb="FF0000")
    p = _paragraph_with_style(doc, "S")
    assert resolve_effective_formatting(p).color_rgb == "FF0000"


def test_color_rgb_strips_leading_hash() -> None:
    doc = Document()
    create_style(doc, "S", color_rgb="#00FF00")
    p = _paragraph_with_style(doc, "S")
    assert resolve_effective_formatting(p).color_rgb == "00FF00"


def test_color_rgb_invalid_raises() -> None:
    doc = Document()
    with pytest.raises(ValueError, match="RRGGBB"):
        create_style(doc, "S", color_rgb="not-a-hex")


def test_alignment_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", alignment="center")
    p = _paragraph_with_style(doc, "S")
    assert resolve_effective_formatting(p).alignment == "center"


def test_outline_level_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", outline_level=1)
    p = _paragraph_with_style(doc, "S")
    assert resolve_effective_formatting(p).outline_level == 1


def test_indent_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", indent_left=720, indent_right=360, indent_first_line=240)
    p = _paragraph_with_style(doc, "S")
    resolved = resolve_effective_formatting(p)
    assert resolved.indent_left == 720
    assert resolved.indent_right == 360
    assert resolved.indent_first_line == 240


def test_indent_first_line_negative_writes_hanging() -> None:
    doc = Document()
    proxy = create_style(doc, "S", indent_left=720, indent_first_line=-360)
    ind = proxy.element.find(qn("w:pPr")).find(qn("w:ind"))
    assert ind.get(qn("w:hanging")) == "360"
    assert ind.get(qn("w:firstLine")) is None
    p = _paragraph_with_style(doc, "S")
    assert resolve_effective_formatting(p).indent_first_line == -360


def test_spacing_before_after_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", spacing_before=120, spacing_after=240)
    p = _paragraph_with_style(doc, "S")
    resolved = resolve_effective_formatting(p)
    assert resolved.spacing_before == 120
    assert resolved.spacing_after == 240


def test_line_spacing_auto_default() -> None:
    doc = Document()
    create_style(doc, "S", line_spacing=2.0)  # default rule auto -> twips=480
    p = _paragraph_with_style(doc, "S")
    resolved = resolve_effective_formatting(p)
    assert resolved.line_spacing == 2.0
    assert resolved.line_spacing_rule == "auto"


def test_line_spacing_with_explicit_rule() -> None:
    doc = Document()
    # Set rule first, then value, so the writer interprets value as twips.
    create_style(doc, "S", line_spacing_rule="exact", line_spacing=480.0)
    p = _paragraph_with_style(doc, "S")
    resolved = resolve_effective_formatting(p)
    assert resolved.line_spacing == 480.0
    assert resolved.line_spacing_rule == "exact"


def test_keep_with_next_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", keep_with_next=True, keep_lines=True, page_break_before=True)
    p = _paragraph_with_style(doc, "S")
    resolved = resolve_effective_formatting(p)
    assert resolved.keep_with_next is True
    assert resolved.keep_lines is True
    assert resolved.page_break_before is True


def test_toggle_bold_true_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", bold=True)
    p = _paragraph_with_style(doc, "S")
    r = p.add_run("text")  # type: ignore[attr-defined]
    assert resolve_effective_formatting(r).bold is True


def test_toggle_bold_false_writes_explicit_val_false() -> None:
    """Per SPEC §5: bold=False writes <w:b w:val="false"/>."""
    doc = Document()
    proxy = create_style(doc, "S", bold=False)
    rpr = proxy.element.find(qn("w:rPr"))
    b = rpr.find(qn("w:b"))
    assert b.get(qn("w:val")) == "false"


def test_underline_and_highlight_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", underline="single", highlight="yellow")
    p = _paragraph_with_style(doc, "S")
    r = p.add_run("text")  # type: ignore[attr-defined]
    resolved = resolve_effective_formatting(r)
    assert resolved.underline == "single"
    assert resolved.highlight == "yellow"


def test_vert_align_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", vert_align="superscript")
    p = _paragraph_with_style(doc, "S")
    r = p.add_run("text")  # type: ignore[attr-defined]
    assert resolve_effective_formatting(r).vert_align == "superscript"


def test_font_name_round_trip() -> None:
    doc = Document()
    create_style(doc, "S", font_name="Arial")
    p = _paragraph_with_style(doc, "S")
    assert resolve_effective_formatting(p).font_name == "Arial"


# issues.md H17: the six toggles the resolver surfaces but the writer
# couldn't previously produce. Each maps a ResolvedFormatting field to its
# rPr tag; round-trip proves the writer emits XML the resolver reads back.
@pytest.mark.parametrize(
    ("prop", "tag"),
    [
        ("cs_bold", "bCs"),
        ("cs_italic", "iCs"),
        ("emboss", "emboss"),
        ("imprint", "imprint"),
        ("outline", "outline"),
        ("shadow", "shadow"),
    ],
)
def test_new_toggle_true_round_trip(prop: str, tag: str) -> None:
    doc = Document()
    create_style(doc, "S", **{prop: True})
    p = _paragraph_with_style(doc, "S")
    r = p.add_run("text")  # type: ignore[attr-defined]
    assert getattr(resolve_effective_formatting(r), prop) is True


@pytest.mark.parametrize(
    ("prop", "tag"),
    [
        ("cs_bold", "bCs"),
        ("cs_italic", "iCs"),
        ("emboss", "emboss"),
        ("imprint", "imprint"),
        ("outline", "outline"),
        ("shadow", "shadow"),
    ],
)
def test_new_toggle_false_writes_explicit_val_false(prop: str, tag: str) -> None:
    doc = Document()
    proxy = create_style(doc, "S", **{prop: False})
    rpr = proxy.element.find(qn("w:rPr"))
    flag = rpr.find(qn(f"w:{tag}"))
    assert flag is not None
    assert flag.get(qn("w:val")) == "false"


# --------------------------------------------------------------------------
# Schema-correct child ordering.
# --------------------------------------------------------------------------


def test_style_children_ordered_correctly() -> None:
    """name, basedOn, uiPriority, qFormat, pPr, rPr — schema order."""
    doc = Document()
    proxy = create_style(
        doc,
        "S",
        based_on="Normal",
        next_style="Normal",
        ui_priority=10,
        q_format=True,
        font_size=14.0,
        alignment="center",
    )
    locals_in_order = [
        # Skip non-w children (none should appear).
        proxy.element[i].tag.split("}", 1)[1]
        for i in range(len(proxy.element))
    ]
    expected_subset = ["name", "basedOn", "next", "uiPriority", "qFormat", "pPr", "rPr"]
    # All of expected_subset should appear in this order in locals_in_order.
    indices = [locals_in_order.index(t) for t in expected_subset]
    assert indices == sorted(indices), f"order broke: {locals_in_order}"


def test_ppr_children_ordered_correctly() -> None:
    """jc must come AFTER spacing/ind in schema order."""
    doc = Document()
    proxy = create_style(
        doc,
        "S",
        alignment="center",
        spacing_before=100,
        indent_left=720,
        outline_level=2,
    )
    ppr = proxy.element.find(qn("w:pPr"))
    locals_in_order = [child.tag.split("}", 1)[1] for child in ppr]
    # Schema order: spacing < ind < jc < outlineLvl.
    expected = ["spacing", "ind", "jc", "outlineLvl"]
    indices = [locals_in_order.index(t) for t in expected]
    assert indices == sorted(indices), f"pPr order broke: {locals_in_order}"


def test_rpr_children_ordered_correctly() -> None:
    """color must precede sz which must precede u in schema order."""
    doc = Document()
    proxy = create_style(
        doc,
        "S",
        bold=True,
        color_rgb="FF0000",
        font_size=14.0,
        underline="single",
    )
    rpr = proxy.element.find(qn("w:rPr"))
    locals_in_order = [child.tag.split("}", 1)[1] for child in rpr]
    # Schema: b < color < sz < u.
    expected = ["b", "color", "sz", "u"]
    indices = [locals_in_order.index(t) for t in expected]
    assert indices == sorted(indices), f"rPr order broke: {locals_in_order}"


# --------------------------------------------------------------------------
# modify_style: change-one-preserve-others, toggles, clearing.
# --------------------------------------------------------------------------


def test_modify_preserves_other_properties() -> None:
    doc = Document()
    create_style(doc, "S", font_size=14.0, alignment="center", bold=True)
    modify_style(doc, "S", font_size=18.0)  # change only font_size
    p = _paragraph_with_style(doc, "S")
    r = p.add_run("text")  # type: ignore[attr-defined]
    resolved = resolve_effective_formatting(r)
    assert resolved.font_size == 18.0
    assert resolved.alignment == "center"
    assert resolved.bold is True


def test_modify_clears_property_with_none() -> None:
    doc = Document()
    create_style(doc, "S", font_size=14.0, color_rgb="FF0000")
    modify_style(doc, "S", color_rgb=None)
    p = _paragraph_with_style(doc, "S")
    resolved = resolve_effective_formatting(p)
    assert resolved.color_rgb is None  # docDefaults provides nothing here
    assert resolved.font_size == 14.0  # other prop preserved


def test_modify_toggle_clear_with_none_removes_element() -> None:
    """Per SPEC §5: bold=None removes <w:b> so XOR with parent resumes."""
    doc = Document()
    proxy = create_style(doc, "S", bold=True)
    modify_style(doc, "S", bold=None)
    rpr = proxy.element.find(qn("w:rPr"))
    if rpr is not None:
        assert rpr.find(qn("w:b")) is None


def test_modify_indent_overwrites_existing() -> None:
    doc = Document()
    create_style(doc, "S", indent_left=720)
    modify_style(doc, "S", indent_left=1440)
    p = _paragraph_with_style(doc, "S")
    assert resolve_effective_formatting(p).indent_left == 1440


def test_modify_missing_raises_by_default() -> None:
    doc = Document()
    with pytest.raises(StyleNotFoundError, match="Nope"):
        modify_style(doc, "Nope", font_size=12.0)


def test_modify_missing_create_falls_through() -> None:
    doc = Document()
    proxy = modify_style(doc, "NewStyle", if_missing="create", font_size=12.0)
    assert proxy.style_id == "NewStyle"
    p = _paragraph_with_style(doc, "NewStyle")
    assert resolve_effective_formatting(p).font_size == 12.0


def test_modify_unknown_property_raises() -> None:
    doc = Document()
    create_style(doc, "S")
    with pytest.raises(UnknownStylePropertyError):
        modify_style(doc, "S", what_is_this=1)  # type: ignore[arg-type]


# --------------------------------------------------------------------------
# apply_style: paragraph, run, cell.
# --------------------------------------------------------------------------


def test_apply_style_to_paragraph() -> None:
    doc = Document()
    create_style(doc, "Big", font_size=20.0)
    p = doc.add_paragraph("hi")
    apply_style(p, "Big")
    assert resolve_effective_formatting(p).font_size == 20.0


def test_apply_style_to_run_writes_rstyle() -> None:
    doc = Document()
    create_style(doc, "RunChar", style_type="character", font_size=20.0)
    p = doc.add_paragraph()
    r = p.add_run("text")
    apply_style(r, "RunChar")
    rstyle = r._r.find(qn("w:rPr")).find(qn("w:rStyle"))
    assert rstyle is not None
    assert rstyle.get(qn("w:val")) == "RunChar"


def test_apply_style_replaces_existing_pstyle() -> None:
    doc = Document()
    create_style(doc, "A", font_size=10.0)
    create_style(doc, "B", font_size=20.0)
    p = doc.add_paragraph()
    apply_style(p, "A")
    apply_style(p, "B")
    # Only one pStyle should survive.
    pstyles = p._p.find(qn("w:pPr")).findall(qn("w:pStyle"))
    assert len(pstyles) == 1
    assert pstyles[0].get(qn("w:val")) == "B"


def test_apply_style_to_cell_styles_each_paragraph() -> None:
    doc = Document()
    create_style(doc, "CellStyle", font_size=14.0)
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.add_paragraph("more text")  # cell now has 2 paragraphs
    apply_style(cell, "CellStyle")
    for p in cell.paragraphs:
        assert resolve_effective_formatting(p).font_size == 14.0


def test_apply_style_unknown_id_raises() -> None:
    doc = Document()
    p = doc.add_paragraph()
    with pytest.raises(StyleNotFoundError, match="Ghost"):
        apply_style(p, "Ghost")


def test_apply_style_rejects_non_target() -> None:
    doc = Document()
    create_style(doc, "S")
    with pytest.raises(TypeError):
        apply_style(42, "S")  # type: ignore[arg-type]


# --------------------------------------------------------------------------
# delete_style: with/without references, force.
# --------------------------------------------------------------------------


def test_delete_style_unreferenced_succeeds() -> None:
    doc = Document()
    create_style(doc, "Doomed")
    delete_style(doc, "Doomed")
    assert list_styles(doc, style_type="paragraph") == [
        s for s in list_styles(doc, style_type="paragraph") if s.style_id != "Doomed"
    ]


def test_delete_style_referenced_by_paragraph_raises() -> None:
    doc = Document()
    create_style(doc, "Used")
    _paragraph_with_style(doc, "Used")
    with pytest.raises(StyleInUseError, match="Used"):
        delete_style(doc, "Used")


def test_delete_style_referenced_by_basedon_raises() -> None:
    doc = Document()
    create_style(doc, "Parent")
    create_style(doc, "Child", based_on="Parent")
    with pytest.raises(StyleInUseError, match="Parent"):
        delete_style(doc, "Parent")


def test_delete_style_force_overrides_reference_check() -> None:
    doc = Document()
    create_style(doc, "Used")
    _paragraph_with_style(doc, "Used")
    delete_style(doc, "Used", force=True)
    # Style is gone; the dangling pStyle is harmless (Word falls back to Normal).
    assert all(s.style_id != "Used" for s in list_styles(doc))


def test_delete_style_unknown_raises() -> None:
    doc = Document()
    with pytest.raises(StyleNotFoundError):
        delete_style(doc, "Nope")


def test_delete_style_with_self_basedon_doesnt_count() -> None:
    """A style whose own basedOn happens to match its id should still be deletable."""
    doc = Document()
    create_style(doc, "S")
    # Manually set basedOn to itself (pathological but not a reference *to* it
    # from another style).
    sub(create_style(doc, "Other").element, "w:basedOn", **{"w:val": "S"})
    # Other references S; S references itself in no way. Delete should fail
    # because Other refs it.
    with pytest.raises(StyleInUseError):
        delete_style(doc, "S")


# --------------------------------------------------------------------------
# ensure_style: idempotent + latent built-in materialisation.
# --------------------------------------------------------------------------


def test_ensure_style_returns_existing() -> None:
    doc = Document()
    proxy1 = create_style(doc, "X", font_size=14.0)
    proxy2 = ensure_style(doc, "X")
    assert proxy2.element is proxy1.element  # same element, no recreation


def test_ensure_style_idempotent_single_definition() -> None:
    doc = Document()
    ensure_style(doc, "X", font_size=14.0)
    ensure_style(doc, "X", font_size=99.0)  # second call is a no-op
    matches = doc.styles.element.findall(qn("w:style"))
    matching = [m for m in matches if m.get(qn("w:styleId")) == "X"]
    assert len(matching) == 1
    # Original definition preserved.
    p = _paragraph_with_style(doc, "X")
    assert resolve_effective_formatting(p).font_size == 14.0


def test_ensure_style_creates_unknown_with_defaults() -> None:
    doc = Document()
    proxy = ensure_style(doc, "BrandNew", font_size=20.0, color_rgb="0000FF")
    assert proxy.style_id == "BrandNew"
    p = _paragraph_with_style(doc, "BrandNew")
    resolved = resolve_effective_formatting(p)
    assert resolved.font_size == 20.0
    assert resolved.color_rgb == "0000FF"


def test_ensure_style_materialises_heading1_when_absent() -> None:
    """ensure_style on a missing built-in id uses the table's defaults."""
    doc = Document()
    # python-docx ships Heading1 already materialised (Word 2007 defaults).
    # Strip it to exercise the genuine latent-materialisation path.
    styles_root = doc.styles.element
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Heading1']")
    if existing is not None:
        styles_root.remove(existing)
    proxy = ensure_style(doc, "Heading1")
    assert proxy.name == "heading 1"
    assert proxy.based_on == "Normal"
    assert proxy.linked_style == "Heading1Char"
    assert proxy.q_format is True
    # Built-in materialisation does NOT mark customStyle.
    assert proxy.element.get(qn("w:customStyle")) is None
    p = _paragraph_with_style(doc, "Heading1")
    resolved = resolve_effective_formatting(p)
    assert resolved.style_id == "Heading1"
    assert resolved.outline_level == 0
    assert resolved.font_size == 16.0  # from built-ins table


def test_ensure_style_returns_existing_heading1_unchanged() -> None:
    """If python-docx already shipped Heading1, ensure_style returns it as-is."""
    doc = Document()
    proxy = ensure_style(doc, "Heading1")
    assert proxy.style_id == "Heading1"
    p = _paragraph_with_style(doc, "Heading1")
    resolved = resolve_effective_formatting(p)
    # Whatever python-docx ships, the resolver must see Heading1 with an
    # outline level (the structural fingerprint of the built-in).
    assert resolved.style_id == "Heading1"
    assert resolved.outline_level == 0


def test_ensure_style_normal_when_absent_marks_default() -> None:
    doc = Document()
    styles_root = doc.styles.element
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Normal']")
    if existing is not None:
        styles_root.remove(existing)
    proxy = ensure_style(doc, "Normal")
    assert proxy.element.get(qn("w:default")) == "1"


def test_ensure_style_all_known_builtins_succeed() -> None:
    """Every entry in the known-built-ins table materialises without error."""
    from docx_plus.styles.modify import _BUILTIN_STYLES

    doc = Document()
    for sid in _BUILTIN_STYLES:
        proxy = ensure_style(doc, sid)
        assert proxy.style_id == sid


def test_builtin_styles_count_matches_documented_total() -> None:
    """Pin the count of `_BUILTIN_STYLES` so docs and code can't drift.

    README, ARCHITECTURE.md §5, docs/index.md, and TEST_GAPS.md I6 all
    cite the same number. If you add or remove an entry, update those
    locations together (the ARCHITECTURE table has per-tier counts that
    must also still sum to the total).
    """
    from docx_plus.styles.modify import _BUILTIN_STYLES

    assert len(_BUILTIN_STYLES) == 107


# --------------------------------------------------------------------------
# list_styles + StyleProxy.
# --------------------------------------------------------------------------


def test_list_styles_returns_styleinfo() -> None:
    doc = Document()
    create_style(doc, "Custom1")
    items = list_styles(doc)
    assert all(isinstance(i, StyleInfo) for i in items)
    assert any(i.style_id == "Custom1" for i in items)


def test_list_styles_filters_by_type() -> None:
    doc = Document()
    create_style(doc, "ParaA")
    create_style(doc, "CharB", style_type="character")
    para_only = list_styles(doc, style_type="paragraph")
    char_only = list_styles(doc, style_type="character")
    assert any(i.style_id == "ParaA" for i in para_only)
    assert all(i.style_id != "CharB" for i in para_only)
    assert any(i.style_id == "CharB" for i in char_only)


def test_list_styles_include_latent_returns_unmaterialised_builtins() -> None:
    """python-docx ships most built-ins; PlaceholderText is one it doesn't."""
    doc = Document()
    items = list_styles(doc, include_latent=True)
    sids = {i.style_id for i in items}
    assert "PlaceholderText" in sids
    ph = next(i for i in items if i.style_id == "PlaceholderText")
    assert ph.is_latent is True


def test_list_styles_include_latent_excludes_already_materialised() -> None:
    doc = Document()
    ensure_style(doc, "PlaceholderText")
    items = list_styles(doc, include_latent=True)
    ph_entries = [i for i in items if i.style_id == "PlaceholderText"]
    assert len(ph_entries) == 1
    assert ph_entries[0].is_latent is False


def test_styleproxy_modify_writes_through() -> None:
    doc = Document()
    proxy = create_style(doc, "P", font_size=12.0)
    proxy.modify(font_size=24.0)
    p = _paragraph_with_style(doc, "P")
    assert resolve_effective_formatting(p).font_size == 24.0


def test_styleproxy_delete_removes_style() -> None:
    doc = Document()
    proxy = create_style(doc, "Goner")
    proxy.delete()
    assert all(s.style_id != "Goner" for s in list_styles(doc))


# --------------------------------------------------------------------------
# find_matching_style: case/space-insensitive lookup.
# --------------------------------------------------------------------------


def test_find_matching_style_exact_id_returns_trivial_match() -> None:
    doc = Document()
    create_style(doc, "MyStyle")
    assert find_matching_style(doc, "MyStyle") == "MyStyle"


def test_find_matching_style_case_insensitive_on_id() -> None:
    doc = Document()
    create_style(doc, "myheading1")
    assert find_matching_style(doc, "MyHeading1") == "myheading1"


def test_find_matching_style_space_insensitive_on_id() -> None:
    """An id like 'Heading 1' (with a space) matches a target 'Heading1'."""
    doc = Document()
    # Strip the python-docx-shipped Heading1 to control the test setup.
    styles_root = doc.styles.element
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Heading1']")
    if existing is not None:
        styles_root.remove(existing)
    # Create a style whose id literally has a space.
    create_style(doc, "Heading 1")
    assert find_matching_style(doc, "Heading1") == "Heading 1"


def test_find_matching_style_via_name_attribute() -> None:
    """Match on w:name when the styleId itself doesn't normalise to target."""
    doc = Document()
    create_style(doc, "Custom42", name="My Big Heading")
    assert find_matching_style(doc, "MyBigHeading") == "Custom42"


def test_find_matching_style_no_match_returns_none() -> None:
    doc = Document()
    assert find_matching_style(doc, "DefinitelyNotAnyStyle") is None


# --------------------------------------------------------------------------
# remap_styles: bulk reconciliation.
# --------------------------------------------------------------------------


def test_remap_styles_no_op_when_all_targets_defined() -> None:
    """If every requested target already exists exactly, mapping is trivial."""
    doc = Document()
    create_style(doc, "Foo")
    result = remap_styles(doc, targets=["Foo"])
    assert result == {"Foo": "Foo"}


def test_remap_styles_uses_explicit_mapping() -> None:
    doc = Document()
    create_style(doc, "ActualHeading")
    result = remap_styles(
        doc,
        targets=["MyHeading"],
        mapping={"MyHeading": "ActualHeading"},
    )
    assert result == {"MyHeading": "ActualHeading"}


def test_remap_styles_explicit_mapping_to_unknown_raises() -> None:
    doc = Document()
    with pytest.raises(StyleNotFoundError, match="DoesNotExist"):
        remap_styles(doc, targets=["X"], mapping={"X": "DoesNotExist"})


def test_remap_styles_matcher_finds_renamed_style() -> None:
    """Doc has 'Heading 1' (with space); remap_styles maps target 'Heading1'."""
    doc = Document()
    styles_root = doc.styles.element
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Heading1']")
    if existing is not None:
        styles_root.remove(existing)
    create_style(doc, "Heading 1")
    result = remap_styles(doc, targets=["Heading1"])
    assert result == {"Heading1": "Heading 1"}


def test_remap_styles_rewrites_body_refs() -> None:
    """A paragraph styled with target id gets rewritten to the resolved id."""
    doc = Document()
    styles_root = doc.styles.element
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Heading1']")
    if existing is not None:
        styles_root.remove(existing)
    create_style(doc, "Heading 1", font_size=20.0)
    # A paragraph references the *missing* "Heading1" id.
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "Heading1"})

    remap_styles(doc, targets=["Heading1"])

    # The reference now points at the matched style and the cascade resolves.
    pstyle = p._p.find(qn("w:pPr")).find(qn("w:pStyle"))
    assert pstyle.get(qn("w:val")) == "Heading 1"
    assert resolve_effective_formatting(p).font_size == 20.0


def test_remap_styles_create_missing_materialises_builtin() -> None:
    doc = Document()
    styles_root = doc.styles.element
    # Strip Heading2 to force the create-missing path.
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Heading2']")
    if existing is not None:
        styles_root.remove(existing)
    result = remap_styles(doc, targets=["Heading2"], create_missing=True)
    assert result == {"Heading2": "Heading2"}
    # And the style was actually materialised.
    assert any(s.style_id == "Heading2" for s in list_styles(doc))


def test_remap_styles_unresolved_target_omitted() -> None:
    """Target with no match, not in built-ins, create_missing=False → omitted."""
    doc = Document()
    result = remap_styles(doc, targets=["NoSuchStyle"], create_missing=False)
    assert "NoSuchStyle" not in result


def test_remap_styles_unresolved_with_create_missing_still_omitted_if_not_builtin() -> None:
    """create_missing only helps for ids in the known-built-ins table."""
    doc = Document()
    result = remap_styles(doc, targets=["NotABuiltin"], create_missing=True)
    assert "NotABuiltin" not in result


def test_remap_styles_create_missing_inherits_from_existing_normal() -> None:
    """The new built-in basedOn=Normal so it picks up the doc's customised Normal."""
    doc = Document()
    styles_root = doc.styles.element
    # Customise Normal with a specific font size.
    normal = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Normal']")
    if normal is not None:
        styles_root.remove(normal)
    create_style(doc, "Normal", custom=False, font_name="MyCorporateFont")
    # Now create Heading1 via remap_styles' create_missing.
    existing_h1 = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Heading1']")
    if existing_h1 is not None:
        styles_root.remove(existing_h1)
    remap_styles(doc, targets=["Heading1"], create_missing=True)
    # A paragraph styled Heading1 inherits the customised Normal's font name.
    p = doc.add_paragraph()
    apply_style(p, "Heading1")
    resolved = resolve_effective_formatting(p)
    assert resolved.font_name == "MyCorporateFont"


def test_remap_styles_default_targets_are_known_builtins() -> None:
    """Calling with no targets considers every entry in the built-ins table."""
    doc = Document()
    result = remap_styles(doc)
    # Every materialised built-in id should appear self-mapped.
    assert "Heading1" in result
    assert result["Heading1"] == "Heading1"


# --------------------------------------------------------------------------
# ensure_style match_existing flag.
# --------------------------------------------------------------------------


def test_ensure_style_match_existing_returns_renamed() -> None:
    """match_existing=True returns the proxy to the renamed style."""
    doc = Document()
    styles_root = doc.styles.element
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Heading1']")
    if existing is not None:
        styles_root.remove(existing)
    create_style(doc, "heading 1")  # renamed variant
    proxy = ensure_style(doc, "Heading1", match_existing=True)
    assert proxy.style_id == "heading 1"


def test_ensure_style_match_existing_false_skips_matcher() -> None:
    """Without match_existing, the builtin table fires even if a rename exists."""
    doc = Document()
    styles_root = doc.styles.element
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='Heading1']")
    if existing is not None:
        styles_root.remove(existing)
    create_style(doc, "heading 1")  # rename exists
    proxy = ensure_style(doc, "Heading1")  # match_existing defaults False
    # The matcher was skipped — we got the built-ins-table materialised one.
    assert proxy.style_id == "Heading1"


def test_ensure_style_match_existing_falls_through_to_builtin_when_no_match() -> None:
    doc = Document()
    styles_root = doc.styles.element
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}='PlaceholderText']")
    if existing is not None:
        styles_root.remove(existing)
    proxy = ensure_style(doc, "PlaceholderText", match_existing=True)
    assert proxy.style_id == "PlaceholderText"


# --------------------------------------------------------------------------
# Full save → reopen round-trip via python-docx.
# --------------------------------------------------------------------------


def test_save_reopen_preserves_created_style(tmp_path: Path) -> None:
    doc = Document()
    create_style(
        doc,
        "Roundtrip",
        font_size=15.0,
        color_rgb="3366CC",
        bold=True,
        alignment="center",
        ui_priority=42,
    )
    p = doc.add_paragraph("Content")
    apply_style(p, "Roundtrip")
    out = tmp_path / "rt.docx"
    doc.save(out)

    doc2 = Document(str(out))
    p2 = doc2.paragraphs[0]
    r2 = p2.add_run("dummy")
    resolved = resolve_effective_formatting(r2)
    assert resolved.style_id == "Roundtrip"
    assert resolved.font_size == 15.0
    assert resolved.color_rgb == "3366CC"
    assert resolved.bold is True
    assert resolved.alignment == "center"


def test_save_reopen_preserves_ensure_style_heading(tmp_path: Path) -> None:
    """SPEC §5 latent materialisation: re-open and resolver still sees Heading1."""
    doc = Document()
    ensure_style(doc, "Heading1")
    p = doc.add_paragraph("Title-ish")
    apply_style(p, "Heading1")
    out = tmp_path / "h1.docx"
    doc.save(out)

    doc2 = Document(str(out))
    p2 = doc2.paragraphs[0]
    resolved = resolve_effective_formatting(p2)
    assert resolved.style_id == "Heading1"
    assert resolved.outline_level == 0


# --------------------------------------------------------------------------
# M11: style-type filtering guards against wrong-type collisions.
# --------------------------------------------------------------------------


def _strip_builtin(doc: Document, style_id: str) -> None:
    styles_root = doc.styles.element
    existing = styles_root.find(f"./{qn('w:style')}[@{qn('w:styleId')}={style_id!r}]")
    if existing is not None:
        styles_root.remove(existing)


def test_find_matching_style_type_filter_skips_wrong_type() -> None:
    """A character style named 'Heading 1' must not match a paragraph request."""
    doc = Document()
    _strip_builtin(doc, "Heading1")
    # A character style whose NAME normalises to the paragraph target.
    create_style(doc, "CharHeading", style_type="character", name="Heading 1")

    # No filter: the char look-alike matches (legacy behaviour, preserved).
    assert find_matching_style(doc, "Heading1") == "CharHeading"
    # Paragraph filter: correctly skipped.
    assert find_matching_style(doc, "Heading1", style_type="paragraph") is None
    # Character filter: matches.
    assert find_matching_style(doc, "Heading1", style_type="character") == "CharHeading"


def test_find_matching_style_type_filter_rejects_wrong_type_exact_id() -> None:
    """Even an exact-id match is skipped when its type is wrong."""
    doc = Document()
    create_style(doc, "MyMark", style_type="character")
    assert find_matching_style(doc, "MyMark", style_type="character") == "MyMark"
    assert find_matching_style(doc, "MyMark", style_type="paragraph") is None


def test_ensure_style_match_existing_ignores_wrong_type_lookalike() -> None:
    """match_existing won't bind a paragraph built-in to a character look-alike (M11)."""
    doc = Document()
    _strip_builtin(doc, "Heading1")
    create_style(doc, "CharHeading", style_type="character", name="Heading 1")

    proxy = ensure_style(doc, "Heading1", match_existing=True)
    # The char look-alike was rejected; the paragraph built-in was materialised.
    assert proxy.style_id == "Heading1"
    info = next(s for s in list_styles(doc) if s.style_id == "Heading1")
    assert info.style_type == "paragraph"


def test_remap_styles_only_rewrites_matching_ref_tag() -> None:
    """A paragraph target resolves through w:pStyle, never w:rStyle (M11)."""
    doc = Document()
    _strip_builtin(doc, "Heading1")
    create_style(doc, "Heading 1")  # paragraph style (default)
    # A run carries an rStyle that happens to share the target id — it must
    # NOT be rewritten, because the resolved style is a paragraph style.
    p = doc.add_paragraph()
    run = p.add_run("x")
    rpr = run._r.get_or_add_rPr()
    sub(rpr, "w:rStyle", **{"w:val": "Heading1"})
    # And a paragraph that legitimately references it via pStyle.
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "Heading1"})

    remap_styles(doc, targets=["Heading1"])

    assert p._p.find(qn("w:pPr")).find(qn("w:pStyle")).get(qn("w:val")) == "Heading 1"
    # The rStyle was left untouched (paragraph style can't go in an rStyle slot).
    assert rpr.find(qn("w:rStyle")).get(qn("w:val")) == "Heading1"


# --------------------------------------------------------------------------
# M12: references inside headers / footers count and are rewritten.
# --------------------------------------------------------------------------


def _styled_header_paragraph(doc: Document, style_id: str) -> object:
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    ppr = p._p.get_or_add_pPr()
    # The default header paragraph already has a pStyle="Header"; replace it
    # so the paragraph references exactly the style under test.
    for existing in ppr.findall(qn("w:pStyle")):
        ppr.remove(existing)
    sub(ppr, "w:pStyle", **{"w:val": style_id})
    return p


def test_delete_style_referenced_in_header_raises() -> None:
    """A style used only in a header is still 'in use' (M12)."""
    doc = Document()
    create_style(doc, "HdrOnly")
    _styled_header_paragraph(doc, "HdrOnly")
    with pytest.raises(StyleInUseError):
        delete_style(doc, "HdrOnly")


def test_remap_styles_rewrites_header_refs() -> None:
    """remap rewrites references inside headers, not just the main body (M12)."""
    doc = Document()
    _strip_builtin(doc, "Heading1")
    create_style(doc, "Heading 1")
    hp = _styled_header_paragraph(doc, "Heading1")

    remap_styles(doc, targets=["Heading1"])

    pstyle = hp._p.find(qn("w:pPr")).find(qn("w:pStyle"))  # type: ignore[attr-defined]
    assert pstyle.get(qn("w:val")) == "Heading 1"
