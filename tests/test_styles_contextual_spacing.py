"""``contextualSpacing`` in the cascade, and the structural edges of the walk.

The numbers Word actually renders live in
``tests/test_contextual_spacing_word_verified.py``. This file covers the
cascade plumbing and the malformed or unusual document shapes the probes
could not express.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from docx_plus.core.oxml import sub
from docx_plus.styles import (
    ParagraphSpacing,
    resolve_effective_formatting,
    resolve_paragraph_spacing,
)


def _style(doc: Document, style_id: str, *, contextual: bool | None = None) -> None:
    style = sub(doc.styles.element, "w:style", **{"w:type": "paragraph", "w:styleId": style_id})
    sub(style, "w:name", **{"w:val": style_id})
    ppr = sub(style, "w:pPr")
    if contextual is not None:
        sub(ppr, "w:contextualSpacing", **({} if contextual else {"w:val": "0"}))


def _para(doc: Document, style_id: str, *, before: int = 0, after: int = 0):
    para = doc.add_paragraph()
    para.add_run("x")
    ppr = para._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": style_id})
    sub(ppr, "w:spacing", **{"w:before": str(before), "w:after": str(after)})
    return para


# --------------------------------------------------------------------------
# The cascade field.
# --------------------------------------------------------------------------


def test_contextual_spacing_resolves_from_the_style() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    assert resolve_effective_formatting(_para(doc, "Ctx")).contextual_spacing is True


def test_contextual_spacing_is_none_when_nothing_declares_it() -> None:
    doc = Document()
    _style(doc, "Plain")
    assert resolve_effective_formatting(_para(doc, "Plain")).contextual_spacing is None


def test_a_direct_val_zero_overrides_the_style() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    para = _para(doc, "Ctx")
    sub(para._p.get_or_add_pPr(), "w:contextualSpacing", **{"w:val": "0"})
    assert resolve_effective_formatting(para).contextual_spacing is False


def test_a_style_val_zero_overrides_an_inherited_on() -> None:
    """It overrides rather than toggling: it is not an ECMA-376 17.7.3 toggle."""
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    child = sub(doc.styles.element, "w:style", **{"w:type": "paragraph", "w:styleId": "Child"})
    sub(child, "w:name", **{"w:val": "Child"})
    sub(child, "w:basedOn", **{"w:val": "Ctx"})
    sub(sub(child, "w:pPr"), "w:contextualSpacing", **{"w:val": "0"})
    assert resolve_effective_formatting(_para(doc, "Child")).contextual_spacing is False


def test_provenance_names_the_layer_that_set_it() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    para = _para(doc, "Ctx")
    resolved = resolve_effective_formatting(para, include_provenance=True)
    assert resolved.provenance["contextual_spacing"].layer == "paragraphStyle"
    sub(para._p.get_or_add_pPr(), "w:contextualSpacing")
    resolved = resolve_effective_formatting(para, include_provenance=True)
    assert resolved.provenance["contextual_spacing"].layer == "directParagraph"


def test_the_flag_does_not_change_the_declared_spacing() -> None:
    """``spacing_before`` / ``spacing_after`` stay what the cascade declares.

    The linter's ``style-drift`` rule compares direct formatting against the
    style, which is a cascade question — folding a neighbour-dependent
    suppression into those fields would make it report the wrong number.
    """
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    _para(doc, "Ctx", before=240)
    resolved = resolve_effective_formatting(first)
    assert resolved.spacing_after == 240
    assert resolve_paragraph_spacing(first).space_below == 0


# --------------------------------------------------------------------------
# Paragraphs with no neighbour on a side.
# --------------------------------------------------------------------------


def test_a_lone_paragraph_keeps_both_declared_values() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    only = _para(doc, "Ctx", before=240, after=360)
    assert resolve_paragraph_spacing(only) == ParagraphSpacing(
        declared_before=240,
        declared_after=360,
        contextual_spacing=True,
        before_suppressed=False,
        after_suppressed=False,
        space_above=240,
        space_below=360,
    )


def test_the_first_and_last_paragraphs_keep_their_outer_edges() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", before=240, after=240)
    last = _para(doc, "Ctx", before=240, after=360)
    assert resolve_paragraph_spacing(first).space_above == 240
    assert resolve_paragraph_spacing(first).before_suppressed is False
    assert resolve_paragraph_spacing(last).space_below == 360
    assert resolve_paragraph_spacing(last).after_suppressed is False


def test_spacing_absent_everywhere_reads_as_zero() -> None:
    doc = Document()
    # The stock template declares w:after in docDefaults; drop it so nothing
    # in the cascade mentions spacing at all.
    default_ppr = (
        doc.styles.element.find(qn("w:docDefaults")).find(qn("w:pPrDefault")).find(qn("w:pPr"))
    )
    default_ppr.remove(default_ppr.find(qn("w:spacing")))
    _style(doc, "Plain")
    para = doc.add_paragraph()
    sub(para._p.get_or_add_pPr(), "w:pStyle", **{"w:val": "Plain"})
    spacing = resolve_paragraph_spacing(para)
    assert (spacing.declared_before, spacing.declared_after) == (0, 0)
    assert spacing.contextual_spacing is False
    assert resolve_effective_formatting(para).spacing_after is None


# --------------------------------------------------------------------------
# What the adjacency walk steps over, and what stops it.
# --------------------------------------------------------------------------


def test_bookmarks_between_paragraphs_are_stepped_over() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    marker = sub(doc.element.body, "w:bookmarkStart", **{"w:id": "1", "w:name": "b"})
    second._p.addprevious(marker)
    second._p.addprevious(sub(doc.element.body, "w:bookmarkEnd", **{"w:id": "1"}))
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0


def test_an_xml_comment_between_paragraphs_is_stepped_over() -> None:
    from lxml import etree

    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    second._p.addprevious(etree.Comment("a note"))
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0


def test_a_nested_content_control_is_still_transparent() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    outer = sub(doc.element.body, "w:sdt")
    outer_content = sub(outer, "w:sdtContent")
    inner = sub(outer_content, "w:sdt")
    inner_content = sub(inner, "w:sdtContent")
    second._p.addprevious(outer)
    inner_content.append(second._p)
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0


def test_a_content_control_with_no_content_element_stops_the_search() -> None:
    """Malformed, so the conservative answer: the declared space is kept."""
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    second._p.addprevious(sub(doc.element.body, "w:sdt"))
    assert resolve_paragraph_spacing(first).space_below == 240
    assert resolve_paragraph_spacing(second).space_above == 240


def test_a_content_control_holding_a_table_stops_the_search() -> None:
    """The table is content, and content between two paragraphs separates them."""
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    sdt = sub(doc.element.body, "w:sdt")
    content = sub(sdt, "w:sdtContent")
    table = doc.add_table(rows=1, cols=1)
    second._p.addprevious(sdt)
    content.append(table._tbl)
    assert resolve_paragraph_spacing(first).space_below == 240
    assert resolve_paragraph_spacing(second).space_above == 240


def test_a_content_control_holding_nothing_but_a_bookmark_stops_the_search() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    sdt = sub(doc.element.body, "w:sdt")
    content = sub(sdt, "w:sdtContent")
    sub(content, "w:bookmarkStart", **{"w:id": "7", "w:name": "b"})
    second._p.addprevious(sdt)
    assert resolve_paragraph_spacing(first).space_below == 240
    assert resolve_paragraph_spacing(second).space_above == 240


def test_a_content_control_holding_a_comment_node_is_stepped_over() -> None:
    from lxml import etree

    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    sdt = sub(doc.element.body, "w:sdt")
    content = sub(sdt, "w:sdtContent")
    content.append(etree.Comment("marker"))
    second._p.addprevious(sdt)
    content.append(second._p)
    assert resolve_paragraph_spacing(first).space_below == 0


def test_content_controls_nested_past_the_guard_stop_the_search() -> None:
    """A malformed part cannot make the walk recurse without end."""
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    outer = sub(doc.element.body, "w:sdt")
    second._p.addprevious(outer)
    container = sub(outer, "w:sdtContent")
    for _ in range(40):
        container = sub(sub(container, "w:sdt"), "w:sdtContent")
    container.append(second._p)
    assert resolve_paragraph_spacing(first).space_below == 240


def test_a_nested_content_control_with_no_content_element_stops_the_search() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    outer = sub(doc.element.body, "w:sdt")
    sub(sub(outer, "w:sdtContent"), "w:sdt")  # inner control, no sdtContent
    second._p.addprevious(outer)
    assert resolve_paragraph_spacing(first).space_below == 240


def test_climbing_out_of_content_controls_nested_past_the_guard_stops() -> None:
    """The mirror of the descent guard: climbing out is bounded too."""
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    outer = sub(doc.element.body, "w:sdt")
    second._p.addprevious(outer)
    container = sub(outer, "w:sdtContent")
    for _ in range(40):
        container = sub(sub(container, "w:sdt"), "w:sdtContent")
    container.append(second._p)
    assert resolve_paragraph_spacing(second).space_above == 240


def test_a_paragraph_in_a_detached_content_control_has_no_neighbours() -> None:
    from lxml import etree

    doc = Document()
    _style(doc, "Ctx", contextual=True)
    para = _para(doc, "Ctx", before=240, after=360)
    orphan = etree.SubElement(etree.Element(qn("w:sdt")), qn("w:sdtContent"))
    doc.element.body.remove(para._p)
    orphan.append(para._p)
    orphan.getparent().remove(orphan)
    spacing = resolve_paragraph_spacing(para)
    assert (spacing.space_above, spacing.space_below) == (240, 360)


def test_a_paragraph_last_inside_a_content_control_looks_outside_it() -> None:
    """Climbing out of ``sdtContent`` is the mirror of descending into it."""
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, "Ctx", before=240)
    sdt = sub(doc.element.body, "w:sdt")
    content = sub(sdt, "w:sdtContent")
    first._p.addprevious(sdt)
    content.append(first._p)
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0


def test_a_detached_paragraph_has_no_neighbours() -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    para = _para(doc, "Ctx", before=240, after=360)
    doc.element.body.remove(para._p)
    spacing = resolve_paragraph_spacing(para)
    assert (spacing.space_above, spacing.space_below) == (240, 360)


@pytest.mark.parametrize("style_id", ["Ctx", "Other"])
def test_suppression_tracks_the_resolved_style_id(style_id: str) -> None:
    doc = Document()
    _style(doc, "Ctx", contextual=True)
    _style(doc, "Other", contextual=True)
    first = _para(doc, "Ctx", after=240)
    second = _para(doc, style_id, before=240)
    expected = 0 if style_id == "Ctx" else 240
    assert resolve_paragraph_spacing(first).space_below == expected
    assert resolve_paragraph_spacing(second).space_above == expected


def test_paragraphs_with_no_style_reference_still_match_each_other() -> None:
    """Two unstyled paragraphs both resolve to the default style."""
    doc = Document()
    defaults = doc.styles.element.find(qn("w:docDefaults"))
    sub(defaults.find(qn("w:pPrDefault")).find(qn("w:pPr")), "w:contextualSpacing")
    first, second = doc.add_paragraph("a"), doc.add_paragraph("b")
    sub(first._p.get_or_add_pPr(), "w:spacing", **{"w:after": "240"})
    sub(second._p.get_or_add_pPr(), "w:spacing", **{"w:before": "240"})
    assert resolve_paragraph_spacing(first).space_below == 0
    assert resolve_paragraph_spacing(second).space_above == 0
