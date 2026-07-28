"""Tests for ``docx-plus lint``.

Drives :func:`docx_plus.cli.main` with an explicit ``argv`` and inspects
stdout, stderr, and the exit code — same shape as ``test_cli.py``. The exit
code carries meaning here that it does not for the other read commands, so
it is asserted throughout.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from docx_plus.cli import main


@pytest.fixture
def messy_doc(tmp_path: Path) -> Path:
    """A document carrying one instance of several different defects."""
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Too deep", style="Heading 3")
    doc.add_paragraph("Body with  two spaces.")
    doc.add_paragraph("trailing space ")
    run = doc.add_paragraph().add_run("redundant")
    run.font.size = Pt(11)
    path = tmp_path / "messy.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def clean_doc(tmp_path: Path) -> Path:
    """A document with nothing to report."""
    doc = Document()
    doc.add_paragraph("Title", style="Heading 1")
    doc.add_paragraph("An ordinary paragraph of prose.")
    path = tmp_path / "clean.docx"
    doc.save(str(path))
    return path


def test_lint_reports_findings_and_exits_nonzero(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """Findings print, and the exit code marks the run as failed for CI."""
    code = main(["lint", str(messy_doc)])
    out = capsys.readouterr().out

    assert code == 1
    assert "heading-level-skip" in out
    assert "double-space" in out
    assert "findings" in out


def test_lint_clean_document_exits_zero(
    clean_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """A clean document says so and exits 0, so the command gates cleanly."""
    code = main(["lint", str(clean_doc)])

    assert code == 0
    assert "No findings." in capsys.readouterr().out


def test_lint_json(messy_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """``--json`` emits the full finding shape."""
    code = main(["lint", str(messy_doc), "--json"])
    records = json.loads(capsys.readouterr().out)

    assert code == 1
    assert records
    record = records[0]
    assert set(record) == {
        "rule",
        "kind",
        "severity",
        "message",
        "location",
        "observed",
        "expected",
        "fixable",
        "adds_content",
    }
    assert set(record["location"]) == {
        "paragraph_index",
        "run_index",
        "style_id",
        "excerpt",
    }


def test_lint_rule_selects_one_rule(messy_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """``--rule`` narrows to a single rule id."""
    main(["lint", str(messy_doc), "--rule", "double-space", "--json"])
    records = json.loads(capsys.readouterr().out)

    assert {r["rule"] for r in records} == {"double-space"}


def test_lint_rule_is_repeatable(messy_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """``--rule`` accumulates."""
    main(
        [
            "lint",
            str(messy_doc),
            "--rule",
            "double-space",
            "--rule",
            "trailing-whitespace",
            "--json",
        ]
    )
    records = json.loads(capsys.readouterr().out)

    assert {r["rule"] for r in records} == {"double-space", "trailing-whitespace"}


def test_lint_rule_by_tag_enables_off_by_default_rules(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """Naming a tag opts into that cluster's off-by-default rules."""
    doc = Document()
    doc.add_paragraph("Content")
    doc.add_paragraph("")
    doc.add_paragraph("")
    path = tmp_path / "blanks.docx"
    doc.save(str(path))

    main(["lint", str(path), "--json"])
    default_rules = {r["rule"] for r in json.loads(capsys.readouterr().out)}

    main(["lint", str(path), "--rule", "whitespace", "--json"])
    tagged_rules = {r["rule"] for r in json.loads(capsys.readouterr().out)}

    assert "stray-empty-paragraph" not in default_rules
    assert "stray-empty-paragraph" in tagged_rules


def test_lint_exclude(messy_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """``--exclude`` drops a rule from the default set."""
    main(["lint", str(messy_doc), "--exclude", "double-space", "--json"])
    records = json.loads(capsys.readouterr().out)

    assert "double-space" not in {r["rule"] for r in records}
    assert records


def test_lint_unknown_rule_errors(messy_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """A typo'd selector fails loudly rather than silently matching nothing."""
    code = main(["lint", str(messy_doc), "--rule", "no-such-rule"])

    assert code == 1
    assert "no-such-rule" in capsys.readouterr().err


def test_lint_no_tables(tmp_path: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """``--no-tables`` skips paragraphs inside cells."""
    doc = Document()
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Inside  a cell."
    path = tmp_path / "table.docx"
    doc.save(str(path))

    assert main(["lint", str(path)]) == 1
    capsys.readouterr()
    assert main(["lint", str(path), "--no-tables"]) == 0


def test_lint_list_rules(capsys: pytest.CaptureFixture[str]) -> None:
    """``--list-rules`` prints the catalogue without needing a document."""
    code = main(["lint", "--list-rules"])
    out = capsys.readouterr().out

    assert code == 0
    assert "double-space" in out
    assert "on by default" in out


def test_lint_list_rules_json(capsys: pytest.CaptureFixture[str]) -> None:
    """The catalogue is machine-readable too, for building a config."""
    code = main(["lint", "--list-rules", "--json"])
    records = json.loads(capsys.readouterr().out)

    assert code == 0
    assert set(records[0]) == {
        "id",
        "kind",
        "severity",
        "description",
        "tags",
        "default_on",
    }


def test_lint_requires_a_file_unless_listing_rules(
    capsys: pytest.CaptureFixture[str],
) -> None:
    """FILE is optional only because ``--list-rules`` needs no document."""
    code = main(["lint"])

    assert code == 1
    assert "FILE is required" in capsys.readouterr().err


def test_lint_missing_file(capsys: pytest.CaptureFixture[str]) -> None:
    """A bad path reports the standard CLI error."""
    assert main(["lint", "no-such-file.docx"]) == 1
    assert "not found" in capsys.readouterr().err


def test_lint_does_not_modify_the_input(messy_doc: Path) -> None:
    """The command is a pure read — the file on disk is byte-identical after."""
    before = messy_doc.read_bytes()
    main(["lint", str(messy_doc)])

    assert messy_doc.read_bytes() == before


def test_lint_output_is_cp1252_safe(messy_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """Text output must survive a default Windows console.

    The same constraint the runnable examples carry — a non-ASCII ellipsis
    or bullet here would raise ``UnicodeEncodeError`` for a Windows user.
    """
    main(["lint", str(messy_doc)])
    out = capsys.readouterr().out

    out.encode("cp1252")


def test_lint_list_rules_output_is_cp1252_safe(capsys: pytest.CaptureFixture[str]) -> None:
    """Same constraint for the catalogue, which prints every description."""
    main(["lint", "--list-rules"])

    capsys.readouterr().out.encode("cp1252")


def test_lint_excerpt_is_quoted_so_whitespace_is_visible(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """A trailing-space finding must show the space it is reporting."""
    main(["lint", str(messy_doc), "--rule", "trailing-whitespace"])
    out = capsys.readouterr().out

    assert '> "trailing space "' in out
