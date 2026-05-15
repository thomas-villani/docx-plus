"""Provenance tracking for ``resolve_effective_formatting``.

When ``include_provenance=True``, each populated field on the resulting
:class:`ResolvedFormatting` carries a :class:`FormattingSource` describing
the cascade layer that produced it. The non-provenance code path must remain
byte-identical to the no-provenance path — SPEC §4 calls out this invariant
explicitly. The two paths are exercised together in
:func:`test_provenance_does_not_change_values`.
"""

from __future__ import annotations

from dataclasses import fields as dc_fields
from pathlib import Path

from docx import Document

from docx_plus.core.oxml import sub
from docx_plus.styles.inspect import (
    FormattingSource,
    resolve_effective_formatting,
)


def _add_paragraph_style(
    doc: Document,
    style_id: str,
    *,
    based_on: str | None = None,
    rpr_children: list[tuple[str, dict[str, str] | None]] | None = None,
) -> None:
    styles_el = doc.styles.element
    s = sub(styles_el, "w:style", **{"w:type": "paragraph", "w:styleId": style_id})
    sub(s, "w:name", **{"w:val": style_id})
    if based_on is not None:
        sub(s, "w:basedOn", **{"w:val": based_on})
    if rpr_children:
        rpr = sub(s, "w:rPr")
        for tag, attrs in rpr_children:
            sub(rpr, tag, **(attrs or {}))


def test_provenance_is_dict_when_requested() -> None:
    doc = Document()
    p = doc.add_paragraph()
    resolved = resolve_effective_formatting(p, include_provenance=True)
    assert isinstance(resolved.provenance, dict)


def test_provenance_only_for_populated_fields() -> None:
    """Fields whose value is None must NOT appear in provenance."""
    doc = Document()
    p = doc.add_paragraph()
    resolved = resolve_effective_formatting(p, include_provenance=True)
    assert resolved.provenance is not None
    for fname, source in resolved.provenance.items():
        value = getattr(resolved, fname)
        assert value is not None, (
            f"field {fname!r} is None but appears in provenance ({source!r})"
        )


def test_doc_defaults_attributed_correctly() -> None:
    doc = Document()
    p = doc.add_paragraph()
    resolved = resolve_effective_formatting(p, include_provenance=True)
    assert resolved.provenance is not None
    source = resolved.provenance["font_size"]
    assert source.layer == "docDefaults"
    assert source.style_id is None


def test_paragraph_style_layer_records_style_id_and_chain_depth() -> None:
    doc = Document()
    _add_paragraph_style(doc, "Big", rpr_children=[("w:sz", {"w:val": "40"})])
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "Big"})

    resolved = resolve_effective_formatting(p, include_provenance=True)
    assert resolved.provenance is not None
    source = resolved.provenance["font_size"]
    assert source.layer == "paragraphStyle"
    assert source.style_id == "Big"
    assert source.chain_depth == 0


def test_chain_depth_increases_with_basedon_distance() -> None:
    """Property set at an ancestor style records its chain_depth."""
    doc = Document()
    # Parent sets size; child sets bold; grandchild empty.
    _add_paragraph_style(doc, "P", rpr_children=[("w:sz", {"w:val": "40"})])
    _add_paragraph_style(doc, "C", based_on="P", rpr_children=[("w:b", None)])
    _add_paragraph_style(doc, "GC", based_on="C")

    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "GC"})

    resolved = resolve_effective_formatting(p, include_provenance=True)
    assert resolved.provenance is not None
    # font_size came from P, which is 2 basedOn hops away.
    assert resolved.provenance["font_size"].style_id == "P"
    assert resolved.provenance["font_size"].chain_depth == 2
    # bold came from C, 1 hop away.
    assert resolved.provenance["bold"].style_id == "C"
    assert resolved.provenance["bold"].chain_depth == 1


def test_direct_paragraph_attributed_correctly() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:jc", **{"w:val": "center"})

    resolved = resolve_effective_formatting(p, include_provenance=True)
    assert resolved.provenance is not None
    assert resolved.provenance["alignment"] == FormattingSource(layer="directParagraph")


def test_direct_run_attributed_correctly() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("text")
    r.bold = True

    resolved = resolve_effective_formatting(r, include_provenance=True)
    assert resolved.provenance is not None
    bold_source = resolved.provenance["bold"]
    assert bold_source.layer == "directRun"


def test_toggle_resolved_flag_set_when_xor_happened() -> None:
    """Bold set in two layers triggers is_toggle_resolved=True at the last."""
    doc = Document()
    _add_paragraph_style(doc, "BoldP", rpr_children=[("w:b", None)])
    _add_paragraph_style(doc, "BoldC", based_on="BoldP", rpr_children=[("w:b", None)])
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "BoldC"})

    resolved = resolve_effective_formatting(p, include_provenance=True)
    assert resolved.provenance is not None
    bold_source = resolved.provenance["bold"]
    assert bold_source.is_toggle_resolved is True


def test_toggle_resolved_flag_unset_when_single_layer() -> None:
    """Bold set in just one layer does NOT count as toggle-resolved."""
    doc = Document()
    _add_paragraph_style(doc, "JustBold", rpr_children=[("w:b", None)])
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    sub(ppr, "w:pStyle", **{"w:val": "JustBold"})

    resolved = resolve_effective_formatting(p, include_provenance=True)
    assert resolved.provenance is not None
    bold_source = resolved.provenance["bold"]
    assert bold_source.is_toggle_resolved is False


def test_provenance_does_not_change_values(multistyle_docx_path: Path) -> None:
    """SPEC §4 invariant: values identical with or without provenance."""
    doc1 = Document(str(multistyle_docx_path))
    doc2 = Document(str(multistyle_docx_path))
    p1 = doc1.paragraphs[0]
    p2 = doc2.paragraphs[0]

    resolved_no = resolve_effective_formatting(p1, include_provenance=False)
    resolved_yes = resolve_effective_formatting(p2, include_provenance=True)

    # Every field except `provenance` itself must match exactly.
    for f in dc_fields(resolved_no):
        if f.name == "provenance":
            continue
        assert getattr(resolved_no, f.name) == getattr(resolved_yes, f.name), (
            f"field {f.name!r} differs between provenance on/off"
        )
