"""Tests for ``docx_plus.fields`` — complex-field insertion and update flag."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus._testing.ooxml_asserts import assert_field_dirty
from docx_plus.core.ns import XML, qn
from docx_plus.core.oxml import sub, xpath
from docx_plus.fields import (
    add_date_field,
    add_field,
    add_page_number_field,
    mark_fields_dirty,
)

# --------------------------------------------------------------------------
# Helpers to walk the run sequence a complex field produces.
# --------------------------------------------------------------------------


def _runs(paragraph) -> list:  # type: ignore[no-untyped-def]
    return list(xpath(paragraph._p, "./w:r"))


def _field_char_types(paragraph) -> list[str | None]:  # type: ignore[no-untyped-def]
    return [fc.get(qn("w:fldCharType")) for fc in xpath(paragraph._p, "./w:r/w:fldChar")]


def _instruction_text(paragraph) -> str:  # type: ignore[no-untyped-def]
    instr = xpath(paragraph._p, "./w:r/w:instrText")
    assert len(instr) == 1, f"expected exactly one w:instrText, got {len(instr)}"
    return instr[0].text or ""


def _result_text(paragraph) -> str:  # type: ignore[no-untyped-def]
    """Find the ``w:t`` between the ``separate`` and ``end`` fldChars."""
    runs = _runs(paragraph)
    sep_idx = None
    end_idx = None
    for i, r in enumerate(runs):
        for fc in r.findall(qn("w:fldChar")):
            ftype = fc.get(qn("w:fldCharType"))
            if ftype == "separate":
                sep_idx = i
            elif ftype == "end":
                end_idx = i
    assert sep_idx is not None and end_idx is not None
    parts: list[str] = []
    for r in runs[sep_idx + 1 : end_idx + 1]:
        for t in r.findall(qn("w:t")):
            parts.append(t.text or "")
    return "".join(parts)


# --------------------------------------------------------------------------
# add_page_number_field — structure, variants, format switch.
# --------------------------------------------------------------------------


def test_page_number_field_emits_five_runs() -> None:
    doc = Document()
    p = doc.add_paragraph("Page ")
    add_page_number_field(p)
    # "Page " is a sixth run holding the label text; ignore it.
    field_runs = _field_char_types(p)
    assert field_runs == ["begin", "separate", "end"]
    # 5 field runs + 1 label run = 6 total.
    assert len(_runs(p)) == 6


def test_page_number_field_default_instruction() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_page_number_field(p)
    assert _instruction_text(p) == " PAGE "


def test_page_number_field_numpages_variant() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_page_number_field(p, field="NUMPAGES")
    assert _instruction_text(p) == " NUMPAGES "


def test_page_number_field_sectionpages_variant() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_page_number_field(p, field="SECTIONPAGES")
    assert _instruction_text(p) == " SECTIONPAGES "


def test_page_number_field_with_format_switch() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_page_number_field(p, format=r"\* ARABIC")
    assert _instruction_text(p) == r" PAGE \* ARABIC "


def test_page_number_field_initial_text_is_one() -> None:
    """Word renders the cached result before recalculating; we seed "1"."""
    doc = Document()
    p = doc.add_paragraph()
    add_page_number_field(p)
    assert _result_text(p) == "1"


def test_page_number_instrtext_has_xml_space_preserve() -> None:
    """``xml:space="preserve"`` keeps Word from collapsing the surrounding spaces."""
    doc = Document()
    p = doc.add_paragraph()
    add_page_number_field(p)
    instr = xpath(p._p, "./w:r/w:instrText")[0]
    assert instr.get(f"{{{XML}}}space") == "preserve"


def test_page_number_field_returns_begin_run() -> None:
    doc = Document()
    p = doc.add_paragraph()
    begin_run = add_page_number_field(p)
    # The return is a <w:r>; its only child is <w:fldChar w:fldCharType="begin"/>.
    fld_chars = begin_run.findall(qn("w:fldChar"))
    assert len(fld_chars) == 1
    assert fld_chars[0].get(qn("w:fldCharType")) == "begin"


# --------------------------------------------------------------------------
# add_date_field — DATE vs CREATEDATE, format passthrough.
# --------------------------------------------------------------------------


def test_date_field_default_instruction() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_date_field(p)
    assert _instruction_text(p) == r' DATE \@ "MMMM d, yyyy" '


def test_date_field_custom_format() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_date_field(p, format="M/d/yyyy")
    assert _instruction_text(p) == r' DATE \@ "M/d/yyyy" '


def test_date_field_no_auto_update_uses_createdate() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_date_field(p, auto_update=False)
    assert _instruction_text(p) == r' CREATEDATE \@ "MMMM d, yyyy" '


def test_date_field_initial_text_is_empty() -> None:
    """Word fills the date on open; offline viewers see an empty result."""
    doc = Document()
    p = doc.add_paragraph()
    add_date_field(p)
    assert _result_text(p) == ""


# --------------------------------------------------------------------------
# add_field — generic passthrough.
# --------------------------------------------------------------------------


def test_generic_field_passthrough() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_field(p, instruction='TOC \\o "1-3" \\h')
    assert _instruction_text(p) == ' TOC \\o "1-3" \\h '


def test_generic_field_strips_user_supplied_spaces() -> None:
    """Trailing/leading whitespace is normalised so we never end up with double spaces."""
    doc = Document()
    p = doc.add_paragraph()
    add_field(p, instruction="  REF Bookmark1  ")
    assert _instruction_text(p) == " REF Bookmark1 "


def test_generic_field_initial_text_renders() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_field(p, instruction="REF Bookmark1", initial_text="(placeholder)")
    assert _result_text(p) == "(placeholder)"


@pytest.mark.parametrize("bad", ["", "   ", "\t", "\n"])
def test_add_field_rejects_empty_instruction(bad: str) -> None:
    """M1 regression: empty / whitespace-only instructions produce silent blanks."""
    doc = Document()
    p = doc.add_paragraph()
    with pytest.raises(ValueError, match="non-empty instruction"):
        add_field(p, instruction=bad)


def test_add_page_number_field_treats_empty_format_as_none() -> None:
    """M2 regression: format="" must not emit double spaces in the instruction."""
    doc = Document()
    p = doc.add_paragraph()
    add_page_number_field(p, format="")
    assert _instruction_text(p) == " PAGE "


def test_add_page_number_field_strips_format_whitespace() -> None:
    """Lead/trail whitespace in format is normalised — no double spaces."""
    doc = Document()
    p = doc.add_paragraph()
    add_page_number_field(p, format=r"  \* ARABIC  ")
    assert _instruction_text(p) == r" PAGE \* ARABIC "


# --------------------------------------------------------------------------
# Multiple fields, save/reopen round-trip.
# --------------------------------------------------------------------------


def test_two_fields_in_one_paragraph() -> None:
    doc = Document()
    p = doc.add_paragraph("Page ")
    add_page_number_field(p)
    p.add_run(" of ")
    add_page_number_field(p, field="NUMPAGES")
    # Two fields = two begin / separate / end triples.
    types = _field_char_types(p)
    assert types == ["begin", "separate", "end", "begin", "separate", "end"]


def test_page_number_field_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("Page ")
    add_page_number_field(p)
    out = tmp_path / "page.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    p2 = reopened.paragraphs[0]
    assert _field_char_types(p2) == ["begin", "separate", "end"]
    assert _instruction_text(p2) == " PAGE "
    assert _result_text(p2) == "1"


def test_date_field_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_date_field(p, format="M/d/yyyy", auto_update=False)
    out = tmp_path / "date.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    p2 = reopened.paragraphs[0]
    assert _instruction_text(p2) == r' CREATEDATE \@ "M/d/yyyy" '


# --------------------------------------------------------------------------
# mark_fields_dirty — happy path, idempotency, schema position.
# --------------------------------------------------------------------------


def test_mark_fields_dirty_inserts_update_fields() -> None:
    doc = Document()
    mark_fields_dirty(doc)
    assert_field_dirty(doc)


def test_mark_fields_dirty_is_idempotent() -> None:
    doc = Document()
    mark_fields_dirty(doc)
    mark_fields_dirty(doc)
    matches = xpath(doc.settings.element, "./w:updateFields")
    assert len(matches) == 1


def test_mark_fields_dirty_updates_existing_false_to_true() -> None:
    doc = Document()
    sub(doc.settings.element, "w:updateFields", **{"w:val": "false"})
    mark_fields_dirty(doc)
    matches = xpath(doc.settings.element, "./w:updateFields")
    assert len(matches) == 1
    assert matches[0].get(qn("w:val")) == "true"


def test_mark_fields_dirty_inserts_before_compat() -> None:
    """``w:compat`` comes after ``w:updateFields`` in CT_Settings.

    python-docx's default ``settings.xml`` already contains ``w:compat`` and
    ``w:rsids``; both are in our anchor list, so a blank-document call should
    place ``w:updateFields`` immediately before the first match (``w:compat``).
    """
    doc = Document()
    mark_fields_dirty(doc)
    compat = doc.settings.element.find(qn("w:compat"))
    assert compat is not None
    prev = compat.getprevious()
    assert prev is not None
    assert prev.tag == qn("w:updateFields")


def test_mark_fields_dirty_appends_when_no_anchor() -> None:
    """With no later-sibling anchor in settings.xml, we append at the end."""
    doc = Document()
    settings = doc.settings.element
    # Strip every anchor child so the fallback (append) path is exercised.
    for child in list(settings):
        if child.tag.startswith("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"):
            local = child.tag.rpartition("}")[2]
            if f"w:{local}" in {
                "w:hdrShapeDefaults",
                "w:footnotePr",
                "w:endnotePr",
                "w:compat",
                "w:docVars",
                "w:rsids",
                "w:mathPr",
                "w:themeFontLang",
                "w:clrSchemeMapping",
                "w:shapeDefaults",
                "w:decimalSymbol",
                "w:listSeparator",
            }:
                settings.remove(child)

    mark_fields_dirty(doc)
    new_element = settings.find(qn("w:updateFields"))
    assert new_element is not None
    assert list(settings)[-1] is new_element


def test_mark_fields_dirty_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("Page ")
    add_page_number_field(p)
    mark_fields_dirty(doc)
    out = tmp_path / "field_dirty.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    assert_field_dirty(reopened)
    p2 = reopened.paragraphs[0]
    assert _field_char_types(p2) == ["begin", "separate", "end"]
