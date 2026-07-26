"""Tests for ``docx_plus.core.oxml``."""

from __future__ import annotations

import pytest
from docx import Document
from lxml import etree

from docx_plus.core.ns import BUILD_NSMAP, W15, W, qn
from docx_plus.core.oxml import (
    _compile_xpath,
    body_document_for,
    build_bookmark,
    el,
    ordered_insert,
    remove,
    sub,
    xpath,
)


def test_el_creates_clark_tag() -> None:
    node = el("w:style")
    assert node.tag == f"{{{W}}}style"


def test_el_translates_namespaced_attribute_keys() -> None:
    node = el("w:style", **{"w:type": "paragraph", "w:styleId": "Foo"})
    assert node.get(qn("w:type")) == "paragraph"
    assert node.get(qn("w:styleId")) == "Foo"


def test_el_keeps_plain_attribute_keys() -> None:
    node = el("w:tag", id="bare")
    assert node.get("id") == "bare"
    assert node.get(qn("w:id")) is None


def test_el_declares_the_build_nsmap_for_document_namespaces() -> None:
    node = el("w:tag")
    assert set(node.nsmap) == set(BUILD_NSMAP) - {"xml"}
    assert "w15" not in node.nsmap


def test_el_declares_only_its_own_prefix_for_extension_namespaces() -> None:
    # w15 belongs to commentsExtended.xml. Declaring the document prefixes on
    # it would put five irrelevant xmlns attributes on every thread entry.
    node = el("w15:commentEx")
    assert node.nsmap == {"w15": W15}


def test_el_rejects_an_unknown_prefix_before_building() -> None:
    with pytest.raises(ValueError, match="unknown namespace prefix"):
        el("xyzzy:thing")


def test_sub_creates_and_appends() -> None:
    parent = el("w:styles")
    child = sub(parent, "w:style", **{"w:styleId": "Foo"})
    assert child in list(parent)
    assert child.getparent() is parent


def test_xpath_returns_matching_elements() -> None:
    parent = el("w:styles")
    sub(parent, "w:style", **{"w:styleId": "A"})
    sub(parent, "w:style", **{"w:styleId": "B"})
    sub(parent, "w:other")
    matches = xpath(parent, "./w:style")
    assert len(matches) == 2
    assert all(m.tag == f"{{{W}}}style" for m in matches)


def test_xpath_with_attribute_predicate() -> None:
    parent = el("w:styles")
    sub(parent, "w:style", **{"w:styleId": "A"})
    sub(parent, "w:style", **{"w:styleId": "B"})
    [match] = xpath(parent, "./w:style[@w:styleId='B']")
    assert match.get(qn("w:styleId")) == "B"


def test_remove_detaches_child() -> None:
    parent = el("w:styles")
    child = sub(parent, "w:style")
    assert child in list(parent)
    remove(child)
    assert child not in list(parent)
    assert child.getparent() is None


def test_remove_on_detached_node_is_noop() -> None:
    orphan = el("w:style")
    remove(orphan)
    assert orphan.getparent() is None


# --------------------------------------------------------------------------
# L11: xpath compiles each distinct expression once and caches it.
# --------------------------------------------------------------------------


def test_xpath_caches_compiled_expression() -> None:
    parent = el("w:styles")
    sub(parent, "w:style")
    first = _compile_xpath("./w:style")
    second = _compile_xpath("./w:style")
    assert first is second  # same compiled object reused


# --------------------------------------------------------------------------
# N4: body_document_for — shared proxy -> Document resolver.
# --------------------------------------------------------------------------


def test_body_document_for_returns_owning_document() -> None:
    doc = Document()
    p = doc.add_paragraph("x")
    # python-docx's DocumentPart.document builds a fresh proxy each call, so
    # compare the underlying element rather than proxy identity.
    assert body_document_for(p).element is doc.element


def test_body_document_for_rejects_non_body_proxy() -> None:
    class _FakePart:
        pass  # no .document attribute -> not the main body

    class _FakeProxy:
        part = _FakePart()

    with pytest.raises(ValueError, match="myop only supports the main document body"):
        body_document_for(_FakeProxy(), operation="myop")


# --------------------------------------------------------------------------
# ordered_insert — promoted out of styles/modify.py in v0.5 so numbering/
# can share it (SPEC §9.1 forbids the sibling import).
# --------------------------------------------------------------------------

_ORDER = ("first", "second", "third")


def test_ordered_insert_places_before_a_later_sibling() -> None:
    parent = el("w:parent")
    sub(parent, "w:third")
    ordered_insert(parent, el("w:second"), _ORDER)
    assert [qname_local(child) for child in parent] == ["second", "third"]


def test_ordered_insert_appends_when_no_later_sibling_exists() -> None:
    parent = el("w:parent")
    sub(parent, "w:first")
    ordered_insert(parent, el("w:third"), _ORDER)
    assert [qname_local(child) for child in parent] == ["first", "third"]


def test_ordered_insert_replaces_an_existing_same_tag() -> None:
    parent = el("w:parent")
    sub(parent, "w:second", **{"w:val": "old"})
    ordered_insert(parent, el("w:second", **{"w:val": "new"}), _ORDER)
    assert len(list(parent)) == 1
    assert parent[0].get(qn("w:val")) == "new"


def test_ordered_insert_appends_a_tag_absent_from_the_order() -> None:
    parent = el("w:parent")
    sub(parent, "w:first")
    ordered_insert(parent, el("w:unknown"), _ORDER)
    assert [qname_local(child) for child in parent] == ["first", "unknown"]


def test_ordered_insert_ignores_comment_nodes() -> None:
    """A comment's ``.tag`` is a callable, not a string — must not crash."""
    parent = el("w:parent")
    parent.append(etree.Comment("a note"))
    sub(parent, "w:third")
    ordered_insert(parent, el("w:first"), _ORDER)
    elements = [child for child in parent if isinstance(child.tag, str)]
    assert [qname_local(child) for child in elements] == ["first", "third"]


def qname_local(node: etree._Element) -> str:
    return etree.QName(node.tag).localname


# --------------------------------------------------------------------------
# build_bookmark — shared emitter so publishing/ can bookmark a caption
# without importing bookmarks/.
# --------------------------------------------------------------------------


def test_build_bookmark_brackets_the_anchors() -> None:
    doc = Document()
    p = doc.add_paragraph()
    first = sub(p._p, "w:r")
    last = sub(p._p, "w:r")

    start, end = build_bookmark(first, last, bookmark_id=7, name="fig_1")

    assert [qname_local(child) for child in p._p] == [
        "bookmarkStart",
        "r",
        "r",
        "bookmarkEnd",
    ]
    assert start.get(qn("w:id")) == "7"
    assert start.get(qn("w:name")) == "fig_1"
    assert end.get(qn("w:id")) == "7"


def test_build_bookmark_accepts_the_same_anchor_twice() -> None:
    doc = Document()
    p = doc.add_paragraph()
    only = sub(p._p, "w:r")

    build_bookmark(only, only, bookmark_id=1, name="solo")

    assert [qname_local(child) for child in p._p] == [
        "bookmarkStart",
        "r",
        "bookmarkEnd",
    ]
