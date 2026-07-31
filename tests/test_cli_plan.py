"""Tests for ``docx-plus plan``.

Drives :func:`docx_plus.cli.main` with an explicit ``argv`` and inspects
stdout and the exit code, like ``test_cli_lint.py``. The exit code carries
meaning here too: ``plan`` gates a pipeline on "is there anything a repair
pass would do", which is a different question from ``lint``'s "is there
anything to report".
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from docx_plus.cli import main
from docx_plus.lint import DEFAULT_PROFILE_NAME


@pytest.fixture
def messy_doc(tmp_path: Path) -> Path:
    """One document carrying a fixable defect, a withheld one, and an unfixable one."""
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Too deep", style="Heading 3")  # unfixable: which way to repair?
    doc.add_paragraph("Body with  two spaces.")  # fixable
    doc.add_paragraph("trailing space ")  # fixable
    doc.add_paragraph("")
    doc.add_paragraph("")  # withheld: deletes content
    doc.add_paragraph("after")
    path = tmp_path / "messy.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def clean_doc(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Title", style="Heading 1")
    doc.add_paragraph("An ordinary paragraph of prose.")
    path = tmp_path / "clean.docx"
    doc.save(str(path))
    return path


def test_plan_lists_the_edits_and_exits_one(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    code = main(["plan", str(messy_doc)])
    out = capsys.readouterr().out

    assert code == 1
    assert "in the order they would be applied" in out
    assert "double-space" in out
    assert "replace-paragraph-text" in out


def test_a_clean_document_plans_nothing_and_exits_zero(
    clean_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    code = main(["plan", str(clean_doc)])
    out = capsys.readouterr().out

    assert code == 0
    assert "No edits." in out
    assert "0 to apply, 0 withheld, 0 dropped, 0 unfixable." in out


def test_findings_nobody_can_repair_do_not_fail_the_gate(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """A skipped outline level is real, and not something a repair pass fixes."""
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Too deep", style="Heading 3")
    path = tmp_path / "outline.docx"
    doc.save(str(path))

    code = main(["plan", str(path)])
    out = capsys.readouterr().out

    assert code == 0
    assert "1 finding(s) with no known repair" in out
    assert "heading-level-skip" in out


def test_content_changing_edits_are_listed_separately(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    main(["plan", str(messy_doc), "--rule", "typography"])
    out = capsys.readouterr().out

    assert "these change what the document contains" in out
    assert "--allow-content" in out


def test_allow_content_promotes_them_into_the_plan(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    main(["plan", str(messy_doc), "--rule", "typography", "--allow-content"])
    out = capsys.readouterr().out

    assert "delete-paragraph" in out
    assert "3 to apply, 0 withheld, 0 dropped, 0 unfixable." in out


def test_conflicting_edits_are_reported_not_hidden(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """A double space that runs up against a comma: both rules claim the space."""
    doc = Document()
    doc.add_paragraph("Alpha  , beta")
    path = tmp_path / "conflict.docx"
    doc.save(str(path))

    main(["plan", str(path)])
    out = capsys.readouterr().out

    assert "conflict(s)" in out
    assert "claimed it first" in out


def test_json_output_is_the_serialized_plan(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    main(["plan", str(messy_doc), "--json"])
    payload = json.loads(capsys.readouterr().out)

    assert set(payload) == {"fixes", "deferred", "conflicts", "unfixable"}
    assert payload["fixes"]
    first = payload["fixes"][0]
    assert set(first["fix"]) == {"summary", "safety", "operations"}
    assert first["fix"]["operations"][0]["op"]


def test_a_property_list_renders_in_the_operation_line(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """The plan has to say *which* properties would be deleted, not just how many."""
    from docx.shared import Pt

    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.font.size = Pt(11)  # already what Normal says
    path = tmp_path / "redundant.docx"
    doc.save(str(path))

    main(["plan", str(path), "--rule", "redundant-direct-formatting"])
    out = capsys.readouterr().out

    assert "[safe ]" in out
    assert "clear-run-properties" in out
    assert "properties=font_size" in out


def test_plan_output_is_cp1252_safe(messy_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """Text output must survive a default Windows console."""
    main(["plan", str(messy_doc), "--allow-content"])

    capsys.readouterr().out.encode("cp1252")


def test_rule_selection_narrows_the_plan(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    main(["plan", str(messy_doc), "--rule", "double-space"])
    out = capsys.readouterr().out

    assert "double-space" in out
    assert "trailing-whitespace" not in out


def test_exclude_removes_a_rule_from_the_plan(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    main(["plan", str(messy_doc), "--exclude", "double-space"])
    out = capsys.readouterr().out

    assert "double-space" not in out


def test_a_missing_file_is_a_clean_error(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    code = main(["plan", str(tmp_path / "nope.docx")])

    assert code == 1
    assert "error:" in capsys.readouterr().err


# --------------------------------------------------------------------------
# Profiles, which both lint and plan take.
# --------------------------------------------------------------------------


def _write_profile(directory: Path, payload: object, name: str = DEFAULT_PROFILE_NAME) -> Path:
    path = directory / name
    path.write_text(json.dumps(payload), encoding="utf-8")
    return path


def test_a_profile_beside_the_document_is_discovered(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    _write_profile(messy_doc.parent, {"rules": {"double-space": {"enabled": False}}})

    main(["plan", str(messy_doc)])

    assert "double-space" not in capsys.readouterr().out


def test_no_profile_ignores_a_discovered_one(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    _write_profile(messy_doc.parent, {"rules": {"double-space": {"enabled": False}}})

    main(["plan", str(messy_doc), "--no-profile"])

    assert "double-space" in capsys.readouterr().out


def test_an_explicit_profile_path_is_used(
    messy_doc: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    elsewhere = tmp_path / "config"
    elsewhere.mkdir()
    path = _write_profile(elsewhere, {"rules": {"double-space": {"enabled": False}}}, "house.json")

    main(["plan", str(messy_doc), "--profile", str(path)])

    assert "double-space" not in capsys.readouterr().out


def test_a_missing_profile_path_is_an_error_not_a_fallback(
    messy_doc: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """Asking for a specific profile and quietly getting a different one is the worst case."""
    code = main(["plan", str(messy_doc), "--profile", str(tmp_path / "absent.json")])

    assert code == 1
    assert "error:" in capsys.readouterr().err


def test_a_broken_profile_is_reported_by_the_cli(
    messy_doc: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    (messy_doc.parent / DEFAULT_PROFILE_NAME).write_text("{", encoding="utf-8")

    code = main(["plan", str(messy_doc)])

    assert code == 1
    assert "InvalidProfileError" in capsys.readouterr().err


def test_lint_takes_a_profile_too(messy_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    """Both commands answer the same question, so both take the same selectors."""
    _write_profile(messy_doc.parent, {"rules": {"double-space": {"severity": "error"}}})

    main(["lint", str(messy_doc), "--json"])
    findings = json.loads(capsys.readouterr().out)

    assert [f["severity"] for f in findings if f["rule"] == "double-space"] == ["error"]


def test_lint_json_reports_fixability(messy_doc: Path, capsys: pytest.CaptureFixture[str]) -> None:
    main(["lint", str(messy_doc), "--json"])
    findings = json.loads(capsys.readouterr().out)

    by_rule = {f["rule"]: f for f in findings}
    assert by_rule["double-space"]["fixable"] is True
    assert by_rule["heading-level-skip"]["fixable"] is False
