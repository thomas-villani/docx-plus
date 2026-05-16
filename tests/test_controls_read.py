"""Tests for ``docx_plus.controls.read``."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from docx_plus.controls import (
    ControlNotFoundError,
    ControlTypeError,
    DuplicateTagError,
    FormBuilder,
    ValueNotInListError,
    clear_control,
    read_controls,
    set_control_value,
)
from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub, xpath


# --------------------------------------------------------------------------
# Helpers.
# --------------------------------------------------------------------------


def _build_form(tmp_path: Path) -> Document:
    """Build a form with one of every control type and reload it."""
    fb = FormBuilder()
    p = fb.doc.add_paragraph("Name: ")
    fb.add_text_control(p, tag="name", alias="Full name")

    p2 = fb.doc.add_paragraph("Region: ")
    fb.add_dropdown(p2, tag="region", items=[("North", "N"), ("South", "S")])

    p3 = fb.doc.add_paragraph("Referral: ")
    fb.add_dropdown(p3, tag="referral", items=["Web", "Friend"], editable=True)

    p4 = fb.doc.add_paragraph("Visit: ")
    fb.add_date_picker(p4, tag="visit", date_format="M/d/yyyy")

    p5 = fb.doc.add_paragraph("Subscribe: ")
    fb.add_checkbox(p5, tag="subscribe")

    out = tmp_path / "form.docx"
    fb.save(out)
    return Document(out)


def _round_trip(doc: Document, tmp_path: Path) -> Document:
    out = tmp_path / "rt.docx"
    doc.save(out)
    return Document(out)


# --------------------------------------------------------------------------
# read_controls — fresh form, placeholder state.
# --------------------------------------------------------------------------


def test_read_controls_all_present(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    controls = read_controls(doc)
    assert set(controls.keys()) == {"name", "region", "referral", "visit", "subscribe"}


def test_read_controls_types(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    controls = read_controls(doc)
    assert controls["name"].control_type == "text"
    assert controls["region"].control_type == "dropdown"
    assert controls["referral"].control_type == "combobox"
    assert controls["visit"].control_type == "date"
    assert controls["subscribe"].control_type == "checkbox"


def test_read_controls_placeholder_state(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    controls = read_controls(doc)
    # text/dropdown/combobox/date are in placeholder state; value is None.
    for key in ("name", "region", "referral", "visit"):
        assert controls[key].is_placeholder is True
        assert controls[key].value is None
    # Checkboxes have no placeholder concept; value is False (default unchecked).
    assert controls["subscribe"].is_placeholder is False
    assert controls["subscribe"].value is False


def test_read_controls_alias_propagated(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    controls = read_controls(doc)
    assert controls["name"].alias == "Full name"
    assert controls["region"].alias is None


# --------------------------------------------------------------------------
# read_controls — by="alias".
# --------------------------------------------------------------------------


def test_read_controls_by_alias_skips_unaliased(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    controls = read_controls(doc, by="alias")
    assert set(controls.keys()) == {"Full name"}
    assert controls["Full name"].tag == "name"


# --------------------------------------------------------------------------
# Duplicate-tag detection.
# --------------------------------------------------------------------------


def test_read_controls_duplicate_tag_raises() -> None:
    fb = FormBuilder()
    p = fb.doc.add_paragraph()
    fb.add_text_control(p, tag="dup")
    fb.add_text_control(p, tag="dup")
    with pytest.raises(DuplicateTagError, match="dup"):
        read_controls(fb.doc)


# --------------------------------------------------------------------------
# set_control_value — round-trips.
# --------------------------------------------------------------------------


def test_set_text_round_trip(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "name", "Ada Lovelace")
    reloaded = _round_trip(doc, tmp_path)
    controls = read_controls(reloaded)
    assert controls["name"].value == "Ada Lovelace"
    assert controls["name"].is_placeholder is False


def test_set_dropdown_by_value(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "region", "N")
    reloaded = _round_trip(doc, tmp_path)
    controls = read_controls(reloaded)
    assert controls["region"].value == "North"  # the displayText is rendered
    assert controls["region"].is_placeholder is False


def test_set_dropdown_by_display_text(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "region", "South")
    reloaded = _round_trip(doc, tmp_path)
    assert read_controls(reloaded)["region"].value == "South"


def test_set_dropdown_no_match_raises(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    with pytest.raises(ValueNotInListError, match="region"):
        set_control_value(doc, "region", "Mars")


def test_set_combobox_freeform_passthrough(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "referral", "LinkedIn")
    reloaded = _round_trip(doc, tmp_path)
    assert read_controls(reloaded)["referral"].value == "LinkedIn"


def test_set_combobox_match_uses_display_text(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "referral", "Web")
    reloaded = _round_trip(doc, tmp_path)
    assert read_controls(reloaded)["referral"].value == "Web"


def test_set_date_round_trip(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "visit", datetime(2026, 5, 15))
    reloaded = _round_trip(doc, tmp_path)
    controls = read_controls(reloaded)
    assert controls["visit"].value == datetime(2026, 5, 15)
    assert controls["visit"].is_placeholder is False


def test_set_checkbox_true(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "subscribe", True)
    reloaded = _round_trip(doc, tmp_path)
    assert read_controls(reloaded)["subscribe"].value is True
    glyph_t = xpath(reloaded.element.body, ".//w:sdt//w:t")[-1]
    assert glyph_t.text == "☒"


def test_set_checkbox_false(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "subscribe", True)
    set_control_value(doc, "subscribe", False)
    reloaded = _round_trip(doc, tmp_path)
    assert read_controls(reloaded)["subscribe"].value is False


# --------------------------------------------------------------------------
# clear_control.
# --------------------------------------------------------------------------


def test_clear_text_control_restores_placeholder(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "name", "Ada")
    clear_control(doc, "name")
    reloaded = _round_trip(doc, tmp_path)
    info = read_controls(reloaded)["name"]
    assert info.is_placeholder is True
    assert info.value is None


def test_clear_dropdown_restores_placeholder(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "region", "N")
    clear_control(doc, "region")
    reloaded = _round_trip(doc, tmp_path)
    info = read_controls(reloaded)["region"]
    assert info.is_placeholder is True


def test_clear_checkbox_unchecks(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    set_control_value(doc, "subscribe", True)
    clear_control(doc, "subscribe")
    reloaded = _round_trip(doc, tmp_path)
    assert read_controls(reloaded)["subscribe"].value is False


# --------------------------------------------------------------------------
# Error cases.
# --------------------------------------------------------------------------


def test_set_unknown_tag_raises(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    with pytest.raises(ControlNotFoundError, match="missing_tag"):
        set_control_value(doc, "missing_tag", "x")


def test_set_text_with_bool_raises(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    with pytest.raises(ControlTypeError):
        set_control_value(doc, "name", True)  # type: ignore[arg-type]


def test_set_checkbox_with_string_raises(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    with pytest.raises(ControlTypeError):
        set_control_value(doc, "subscribe", "yes")  # type: ignore[arg-type]


def test_set_date_with_string_raises(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    with pytest.raises(ControlTypeError):
        set_control_value(doc, "visit", "2026-05-15")  # type: ignore[arg-type]


def test_clear_unknown_tag_raises(tmp_path: Path) -> None:
    doc = _build_form(tmp_path)
    with pytest.raises(ControlNotFoundError):
        clear_control(doc, "missing")


# --------------------------------------------------------------------------
# Read on externally-built docs (no FormBuilder).
# --------------------------------------------------------------------------


def test_read_external_form(existing_form_docx_path: Path) -> None:
    doc = Document(existing_form_docx_path)
    controls = read_controls(doc)
    assert set(controls.keys()) == {"name", "region", "subscribe"}
    assert controls["name"].control_type == "text"
    assert controls["name"].value == "Ada Lovelace"
    assert controls["name"].is_placeholder is False
    assert controls["region"].control_type == "dropdown"
    assert controls["region"].is_placeholder is True
    assert controls["region"].alias == "Region selector"
    assert controls["subscribe"].control_type == "checkbox"
    assert controls["subscribe"].value is True


def test_set_external_form_value(
    existing_form_docx_path: Path, tmp_path: Path
) -> None:
    doc = Document(existing_form_docx_path)
    set_control_value(doc, "region", "N")
    out = tmp_path / "modified.docx"
    doc.save(out)
    reloaded = Document(out)
    assert read_controls(reloaded)["region"].value == "North"


# --------------------------------------------------------------------------
# Skip rich-text SDTs gracefully.
# --------------------------------------------------------------------------


def test_unrecognised_sdt_is_skipped(tmp_path: Path) -> None:
    """Rich-text SDTs (no recognised type marker) must be silently skipped."""
    doc = Document()
    p = doc.add_paragraph()
    sdt = sub(p._p, "w:sdt")
    pr = sub(sdt, "w:sdtPr")
    sub(pr, "w:tag", **{"w:val": "rich"})
    sub(pr, "w:id", **{"w:val": "1"})
    # No type marker.
    sub(sdt, "w:sdtContent")

    controls = read_controls(doc)
    assert controls == {}
