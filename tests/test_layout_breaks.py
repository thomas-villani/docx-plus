"""Tests for ``docx_plus.layout.insert_section_break``."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.section import Section

from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath
from docx_plus.layout import insert_section_break


def test_insert_section_break_returns_section() -> None:
    doc = Document()
    doc.add_paragraph("a")
    p = doc.add_paragraph("b")
    doc.add_paragraph("c")
    result = insert_section_break(p)
    assert isinstance(result, Section)


def test_insert_section_break_doubles_section_count() -> None:
    doc = Document()
    doc.add_paragraph("a")
    p = doc.add_paragraph("b")
    doc.add_paragraph("c")
    initial = len(doc.sections)
    insert_section_break(p)
    assert len(doc.sections) == initial + 1


def test_insert_section_break_writes_pPr_sectPr_on_target() -> None:
    doc = Document()
    doc.add_paragraph("a")
    p = doc.add_paragraph("b")
    insert_section_break(p)
    nested = xpath(p._p, "./w:pPr/w:sectPr")
    assert len(nested) == 1


def test_insert_section_break_default_type_is_nextPage() -> None:
    doc = Document()
    p = doc.add_paragraph("b")
    insert_section_break(p)
    type_el = xpath(p._p, "./w:pPr/w:sectPr/w:type")
    assert len(type_el) == 1
    assert type_el[0].get(qn("w:val")) == "nextPage"


def test_insert_section_break_continuous() -> None:
    doc = Document()
    p = doc.add_paragraph("b")
    insert_section_break(p, start_type="continuous")
    type_el = xpath(p._p, "./w:pPr/w:sectPr/w:type")[0]
    assert type_el.get(qn("w:val")) == "continuous"


def test_insert_section_break_even_page() -> None:
    doc = Document()
    p = doc.add_paragraph("b")
    insert_section_break(p, start_type="evenPage")
    type_el = xpath(p._p, "./w:pPr/w:sectPr/w:type")[0]
    assert type_el.get(qn("w:val")) == "evenPage"


def test_insert_section_break_inherits_page_size() -> None:
    """The cloned sectPr should carry the original's pgSz."""
    doc = Document()
    p = doc.add_paragraph("b")
    insert_section_break(p)
    pg_sz_nested = xpath(p._p, "./w:pPr/w:sectPr/w:pgSz")
    pg_sz_trailing = xpath(doc.element.body, "./w:sectPr/w:pgSz")
    assert len(pg_sz_nested) == 1
    assert len(pg_sz_trailing) == 1
    # Same width / height — properties were cloned.
    assert pg_sz_nested[0].get(qn("w:w")) == pg_sz_trailing[0].get(qn("w:w"))


def test_insert_section_break_returned_section_is_mutable() -> None:
    """The returned Section should let callers mutate the new section."""
    from docx.shared import Inches

    doc = Document()
    p = doc.add_paragraph("b")
    new_section = insert_section_break(p)
    new_section.left_margin = Inches(2)
    # Mutation lands in the paragraph-nested sectPr. ``w:left`` stores
    # twips (1/1440 inch), so 2 inches == 2880.
    pg_mar = xpath(p._p, "./w:pPr/w:sectPr/w:pgMar")
    assert len(pg_mar) == 1
    assert pg_mar[0].get(qn("w:left")) == "2880"


def test_insert_section_break_idempotent_replaces_type() -> None:
    doc = Document()
    p = doc.add_paragraph("b")
    insert_section_break(p, start_type="nextPage")
    insert_section_break(p, start_type="continuous")
    type_els = xpath(p._p, "./w:pPr/w:sectPr/w:type")
    assert len(type_els) == 1
    assert type_els[0].get(qn("w:val")) == "continuous"


def test_insert_section_break_round_trip(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("a")
    p = doc.add_paragraph("b")
    doc.add_paragraph("c")
    insert_section_break(p, start_type="continuous")
    out = tmp_path / "break.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    assert len(reopened.sections) == 2
    p_split = reopened.paragraphs[1]
    type_el = xpath(p_split._p, "./w:pPr/w:sectPr/w:type")[0]
    assert type_el.get(qn("w:val")) == "continuous"


def test_insert_section_break_places_type_after_header_reference() -> None:
    """M7: w:type must follow headerReference per ECMA-376 17.6.17.

    The pre-fix code jammed w:type to position 0, ahead of any header
    references the cloned sectPr carried — schema-invalid output.
    """
    import lxml.etree as ET

    doc = Document()
    # Give the section a header so its sectPr carries a headerReference.
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    header.paragraphs[0].add_run("H")

    doc.add_paragraph("intro")
    p = doc.add_paragraph("split")
    doc.add_paragraph("after")
    insert_section_break(p, start_type="continuous")

    sect_pr = xpath(p._p, "./w:pPr/w:sectPr")[0]
    locals_ = [ET.QName(c.tag).localname for c in sect_pr]
    assert "headerReference" in locals_ and "type" in locals_
    assert locals_.index("headerReference") < locals_.index("type")
    # And w:type still precedes pgSz.
    assert locals_.index("type") < locals_.index("pgSz")


def test_insert_section_break_requires_body_parent() -> None:
    """A paragraph outside <w:body> — here a real header paragraph — is rejected.

    The header paragraph's element is parented to ``<w:hdr>``, not
    ``<w:body>``, so the body-parent guard fires on a genuine non-body
    proxy rather than a hand-built fake that the code never consults.
    """
    doc = Document()
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]

    with pytest.raises(ValueError, match="main document body"):
        insert_section_break(header_paragraph)
