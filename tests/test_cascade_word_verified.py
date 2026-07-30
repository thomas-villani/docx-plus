"""The toggle rule as measured against live Microsoft Word.

Every expected value in this file was read out of Word over COM, not derived
from the spec: a probe document was built for each cascade shape, opened in
Word, and its effective formatting recorded. ECMA-376 17.7.3's prose is
ambiguous enough that several readings survive a casual sample, and the
resolver previously shipped one of the wrong ones.

The table is deliberately exhaustive over the *boundaries* rather than over
the properties — bold and italic stand in for all twelve toggles (which
:mod:`tests.test_cascade_toggles` checks behave alike). What varies here is
where in the cascade the values sit, because that is what the rule is about.

Re-measuring: build a document with the same shapes, open it in Word, and
read ``wordlive read format --anchor-id para:N``.
"""

from __future__ import annotations

import pytest
from docx import Document

from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub
from docx_plus.styles import resolve_effective_formatting

# Each case: (id, docDefaults value, table style, paragraph style,
# character style, direct run, what Word rendered).
#
# ``None`` means the layer says nothing; True/False mean it specifies
# ``<w:b/>`` or ``<w:b w:val="0"/>`` respectively.
Spec = bool | None
CASES: list[tuple[str, Spec, Spec, Spec, Spec, Spec, bool]] = [
    # ---- nothing, and one level alone -----------------------------------
    ("nothing-anywhere", None, None, None, None, None, False),
    ("paragraph-style-only", None, None, True, None, None, True),
    ("character-style-only", None, None, None, True, None, True),
    ("table-style-only", None, True, None, None, None, True),
    ("doc-defaults-only", True, None, None, None, None, True),
    # ---- two style levels cancel ----------------------------------------
    ("table-and-paragraph", None, True, True, None, None, False),
    ("table-and-character", None, True, None, True, None, False),
    ("paragraph-and-character", None, None, True, True, None, False),
    # ---- three style levels: odd parity -----------------------------------
    ("all-three-levels", None, True, True, True, None, True),
    # ---- an "off" matching the base is inert ------------------------------
    ("paragraph-on-character-off", None, None, True, False, None, True),
    ("table-on-paragraph-off", None, True, False, None, None, True),
    ("table-on-character-off", None, True, None, False, None, True),
    ("paragraph-off-character-on", None, None, False, True, None, True),
    ("table-off-paragraph-on", None, False, True, None, None, True),
    ("paragraph-off-alone", None, None, False, None, None, False),
    ("character-off-alone", None, None, None, False, None, False),
    ("table-off-alone", None, False, None, None, None, False),
    ("table-on-paragraph-off-character-on", None, True, False, True, None, False),
    # ---- docDefaults is the base the levels count against -----------------
    ("base-on-paragraph-restates", True, None, True, None, None, True),
    ("base-on-character-restates", True, None, None, True, None, True),
    ("base-on-two-levels-restate", True, None, True, True, None, True),
    ("base-on-paragraph-contradicts", True, None, False, None, None, False),
    ("base-on-character-contradicts", True, None, None, False, None, False),
    ("base-on-paragraph-on-character-off", True, None, True, False, None, False),
    # ---- direct formatting is absolute ------------------------------------
    ("direct-on-over-nothing", None, None, None, None, True, True),
    ("direct-on-over-bold-paragraph", None, None, True, None, True, True),
    ("direct-on-over-bold-character", None, None, None, True, True, True),
    ("direct-off-over-bold-paragraph", None, None, True, None, False, False),
    ("direct-on-over-cancelled-pair", None, None, True, True, True, True),
    ("direct-on-over-italic-base", True, None, None, None, True, True),
]


def _spec(parent: object, tag: str, value: Spec) -> None:
    """Write ``<w:b/>`` / ``<w:b w:val="0"/>``, or nothing at all."""
    if value is None:
        return
    sub(parent, tag, **({} if value else {"w:val": "0"}))  # type: ignore[arg-type]


def _build(
    tag: str,
    base: Spec,
    table: Spec,
    paragraph: Spec,
    character: Spec,
    direct: Spec,
) -> object:
    """Assemble a document matching one row and return the run to resolve."""
    doc = Document()
    for stray in list(doc.paragraphs):
        stray._p.getparent().remove(stray._p)

    defaults = doc.styles.element.find(qn("w:docDefaults"))
    _spec(defaults.find(qn("w:rPrDefault")).find(qn("w:rPr")), tag, base)

    para_style = sub(
        doc.styles.element,
        "w:style",
        **{"w:type": "paragraph", "w:styleId": "P", "w:customStyle": "1"},
    )
    sub(para_style, "w:name", **{"w:val": "P"})
    sub(para_style, "w:basedOn", **{"w:val": "Normal"})
    _spec(sub(para_style, "w:rPr"), tag, paragraph)

    char_style = sub(
        doc.styles.element,
        "w:style",
        **{"w:type": "character", "w:styleId": "C", "w:customStyle": "1"},
    )
    sub(char_style, "w:name", **{"w:val": "C"})
    _spec(sub(char_style, "w:rPr"), tag, character)

    if table is not None:
        table_style = sub(
            doc.styles.element,
            "w:style",
            **{"w:type": "table", "w:styleId": "T", "w:customStyle": "1"},
        )
        sub(table_style, "w:name", **{"w:val": "T"})
        _spec(sub(table_style, "w:rPr"), tag, table)
        container = doc.add_table(rows=1, cols=1)
        # Written straight into tblPr rather than via ``table.style``:
        # python-docx's style factory only understands elements it built
        # itself, and these styles are assembled as raw XML.
        tbl_pr = container._tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = sub(container._tbl, "w:tblPr")
        sub(tbl_pr, "w:tblStyle", **{"w:val": "T"})
        cell = container.cell(0, 0)
        para = cell.add_paragraph()
        for stray in list(cell.paragraphs):
            if stray._p is not para._p:
                stray._p.getparent().remove(stray._p)
    else:
        para = doc.add_paragraph()

    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "P"})
    run = para.add_run("text")
    if character is not None:
        sub(run._r.get_or_add_rPr(), "w:rStyle", **{"w:val": "C"})
    _spec(run._r.get_or_add_rPr(), tag, direct)
    return run


@pytest.mark.parametrize(
    ("case_id", "base", "table", "paragraph", "character", "direct", "expected"),
    CASES,
    ids=[case[0] for case in CASES],
)
def test_matches_word(
    case_id: str,
    base: Spec,
    table: Spec,
    paragraph: Spec,
    character: Spec,
    direct: Spec,
    expected: bool,
) -> None:
    run = _build("w:b", base, table, paragraph, character, direct)
    resolved = resolve_effective_formatting(run)

    # Word reports an unset toggle as off; the resolver distinguishes the
    # two, so normalise before comparing against what Word rendered.
    assert bool(resolved.bold) is expected


@pytest.mark.parametrize(
    ("case_id", "base", "table", "paragraph", "character", "direct", "expected"),
    CASES,
    ids=[case[0] for case in CASES],
)
def test_italic_behaves_identically_to_bold(
    case_id: str,
    base: Spec,
    table: Spec,
    paragraph: Spec,
    character: Spec,
    direct: Spec,
    expected: bool,
) -> None:
    """The rule is a property of toggles, not of ``w:b`` in particular."""
    run = _build("w:i", base, table, paragraph, character, direct)

    assert bool(resolve_effective_formatting(run).italic) is expected


def test_unset_stays_none_rather_than_false() -> None:
    """The one place the resolver deliberately says more than Word can.

    Word's object model cannot distinguish "nothing set bold" from "something
    set it off"; the resolver can, and the linter needs it to.
    """
    run = _build("w:b", None, None, None, None, None)

    assert resolve_effective_formatting(run).bold is None


def test_a_link_partner_contributes_nothing() -> None:
    """Word never consults the ``w:link`` partner when rendering a run."""
    doc = Document()
    para_style = sub(
        doc.styles.element,
        "w:style",
        **{"w:type": "paragraph", "w:styleId": "Linked", "w:customStyle": "1"},
    )
    sub(para_style, "w:name", **{"w:val": "Linked"})
    sub(para_style, "w:basedOn", **{"w:val": "Normal"})
    sub(para_style, "w:link", **{"w:val": "LinkedChar"})

    char_style = sub(
        doc.styles.element,
        "w:style",
        **{"w:type": "character", "w:styleId": "LinkedChar", "w:customStyle": "1"},
    )
    sub(char_style, "w:name", **{"w:val": "Linked Char"})
    sub(char_style, "w:link", **{"w:val": "Linked"})
    linked_rpr = sub(char_style, "w:rPr")
    sub(linked_rpr, "w:b")
    sub(linked_rpr, "w:sz", **{"w:val": "48"})

    para = doc.add_paragraph()
    sub(sub(para._p, "w:pPr"), "w:pStyle", **{"w:val": "Linked"})
    resolved = resolve_effective_formatting(para.add_run("text"))

    assert resolved.bold is None
    assert resolved.font_size == 11.0


def test_a_numbering_levels_rpr_contributes_nothing() -> None:
    """``w:lvl/w:rPr`` formats the glyph; ``w:lvl/w:pPr`` formats the paragraph."""
    doc = Document()
    numbering = doc.part.numbering_part.element
    abstract = sub(numbering, "w:abstractNum", **{"w:abstractNumId": "700"})
    sub(abstract, "w:multiLevelType", **{"w:val": "hybridMultilevel"})
    lvl = sub(abstract, "w:lvl", **{"w:ilvl": "0"})
    sub(lvl, "w:start", **{"w:val": "1"})
    sub(lvl, "w:numFmt", **{"w:val": "decimal"})
    sub(lvl, "w:lvlText", **{"w:val": "%1."})
    sub(lvl, "w:lvlJc", **{"w:val": "left"})
    sub(sub(lvl, "w:pPr"), "w:ind", **{"w:left": "720"})
    sub(sub(lvl, "w:rPr"), "w:b")
    num = sub(numbering, "w:num", **{"w:numId": "700"})
    sub(num, "w:abstractNumId", **{"w:val": "700"})

    para = doc.add_paragraph("item")
    num_pr = sub(sub(para._p, "w:pPr"), "w:numPr")
    sub(num_pr, "w:ilvl", **{"w:val": "0"})
    sub(num_pr, "w:numId", **{"w:val": "700"})

    resolved = resolve_effective_formatting(para)

    assert resolved.indent_left == 720
    assert resolved.bold is None


def test_a_stock_heading_is_bold_at_both_targets() -> None:
    """The shape every one of these rules exists to get right."""
    doc = Document()
    para = doc.add_paragraph("Heading", style="Heading 1")

    assert resolve_effective_formatting(para).bold is True
    assert resolve_effective_formatting(para.runs[0]).bold is True
    assert resolve_effective_formatting(para.runs[0]).font_size == 14.0
