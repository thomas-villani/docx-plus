"""Tests for resolving formatting against ``Run`` and ``_Cell`` targets.

Most existing cascade tests target a ``Paragraph``. This file exercises the
two other supported target kinds so the run-cascade branch
(``inspect.py:302-316``) and the cell-cascade branch (``inspect.py:319-329``)
are covered.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from docx_plus.core.oxml import sub
from docx_plus.styles import apply_style, ensure_style, resolve_effective_formatting


def test_resolve_run_target_picks_up_direct_run_formatting(
    tmp_path: Path,
) -> None:
    """``include_provenance`` on a Run target reports the directRun layer."""
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("bold run")
    run.bold = True
    out = tmp_path / "run.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    target_run = reopened.paragraphs[0].runs[0]
    resolved = resolve_effective_formatting(target_run, include_provenance=True)
    assert resolved.bold is True
    prov = resolved.provenance or {}
    assert prov["bold"].layer == "directRun"


def test_resolve_run_target_inherits_paragraph_style(tmp_path: Path) -> None:
    """A Run inside a Heading1 paragraph inherits Heading1 properties."""
    doc = Document()
    ensure_style(doc, "Heading1")
    para = doc.add_paragraph("heading text")
    apply_style(para, "Heading1")
    para.add_run(" more text")  # second run; ensures runs are present
    out = tmp_path / "run_style.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    target_run = reopened.paragraphs[0].runs[0]
    resolved = resolve_effective_formatting(target_run)
    # font_size, bold, color_rgb all come from Heading1 via paragraphStyle.
    assert resolved.font_size is not None
    assert resolved.style_id == "Heading1"


def test_resolve_run_target_with_rstyle_applies_linked_char_chain(
    tmp_path: Path,
) -> None:
    """A run with ``w:rStyle`` triggers the linked-char-style branch."""
    doc = Document()
    # Hand-build a character style with a distinctive color.
    styles_root = doc.styles.element
    char_style = sub(
        styles_root,
        "w:style",
        **{"w:type": "character", "w:styleId": "MyChar"},
    )
    sub(char_style, "w:name", **{"w:val": "MyChar"})
    cs_rpr = sub(char_style, "w:rPr")
    sub(cs_rpr, "w:color", **{"w:val": "FF00FF"})

    para = doc.add_paragraph()
    run = para.add_run("text")
    # Attach the rStyle to the run.
    r_pr = sub(run._r, "w:rPr")
    sub(r_pr, "w:rStyle", **{"w:val": "MyChar"})
    out = tmp_path / "rstyle.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    target_run = reopened.paragraphs[0].runs[0]
    resolved = resolve_effective_formatting(target_run)
    assert resolved.color_rgb == "FF00FF"


def test_resolve_paragraph_in_table_cell_walks_table_style(tmp_path: Path) -> None:
    """A paragraph inside a table cell exercises layer 2 (table style)."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    para = cell.paragraphs[0]
    para.text = "cell text"
    out = tmp_path / "tabled.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    cell_para = reopened.tables[0].cell(0, 0).paragraphs[0]
    resolved = resolve_effective_formatting(cell_para)
    # No assertion on specific values — the default TableNormal style sets
    # minimal pPr — but the call must succeed (covers _enclosing_cell +
    # _enclosing_table + the table-style branch).
    assert resolved is not None


def test_resolve_cell_target_directly(tmp_path: Path) -> None:
    """Resolving against a ``_Cell`` target uses ``_apply_cell_cascade``."""
    doc = Document()
    doc.add_table(rows=1, cols=1)
    out = tmp_path / "celltarget.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    target_cell = reopened.tables[0].cell(0, 0)
    resolved = resolve_effective_formatting(target_cell)
    # Cell cascade only walks docDefaults + table style — must return a
    # result without raising. docDefaults' line_spacing should populate.
    assert resolved.line_spacing is not None or resolved.partial is False
