"""Tests for ``docx_plus.numbering`` — custom list definitions."""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from docx.document import Document as DocumentObject
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

from docx_plus._testing.ooxml_asserts import assert_numbering_well_formed
from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub, xpath
from docx_plus.core.parts import NUMBERING_SPEC, get_or_create_part
from docx_plus.numbering import (
    MAX_LEVELS,
    AbstractNumIdRegistry,
    InvalidLevelError,
    LevelDefinition,
    ListDefinitionNotFoundError,
    NumIdRegistry,
    apply_list,
    define_bullet_list,
    define_list_definition,
    define_numbered_list,
    read_list_definitions,
    remove_list,
    restart_list,
)
from docx_plus.styles import resolve_effective_formatting

# --------------------------------------------------------------------------
# Helpers.
# --------------------------------------------------------------------------


def _bare_doc() -> DocumentObject:
    """A document with no numbering part.

    python-docx's bundled template ships nine ``abstractNum`` entries and
    nine ``num`` instances, which would otherwise sit in front of
    everything these tests assert on.
    """
    doc = Document()
    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype == RT.NUMBERING:
            doc.part.drop_rel(rid)
    return doc


def _numbering_root(doc: DocumentObject) -> etree._Element:
    _, root = get_or_create_part(doc, NUMBERING_SPEC)
    return root


def _children(node: etree._Element) -> list[str]:
    return [etree.QName(child.tag).localname for child in node if isinstance(child.tag, str)]


def _abstract(doc: DocumentObject, index: int = 0) -> etree._Element:
    return xpath(_numbering_root(doc), "./w:abstractNum")[index]


def _num_pr(paragraph) -> etree._Element | None:
    ppr = paragraph._p.find(qn("w:pPr"))
    return None if ppr is None else ppr.find(qn("w:numPr"))


def _round_trip(doc: DocumentObject) -> DocumentObject:
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Document(buffer)


# --------------------------------------------------------------------------
# LevelDefinition validation. A bad numFmt or lvlText renders wrong rather
# than failing, so these are checked at construction.
# --------------------------------------------------------------------------


def test_level_definition_defaults_to_a_decimal_list() -> None:
    level = LevelDefinition()
    assert (level.fmt, level.text, level.start, level.suffix) == ("decimal", "%1.", 1, "tab")


@pytest.mark.parametrize("fmt", ["decimal", "bullet", "lowerRoman", "upperLetter", "none"])
def test_level_definition_accepts_st_number_formats(fmt: str) -> None:
    assert LevelDefinition(fmt=fmt).fmt == fmt


@pytest.mark.parametrize("fmt", ["numeric", "arabic", "Decimal", ""])
def test_level_definition_rejects_unknown_number_formats(fmt: str) -> None:
    with pytest.raises(InvalidLevelError, match="ST_NumberFormat"):
        LevelDefinition(fmt=fmt)


def test_level_definition_rejects_unknown_suffix() -> None:
    with pytest.raises(InvalidLevelError, match="suffix"):
        LevelDefinition(suffix="comma")  # type: ignore[arg-type]


def test_level_definition_rejects_unknown_justification() -> None:
    with pytest.raises(InvalidLevelError, match="justify"):
        LevelDefinition(justify="justified")  # type: ignore[arg-type]


def test_level_definition_rejects_negative_start() -> None:
    with pytest.raises(InvalidLevelError, match="non-negative"):
        LevelDefinition(start=-1)


def test_level_definition_rejects_out_of_range_restart_after() -> None:
    with pytest.raises(InvalidLevelError, match="restart_after"):
        LevelDefinition(restart_after=10)


def test_level_definition_rejects_hanging_without_indent() -> None:
    """Word ignores a ``w:ind`` carrying only ``w:hanging``."""
    with pytest.raises(InvalidLevelError, match="hanging has no effect"):
        LevelDefinition(hanging=360)


# --------------------------------------------------------------------------
# define_list_definition — the primitive.
# --------------------------------------------------------------------------


def test_define_creates_the_numbering_part_when_absent() -> None:
    doc = _bare_doc()
    with pytest.raises(KeyError):
        doc.part.part_related_by(RT.NUMBERING)

    define_list_definition(doc, levels=[LevelDefinition()])

    assert doc.part.part_related_by(RT.NUMBERING) is not None


def test_define_returns_the_instance_id_not_the_abstract_id() -> None:
    doc = _bare_doc()
    num_id = define_list_definition(doc, levels=[LevelDefinition()])

    root = _numbering_root(doc)
    nums = xpath(root, "./w:num")
    assert len(nums) == 1
    assert nums[0].get(qn("w:numId")) == str(num_id)


def test_define_points_the_instance_at_the_abstract_definition() -> None:
    doc = _bare_doc()
    define_list_definition(doc, levels=[LevelDefinition()])
    root = _numbering_root(doc)

    abstract_id = xpath(root, "./w:abstractNum")[0].get(qn("w:abstractNumId"))
    ref = xpath(root, "./w:num")[0].find(qn("w:abstractNumId"))
    assert ref is not None
    assert ref.get(qn("w:val")) == abstract_id


def test_abstract_num_precedes_every_num() -> None:
    """CT_Numbering is ``numPicBullet*, abstractNum*, num*``.

    Nothing in python-docx inserts a ``w:abstractNum``, and the repo's own
    fixture builder appends one *after* the ``w:num`` entries — it
    round-trips only because consumers are lenient. Word is not always.
    """
    doc = _bare_doc()
    define_list_definition(doc, levels=[LevelDefinition()])
    define_list_definition(doc, levels=[LevelDefinition()])

    order = _children(_numbering_root(doc))
    assert order == ["abstractNum", "abstractNum", "num", "num"]


def test_lvl_children_are_in_schema_order() -> None:
    doc = _bare_doc()
    define_list_definition(
        doc,
        levels=[
            LevelDefinition(
                fmt="bullet",
                text="-",
                indent=720,
                hanging=360,
                suffix="space",
                restart_after=0,
                font="Symbol",
            )
        ],
    )
    lvl = xpath(_abstract(doc), "./w:lvl")[0]
    # ECMA-376 17.9.6: start, numFmt, lvlRestart, ..., suff, lvlText, ...,
    # lvlJc, pPr, rPr.
    assert _children(lvl) == [
        "start",
        "numFmt",
        "lvlRestart",
        "suff",
        "lvlText",
        "lvlJc",
        "pPr",
        "rPr",
    ]


def test_abstract_num_children_are_in_schema_order() -> None:
    doc = _bare_doc()
    define_list_definition(
        doc,
        levels=[LevelDefinition(), LevelDefinition(text="%2.")],
        name="Steps",
        style_link="MyListStyle",
    )
    assert _children(_abstract(doc)) == ["multiLevelType", "name", "styleLink", "lvl", "lvl"]


def test_suffix_tab_is_omitted_as_the_schema_default() -> None:
    doc = _bare_doc()
    define_list_definition(doc, levels=[LevelDefinition(suffix="tab")])
    assert "suff" not in _children(xpath(_abstract(doc), "./w:lvl")[0])


def test_multi_level_type_defaults_by_level_count() -> None:
    single = _bare_doc()
    define_list_definition(single, levels=[LevelDefinition()])
    assert _child_val(_abstract(single), "w:multiLevelType") == "singleLevel"

    multi = _bare_doc()
    define_list_definition(multi, levels=[LevelDefinition(), LevelDefinition(text="%2.")])
    assert _child_val(_abstract(multi), "w:multiLevelType") == "multilevel"


def _child_val(parent: etree._Element, tag: str) -> str | None:
    child = parent.find(qn(tag))
    return None if child is None else child.get(qn("w:val"))


def test_define_writes_indent_and_font() -> None:
    doc = _bare_doc()
    define_list_definition(
        doc,
        levels=[LevelDefinition(fmt="bullet", text="*", indent=1080, hanging=270, font="Symbol")],
    )
    lvl = xpath(_abstract(doc), "./w:lvl")[0]
    ind = lvl.find(f"{qn('w:pPr')}/{qn('w:ind')}")
    fonts = lvl.find(f"{qn('w:rPr')}/{qn('w:rFonts')}")
    assert ind is not None and fonts is not None
    assert ind.get(qn("w:left")) == "1080"
    assert ind.get(qn("w:hanging")) == "270"
    assert fonts.get(qn("w:ascii")) == "Symbol"
    # Without an explicit hint Word may substitute a theme font and render
    # the bullet as a Latin letter.
    assert fonts.get(qn("w:hint")) == "default"


def test_define_rejects_an_empty_level_list() -> None:
    with pytest.raises(InvalidLevelError, match="at least one level"):
        define_list_definition(_bare_doc(), levels=[])


def test_define_rejects_more_than_nine_levels() -> None:
    with pytest.raises(InvalidLevelError, match="at most 9 levels"):
        define_list_definition(_bare_doc(), levels=[LevelDefinition()] * (MAX_LEVELS + 1))


def test_define_rejects_a_placeholder_deeper_than_its_own_level() -> None:
    """``%3`` on level 0 references a counter that does not exist yet."""
    with pytest.raises(InvalidLevelError, match="can only interpolate counters"):
        define_list_definition(_bare_doc(), levels=[LevelDefinition(text="%3.")])


def test_define_allows_cumulative_placeholders_at_depth() -> None:
    """The legal-outline shape: level 1 may reference %1 and %2."""
    doc = _bare_doc()
    define_list_definition(
        doc,
        levels=[LevelDefinition(text="%1."), LevelDefinition(text="%1.%2.")],
    )
    texts = [_child_val(lvl, "w:lvlText") for lvl in xpath(_abstract(doc), "./w:lvl")]
    assert texts == ["%1.", "%1.%2."]


def test_define_does_not_check_placeholders_on_a_bullet_level() -> None:
    """For ``fmt="bullet"`` the lvlText is a literal glyph, not a pattern."""
    doc = _bare_doc()
    define_list_definition(doc, levels=[LevelDefinition(fmt="bullet", text="100%9")])
    assert _child_val(xpath(_abstract(doc), "./w:lvl")[0], "w:lvlText") == "100%9"


def test_define_writes_a_num_style_link() -> None:
    """The inverse of ``style_link``: defer numbering to a style."""
    doc = _bare_doc()
    define_list_definition(doc, levels=[LevelDefinition()], num_style_link="MyListStyle")
    assert _child_val(_abstract(doc), "w:numStyleLink") == "MyListStyle"
    assert read_list_definitions(doc)[0].num_style_link == "MyListStyle"


def test_define_rejects_both_style_links() -> None:
    with pytest.raises(ValueError, match="two halves"):
        define_list_definition(
            _bare_doc(),
            levels=[LevelDefinition()],
            style_link="A",
            num_style_link="B",
        )


def test_define_rejects_an_unknown_multi_level_type() -> None:
    with pytest.raises(ValueError, match="multi_level_type"):
        define_list_definition(
            _bare_doc(),
            levels=[LevelDefinition()],
            multi_level_type="outline",  # type: ignore[arg-type]
        )


def test_define_does_not_collide_with_the_template_definitions() -> None:
    """A fresh Document already carries abstractNumId 0-8 and numId 1-9."""
    doc = Document()
    num_id = define_list_definition(doc, levels=[LevelDefinition()])
    assert num_id == 10

    root = _numbering_root(doc)
    ids = [num.get(qn("w:numId")) for num in xpath(root, "./w:num")]
    assert len(ids) == len(set(ids)), "numId collision with the template entries"


def test_define_shares_registries_across_calls() -> None:
    doc = _bare_doc()
    nums = NumIdRegistry(doc)
    abstracts = AbstractNumIdRegistry(doc)
    ids = [
        define_list_definition(
            doc,
            levels=[LevelDefinition()],
            num_registry=nums,
            abstract_registry=abstracts,
        )
        for _ in range(3)
    ]
    assert ids == [1, 2, 3]


# --------------------------------------------------------------------------
# Presets.
# --------------------------------------------------------------------------


def test_bullet_preset_uses_words_glyph_and_font_cycle() -> None:
    doc = _bare_doc()
    define_bullet_list(doc, levels=3)
    levels = xpath(_abstract(doc), "./w:lvl")

    glyphs = [_child_val(lvl, "w:lvlText") for lvl in levels]
    assert glyphs == ["", "o", ""]

    fonts = [
        lvl.find(f"{qn('w:rPr')}/{qn('w:rFonts')}").get(qn("w:ascii"))  # type: ignore[union-attr]
        for lvl in levels
    ]
    assert fonts == ["Symbol", "Courier New", "Wingdings"]


def test_bullet_preset_cycles_past_three_levels() -> None:
    doc = _bare_doc()
    define_bullet_list(doc, levels=5)
    glyphs = [_child_val(lvl, "w:lvlText") for lvl in xpath(_abstract(doc), "./w:lvl")]
    assert glyphs[3] == glyphs[0]
    assert glyphs[4] == glyphs[1]


def test_numbered_preset_uses_words_format_cycle() -> None:
    doc = _bare_doc()
    define_numbered_list(doc, levels=3)
    levels = xpath(_abstract(doc), "./w:lvl")

    assert [_child_val(lvl, "w:numFmt") for lvl in levels] == [
        "decimal",
        "lowerLetter",
        "lowerRoman",
    ]
    # Each level's counter stands alone — %1., %2., %3. — rather than
    # accumulating into 1.1.1.
    assert [_child_val(lvl, "w:lvlText") for lvl in levels] == ["%1.", "%2.", "%3."]


def test_presets_indent_each_level_by_the_step() -> None:
    doc = _bare_doc()
    define_numbered_list(doc, levels=3, indent_step=500, hanging=250)
    indents = [
        lvl.find(f"{qn('w:pPr')}/{qn('w:ind')}").get(qn("w:left"))  # type: ignore[union-attr]
        for lvl in xpath(_abstract(doc), "./w:lvl")
    ]
    assert indents == ["500", "1000", "1500"]


@pytest.mark.parametrize("levels", [0, -1, MAX_LEVELS + 1])
def test_presets_reject_out_of_range_level_counts(levels: int) -> None:
    with pytest.raises(InvalidLevelError, match="levels must be between"):
        define_bullet_list(_bare_doc(), levels=levels)


# --------------------------------------------------------------------------
# apply_list / remove_list.
# --------------------------------------------------------------------------


def test_apply_list_writes_ilvl_and_num_id() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc, levels=2)
    para = doc.add_paragraph("item")
    apply_list(para, num_id, level=1)

    num_pr = _num_pr(para)
    assert num_pr is not None
    assert _child_val(num_pr, "w:ilvl") == "1"
    assert _child_val(num_pr, "w:numId") == str(num_id)
    assert _children(num_pr) == ["ilvl", "numId"]


def test_apply_list_places_num_pr_in_schema_position() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc)
    para = doc.add_paragraph("item")
    # w:ind and w:jc both sort after w:numPr in CT_PPr.
    ppr = para._p.get_or_add_pPr()
    sub(ppr, "w:ind", **{"w:left": "720"})
    sub(ppr, "w:jc", **{"w:val": "left"})

    apply_list(para, num_id)

    assert _children(ppr) == ["numPr", "ind", "jc"]


def test_apply_list_is_idempotent() -> None:
    doc = _bare_doc()
    first = define_numbered_list(doc)
    second = define_bullet_list(doc)
    para = doc.add_paragraph("item")

    apply_list(para, first)
    apply_list(para, second, level=2)

    ppr = para._p.find(qn("w:pPr"))
    assert ppr is not None
    assert len(ppr.findall(qn("w:numPr"))) == 1
    num_pr = _num_pr(para)
    assert num_pr is not None
    assert _child_val(num_pr, "w:numId") == str(second)
    assert _child_val(num_pr, "w:ilvl") == "2"


@pytest.mark.parametrize("level", [-1, MAX_LEVELS, 99])
def test_apply_list_rejects_out_of_range_levels(level: int) -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc)
    with pytest.raises(ValueError, match="level is zero-based"):
        apply_list(doc.add_paragraph("x"), num_id, level=level)


def test_apply_list_does_not_validate_the_num_id() -> None:
    """A dangling reference is legal — Word renders it unnumbered."""
    doc = _bare_doc()
    para = doc.add_paragraph("orphan")
    apply_list(para, 4242)
    num_pr = _num_pr(para)
    assert num_pr is not None
    assert _child_val(num_pr, "w:numId") == "4242"


def test_remove_list_drops_the_num_pr() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc)
    para = doc.add_paragraph("item")
    apply_list(para, num_id)

    remove_list(para)

    assert _num_pr(para) is None


def test_remove_list_is_idempotent_on_a_plain_paragraph() -> None:
    doc = _bare_doc()
    para = doc.add_paragraph("plain")
    remove_list(para)
    remove_list(para)
    assert _num_pr(para) is None


def test_remove_list_can_suppress_style_numbering() -> None:
    """numId 0 is the sentinel for "definitely not numbered"."""
    doc = _bare_doc()
    para = doc.add_paragraph("item")
    apply_list(para, define_numbered_list(doc), level=2)

    remove_list(para, suppress_style_numbering=True)

    num_pr = _num_pr(para)
    assert num_pr is not None
    assert _child_val(num_pr, "w:numId") == "0"
    # A depth is meaningless without a list.
    assert _children(num_pr) == ["numId"]


# --------------------------------------------------------------------------
# restart_list. Restarting is not a paragraph property in OOXML — it is a
# second w:num over the same w:abstractNum.
# --------------------------------------------------------------------------


def test_restart_list_creates_a_second_instance_of_one_definition() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc)
    new_id = restart_list(doc.add_paragraph("one again"), num_id)

    assert new_id != num_id
    root = _numbering_root(doc)
    assert len(xpath(root, "./w:abstractNum")) == 1, "should reuse the definition"
    assert len(xpath(root, "./w:num")) == 2


def test_restart_list_writes_a_start_override() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc, levels=3)
    new_id = restart_list(doc.add_paragraph("x"), num_id, level=1, start=5)

    num = xpath(_numbering_root(doc), "./w:num[@w:numId=$n]", n=str(new_id))[0]
    override = num.find(qn("w:lvlOverride"))
    assert override is not None
    assert override.get(qn("w:ilvl")) == "1"
    assert _child_val(override, "w:startOverride") == "5"


def test_restart_list_applies_the_new_id_to_the_paragraph() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc)
    para = doc.add_paragraph("restarted")
    new_id = restart_list(para, num_id)

    num_pr = _num_pr(para)
    assert num_pr is not None
    assert _child_val(num_pr, "w:numId") == str(new_id)


def test_restart_list_rejects_a_negative_start() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc)
    with pytest.raises(ValueError, match="start must be non-negative"):
        restart_list(doc.add_paragraph("x"), num_id, start=-1)


def test_restart_list_raises_for_an_unknown_definition() -> None:
    doc = _bare_doc()
    with pytest.raises(ListDefinitionNotFoundError):
        restart_list(doc.add_paragraph("x"), 999)


def test_restart_list_error_subclasses_key_error() -> None:
    doc = _bare_doc()
    with pytest.raises(KeyError):
        restart_list(doc.add_paragraph("x"), 999)


def test_restart_list_raises_when_the_instance_has_no_abstract_reference() -> None:
    doc = _bare_doc()
    root = _numbering_root(doc)
    sub(root, "w:num", **{"w:numId": "7"})  # malformed: no w:abstractNumId

    with pytest.raises(ListDefinitionNotFoundError, match="abstractNumId"):
        restart_list(doc.add_paragraph("x"), 7)


def test_restart_list_keeps_nums_after_abstract_nums() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc)
    restart_list(doc.add_paragraph("x"), num_id)
    assert _children(_numbering_root(doc)) == ["abstractNum", "num", "num"]


# --------------------------------------------------------------------------
# read_list_definitions.
# --------------------------------------------------------------------------


def test_read_returns_empty_without_a_numbering_part() -> None:
    assert read_list_definitions(_bare_doc()) == []


def test_read_does_not_fabricate_the_part() -> None:
    doc = _bare_doc()
    read_list_definitions(doc)
    with pytest.raises(KeyError):
        doc.part.part_related_by(RT.NUMBERING)


def test_read_round_trips_a_definition() -> None:
    doc = _bare_doc()
    num_id = define_list_definition(
        doc,
        levels=[
            LevelDefinition(
                fmt="upperRoman",
                text="%1)",
                start=3,
                indent=720,
                hanging=360,
                justify="right",
                suffix="space",
                restart_after=0,
                font="Cambria",
            )
        ],
        name="Exhibits",
    )

    definition = read_list_definitions(doc)[0]
    assert definition.num_id == num_id
    assert definition.name == "Exhibits"
    assert definition.multi_level_type == "singleLevel"

    level = definition.levels[0]
    assert level.level == 0
    assert level.fmt == "upperRoman"
    assert level.text == "%1)"
    assert level.start == 3
    assert level.indent == 720
    assert level.hanging == 360
    assert level.justify == "right"
    assert level.suffix == "space"
    assert level.restart_after == 0
    assert level.font == "Cambria"


def test_read_reports_start_overrides() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc, levels=2)
    new_id = restart_list(doc.add_paragraph("x"), num_id, level=1, start=4)

    by_id = {d.num_id: d for d in read_list_definitions(doc)}
    assert by_id[num_id].start_overrides == ()
    assert by_id[new_id].start_overrides == ((1, 4),)
    # Both instances share one definition, so both report its levels.
    assert by_id[num_id].abstract_id == by_id[new_id].abstract_id
    assert len(by_id[new_id].levels) == 2


def test_read_sees_the_bundled_template_definitions() -> None:
    """A fresh Document is not a blank slate for numbering."""
    definitions = read_list_definitions(Document())
    assert len(definitions) == 9
    assert {d.abstract_id for d in definitions} == set(range(9))


def test_read_tolerates_a_dangling_abstract_reference() -> None:
    doc = _bare_doc()
    root = _numbering_root(doc)
    num = sub(root, "w:num", **{"w:numId": "5"})
    sub(num, "w:abstractNumId", **{"w:val": "404"})

    definition = read_list_definitions(doc)[0]
    assert definition.num_id == 5
    assert definition.abstract_id == 404
    assert definition.levels == ()


def test_read_skips_an_instance_with_no_id() -> None:
    doc = _bare_doc()
    sub(_numbering_root(doc), "w:num")
    assert read_list_definitions(doc) == []


def test_read_tolerates_unparseable_ids() -> None:
    doc = _bare_doc()
    root = _numbering_root(doc)
    num = sub(root, "w:num", **{"w:numId": "3"})
    sub(num, "w:abstractNumId", **{"w:val": "not-a-number"})

    definition = read_list_definitions(doc)[0]
    assert definition.num_id == 3
    assert definition.abstract_id is None


def test_read_skips_a_level_with_no_ilvl() -> None:
    """A ``w:lvl`` without a depth cannot be placed, so it is dropped."""
    doc = _bare_doc()
    root = _numbering_root(doc)
    abstract = sub(root, "w:abstractNum", **{"w:abstractNumId": "0"})
    sub(abstract, "w:lvl")  # no w:ilvl
    sub(abstract, "w:lvl", **{"w:ilvl": "0"})
    num = sub(root, "w:num", **{"w:numId": "1"})
    sub(num, "w:abstractNumId", **{"w:val": "0"})

    levels = read_list_definitions(doc)[0].levels
    assert [lvl.level for lvl in levels] == [0]


def test_read_omits_absent_level_children_as_none() -> None:
    """``None`` means the element is absent, not that its value is zero."""
    doc = _bare_doc()
    root = _numbering_root(doc)
    abstract = sub(root, "w:abstractNum", **{"w:abstractNumId": "0"})
    sub(abstract, "w:lvl", **{"w:ilvl": "0"})
    num = sub(root, "w:num", **{"w:numId": "1"})
    sub(num, "w:abstractNumId", **{"w:val": "0"})

    level = read_list_definitions(doc)[0].levels[0]
    assert level.level == 0
    assert level.fmt is None
    assert level.start is None
    assert level.indent is None
    assert level.font is None


# --------------------------------------------------------------------------
# Round-trip.
# --------------------------------------------------------------------------


def test_definition_survives_save_and_reopen(tmp_path: Path) -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc, levels=2)
    apply_list(doc.add_paragraph("one"), num_id)
    apply_list(doc.add_paragraph("nested"), num_id, level=1)

    out = tmp_path / "lists.docx"
    doc.save(str(out))
    reopened = Document(str(out))

    definitions = read_list_definitions(reopened)
    assert [d.num_id for d in definitions] == [num_id]
    assert [lvl.fmt for lvl in definitions[0].levels] == ["decimal", "lowerLetter"]

    levels = [
        _child_val(num_pr, "w:ilvl") for num_pr in xpath(reopened.element.body, ".//w:pPr/w:numPr")
    ]
    assert levels == ["0", "1"]


def test_restarted_sequence_survives_a_round_trip() -> None:
    doc = _bare_doc()
    num_id = define_numbered_list(doc)
    apply_list(doc.add_paragraph("one"), num_id)
    new_id = restart_list(doc.add_paragraph("one again"), num_id, start=1)

    reopened = _round_trip(doc)
    by_id = {d.num_id: d for d in read_list_definitions(reopened)}
    assert by_id[new_id].start_overrides == ((0, 1),)


def test_generated_numbering_is_well_formed() -> None:
    doc = _bare_doc()
    first = define_numbered_list(doc, levels=3)
    define_bullet_list(doc, levels=2)
    restart_list(doc.add_paragraph("x"), first)

    assert_numbering_well_formed(doc)
    assert_numbering_well_formed(_round_trip(doc))


def test_well_formed_assertion_catches_misordered_children() -> None:
    """Guards the guard: the repo's own fixture builder gets this wrong."""
    doc = _bare_doc()
    root = _numbering_root(doc)
    num = sub(root, "w:num", **{"w:numId": "1"})
    sub(num, "w:abstractNumId", **{"w:val": "0"})
    sub(root, "w:abstractNum", **{"w:abstractNumId": "0"})  # appended after the num

    with pytest.raises(AssertionError, match="must precede"):
        assert_numbering_well_formed(doc)


def test_a_definition_added_to_the_stock_template_round_trips(tmp_path: Path) -> None:
    """The realistic path: a document that already has numbering.xml."""
    doc = Document()
    num_id = define_bullet_list(doc, levels=2)
    apply_list(doc.add_paragraph("bullet"), num_id)

    out = tmp_path / "mixed.docx"
    doc.save(str(out))
    reopened = Document(str(out))

    mine = [d for d in read_list_definitions(reopened) if d.num_id == num_id]
    assert len(mine) == 1
    assert mine[0].levels[0].fmt == "bullet"
    # The template's nine definitions are untouched.
    assert len(read_list_definitions(reopened)) == 10


# --------------------------------------------------------------------------
# Registries.
# --------------------------------------------------------------------------


def test_registries_seed_from_an_existing_part() -> None:
    doc = Document()
    assert NumIdRegistry(doc).issued() == frozenset(range(1, 10))
    assert AbstractNumIdRegistry(doc).issued() == frozenset(range(9))


def test_registries_are_empty_without_a_part() -> None:
    doc = _bare_doc()
    assert NumIdRegistry(doc).issued() == frozenset()
    assert AbstractNumIdRegistry(doc).issued() == frozenset()


def test_registry_construction_does_not_create_the_part() -> None:
    doc = _bare_doc()
    NumIdRegistry(doc)
    with pytest.raises(KeyError):
        doc.part.part_related_by(RT.NUMBERING)


def test_abstract_registry_allocates_from_zero() -> None:
    """Unlike every w:id namespace, abstractNumId 0 is legal."""
    assert AbstractNumIdRegistry(_bare_doc()).next_sequential() == 0


def test_num_registry_allocates_from_one() -> None:
    """numId 0 is the "no numbering" sentinel, never a definition id."""
    assert NumIdRegistry(_bare_doc()).next_sequential() == 1


# --------------------------------------------------------------------------
# Interaction with the style cascade, which reads numbering.xml as layer 4.
# --------------------------------------------------------------------------


def test_cascade_resolves_formatting_from_a_generated_definition() -> None:
    doc = _bare_doc()
    num_id = define_list_definition(doc, levels=[LevelDefinition(indent=1440, hanging=360)])
    para = doc.add_paragraph("item")
    apply_list(para, num_id)

    resolved = resolve_effective_formatting(para, include_provenance=True)
    assert resolved.num_id == num_id
    assert resolved.num_level == 0
    assert resolved.indent_left == 1440
    assert (resolved.provenance or {})["indent_left"].layer == "numbering"
