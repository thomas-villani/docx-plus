"""Tests for comment durable ids (`commentsIds.xml`) and author presence.

The URIs and the on-disk shapes asserted here were verified against a
file authored by Word 2016 itself, not inferred from the spec.
"""

from __future__ import annotations

import re

import pytest
from docx import Document

from docx_plus._testing.ooxml_asserts import assert_durable_ids_well_formed
from docx_plus.comments import (
    LOCAL_PROVIDER,
    AuthorPresence,
    DurableIdRegistry,
    _ids,
    add_comment,
    clear_all_comments,
    clear_author_presence,
    delete_comment,
    edit_comment,
    read_author_presence,
    read_comments,
    reply_to_comment,
    set_author_presence,
)
from docx_plus.core.ns import qn
from docx_plus.core.oxml import el, xpath
from docx_plus.core.parts import (
    CT_COMMENTS_IDS,
    CT_PEOPLE,
    RT_COMMENTS_IDS,
    RT_PEOPLE,
)

LONG_HEX = re.compile(r"^[0-9A-F]{8}$")


@pytest.fixture
def doc():
    """A document with one paragraph ready to be commented."""
    document = Document()
    document.add_paragraph("Some reviewed text here.")
    return document


def id_entries(document):
    """Every `<w16cid:commentId>` in the document, or [] if the part is absent."""
    root = _ids.ids_root(document)
    return [] if root is None else xpath(root, "./w16cid:commentId")


# ---------------------------------------------------------------------------
# The URIs, pinned against what Word actually writes.
# ---------------------------------------------------------------------------


class TestPartUris:
    """Verified against a Word 2016-authored file, unzipped and read."""

    def test_comments_ids_content_type(self):
        assert CT_COMMENTS_IDS == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsIds+xml"
        )

    def test_comments_ids_relationship(self):
        assert RT_COMMENTS_IDS == (
            "http://schemas.microsoft.com/office/2016/09/relationships/commentsIds"
        )

    def test_people_content_type(self):
        assert CT_PEOPLE == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.people+xml"
        )

    def test_people_relationship(self):
        assert RT_PEOPLE == "http://schemas.microsoft.com/office/2011/relationships/people"


# ---------------------------------------------------------------------------
# Durable ids.
# ---------------------------------------------------------------------------


class TestDurableIdsOnWrite:
    def test_add_comment_writes_an_entry(self, doc):
        add_comment(doc.paragraphs[0], "note", author="R")
        entries = id_entries(doc)
        assert len(entries) == 1
        assert LONG_HEX.match(entries[0].get(qn("w16cid:durableId")))
        assert_durable_ids_well_formed(doc)

    def test_durable_id_is_hex_not_decimal(self, doc):
        """Word writes ST_LongHexNumber here; the plan had assumed decimal."""
        add_comment(doc.paragraphs[0], "note")
        raw = id_entries(doc)[0].get(qn("w16cid:durableId"))
        assert LONG_HEX.match(raw), raw
        assert raw.upper() == raw

    def test_entry_keys_off_the_thread_key(self, doc):
        ref = add_comment(doc.paragraphs[0], "note")
        body_para_id = xpath(ref.body_element, ".//w:p")[-1].get(qn("w14:paraId"))
        assert id_entries(doc)[0].get(qn("w16cid:paraId")) == body_para_id

    def test_reply_gets_its_own(self, doc):
        root = add_comment(doc.paragraphs[0], "root")
        reply_to_comment(doc, root.comment_id, "reply")
        assert len(id_entries(doc)) == 2
        assert_durable_ids_well_formed(doc)

    def test_many_comments_stay_unique(self, doc):
        for index in range(25):
            add_comment(doc.add_paragraph(f"para {index}"), f"note {index}")
        assert len(id_entries(doc)) == 25
        assert_durable_ids_well_formed(doc)

    def test_shared_registry_stays_unique(self, doc):
        registry = DurableIdRegistry(doc)
        for index in range(10):
            add_comment(
                doc.add_paragraph(f"para {index}"),
                f"note {index}",
                durable_id_registry=registry,
            )
        assert_durable_ids_well_formed(doc)

    def test_registry_seeds_from_an_existing_part(self, doc):
        add_comment(doc.paragraphs[0], "first")
        existing = id_entries(doc)[0].get(qn("w16cid:durableId"))
        assert int(existing, 16) in DurableIdRegistry(doc)._issued

    def test_no_part_until_a_comment_exists(self, doc):
        assert _ids.ids_root(doc) is None


class TestDurableIdStability:
    def test_survives_a_save_and_reload(self, doc, tmp_path):
        add_comment(doc.paragraphs[0], "note")
        before = read_comments(doc)[0].durable_id
        path = tmp_path / "durable.docx"
        doc.save(path)
        assert read_comments(Document(path))[0].durable_id == before

    def test_edit_comment_does_not_reissue(self, doc):
        ref = add_comment(doc.paragraphs[0], "note")
        before = read_comments(doc)[0].durable_id
        edit_comment(doc, ref.comment_id, "rewritten entirely")
        assert read_comments(doc)[0].durable_id == before

    def test_upsert_is_idempotent(self, doc):
        add_comment(doc.paragraphs[0], "note")
        para_id = id_entries(doc)[0].get(qn("w16cid:paraId"))
        first = _ids.upsert_comment_id(doc, para_id)
        second = _ids.upsert_comment_id(doc, para_id)
        assert first == second
        assert len(id_entries(doc)) == 1

    def test_empty_thread_key_is_a_no_op(self, doc):
        assert _ids.upsert_comment_id(doc, "") is None


class TestDurableIdsOnRemoval:
    def test_delete_comment_drops_the_entry(self, doc):
        first = add_comment(doc.paragraphs[0], "one")
        add_comment(doc.add_paragraph("second para"), "two")
        delete_comment(doc, first.comment_id)
        assert len(id_entries(doc)) == 1

    def test_delete_cascades_to_replies(self, doc):
        root = add_comment(doc.paragraphs[0], "root")
        reply_to_comment(doc, root.comment_id, "reply")
        delete_comment(doc, root.comment_id)
        assert id_entries(doc) == []

    def test_clear_all_empties_the_part(self, doc):
        add_comment(doc.paragraphs[0], "note")
        clear_all_comments(doc)
        assert id_entries(doc) == []
        assert _ids.ids_root(doc) is not None

    def test_clear_all_with_remove_part_drops_it(self, doc):
        add_comment(doc.paragraphs[0], "note")
        clear_all_comments(doc, remove_part=True)
        assert _ids.ids_root(doc) is None

    def test_drop_is_idempotent(self, doc):
        add_comment(doc.paragraphs[0], "note")
        para_id = id_entries(doc)[0].get(qn("w16cid:paraId"))
        _ids.drop_comment_id(doc, para_id)
        _ids.drop_comment_id(doc, para_id)
        assert id_entries(doc) == []

    def test_drop_without_a_part_is_a_no_op(self, doc):
        _ids.drop_comment_id(doc, "DEADBEEF")

    def test_drop_with_an_empty_key_is_a_no_op(self, doc):
        add_comment(doc.paragraphs[0], "note")
        _ids.drop_comment_id(doc, "")
        assert len(id_entries(doc)) == 1

    def test_clear_without_a_part_is_a_no_op(self, doc):
        clear_all_comments(doc)


class TestDurableIdReads:
    def test_read_comments_reports_it(self, doc):
        add_comment(doc.paragraphs[0], "note")
        assert LONG_HEX.match(read_comments(doc)[0].durable_id)

    def test_none_when_the_part_is_absent(self, doc):
        add_comment(doc.paragraphs[0], "note")
        clear_all_comments(doc, remove_part=True)
        add_comment(doc.paragraphs[0], "note again")
        # Re-added, so present again; strip the part to prove the fallback.
        for rid, rel in list(doc.part.rels.items()):
            if rel.reltype == RT_COMMENTS_IDS:
                doc.part.drop_rel(rid)
        assert read_comments(doc)[0].durable_id is None

    def test_map_is_empty_without_a_part(self, doc):
        assert _ids.durable_id_map(doc) == {}

    def test_map_skips_incomplete_entries(self, doc):
        add_comment(doc.paragraphs[0], "note")
        root = _ids.ids_root(doc)
        root.append(el("w16cid:commentId", **{"w16cid:paraId": "AAAAAAAA"}))
        root.append(el("w16cid:commentId", **{"w16cid:durableId": "BBBBBBBB"}))
        assert len(_ids.durable_id_map(doc)) == 1

    def test_find_returns_none_for_an_unknown_key(self, doc):
        add_comment(doc.paragraphs[0], "note")
        assert _ids.find_comment_id(_ids.ids_root(doc), "00000001") is None


# ---------------------------------------------------------------------------
# Author presence.
# ---------------------------------------------------------------------------


class TestSetAuthorPresence:
    def test_creates_the_part_and_entry(self, doc):
        result = set_author_presence(doc, "Reviewer")
        assert result == AuthorPresence("Reviewer", LOCAL_PROVIDER, "Reviewer")
        assert read_author_presence(doc) == [result]

    def test_local_provider_default(self):
        assert LOCAL_PROVIDER == "None"

    def test_custom_provider_and_user_id(self, doc):
        result = set_author_presence(
            doc, "Author", provider_id="AD", user_id="S::author@example.com::abc"
        )
        assert result == AuthorPresence("Author", "AD", "S::author@example.com::abc")

    def test_empty_author_is_a_no_op(self, doc):
        assert set_author_presence(doc, "") is None
        assert read_author_presence(doc) == []

    def test_empty_author_does_not_create_the_part(self, doc):
        set_author_presence(doc, "")
        with pytest.raises(KeyError):
            doc.part.part_related_by(RT_PEOPLE)

    def test_idempotent_replace(self, doc):
        set_author_presence(doc, "Reviewer")
        set_author_presence(doc, "Reviewer", provider_id="AD", user_id="new")
        people = read_author_presence(doc)
        assert len(people) == 1
        assert people[0] == AuthorPresence("Reviewer", "AD", "new")

    def test_several_authors_coexist(self, doc):
        set_author_presence(doc, "One")
        set_author_presence(doc, "Two")
        assert [p.author for p in read_author_presence(doc)] == ["One", "Two"]

    def test_survives_a_save_and_reload(self, doc, tmp_path):
        set_author_presence(doc, "Reviewer", provider_id="AD", user_id="uid")
        path = tmp_path / "people.docx"
        doc.save(path)
        assert read_author_presence(Document(path)) == [AuthorPresence("Reviewer", "AD", "uid")]


class TestReadAuthorPresence:
    def test_empty_without_a_part(self, doc):
        assert read_author_presence(doc) == []

    def test_add_comment_does_not_write_the_part(self, doc):
        """Presence is cosmetic and needs an identity we cannot invent."""
        add_comment(doc.paragraphs[0], "note", author="Reviewer")
        assert read_author_presence(doc) == []

    def test_entry_without_presence_info_reads_as_none(self, doc):
        set_author_presence(doc, "Reviewer")
        people_root = doc.part.part_related_by(RT_PEOPLE).element
        for person in xpath(people_root, "./w15:person"):
            people_root.remove(person)
        people_root.append(el("w15:person", **{"w15:author": "Bare"}))
        assert read_author_presence(doc) == [AuthorPresence("Bare", None, None)]

    def test_entry_without_an_author_is_skipped(self, doc):
        set_author_presence(doc, "Reviewer")
        people_root = doc.part.part_related_by(RT_PEOPLE).element
        people_root.append(el("w15:person"))
        assert [p.author for p in read_author_presence(doc)] == ["Reviewer"]


class TestClearAuthorPresence:
    def test_empties_but_keeps_the_part(self, doc):
        set_author_presence(doc, "Reviewer")
        clear_author_presence(doc)
        assert read_author_presence(doc) == []
        doc.part.part_related_by(RT_PEOPLE)  # still related

    def test_remove_part_drops_it(self, doc):
        set_author_presence(doc, "Reviewer")
        clear_author_presence(doc, remove_part=True)
        with pytest.raises(KeyError):
            doc.part.part_related_by(RT_PEOPLE)

    def test_without_a_part_is_a_no_op(self, doc):
        clear_author_presence(doc)

    def test_not_pruned_by_deleting_a_comment(self, doc):
        """Word keeps stale authors; ref-counting them is the caller's call."""
        ref = add_comment(doc.paragraphs[0], "note", author="Reviewer")
        set_author_presence(doc, "Reviewer")
        delete_comment(doc, ref.comment_id)
        assert [p.author for p in read_author_presence(doc)] == ["Reviewer"]

    def test_not_pruned_by_clear_all_comments(self, doc):
        add_comment(doc.paragraphs[0], "note", author="Reviewer")
        set_author_presence(doc, "Reviewer")
        clear_all_comments(doc, remove_part=True)
        assert [p.author for p in read_author_presence(doc)] == ["Reviewer"]


# ---------------------------------------------------------------------------
# The assertion helper itself.
# ---------------------------------------------------------------------------


class TestDurableIdAssertion:
    def test_passes_without_a_part(self, doc):
        assert_durable_ids_well_formed(doc)

    def test_rejects_a_duplicate_durable_id(self, doc):
        add_comment(doc.paragraphs[0], "note")
        root = _ids.ids_root(doc)
        clone = root[0].get(qn("w16cid:durableId"))
        root.append(
            el("w16cid:commentId", **{"w16cid:paraId": "AAAAAAAA", "w16cid:durableId": clone})
        )
        with pytest.raises(AssertionError, match="duplicate w16cid:durableId"):
            assert_durable_ids_well_formed(doc)

    def test_rejects_a_decimal_durable_id(self, doc):
        add_comment(doc.paragraphs[0], "note")
        _ids.ids_root(doc).append(
            el(
                "w16cid:commentId",
                **{"w16cid:paraId": "AAAAAAAA", "w16cid:durableId": "123456789"},
            )
        )
        with pytest.raises(AssertionError, match="ST_LongHexNumber"):
            assert_durable_ids_well_formed(doc)

    def test_rejects_a_missing_attribute(self, doc):
        add_comment(doc.paragraphs[0], "note")
        _ids.ids_root(doc).append(el("w16cid:commentId", **{"w16cid:paraId": "AAAAAAAA"}))
        with pytest.raises(AssertionError, match="missing w16cid:durableId"):
            assert_durable_ids_well_formed(doc)

    def test_rejects_a_duplicate_para_id(self, doc):
        add_comment(doc.paragraphs[0], "note")
        root = _ids.ids_root(doc)
        para_id = root[0].get(qn("w16cid:paraId"))
        root.append(
            el(
                "w16cid:commentId",
                **{"w16cid:paraId": para_id, "w16cid:durableId": "AAAAAAAA"},
            )
        )
        with pytest.raises(AssertionError, match="duplicate w16cid:paraId"):
            assert_durable_ids_well_formed(doc)
