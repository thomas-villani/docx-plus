"""Tests for ``docx_plus.revisions`` — tracked-change authoring, reading,
accept/reject, the track-changes settings toggle, and the revision-id
registry."""

from __future__ import annotations

import datetime as dt

import pytest
from docx import Document

from docx_plus.core.ids import DuplicateIdError
from docx_plus.core.ns import qn
from docx_plus.core.oxml import el, sub, xpath
from docx_plus.revisions import (
    RevisionIdRegistry,
    RevisionNotFoundError,
    RevisionRef,
    TrackedChange,
    accept_all_revisions,
    accept_revision,
    disable_track_changes,
    enable_track_changes,
    mark_deletion,
    mark_insertion,
    read_revisions,
    reject_all_revisions,
    reject_revision,
)

# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------

_DATE = "2020-01-01T00:00:00Z"


def _body(doc):
    return doc.element.body


def _ins(doc, rid=None):
    expr = ".//w:ins[@w:id=$rid]" if rid is not None else ".//w:ins"
    return xpath(_body(doc), expr, rid=str(rid)) if rid is not None else xpath(_body(doc), expr)


def _del(doc, rid=None):
    expr = ".//w:del[@w:id=$rid]" if rid is not None else ".//w:del"
    return xpath(_body(doc), expr, rid=str(rid)) if rid is not None else xpath(_body(doc), expr)


def _settings_tags(doc):
    return [c.tag.rpartition("}")[2] for c in doc.settings.element]


def _add_rpr_change(run, *, rid, bold_now=True):
    """Give ``run`` a rPrChange: new props bold, recorded old props empty."""
    rpr = run._r.get_or_add_rPr()
    if bold_now:
        sub(rpr, "w:b")
    change = sub(rpr, "w:rPrChange", **{"w:id": str(rid), "w:author": "X", "w:date": _DATE})
    sub(change, "w:rPr")  # recorded prior props: not bold


def _add_ppr_change(paragraph, *, rid):
    """Give ``paragraph`` a pPrChange: new jc=center, recorded old jc=left."""
    ppr = paragraph._p.get_or_add_pPr()
    sub(ppr, "w:jc", **{"w:val": "center"})
    change = sub(ppr, "w:pPrChange", **{"w:id": str(rid), "w:author": "X", "w:date": _DATE})
    old = sub(change, "w:pPr")
    sub(old, "w:jc", **{"w:val": "left"})


def _add_paragraph_mark_ins(paragraph, *, rid):
    """Mark ``paragraph``'s end mark as an inserted revision (pPr/rPr/ins)."""
    ppr = paragraph._p.get_or_add_pPr()
    rpr = ppr.find(qn("w:rPr"))
    if rpr is None:
        rpr = sub(ppr, "w:rPr")
    sub(rpr, "w:ins", **{"w:id": str(rid), "w:author": "X", "w:date": _DATE})


# --------------------------------------------------------------------------
# RevisionIdRegistry — single shared id namespace
# --------------------------------------------------------------------------


def test_registry_seeds_from_existing_ins_and_del() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("a")
    r2 = p.add_run("b")
    mark_insertion(r1)  # fresh registry each -> random ids
    mark_deletion(r2)
    seeded = RevisionIdRegistry(doc).issued()
    existing = {int(e.get(qn("w:id"))) for e in _ins(doc) + _del(doc)}
    assert existing.issubset(seeded)


def test_registry_cross_type_collision_one_namespace() -> None:
    """A w:ins id blocks reuse by a w:del — all revisions share one namespace."""
    doc = Document()
    p = doc.add_paragraph()
    ref = mark_insertion(p.add_run("x"))
    reg = RevisionIdRegistry(doc)
    with pytest.raises(DuplicateIdError):
        reg.reserve(ref.revision_id)


def test_registry_shared_across_marks_gives_unique_ids() -> None:
    doc = Document()
    p = doc.add_paragraph()
    reg = RevisionIdRegistry(doc)
    ids = {mark_insertion(p.add_run(c), id_registry=reg).revision_id for c in "abcde"}
    assert len(ids) == 5


def test_registry_seeds_from_property_change_markers() -> None:
    doc = Document()
    p = doc.add_paragraph()
    _add_rpr_change(p.add_run("x"), rid=4242)
    assert 4242 in RevisionIdRegistry(doc).issued()


# --------------------------------------------------------------------------
# settings toggle
# --------------------------------------------------------------------------


def test_enable_track_changes_writes_one_element() -> None:
    doc = Document()
    enable_track_changes(doc)
    assert _settings_tags(doc).count("trackChanges") == 1


def test_enable_track_changes_idempotent() -> None:
    doc = Document()
    enable_track_changes(doc)
    enable_track_changes(doc)
    assert _settings_tags(doc).count("trackChanges") == 1


def test_enable_track_changes_schema_position() -> None:
    doc = Document()
    enable_track_changes(doc)
    tags = _settings_tags(doc)
    assert tags.index("trackChanges") < tags.index("defaultTabStop")


def test_enable_normalizes_existing_false_value() -> None:
    doc = Document()
    settings = doc.settings.element
    settings.append(el("w:trackChanges", **{"w:val": "false"}))
    enable_track_changes(doc)
    elems = settings.findall(qn("w:trackChanges"))
    assert len(elems) == 1
    assert elems[0].get(qn("w:val")) is None


def test_disable_track_changes_removes_all() -> None:
    doc = Document()
    enable_track_changes(doc)
    disable_track_changes(doc)
    assert "trackChanges" not in _settings_tags(doc)


def test_disable_track_changes_idempotent() -> None:
    doc = Document()
    disable_track_changes(doc)  # never enabled
    assert "trackChanges" not in _settings_tags(doc)


# --------------------------------------------------------------------------
# mark_insertion / mark_deletion — wrapping
# --------------------------------------------------------------------------


def test_mark_insertion_wraps_run_in_ins() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("hello")
    ref = mark_insertion(run, author="Alice")
    assert isinstance(ref, RevisionRef)
    wraps = _ins(doc, ref.revision_id)
    assert len(wraps) == 1
    assert wraps[0].get(qn("w:author")) == "Alice"
    # the run now lives inside the w:ins
    assert run._r.getparent() is wraps[0]
    # w:t preserved (not retagged)
    assert xpath(wraps[0], ".//w:t")


def test_mark_deletion_retags_t_to_deltext() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("gone")
    ref = mark_deletion(run, author="Bob")
    wrap = _del(doc, ref.revision_id)[0]
    assert not xpath(wrap, ".//w:t")
    deltexts = xpath(wrap, ".//w:delText")
    assert len(deltexts) == 1
    assert deltexts[0].text == "gone"


def test_mark_emits_iso_utc_timestamp() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = mark_insertion(p.add_run("x"))
    date = _ins(doc, ref.revision_id)[0].get(qn("w:date"))
    assert date is not None and date.endswith("Z")
    parsed = dt.datetime.fromisoformat(date.replace("Z", "+00:00"))
    assert parsed.tzinfo is not None


def test_mark_accepts_explicit_naive_date_as_utc() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = mark_insertion(p.add_run("x"), date=dt.datetime(2021, 6, 1, 12, 0, 0))
    date = _ins(doc, ref.revision_id)[0].get(qn("w:date"))
    assert date == "2021-06-01T12:00:00.000Z"


def test_mark_insertion_paragraph_wraps_all_runs() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("one ")
    p.add_run("two ")
    p.add_run("three")
    ref = mark_insertion(p)
    wrap = _ins(doc, ref.revision_id)[0]
    assert len(xpath(wrap, "./w:r")) == 3


def test_mark_insertion_range_wraps_inclusive() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("first ")
    p.add_run("middle ")
    r3 = p.add_run("last")
    ref = mark_insertion((r1, r3))
    wrap = _ins(doc, ref.revision_id)[0]
    assert len(xpath(wrap, "./w:r")) == 3


def test_mark_paragraph_with_no_runs_raises() -> None:
    doc = Document()
    p = doc.add_paragraph()
    with pytest.raises(ValueError, match="at least one run"):
        mark_insertion(p)


def test_mark_range_reversed_raises() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("a ")
    r2 = p.add_run("b")
    with pytest.raises(ValueError, match="reversed"):
        mark_insertion((r2, r1))


def test_mark_range_across_paragraphs_raises() -> None:
    doc = Document()
    p1 = doc.add_paragraph()
    r1 = p1.add_run("a")
    p2 = doc.add_paragraph()
    r2 = p2.add_run("b")
    with pytest.raises(ValueError, match="single paragraph"):
        mark_insertion((r1, r2))


def test_mark_bad_target_type_raises() -> None:
    with pytest.raises(TypeError, match="Run, Paragraph"):
        mark_insertion("not a run")  # type: ignore[arg-type]


def test_marked_runs_invisible_to_python_docx_runs() -> None:
    """Runs wrapped in w:ins are not exposed via paragraph.runs."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("only run")
    assert len(p.runs) == 1
    mark_insertion(p.runs[0])
    assert len(p.runs) == 0  # now nested inside w:ins
    # but read_revisions still sees it
    assert len(read_revisions(doc)) == 1


def test_mark_preserves_whitespace_in_deltext() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("  spaced  ")
    run._r.find(qn("w:t")).set(qn("xml:space"), "preserve")
    ref = mark_deletion(run)
    deltext = xpath(_del(doc, ref.revision_id)[0], ".//w:delText")[0]
    assert deltext.get(qn("xml:space")) == "preserve"
    assert deltext.text == "  spaced  "


# --------------------------------------------------------------------------
# read_revisions
# --------------------------------------------------------------------------


def test_read_revisions_empty_document() -> None:
    assert read_revisions(Document()) == []


def test_read_revisions_reports_insertion_and_deletion() -> None:
    doc = Document()
    p = doc.add_paragraph()
    mark_insertion(p.add_run("ins text"), author="A")
    mark_deletion(p.add_run("del text"), author="B")
    revs = read_revisions(doc)
    by_type = {r.revision_type: r for r in revs}
    assert isinstance(revs[0], TrackedChange)
    assert by_type["insertion"].text == "ins text"
    assert by_type["insertion"].author == "A"
    assert by_type["deletion"].text == "del text"
    assert by_type["deletion"].author == "B"
    assert all(r.paragraph_index == 0 for r in revs)


def test_read_revisions_document_order() -> None:
    doc = Document()
    p = doc.add_paragraph()
    reg = RevisionIdRegistry(doc)
    mark_insertion(p.add_run("1"), id_registry=reg)
    mark_deletion(p.add_run("2"), id_registry=reg)
    mark_insertion(p.add_run("3"), id_registry=reg)
    assert [r.text for r in read_revisions(doc)] == ["1", "2", "3"]


def test_read_revisions_parses_timestamp() -> None:
    doc = Document()
    p = doc.add_paragraph()
    mark_insertion(p.add_run("x"))
    ts = read_revisions(doc)[0].timestamp
    assert ts is not None and ts.tzinfo is not None


def test_read_revisions_format_run_change() -> None:
    doc = Document()
    p = doc.add_paragraph()
    _add_rpr_change(p.add_run("x"), rid=11)
    rev = read_revisions(doc)[0]
    assert rev.revision_type == "format_run"
    assert rev.revision_id == 11
    assert rev.text == ""


def test_read_revisions_format_paragraph_change() -> None:
    doc = Document()
    p = doc.add_paragraph("body")
    _add_ppr_change(p, rid=12)
    rev = next(r for r in read_revisions(doc) if r.revision_type == "format_paragraph")
    assert rev.revision_id == 12


def test_read_revisions_move_wrappers() -> None:
    doc = Document()
    p = doc.add_paragraph()
    mf = sub(p._p, "w:moveFrom", **{"w:id": "20", "w:author": "M", "w:date": _DATE})
    r = sub(mf, "w:r")
    sub(r, "w:t").text = "moved"
    mt = sub(p._p, "w:moveTo", **{"w:id": "21", "w:author": "M", "w:date": _DATE})
    r2 = sub(mt, "w:r")
    sub(r2, "w:t").text = "moved"
    types = {r.revision_type for r in read_revisions(doc)}
    assert "move_from" in types
    assert "move_to" in types


def test_read_revisions_paragraph_mark_insertion() -> None:
    doc = Document()
    p = doc.add_paragraph("text")
    _add_paragraph_mark_ins(p, rid=30)
    rev = next(r for r in read_revisions(doc) if r.revision_id == 30)
    assert rev.revision_type == "paragraph_mark_insertion"
    assert rev.text == ""


# --------------------------------------------------------------------------
# accept / reject round-trips
# --------------------------------------------------------------------------


def test_accept_insertion_keeps_text() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = mark_insertion(p.add_run("stays"))
    accept_revision(doc, ref.revision_id)
    assert not _ins(doc)
    assert p.text == "stays"


def test_reject_insertion_drops_text() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("keep ")
    ref = mark_insertion(p.add_run("drop"))
    reject_revision(doc, ref.revision_id)
    assert not _ins(doc)
    assert p.text == "keep "


def test_accept_deletion_removes_text() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("keep ")
    ref = mark_deletion(p.add_run("gone"))
    accept_revision(doc, ref.revision_id)
    assert not _del(doc)
    assert p.text == "keep "


def test_reject_deletion_restores_live_text() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = mark_deletion(p.add_run("back"))
    reject_revision(doc, ref.revision_id)
    assert not _del(doc)
    assert not xpath(_body(doc), ".//w:delText")
    assert p.text == "back"


def test_reject_deletion_only_affects_target() -> None:
    """Rejecting one deletion must not retag another deletion's delText."""
    doc = Document()
    p = doc.add_paragraph()
    reg = RevisionIdRegistry(doc)
    ref1 = mark_deletion(p.add_run("first"), id_registry=reg)
    mark_deletion(p.add_run("second"), id_registry=reg)
    reject_revision(doc, ref1.revision_id)
    # the second deletion is untouched -> still has delText
    assert len(xpath(_body(doc), ".//w:delText")) == 1


def test_accept_format_run_change_keeps_new_props() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("x")
    _add_rpr_change(run, rid=40)
    accept_revision(doc, 40)
    rpr = run._r.find(qn("w:rPr"))
    assert rpr.find(qn("w:rPrChange")) is None
    assert rpr.find(qn("w:b")) is not None  # new (bold) kept


def test_reject_format_run_change_restores_old_props() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("x")
    _add_rpr_change(run, rid=41)
    reject_revision(doc, 41)
    rpr = run._r.find(qn("w:rPr"))
    assert rpr.find(qn("w:rPrChange")) is None
    assert rpr.find(qn("w:b")) is None  # reverted to old (not bold)


def test_reject_format_paragraph_change_restores_old_jc() -> None:
    doc = Document()
    p = doc.add_paragraph("x")
    _add_ppr_change(p, rid=42)
    reject_revision(doc, 42)
    ppr = p._p.find(qn("w:pPr"))
    assert ppr.find(qn("w:pPrChange")) is None
    assert ppr.find(qn("w:jc")).get(qn("w:val")) == "left"


def test_accept_revision_unknown_id_raises_keyerror() -> None:
    doc = Document()
    with pytest.raises(RevisionNotFoundError):
        accept_revision(doc, 999999)
    with pytest.raises(KeyError):  # subclass relationship
        reject_revision(doc, 999999)


# --------------------------------------------------------------------------
# accept_all / reject_all
# --------------------------------------------------------------------------


def test_accept_all_resolves_everything() -> None:
    doc = Document()
    p = doc.add_paragraph()
    reg = RevisionIdRegistry(doc)
    p.add_run("A ")
    mark_insertion(p.add_run("B "), id_registry=reg)
    mark_deletion(p.add_run("C "), id_registry=reg)
    accept_all_revisions(doc)
    assert read_revisions(doc) == []
    assert p.text == "A B "


def test_reject_all_resolves_everything() -> None:
    doc = Document()
    p = doc.add_paragraph()
    reg = RevisionIdRegistry(doc)
    p.add_run("A ")
    mark_insertion(p.add_run("B "), id_registry=reg)
    mark_deletion(p.add_run("C "), id_registry=reg)
    reject_all_revisions(doc)
    assert read_revisions(doc) == []
    assert p.text == "A C "


def test_accept_all_idempotent_on_clean_doc() -> None:
    doc = Document()
    doc.add_paragraph("nothing to resolve")
    accept_all_revisions(doc)  # no error
    assert read_revisions(doc) == []


def test_accept_all_handles_nested_revision() -> None:
    """A w:ins nested inside a w:del is resolved without error."""
    doc = Document()
    p = doc.add_paragraph()
    outer = sub(p._p, "w:del", **{"w:id": "60", "w:author": "X", "w:date": _DATE})
    inner = sub(outer, "w:ins", **{"w:id": "61", "w:author": "X", "w:date": _DATE})
    r = sub(inner, "w:r")
    sub(r, "w:t").text = "nested"
    accept_all_revisions(doc)
    assert read_revisions(doc) == []


def test_paragraph_mark_revision_safe_fallback() -> None:
    """Accepting a paragraph-mark revision drops the mark without corrupting."""
    doc = Document()
    p = doc.add_paragraph("text")
    _add_paragraph_mark_ins(p, rid=70)
    accept_revision(doc, 70)
    # mark gone, paragraph text intact, no orphaned ins
    assert not _ins(doc)
    assert p.text == "text"


def test_full_round_trip_through_save(tmp_path) -> None:
    doc = Document()
    enable_track_changes(doc)
    p = doc.add_paragraph()
    reg = RevisionIdRegistry(doc)
    p.add_run("keep ")
    mark_insertion(p.add_run("ins "), author="A", id_registry=reg)
    mark_deletion(p.add_run("del"), author="B", id_registry=reg)
    out = tmp_path / "rev.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    revs = read_revisions(reopened)
    assert {r.revision_type for r in revs} == {"insertion", "deletion"}
    accept_all_revisions(reopened)
    assert reopened.paragraphs[0].text == "keep ins "
