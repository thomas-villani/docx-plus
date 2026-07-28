"""Tests for the document-wide cascade sweep (``iter_resolved_paragraphs``).

The sweep's contract is that it is *exactly* the per-target resolver run over
every paragraph, only with the document-level lookups shared. So the load-
bearing test here is equivalence: whatever ``resolve_effective_formatting``
says about a target, the sweep must say too. The rest cover the walk itself —
document order, table descent, and the merged-cell case where python-docx
hands back the same cell more than once.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from docx_plus.core.ns import qn
from docx_plus.core.oxml import sub
from docx_plus.styles import (
    StyleCascadeError,
    iter_resolved_paragraphs,
    resolve_effective_formatting,
)


def test_sweep_matches_per_target_resolver(multistyle_docx_path: Path) -> None:
    """Every paragraph and run resolves identically through both paths."""
    doc = Document(str(multistyle_docx_path))

    for resolved in iter_resolved_paragraphs(doc):
        assert resolved.formatting == resolve_effective_formatting(resolved.paragraph)
        for resolved_run in resolved.runs:
            assert resolved_run.formatting == resolve_effective_formatting(resolved_run.run)


def test_sweep_matches_per_target_resolver_with_numbering(
    numbered_docx_path: Path,
) -> None:
    """Equivalence holds for the numbering layer, which the cache also memoizes."""
    doc = Document(str(numbered_docx_path))

    for resolved in iter_resolved_paragraphs(doc):
        assert resolved.formatting == resolve_effective_formatting(resolved.paragraph)


def test_sweep_matches_per_target_resolver_with_provenance(
    multistyle_docx_path: Path,
) -> None:
    """Caching must not perturb provenance — the style_id / chain_depth attribution."""
    doc = Document(str(multistyle_docx_path))

    for resolved in iter_resolved_paragraphs(doc, include_provenance=True):
        expected = resolve_effective_formatting(resolved.paragraph, include_provenance=True)
        assert resolved.formatting == expected


def test_sweep_yields_document_order_including_tables() -> None:
    """Paragraphs and table cells interleave in document order.

    ``doc.paragraphs`` drops the tables and ``doc.tables`` drops the
    ordering, so neither alone can express this.
    """
    doc = Document()
    doc.add_paragraph("before")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "left"
    table.cell(0, 1).text = "right"
    doc.add_paragraph("after")

    assert [r.text for r in iter_resolved_paragraphs(doc)] == [
        "before",
        "left",
        "right",
        "after",
    ]


def test_sweep_indices_are_sequential() -> None:
    """``index`` counts what the sweep yields, gaplessly."""
    doc = Document()
    for i in range(4):
        doc.add_paragraph(f"p{i}")

    assert [r.index for r in iter_resolved_paragraphs(doc)] == [0, 1, 2, 3]


def test_sweep_excluding_tables_matches_doc_paragraphs() -> None:
    """With ``include_tables=False`` the indexing lines up with ``doc.paragraphs``."""
    doc = Document()
    doc.add_paragraph("before")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "in cell"
    doc.add_paragraph("after")

    swept = list(iter_resolved_paragraphs(doc, include_tables=False))

    assert [r.text for r in swept] == [p.text for p in doc.paragraphs]
    assert all(r.paragraph._p is doc.paragraphs[r.index]._p for r in swept)


def test_sweep_reports_table_depth() -> None:
    """``table_depth`` counts nesting; ``in_table`` is the boolean shorthand."""
    doc = Document()
    doc.add_paragraph("body")
    outer = doc.add_table(rows=1, cols=1)
    outer.cell(0, 0).text = "outer"
    inner = outer.cell(0, 0).add_table(rows=1, cols=1)
    inner.cell(0, 0).text = "inner"

    depths = {r.text: r.table_depth for r in iter_resolved_paragraphs(doc)}

    assert depths["body"] == 0
    assert depths["outer"] == 1
    assert depths["inner"] == 2
    assert [r.in_table for r in iter_resolved_paragraphs(doc) if r.text == "body"] == [False]


def test_sweep_does_not_double_count_merged_cells() -> None:
    """A horizontally merged cell is yielded once, not once per column it spans.

    ``row.cells`` returns the same ``_Cell`` object for every grid position a
    merge covers, so a naive walk reports the spanned cell's content twice.
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = f"c{i}"
    table.rows[0].cells[0].merge(table.rows[0].cells[1])

    texts = [r.text for r in iter_resolved_paragraphs(doc)]

    # The merge concatenates c0 and c1 into one cell holding two paragraphs.
    assert texts.count("c0") == 1
    assert texts.count("c1") == 1
    assert texts.count("c2") == 1


def test_sweep_can_skip_runs() -> None:
    """``include_runs=False`` skips the run walk, which is most of the work."""
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("alpha")
    para.add_run("beta")

    with_runs = list(iter_resolved_paragraphs(doc))
    without = list(iter_resolved_paragraphs(doc, include_runs=False))

    assert len(with_runs[0].runs) == 2
    assert without[0].runs == ()
    # Paragraph-level resolution is unaffected by the flag.
    assert without[0].formatting == with_runs[0].formatting


def test_sweep_run_indices_are_per_paragraph() -> None:
    """Run ``index`` restarts within each paragraph."""
    doc = Document()
    for _ in range(2):
        para = doc.add_paragraph()
        para.add_run("a")
        para.add_run("b")

    assert [[r.index for r in p.runs] for p in iter_resolved_paragraphs(doc)] == [
        [0, 1],
        [0, 1],
    ]


def test_sweep_is_lazy() -> None:
    """The sweep yields as it walks, so a caller can stop early.

    Guards the streaming contract: a generator that materialised everything
    first would still pass every other test here.
    """
    doc = Document()
    for i in range(50):
        doc.add_paragraph(f"p{i}")

    iterator = iter_resolved_paragraphs(doc)
    first = next(iterator)

    assert first.index == 0
    assert first.text == "p0"


def test_sweep_propagates_cascade_errors() -> None:
    """A style cycle raises out of the sweep rather than being swallowed."""
    doc = Document()
    style_el = sub(doc.styles.element, "w:style", **{"w:type": "paragraph", "w:styleId": "Loop"})
    sub(style_el, "w:name", **{"w:val": "Loop"})
    sub(style_el, "w:basedOn", **{"w:val": "Loop"})

    para = doc.add_paragraph("cycles")
    ppr = sub(para._p, "w:pPr")
    sub(ppr, "w:pStyle", **{"w:val": "Loop"})

    with pytest.raises(StyleCascadeError, match="cycle"):
        list(iter_resolved_paragraphs(doc))


def test_sweep_resolves_style_supplied_numbering() -> None:
    """The sweep sees style-supplied numbering, which the linter's rules turn on."""
    doc = Document()
    doc.add_paragraph("bulleted", style="List Bullet")
    doc.add_paragraph("plain")

    swept = list(iter_resolved_paragraphs(doc))

    assert swept[0].formatting.num_id == 1
    assert swept[1].formatting.num_id is None


def test_sweep_of_empty_document_yields_nothing() -> None:
    """A document with no body paragraphs sweeps cleanly."""
    doc = Document()
    for para in list(doc.paragraphs):
        para._p.getparent().remove(para._p)

    assert list(iter_resolved_paragraphs(doc)) == []


def test_sweep_handles_paragraph_with_no_runs() -> None:
    """An empty paragraph resolves with an empty run tuple, not an error."""
    doc = Document()
    doc.add_paragraph("")

    swept = list(iter_resolved_paragraphs(doc))

    assert len(swept) == 1
    assert swept[0].runs == ()
    assert swept[0].formatting.font_size is not None  # docDefaults still applied


def test_cache_is_not_shared_between_sweeps() -> None:
    """Each sweep builds its own cache, so a document edited between two runs re-resolves.

    The cache memoizes on the assumption that styles.xml does not change
    *during* a walk; it must not outlive one.
    """
    doc = Document()
    para = doc.add_paragraph("text")

    before = next(iter_resolved_paragraphs(doc)).formatting
    assert before.bold is not True

    rpr = sub(para._p, "w:pPr")
    sub(sub(rpr, "w:rPr"), "w:b")
    doc.styles.element.find(qn("w:docDefaults"))  # touch, no mutation

    after = next(iter_resolved_paragraphs(doc)).formatting
    assert after.bold is True


# ---------------------------------------------------------------------------
# Baselines: the same target resolved without its own direct formatting.
# ---------------------------------------------------------------------------


def test_baseline_is_absent_by_default() -> None:
    """It roughly doubles the resolve work, so it is opt-in."""
    doc = Document()
    doc.add_paragraph().add_run("text")

    resolved = next(iter_resolved_paragraphs(doc))

    assert resolved.baseline is None
    assert resolved.runs[0].baseline is None


def test_paragraph_baseline_excludes_its_own_ppr() -> None:
    doc = Document()
    para = doc.add_paragraph("text")
    para.paragraph_format.space_after = Pt(18)

    resolved = next(iter_resolved_paragraphs(doc, include_baseline=True))

    assert resolved.formatting.spacing_after == 18 * 20
    assert resolved.baseline is not None
    assert resolved.baseline.spacing_after != resolved.formatting.spacing_after


def test_run_baseline_excludes_its_own_rpr() -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("text")
    run.font.size = Pt(24)

    resolved = next(iter_resolved_paragraphs(doc, include_baseline=True))

    assert resolved.runs[0].formatting.font_size == 24.0
    assert resolved.runs[0].baseline is not None
    assert resolved.runs[0].baseline.font_size == 11.0


def test_baselines_match_the_per_target_resolver() -> None:
    """Same equivalence contract as the full resolve, one layer down."""
    doc = Document()
    para = doc.add_paragraph("text", style="Heading 1")
    para.paragraph_format.space_after = Pt(18)
    para.runs[0].font.size = Pt(24)

    resolved = next(iter_resolved_paragraphs(doc, include_baseline=True))

    assert resolved.baseline == resolve_effective_formatting(para, stop_below="directParagraph")
    assert resolved.runs[0].baseline == resolve_effective_formatting(
        para.runs[0], stop_below="directRun"
    )


def test_baseline_is_skipped_for_runs_when_runs_are() -> None:
    doc = Document()
    doc.add_paragraph().add_run("text")

    resolved = next(iter_resolved_paragraphs(doc, include_baseline=True, include_runs=False))

    assert resolved.runs == ()
    assert resolved.baseline is not None
