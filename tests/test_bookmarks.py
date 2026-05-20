"""Tests for ``docx_plus.bookmarks`` — anchored bookmarks, reading, deletion."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_plus.bookmarks import (
    BookmarkIdRegistry,
    BookmarkInfo,
    BookmarkRef,
    add_bookmark,
    delete_bookmark,
    read_bookmarks,
)
from docx_plus.core.ids import DuplicateIdError
from docx_plus.core.ns import qn
from docx_plus.core.oxml import xpath


def _body(doc):
    return doc.element.body


def _starts(doc, name=None):
    if name is None:
        return xpath(_body(doc), ".//w:bookmarkStart")
    return xpath(_body(doc), ".//w:bookmarkStart[@w:name=$n]", n=name)


def _ends(doc):
    return xpath(_body(doc), ".//w:bookmarkEnd")


# --------------------------------------------------------------------------
# add_bookmark — body anchoring.
# --------------------------------------------------------------------------


def test_add_bookmark_on_run_writes_paired_markers() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("anchored")
    ref = add_bookmark(run, "test_bm")
    assert isinstance(ref, BookmarkRef)
    assert len(_starts(doc, "test_bm")) == 1
    assert len(_ends(doc)) == 1
    # ids match
    bid = _starts(doc, "test_bm")[0].get(qn("w:id"))
    assert _ends(doc)[0].get(qn("w:id")) == bid


def test_add_bookmark_brackets_the_run() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("hi")
    add_bookmark(run, "bm")
    types = [c.tag.rpartition("}")[2] for c in p._p]
    bs = types.index("bookmarkStart")
    be = types.index("bookmarkEnd")
    r = types.index("r", bs)
    assert bs < r < be


def test_add_bookmark_on_paragraph_wraps_all_runs() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("a ")
    p.add_run("b ")
    p.add_run("c")
    add_bookmark(p, "para_bm")
    types = [c.tag.rpartition("}")[2] for c in p._p]
    bs = types.index("bookmarkStart")
    be = types.index("bookmarkEnd")
    runs_between = [t for t in types[bs + 1 : be] if t == "r"]
    assert len(runs_between) == 3


def test_add_bookmark_on_paragraph_with_no_runs_raises() -> None:
    doc = Document()
    p = doc.add_paragraph()
    with pytest.raises(ValueError, match="at least one run"):
        add_bookmark(p, "bm")


def test_add_bookmark_on_run_range() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("first ")
    p.add_run("middle ")
    r3 = p.add_run("last")
    add_bookmark((r1, r3), "range_bm")
    types = [c.tag.rpartition("}")[2] for c in p._p]
    bs = types.index("bookmarkStart")
    be = types.index("bookmarkEnd")
    runs_between = [t for t in types[bs + 1 : be] if t == "r"]
    assert len(runs_between) == 3


def test_add_bookmark_rejects_invalid_name() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("x")
    with pytest.raises(ValueError, match="bookmark name"):
        add_bookmark(run, "has spaces")
    with pytest.raises(ValueError, match="bookmark name"):
        add_bookmark(run, "1starts_with_digit")
    with pytest.raises(ValueError, match="bookmark name"):
        add_bookmark(run, "a" * 41)


def test_add_bookmark_accepts_valid_names() -> None:
    doc = Document()
    p = doc.add_paragraph()
    for name in ("ok", "_underscore", "a1", "a_b_c", "Foo_Bar_42"):
        run = p.add_run(name)
        add_bookmark(run, name)


def test_add_bookmark_rejects_bad_target() -> None:
    with pytest.raises(TypeError, match="Run, Paragraph"):
        add_bookmark("not a run", "bm")  # type: ignore[arg-type]


def test_add_bookmark_returns_unique_ids() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("a")
    r2 = p.add_run("b")
    a = add_bookmark(r1, "one")
    b = add_bookmark(r2, "two")
    assert a.bookmark_id != b.bookmark_id


# --------------------------------------------------------------------------
# BookmarkIdRegistry — seeding.
# --------------------------------------------------------------------------


def test_registry_seeds_from_existing_bookmarks() -> None:
    doc = Document()
    p = doc.add_paragraph()
    ref = add_bookmark(p.add_run("x"), "bm")
    reg = BookmarkIdRegistry(doc)
    with pytest.raises(DuplicateIdError):
        reg.reserve(ref.bookmark_id)


def test_shared_registry_produces_disjoint_ids() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("a")
    r2 = p.add_run("b")
    reg = BookmarkIdRegistry(doc)
    a = add_bookmark(r1, "one", id_registry=reg)
    b = add_bookmark(r2, "two", id_registry=reg)
    assert a.bookmark_id != b.bookmark_id


# --------------------------------------------------------------------------
# read_bookmarks
# --------------------------------------------------------------------------


def test_read_bookmarks_empty() -> None:
    assert read_bookmarks(Document()) == []


def test_read_bookmarks_returns_anchored_text() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("prefix ")
    target = p.add_run("ANCHORED")
    p.add_run(" suffix")
    add_bookmark(target, "mybm")

    bookmarks = read_bookmarks(doc)
    assert len(bookmarks) == 1
    only = bookmarks[0]
    assert isinstance(only, BookmarkInfo)
    assert only.name == "mybm"
    assert only.anchored_text == "ANCHORED"
    assert only.paragraph_index == 0


def test_read_bookmarks_paragraph_index_tracks_position() -> None:
    doc = Document()
    doc.add_paragraph("first")
    doc.add_paragraph("second")
    p3 = doc.add_paragraph()
    add_bookmark(p3.add_run("third"), "bm")
    info = read_bookmarks(doc)[0]
    assert info.paragraph_index == 2


def test_read_bookmarks_preserves_document_order() -> None:
    doc = Document()
    p = doc.add_paragraph()
    refs = [add_bookmark(p.add_run(f"r{i}"), f"bm_{i}") for i in range(3)]
    bookmarks = read_bookmarks(doc)
    assert [b.bookmark_id for b in bookmarks] == [r.bookmark_id for r in refs]


def test_read_bookmarks_handles_orphan_start_without_end() -> None:
    """A bookmarkStart with no matching bookmarkEnd reads as anchored_text=''."""
    doc = Document()
    p = doc.add_paragraph()
    add_bookmark(p.add_run("x"), "bm")
    # Strip the end marker.
    from docx_plus.core.oxml import remove

    for end in _ends(doc):
        remove(end)
    info = read_bookmarks(doc)[0]
    assert info.anchored_text == ""


# --------------------------------------------------------------------------
# delete_bookmark
# --------------------------------------------------------------------------


def test_delete_bookmark_removes_pair() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_bookmark(p.add_run("x"), "kill")
    delete_bookmark(doc, "kill")
    assert _starts(doc, "kill") == []
    assert _ends(doc) == []


def test_delete_bookmark_is_idempotent_for_missing() -> None:
    delete_bookmark(Document(), "nonexistent")  # no-op


def test_delete_bookmark_leaves_others_intact() -> None:
    doc = Document()
    p = doc.add_paragraph()
    add_bookmark(p.add_run("a"), "keep")
    add_bookmark(p.add_run("b"), "drop")
    delete_bookmark(doc, "drop")
    assert _starts(doc, "keep")
    assert _starts(doc, "drop") == []


# --------------------------------------------------------------------------
# Round-trip
# --------------------------------------------------------------------------


def test_bookmark_round_trip(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("before ")
    p.add_run("INSIDE")
    p.add_run(" after")
    target = [r for r in p.runs if r.text == "INSIDE"][0]
    add_bookmark(target, "round")
    out = tmp_path / "bm.docx"
    doc.save(str(out))

    reopened = Document(str(out))
    bookmarks = read_bookmarks(reopened)
    assert len(bookmarks) == 1
    assert bookmarks[0].name == "round"
    assert bookmarks[0].anchored_text == "INSIDE"
